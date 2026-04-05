#!/usr/bin/env python3
"""
Generate "Cómo Montar una Dark Kitchen" guide deliverables (4 files).
AI Chef Pro — aichef.pro
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment, numbers
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "public", "dl", "guia-dark-kitchen"
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─── THEME COLORS ───────────────────────────────────────
GOLD = "FFD700"
GOLD_RGB = RGBColor(0xFF, 0xD7, 0x00)
DARK_BG = "1A1A1A"
HEADER_BG = "2D2D2D"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"
MEDIUM_GRAY = "E0E0E0"
INPUT_COLOR = "E8F5E9"

# ─── EXCEL STYLES ──────────────────────────────────────
title_font_xl = Font(name="Calibri", size=16, bold=True, color=GOLD)
subtitle_font_xl = Font(name="Calibri", size=11, color="888888", italic=True)
header_font_xl = Font(name="Calibri", size=11, bold=True, color=WHITE)
section_font_xl = Font(name="Calibri", size=12, bold=True, color=GOLD)
data_font_xl = Font(name="Calibri", size=11)
bold_font_xl = Font(name="Calibri", size=11, bold=True)
formula_font_xl = Font(name="Calibri", size=11, color="1565C0", bold=True)

header_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
gold_fill = PatternFill(start_color=GOLD, end_color=GOLD, fill_type="solid")
input_fill = PatternFill(start_color=INPUT_COLOR, end_color=INPUT_COLOR, fill_type="solid")
light_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")

thin_border = Border(
    left=Side(style="thin", color=MEDIUM_GRAY),
    right=Side(style="thin", color=MEDIUM_GRAY),
    top=Side(style="thin", color=MEDIUM_GRAY),
    bottom=Side(style="thin", color=MEDIUM_GRAY),
)

center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
currency_format = '#,##0.00 €'
pct_format = '0.0%'

BRAND_LINE = "AI Chef Pro · aichef.pro — Cómo Montar una Dark Kitchen"


# ═══════════════════════════════════════════════════════════
# HELPER: style Excel sheet
# ═══════════════════════════════════════════════════════════
def style_header_row(ws, row, col_count):
    for c in range(1, col_count + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = header_font_xl
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align


def style_data_cell(ws, row, col, value=None, font=None, fill=None, fmt=None, align=None):
    cell = ws.cell(row=row, column=col)
    if value is not None:
        cell.value = value
    cell.font = font or data_font_xl
    cell.border = thin_border
    cell.alignment = align or left_align
    if fill:
        cell.fill = fill
    if fmt:
        cell.number_format = fmt
    return cell


def add_title_block(ws, title, subtitle, merge_to=8):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=merge_to)
    c = ws.cell(row=1, column=1, value=title)
    c.font = title_font_xl
    c.alignment = Alignment(horizontal="left", vertical="center")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=merge_to)
    c2 = ws.cell(row=2, column=1, value=subtitle)
    c2.font = subtitle_font_xl
    ws.row_dimensions[1].height = 30
    ws.row_dimensions[2].height = 20


# ═══════════════════════════════════════════════════════════
# FILE 1: DOCX — Guía completa (~40 pages)
# ═══════════════════════════════════════════════════════════
def generate_guide_docx():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Default style ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    # Heading 1 style
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(20)
    h1.font.color.rgb = GOLD_RGB
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2 style
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 style
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    h3.font.bold = True

    def add_tip_box(text):
        """Add a 'Consejo del Chef' callout."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        # Create a shaded paragraph
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFF8E1"/>')
        pPr.append(shading)
        run = p.add_run("🍳 Consejo del Chef: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0xD4, 0xA0, 0x00)
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.italic = True

    def add_bullet(text):
        p = doc.add_paragraph(text, style='List Bullet')
        return p

    def add_numbered(text):
        p = doc.add_paragraph(text, style='List Number')
        return p

    # ════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CÓMO MONTAR UNA\nDARK KITCHEN EN ESPAÑA")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = GOLD_RGB

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Guía Completa")
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("De la Idea a la Primera Entrega en 90 Días")
    run3.font.size = Pt(14)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for _ in range(4):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Chef John Guerrero")
    run4.font.size = Pt(14)
    run4.font.bold = True

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("AI Chef Pro · aichef.pro")
    run5.font.size = Pt(12)
    run5.font.color.rgb = GOLD_RGB

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run6 = p6.add_run("29 años en alta hostelería · 15 años en consultoría gastronómica")
    run6.font.size = Pt(10)
    run6.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ════════════════════════════════════════════
    doc.add_heading("Índice de Contenidos", level=1)
    toc_items = [
        "1. ¿Qué es una Dark Kitchen?",
        "2. El Mercado Delivery en España",
        "3. Modelos de Negocio",
        "4. Estudio de Viabilidad y Plan Financiero",
        "5. Requisitos Legales en España",
        "6. APPCC y Seguridad Alimentaria",
        "7. Ubicación y Local",
        "8. Diseño y Layout de Cocina",
        "9. Equipamiento y Proveedores",
        "10. Tecnología y Gestión de Pedidos",
        "11. Packaging y Sostenibilidad",
        "12. Marketing, Lanzamiento y Errores Comunes",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 1
    # ════════════════════════════════════════════
    doc.add_heading("1. ¿Qué es una Dark Kitchen?", level=1)

    doc.add_paragraph(
        "Una dark kitchen, también conocida como cocina fantasma, cocina virtual o cloud kitchen, "
        "es un establecimiento de restauración diseñado exclusivamente para la preparación de comida "
        "destinada al reparto a domicilio. A diferencia de un restaurante tradicional, no cuenta con "
        "sala de comedor ni zona de atención al público. Su único objetivo es producir platos de "
        "calidad con la máxima eficiencia operativa para las plataformas de delivery."
    )

    doc.add_heading("Origen y evolución del concepto", level=2)
    doc.add_paragraph(
        "El concepto de dark kitchen nació en el Reino Unido y Estados Unidos alrededor de 2015, "
        "impulsado por el crecimiento exponencial de las plataformas de reparto como Deliveroo, "
        "Uber Eats y DoorDash. En España, el fenómeno cobró fuerza a partir de 2018, cuando Glovo, "
        "Uber Eats y Just Eat consolidaron su presencia en las principales ciudades. La pandemia "
        "de COVID-19 en 2020 aceleró esta tendencia de manera radical: muchos restaurantes que "
        "cerraron su sala descubrieron que el modelo delivery-only era viable y, en muchos casos, "
        "más rentable que mantener un local con servicio en mesa."
    )
    doc.add_paragraph(
        "Hoy, las dark kitchens representan uno de los segmentos con mayor crecimiento dentro de "
        "la industria de la hostelería. Según datos de Euromonitor, el mercado global de cocinas "
        "fantasma alcanzará los 1,5 billones de dólares en 2030. En España, se estima que hay más "
        "de 2.000 dark kitchens operativas, concentradas principalmente en Madrid, Barcelona, "
        "Valencia, Sevilla y Málaga."
    )

    doc.add_heading("Diferencias clave con un restaurante tradicional", level=2)
    doc.add_paragraph(
        "La diferencia fundamental radica en la ausencia total de espacio para comensales. Esto "
        "tiene implicaciones profundas en todos los aspectos del negocio:"
    )
    add_bullet("Inversión inicial: entre un 50% y un 70% menor que un restaurante con sala.")
    add_bullet("Personal: se elimina la necesidad de camareros, maîtres y personal de sala. "
               "El equipo se compone exclusivamente de cocineros y un encargado de expedición.")
    add_bullet("Ubicación: no se necesita una zona de alto tránsito peatonal. Se puede operar "
               "desde un polígono industrial o una zona comercial secundaria con alquileres más bajos.")
    add_bullet("Licencia de actividad: los requisitos son generalmente menos exigentes, ya que "
               "no hay atención presencial al público, aunque varía según la comunidad autónoma.")
    add_bullet("Horarios: la cocina opera exclusivamente durante las franjas de mayor demanda "
               "de delivery (12:00-15:30 y 19:30-23:00), lo que optimiza los costes laborales.")
    add_bullet("Marketing: toda la captación de clientes se realiza a través de plataformas "
               "digitales y redes sociales, no mediante la visibilidad física del local.")

    add_tip_box(
        "No subestimes la importancia de elegir bien las franjas horarias. El 80% de la "
        "facturación de una dark kitchen se concentra en solo 6 horas al día. Organiza tu "
        "mise en place y tus turnos en torno a esos picos."
    )

    doc.add_heading("Tipos de dark kitchen", level=2)
    doc.add_paragraph(
        "Dentro del ecosistema de cocinas fantasma, existen varios subtipos que conviene "
        "distinguir antes de tomar cualquier decisión de inversión:"
    )
    add_bullet("Dark kitchen independiente: un emprendedor o chef monta su propia cocina "
               "y opera una o varias marcas virtuales. Es el modelo más común y el que "
               "abordaremos en profundidad en esta guía.")
    add_bullet("Cocina compartida (shared kitchen): un operador alquila espacios de cocina "
               "equipados a diferentes marcas. Empresas como Cocinas Ocultas o Kitchen United "
               "operan bajo este modelo en España.")
    add_bullet("Cocina de restaurante reconvertida: restaurantes tradicionales que destinan "
               "parte de su cocina exclusivamente a marcas virtuales para delivery.")
    add_bullet("Franquicia virtual: marcas de delivery que venden su concepto, recetas y "
               "know-how a operadores que producen desde sus propias cocinas.")

    doc.add_paragraph(
        "Cada modelo tiene ventajas e inconvenientes en términos de inversión, control operativo "
        "y escalabilidad. En el capítulo 3 analizaremos cada uno en detalle con sus pros y contras "
        "para que puedas elegir el que mejor se adapta a tu situación."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 2
    # ════════════════════════════════════════════
    doc.add_heading("2. El Mercado Delivery en España", level=1)

    doc.add_paragraph(
        "España se ha consolidado como uno de los mercados de delivery más dinámicos de Europa. "
        "Comprender su estructura, sus cifras y sus tendencias es fundamental antes de embarcarse "
        "en un proyecto de dark kitchen. Este capítulo analiza el estado actual del mercado, las "
        "plataformas dominantes y las oportunidades que existen para nuevos operadores."
    )

    doc.add_heading("Datos del mercado español", level=2)
    doc.add_paragraph(
        "El sector del food delivery en España ha experimentado un crecimiento sostenido. Según "
        "datos de la consultora NPD Group y la Asociación Española de Food Delivery, el mercado "
        "generó más de 4.200 millones de euros en 2024, con un crecimiento interanual del 8%. "
        "Las proyecciones para 2026 sitúan el volumen por encima de los 5.000 millones de euros."
    )
    doc.add_paragraph("Los indicadores clave del mercado son:")
    add_bullet("Ticket medio: entre 18€ y 25€, dependiendo de la ciudad y el tipo de cocina.")
    add_bullet("Frecuencia de pedido: los usuarios activos realizan una media de 3,2 pedidos "
               "por semana en las grandes ciudades.")
    add_bullet("Penetración: el 45% de los hogares urbanos españoles ha utilizado alguna "
               "plataforma de delivery en el último mes.")
    add_bullet("Franjas horarias principales: almuerzo (12:30-14:30) y cena (20:00-22:30), "
               "con un crecimiento emergente del desayuno/brunch los fines de semana.")

    doc.add_heading("Las plataformas dominantes", level=2)
    doc.add_paragraph(
        "El mercado español está dominado por tres grandes plataformas que concentran más del "
        "90% de los pedidos online. Conocer las particularidades de cada una es esencial para "
        "maximizar tu visibilidad y tus márgenes:"
    )

    doc.add_heading("Glovo", level=3)
    doc.add_paragraph(
        "Líder en España con presencia en más de 300 ciudades. Glovo ofrece el modelo de "
        "marketplace más completo, con Glovo Food, Glovo Market y Glovo Express. Su comisión "
        "estándar oscila entre el 25% y el 35% del valor del pedido, dependiendo del plan "
        "contratado. Destaca por su algoritmo de posicionamiento que premia el tiempo de "
        "preparación rápido, las buenas valoraciones y la aceptación de pedidos."
    )

    doc.add_heading("Uber Eats", level=3)
    doc.add_paragraph(
        "Segunda plataforma por cuota de mercado, especialmente fuerte en Madrid y Barcelona. "
        "Su comisión se sitúa entre el 30% y el 35%. Ofrece herramientas de marketing integradas "
        "como Uber Eats Ads y promociones patrocinadas. Su punto fuerte es la integración con "
        "el ecosistema Uber, que le permite ofrecer entregas rápidas y una base de usuarios "
        "amplia y con alto poder adquisitivo."
    )

    doc.add_heading("Just Eat", level=3)
    doc.add_paragraph(
        "Históricamente la primera plataforma en España, con una base de restaurantes muy amplia. "
        "Ofrece dos modelos: marketplace (el restaurante gestiona su propio reparto) con comisiones "
        "del 14-16%, y logística completa con comisiones del 25-30%. Es especialmente fuerte en "
        "ciudades medianas y zonas residenciales."
    )

    doc.add_heading("Ciudades con mayor potencial", level=2)
    doc.add_paragraph(
        "No todas las ciudades españolas ofrecen las mismas oportunidades para una dark kitchen. "
        "Los factores clave son la densidad de población, el nivel de renta, la penetración de "
        "plataformas y la competencia existente:"
    )
    add_bullet("Madrid: el mayor mercado, con alta competencia pero también alta demanda. "
               "Zonas como Vallecas, Usera y Carabanchel ofrecen costes de local más bajos "
               "y buena cobertura de riders.")
    add_bullet("Barcelona: segundo mercado en volumen, muy competitivo en el centro pero con "
               "oportunidades en el cinturón metropolitano (Hospitalet, Badalona, Santa Coloma).")
    add_bullet("Valencia: mercado en crecimiento rápido, con costes operativos significativamente "
               "menores que Madrid y Barcelona. Excelente relación demanda/competencia.")
    add_bullet("Sevilla y Málaga: mercados emergentes con gran potencial, especialmente "
               "en temporada alta turística. Costes de local muy competitivos.")
    add_bullet("Ciudades universitarias (Salamanca, Granada, Santiago): demanda constante "
               "de delivery económico por parte de estudiantes, con temporada baja en verano.")

    add_tip_box(
        "Antes de elegir tu ubicación, analiza la densidad de restaurantes en delivery dentro "
        "de tu radio objetivo usando las propias apps. Si en un radio de 2 km hay más de 200 "
        "opciones, la competencia será feroz. Busca zonas con demanda alta pero oferta "
        "limitada: ahí está la verdadera oportunidad."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 3
    # ════════════════════════════════════════════
    doc.add_heading("3. Modelos de Negocio", level=1)

    doc.add_paragraph(
        "Elegir el modelo de negocio adecuado es la primera gran decisión estratégica. No existe "
        "un modelo universalmente mejor: cada uno tiene ventajas, riesgos y niveles de inversión "
        "diferentes. En este capítulo analizamos los cuatro modelos principales con sus pros y "
        "contras para que puedas tomar una decisión informada."
    )

    doc.add_heading("Modelo 1: Marca única (Single Brand)", level=2)
    doc.add_paragraph(
        "Es el modelo más sencillo y el recomendado para emprendedores que se inician en el "
        "mundo del delivery. Consiste en operar una única marca desde tu cocina, especializándote "
        "en un tipo de cocina concreto (por ejemplo, hamburguesas gourmet, poke bowls o comida "
        "mexicana). La ventaja principal es la simplicidad operativa: una sola carta, un solo "
        "flujo de trabajo y un enfoque claro en la calidad de un producto."
    )
    doc.add_paragraph("Ventajas:")
    add_bullet("Menor complejidad operativa: un solo menú, un solo flujo de producción.")
    add_bullet("Inversión inicial más baja: menos ingredientes, menos equipamiento específico.")
    add_bullet("Más fácil de mantener la calidad consistente.")
    add_bullet("Ideal para validar el concepto antes de escalar.")
    doc.add_paragraph("Inconvenientes:")
    add_bullet("Dependencia total de un único concepto gastronómico.")
    add_bullet("Si la marca no funciona, no hay alternativa inmediata.")
    add_bullet("Menor facturación potencial que un modelo multimarca.")

    doc.add_heading("Modelo 2: Multimarca (Multi Brand)", level=2)
    doc.add_paragraph(
        "El modelo multimarca consiste en operar varias marcas virtuales desde la misma cocina. "
        "Por ejemplo, una dark kitchen puede tener simultáneamente una marca de hamburguesas, "
        "otra de sushi y otra de ensaladas saludables. Cada marca aparece como un restaurante "
        "independiente en las plataformas de delivery, lo que permite captar diferentes segmentos "
        "de demanda y multiplicar la visibilidad."
    )
    doc.add_paragraph(
        "Este es el modelo que genera mayor facturación por metro cuadrado, pero también el más "
        "exigente en términos de gestión. Requiere una cocina bien diseñada con zonas diferenciadas, "
        "un sistema de gestión de pedidos robusto y un equipo capaz de producir platos de diferentes "
        "estilos simultáneamente sin sacrificar calidad ni tiempos de preparación."
    )
    doc.add_paragraph("Ventajas:")
    add_bullet("Diversificación del riesgo: si una marca no funciona, las otras compensan.")
    add_bullet("Mayor facturación por metro cuadrado de cocina.")
    add_bullet("Aprovechamiento de ingredientes comunes entre marcas.")
    add_bullet("Mayor visibilidad en las plataformas (más restaurantes = más exposición).")
    doc.add_paragraph("Inconvenientes:")
    add_bullet("Mayor complejidad operativa y de gestión de inventario.")
    add_bullet("Necesidad de más personal cualificado.")
    add_bullet("Inversión inicial más alta en equipamiento.")
    add_bullet("Riesgo de canibalización entre marcas propias.")

    doc.add_heading("Modelo 3: Cocina compartida (Shared Kitchen)", level=2)
    doc.add_paragraph(
        "En este modelo, un operador principal equipa y gestiona un espacio de cocina grande que "
        "se divide en estaciones independientes. Cada estación se alquila a un operador diferente "
        "que produce sus propios platos. Es el modelo de negocio de empresas como Cocinas Ocultas "
        "en Madrid o Cuyna en Barcelona."
    )
    doc.add_paragraph(
        "Para el emprendedor individual, la cocina compartida ofrece una forma de empezar con "
        "una inversión mínima. El alquiler mensual suele incluir equipamiento, suministros básicos, "
        "licencias y, en muchos casos, soporte tecnológico. Los precios oscilan entre 1.500€ y "
        "4.000€ mensuales por estación, dependiendo de la ciudad y el tamaño."
    )

    doc.add_heading("Modelo 4: Franquicia virtual", level=2)
    doc.add_paragraph(
        "La franquicia virtual es un modelo relativamente nuevo en el que una marca de delivery "
        "consolidada licencia su concepto, sus recetas estandarizadas y su know-how a operadores "
        "que producen desde sus propias cocinas. Marcas como Taster, Not So Dark o Keatz han "
        "operado bajo este modelo en Europa."
    )
    doc.add_paragraph(
        "El franquiciado se beneficia de una marca ya posicionada en las plataformas, con "
        "valoraciones y visibilidad existentes. A cambio, paga un canon de entrada y un royalty "
        "sobre las ventas (habitualmente entre el 5% y el 10%). Es una opción interesante para "
        "operadores que ya tienen una cocina y quieren añadir una marca adicional sin desarrollarla "
        "desde cero."
    )

    add_tip_box(
        "Mi recomendación para empezadores: comienza con una marca única, perfecciona tu "
        "operación durante 3-6 meses, y solo entonces valora lanzar una segunda marca. "
        "La calidad y la velocidad de preparación son los factores que más influyen en tu "
        "posicionamiento en las plataformas, y es mucho más fácil dominarlos con un solo concepto."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 4
    # ════════════════════════════════════════════
    doc.add_heading("4. Estudio de Viabilidad y Plan Financiero", level=1)

    doc.add_paragraph(
        "El plan financiero es la columna vertebral de cualquier proyecto de dark kitchen. Sin "
        "números claros, estarás navegando a ciegas. Este capítulo te guía paso a paso en la "
        "elaboración de un estudio de viabilidad realista, con cifras basadas en la experiencia "
        "real del mercado español."
    )

    doc.add_heading("Inversión inicial: ¿cuánto necesitas?", level=2)
    doc.add_paragraph(
        "La inversión inicial para montar una dark kitchen en España varía significativamente "
        "según el modelo elegido, la ciudad y el estado del local. Como referencia, estos son "
        "los rangos habituales:"
    )
    add_bullet("Dark kitchen básica (marca única, 40-60 m²): 30.000€ - 50.000€")
    add_bullet("Dark kitchen multimarca (2-3 marcas, 80-120 m²): 50.000€ - 80.000€")
    add_bullet("Cocina compartida (alquiler de estación): 5.000€ - 15.000€ de entrada")

    doc.add_paragraph("Desglose típico de la inversión para una dark kitchen básica:")
    add_bullet("Obra civil y acondicionamiento: 8.000€ - 15.000€")
    add_bullet("Equipamiento de cocina: 10.000€ - 20.000€")
    add_bullet("Instalaciones (electricidad, agua, gas, ventilación): 5.000€ - 10.000€")
    add_bullet("Licencias y trámites legales: 2.000€ - 4.000€")
    add_bullet("Tecnología (tablets, software, KDS): 1.000€ - 3.000€")
    add_bullet("Marketing de lanzamiento: 2.000€ - 5.000€")
    add_bullet("Capital circulante (2 meses): 5.000€ - 10.000€")

    doc.add_heading("Costes fijos y variables mensuales", level=2)
    doc.add_paragraph("Los costes mensuales de operación se dividen en dos categorías:")
    doc.add_paragraph("Costes fijos (no varían con el volumen de pedidos):")
    add_bullet("Alquiler del local: 800€ - 2.500€/mes (según ciudad y zona)")
    add_bullet("Suministros (luz, agua, gas): 400€ - 800€/mes")
    add_bullet("Seguros: 100€ - 200€/mes")
    add_bullet("Software y tecnología: 100€ - 300€/mes")
    add_bullet("Gestoría y asesoría: 150€ - 300€/mes")
    add_bullet("Limpieza y desratización: 100€ - 200€/mes")

    doc.add_paragraph("Costes variables (aumentan con el volumen):")
    add_bullet("Coste de materia prima (food cost): 28-35% de la facturación")
    add_bullet("Comisiones de plataformas: 25-35% del valor del pedido")
    add_bullet("Packaging: 0,50€ - 1,50€ por pedido")
    add_bullet("Personal de cocina: depende del volumen, mínimo 1 cocinero + 1 ayudante")
    add_bullet("Marketing en plataformas: 3-8% de la facturación")

    doc.add_heading("Cálculo del punto de equilibrio (break-even)", level=2)
    doc.add_paragraph(
        "El punto de equilibrio es el número mínimo de pedidos diarios que necesitas para cubrir "
        "todos tus costes. Para una dark kitchen básica con los siguientes parámetros:"
    )
    add_bullet("Ticket medio: 20€")
    add_bullet("Comisión media de plataforma: 30%")
    add_bullet("Food cost: 30% sobre el precio de venta")
    add_bullet("Packaging por pedido: 0,80€")
    add_bullet("Costes fijos mensuales totales: 4.500€")

    doc.add_paragraph(
        "Ingreso neto por pedido = 20€ - 6€ (comisión) - 6€ (food cost) - 0,80€ (packaging) = 7,20€. "
        "Para cubrir 4.500€ de costes fijos: 4.500 / 7,20 = 625 pedidos/mes = ~21 pedidos/día. "
        "Esto significa que necesitas una media de 21 pedidos diarios solo para llegar al punto de "
        "equilibrio. A partir del pedido 22, empiezas a generar beneficio."
    )

    doc.add_heading("Márgenes reales tras comisiones", level=2)
    doc.add_paragraph(
        "Aquí es donde muchos emprendedores se llevan una sorpresa desagradable. Las comisiones "
        "de las plataformas de delivery son el mayor coste variable de una dark kitchen, y pueden "
        "hacer que un negocio aparentemente rentable sea deficitario si no se gestionan bien los márgenes."
    )
    doc.add_paragraph(
        "Un escenario realista para una dark kitchen que facture 12.000€ brutos mensuales: "
        "tras descontar comisiones de plataforma (3.600€), food cost (3.600€), packaging (600€) "
        "y costes fijos (4.500€), el beneficio operativo (EBITDA) queda en aproximadamente -300€. "
        "Sí, con 12.000€ de facturación bruta puedes estar perdiendo dinero."
    )
    doc.add_paragraph(
        "La clave está en facturar por encima de los 15.000€-18.000€ brutos mensuales para "
        "alcanzar un EBITDA positivo significativo. Con 20.000€ de facturación y buena gestión, "
        "un EBITDA del 10-15% es realista, lo que equivale a 2.000€-3.000€ mensuales de beneficio."
    )

    add_tip_box(
        "Negocia siempre las comisiones con las plataformas. Si demuestras un volumen de "
        "pedidos creciente, puedes conseguir rebajas de 2-5 puntos porcentuales. Esos puntos "
        "van directamente a tu margen. También considera el canal propio de pedidos (web/app "
        "propia) para el 15-20% de tus ventas: ahí la comisión es cero."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 5
    # ════════════════════════════════════════════
    doc.add_heading("5. Requisitos Legales en España", level=1)

    doc.add_paragraph(
        "Cumplir con la normativa legal es imprescindible y no negociable. Una dark kitchen, "
        "aunque no tenga sala de comedor, es un establecimiento alimentario sujeto a regulaciones "
        "estrictas. La burocracia española puede parecer abrumadora, pero si sigues esta guía "
        "paso a paso, podrás completar todos los trámites en 30-60 días."
    )

    doc.add_heading("Forma jurídica: autónomo vs. sociedad limitada", level=2)
    doc.add_paragraph(
        "La primera decisión es elegir la forma jurídica. Para una dark kitchen, las opciones "
        "habituales son darse de alta como autónomo (trabajador por cuenta propia) o constituir "
        "una sociedad limitada (SL). Cada opción tiene implicaciones fiscales y de responsabilidad "
        "diferentes."
    )
    doc.add_paragraph(
        "Si empiezas solo o con un socio y tu inversión es moderada (menos de 30.000€), el "
        "alta como autónomo es la opción más rápida y económica. La cuota de autónomo con tarifa "
        "plana es de 80€/mes durante los primeros 12 meses, y 230€/mes el segundo año. Sin embargo, "
        "el autónomo responde con todo su patrimonio personal ante deudas del negocio."
    )
    doc.add_paragraph(
        "Si la inversión es mayor, tienes socios o quieres limitar tu responsabilidad, la SL "
        "es más adecuada. El capital social mínimo es de 1€ (desde la reforma de la Ley Crea y "
        "Crece), aunque se recomienda aportar al menos 3.000€. El coste de constitución ronda "
        "los 600€-1.200€ con notario, registro y gestoría."
    )

    doc.add_heading("Licencia de actividad municipal", level=2)
    doc.add_paragraph(
        "Toda dark kitchen necesita una licencia de actividad (también llamada licencia de "
        "apertura o comunicación previa de actividad) del ayuntamiento donde se ubique el local. "
        "Los requisitos varían significativamente entre municipios, pero en general necesitarás:"
    )
    add_numbered("Proyecto técnico firmado por un ingeniero o arquitecto técnico.")
    add_numbered("Certificado de compatibilidad urbanística (que el local permite actividad "
                 "de restauración/obrador).")
    add_numbered("Certificado de instalación eléctrica (boletín eléctrico).")
    add_numbered("Certificado de instalación de gas (si procede).")
    add_numbered("Proyecto de ventilación y extracción de humos.")
    add_numbered("Plan de emergencia y evacuación.")
    add_numbered("Plano de distribución del local con zonas diferenciadas.")
    doc.add_paragraph(
        "El coste del proyecto técnico y la tramitación oscila entre 1.500€ y 3.500€. El plazo "
        "medio de obtención es de 30-60 días, aunque en algunos ayuntamientos la declaración "
        "responsable permite iniciar la actividad de forma inmediata mientras se tramita."
    )

    doc.add_heading("Registro sanitario", level=2)
    doc.add_paragraph(
        "Según el Reglamento (CE) 852/2004 y el Real Decreto 191/2011, todo establecimiento "
        "que elabore o manipule alimentos debe estar inscrito en el Registro General Sanitario "
        "de Empresas Alimentarias y Alimentos (RGSEAA). El trámite se realiza ante la autoridad "
        "sanitaria de la comunidad autónoma correspondiente."
    )
    doc.add_paragraph(
        "El registro sanitario es gratuito en la mayoría de comunidades autónomas, pero requiere "
        "una inspección previa del local por parte de los técnicos de sanidad, que verificarán "
        "que cumples con las normas de higiene, instalaciones y equipamiento. Es fundamental "
        "tener preparado tu plan APPCC antes de solicitar el registro."
    )

    doc.add_heading("Obligaciones fiscales", level=2)
    add_bullet("Alta en el IAE (Impuesto de Actividades Económicas): epígrafe 671.4 o similar "
               "para servicios de alimentación.")
    add_bullet("Alta censal en Hacienda (modelo 036/037).")
    add_bullet("IVA: las entregas de comida a domicilio tributan al 10% (tipo reducido).")
    add_bullet("Declaraciones trimestrales: IVA (modelo 303), IRPF (modelo 130 si autónomo), "
               "IS (modelo 202 si SL).")
    add_bullet("Libros contables obligatorios y facturación electrónica (obligatoria desde 2026).")

    doc.add_heading("Seguros obligatorios y recomendados", level=2)
    add_bullet("Seguro de responsabilidad civil: obligatorio, cubre daños a terceros por "
               "intoxicaciones o problemas derivados de la actividad. Prima anual: 300€-600€.")
    add_bullet("Seguro de local/continente: cubre daños por incendio, inundación, robo. "
               "Prima anual: 200€-400€.")
    add_bullet("Seguro de pérdida de mercancía: especialmente importante si almacenas "
               "stock perecedero. Prima anual: 100€-200€.")

    add_tip_box(
        "Contrata un gestor especializado en hostelería desde el primer día. Los 200-300€/mes "
        "que te cuesta se amortizan rápidamente al evitar errores fiscales, multas y pérdida "
        "de tiempo con la burocracia. Pide referencias a otros hosteleros de tu zona."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 6
    # ════════════════════════════════════════════
    doc.add_heading("6. APPCC y Seguridad Alimentaria", level=1)

    doc.add_paragraph(
        "El sistema de Análisis de Peligros y Puntos de Control Crítico (APPCC) es obligatorio "
        "para todo establecimiento que manipule alimentos en la Unión Europea, según el "
        "Reglamento (CE) 852/2004. Una dark kitchen no es ninguna excepción. De hecho, al "
        "no haber consumo inmediato del plato (hay un tiempo de transporte), los controles de "
        "temperatura y seguridad alimentaria son aún más críticos."
    )

    doc.add_heading("Los 7 principios del APPCC", level=2)
    add_numbered("Identificar los peligros: biológicos (bacterias, virus), químicos (productos "
                 "de limpieza, alérgenos) y físicos (objetos extraños, cristales).")
    add_numbered("Determinar los Puntos de Control Crítico (PCC): las etapas del proceso donde "
                 "un peligro puede prevenirse o eliminarse (por ejemplo, la cocción a temperatura "
                 "adecuada).")
    add_numbered("Establecer límites críticos: valores máximos o mínimos que no deben superarse "
                 "(ejemplo: temperatura de cocción interna ≥ 75°C).")
    add_numbered("Establecer un sistema de vigilancia: procedimientos para monitorizar cada PCC "
                 "(quién mide, cuándo, con qué instrumento).")
    add_numbered("Establecer medidas correctivas: qué hacer cuando un PCC se desvía de los "
                 "límites (ejemplo: desechar producto que ha superado 2 horas a temperatura ambiente).")
    add_numbered("Establecer procedimientos de verificación: auditorías periódicas para confirmar "
                 "que el sistema funciona correctamente.")
    add_numbered("Establecer un sistema de documentación: registros de todas las mediciones, "
                 "incidencias y acciones correctivas.")

    doc.add_heading("Trazabilidad y alérgenos", level=2)
    doc.add_paragraph(
        "La trazabilidad es la capacidad de rastrear un alimento a lo largo de toda la cadena "
        "de producción. En una dark kitchen debes poder identificar, para cada plato entregado, "
        "qué ingredientes se utilizaron, de qué proveedor procedían y en qué fecha se recibieron."
    )
    doc.add_paragraph(
        "En cuanto a alérgenos, el Reglamento (UE) 1169/2011 obliga a informar sobre los 14 "
        "alérgenos de declaración obligatoria en todos los platos. En delivery, esta información "
        "debe aparecer en la descripción del plato en la plataforma y, preferiblemente, en una "
        "etiqueta adherida al envase. Los 14 alérgenos son: gluten, crustáceos, huevos, pescado, "
        "cacahuetes, soja, lácteos, frutos de cáscara, apio, mostaza, sésamo, sulfitos, altramuces "
        "y moluscos."
    )

    doc.add_heading("Control de temperaturas", level=2)
    doc.add_paragraph(
        "El control de temperaturas es el punto más crítico en una dark kitchen. Las temperaturas "
        "de referencia que debes monitorizar diariamente son:"
    )
    add_bullet("Recepción de mercancía refrigerada: ≤ 4°C.")
    add_bullet("Recepción de mercancía congelada: ≤ -18°C.")
    add_bullet("Almacenamiento en cámara frigorífica: 0°C a 4°C.")
    add_bullet("Almacenamiento en congelador: ≤ -18°C.")
    add_bullet("Temperatura de cocción interna: ≥ 75°C (salvo excepciones documentadas).")
    add_bullet("Enfriamiento rápido: de 60°C a 10°C en menos de 2 horas.")
    add_bullet("Mantenimiento en caliente (servicio): ≥ 65°C.")
    add_bullet("Temperatura del plato en el momento de expedición al rider: ≥ 60°C para "
               "calientes, ≤ 8°C para fríos.")

    doc.add_heading("Registros diarios obligatorios", level=2)
    doc.add_paragraph(
        "Los registros que debes cumplimentar diariamente en tu dark kitchen incluyen:"
    )
    add_bullet("Registro de temperaturas de cámaras y congeladores (mínimo 2 veces al día).")
    add_bullet("Registro de recepción de mercancías (proveedor, temperatura, estado del producto).")
    add_bullet("Registro de limpieza y desinfección (zonas limpiadas, producto utilizado, responsable).")
    add_bullet("Registro de incidencias y acciones correctivas.")
    add_bullet("Registro de control de plagas (revisiones del contrato de DDD).")

    add_tip_box(
        "Digitaliza todos tus registros de APPCC desde el primer día. Apps como FoodDocs, "
        "Andy o incluso una simple hoja de cálculo compartida en Google Drive te ahorrarán "
        "horas de papeleo y te permitirán demostrar cumplimiento ante cualquier inspección "
        "de sanidad."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 7
    # ════════════════════════════════════════════
    doc.add_heading("7. Ubicación y Local", level=1)

    doc.add_paragraph(
        "La elección de la ubicación es una de las decisiones más importantes y difíciles de "
        "revertir. A diferencia de un restaurante tradicional, donde la visibilidad y el tránsito "
        "peatonal son fundamentales, una dark kitchen puede ubicarse en zonas secundarias o "
        "industriales, priorizando otros factores como el coste, la accesibilidad para riders "
        "y las condiciones del local."
    )

    doc.add_heading("Zona industrial vs. zona comercial", level=2)
    doc.add_paragraph(
        "Muchos emprendedores se plantean la duda: ¿es mejor un local en zona industrial "
        "(más barato, más espacio) o en zona comercial (más cerca del cliente final)? La "
        "respuesta depende de tu modelo de negocio y de la plataforma con la que trabajes."
    )
    doc.add_paragraph(
        "Las plataformas de delivery asignan un radio de cobertura a cada restaurante, "
        "habitualmente entre 3 y 5 km. Si te ubicas en un polígono industrial alejado de "
        "las zonas residenciales, tu radio de cobertura puede no alcanzar suficientes clientes "
        "potenciales. Además, los riders tardarán más en recoger los pedidos, lo que penaliza "
        "tu posicionamiento en el algoritmo de la plataforma."
    )
    doc.add_paragraph(
        "La ubicación ideal para una dark kitchen es una zona comercial secundaria o una calle "
        "de bajo tránsito dentro de una zona residencial densa. Así consigues un alquiler "
        "razonable (30-50% menos que una ubicación prime), buena cobertura de clientes en tu "
        "radio y tiempos de recogida cortos para los riders."
    )

    doc.add_heading("Requisitos mínimos del local", level=2)
    doc.add_paragraph("El local debe cumplir unos requisitos mínimos para poder operar como dark kitchen:")
    add_bullet("Superficie: mínimo 40 m² para marca única, 80-120 m² para multimarca.")
    add_bullet("Altura libre: mínimo 2,70 metros (para instalación de campanas extractoras).")
    add_bullet("Ventilación: salida de humos a cubierta (obligatorio) con campana y filtros.")
    add_bullet("Instalación eléctrica: potencia mínima de 15 kW (monofásica) o 25 kW (trifásica).")
    add_bullet("Agua caliente y fría: con suficiente caudal para limpieza continua.")
    add_bullet("Suelo: antideslizante, impermeable y de fácil limpieza (resina epoxi o gres técnico).")
    add_bullet("Paredes: revestimiento liso, lavable e impermeable hasta al menos 2 metros de altura.")
    add_bullet("Acceso exterior: zona de carga/descarga para proveedores y espacio de espera "
               "para riders (idealmente con marquesina).")

    doc.add_heading("Consejos para negociar el alquiler", level=2)
    doc.add_paragraph(
        "El alquiler será tu segundo mayor coste fijo tras la nómina. Negociarlo bien puede "
        "marcar la diferencia entre un negocio rentable y uno que pierde dinero. Algunos consejos "
        "prácticos:"
    )
    add_bullet("Negocia un período de carencia de 2-3 meses para la fase de obra y obtención "
               "de licencias. No deberías pagar alquiler mientras no puedes operar.")
    add_bullet("Solicita un contrato de larga duración (mínimo 5 años) para asegurar la inversión "
               "en el acondicionamiento del local.")
    add_bullet("Incluye una cláusula de revisión de renta vinculada al IPC con un tope máximo anual.")
    add_bullet("Verifica que el propietario permite la actividad de dark kitchen: no todos los "
               "contratos de alquiler de locales comerciales contemplan actividad de cocina.")
    add_bullet("Si el local necesita obra significativa, negocia que el propietario asuma parte "
               "del coste o lo descuente del alquiler durante los primeros meses.")

    add_tip_box(
        "Antes de firmar nada, consulta el plan urbanístico municipal (PGOU) de tu ayuntamiento "
        "para confirmar que la actividad de obrador/cocina industrial está permitida en esa "
        "parcela. Muchos emprendedores han perdido meses y miles de euros al descubrir que su "
        "local no admitía la actividad después de firmado el contrato."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 8
    # ════════════════════════════════════════════
    doc.add_heading("8. Diseño y Layout de Cocina", level=1)

    doc.add_paragraph(
        "El diseño del layout de una cocina de delivery es radicalmente diferente al de un "
        "restaurante tradicional. En un restaurante, el flujo va de cocina a sala; en una dark "
        "kitchen, el flujo va de recepción de mercancía a zona de expedición para riders. Cada "
        "metro cuadrado debe estar optimizado para la eficiencia, la velocidad y la seguridad "
        "alimentaria."
    )

    doc.add_heading("Principios del flujo de trabajo delivery-only", level=2)
    doc.add_paragraph(
        "El diseño de la cocina debe seguir el principio de flujo lineal o en U, evitando "
        "cruces entre zonas sucias (recepción, lavado) y zonas limpias (preparación, emplatado). "
        "El recorrido ideal de un producto es:"
    )
    add_numbered("Recepción y almacenamiento: zona de descarga → inspección → cámaras/almacén.")
    add_numbered("Preparación previa (mise en place): zona de corte, pelado, porcionado.")
    add_numbered("Cocción: zona caliente con planchas, freidoras, hornos, woks.")
    add_numbered("Emplatado/montaje: zona limpia donde se ensamblan los pedidos.")
    add_numbered("Expedición: zona de sellado, etiquetado y entrega a riders.")

    doc.add_paragraph(
        "La zona de expedición es especialmente importante en una dark kitchen. Debe estar "
        "claramente separada de la cocina, con un mostrador o estantería térmica donde los "
        "pedidos terminados esperan a ser recogidos por los riders. Idealmente, los riders "
        "no deberían entrar en la zona de cocina."
    )

    doc.add_heading("Diseño multimarca: cómo compartir espacio sin caos", level=2)
    doc.add_paragraph(
        "Si operas varias marcas desde la misma cocina, el diseño debe contemplar estaciones "
        "de trabajo independientes para cada marca, pero compartiendo equipamiento cuando sea "
        "posible. Por ejemplo:"
    )
    add_bullet("Una marca de hamburguesas y una de alitas pueden compartir freidora y plancha.")
    add_bullet("Una marca de poke bowls y una de ensaladas pueden compartir la zona de "
               "preparación fría.")
    add_bullet("El almacenamiento y la expedición son siempre compartidos.")

    doc.add_paragraph(
        "La clave es que cada marca tenga su propia zona de montaje/emplatado, con los "
        "ingredientes específicos a mano, para no perder tiempo buscando componentes. Utiliza "
        "códigos de colores en bandejas y etiquetas para evitar confusiones entre marcas."
    )

    doc.add_heading("Zonas obligatorias y opcionales", level=2)
    doc.add_paragraph("Según la normativa sanitaria, tu cocina debe tener como mínimo:")
    add_bullet("Zona de recepción de mercancías (diferenciada de la zona de expedición).")
    add_bullet("Zona de almacenamiento en seco, refrigerado y congelado.")
    add_bullet("Zona de preparación/manipulación (con diferenciación crudo/cocido si es posible).")
    add_bullet("Zona de cocción.")
    add_bullet("Zona de lavado de utensilios y vajilla (con lavavajillas industrial).")
    add_bullet("Zona de lavado de manos (independiente del fregadero de utensilios).")
    add_bullet("Vestuario para el personal (si hay más de 2 trabajadores).")
    add_bullet("Almacén de productos de limpieza (separado de alimentos).")
    add_bullet("Cubo de basura con tapa y pedal en cada zona de trabajo.")

    add_tip_box(
        "Invierte en un buen sistema de organización visual (estanterías etiquetadas, zonas "
        "marcadas con cinta de color en el suelo, estaciones de trabajo identificadas). En "
        "momentos de máxima presión durante el servicio, la organización visual evita errores, "
        "reduce el estrés y mejora los tiempos de preparación."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 9
    # ════════════════════════════════════════════
    doc.add_heading("9. Equipamiento y Proveedores", level=1)

    doc.add_paragraph(
        "El equipamiento de una dark kitchen debe priorizar la eficiencia, la velocidad y la "
        "durabilidad. No necesitas el equipamiento más caro del mercado, pero tampoco puedes "
        "escatimar en elementos críticos que afectan a la calidad del producto y a los tiempos "
        "de preparación. Este capítulo incluye una lista detallada con precios orientativos "
        "del mercado español."
    )

    doc.add_heading("Equipamiento esencial de cocina caliente", level=2)
    add_bullet("Horno de convección profesional (6-10 bandejas GN 1/1): 2.500€ - 5.000€. "
               "Imprescindible para hornear, regenerar y gratinar. Marcas recomendadas: Rational, "
               "Unox, Fagor.")
    add_bullet("Plancha/grill industrial (60-80 cm): 800€ - 1.500€. Superficie lisa + estriada "
               "para hamburguesas, carnes y verduras.")
    add_bullet("Freidora industrial (1-2 cubas, 8-12 litros): 800€ - 1.500€. Con termostato "
               "digital y cestillo de seguridad.")
    add_bullet("Cocina de inducción (2-4 fuegos): 600€ - 1.200€. Más eficiente y segura que el gas, "
               "ideal para dark kitchens con ventilación limitada.")
    add_bullet("Campana extractora con filtros antigrasa: 1.200€ - 3.000€. Dimensionada según "
               "la superficie de cocción.")
    add_bullet("Baño maría o marmita: 300€ - 600€. Para mantener salsas y bases en temperatura.")

    doc.add_heading("Equipamiento de frío y almacenamiento", level=2)
    add_bullet("Cámara frigorífica o armario refrigerado (700-1.400 litros): 1.800€ - 3.500€.")
    add_bullet("Congelador arcón o armario congelador: 800€ - 1.500€.")
    add_bullet("Mesa refrigerada (2-3 puertas): 1.500€ - 2.800€. Ideal como estación de "
               "trabajo con ingredientes a mano.")
    add_bullet("Estanterías de acero inoxidable (4-5 niveles): 150€ - 300€ por unidad. Necesitarás "
               "al menos 3-4 unidades para almacenamiento en seco.")
    add_bullet("Contenedores GN de policarbonato con tapa: 200€ - 400€ el lote completo. "
               "Fundamentales para la organización de mise en place.")

    doc.add_heading("Equipamiento de preparación", level=2)
    add_bullet("Mesa de trabajo de acero inoxidable (1,5 m x 0,7 m): 300€ - 500€ por unidad.")
    add_bullet("Cortadora de fiambres profesional: 400€ - 800€.")
    add_bullet("Robot de cocina / cutter: 500€ - 1.200€. Para salsas, cremas y masas.")
    add_bullet("Batidora de brazo industrial: 200€ - 400€.")
    add_bullet("Envasadora al vacío: 300€ - 800€. Esencial para conservación y mise en place avanzada.")

    doc.add_heading("Equipamiento de limpieza", level=2)
    add_bullet("Lavavajillas industrial (cesta 50x50): 1.500€ - 3.000€. Con ciclos de 2-3 minutos.")
    add_bullet("Fregadero doble de acero inoxidable: 300€ - 500€.")
    add_bullet("Lavamanos con accionamiento no manual (pedal o sensor): 150€ - 300€.")

    doc.add_heading("Nuevo vs. segunda mano", level=2)
    doc.add_paragraph(
        "Comprar equipamiento de segunda mano puede ahorrarte un 40-60% en la inversión inicial. "
        "Portales como Wallapop, Milanuncios y webs especializadas como Hostelmarkt o Equipos de "
        "Hostelería Usados ofrecen maquinaria en buen estado procedente de restaurantes que han "
        "cerrado. Sin embargo, hay que tener precaución:"
    )
    add_bullet("Siempre comprueba el funcionamiento in situ antes de comprar.")
    add_bullet("Verifica que el equipamiento cumple la normativa CE vigente.")
    add_bullet("No compres de segunda mano: lavavajillas (pieza crítica que debe funcionar "
               "perfectamente) ni sistemas de refrigeración (un fallo puede suponer pérdida "
               "de mercancía y problemas sanitarios).")
    add_bullet("Sí puedes comprar de segunda mano: mesas de trabajo, estanterías, pequeño "
               "menaje y utensilios, hornos (si tienen menos de 5 años de uso).")

    add_tip_box(
        "Haz una lista priorizada del equipamiento: primero lo imprescindible para abrir, "
        "después lo deseable. Puedes empezar con el mínimo viable y añadir equipamiento "
        "según crece tu facturación. No inviertas en un horno Rational de 8.000€ si con uno "
        "de 3.000€ puedes cubrir tu producción durante los primeros 6 meses."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 10
    # ════════════════════════════════════════════
    doc.add_heading("10. Tecnología y Gestión de Pedidos", level=1)

    doc.add_paragraph(
        "La tecnología es el sistema nervioso de una dark kitchen. A diferencia de un restaurante "
        "tradicional donde el camarero lleva la comanda a la cocina, en una dark kitchen los "
        "pedidos llegan simultáneamente desde múltiples plataformas digitales. Sin un sistema "
        "tecnológico robusto, el caos es inevitable. Este capítulo cubre las herramientas "
        "esenciales y cómo integrarlas."
    )

    doc.add_heading("El problema de las tablets múltiples", level=2)
    doc.add_paragraph(
        "Cada plataforma de delivery (Glovo, Uber Eats, Just Eat) requiere su propia tablet "
        "para recibir y gestionar pedidos. Si operas con las tres plataformas principales, "
        "tendrás mínimo tres tablets sonando simultáneamente durante el servicio. Si además "
        "operas varias marcas, la cantidad de tablets se multiplica. Este sistema es inviable "
        "a partir de un cierto volumen."
    )

    doc.add_heading("Integradores de pedidos (aggregators)", level=2)
    doc.add_paragraph(
        "La solución al caos de las tablets son los integradores de pedidos. Estas plataformas "
        "reciben todos los pedidos de todas las plataformas en un único panel centralizado, "
        "los organizan por tiempo de preparación y los envían a una impresora de tickets o a "
        "un KDS (Kitchen Display System). Las opciones más populares en España son:"
    )
    add_bullet("Ordatic: integrador español, muy popular en el mercado nacional. Integra Glovo, "
               "Uber Eats, Just Eat y canal propio. Precio: desde 69€/mes por establecimiento.")
    add_bullet("Deliverect: integrador internacional con buena presencia en España. Ofrece "
               "gestión de menú centralizada para actualizar precios y disponibilidad en todas "
               "las plataformas a la vez. Precio: desde 79€/mes.")
    add_bullet("Hubrise: integrador técnico que conecta plataformas de delivery con TPV y "
               "sistemas de gestión. Más orientado a cadenas. Precio: desde 25€/mes por punto.")
    add_bullet("Otter (by Uber): integrador gratuito para restaurantes que operan en Uber Eats, "
               "con integraciones adicionales de pago para otras plataformas.")

    doc.add_heading("KDS (Kitchen Display System)", level=2)
    doc.add_paragraph(
        "Un KDS es una pantalla instalada en la cocina que muestra los pedidos en tiempo real, "
        "organizados por prioridad y tiempo de preparación. Sustituye a la impresora de tickets "
        "tradicional y aporta ventajas significativas:"
    )
    add_bullet("Visualización clara de todos los pedidos pendientes con temporizadores.")
    add_bullet("Código de colores según urgencia (verde = a tiempo, amarillo = ajusto, rojo = tarde).")
    add_bullet("Posibilidad de marcar pedidos como preparados con un toque.")
    add_bullet("Historial de tiempos de preparación para análisis de eficiencia.")
    add_bullet("Sin papel, sin tinta, sin atascos de impresora en pleno servicio.")

    doc.add_heading("Software de gestión de inventario y costes", level=2)
    doc.add_paragraph(
        "Controlar el food cost en una dark kitchen es vital para la supervivencia del negocio. "
        "Con márgenes tan ajustados (recuerda, un 30% se va en comisiones de plataforma), cada "
        "céntimo de desperdicio cuenta. Herramientas recomendadas:"
    )
    add_bullet("Marketman: gestión de inventario, pedidos a proveedores y food cost en tiempo "
               "real. Integración con TPV. Precio: desde 199€/mes.")
    add_bullet("Apicbase: plataforma de gestión de back-of-house que incluye recetas, fichas "
               "técnicas, inventario y APPCC digital. Precio: desde 149€/mes.")
    add_bullet("Tspoonlab: herramienta española para gestión de recetas, escandallos y etiquetado "
               "nutricional. Precio: desde 40€/mes.")
    add_bullet("Hojas de cálculo avanzadas: para operaciones pequeñas, un buen Excel con "
               "escandallos y control de inventario semanal puede ser suficiente (consulta nuestro "
               "Kit de Escandallos Pro en aichef.pro).")

    add_tip_box(
        "La tecnología debe simplificar tu operación, no complicarla. Empieza con lo mínimo "
        "viable: un integrador de pedidos y un buen Excel para el food cost. Solo cuando tu "
        "volumen supere los 50 pedidos diarios, invierte en KDS y software avanzado de inventario."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 11
    # ════════════════════════════════════════════
    doc.add_heading("11. Packaging y Sostenibilidad", level=1)

    doc.add_paragraph(
        "El packaging es la primera impresión física que el cliente tiene de tu marca. En una "
        "dark kitchen, donde no hay experiencia en sala ni interacción personal, el envase "
        "se convierte en tu embajador. Un buen packaging no solo protege la comida durante "
        "el transporte: comunica calidad, profesionalidad y valores de marca. Además, la "
        "regulación medioambiental cada vez es más estricta, y adaptarse es obligatorio."
    )

    doc.add_heading("Requisitos del packaging para delivery", level=2)
    doc.add_paragraph(
        "El envase ideal para delivery debe cumplir simultáneamente varias funciones que en un "
        "restaurante tradicional son irrelevantes:"
    )
    add_bullet("Mantener la temperatura: caliente lo caliente, frío lo frío. Los envases con "
               "doble pared o cámara de aire son más efectivos que los de una sola capa.")
    add_bullet("Evitar fugas y derrames: el transporte en mochila de rider es violento. "
               "Los envases deben tener cierre hermético o, como mínimo, precintado.")
    add_bullet("Preservar la textura: separar elementos crujientes de húmedos. Las patatas "
               "fritas que llegan blandas son la queja número uno en delivery.")
    add_bullet("Permitir la ventilación adecuada: los envases completamente herméticos pueden "
               "condensar vapor y ablandar productos fritos o crujientes. Pequeños orificios "
               "de ventilación pueden ser la solución.")
    add_bullet("Ser apilable y estable: los riders apilan los pedidos en la mochila. Envases "
               "con base ancha y tapa plana se apilan mejor.")

    doc.add_heading("Tipos de envases y costes aproximados", level=2)
    add_bullet("Envases de cartón kraft con tapa (hamburguesas, bowls): 0,15€ - 0,35€ por unidad.")
    add_bullet("Envases de bagazo de caña de azúcar (alternativa ecológica): 0,20€ - 0,40€.")
    add_bullet("Envases de PP apto para microondas (comida caliente): 0,10€ - 0,25€.")
    add_bullet("Envases de rPET (ensaladas, poke, comida fría): 0,12€ - 0,30€.")
    add_bullet("Bolsas de papel kraft con asa: 0,08€ - 0,15€.")
    add_bullet("Pegatinas de precinto con logo: 0,02€ - 0,05€ por unidad (pedido mínimo 1.000).")
    add_bullet("Servilletas, cubiertos, salseras: 0,10€ - 0,20€ por pedido.")

    doc.add_heading("Normativa de plásticos en España (2026)", level=2)
    doc.add_paragraph(
        "La regulación medioambiental sobre envases de plástico se ha endurecido significativamente "
        "en los últimos años. Las normas clave que afectan a las dark kitchens son:"
    )
    add_bullet("Real Decreto 1055/2022 de envases y residuos de envases: obliga a ofrecer "
               "alternativas reutilizables y establece objetivos de reducción de plástico de "
               "un solo uso.")
    add_bullet("Impuesto al plástico: 0,45€ por kg de plástico no reciclado en envases. "
               "Esto encarece los envases de plástico virgen en un 5-10%.")
    add_bullet("Prohibición de plásticos de un solo uso: cubiertos, pajitas, agitadores y "
               "recipientes de poliestireno expandido (EPS) están prohibidos desde 2023.")
    add_bullet("Obligación de cobrar los envases reutilizables: a partir de 2026, los "
               "establecimientos de comida a domicilio deberán ofrecer la opción de envases "
               "reutilizables con sistema de depósito.")

    doc.add_heading("Branding en el packaging", level=2)
    doc.add_paragraph(
        "El packaging es tu única oportunidad de hacer marca fuera de la plataforma digital. "
        "Aprovéchalo al máximo con estos elementos:"
    )
    add_bullet("Logo y colores de marca en la bolsa y en al menos un envase.")
    add_bullet("Tarjeta de agradecimiento con código QR a tu web/redes sociales.")
    add_bullet("Código de descuento para el próximo pedido directo (sin comisión de plataforma).")
    add_bullet("Información nutricional y de alérgenos visible en el envase.")
    add_bullet("Un toque personal: una frase, un dibujo, algo que humanice la experiencia y "
               "genere una conexión emocional con el cliente.")

    add_tip_box(
        "Invierte en una prueba de transporte antes de elegir tus envases definitivos. Prepara "
        "tus platos más delicados, empaquétalos, mételos en una mochila de rider y conduce "
        "15 minutos en moto por tu ciudad. Al abrir, verás exactamente qué funciona y qué no. "
        "Esta prueba te ahorrará cientos de reseñas negativas."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CHAPTER 12
    # ════════════════════════════════════════════
    doc.add_heading("12. Marketing, Lanzamiento y Errores Comunes", level=1)

    doc.add_paragraph(
        "Has montado tu cocina, tienes tus licencias, tu APPCC está listo y tu menú es espectacular. "
        "Ahora viene la parte que muchos chef-emprendedores subestiman: el marketing y el "
        "lanzamiento. En delivery, la visibilidad lo es todo. Si no apareces en las primeras "
        "posiciones de la plataforma, simplemente no existes. Este capítulo te guía en el "
        "lanzamiento y te advierte de los errores que debes evitar a toda costa."
    )

    doc.add_heading("Onboarding en plataformas: paso a paso", level=2)
    doc.add_paragraph(
        "El proceso de alta en cada plataforma es similar, pero los detalles importan porque "
        "determinan tu posicionamiento inicial:"
    )
    add_numbered("Registro en el portal de partners de cada plataforma (Glovo Partners, "
                 "Uber Eats Restaurant Dashboard, Just Eat Partner).")
    add_numbered("Envío de documentación: CIF/NIF, licencia de actividad, registro sanitario, "
                 "seguro de responsabilidad civil, datos bancarios.")
    add_numbered("Configuración del menú: nombre de los platos, descripciones atractivas, "
                 "precios, fotografías, alérgenos, opciones de personalización.")
    add_numbered("Configuración operativa: horarios, tiempo de preparación estimado, radio "
                 "de cobertura, pedido mínimo.")
    add_numbered("Sesión de fotos profesional: es el factor que más influye en la tasa de "
                 "conversión. Las plataformas suelen ofrecer un servicio de fotografía gratuito "
                 "o subvencionado.")
    add_numbered("Periodo de prueba (1-2 semanas): la plataforma monitoriza tu rendimiento "
                 "(aceptación de pedidos, tiempos de preparación, valoraciones).")

    doc.add_heading("Fotografía gastronómica para delivery", level=2)
    doc.add_paragraph(
        "Las fotografías de tus platos son el equivalente al escaparate de una tienda. Según "
        "datos de las propias plataformas, los restaurantes con fotografías profesionales "
        "reciben un 65% más de pedidos que los que usan fotos de móvil. Consejos para conseguir "
        "buenas fotos sin arruinarte:"
    )
    add_bullet("Contrata un fotógrafo especializado en gastronomía. Precio orientativo: "
               "200€-500€ por una sesión de 15-20 platos. Es la mejor inversión en marketing.")
    add_bullet("Si haces las fotos tú mismo: usa luz natural (junto a una ventana), fondo "
               "neutro, ángulo cenital (desde arriba) y asegúrate de que el plato esté recién "
               "preparado y con su mejor aspecto.")
    add_bullet("Incluye ingredientes visibles que despierten el apetito: el queso fundido "
               "cayendo, las hierbas frescas como toque final, la salsa brillante.")
    add_bullet("Utiliza herramientas de IA como Midjourney o DALL·E para generar imágenes "
               "de referencia y planificar la composición antes de la sesión real.")

    doc.add_heading("Estrategia de lanzamiento (primeros 30 días)", level=2)
    doc.add_paragraph(
        "Los primeros 30 días son críticos. Las plataformas suelen dar un impulso de visibilidad "
        "a los nuevos restaurantes (el llamado 'new restaurant boost'), que dura entre 2 y 4 "
        "semanas. Aprovéchalo al máximo:"
    )
    add_bullet("Semana 1: activa todas las plataformas simultáneamente con una oferta de "
               "lanzamiento agresiva (2x1, envío gratis, 30% de descuento).")
    add_bullet("Semana 2: reduce la oferta al 20% de descuento y centra esfuerzos en obtener "
               "reseñas positivas. Cada pedido que salga perfecto es una reseña potencial.")
    add_bullet("Semana 3: activa campañas de marketing dentro de las plataformas (Glovo Ads, "
               "Uber Eats Ads). Destina un 5-8% de tu facturación a publicidad.")
    add_bullet("Semana 4: analiza métricas, ajusta precios y carta según los datos de las "
               "primeras semanas. Elimina los platos con baja demanda y potencia los bestsellers.")

    doc.add_heading("Los 10 errores fatales de una dark kitchen", level=2)
    doc.add_paragraph(
        "Después de asesorar a decenas de proyectos de dark kitchen, estos son los errores "
        "que he visto repetirse una y otra vez y que pueden hundir tu negocio:"
    )
    add_numbered("No hacer números antes de empezar. Muchos emprendedores subestiman las "
                 "comisiones de las plataformas (25-35%) y se encuentran con márgenes negativos.")
    add_numbered("Carta demasiado amplia. Un menú de 40 platos es inviable en una dark kitchen. "
                 "Máximo 15-20 referencias por marca, optimizadas para delivery.")
    add_numbered("Fotos de mala calidad. Es el factor número uno que determina si un cliente "
                 "hace clic en tu restaurante o pasa de largo.")
    add_numbered("No controlar el food cost semanalmente. El food cost debe revisarse cada "
                 "semana. Un desvío de 3 puntos porcentuales puede significar pérdidas mensuales.")
    add_numbered("Ignorar las valoraciones y reseñas. Responde a TODAS las reseñas, especialmente "
                 "las negativas. Las plataformas valoran la interacción del restaurante.")
    add_numbered("Depender de una sola plataforma. Si Glovo cambia sus condiciones o te "
                 "desactiva, pierdes el 100% de tus ingresos. Diversifica siempre.")
    add_numbered("Ahorrar en packaging. El envase barato que se abre durante el transporte "
                 "genera reseñas de una estrella que te hunden en el ranking.")
    add_numbered("No invertir en marketing de lanzamiento. El boost de nuevo restaurante es "
                 "una oportunidad única. Si no la aprovechas con promociones agresivas, la pierdes.")
    add_numbered("Subestimar la importancia del tiempo de preparación. Cada minuto extra de "
                 "espera del rider penaliza tu posicionamiento. Optimiza tu flujo de trabajo "
                 "hasta que el tiempo medio sea inferior a 12 minutos.")
    add_numbered("No tener plan B. ¿Qué pasa si se avería tu cámara frigorífica un viernes "
                 "por la noche? Ten un plan de contingencia para las situaciones más comunes: "
                 "avería de equipos, falta de personal, corte de suministro eléctrico.")

    add_tip_box(
        "El éxito de una dark kitchen se mide en tres KPIs fundamentales: tiempo medio de "
        "preparación (objetivo: <12 minutos), valoración media en plataformas (objetivo: >4,5 "
        "estrellas) y food cost (objetivo: <32%). Si controlas estos tres números semanalmente, "
        "tu negocio será viable. Si los ignoras, estás en peligro."
    )

    # ════════════════════════════════════════════
    # FINAL PAGE
    # ════════════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        doc.add_paragraph()

    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_final.add_run("Guía elaborada por Chef John Guerrero")
    run_f.font.size = Pt(14)
    run_f.font.bold = True

    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = p_brand.add_run("AI Chef Pro · aichef.pro")
    run_b.font.size = Pt(12)
    run_b.font.color.rgb = GOLD_RGB

    p_exp = doc.add_paragraph()
    p_exp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_e = p_exp.add_run("29 años en alta hostelería · 15 años en consultoría gastronómica")
    run_e.font.size = Pt(10)
    run_e.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p_copy = doc.add_paragraph()
    p_copy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_copy.add_run("© 2026 AI Chef Pro. Todos los derechos reservados.")
    run_c.font.size = Pt(9)
    run_c.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ct = p_contact.add_run("Contacto: info@aichef.pro")
    run_ct.font.size = Pt(10)
    run_ct.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    path = os.path.join(OUTPUT_DIR, "guia-como-montar-dark-kitchen.docx")
    doc.save(path)
    print(f"✅ DOCX generado: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# FILE 2: Checklist apertura legal (35 items)
# ═══════════════════════════════════════════════════════════
def generate_checklist_legal():
    wb = Workbook()
    ws = wb.active
    ws.title = "Checklist Apertura Legal"

    add_title_block(ws, "Checklist de Apertura Legal — Dark Kitchen",
                    BRAND_LINE, merge_to=8)

    headers = ["#", "Categoría", "Tarea", "Responsable", "Fecha Límite",
               "Coste Estimado (€)", "Estado", "Notas"]
    row = 4
    for c, h in enumerate(headers, 1):
        ws.cell(row=row, column=c, value=h)
    style_header_row(ws, row, len(headers))

    items = [
        # Constitución Empresa
        ("Constitución Empresa", "Decidir forma jurídica (autónomo o SL)", "Titular", 300),
        ("Constitución Empresa", "Reservar denominación social en el Registro Mercantil (si SL)", "Gestor", 50),
        ("Constitución Empresa", "Escritura pública ante notario (si SL)", "Notario", 400),
        ("Constitución Empresa", "Inscripción en el Registro Mercantil (si SL)", "Gestor", 200),
        ("Constitución Empresa", "Obtención del CIF definitivo", "Gestor", 0),
        # Licencias Municipales
        ("Licencias Municipales", "Consultar compatibilidad urbanística del local", "Arquitecto", 150),
        ("Licencias Municipales", "Contratar arquitecto/ingeniero para proyecto técnico", "Titular", 2000),
        ("Licencias Municipales", "Presentar declaración responsable o solicitud de licencia de actividad", "Arquitecto", 500),
        ("Licencias Municipales", "Obtener boletín eléctrico", "Electricista", 200),
        ("Licencias Municipales", "Obtener certificado de instalación de gas (si procede)", "Instalador", 150),
        ("Licencias Municipales", "Presentar proyecto de extracción de humos y ventilación", "Arquitecto", 0),
        # Sanidad
        ("Sanidad", "Elaborar plan APPCC completo", "Consultor", 800),
        ("Sanidad", "Inscripción en el RGSEAA (Registro Sanitario)", "Titular", 0),
        ("Sanidad", "Solicitar inspección sanitaria previa a la apertura", "Titular", 0),
        ("Sanidad", "Preparar registros de trazabilidad y control de alérgenos", "Consultor", 200),
        ("Sanidad", "Contratar servicio de Desinsectación, Desratización y Desinfección (DDD)", "Titular", 300),
        ("Sanidad", "Obtener certificados de formación en manipulación de alimentos", "Empleados", 50),
        ("Sanidad", "Implementar sistema de registro de temperaturas", "Titular", 100),
        # Seguros
        ("Seguros", "Contratar seguro de responsabilidad civil", "Corredor", 400),
        ("Seguros", "Contratar seguro de local/continente", "Corredor", 300),
        ("Seguros", "Contratar seguro de pérdida de mercancía (opcional)", "Corredor", 150),
        ("Seguros", "Contratar seguro de accidentes laborales (convenio hostelería)", "Corredor", 200),
        # Hacienda / AEAT
        ("Hacienda/AEAT", "Alta censal en Hacienda (modelo 036 o 037)", "Gestor", 0),
        ("Hacienda/AEAT", "Alta en el IAE (epígrafe 671.4 o similar)", "Gestor", 0),
        ("Hacienda/AEAT", "Alta como autónomo en la Seguridad Social (si procede)", "Gestor", 80),
        ("Hacienda/AEAT", "Alta de trabajadores en la Seguridad Social", "Gestor", 0),
        ("Hacienda/AEAT", "Configurar facturación electrónica (obligatoria 2026)", "Gestor", 200),
        ("Hacienda/AEAT", "Apertura de cuenta bancaria empresarial", "Titular", 0),
        # Protección de Datos
        ("Protección de Datos", "Registro de actividades de tratamiento de datos (RGPD)", "Consultor", 300),
        ("Protección de Datos", "Redactar política de privacidad para web/app", "Consultor", 0),
        ("Protección de Datos", "Implementar sistema de gestión de consentimientos", "Titular", 100),
        ("Protección de Datos", "Videovigilancia: cartelería LOPD y registro de cámaras", "Consultor", 50),
        # APPCC
        ("APPCC", "Calibrar termómetros de sonda y cámaras", "Titular", 100),
        ("APPCC", "Elaborar plano de puntos de control crítico (PCC)", "Consultor", 0),
        ("APPCC", "Crear archivador/sistema digital de registros APPCC", "Titular", 50),
    ]

    # Data validation for Estado
    dv = DataValidation(type="list", formula1='"Pendiente,En Proceso,Completado"',
                        allow_blank=True, showDropDown=False)
    dv.error = "Seleccione un estado válido"
    dv.errorTitle = "Estado inválido"
    ws.add_data_validation(dv)

    for i, (cat, tarea, resp, coste) in enumerate(items, 1):
        r = row + i
        style_data_cell(ws, r, 1, i, align=center_align)
        style_data_cell(ws, r, 2, cat)
        style_data_cell(ws, r, 3, tarea)
        style_data_cell(ws, r, 4, resp, align=center_align)
        style_data_cell(ws, r, 5, None, align=center_align)  # Fecha
        ws.cell(row=r, column=5).number_format = 'DD/MM/YYYY'
        style_data_cell(ws, r, 6, coste if coste > 0 else 0, fmt=currency_format, align=center_align)
        style_data_cell(ws, r, 7, "Pendiente", align=center_align)
        dv.add(ws.cell(row=r, column=7))
        style_data_cell(ws, r, 8, "")
        # Alternate row shading
        if i % 2 == 0:
            for c in range(1, 9):
                ws.cell(row=r, column=c).fill = light_fill

    # Column widths
    widths = [5, 22, 55, 14, 14, 18, 14, 30]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Total row
    total_row = row + len(items) + 1
    ws.cell(row=total_row, column=5, value="TOTAL").font = bold_font_xl
    ws.cell(row=total_row, column=5).alignment = center_align
    ws.cell(row=total_row, column=6).value = f"=SUM(F5:F{total_row-1})"
    ws.cell(row=total_row, column=6).font = formula_font_xl
    ws.cell(row=total_row, column=6).number_format = currency_format
    ws.cell(row=total_row, column=6).alignment = center_align

    path = os.path.join(OUTPUT_DIR, "checklist-apertura-legal.xlsx")
    wb.save(path)
    print(f"✅ Checklist legal generado: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# FILE 3: Checklist equipamiento y obra (40 items)
# ═══════════════════════════════════════════════════════════
def generate_checklist_equipamiento():
    wb = Workbook()
    ws = wb.active
    ws.title = "Checklist Equipamiento"

    add_title_block(ws, "Checklist de Equipamiento y Obra — Dark Kitchen",
                    BRAND_LINE, merge_to=9)

    headers = ["#", "Zona", "Equipo/Concepto", "Proveedor", "Presupuesto (€)",
               "Real (€)", "Desviación (%)", "Estado", "Notas"]
    row = 4
    for c, h in enumerate(headers, 1):
        ws.cell(row=row, column=c, value=h)
    style_header_row(ws, row, len(headers))

    items = [
        # Obra Civil
        ("Obra Civil", "Acondicionamiento general del local (suelos, paredes, techos)", "", 6000),
        ("Obra Civil", "Revestimiento de suelos (resina epoxi antideslizante)", "", 2500),
        ("Obra Civil", "Alicatado de paredes hasta 2 m de altura", "", 1800),
        ("Obra Civil", "Puerta de acceso de riders con cierre automático", "", 500),
        ("Obra Civil", "Pintura y acabados generales", "", 800),
        ("Obra Civil", "Aseo y vestuario para personal", "", 1200),
        # Instalaciones
        ("Instalaciones", "Instalación eléctrica completa (cuadro, cableado, tomas)", "", 3500),
        ("Instalaciones", "Instalación de fontanería (agua fría y caliente)", "", 1500),
        ("Instalaciones", "Instalación de gas (acometida y canalizaciones)", "", 1200),
        ("Instalaciones", "Sistema de extracción y ventilación (campana + conductos)", "", 3500),
        ("Instalaciones", "Sistema contra incendios (extintores, BIE, señalización)", "", 600),
        ("Instalaciones", "Iluminación LED para cocina (norma mínima 500 lux)", "", 800),
        # Cocina Caliente
        ("Cocina Caliente", "Horno de convección profesional (10 bandejas GN 1/1)", "Rational/Unox", 3500),
        ("Cocina Caliente", "Plancha/grill industrial 80 cm (lisa + estriada)", "Repagas", 1200),
        ("Cocina Caliente", "Freidora industrial doble cuba (2x10L)", "Fagor", 1200),
        ("Cocina Caliente", "Cocina de inducción 4 fuegos", "Hendi", 1000),
        ("Cocina Caliente", "Campana extractora con filtros antigrasa", "Halton", 2500),
        ("Cocina Caliente", "Salamandra/gratinador", "Repagas", 600),
        ("Cocina Caliente", "Baño maría 3 cubetas GN", "", 350),
        # Cocina Fría
        ("Cocina Fría", "Cámara frigorífica walk-in (6 m³)", "Infricool", 2800),
        ("Cocina Fría", "Armario congelador vertical (700L)", "Coreco", 1500),
        ("Cocina Fría", "Mesa refrigerada 3 puertas", "Infricool", 2200),
        ("Cocina Fría", "Cortadora de fiambres profesional", "Sammic", 600),
        ("Cocina Fría", "Envasadora al vacío de campana", "Sammic", 700),
        # Almacenamiento
        ("Almacenamiento", "Estanterías de acero inoxidable (x4 unidades)", "Distform", 900),
        ("Almacenamiento", "Contenedores GN de policarbonato con tapa (lote completo)", "", 350),
        ("Almacenamiento", "Estantería para almacén seco con ruedas", "", 250),
        ("Almacenamiento", "Contenedores de ingredientes con dosificador (x10)", "", 150),
        # Expedición
        ("Expedición", "Estantería térmica de expedición con calor y frío", "", 1200),
        ("Expedición", "Mesa de emplatado/montaje de acero inoxidable", "Distform", 450),
        ("Expedición", "Dispensador de bolsas y envases de delivery", "", 200),
        ("Expedición", "Precintadora/selladora de envases", "", 300),
        ("Expedición", "Impresora de etiquetas para pedidos", "Brother", 250),
        # Tecnología
        ("Tecnología", "Tablets para plataformas de delivery (x3)", "", 600),
        ("Tecnología", "KDS (Kitchen Display System) pantalla + soporte", "Ordatic", 500),
        ("Tecnología", "Router WiFi de alta velocidad (fibra)", "", 100),
        ("Tecnología", "Software integrador de pedidos (12 meses)", "Ordatic/Deliverect", 900),
        ("Tecnología", "TPV con software de gestión", "", 800),
        ("Tecnología", "Cámaras de seguridad (4 unidades + grabador)", "", 400),
        ("Tecnología", "Sistema de música ambiental (opcional)", "", 150),
    ]

    dv = DataValidation(type="list", formula1='"Pendiente,Pedido,Recibido,Instalado"',
                        allow_blank=True, showDropDown=False)
    ws.add_data_validation(dv)

    for i, (zona, equipo, prov, presup) in enumerate(items, 1):
        r = row + i
        style_data_cell(ws, r, 1, i, align=center_align)
        style_data_cell(ws, r, 2, zona)
        style_data_cell(ws, r, 3, equipo)
        style_data_cell(ws, r, 4, prov, align=center_align)
        style_data_cell(ws, r, 5, presup, fmt=currency_format, align=center_align, fill=input_fill)
        style_data_cell(ws, r, 6, None, fmt=currency_format, align=center_align, fill=input_fill)
        # Deviation formula
        cell_dev = style_data_cell(ws, r, 7, f'=IF(E{r}=0,"",((F{r}-E{r})/E{r}))', fmt=pct_format, align=center_align)
        style_data_cell(ws, r, 8, "Pendiente", align=center_align)
        dv.add(ws.cell(row=r, column=8))
        style_data_cell(ws, r, 9, "")
        if i % 2 == 0:
            for c in range(1, 10):
                if c not in (5, 6):
                    ws.cell(row=r, column=c).fill = light_fill

    # Total row
    total_row = row + len(items) + 1
    ws.cell(row=total_row, column=4, value="TOTAL").font = bold_font_xl
    ws.cell(row=total_row, column=4).alignment = center_align
    ws.cell(row=total_row, column=5).value = f"=SUM(E5:E{total_row-1})"
    ws.cell(row=total_row, column=5).font = formula_font_xl
    ws.cell(row=total_row, column=5).number_format = currency_format
    ws.cell(row=total_row, column=5).alignment = center_align
    ws.cell(row=total_row, column=6).value = f"=SUM(F5:F{total_row-1})"
    ws.cell(row=total_row, column=6).font = formula_font_xl
    ws.cell(row=total_row, column=6).number_format = currency_format
    ws.cell(row=total_row, column=6).alignment = center_align
    ws.cell(row=total_row, column=7).value = f'=IF(E{total_row}=0,"",((F{total_row}-E{total_row})/E{total_row}))'
    ws.cell(row=total_row, column=7).font = formula_font_xl
    ws.cell(row=total_row, column=7).number_format = pct_format
    ws.cell(row=total_row, column=7).alignment = center_align

    # Column widths
    widths = [5, 18, 55, 18, 16, 14, 14, 14, 30]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    path = os.path.join(OUTPUT_DIR, "checklist-equipamiento-obra.xlsx")
    wb.save(path)
    print(f"✅ Checklist equipamiento generado: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# FILE 4: Calculadora de viabilidad (3 tabs)
# ═══════════════════════════════════════════════════════════
def generate_calculadora():
    wb = Workbook()

    # ── TAB 1: Inversión Inicial ──────────────────────────
    ws1 = wb.active
    ws1.title = "Inversión Inicial"

    add_title_block(ws1, "Inversión Inicial — Dark Kitchen", BRAND_LINE, merge_to=5)

    headers1 = ["Categoría", "Concepto", "Presupuesto (€)", "Real (€)", "Notas"]
    for c, h in enumerate(headers1, 1):
        ws1.cell(row=4, column=c, value=h)
    style_header_row(ws1, 4, len(headers1))

    inversion_items = [
        ("Obra Civil", "Acondicionamiento del local (suelos, paredes, techos)", 8000),
        ("Obra Civil", "Instalación eléctrica", 3500),
        ("Obra Civil", "Fontanería y saneamiento", 1500),
        ("Obra Civil", "Instalación de gas", 1200),
        ("Obra Civil", "Extracción de humos y ventilación", 3500),
        ("Obra Civil", "Sistema contra incendios", 600),
        ("Equipamiento", "Horno de convección profesional", 3500),
        ("Equipamiento", "Plancha industrial", 1200),
        ("Equipamiento", "Freidora industrial", 1200),
        ("Equipamiento", "Cocina de inducción", 1000),
        ("Equipamiento", "Cámara frigorífica", 2800),
        ("Equipamiento", "Congelador industrial", 1500),
        ("Equipamiento", "Mesa refrigerada", 2200),
        ("Equipamiento", "Lavavajillas industrial", 2000),
        ("Equipamiento", "Mesas de trabajo, estanterías y menaje", 2000),
        ("Tecnología", "Tablets para plataformas (x3)", 600),
        ("Tecnología", "KDS + integrador de pedidos", 1200),
        ("Tecnología", "TPV y software de gestión", 800),
        ("Tecnología", "WiFi + cámaras de seguridad", 500),
        ("Licencias", "Proyecto técnico (arquitecto/ingeniero)", 2500),
        ("Licencias", "Tasas municipales de licencia", 500),
        ("Licencias", "Plan APPCC + consultoría sanitaria", 800),
        ("Licencias", "Constitución de empresa (notario + registro)", 700),
        ("Seguros", "Seguro de responsabilidad civil (1er año)", 400),
        ("Seguros", "Seguro de local y contenido (1er año)", 300),
        ("Marketing", "Sesión de fotografía profesional", 400),
        ("Marketing", "Diseño de marca y packaging", 1500),
        ("Marketing", "Campañas de lanzamiento (1er mes)", 2000),
        ("Capital Circulante", "Stock inicial de materias primas", 2000),
        ("Capital Circulante", "Fondo de maniobra (2 meses de gastos fijos)", 8000),
    ]

    for i, (cat, concepto, presup) in enumerate(inversion_items, 1):
        r = 4 + i
        style_data_cell(ws1, r, 1, cat)
        style_data_cell(ws1, r, 2, concepto)
        style_data_cell(ws1, r, 3, presup, fmt=currency_format, align=center_align, fill=input_fill)
        style_data_cell(ws1, r, 4, None, fmt=currency_format, align=center_align, fill=input_fill)
        style_data_cell(ws1, r, 5, "")
        if i % 2 == 0:
            for c in [1, 2, 5]:
                ws1.cell(row=r, column=c).fill = light_fill

    total_r = 4 + len(inversion_items) + 1
    ws1.cell(row=total_r, column=2, value="INVERSIÓN TOTAL").font = Font(name="Calibri", size=12, bold=True, color=GOLD)
    ws1.cell(row=total_r, column=3).value = f"=SUM(C5:C{total_r-1})"
    ws1.cell(row=total_r, column=3).font = Font(name="Calibri", size=12, bold=True, color="1565C0")
    ws1.cell(row=total_r, column=3).number_format = currency_format
    ws1.cell(row=total_r, column=3).alignment = center_align
    ws1.cell(row=total_r, column=4).value = f"=SUM(D5:D{total_r-1})"
    ws1.cell(row=total_r, column=4).font = Font(name="Calibri", size=12, bold=True, color="1565C0")
    ws1.cell(row=total_r, column=4).number_format = currency_format
    ws1.cell(row=total_r, column=4).alignment = center_align

    widths1 = [22, 50, 18, 16, 30]
    for i, w in enumerate(widths1, 1):
        ws1.column_dimensions[get_column_letter(i)].width = w

    # ── TAB 2: P&L Mensual ────────────────────────────────
    ws2 = wb.create_sheet("P&L Mensual")

    add_title_block(ws2, "Cuenta de Resultados Mensual — Dark Kitchen", BRAND_LINE, merge_to=4)

    # Input section
    r = 4
    ws2.cell(row=r, column=1, value="PARÁMETROS DE ENTRADA").font = section_font_xl
    r += 1

    inputs = [
        ("Pedidos por día", 35, ""),
        ("Ticket medio (€)", 22, currency_format),
        ("Días operativos al mes", 30, ""),
        ("Comisión media plataformas (%)", 0.30, pct_format),
        ("Food cost (%)", 0.30, pct_format),
        ("Coste packaging por pedido (€)", 0.80, currency_format),
    ]

    for i, (label, val, fmt) in enumerate(inputs):
        row_i = r + i
        style_data_cell(ws2, row_i, 1, label, font=bold_font_xl)
        c = style_data_cell(ws2, row_i, 2, val, fill=input_fill, align=center_align)
        if fmt:
            c.number_format = fmt

    # Calculated P&L
    r_pl = r + len(inputs) + 2
    ws2.cell(row=r_pl, column=1, value="CUENTA DE RESULTADOS").font = section_font_xl

    # References to input cells
    pedidos_cell = f"B{r}"
    ticket_cell = f"B{r+1}"
    dias_cell = f"B{r+2}"
    comision_cell = f"B{r+3}"
    foodcost_cell = f"B{r+4}"
    packaging_cell = f"B{r+5}"

    pl_items = [
        ("Ingresos brutos", f"={pedidos_cell}*{ticket_cell}*{dias_cell}", True),
        ("(-) Comisiones de plataformas", f"=-{pedidos_cell}*{ticket_cell}*{dias_cell}*{comision_cell}", False),
        ("= Ingresos netos", f"={pedidos_cell}*{ticket_cell}*{dias_cell}*(1-{comision_cell})", True),
        ("", "", False),
        ("(-) Coste de materia prima (food cost)", f"=-{pedidos_cell}*{ticket_cell}*{dias_cell}*{foodcost_cell}", False),
        ("(-) Packaging", f"=-{pedidos_cell}*{dias_cell}*{packaging_cell}", False),
        ("= Margen bruto", f"={pedidos_cell}*{ticket_cell}*{dias_cell}*(1-{comision_cell}-{foodcost_cell})-{pedidos_cell}*{dias_cell}*{packaging_cell}", True),
    ]

    # Fixed costs
    fixed_costs = [
        ("Alquiler del local", 1500),
        ("Suministros (luz, agua, gas)", 600),
        ("Personal (1 jefe cocina + 1 ayudante)", 3800),
        ("Seguros", 60),
        ("Gestoría/asesoría", 200),
        ("Software y tecnología", 150),
        ("Limpieza y DDD", 120),
        ("Marketing en plataformas (5% ingresos)", None),
    ]

    r_current = r_pl + 1
    for label, formula, is_bold in pl_items:
        if label:
            style_data_cell(ws2, r_current, 1, label, font=bold_font_xl if is_bold else data_font_xl)
            ws2.cell(row=r_current, column=2).value = formula
            ws2.cell(row=r_current, column=2).number_format = currency_format
            ws2.cell(row=r_current, column=2).font = formula_font_xl if is_bold else data_font_xl
            ws2.cell(row=r_current, column=2).alignment = center_align
            ws2.cell(row=r_current, column=2).border = thin_border
            ws2.cell(row=r_current, column=1).border = thin_border
        r_current += 1

    r_current += 1
    ws2.cell(row=r_current, column=1, value="COSTES FIJOS MENSUALES").font = section_font_xl
    r_current += 1
    fixed_start = r_current

    ingresos_brutos_row = r_pl + 1  # Row of "Ingresos brutos"

    for label, amount in fixed_costs:
        style_data_cell(ws2, r_current, 1, label)
        if amount is not None:
            style_data_cell(ws2, r_current, 2, amount, fmt=currency_format, align=center_align, fill=input_fill)
        else:
            # Marketing = 5% of gross revenue
            ws2.cell(row=r_current, column=2).value = f"=B{ingresos_brutos_row}*0.05"
            ws2.cell(row=r_current, column=2).number_format = currency_format
            ws2.cell(row=r_current, column=2).font = data_font_xl
            ws2.cell(row=r_current, column=2).alignment = center_align
            ws2.cell(row=r_current, column=2).border = thin_border
        r_current += 1

    fixed_end = r_current - 1
    r_current += 1
    style_data_cell(ws2, r_current, 1, "TOTAL COSTES FIJOS", font=bold_font_xl)
    ws2.cell(row=r_current, column=2).value = f"=SUM(B{fixed_start}:B{fixed_end})"
    ws2.cell(row=r_current, column=2).font = formula_font_xl
    ws2.cell(row=r_current, column=2).number_format = currency_format
    ws2.cell(row=r_current, column=2).alignment = center_align
    total_fixed_row = r_current

    # Find margen bruto row
    margen_bruto_row = r_pl + 7  # 7th item in pl_items

    r_current += 2
    style_data_cell(ws2, r_current, 1, "EBITDA MENSUAL", font=Font(name="Calibri", size=14, bold=True, color=GOLD))
    ws2.cell(row=r_current, column=2).value = f"=B{margen_bruto_row}-B{total_fixed_row}"
    ws2.cell(row=r_current, column=2).font = Font(name="Calibri", size=14, bold=True, color="1565C0")
    ws2.cell(row=r_current, column=2).number_format = currency_format
    ws2.cell(row=r_current, column=2).alignment = center_align
    ebitda_row = r_current

    r_current += 1
    style_data_cell(ws2, r_current, 1, "Margen EBITDA (%)", font=bold_font_xl)
    ws2.cell(row=r_current, column=2).value = f"=IF(B{ingresos_brutos_row}=0,0,B{ebitda_row}/B{ingresos_brutos_row})"
    ws2.cell(row=r_current, column=2).font = formula_font_xl
    ws2.cell(row=r_current, column=2).number_format = pct_format
    ws2.cell(row=r_current, column=2).alignment = center_align

    ws2.column_dimensions['A'].width = 45
    ws2.column_dimensions['B'].width = 20
    ws2.column_dimensions['C'].width = 15
    ws2.column_dimensions['D'].width = 15

    # ── TAB 3: Punto de Equilibrio ────────────────────────
    ws3 = wb.create_sheet("Punto de Equilibrio")

    add_title_block(ws3, "Análisis del Punto de Equilibrio — Dark Kitchen", BRAND_LINE, merge_to=5)

    r = 4
    ws3.cell(row=r, column=1, value="CÁLCULO DEL BREAK-EVEN").font = section_font_xl
    r += 1

    # Reference P&L sheet
    be_items = [
        ("Ticket medio (€)", f"='P&L Mensual'!B{5+1}", currency_format),
        ("Comisión media plataformas (%)", f"='P&L Mensual'!B{5+3}", pct_format),
        ("Food cost (%)", f"='P&L Mensual'!B{5+4}", pct_format),
        ("Packaging por pedido (€)", f"='P&L Mensual'!B{5+5}", currency_format),
        ("Total costes fijos mensuales (€)", f"='P&L Mensual'!B{total_fixed_row}", currency_format),
    ]

    for i, (label, formula, fmt) in enumerate(be_items):
        ri = r + i
        style_data_cell(ws3, ri, 1, label, font=bold_font_xl)
        ws3.cell(row=ri, column=2).value = formula
        ws3.cell(row=ri, column=2).number_format = fmt
        ws3.cell(row=ri, column=2).font = data_font_xl
        ws3.cell(row=ri, column=2).alignment = center_align
        ws3.cell(row=ri, column=2).border = thin_border

    r_calc = r + len(be_items) + 1
    ticket_ref = f"B{r}"
    comision_ref = f"B{r+1}"
    fc_ref = f"B{r+2}"
    pkg_ref = f"B{r+3}"
    fixed_ref = f"B{r+4}"

    ws3.cell(row=r_calc, column=1, value="RESULTADOS").font = section_font_xl
    r_calc += 1

    # Contribution margin per order
    style_data_cell(ws3, r_calc, 1, "Margen de contribución por pedido (€)", font=bold_font_xl)
    ws3.cell(row=r_calc, column=2).value = f"={ticket_ref}*(1-{comision_ref}-{fc_ref})-{pkg_ref}"
    ws3.cell(row=r_calc, column=2).number_format = currency_format
    ws3.cell(row=r_calc, column=2).font = formula_font_xl
    ws3.cell(row=r_calc, column=2).alignment = center_align
    margin_row = r_calc

    r_calc += 1
    style_data_cell(ws3, r_calc, 1, "Pedidos necesarios al mes (break-even)", font=bold_font_xl)
    ws3.cell(row=r_calc, column=2).value = f"=ROUNDUP({fixed_ref}/B{margin_row},0)"
    ws3.cell(row=r_calc, column=2).font = Font(name="Calibri", size=14, bold=True, color="D32F2F")
    ws3.cell(row=r_calc, column=2).number_format = '#,##0'
    ws3.cell(row=r_calc, column=2).alignment = center_align

    r_calc += 1
    style_data_cell(ws3, r_calc, 1, "Pedidos necesarios al día (30 días)", font=bold_font_xl)
    ws3.cell(row=r_calc, column=2).value = f"=ROUNDUP(B{r_calc-1}/30,0)"
    ws3.cell(row=r_calc, column=2).font = Font(name="Calibri", size=14, bold=True, color="D32F2F")
    ws3.cell(row=r_calc, column=2).number_format = '#,##0'
    ws3.cell(row=r_calc, column=2).alignment = center_align

    r_calc += 1
    style_data_cell(ws3, r_calc, 1, "Facturación mensual mínima (€)", font=bold_font_xl)
    ws3.cell(row=r_calc, column=2).value = f"=B{r_calc-2}*{ticket_ref}"
    ws3.cell(row=r_calc, column=2).font = Font(name="Calibri", size=14, bold=True, color="D32F2F")
    ws3.cell(row=r_calc, column=2).number_format = currency_format
    ws3.cell(row=r_calc, column=2).alignment = center_align

    # 3 Scenarios
    r_sc = r_calc + 3
    ws3.cell(row=r_sc, column=1, value="ESCENARIOS DE FACTURACIÓN").font = section_font_xl
    r_sc += 1

    sc_headers = ["", "Pesimista", "Realista", "Optimista"]
    for c, h in enumerate(sc_headers, 1):
        ws3.cell(row=r_sc, column=c, value=h)
    style_header_row(ws3, r_sc, 4)
    ws3.cell(row=r_sc, column=1).fill = PatternFill()  # clear first cell

    scenarios = [
        ("Pedidos/día", 20, 40, 65),
        ("Ticket medio (€)", f"={ticket_ref}", f"={ticket_ref}", f"={ticket_ref}"),
    ]

    r_sc += 1
    for label, pes, rea, opt in scenarios:
        style_data_cell(ws3, r_sc, 1, label, font=bold_font_xl)
        for col, val in [(2, pes), (3, rea), (4, opt)]:
            c = style_data_cell(ws3, r_sc, col, val, align=center_align, fill=input_fill)
            if isinstance(val, str):
                c.number_format = currency_format
        r_sc += 1

    pedidos_row = r_sc - 2
    ticket_row = r_sc - 1

    calc_rows = [
        ("Facturación bruta mensual (€)", lambda col: f"={get_column_letter(col)}{pedidos_row}*{get_column_letter(col)}{ticket_row}*30", currency_format),
        ("(-) Comisiones plataformas (€)", lambda col: f"=-{get_column_letter(col)}{r_sc-1}*{comision_ref}", currency_format),
        ("(-) Food cost (€)", lambda col: f"=-{get_column_letter(col)}{r_sc-1}*{fc_ref}", currency_format),
        ("(-) Packaging (€)", lambda col: f"=-{get_column_letter(col)}{pedidos_row}*30*{pkg_ref}", currency_format),
        ("= Margen bruto (€)", lambda col: f"=SUM({get_column_letter(col)}{r_sc}:{get_column_letter(col)}{r_sc+3})", currency_format),
        ("(-) Costes fijos (€)", lambda col: f"=-{fixed_ref}", currency_format),
        ("= EBITDA mensual (€)", lambda col: f"={get_column_letter(col)}{r_sc+4}+{get_column_letter(col)}{r_sc+5}", currency_format),
        ("Margen EBITDA (%)", lambda col: f"=IF({get_column_letter(col)}{r_sc}=0,0,{get_column_letter(col)}{r_sc+6}/{get_column_letter(col)}{r_sc})", pct_format),
    ]

    for label, formula_fn, fmt in calc_rows:
        style_data_cell(ws3, r_sc, 1, label, font=bold_font_xl if "=" in label or "EBITDA" in label else data_font_xl)
        for col in [2, 3, 4]:
            ws3.cell(row=r_sc, column=col).value = formula_fn(col)
            ws3.cell(row=r_sc, column=col).number_format = fmt
            ws3.cell(row=r_sc, column=col).font = formula_font_xl
            ws3.cell(row=r_sc, column=col).alignment = center_align
            ws3.cell(row=r_sc, column=col).border = thin_border
        r_sc += 1

    ws3.column_dimensions['A'].width = 42
    ws3.column_dimensions['B'].width = 20
    ws3.column_dimensions['C'].width = 20
    ws3.column_dimensions['D'].width = 20
    ws3.column_dimensions['E'].width = 15

    path = os.path.join(OUTPUT_DIR, "calculadora-viabilidad-dark-kitchen.xlsx")
    wb.save(path)
    print(f"✅ Calculadora de viabilidad generada: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("🚀 Generando archivos de la Guía Dark Kitchen...\n")
    f1 = generate_guide_docx()
    f2 = generate_checklist_legal()
    f3 = generate_checklist_equipamiento()
    f4 = generate_calculadora()

    print("\n📊 Verificación de archivos:")
    for f in [f1, f2, f3, f4]:
        size = os.path.getsize(f)
        if size > 1024 * 1024:
            print(f"   {os.path.basename(f)}: {size / (1024*1024):.1f} MB")
        else:
            print(f"   {os.path.basename(f)}: {size / 1024:.1f} KB")

    print("\n✅ ¡Todos los archivos generados correctamente!")
