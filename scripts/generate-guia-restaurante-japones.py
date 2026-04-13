#!/usr/bin/env python3
"""
Generate "Cómo Montar un Restaurante Japonés 60 Plazas" guide deliverables.
AI Chef Pro — aichef.pro
19 files: 1 DOCX guide + 8 Excel templates + 6 Excel checklists + 2 DOCX models + 1 PDF placeholder
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "public", "dl", "guia-restaurante-japones"
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

GOLD = "FFD700"
GOLD_RGB = RGBColor(0xFF, 0xD7, 0x00)
HEADER_BG = "2D2D2D"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"
MEDIUM_GRAY = "E0E0E0"
INPUT_COLOR = "E8F5E9"
BRAND = "AI Chef Pro · aichef.pro — Restaurante Japonés"

title_font = Font(name="Calibri", size=16, bold=True, color=GOLD)
sub_font = Font(name="Calibri", size=11, color="888888", italic=True)
hdr_font = Font(name="Calibri", size=11, bold=True, color=WHITE)
sec_font = Font(name="Calibri", size=12, bold=True, color=GOLD)
dat_font = Font(name="Calibri", size=11)
bld_font = Font(name="Calibri", size=11, bold=True)
frm_font = Font(name="Calibri", size=11, color="1565C0", bold=True)
brn_font = Font(name="Calibri", size=9, color="888888", italic=True)

hdr_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
inp_fill = PatternFill(start_color=INPUT_COLOR, end_color=INPUT_COLOR, fill_type="solid")
lgt_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")
sec_fill = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")

thin_b = Border(
    left=Side(style="thin", color=MEDIUM_GRAY), right=Side(style="thin", color=MEDIUM_GRAY),
    top=Side(style="thin", color=MEDIUM_GRAY), bottom=Side(style="thin", color=MEDIUM_GRAY),
)
c_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
l_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
cur_fmt = '#,##0.00 €'
pct_fmt = '0.0%'


def shr(ws, row, ncols):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin_b; cell.alignment = c_align

def sdc(ws, r, c, v=None, font=None, fill=None, fmt=None, align=None):
    cell = ws.cell(row=r, column=c)
    if v is not None: cell.value = v
    cell.font = font or dat_font; cell.border = thin_b; cell.alignment = align or l_align
    if fill: cell.fill = fill
    if fmt: cell.number_format = fmt
    return cell

def title_block(ws, t, ncols=8):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    ws.cell(row=1, column=1, value=t).font = title_font
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
    ws.cell(row=2, column=1, value=BRAND).font = sub_font
    ws.row_dimensions[1].height = 30

def brand_footer(ws, row, ncols):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
    ws.cell(row=row, column=1, value=BRAND).font = brn_font
    ws.cell(row=row, column=1).alignment = c_align

def instr_sheet(wb, t, items):
    ws = wb.active; ws.title = "Instrucciones"; ws.sheet_properties.tabColor = "FFD700"
    title_block(ws, t, 3)
    r = 4
    ws.cell(row=r, column=1, value="Instrucciones de uso").font = sec_font; r += 1
    for i, x in enumerate(items, 1):
        ws.cell(row=r, column=1, value=f"{i}. {x}").font = dat_font; r += 1
    r += 1
    ws.cell(row=r, column=1, value="Celdas verdes = campos editables").font = dat_font
    ws.cell(row=r, column=1).fill = inp_fill
    ws.column_dimensions['A'].width = 80

def checklist_ws(ws, t, items):
    hdrs = ["#", "Categoría", "Tarea / Ítem", "Responsable", "Fecha Límite", "Estado", "Coste Est. (€)", "Notas"]
    title_block(ws, t)
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8); ws.freeze_panes = f"A{r+1}"
    ws.auto_filter.ref = f"A{r}:H{r}"
    dv = DataValidation(type="list", formula1='"Pendiente,En Curso,Completado"', allow_blank=True)
    ws.add_data_validation(dv)
    widths = [6, 22, 45, 18, 14, 14, 14, 30]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    for idx, (cat, tarea, resp, coste) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, tarea)
        sdc(ws, r, 4, resp); sdc(ws, r, 5, None, fill=inp_fill)
        sdc(ws, r, 6, "Pendiente", fill=inp_fill); dv.add(ws.cell(row=r, column=6))
        sdc(ws, r, 7, coste, fmt=cur_fmt, fill=inp_fill); sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1; brand_footer(ws, r, 8)


# ═══════════════════════════════════════════════════════════
# 1. GUIDE DOCX
# ═══════════════════════════════════════════════════════════
def gen_guide_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    def add_bullet(txt):
        doc.add_paragraph(txt, style='List Bullet')
    def add_num(txt):
        doc.add_paragraph(txt, style='List Number')
    def tip(txt):
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Consejo del Chef: {txt}")
        run.font.italic = True; run.font.color.rgb = GOLD_RGB

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cómo Montar un Restaurante Japonés"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("60 Plazas de Aforo — Guía España 2026"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Chef John Guerrero\nAI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # CH1
    doc.add_heading("1. Qué es un Restaurante Japonés Auténtico", level=1)
    doc.add_paragraph(
        "La cocina japonesa es Patrimonio Cultural Inmaterial de la UNESCO desde 2013 y una de las más "
        "respetadas del mundo. Desde los sushi-ya tradicionales de Tokio hasta las izakayas de barrio, "
        "los templos del ramen, los yakitori-ya y las experiencias omakase de alta gama, la gastronomía "
        "japonesa abarca un universo enorme de conceptos, técnicas e ingredientes. Lejos del 'japonés "
        "buffet' genérico, un restaurante japonés auténtico exige materia prima de máxima calidad, "
        "dominio técnico del itamae (chef de sushi) y respeto absoluto por el producto."
    )
    doc.add_paragraph(
        "En España, el mercado de restauración japonesa ha evolucionado enormemente en la última década. "
        "Ha pasado del buffet libre de sushi a propuestas de autor con itamae formados en Japón, "
        "omakase de 120-180€ por persona, izakayas especializadas en robatayaki y templos del ramen "
        "con 90 minutos de espera. Un restaurante japonés auténtico con un itamae cualificado, "
        "pescado sashimi-grade y una barra de sake y whisky japonés tiene un posicionamiento "
        "diferencial potentísimo y márgenes de coctelería y sake del 75-85%."
    )
    doc.add_heading("La diversidad japonesa: mucho más que sushi", level=2)
    add_bullet("Sushi-ya (sushi y sashimi): nigiri, maki, temaki, sashimi. Producto fresco de máxima calidad como eje.")
    add_bullet("Ramen-ya: ramen tonkotsu, shoyu, miso, shio. Caldos de 12-24h, elaboración artesanal obligatoria.")
    add_bullet("Izakaya: cocina japonesa informal para compartir con cerveza/sake. Yakitori, karaage, tsukemono, agemono.")
    add_bullet("Omakase (alta gama): menú degustación de 12-20 pases a cargo del itamae. Ticket 80-180€.")
    add_bullet("Robatayaki: parrilla de carbón binchotan a la vista, ingredientes premium (wagyu, gambas, verduras).")
    add_bullet("Tempura-ya: especialistas en tempura, aceite limpio, masa al momento, técnicas de fritura precisas.")
    add_bullet("Yakiniku: parrilla japonesa en mesa (carne wagyu, kalbi). Muy rentable, alto ticket.")
    add_bullet("Ticket medio auténtico: 35-55€ (izakaya/casual) o 80-150€ (omakase premium).")
    tip("La autenticidad y el itamae son tu mayor ventaja competitiva. Un itamae formado en Japón marca la diferencia entre un 'japonés de supermercado' y un restaurante con lista de espera. Si consigues un itamae cualificado y pescado sashimi-grade diario de Mercamadrid o Mercabarna, ya estás por encima del 80% de los japoneses en España.")
    doc.add_page_break()

    # CH2
    doc.add_heading("2. El Mercado de la Cocina Japonesa en España 2026", level=1)
    doc.add_paragraph(
        "La cocina japonesa es la segunda cocina asiática más demandada en España después de la china, "
        "pero con mucho mayor ticket medio. El consumo de sushi en España se ha multiplicado por 5 en "
        "la última década, y la demanda de experiencias omakase, ramen auténtico e izakayas de autor "
        "crece a doble dígito anual."
    )
    doc.add_heading("Datos clave del sector", level=2)
    add_bullet("Restaurantes japoneses en España: +2.400 (2025), crecimiento del 35% en 5 años.")
    add_bullet("Ticket medio: 35-55€ (izakaya/casual) vs 80-150€ (omakase/premium) vs 18-28€ (ramen-ya).")
    add_bullet("Ciudades con mayor demanda: Madrid, Barcelona, Valencia, Málaga, San Sebastián, Bilbao, Palma.")
    add_bullet("Público: 28-55 años, urbano, foodie, alto poder adquisitivo, viajero (muchos han visitado Japón).")
    add_bullet("Tendencia clave: omakase y sake premium en auge. Los omakase de 120-180€ llenan reserva con 3 semanas de antelación.")
    add_bullet("Ramen explosivo: los ramen-ya auténticos en Madrid y Barcelona tienen colas diarias de 60-90 minutos.")
    add_bullet("Delivery japonés: el sushi es el rey del delivery asiático. Ramen y platos calientes requieren packaging especial.")
    tip("El 70% de los restaurantes japoneses en España son 'japoneses chinos' — propiedad de empresarios chinos, con personal chino y sushi prefabricado. Si apuestas por un equipo japonés o formado en Japón y materia prima real (pescado sashimi-grade, arroz Koshihikari, alga nori premium), te sitúas en un segmento premium con demanda insaciable.")
    doc.add_page_break()

    # CH3
    doc.add_heading("3. Modelos de Negocio: Sushi-ya, Ramen, Izakaya, Omakase", level=1)
    doc.add_paragraph("El restaurante japonés admite múltiples modelos muy diferenciados. La elección del modelo determina la inversión, la operativa diaria, el perfil del equipo y tu posicionamiento en el mercado.")
    doc.add_heading("Modelo 1: Sushi-ya (Sushi y Sashimi)", level=2)
    doc.add_paragraph("Especializado en sushi, sashimi, nigiri y maki. Barra de sushi visible donde el itamae trabaja delante del cliente. Ticket 45-80€. Requiere itamae cualificado, pescado sashimi-grade diario y arroz Koshihikari. Es el modelo más prestigioso y técnicamente exigente.")
    doc.add_heading("Modelo 2: Ramen-ya", level=2)
    doc.add_paragraph("Especializado en ramen. Caldos de 12-24h (tonkotsu, shoyu, miso, shio), fideos artesanales o de proveedor japonés, toppings variados (chashu, ajitama, menma, negi). Ticket 18-28€. Alta rotación, margen brillante, menor inversión en producto fresco que un sushi-ya. Ideal para concepto rápido-premium.")
    doc.add_heading("Modelo 3: Izakaya (Gastrobar Japonés)", level=2)
    doc.add_paragraph("Cocina japonesa informal para compartir con sake y cerveza. Yakitori (brochetas de pollo), karaage, gyoza, tempura, agemono, tsukemono. Ambiente casual, alto consumo de bebida. Ticket 35-50€. Modelo muy rentable con equipo menos costoso que un sushi-ya puro.")
    doc.add_heading("Modelo 4: Omakase / Alta Gama", level=2)
    doc.add_paragraph("Barra de 8-12 asientos, menú degustación de 12-20 pases a cargo del itamae. Ticket 80-180€. Requiere itamae de máximo nivel (idealmente formado en Japón), pescado premium (toro, uni, anago). Modelo más exclusivo, márgenes altos pero muy dependiente del talento del itamae.")
    doc.add_heading("Modelo 5: Robatayaki", level=2)
    doc.add_paragraph("Parrilla robata de carbón binchotan a la vista, platos servidos con pala larga de madera. Ingredientes premium: wagyu, gambas, verduras de temporada, setas, pescado entero. Ticket 60-100€. Inversión alta en extracción y robata, pero experiencial y muy instagrameable.")
    doc.add_heading("Modelo 6: Mixto (Sushi + Ramen + Izakaya)", level=2)
    doc.add_paragraph("Restaurante japonés completo con barra de sushi, estación de ramen y cocina izakaya. Carta amplia 40-60 platos. Ticket 35-55€. Modelo más replicable y estable en España, aunque requiere cocina más compleja y equipo más grande. El más común en apertura.")
    doc.add_heading("Modelo 7: Dark Kitchen Japonesa", level=2)
    doc.add_paragraph("Solo delivery: menú optimizado para sushi, bowls de poke/chirashi, ramen que viaje (packaging separado), katsu, gyoza. Sin sala. Inversión 50-90K€. Ideal para testear mercado antes de abrir local físico.")
    doc.add_page_break()

    # CH4
    doc.add_heading("4. Estudio de Viabilidad y Plan Financiero", level=1)
    doc.add_paragraph(
        "La inversión total para un restaurante japonés de 60 plazas en España oscila "
        "entre 250.000€ y 500.000€ dependiendo de la ubicación, nivel de acabados, "
        "si incluyes barra de sushi tradicional y el grado de equipamiento especializado "
        "japonés (suihanki, vitrinas refrigeradas de sashimi, robata, teppanyaki, cuchillos yanagiba)."
    )
    doc.add_heading("Desglose de inversión típica", level=2)
    add_bullet("Obra civil y acondicionamiento: 70.000-150.000€")
    add_bullet("Equipamiento cocina japonesa (suihanki, vitrina sashimi, ramen cooker, robata, teppanyaki, cuchillos): 55.000-110.000€")
    add_bullet("Barra de sushi artesanal + vitrina refrigerada sashimi: 15.000-35.000€")
    add_bullet("Barra de sake y whisky japonés (instalación + stock inicial 30-50 refs): 8.000-18.000€")
    add_bullet("Mobiliario sala + decoración japonesa (madera, piedra, lámparas washi, bambú, minimalismo): 25.000-55.000€")
    add_bullet("Licencias y permisos (incluye extracción reforzada robata, registro sanitario sashimi): 6.000-14.000€")
    add_bullet("Marketing pre-apertura: 5.000-12.000€")
    add_bullet("Stock inicial (pescado sashimi-grade + arroz Koshihikari + ajíes + sake + algas): 12.000-25.000€")
    add_bullet("Fondo de maniobra (3 meses): 45.000-95.000€")
    add_bullet("Tecnología (TPV, web, delivery, pantallas cocina): 4.000-10.000€")
    doc.add_heading("Ratios financieros objetivo", level=2)
    add_bullet("Food cost: 30-35% sobre ventas (pescado fresco premium e importación sube ligeramente el coste)")
    add_bullet("Coste de personal: 30-38% sobre ventas (el itamae es caro)")
    add_bullet("Alquiler: máximo 8-10% sobre ventas")
    add_bullet("EBITDA objetivo: 10-18%")
    add_bullet("Break-even: mes 10-16 (depende de ubicación, ticket y tipo de modelo)")
    tip("El margen en sake, whisky japonés y coctelería compensa el food cost del pescado fresco. Un sake premium cuesta 6-8€ por copa y lo vendes a 15-22€. Un highball de whisky japonés cuesta 1.50€ y se vende a 9-12€. La barra de sake y whisky puede representar el 20-30% de tu facturación con márgenes del 75-85%.")
    doc.add_page_break()

    # CH5
    doc.add_heading("5. Requisitos Legales en España + Normativa Anisakis", level=1)
    doc.add_paragraph("Abrir un restaurante japonés en España requiere los mismos trámites que cualquier restaurante, MÁS una atención especial a la normativa de congelación preventiva para pescado consumido crudo (anisakis), trazabilidad de pescado y, si importas productos japoneses directamente, trámites de comercio exterior.")
    doc.add_heading("Trámites generales", level=2)
    add_num("Constitución de sociedad (SL) o alta como autónomo")
    add_num("Licencia de actividad / declaración responsable (según municipio)")
    add_num("Registro sanitario autonómico")
    add_num("Alta en Hacienda: IAE epígrafe 671.4 o similar")
    add_num("Inscripción Seguridad Social empresa")
    add_num("Seguro de responsabilidad civil (mínimo 600.000€ por riesgo sanitario del pescado crudo)")
    add_num("Plan APPCC con PCC específico para pescado crudo (sashimi/sushi)")
    add_num("Registro RGSEAA")
    add_num("Licencia de terraza (si aplica)")
    add_num("Licencia de música / SGAE")
    add_num("Libro de reclamaciones")
    doc.add_heading("Normativa ANISAKIS (OBLIGATORIA — crítica para sushi y sashimi)", level=2)
    add_bullet("El Real Decreto 1420/2006 OBLIGA a congelar preventivamente todo pescado destinado a ser consumido crudo, marinado o poco hecho a -20°C durante al menos 24 horas (o -35°C durante 15 h).")
    add_bullet("Aplica a: sashimi, sushi, nigiri, tataki, marinados en cítricos o vinagre.")
    add_bullet("El congelador debe ser específico y alcanzar la temperatura certificada. Conservar registros de congelación por lote.")
    add_bullet("Excepciones: pescado de acuicultura alimentado en condiciones controladas SIN riesgo de anisakis puede venderse fresco (requiere certificado del proveedor).")
    add_bullet("Información al cliente: obligación de informar en carta que el pescado ha sido congelado preventivamente según normativa.")
    add_bullet("Sanciones: multas de 5.000-60.000€ por incumplimiento + cierre cautelar en caso de caso positivo de anisakiasis.")
    doc.add_heading("Requisitos específicos de importación de Japón", level=2)
    add_bullet("Registro como operador de comercio exterior (EORI) si importas directamente")
    add_bullet("Certificado SOIVRE para productos alimentarios importados de fuera de la UE")
    add_bullet("Etiquetado en español obligatorio para todos los productos importados (sake, salsa soja, mirin, panko)")
    add_bullet("Impuestos especiales para bebidas alcohólicas (sake, shochu, whisky japonés): IIEE")
    add_bullet("Alternativa más sencilla: comprar a distribuidores especializados en España (Japan Sushi Express, Oriental Gourmet, Japonshop, Comercial Minamoto)")
    tip("La normativa anisakis NO es negociable. Compra un congelador certificado que alcance -20°C real, documenta cada lote, e incluye en el plan APPCC un PCC (Punto de Control Crítico) específico para el pescado crudo. Guarda registros al menos 12 meses. Un caso positivo de anisakiasis puede cerrar tu restaurante y generar responsabilidad civil millonaria.")
    doc.add_page_break()

    # CH6
    doc.add_heading("6. APPCC y Seguridad Alimentaria para Cocina Japonesa", level=1)
    doc.add_paragraph("El plan APPCC de un restaurante japonés tiene las mismas bases que cualquier restaurante, pero con puntos de control críticos (PCC) adicionales para sashimi/sushi (pescado crudo), arroz sushi (temperatura y acidificación), y productos importados.")
    doc.add_heading("Prerrequisitos estándar", level=2)
    add_bullet("Plan de limpieza y desinfección (frecuencias, productos, responsables)")
    add_bullet("Control de temperaturas (cámaras, servicio, transporte)")
    add_bullet("Plan de control de plagas (DDD)")
    add_bullet("Trazabilidad de materias primas (proveedores, lotes, fechas, documentación de congelación del pescado)")
    add_bullet("Plan de formación del personal (manipulador de alimentos, específico pescado crudo)")
    add_bullet("Control de agua potable")
    add_bullet("Plan de gestión de residuos y aceite usado (tempura)")
    add_bullet("Plan de alérgenos (14 alérgenos + especial atención a soja, sésamo, crustáceos, gluten)")
    doc.add_heading("Puntos críticos específicos de cocina japonesa", level=2)
    add_bullet("Pescado para sashimi/sushi: CONGELACIÓN PREVIA obligatoria -20°C/24h para eliminar anisakis. NO es opcional. Documentación de cada lote.")
    add_bullet("Arroz de sushi (shari): tras cocer y aderezar con vinagre (sushi-zu), mantener en ohitsu o equipo certificado entre 20-25°C máximo 4 horas. Nunca refrigerar ni dejar a temperatura ambiente más tiempo.")
    add_bullet("Descongelación de pescado: SIEMPRE en cámara (0-4°C) durante la noche. Nunca a temperatura ambiente ni con agua.")
    add_bullet("Separación estricta de utensilios para pescado crudo: cuchillos yanagiba exclusivos, tablas dedicadas, superficies limpias.")
    add_bullet("Vitrinas refrigeradas de sashimi: mantenimiento constante 0-4°C, rotación máxima de pescado expuesto cada 2-3 horas.")
    add_bullet("Caldos de ramen (tonkotsu, shoyu): riesgo bacteriano en cocciones largas. Enfriamiento rápido (<4°C en <2h) y recalentamiento >75°C.")
    add_bullet("Huevos crudos (tamago en nigiri, sushi): usar huevos pasteurizados o cumplir estrictos protocolos de fecha.")
    add_bullet("Tempura (aceite): filtrado diario, cambio según índice de polaridad, temperatura controlada 170-180°C.")
    tip("El arroz de sushi (shari) es tu segundo gran riesgo sanitario después del pescado crudo. La temperatura debe mantenerse entre 20-25°C (cálido pero no caliente). Si lo refrigeras pierde textura y si lo dejas a temperatura ambiente más de 4 horas se convierte en caldo de cultivo. Usa un ohitsu profesional o equipo certificado y registra las temperaturas cada 2 horas.")
    doc.add_page_break()

    # CH7
    doc.add_heading("7. Ubicación y Local", level=1)
    doc.add_paragraph("La ubicación de un restaurante japonés es crucial. El alto ticket medio, la dependencia de pescado fresco diario y la necesidad de un público sofisticado acotan las zonas viables a barrios urbanos premium y zonas gastronómicas consolidadas.")
    doc.add_heading("Criterios de selección", level=2)
    add_bullet("Tráfico peatonal: mínimo 300 personas/hora en horario punta, especialmente público 28-55 años")
    add_bullet("Visibilidad: fachada minimalista con identidad japonesa sutil (noren, kanji, madera natural)")
    add_bullet("Metros cuadrados: 140-200 m² para 60 plazas (ratio 2.3-3.3 m²/plaza — los japoneses priorizan confort)")
    add_bullet("Zona: gastronómica premium + afterwork + público foodie. Muy mal en zonas populares de bajo ticket.")
    add_bullet("Competencia: analiza los 500m alrededor — ideal si hay otros restaurantes premium no japoneses")
    add_bullet("Salida de humos: imprescindible y reforzada. La robata de carbón binchotan, el teppanyaki y la tempura generan mucho humo.")
    add_bullet("Proximidad a mercados de abastos: ventaja enorme para pescado sashimi-grade diario (Mercamadrid, Mercabarna, lonjas)")
    add_bullet("Accesibilidad: acceso para personas con movilidad reducida (obligatorio por ley)")
    add_bullet("Barrios donde funciona muy bien: Chamberí/Chamartín (Madrid), Eixample/Sarrià (Barcelona), Ruzafa (Valencia), Indautxu (Bilbao), Ensanche (San Sebastián)")
    tip("El japonés premium depende menos del tráfico peatonal y más de la reputación y las reservas. Un omakase en una calle secundaria con 14 plazas y lista de espera factura más por m² que un japonés buffet en zona comercial. Prioriza zonas donde tu público objetivo (foodies 30-55 años con alto poder adquisitivo) vive y come.")
    doc.add_page_break()

    # CH8
    doc.add_heading("8. Diseño de Cocina Japonesa Profesional", level=1)
    doc.add_paragraph("La cocina de un restaurante japonés tiene estaciones muy especializadas que no existen en una cocina convencional europea: la barra de sushi del itamae, la estación del ramen, la parrilla robata, la zona de tempura y la zona de arroz sushi (shari).")
    doc.add_heading("Zonas de la cocina japonesa", level=2)
    add_bullet("Barra de sushi (sushi bar): zona de trabajo del itamae, visible al cliente. Mesa refrigerada para sashimi, tabla de corte grande, cuchillos yanagiba, deba, usuba. Temperatura ambiente controlada.")
    add_bullet("Estación de arroz sushi (shari): suihanki (arrocera profesional japonesa), ohitsu (recipiente de cedro), abanico uchiwa. Aderezo con sushi-zu (vinagre + azúcar + sal).")
    add_bullet("Estación de ramen (ramen station): ramen cooker o pasta cooker profesional (8-12 cestas), ollas de caldo de 50-100 litros, zona de montaje de bowls con toppings.")
    add_bullet("Parrilla robata / robatayaki: parrilla de carbón binchotan, zona de brochetas, pala larga de madera para servir. Requiere extracción extrema.")
    add_bullet("Plancha teppanyaki: plancha japonesa de gran superficie para carnes, verduras, yakisoba. Visible al cliente en algunos modelos.")
    add_bullet("Zona caliente: wok para saltados, freidora para tempura, fogones para guisos.")
    add_bullet("Zona fría / preparación: corte de verduras, preparación de salsas (ponzu, teriyaki, tare), mise en place.")
    add_bullet("Almacenamiento: cámaras frigoríficas (2) con zona específica para pescado sashimi-grade, cámara de congelación CERTIFICADA -20°C, economato seco para arroz, algas, salsas.")
    add_bullet("Zona de lavado: tren de lavado, zona de limpieza.")
    add_bullet("Superficie mínima cocina: 55-70 m² (ratio cocina/sala 1:2.5 — la cocina japonesa necesita más espacio)")
    tip("La barra de sushi a la vista del cliente es tu mayor herramienta de marketing y tu mejor vitrina. Ver al itamae cortando pescado con un yanagiba de 30 cm, montando nigiris a mano, sirviendo directamente al cliente — es teatro en vivo. Invierte en una barra de sushi profesional con vitrina refrigerada delantera y deja que tu itamae sea el protagonista. Las fotos en redes se disparan.")
    doc.add_page_break()

    # CH9
    doc.add_heading("9. Equipamiento Específico de Cocina Japonesa", level=1)
    doc.add_paragraph("Además del equipamiento estándar de cocina profesional, un restaurante japonés necesita equipos específicos, muchos de ellos importados directamente de Japón o distribuidos por especialistas.")
    doc.add_heading("Equipamiento específico japonés con costes orientativos", level=2)
    add_bullet("Suihanki (arrocera japonesa profesional, 5-10 litros): 1.800-4.500€")
    add_bullet("Vitrina refrigerada de sashimi (barra de sushi): 4.500-9.000€")
    add_bullet("Ramen cooker / pasta cooker profesional (8-12 cestas): 3.500-8.000€")
    add_bullet("Ollas de caldo 50-100L para tonkotsu/shoyu (2-3 uds): 1.500-3.500€")
    add_bullet("Parrilla robata de carbón binchotan: 4.000-9.000€")
    add_bullet("Plancha teppanyaki profesional: 3.000-7.000€")
    add_bullet("Cuchillos japoneses profesionales (yanagiba, deba, usuba): 1.500-4.500€ (set completo itamae)")
    add_bullet("Tabla de corte madera de Hiba / sintética para sushi: 300-800€")
    add_bullet("Ohitsu (recipiente de cedro para shari): 400-1.200€")
    add_bullet("Maquina de sake caliente (tokkuri warmer): 400-900€")
    add_bullet("Vitrina refrigerada de sake y whisky japonés: 3.000-6.000€")
    doc.add_heading("Equipamiento estándar de cocina", level=2)
    add_bullet("Cocina de gas 6 fuegos + horno: 3.500-6.000€")
    add_bullet("Wok profesional alta potencia (para yakisoba, saltados): 1.500-3.500€")
    add_bullet("Freidora doble (tempura, karaage, agemono): 1.500-3.000€")
    add_bullet("Horno mixto (Rational / Unox): 8.000-15.000€")
    add_bullet("Cámaras frigoríficas (2 puertas): 2.500-4.000€")
    add_bullet("Cámara congelación CERTIFICADA -20°C (imprescindible para anisakis): 3.000-5.500€")
    add_bullet("Tren de lavado: 4.000-8.000€")
    add_bullet("Campana extractora REFORZADA + filtros (robata genera humo extremo): 6.000-12.000€")
    add_bullet("Total equipamiento cocina: 55.000-110.000€")
    tip("La robata de carbón binchotan es el elemento más espectacular de un restaurante japonés premium — pero también el más exigente en extracción. El binchotan alcanza 1.000°C y genera humo intenso. Invierte en extracción sobredimensionada (mínimo 12.000 m³/h) y filtros de grado profesional. Una instalación de extracción deficiente te cerrará el restaurante por quejas vecinales en menos de un mes.")
    doc.add_page_break()

    # CH10
    doc.add_heading("10. Diseño de Sala y Decoración Japonesa", level=1)
    doc.add_paragraph("La decoración de un restaurante japonés debe reflejar los principios estéticos de 'wabi-sabi' (belleza en la imperfección), 'ma' (el vacío significativo) y 'shibui' (elegancia sobria). Lejos del cliché de farolillos rojos y dragones — eso es decoración china. El japonés auténtico es minimalismo, madera, piedra, bambú, luz tenue y silencio.")
    doc.add_heading("Distribución tipo (60 plazas)", level=2)
    add_bullet("Barra de sushi (sushi bar): 8-12 plazas frente al itamae — lo más prestigioso")
    add_bullet("Mesas de 2: 5-7 unidades (10-14 plazas) — parejas, cenas íntimas")
    add_bullet("Mesas de 4: 6-8 unidades (24-32 plazas) — grupos, familias")
    add_bullet("Zona privada / tatami (si aplica): 4-8 plazas — experiencia premium")
    add_bullet("Distancia entre mesas: 70-90 cm (los japoneses priorizan confort y silencio)")
    doc.add_heading("Elementos clave de la decoración japonesa auténtica", level=2)
    add_bullet("Madera natural: mesas, barra de sushi y detalles en madera de roble, cedro o haya. Clave en la estética japonesa.")
    add_bullet("Piedra y grava: detalles en piedra natural, pequeños jardines zen (karesansui) en zonas visibles.")
    add_bullet("Bambú: pantallas divisorias, lámparas, detalles decorativos. Evoca templos y casas de té.")
    add_bullet("Colores: paleta neutra de madera natural, blanco roto, negro mate, gris piedra. Acentos en rojo vermillón (torii) o azul índigo muy puntuales.")
    add_bullet("Iluminación: tenue, cálida (2400-2700K), lámparas washi (papel de arroz), puntos de luz dirigida a la barra y a las mesas.")
    add_bullet("Noren (cortina corta de tela sobre la puerta): signo universal de entrada a un local japonés. Imprescindible.")
    add_bullet("Vajilla: cerámica japonesa artesanal, asimetría, texturas rugosas (wabi-sabi). Platos de pizarra negra, piedra, madera para sushi.")
    add_bullet("Ausencia de decoración superflua: el 'ma' (espacio vacío) es parte de la decoración. Paredes limpias, menos objetos.")
    add_bullet("Música: jazz japonés, música ambient, nada estridente. Volumen bajo.")
    tip("Menos es más en decoración japonesa. Una pared de madera natural bien iluminada, un jardín zen en piedra, unas lámparas washi y una barra de sushi impecable crean un ambiente sofisticado y fotogénico. Evita las geishas en papel, los samuráis de plástico y los farolillos de papel — eso es un japonés de feria medieval, no un restaurante premium.")
    doc.add_page_break()

    # CH11
    doc.add_heading("11. Barra de Sake y Whisky Japonés (30-50 Referencias)", level=1)
    doc.add_paragraph("La barra de sake y whisky japonés es el elemento diferencial que eleva un restaurante japonés de casual a experiencial. El sake premium (junmai daiginjo), el whisky japonés (Yamazaki, Hakushu, Nikka) y el shochu son productos de muy alto margen y enorme potencial de upselling.")
    doc.add_heading("Estructura de la barra (30-50 referencias)", level=2)
    add_bullet("Sake Junmai (6-8 refs): sake puro de arroz, sin alcohol añadido. Estilos limpios y afrutados.")
    add_bullet("Sake Junmai Ginjo (4-6 refs): arroz pulido al 60%, aromático y elegante.")
    add_bullet("Sake Junmai Daiginjo (4-6 refs): arroz pulido al 50% o menos, máximo nivel. Premium.")
    add_bullet("Sake Nigori (2-3 refs): sake sin filtrar, lechoso, dulce. Ideal para postres.")
    add_bullet("Sake Sparkling (2-3 refs): sake espumoso, tipo champán. Muy popular.")
    add_bullet("Whisky japonés (6-10 refs): Yamazaki, Hakushu, Hibiki, Nikka, Chichibu. Cada vez más difíciles de encontrar — gestionar stock con antelación.")
    add_bullet("Shochu (3-5 refs): destilado japonés de arroz, cebada o batata. Alternativa al sake.")
    add_bullet("Cervezas japonesas (3-4 refs): Asahi, Sapporo, Kirin, Yebisu.")
    add_bullet("Cócteles japoneses (8-12): highball de whisky, sake sour, yuzu gin tonic, umeshu soda, matcha cocktail, sakura martini.")
    doc.add_heading("Inversión y márgenes", level=2)
    add_bullet("Stock inicial barra (30-50 refs): 6.000-15.000€")
    add_bullet("Vitrina refrigerada / estantería para botellas: 3.000-6.000€")
    add_bullet("Tokkuri (jarra) y ochoko (copa) para sake: 300-800€")
    add_bullet("Margen sake premium: 70-80% (compra 25€/botella → venta 12-16€/copa de 120ml)")
    add_bullet("Margen whisky japonés: 75-85% (un highball cuesta 1.50€ y se vende a 9-12€)")
    add_bullet("Margen cócteles: 80-85%")
    add_bullet("Objetivo: que la barra represente 20-30% de la facturación total")
    tip("Ofrece vuelos de sake (3 copas de diferentes tipos: junmai, ginjo, daiginjo) por 18-25€. Es la forma más rentable de vender sake premium: el cliente aprende, disfruta y acaba pidiendo una botella de daiginjo para la mesa. Margen brutal. Lo mismo con whisky japonés: vuelo de Yamazaki 12 + Hakushu 12 + Hibiki Harmony por 30-40€.")
    doc.add_page_break()

    # CH12
    doc.add_heading("12. Brigada de Cocina (8-12 personas)", level=1)
    doc.add_paragraph("La brigada de cocina de un restaurante japonés tiene roles específicos que no existen en una cocina europea convencional: el itamae (chef de sushi cabeza), los sushi chef, el ramen cook, el robata cook y el tempura cook son puestos muy especializados que requieren formación específica, a menudo en Japón.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Itamae / Jefe de cocina (idealmente formado en Japón, mínimo 5 años de experiencia en sushi): 1 persona — 3.000-4.500€ brutos/mes")
    add_bullet("Sushi chef / Segundo itamae (estación de sushi y sashimi): 1 persona — 2.200-2.800€ brutos/mes")
    add_bullet("Ramen cook (estación de ramen, caldos, montaje bowls): 1 persona — 1.900-2.400€ brutos/mes")
    add_bullet("Robata/Yakitori cook (parrilla binchotan): 1 persona — 1.900-2.400€ brutos/mes")
    add_bullet("Tempura/Hot kitchen cook (fritos, katsu, agemono, teppanyaki): 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Cocineros (guisos, saltados, preparación): 1-2 personas — 1.600-2.000€ brutos/mes cada uno")
    add_bullet("Preparación pescado / escamado / limpieza: 1 persona — 1.500-1.800€ brutos/mes")
    add_bullet("Ayudante / office: 1-2 personas — 1.300-1.600€ brutos/mes")
    add_bullet("Coste total cocina: 17.000-28.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Contratación del itamae", level=2)
    doc.add_paragraph("El itamae es el alma del restaurante japonés y la decisión más importante que tomarás. Necesita 5+ años de experiencia con sushi (idealmente formación en Japón), dominio técnico absoluto (corte yanagiba, montaje de nigiri, preparación de shari), conocimiento profundo del pescado, y sensibilidad estética. Un itamae cualificado en España cuesta 3.000-4.500€/mes y la lista de espera es larga. Alternativa: contratar un sushi chef español/europeo con formación específica (cursos Tokyo Sushi Academy, Hattori Nutrition College) — más barato pero con resultado variable.")
    tip("Busca itamae en España a través de redes de chefs japoneses (Asociación Japonesa de España, comunidades japonesas de Madrid y Barcelona, LinkedIn, Hostelwork). Si no encuentras, plantéate traer un itamae de Japón con visado de trabajo — el proceso es complejo pero hay empresas especializadas. Un itamae auténtico es tu mayor inversión y tu mayor retorno.")
    doc.add_page_break()

    # CH13
    doc.add_heading("13. Equipo de Sala (5-8 personas)", level=1)
    doc.add_paragraph("El equipo de sala de un restaurante japonés necesita conocer la cultura gastronómica japonesa para explicar platos, recomendar sakes, servir té japonés correctamente y crear la experiencia completa. En un omakase, el equipo de sala es casi tan importante como el itamae.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Encargado/a de sala (con conocimientos de sake y protocolo japonés): 1 persona — 2.200-2.800€ brutos/mes")
    add_bullet("Sommelier de sake y whisky japonés: 1 persona — 2.200-2.800€ brutos/mes (a veces el mismo encargado)")
    add_bullet("Camareros: 3-4 personas — 1.500-1.800€ brutos/mes cada uno")
    add_bullet("Runner / ayudante: 1-2 personas — 1.300-1.500€ brutos/mes")
    add_bullet("Coste total sala: 9.000-15.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Claves del servicio en un restaurante japonés", level=2)
    add_bullet("Conocimiento de sake: todo el equipo debe saber distinguir Junmai, Ginjo, Daiginjo, Nigori y recomendar maridajes.")
    add_bullet("Protocolo japonés: servir el sake sin llenar la copa del propio camarero, presentar botellas con ambas manos, saludo inicial 'irasshaimase'.")
    add_bullet("Explicar platos: muchos clientes no saben qué es un chawanmushi, un uni, un anago o un chirashi. El camarero debe saber describirlo con pasión y precisión.")
    add_bullet("Orden tradicional japonés: sashimi primero (sabor limpio), luego platos calientes, sushi al final para no alterar paladar.")
    add_bullet("Té japonés: ofrecer té genmaicha, sencha o hojicha gratuito al inicio. Es cortesía japonesa.")
    add_bullet("Upselling de bebidas: sake maridado, whisky japonés, sake nigori con postre. El sommelier de sake es clave para el ticket medio.")
    add_bullet("Respeto al itamae: en un sushi-ya, el camarero no interrumpe al itamae cuando trabaja. Los pedidos se pasan con protocolo.")
    tip("La formación del sommelier de sake es una inversión que se multiplica. Un sommelier que sabe contar la historia del sake, explica los grados de pulido del arroz y sugiere maridajes precisos vende el triple. Existen cursos oficiales de sake sommelier (WSET Sake Level 3, Sake School of America) que son altamente valorados.")
    doc.add_page_break()

    # CH14
    doc.add_heading("14. Menú Engineering: Sushi, Ramen, Izakaya, Robata", level=1)
    doc.add_paragraph("El diseño de la carta de un restaurante japonés debe equilibrar autenticidad, margen y accesibilidad. La carta debe educar al cliente no experto sin aburrir al experto, y ofrecer platos de todos los rangos de precio.")
    doc.add_heading("Estructura de carta japonesa", level=2)
    add_bullet("Sashimi: 5-8 opciones (salmón, atún, toro, lubina, hamachi, vieira, pulpo, uni)")
    add_bullet("Nigiri (individuales o pares): 10-15 opciones")
    add_bullet("Maki y uramaki: 8-12 opciones (futomaki, spicy tuna, california, dragon roll)")
    add_bullet("Izakaya / entrantes calientes: 8-12 opciones (gyoza, karaage, edamame, tempura, takoyaki)")
    add_bullet("Ramen: 3-5 opciones (tonkotsu, shoyu, miso, shio, tsukemen)")
    add_bullet("Robata / Yakitori: 6-10 opciones (wagyu, pollo, anguila, gambas, verduras)")
    add_bullet("Principales calientes: 4-6 opciones (katsu, tempura moriawase, yakiniku, teriyaki)")
    add_bullet("Postres: 3-4 opciones (mochi, dorayaki, matcha tiramisu, helado de yuzu)")
    add_bullet("Total carta: 45-65 platos (menos en omakase, más en mixto)")
    doc.add_heading("Menú engineering: clasificación de platos japoneses", level=2)
    add_bullet("Stars (alta popularidad + alto margen): Sashimi mixto, Highball, Gyoza, Ramen tonkotsu → destacar siempre")
    add_bullet("Plowhorses (alta popularidad + bajo margen): Nigiri salmón, Uramaki California → subir precio o ajustar porción")
    add_bullet("Puzzles (baja popularidad + alto margen): Uni nigiri, Tataki wagyu, Omakase → mejorar descripción y fotografía")
    add_bullet("Dogs (baja popularidad + bajo margen): Ensaladas genéricas, sopas de miso mal presentadas → eliminar o reinventar")
    tip("El sashimi mixto de 12 piezas es tu Star absoluto. El coste de un sashimi moriawase premium es 7-10€ y lo vendes a 28-36€. Con un sashimi + ramen + sake por persona ya tienes un ticket de 55€+ con márgenes excelentes. Es el plato con mejor ratio margen/satisfacción/foto de Instagram de toda la carta.")
    doc.add_page_break()

    # CH15
    doc.add_heading("15. Proveedores de Producto Japonés en España", level=1)
    doc.add_paragraph("Encontrar ingredientes auténticos japoneses en España es relativamente fácil gracias a distribuidores especializados y la red de tiendas asiáticas en Madrid y Barcelona.")
    doc.add_heading("Ingredientes clave y dónde encontrarlos", level=2)
    add_bullet("Pescado sashimi-grade: Mercamadrid, Mercabarna, lonjas de San Sebastián y Galicia. Proveedores especializados en pescado para sushi (Albert Rovira, Sakura Fish, Pescados Oceanic).")
    add_bullet("Arroz Koshihikari (o equivalente de sushi): distribuidores japoneses en España (Japan Sushi Express, Comercial Minamoto). Origen España, Italia o Japón.")
    add_bullet("Alga nori premium (yaki nori): distribuidores japoneses. Grado A (oro) es el mejor. Importación directa de Japón.")
    add_bullet("Salsa de soja (shoyu): Kikkoman es estándar. Premium: Yamasa, Yamaroku, Kishibori. Distribuidores japoneses.")
    add_bullet("Mirin, sake de cocina, vinagre de arroz: Kikkoman, Mitsukan. Distribuidores japoneses o tiendas asiáticas.")
    add_bullet("Wasabi: la mayoría es wasabi falso (rábano picante + colorante). Wasabi real (Wasabia japonica) fresco es MUY caro (300-600€/kg). Alternativa: pasta de wasabi premium real (congelada).")
    add_bullet("Panko (pan rallado japonés): distribuidores japoneses, tiendas asiáticas.")
    add_bullet("Miso (blanco, rojo, mixto): distribuidores japoneses, tiendas especializadas.")
    add_bullet("Carbón binchotan (para robata): importación desde Japón o alternativas de calidad de Laos/Vietnam. Distribuidores especializados.")
    add_bullet("Sake y whisky japonés: Distribuidores como Japan Sake (Madrid), Sake Distribución España, BlackTears Spirits. Algunos whiskies japoneses están en alocación por escasez.")
    add_bullet("Yuzu, shiso, mitsuba: difíciles frescos en España. Productores locales (Mercabarna), o congelado/en pasta.")
    add_bullet("Wagyu español y japonés: productores de wagyu en España (Finca Santa Rosalía, Luismi), importación de A5 japonés para omakase premium.")
    add_bullet("Hamachi, uni, anago: importación directa de Japón (vuelos diarios) a través de distribuidores especializados.")
    tip("Establece relación con un proveedor de pescado sashimi-grade diario ANTES de abrir. Sin pescado fresco diario de máxima calidad no hay restaurante japonés viable. Los proveedores serios tienen rutas que van directamente de lonja/Mercamadrid a tu restaurante cada mañana. Firma un acuerdo de suministro y haz un pedido tipo semanal.")
    doc.add_page_break()

    # CH16
    doc.add_heading("16. 15 Recetas Base con Food Cost", level=1)
    doc.add_paragraph("Estas 15 recetas forman la base de la carta de un restaurante japonés. Cada una incluye food cost estimado y PVP sugerido para mantener márgenes del 30-35%.")
    recipes = [
        ("Sashimi moriawase (12 piezas mixtas)", "8.50€", "32.00€", "27%"),
        ("Nigiri salmón (2 piezas)", "1.80€", "6.50€", "28%"),
        ("Nigiri toro / atún graso (2 piezas)", "5.00€", "15.00€", "33%"),
        ("Uramaki spicy tuna (8 piezas)", "3.20€", "13.00€", "25%"),
        ("Ramen tonkotsu completo", "3.80€", "14.00€", "27%"),
        ("Ramen shoyu con chashu", "3.40€", "13.00€", "26%"),
        ("Gyoza de cerdo (6 uds)", "1.50€", "7.50€", "20%"),
        ("Karaage de pollo (250g)", "2.20€", "10.00€", "22%"),
        ("Tempura moriawase (8 piezas)", "3.50€", "14.00€", "25%"),
        ("Yakitori pollo tare (3 brochetas)", "1.80€", "9.00€", "20%"),
        ("Yakitori wagyu (2 brochetas)", "6.50€", "22.00€", "30%"),
        ("Katsu de cerdo con arroz", "3.00€", "13.00€", "23%"),
        ("Chirashi bowl (salmón, atún, hamachi)", "6.00€", "22.00€", "27%"),
        ("Matcha tiramisu", "1.80€", "8.00€", "23%"),
        ("Dorayaki de matcha (2 uds)", "1.20€", "6.50€", "18%"),
    ]
    for name, fc, pvp, pct in recipes:
        doc.add_heading(name, level=2)
        add_bullet(f"Food cost: {fc}")
        add_bullet(f"PVP sugerido (sin IVA): {pvp}")
        add_bullet(f"% Food cost: {pct}")
    tip("El sashimi moriawase es tu plato con mejor margen absoluto y el más instagrameable. Un sashimi premium cuesta 8-10€ y lo vendes a 32€. Si el cliente pide sashimi + ramen + highball de whisky japonés, tienes un ticket de 55€ con un food cost total del 28%. Es el plato rey de la carta.")
    doc.add_page_break()

    # CH17
    doc.add_heading("17. Delivery Optimizado para Cocina Japonesa", level=1)
    doc.add_paragraph("La cocina japonesa tiene platos que viajan muy bien (sushi, bowls, katsu) y otros que no (ramen caliente, tempura). Adaptar el menú de delivery es clave para mantener la reputación de calidad.")
    doc.add_heading("Plataformas y canales", level=2)
    add_bullet("Glovo, Uber Eats, Just Eat: comisión 25-35% — útiles para visibilidad inicial")
    add_bullet("Delivery propio (web/app): comisión 0% pero requiere repartidores propios o Stuart/Paack")
    add_bullet("Take away en local: margen completo, packaging incluido en precio")
    doc.add_heading("Menú delivery japonés optimizado", level=2)
    add_bullet("Sushi boxes (12-24 piezas mixtas): producto estrella del delivery japonés. Caja negra mate con compartimentos, salsa soja en mini-bote, wasabi, gari.")
    add_bullet("Chirashi bowls: base de arroz + sashimi + verduras + salsa. Formato bowl cerrado, mantiene presentación.")
    add_bullet("Poke bowls estilo japonés: arroz, pescado marinado, aguacate, edamame, cebolla crujiente. Viaja perfecto.")
    add_bullet("Katsudon / Oyakodon bowls: arroz + katsu o pollo + huevo + salsa. Envases con tapa hermética.")
    add_bullet("Gyoza y karaage: fritos que aguantan 20-25 min. Papel absorbente en envase.")
    add_bullet("Ramen delivery: SOLO si entrega < 15 min. Caldo separado del fideo para que no se sobre-cocine. Cliente ensambla en casa.")
    add_bullet("NO incluir en delivery: tempura (se ablanda), sashimi puro sin cama de hielo (riesgo sanitario), sopa miso caliente.")
    add_bullet("Packaging: cajas negras de sushi con logotipo, mini-botes de salsa soja, palillos, servilleta japonesa. Coste 1.20-2.80€/pedido.")
    tip("El 'sushi box' premium es la mina de oro del delivery japonés. Incluye 18-24 piezas variadas, edamame, gyoza y te va a 35-45€. Margen del 65% con packaging premium. Plantea también 'sushi para compartir' (48-60 piezas) para cenas en casa/grupos. Ticket 60-90€ con food cost del 27%.")
    doc.add_page_break()

    # CH18
    doc.add_heading("18. Marketing y Calendario Cultural Japonés", level=1)
    doc.add_paragraph("Un restaurante japonés tiene un calendario cultural con eventos que generan picos de facturación. El Hanami (floración del cerezo en marzo-abril), el Día del Sushi Internacional (18 junio), Navidad japonesa y el Año Nuevo japonés son oportunidades excelentes.")
    doc.add_heading("Eventos culturales para marketing", level=2)
    add_bullet("Hanami (marzo-abril): celebración de la floración del cerezo. Menú sakura especial, decoración con ramas de cerezo, cócteles de yuzu y sakura. Genera mucho contenido visual.")
    add_bullet("Día del Sushi Internacional (18 junio): promociones de sushi, menús omakase degustación, evento con itamae explicando técnicas. Excelente para redes sociales.")
    add_bullet("Obon (mediados agosto): festival tradicional japonés, menús especiales, eventos culturales con música y taiko.")
    add_bullet("Año Nuevo japonés (Oshogatsu, 1 enero): menú osechi ryori (comida tradicional de año nuevo), soba de toshikoshi (31 diciembre).")
    add_bullet("Día de la cerveza japonesa (1 agosto): promociones de cervezas Asahi, Sapporo, Kirin con platos izakaya.")
    add_bullet("Día del whisky japonés: eventos de cata y maridaje. Con la escasez actual, generan enorme interés.")
    doc.add_heading("Canales de marketing", level=2)
    add_num("Instagram: contenido visual de sushi, ramen, itamae en barra. Reels del itamae cortando pescado son virales.")
    add_num("TikTok: vídeos del itamae preparando nigiri, ramen, montaje de sashimi. Alto potencial viral.")
    add_num("Google My Business: ficha completa con fotos profesionales, menú, reseñas.")
    add_num("TheFork / TripAdvisor: perfiles con fotos y respuesta a reseñas (crítico para japonés premium).")
    add_num("Web propia: menú, historia, reservas online, carta de sakes y whiskies.")
    doc.add_heading("Presupuesto marketing mensual", level=2)
    add_bullet("Google Ads local: 300-700€/mes")
    add_bullet("Instagram/Facebook Ads: 300-700€/mes")
    add_bullet("Fotógrafo food: 300-500€/sesión trimestral")
    add_bullet("Community manager (si se externaliza): 400-800€/mes")
    add_bullet("Total recomendado: 800-2.000€/mes (2-3% facturación)")
    tip("El contenido visual de un sushi-ya se vende solo. Invierte en fotografía profesional de platos, un vídeo del itamae cortando y montando nigiri, y publica Reels constantemente. Los japoneses premium tienen el mejor ROI en Instagram de todo el sector hostelería. Un buen Reel puede traer 50-100 reservas en 48h.")
    doc.add_page_break()

    # CH19
    doc.add_heading("19. Tecnología para Restaurante Japonés", level=1)
    doc.add_paragraph("La tecnología en un restaurante japonés debe ser invisible pero fiable. En un omakase con reservas de 3 semanas, no puedes permitirte un error de gestión. En un ramen-ya con cola diaria, necesitas gestión de espera digital.")
    doc.add_heading("Stack tecnológico recomendado", level=2)
    add_bullet("TPV: Lightspeed, Square, Revo. Cloud, con reporting en tiempo real. 60-150€/mes.")
    add_bullet("Reservas: TheFork, CoverManager, Resy. Gestión de turnos, depósitos para omakase, política de no-shows. 0-150€/mes.")
    add_bullet("Delivery: integrador (Ordatic, Deliverect) que centraliza Glovo + Uber Eats + Just Eat. 80-150€/mes.")
    add_bullet("Contabilidad: Holded, Quipu. Facturas automáticas, conexión con asesoría. 30-60€/mes.")
    add_bullet("RRHH y turnos: Factorial, Kenjo. Cuadrantes, fichajes, nóminas. 4-8€/empleado/mes.")
    add_bullet("Gestión de cola (si ramen-ya con espera): Waitwhile, Qminder. Clientes dan su nombre y reciben notificación. 30-80€/mes.")
    add_bullet("Web y carta digital: QR con carta actualizable, sección de sakes. WordPress o Squarespace. 10-30€/mes.")
    add_bullet("WiFi profesional: router dual-band, red separada para clientes y cocina. 50-80€/mes.")
    tip("En un omakase, la gestión de reservas con depósito es fundamental. Cobra 30-50€ por persona como depósito de reserva (reembolsable contra la cena) para evitar no-shows. Los no-shows en un omakase de 14 plazas pueden destruir tu facturación del día. CoverManager y Resy permiten gestionar esto automáticamente.")
    doc.add_page_break()

    # CH20
    doc.add_heading("20. Plan de Acción: De la Idea a la Inauguración", level=1)
    doc.add_paragraph("Abrir un restaurante japonés en España es un proyecto de 12-16 meses desde la primera idea hasta la inauguración. Este capítulo resume el cronograma y los hitos clave.")
    doc.add_heading("Fases del proyecto", level=2)
    add_bullet("Meses 1-3: Plan de negocio, búsqueda de financiación, constitución de empresa, estudio de mercado.")
    add_bullet("Meses 3-5: Búsqueda y contratación de local, negociación de alquiler, licencia de actividad.")
    add_bullet("Meses 4-7: Obra civil, instalaciones (especialmente extracción reforzada para robata), decoración japonesa.")
    add_bullet("Meses 5-7: Equipamiento de cocina japonesa (suihanki, vitrina sashimi, ramen cooker, robata, cuchillos japoneses).")
    add_bullet("Meses 5-8: Búsqueda y reclutamiento de itamae (la decisión más crítica, puede tardar meses).")
    add_bullet("Meses 5-8: Establecer relaciones con proveedores de pescado sashimi-grade y distribuidores japoneses.")
    add_bullet("Meses 8-10: Contratación y formación del resto del equipo (sushi chef, ramen cook, robata cook, sommelier sake).")
    add_bullet("Meses 9-11: Branding, web, redes sociales, fotografía profesional.")
    add_bullet("Mes 11-12: Soft opening con amigos, familia, influencers y prensa gastronómica.")
    add_bullet("Mes 12-13: Ajustes finales e INAUGURACIÓN.")
    add_bullet("Mes 13-16: Seguimiento, ajustes de carta y optimización de costes.")
    tip("No abras sin haber hecho mínimo 4 soft openings. Necesitas probar la barra de sushi con el itamae a pleno rendimiento, los tiempos de ramen, el flujo de sala con camareros nuevos, la coctelería de sake y whisky y la operativa de delivery. Cada soft opening te enseña fallos que no puedes detectar en papel. Y con un japonés premium no puedes permitirte un mal servicio el primer día — el boca a boca es brutal.")

    # Save
    path = os.path.join(OUTPUT_DIR, "guia-restaurante-japones.docx")
    doc.save(path)
    print(f"✓ {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 2. EXCEL TEMPLATES (8)
# ═══════════════════════════════════════════════════════════

def gen_plan_financiero():
    wb = Workbook()
    instr_sheet(wb, "Plan Financiero 3 Años — Restaurante Japonés 60 Plazas", [
        "Rellena las celdas verdes con los datos de tu proyecto.",
        "Las fórmulas se recalculan automáticamente.",
        "Pestaña 'Inversión' = CAPEX inicial desglosado.",
        "Pestaña 'P&L Mensual' = cuenta de resultados mes a mes.",
        "Pestaña 'Proyección 3 Años' = evolución anual.",
    ])
    # Inversión sheet
    ws = wb.create_sheet("Inversión"); ws.sheet_properties.tabColor = "4CAF50"
    title_block(ws, "Inversión Inicial — Restaurante Japonés 60 Plazas")
    hdrs = ["Categoría", "Partida", "Coste Estimado (€)", "Coste Real (€)", "Diferencia", "Notas"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 6)
    widths = [25, 40, 18, 18, 18, 30]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    items = [
        ("Obra Civil", "Acondicionamiento local", 110000),
        ("Obra Civil", "Instalaciones eléctricas y gas", 14000),
        ("Obra Civil", "Fontanería y saneamiento", 9000),
        ("Obra Civil", "Climatización / HVAC", 10000),
        ("Obra Civil", "Extracción reforzada (robata, teppanyaki, tempura)", 12000),
        ("Cocina JP", "Suihanki (arrocera japonesa profesional 5-10L)", 3200),
        ("Cocina JP", "Vitrina refrigerada de sashimi para barra sushi", 6500),
        ("Cocina JP", "Ramen cooker / pasta cooker profesional", 5500),
        ("Cocina JP", "Ollas de caldo 50-100L tonkotsu/shoyu (2 uds)", 2500),
        ("Cocina JP", "Parrilla robata carbón binchotan", 6500),
        ("Cocina JP", "Plancha teppanyaki profesional", 4500),
        ("Cocina JP", "Cuchillos japoneses set itamae (yanagiba, deba, usuba)", 2800),
        ("Cocina JP", "Ohitsu (recipiente cedro para shari) + accesorios", 800),
        ("Cocina JP", "Maquina sake caliente (tokkuri warmer)", 600),
        ("Cocina Estándar", "Horno mixto Rational/Unox", 12000),
        ("Cocina Estándar", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina Estándar", "Wok profesional alta potencia (yakisoba)", 2500),
        ("Cocina Estándar", "Freidora doble (tempura, karaage)", 2500),
        ("Cocina Estándar", "Cámaras frigoríficas (2)", 6000),
        ("Cocina Estándar", "Congelador CERTIFICADO -20°C (anisakis)", 4500),
        ("Cocina Estándar", "Tren de lavado", 5500),
        ("Cocina Estándar", "Campana extractora REFORZADA (robata)", 10000),
        ("Cocina Estándar", "Menaje y utensilios japoneses", 3500),
        ("Barra Sushi", "Barra de sushi profesional madera/vitrina", 9000),
        ("Barra Sake", "Vitrina refrigerada sake y whisky japonés", 4500),
        ("Barra Sake", "Utensilios coctelería japonesa", 700),
        ("Barra Sake", "Stock inicial sake y whisky (30-50 refs)", 10000),
        ("Barra Sake", "Máquina de hielo", 1800),
        ("Sala", "Mobiliario interior mesas madera + sillas (60 plazas)", 22000),
        ("Sala", "Barra sushi taburetes altos", 4000),
        ("Decoración", "Madera natural en paredes de acento", 8000),
        ("Decoración", "Noren (cortina entrada) + detalles textiles", 800),
        ("Decoración", "Lámparas washi (papel de arroz) y iluminación cálida", 3500),
        ("Decoración", "Jardín zen, bambú, piedras", 2500),
        ("Decoración", "Vajilla cerámica japonesa artesanal", 4500),
        ("Decoración", "Carbonilla / grava / detalles piedra", 1200),
        ("Terraza", "Mobiliario exterior (si aplica)", 5000),
        ("Tecnología", "TPV + pantallas cocina", 3500),
        ("Tecnología", "Web + carta QR + WiFi + sistema reservas", 3500),
        ("Licencias", "Licencia actividad + terraza + proyecto técnico", 12000),
        ("Licencias", "Registro EORI / importación (si aplica)", 500),
        ("Marketing", "Branding japonés + diseño + logotipo", 5000),
        ("Marketing", "Fotos + vídeo + campaña pre-apertura", 5000),
        ("Stock", "Materia prima inicial (pescado, arroz, algas, salsas)", 15000),
        ("Stock", "Bebidas (cerveza japonesa, refrescos, té)", 5000),
        ("Maniobra", "Fondo de maniobra (3 meses)", 70000),
    ]
    r += 1
    start_r = r
    for cat, partida, coste in items:
        sdc(ws, r, 1, cat); sdc(ws, r, 2, partida)
        sdc(ws, r, 3, coste, fmt=cur_fmt); sdc(ws, r, 4, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}-C{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 2, "TOTAL INVERSIÓN", font=bld_font)
    sdc(ws, r, 3, f"=SUM(C{start_r}:C{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=D{r}-C{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 6)

    # P&L Mensual sheet
    ws2 = wb.create_sheet("P&L Mensual"); ws2.sheet_properties.tabColor = "2196F3"
    title_block(ws2, "P&L Mensual — Restaurante Japonés 60 Plazas", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic", "TOTAL"]
    r = 4
    for i, m in enumerate(months, 1): ws2.cell(row=r, column=i, value=m)
    shr(ws2, r, 14)
    ws2.column_dimensions['A'].width = 28
    for i in range(2, 15): ws2.column_dimensions[get_column_letter(i)].width = 12
    ws2.freeze_panes = f"B{r+1}"
    rows_data = [
        ("INGRESOS", None, True),
        ("Ventas sala (sushi + ramen + izakaya)", 45000, False),
        ("Barra sake + whisky + coctelería", 15000, False),
        ("Delivery (sushi boxes, bowls, katsu)", 10000, False),
        ("Bebidas (cerveza, refrescos, té)", 8000, False),
        ("TOTAL INGRESOS", "=SUM", True),
        ("", None, False),
        ("COSTES VARIABLES", None, True),
        ("Food cost (33%)", "=0.33*", False),
        ("Packaging delivery", 700, False),
        ("TOTAL COSTES VARIABLES", "=SUM_CV", True),
        ("", None, False),
        ("COSTES FIJOS", None, True),
        ("Alquiler", 6500, False),
        ("Personal cocina (itamae, sushi chef, ramen, robata, hot)", 22000, False),
        ("Personal sala (sommelier sake, camareros)", 13000, False),
        ("Suministros (luz, gas, agua)", 3500, False),
        ("Seguros", 500, False),
        ("Gestoría + contabilidad", 500, False),
        ("Marketing", 1200, False),
        ("Tecnología (TPV, delivery, reservas)", 450, False),
        ("Mantenimiento", 400, False),
        ("Reposición stock sake/whisky", 1800, False),
        ("Varios / imprevistos", 700, False),
        ("TOTAL COSTES FIJOS", "=SUM_CF", True),
        ("", None, False),
        ("EBITDA", "=EBITDA", True),
        ("% EBITDA", "=%EBITDA", True),
    ]
    r += 1
    for label, val, bold in rows_data:
        sdc(ws2, r, 1, label, font=bld_font if bold else dat_font)
        if val and val not in ("=SUM", "=SUM_CV", "=SUM_CF", "=EBITDA", "=%EBITDA") and not str(val).startswith("=0."):
            for c in range(2, 14):
                sdc(ws2, r, c, val, fill=inp_fill, fmt=cur_fmt)
            sdc(ws2, r, 14, f"=SUM(B{r}:M{r})", font=bld_font, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws2, r, 14)

    path = os.path.join(OUTPUT_DIR, "plan-financiero-3-anos.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_capex():
    wb = Workbook()
    instr_sheet(wb, "Calculadora CAPEX — Restaurante Japonés 60 Plazas", [
        "Introduce tus costes reales en las celdas verdes.",
        "La columna 'Diferencia' se calcula automáticamente.",
        "Usa la columna 'Prioridad' para planificar fases de inversión.",
    ])
    ws = wb.create_sheet("CAPEX Desglosado"); ws.sheet_properties.tabColor = "FF9800"
    title_block(ws, "Calculadora CAPEX — Inversión 250K-500K€")
    hdrs = ["#", "Categoría", "Partida", "Estimado (€)", "Real (€)", "Diferencia", "Prioridad", "Proveedor"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8)
    widths = [5, 20, 40, 14, 14, 14, 12, 25]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    dv = DataValidation(type="list", formula1='"Alta,Media,Baja"', allow_blank=True)
    ws.add_data_validation(dv)
    items = [
        ("Obra", "Acondicionamiento general", 110000),
        ("Obra", "Electricidad + gas", 14000),
        ("Obra", "Fontanería", 9000),
        ("Obra", "Climatización", 10000),
        ("Obra", "Extracción reforzada robata/teppan/tempura", 12000),
        ("Cocina JP", "Suihanki arrocera japonesa profesional", 3200),
        ("Cocina JP", "Vitrina refrigerada sashimi (barra sushi)", 6500),
        ("Cocina JP", "Ramen cooker profesional", 5500),
        ("Cocina JP", "Ollas caldo 50-100L (2 uds)", 2500),
        ("Cocina JP", "Parrilla robata carbón binchotan", 6500),
        ("Cocina JP", "Plancha teppanyaki profesional", 4500),
        ("Cocina JP", "Cuchillos japoneses set (yanagiba, deba, usuba)", 2800),
        ("Cocina JP", "Ohitsu cedro + accesorios", 800),
        ("Cocina", "Horno mixto Rational/Unox", 12000),
        ("Cocina", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina", "Wok profesional alta potencia", 2500),
        ("Cocina", "Freidora doble (tempura, karaage)", 2500),
        ("Cocina", "Cámaras frigoríficas (2)", 6000),
        ("Cocina", "Congelador CERTIFICADO -20°C anisakis", 4500),
        ("Cocina", "Tren de lavado", 5500),
        ("Cocina", "Campana extractora REFORZADA robata", 10000),
        ("Cocina", "Menaje y utensilios japoneses", 3500),
        ("Barra Sushi", "Barra sushi profesional madera", 9000),
        ("Barra Sake", "Vitrina refrigerada sake/whisky japonés", 4500),
        ("Barra Sake", "Utensilios coctelería + máquina hielo", 2500),
        ("Barra Sake", "Stock inicial sake/whisky (30-50 refs)", 10000),
        ("Sala", "Mesas madera, sillas, barra (60 plazas)", 26000),
        ("Decoración", "Madera, noren, washi, jardín zen, cerámica", 20000),
        ("Terraza", "Mobiliario exterior", 5000),
        ("Tech", "TPV + pantallas cocina + reservas", 7000),
        ("Licencias", "Licencia actividad + terraza + proyecto", 12500),
        ("Marketing", "Branding + fotos + vídeo + pre-apertura", 10000),
        ("Stock", "Materia prima + pescado + arroz + salsas", 15000),
        ("Stock", "Bebidas (cerveza, té, refrescos)", 5000),
        ("Maniobra", "Fondo 3 meses", 70000),
    ]
    r += 1; start_r = r
    for idx, (cat, partida, est) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, partida)
        sdc(ws, r, 4, est, fmt=cur_fmt); sdc(ws, r, 5, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 6, f"=E{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 7, "Alta", fill=inp_fill); dv.add(ws.cell(row=r, column=7))
        sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 3, "TOTAL", font=bld_font)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=SUM(E{start_r}:E{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 6, f"=E{r}-D{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 8)
    path = os.path.join(OUTPUT_DIR, "calculadora-capex.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_pl_mensual():
    wb = Workbook()
    instr_sheet(wb, "P&L Mensual con 3 Escenarios — Restaurante Japonés", [
        "3 pestañas: Pesimista, Realista, Optimista.",
        "Modifica los ingresos en cada escenario.",
        "Los costes fijos son iguales en los 3 escenarios.",
    ])
    for scenario, factor, color in [("Pesimista", 0.75, "F44336"), ("Realista", 1.0, "4CAF50"), ("Optimista", 1.25, "2196F3")]:
        ws = wb.create_sheet(scenario); ws.sheet_properties.tabColor = color
        title_block(ws, f"P&L Mensual — Escenario {scenario}", 4)
        hdrs = ["Concepto", "Mensual (€)", "Anual (€)", "% s/Ventas"]
        r = 4
        for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
        shr(ws, r, 4)
        ws.column_dimensions['A'].width = 35
        for i in range(2, 5): ws.column_dimensions[get_column_letter(i)].width = 16
        base_ventas = int(78000 * factor)
        rows = [
            ("INGRESOS", None, True),
            ("Ventas sala (sushi + ramen + izakaya)", int(45000 * factor)),
            ("Barra sake + whisky japonés", int(15000 * factor)),
            ("Delivery (sushi boxes, bowls)", int(10000 * factor)),
            ("Bebidas (cerveza, té, refrescos)", int(8000 * factor)),
            ("TOTAL INGRESOS", base_ventas, True),
            ("", None, False),
            ("COSTES VARIABLES", None, True),
            ("Food cost (33%)", int(base_ventas * 0.33)),
            ("Packaging delivery", 700),
            ("TOTAL COSTES VARIABLES", int(base_ventas * 0.33 + 700), True),
            ("", None, False),
            ("COSTES FIJOS", None, True),
            ("Alquiler", 6500),
            ("Personal cocina (itamae + equipo)", 22000),
            ("Personal sala", 13000),
            ("Suministros", 3500),
            ("Seguros", 500),
            ("Gestoría", 500),
            ("Marketing", 1200),
            ("Tecnología", 450),
            ("Mantenimiento", 400),
            ("Reposición sake/whisky", 1800),
            ("Varios", 700),
            ("TOTAL COSTES FIJOS", 50550, True),
            ("", None, False),
            ("EBITDA", int(base_ventas - base_ventas * 0.33 - 700 - 50550), True),
        ]
        r += 1
        for item in rows:
            if len(item) == 3:
                label, val, bold = item
            else:
                label, val = item; bold = False
            sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
            if val is not None:
                sdc(ws, r, 2, val, font=bld_font if bold else dat_font, fmt=cur_fmt)
                sdc(ws, r, 3, val * 12, font=bld_font if bold else dat_font, fmt=cur_fmt)
            r += 1
        r += 1; brand_footer(ws, r, 4)
    path = os.path.join(OUTPUT_DIR, "pl-mensual-escenarios.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cash_flow():
    wb = Workbook()
    instr_sheet(wb, "Cash Flow y Break-Even — Restaurante Japonés 60 Plazas", [
        "Pestaña 'Cash Flow' = flujo de caja mensual 12 meses.",
        "Pestaña 'Break-Even' = calculadora de punto de equilibrio.",
        "Rellena las celdas verdes con tus datos reales.",
    ])
    ws = wb.create_sheet("Cash Flow"); ws.sheet_properties.tabColor = "009688"
    title_block(ws, "Cash Flow 12 Meses — Restaurante Japonés", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
    r = 4
    for i, m in enumerate(months, 1): ws.cell(row=r, column=i, value=m)
    shr(ws, r, 13)
    ws.column_dimensions['A'].width = 28
    for i in range(2, 14): ws.column_dimensions[get_column_letter(i)].width = 12
    concepts = [
        ("Saldo inicial", True), ("", False),
        ("ENTRADAS", True), ("Ventas sala (sushi + ramen + izakaya)", False), ("Barra sake/whisky", False),
        ("Delivery (sushi boxes, bowls)", False), ("Bebidas + té", False), ("Total entradas", True), ("", False),
        ("SALIDAS", True), ("Pescado sashimi-grade + arroz + importados", False), ("Personal cocina (itamae)", False),
        ("Personal sala", False), ("Alquiler", False), ("Suministros", False), ("Marketing", False),
        ("Reposición stock sake/whisky", False), ("Tecnología", False), ("Otros gastos", False),
        ("Total salidas", True), ("", False),
        ("Flujo neto mensual", True), ("Saldo acumulado", True),
    ]
    r += 1
    for label, bold in concepts:
        sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
        for c in range(2, 14):
            sdc(ws, r, c, None, fill=inp_fill if not bold and label else None, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws, r, 13)

    # Break-even
    ws2 = wb.create_sheet("Break-Even"); ws2.sheet_properties.tabColor = "FF5722"
    title_block(ws2, "Calculadora Break-Even — Restaurante Japonés", 4)
    r = 4
    params = [
        ("Ticket medio (€)", 45),
        ("Comensales/día promedio", 65),
        ("Días abierto/mes", 26),
        ("Food cost (%)", 0.33),
        ("Costes fijos mensuales (€)", 50550),
    ]
    for label, val in params:
        r += 1
        sdc(ws2, r, 1, label, font=bld_font); sdc(ws2, r, 2, val, fill=inp_fill, fmt=cur_fmt if isinstance(val, int) else pct_fmt)
    r += 2
    sdc(ws2, r, 1, "Facturación mensual estimada", font=bld_font)
    sdc(ws2, r, 2, "=B5*B6*B7", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Margen contribución mensual", font=bld_font)
    sdc(ws2, r, 2, "=B12*(1-B8)", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Break-Even (meses)", font=bld_font)
    ws2.column_dimensions['A'].width = 35; ws2.column_dimensions['B'].width = 20
    r += 2; brand_footer(ws2, r, 4)
    path = os.path.join(OUTPUT_DIR, "cash-flow-break-even.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_escandallo():
    wb = Workbook()
    instr_sheet(wb, "Escandallo Maestro — Restaurante Japonés", [
        "Una ficha técnica por plato.",
        "Introduce ingredientes, cantidades y precios.",
        "El food cost se calcula automáticamente.",
    ])
    ws = wb.create_sheet("Escandallo"); ws.sheet_properties.tabColor = "E91E63"
    title_block(ws, "Escandallo Maestro — Fichas Técnicas de Platos Japoneses")
    hdrs = ["#", "Ingrediente", "Cantidad (g/ml)", "Precio/Kg (€)", "Coste (€)", "Merma (%)", "Coste Real (€)"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 7)
    widths = [5, 30, 16, 14, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "NOMBRE DEL PLATO:", font=bld_font)
    sdc(ws, r, 3, "Ejemplo: Sashimi Moriawase 12 piezas", fill=inp_fill)
    r += 1
    ingredients = [
        ("Salmón sashimi-grade (previamente congelado -20°C/24h)", 80, 22.00, 15),
        ("Atún rojo sashimi-grade (previamente congelado)", 60, 55.00, 20),
        ("Hamachi (pez limón) sashimi-grade", 50, 38.00, 18),
        ("Lubina sashimi-grade", 40, 28.00, 15),
        ("Vieira sashimi-grade", 35, 45.00, 10),
        ("Pulpo cocido", 30, 24.00, 5),
        ("Rábano daikon (guarnición juliana)", 40, 3.50, 15),
        ("Shiso hoja fresca", 3, 80.00, 0),
        ("Wasabi pasta real", 4, 180.00, 0),
        ("Jengibre encurtido (gari)", 10, 12.00, 0),
        ("Salsa soja premium", 20, 18.00, 0),
    ]
    start_r = r
    for idx, (ing, cant, precio, merma) in enumerate(ingredients, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, ing, fill=inp_fill)
        sdc(ws, r, 3, cant, fill=inp_fill)
        sdc(ws, r, 4, precio, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=(C{r}/1000)*D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, merma / 100, fill=inp_fill, fmt=pct_fmt)
        sdc(ws, r, 7, f"=E{r}*(1+F{r})", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 2, "COSTE TOTAL PLATO", font=bld_font)
    sdc(ws, r, 7, f"=SUM(G{start_r}:G{r-2})", font=bld_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP sugerido (food cost 33%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}/0.33", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP con IVA (10%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}*1.10", font=frm_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 7)
    path = os.path.join(OUTPUT_DIR, "escandallo-maestro-japones.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_menu_engineering():
    wb = Workbook()
    instr_sheet(wb, "Menú Engineering Matrix — Restaurante Japonés", [
        "Introduce los platos de tu carta con ventas y food cost.",
        "La matrix clasifica automáticamente cada plato.",
        "Stars = mantener. Plowhorses = subir precio. Puzzles = promocionar. Dogs = eliminar.",
    ])
    ws = wb.create_sheet("Matrix"); ws.sheet_properties.tabColor = "9C27B0"
    title_block(ws, "Menú Engineering Matrix — Restaurante Japonés", 10)
    hdrs = ["#", "Plato", "PVP (€)", "Food Cost (€)", "% Food Cost", "Uds. Vendidas/Mes", "Margen Unit. (€)", "Margen Total (€)", "Popularidad", "Clasificación"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    widths = [5, 30, 10, 12, 10, 16, 12, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    platos = [
        ("Sashimi moriawase 12 piezas", 32.00, 8.50, 140),
        ("Nigiri salmón (2 uds)", 6.50, 1.80, 220),
        ("Nigiri toro (2 uds)", 15.00, 5.00, 60),
        ("Uramaki spicy tuna (8)", 13.00, 3.20, 130),
        ("Ramen tonkotsu", 14.00, 3.80, 180),
        ("Ramen shoyu chashu", 13.00, 3.40, 120),
        ("Gyoza cerdo (6)", 7.50, 1.50, 150),
        ("Karaage de pollo", 10.00, 2.20, 110),
        ("Tempura moriawase (8)", 14.00, 3.50, 90),
        ("Yakitori pollo tare (3)", 9.00, 1.80, 130),
        ("Yakitori wagyu (2)", 22.00, 6.50, 50),
        ("Katsu cerdo + arroz", 13.00, 3.00, 85),
        ("Chirashi bowl", 22.00, 6.00, 75),
        ("Highball whisky japonés", 10.00, 1.50, 260),
        ("Sake daiginjo (copa)", 15.00, 4.00, 90),
    ]
    r += 1
    for idx, (plato, pvp, fc, uds) in enumerate(platos, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, plato, fill=inp_fill)
        sdc(ws, r, 3, pvp, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 4, fc, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}/C{r}", font=frm_font, fmt=pct_fmt)
        sdc(ws, r, 6, uds, fill=inp_fill)
        sdc(ws, r, 7, f"=C{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 8, f"=G{r}*F{r}", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "menu-engineering-matrix.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_ticket():
    wb = Workbook()
    instr_sheet(wb, "Calculadora Ticket Medio + Margen Sake/Whisky", [
        "Simula diferentes escenarios de ticket medio.",
        "Incluye pestaña separada para margen de barra de sake y whisky japonés.",
        "Rellena las celdas verdes con tus datos.",
    ])
    ws = wb.create_sheet("Ticket Medio"); ws.sheet_properties.tabColor = "00BCD4"
    title_block(ws, "Calculadora Ticket Medio — Restaurante Japonés", 4)
    r = 4
    params = [
        ("% clientes que piden sashimi/sushi para compartir", 0.70),
        ("Precio medio sashimi/sushi (€)", 22.00),
        ("Precio medio ramen/principal (€)", 15.00),
        ("% clientes que piden postre", 0.25),
        ("Precio medio postre (€)", 8.00),
        ("% clientes que piden sake/whisky", 0.50),
        ("Precio medio sake/whisky (€)", 12.00),
        ("% clientes que piden cerveza/té", 0.55),
        ("Precio medio cerveza/té (€)", 5.00),
    ]
    r += 1
    for label, val in params:
        sdc(ws, r, 1, label, font=dat_font)
        sdc(ws, r, 2, val, fill=inp_fill, fmt=pct_fmt if isinstance(val, float) and val < 1 else cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 1, "TICKET MEDIO ESTIMADO", font=bld_font)
    sdc(ws, r, 2, None, font=bld_font, fmt=cur_fmt)
    ws.column_dimensions['A'].width = 50; ws.column_dimensions['B'].width = 18

    # Margen sake/whisky tab
    ws2 = wb.create_sheet("Margen Sake-Whisky"); ws2.sheet_properties.tabColor = "FFC107"
    title_block(ws2, "Simulador Margen Barra de Sake y Whisky Japonés", 4)
    r = 4
    sdc(ws2, r+1, 1, "Precio medio copa sake (€)", font=bld_font); sdc(ws2, r+1, 2, 12.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+2, 1, "Coste medio copa sake (€)", font=bld_font); sdc(ws2, r+2, 2, 3.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+3, 1, "Margen por copa sake (€)", font=bld_font); sdc(ws2, r+3, 2, f"=B{r+1}-B{r+2}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+4, 1, "% Margen sake", font=bld_font); sdc(ws2, r+4, 2, f"=B{r+3}/B{r+1}", font=frm_font, fmt=pct_fmt)
    sdc(ws2, r+5, 1, "Precio medio highball whisky japonés (€)", font=bld_font); sdc(ws2, r+5, 2, 10.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+6, 1, "Coste medio highball (€)", font=bld_font); sdc(ws2, r+6, 2, 1.50, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+7, 1, "Margen por highball (€)", font=bld_font); sdc(ws2, r+7, 2, f"=B{r+5}-B{r+6}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+8, 1, "Copas sake vendidas/día", font=bld_font); sdc(ws2, r+8, 2, 25, fill=inp_fill)
    sdc(ws2, r+9, 1, "Highballs vendidos/día", font=bld_font); sdc(ws2, r+9, 2, 40, fill=inp_fill)
    sdc(ws2, r+10, 1, "Días/mes", font=bld_font); sdc(ws2, r+10, 2, 26, fill=inp_fill)
    sdc(ws2, r+11, 1, "Facturación barra sake/whisky/mes (€)", font=bld_font)
    sdc(ws2, r+11, 2, f"=(B{r+1}*B{r+8}+B{r+5}*B{r+9})*B{r+10}", font=frm_font, fmt=cur_fmt)
    ws2.column_dimensions['A'].width = 42; ws2.column_dimensions['B'].width = 18
    path = os.path.join(OUTPUT_DIR, "calculadora-ticket-medio.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cronograma():
    wb = Workbook()
    instr_sheet(wb, "Cronograma Apertura Gantt 12 Meses — Restaurante Japonés", [
        "Fases y tareas para abrir un restaurante japonés.",
        "Marca las celdas con 'X' para indicar el mes activo.",
        "Incluye fase de búsqueda de itamae y proveedores de pescado sashimi-grade.",
    ])
    ws = wb.create_sheet("Gantt"); ws.sheet_properties.tabColor = "795548"
    title_block(ws, "Cronograma de Apertura — Restaurante Japonés 60 Plazas", 15)
    hdrs = ["#", "Fase", "Tarea", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "M8", "M9", "M10", "M11", "M12"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 15)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 45
    for i in range(4, 16): ws.column_dimensions[get_column_letter(i)].width = 5
    tasks = [
        ("Planificación", "Estudio de mercado y viabilidad", [1, 2]),
        ("Planificación", "Plan financiero y búsqueda financiación", [1, 2, 3]),
        ("Planificación", "Constitución empresa / alta autónomo", [2, 3]),
        ("Itamae", "Búsqueda y reclutamiento de itamae (crítico)", [2, 3, 4, 5]),
        ("Importación", "Contacto distribuidores productos japoneses", [2, 3]),
        ("Proveedores", "Selección proveedor pescado sashimi-grade diario", [3, 4]),
        ("Local", "Búsqueda y selección de local", [2, 3, 4]),
        ("Local", "Negociación alquiler y contrato", [3, 4]),
        ("Licencias", "Licencia actividad / declaración responsable", [3, 4, 5]),
        ("Licencias", "Proyecto técnico y visado", [3, 4]),
        ("Obra", "Acondicionamiento y obra civil", [4, 5, 6, 7]),
        ("Obra", "Instalaciones (electricidad, gas, extracción reforzada robata)", [5, 6, 7]),
        ("Obra", "Decoración japonesa (madera, washi, jardín zen)", [6, 7]),
        ("Equipamiento", "Pedido equipo cocina japonesa (suihanki, vitrina, robata)", [5, 6]),
        ("Equipamiento", "Instalación barra sushi + vitrina sashimi", [6, 7]),
        ("Equipamiento", "Recepción e instalación equipos", [7, 8]),
        ("Equipamiento", "Mobiliario sala + barra sake", [6, 7]),
        ("Equipo", "Selección sushi chef + ramen cook + sommelier sake", [7, 8, 9]),
        ("Equipo", "Selección y contratación resto equipo", [8, 9]),
        ("Equipo", "Formación equipo (técnicas sushi + sake + protocolo)", [9, 10]),
        ("Marketing", "Branding, web, redes sociales", [6, 7, 8]),
        ("Marketing", "Fotos profesionales de platos (sushi, ramen)", [9]),
        ("Marketing", "Campaña pre-apertura", [9, 10]),
        ("APPCC", "Plan APPCC + registro sanitario (pescado crudo + anisakis)", [7, 8]),
        ("Pre-apertura", "Soft opening (amigos, prensa, influencers)", [10, 11]),
        ("Pre-apertura", "Ajustes finales", [11]),
        ("Apertura", "INAUGURACIÓN", [12]),
        ("Post-apertura", "Seguimiento y ajustes", [12]),
    ]
    r += 1
    for idx, (fase, tarea, meses) in enumerate(tasks, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, fase); sdc(ws, r, 3, tarea)
        for m in meses:
            sdc(ws, r, m + 3, "X", font=bld_font, fill=sec_fill, align=c_align)
        r += 1
    r += 1; brand_footer(ws, r, 15)
    path = os.path.join(OUTPUT_DIR, "cronograma-apertura-gantt.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_plantilla_turnos():
    wb = Workbook()
    instr_sheet(wb, "Plantilla Turnos Brigada — Restaurante Japonés", [
        "Cuadrante semanal para todo el equipo (cocina + sala).",
        "Rellena los turnos con: M (mañana), T (tarde), P (partido), L (libre).",
        "Incluye puestos específicos: itamae, sushi chef, ramen cook, robata cook, sommelier sake.",
    ])
    ws = wb.create_sheet("Turnos Semana"); ws.sheet_properties.tabColor = "607D8B"
    title_block(ws, "Cuadrante Semanal — Restaurante Japonés 60 Plazas", 10)
    hdrs = ["#", "Nombre", "Puesto", "Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 28
    for i in range(4, 11): ws.column_dimensions[get_column_letter(i)].width = 8
    dv = DataValidation(type="list", formula1='"M,T,P,L,V"', allow_blank=True)
    ws.add_data_validation(dv)
    staff = [
        ("Itamae / Jefe de cocina", "P"), ("Sushi chef (segundo)", "P"),
        ("Ramen cook", "P"), ("Robata / Yakitori cook", "T"),
        ("Tempura / Hot kitchen cook", "P"), ("Cocinero 1 (izakaya)", "P"),
        ("Preparación pescado", "M"), ("Ayudante cocina / office", "P"),
        ("Office", "P"),
        ("Encargado sala / Sommelier sake", "P"), ("Camarero 1", "P"),
        ("Camarero 2", "P"), ("Camarero 3", "T"),
        ("Runner", "P"),
    ]
    r += 1
    for idx, (puesto, turno_default) in enumerate(staff, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, None, fill=inp_fill); sdc(ws, r, 3, puesto)
        for c in range(4, 11):
            sdc(ws, r, c, turno_default, fill=inp_fill, align=c_align)
            dv.add(ws.cell(row=r, column=c))
        r += 1
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "M=Mañana, T=Tarde, P=Partido, L=Libre, V=Vacaciones", font=brn_font)
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "plantilla-turnos-brigada.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 3. CHECKLISTS (6)
# ═══════════════════════════════════════════════════════════

def gen_checklist_legal():
    wb = Workbook(); ws = wb.active
    items = [
        ("Empresa", "Decidir forma jurídica (SL, autónomo, cooperativa)", "Socio/Gestor", 500),
        ("Empresa", "Constituir sociedad ante notario", "Socio/Gestor", 600),
        ("Empresa", "Inscripción en Registro Mercantil", "Gestor", 200),
        ("Empresa", "Obtener CIF provisional", "Gestor", 0),
        ("Empresa", "Alta IAE epígrafe 671.4 o similar", "Gestor", 0),
        ("Empresa", "Alta censal en Hacienda (modelo 036/037)", "Gestor", 0),
        ("Empresa", "Inscripción Seguridad Social empresa", "Gestor", 0),
        ("Empresa", "Abrir cuenta bancaria empresa", "Socio", 0),
        ("Licencias", "Licencia de actividad / declaración responsable", "Arquitecto", 2500),
        ("Licencias", "Proyecto técnico y visado colegial", "Arquitecto", 3500),
        ("Licencias", "Licencia de obras (si aplica)", "Arquitecto", 2000),
        ("Licencias", "Licencia de terraza", "Socio", 1500),
        ("Licencias", "Licencia de música / SGAE / AGEDI", "Socio", 500),
        ("Anisakis", "Compra congelador certificado -20°C (obligatorio para sushi/sashimi)", "Socio", 4500),
        ("Anisakis", "Documento protocolo congelación preventiva pescado crudo", "Jefe cocina", 0),
        ("Anisakis", "Registros de congelación por lote de pescado (archivo 12 meses)", "Itamae", 0),
        ("Anisakis", "Información al cliente en carta sobre congelación preventiva", "Socio", 0),
        ("Anisakis", "Certificados de proveedores pescado de acuicultura (excepción)", "Jefe cocina", 0),
        ("Importación", "Registro EORI (operador comercio exterior) si aplica", "Gestor", 0),
        ("Importación", "Certificado SOIVRE para alimentos importados extra-UE", "Gestor", 300),
        ("Importación", "Etiquetado en español de productos japoneses importados", "Proveedor", 200),
        ("Importación", "Impuestos especiales bebidas alcohólicas (IIEE) sake/whisky", "Gestor", 0),
        ("Sanidad", "Registro General Sanitario (RGSEAA)", "Gestor", 0),
        ("Sanidad", "Plan APPCC documentado (adaptado pescado crudo + arroz sushi)", "Consultor", 2000),
        ("Sanidad", "Formación manipulador alimentos (equipo — especial pescado crudo)", "Consultor", 500),
        ("Sanidad", "Contrato empresa DDD (desratización)", "Socio", 600),
        ("Sanidad", "Control de aguas", "Socio", 200),
        ("Sanidad", "Plan de alérgenos (14 alérgenos + soja, sésamo)", "Jefe cocina", 0),
        ("Seguros", "Seguro responsabilidad civil MÍNIMO 600K€ (riesgo sanitario pescado crudo)", "Socio", 1200),
        ("Seguros", "Seguro de local (incendio, robo, daños)", "Socio", 700),
        ("Seguros", "Seguro de accidentes convenio", "Gestor", 400),
        ("Laboral", "Contratos de trabajo según convenio", "Gestor", 0),
        ("Laboral", "Contrato específico del itamae (cláusulas de propiedad intelectual recetas)", "Gestor", 0),
        ("Laboral", "Prevención de riesgos laborales", "SPA", 500),
        ("Laboral", "Calendario laboral y cuadrante turnos", "Encargado", 0),
        ("Protección datos", "Registro LOPD / RGPD", "Gestor", 300),
        ("Protección datos", "Política de cookies web", "Socio", 0),
        ("Protección datos", "Videovigilancia (si aplica): cartelería + registro", "Socio", 100),
        ("Otros", "Libro de reclamaciones", "Socio", 30),
        ("Otros", "Cartel de horarios en puerta", "Socio", 0),
        ("Otros", "Aforo máximo visible", "Socio", 0),
        ("Otros", "Señalética salida emergencia + extintores", "Arquitecto", 600),
    ]
    checklist_ws(ws, "Checklist Legal — Restaurante Japonés España", items)
    path = os.path.join(OUTPUT_DIR, "checklist-legal.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_equipamiento():
    wb = Workbook(); ws = wb.active
    items = [
        ("Cocina JP", "Suihanki (arrocera japonesa profesional 5-10L)", "Itamae", 3200),
        ("Cocina JP", "Vitrina refrigerada de sashimi (barra de sushi)", "Itamae", 6500),
        ("Cocina JP", "Ramen cooker / pasta cooker profesional 8-12 cestas", "Jefe cocina", 5500),
        ("Cocina JP", "Ollas de caldo 50-100L tonkotsu/shoyu (2 uds)", "Jefe cocina", 2500),
        ("Cocina JP", "Parrilla robata de carbón binchotan", "Jefe cocina", 6500),
        ("Cocina JP", "Plancha teppanyaki profesional", "Jefe cocina", 4500),
        ("Cocina JP", "Set cuchillos japoneses (yanagiba, deba, usuba, santoku)", "Itamae", 2800),
        ("Cocina JP", "Tabla corte madera Hiba / sintética sushi (2 uds)", "Itamae", 500),
        ("Cocina JP", "Ohitsu (recipiente cedro) + abanico uchiwa", "Itamae", 800),
        ("Cocina JP", "Maquina sake caliente (tokkuri warmer)", "Encargado", 600),
        ("Cocina JP", "Vitrina refrigerada sake y whisky japonés barra", "Encargado", 4500),
        ("Cocción", "Cocina gas 6 fuegos + horno", "Jefe cocina", 5000),
        ("Cocción", "Horno mixto (Rational / Unox)", "Jefe cocina", 12000),
        ("Cocción", "Wok profesional alta potencia (yakisoba, saltados)", "Jefe cocina", 2500),
        ("Cocción", "Freidora doble cuba (tempura, karaage, agemono)", "Jefe cocina", 2500),
        ("Cocción", "Salamandra", "Jefe cocina", 1200),
        ("Frío", "Cámara frigorífica 2 puertas (pescado fresco)", "Jefe cocina", 4000),
        ("Frío", "Cámara congelación CERTIFICADA -20°C (anisakis obligatorio)", "Itamae", 4500),
        ("Frío", "Mesa refrigerada bajo barra sushi", "Itamae", 3000),
        ("Frío", "Botellero refrigerado barra sake", "Encargado", 2000),
        ("Preparación", "Mesa de trabajo acero inox (2-3 uds)", "Jefe cocina", 2000),
        ("Preparación", "Robot de cocina / cutter", "Jefe cocina", 1500),
        ("Preparación", "Escamadora eléctrica de pescado", "Jefe cocina", 400),
        ("Lavado", "Tren de lavado / lavavajillas industrial", "Jefe cocina", 5500),
        ("Lavado", "Fregadero doble seno inox", "Jefe cocina", 800),
        ("Extracción", "Campana extractora REFORZADA + filtros (robata + teppan + tempura)", "Instalador", 10000),
        ("Extracción", "Salida de humos reglamentaria (sobredimensionada)", "Instalador", 6000),
        ("Menaje", "Ollas grandes para ramen, caldos, arroz", "Jefe cocina", 1200),
        ("Menaje", "Sartén wok adicional (yakisoba)", "Jefe cocina", 300),
        ("Menaje", "Cuchillos adicionales y afiladores", "Jefe cocina", 800),
        ("Menaje", "Gastronormas GN 1/1, 1/2, 1/3", "Jefe cocina", 600),
        ("Menaje", "Tablas de corte por colores (separación crudo/cocinado)", "Jefe cocina", 150),
        ("Almacenamiento", "Estanterías inox economato seco (arroz, algas, salsas)", "Jefe cocina", 800),
        ("Almacenamiento", "Contenedores herméticos para arroz Koshihikari", "Jefe cocina", 300),
        ("Barra", "Utensilios coctelería japonesa (shaker, jiggers, stirrer)", "Encargado", 600),
        ("Barra", "Máquina de hielo (alto consumo por highballs)", "Encargado", 2200),
    ]
    checklist_ws(ws, "Checklist Equipamiento Cocina Japonesa", items)
    path = os.path.join(OUTPUT_DIR, "checklist-equipamiento-cocina-japonesa.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_appcc():
    wb = Workbook(); ws = wb.active
    items = [
        ("Prerrequisitos", "Plan de limpieza y desinfección documentado", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan DDD contratado (empresa externa)", "Socio", 600),
        ("Prerrequisitos", "Plan de control de agua potable", "Socio", 200),
        ("Prerrequisitos", "Plan de formación del personal (especial pescado crudo)", "Jefe cocina", 500),
        ("Prerrequisitos", "Plan de gestión de residuos (incluye aceite tempura)", "Socio", 0),
        ("Prerrequisitos", "Plan de mantenimiento de instalaciones", "Socio", 0),
        ("Prerrequisitos", "Plan de trazabilidad (registro proveedores + lotes)", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan de control de alérgenos (14 + soja, sésamo)", "Jefe cocina", 0),
        ("Anisakis (PCC)", "CONGELACIÓN PREVIA obligatoria -20°C/24h todo pescado crudo", "Itamae", 0),
        ("Anisakis (PCC)", "Congelador certificado con registro de temperatura automático", "Itamae", 4500),
        ("Anisakis (PCC)", "Registro de congelación por lote (hora entrada/salida/temperatura)", "Itamae", 0),
        ("Anisakis (PCC)", "Archivo de registros mínimo 12 meses", "Jefe cocina", 0),
        ("Anisakis (PCC)", "Certificado proveedor para pescado de acuicultura (excepción)", "Jefe cocina", 0),
        ("Anisakis (PCC)", "Información al cliente en carta: 'pescado congelado preventivamente'", "Socio", 0),
        ("Arroz sushi (PCC)", "Temperatura shari entre 20-25°C (no refrigerar, no calentar)", "Itamae", 0),
        ("Arroz sushi (PCC)", "Máxima permanencia 4 horas en ohitsu", "Itamae", 0),
        ("Arroz sushi (PCC)", "Registro de temperatura cada 2 horas", "Sushi chef", 0),
        ("Arroz sushi (PCC)", "Acidificación con sushi-zu (vinagre + azúcar + sal) verificada", "Itamae", 0),
        ("Temperaturas", "Termómetro sonda calibrado", "Jefe cocina", 50),
        ("Temperaturas", "Registro diario temperaturas cámaras", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas servicio (caliente >65°C)", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas recepción mercancía pescado", "Jefe cocina", 0),
        ("Temperaturas", "Control cadena de frío (0-4°C refrigerado, -20°C congelado)", "Jefe cocina", 0),
        ("Temperaturas", "Control temperatura caldos ramen (enfriamiento rápido <4°C)", "Ramen cook", 0),
        ("Pescado crudo", "Descongelación SIEMPRE en cámara 0-4°C (nunca a temp. ambiente)", "Itamae", 0),
        ("Pescado crudo", "Trazabilidad completa del pescado (proveedor, lonja, fecha)", "Jefe cocina", 0),
        ("Pescado crudo", "Vitrina de sashimi a 0-4°C con control continuo", "Itamae", 0),
        ("Pescado crudo", "Rotación pescado expuesto en vitrina cada 2-3 horas", "Sushi chef", 0),
        ("Pescado crudo", "Cuchillos y tablas exclusivos para pescado crudo", "Itamae", 0),
        ("Prod. Importados", "Etiquetado en español (sake, soja, mirin, algas, panko)", "Jefe cocina", 0),
        ("Prod. Importados", "Trazabilidad de origen productos japoneses", "Jefe cocina", 0),
        ("Prod. Importados", "Almacenamiento separado productos secos (arroz, algas)", "Cocinero", 0),
        ("Prod. Importados", "Registro de lotes de sake/whisky importado", "Encargado", 0),
        ("Tempura", "Filtrado aceite diario, registro de cambio por polaridad", "Tempura cook", 0),
        ("Tempura", "Temperatura aceite controlada 170-180°C", "Tempura cook", 0),
        ("Higiene", "Lavamanos no manual en cocina con jabón y papel", "Socio", 300),
        ("Higiene", "Vestuarios con taquillas para personal", "Socio", 500),
        ("Higiene", "Uniformes cocina (itamae uniforme blanco tradicional + equipo)", "Socio", 800),
        ("Higiene", "Protocolo lavado de manos (cartelería)", "Jefe cocina", 0),
        ("Higiene", "Contenedores de basura con tapa y pedal", "Socio", 200),
        ("Almacenamiento", "Sistema FIFO/FEFO en cámaras y economato", "Cocinero", 0),
        ("Almacenamiento", "Separación crudo/cocinado en cámaras (especialmente pescado)", "Cocinero", 0),
        ("Documentación", "Plan APPCC escrito y disponible en cocina", "Socio", 2000),
        ("Documentación", "Registros de control actualizados diariamente", "Jefe cocina", 0),
        ("Documentación", "Certificados manipulador alimentos de todo el personal", "Socio", 0),
        ("Documentación", "Contrato DDD visible y registros de actuaciones", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist APPCC — Restaurante Japonés (Sashimi + Sushi + Anisakis)", items)
    path = os.path.join(OUTPUT_DIR, "checklist-appcc.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_sala():
    wb = Workbook(); ws = wb.active
    items = [
        ("Distribución", "Plano de sala con posición de mesas y barra sushi definitivo", "Socio/Arquitecto", 0),
        ("Distribución", "60 plazas distribuidas (barra sushi 8-12 + mesas)", "Socio", 0),
        ("Distribución", "Distancia entre mesas: mínimo 70-90 cm (confort japonés)", "Arquitecto", 0),
        ("Distribución", "Flujo de servicio definido (entrada → sala → barra → aseos)", "Socio", 0),
        ("Distribución", "Zona de espera / recepción con noren", "Socio", 800),
        ("Distribución", "Accesibilidad PMR (rampa, aseo adaptado)", "Arquitecto", 2000),
        ("Barra Sushi", "Barra de sushi profesional madera + vitrina sashimi delantera", "Interiorista", 10000),
        ("Barra Sushi", "Taburetes altos cómodos (8-12 plazas)", "Socio", 3500),
        ("Barra Sushi", "Iluminación dirigida al itamae (teatro)", "Electricista", 1200),
        ("Barra Sushi", "Carta de sushi/omakase enmarcada", "Diseñador", 400),
        ("Barra Sake", "Vitrina iluminada trasera para sake y whisky japonés", "Interiorista", 5000),
        ("Barra Sake", "Iluminación ambiental barra (retroiluminación botellas)", "Electricista", 1000),
        ("Decoración JP", "Revestimientos de madera natural en paredes de acento", "Interiorista", 8000),
        ("Decoración JP", "Noren (cortina entrada) de tela japonesa", "Interiorista", 400),
        ("Decoración JP", "Jardín zen / karesansui (piedras, grava, bambú)", "Interiorista", 2500),
        ("Decoración JP", "Paleta minimalista: madera, blanco roto, negro mate, gris piedra", "Interiorista", 0),
        ("Decoración JP", "Vajilla cerámica japonesa artesanal (asimétrica, wabi-sabi)", "Socio", 4500),
        ("Decoración JP", "Bambú decorativo (pantallas, lámparas, detalles)", "Socio", 1200),
        ("Decoración JP", "Detalles piedra natural (pared acento, barra)", "Interiorista", 2500),
        ("Iluminación", "Iluminación general cálida 2400-2700K (tenue)", "Electricista", 2500),
        ("Iluminación", "Lámparas washi (papel de arroz) sobre mesas", "Socio", 2500),
        ("Iluminación", "Iluminación regulable por zonas (dimmer)", "Electricista", 600),
        ("Acústica", "Paneles absorbentes para ambiente silencioso (cultura japonesa)", "Interiorista", 2000),
        ("Acústica", "Sistema de música (jazz japonés, ambient, volumen bajo)", "Socio", 800),
        ("Mobiliario", "Mesas de madera natural maciza (roble/cedro/haya)", "Socio", 14000),
        ("Mobiliario", "Sillas tapizadas en tonos neutros", "Socio", 8000),
        ("Mobiliario", "Zona tatami privada (si aplica)", "Interiorista", 3500),
        ("Terraza", "Mobiliario exterior en madera/bambú natural", "Socio", 5000),
        ("Aseos", "Aseos diferenciados y aseo PMR", "Arquitecto", 3500),
        ("Aseos", "Decoración japonesa en aseos (madera, detalles zen)", "Socio", 600),
        ("Señalética", "Señalética interior con identidad japonesa minimalista (kanji)", "Diseñador", 400),
        ("Señalética", "Rótulo exterior / fachada con identidad japonesa sutil", "Socio", 3000),
    ]
    checklist_ws(ws, "Checklist Diseño Sala Japonesa + Barra Sushi + Barra Sake", items)
    path = os.path.join(OUTPUT_DIR, "checklist-diseno-sala-japonesa.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_contratacion():
    wb = Workbook(); ws = wb.active
    items = [
        ("Perfiles", "Definir organigrama completo (cocina + sala)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Itamae / Jefe cocina (5+ años sushi, idealmente Japón)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Sushi chef / Segundo itamae", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Ramen cook (caldos, montaje bowls)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Robata cook (parrilla binchotan, yakitori)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Tempura / Hot kitchen cook", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Cocinero izakaya (gyoza, karaage, agemono)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Preparación pescado / escamado", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Ayudante / office", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Encargado sala / Sommelier sake", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Camarero (con conocimiento sake)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Runner", "Socio", 0),
        ("Selección", "Publicar ofertas (InfoJobs, Hostelwork, LinkedIn)", "Socio", 300),
        ("Selección", "Búsqueda itamae en comunidades japonesas España (Madrid, Barcelona)", "Socio", 0),
        ("Selección", "Búsqueda itamae internacional (si procede — visado)", "Socio", 2000),
        ("Selección", "Búsqueda sushi chef con formación específica", "Socio", 0),
        ("Selección", "Búsqueda sommelier sake certificado (WSET Sake / Sake School)", "Socio", 0),
        ("Selección", "Selección resto equipo cocina y sala", "Jefe cocina", 0),
        ("Selección", "Pruebas prácticas cocina (elaborar nigiri, ramen caldo, yakitori)", "Jefe cocina", 0),
        ("Documentación", "Contratos según convenio hostelería CCAA", "Gestor", 0),
        ("Documentación", "Contrato específico del itamae (cláusulas premium)", "Gestor", 200),
        ("Documentación", "Alta Seguridad Social", "Gestor", 0),
        ("Documentación", "Certificado manipulador alimentos (especial pescado crudo)", "Empleado", 0),
        ("Documentación", "Prevención riesgos laborales (formación)", "SPA", 500),
        ("Onboarding", "Manual de bienvenida / handbook del restaurante", "Socio", 0),
        ("Onboarding", "Formación en carta japonesa y alérgenos (soja, sésamo, gluten)", "Jefe cocina", 0),
        ("Onboarding", "Formación en sake y whisky japonés (todo equipo sala)", "Sommelier", 0),
        ("Onboarding", "Formación en TPV y sistema de pedidos", "Encargado", 0),
        ("Onboarding", "Formación APPCC (especial pescado crudo + anisakis + arroz sushi)", "Jefe cocina", 0),
        ("Onboarding", "Formación protocolo japonés sala (irasshaimase, servicio)", "Encargado", 0),
        ("Onboarding", "Ensayos de servicio (soft opening)", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist Contratación — Restaurante Japonés", items)
    path = os.path.join(OUTPUT_DIR, "checklist-contratacion.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_marketing():
    wb = Workbook(); ws = wb.active
    items = [
        ("Branding", "Nombre del restaurante definitivo (verificar disponibilidad)", "Socio", 0),
        ("Branding", "Logo profesional con identidad japonesa minimalista", "Diseñador", 700),
        ("Branding", "Manual de marca (paleta neutra, tipografía, estilo wabi-sabi)", "Diseñador", 400),
        ("Branding", "Diseño de carta / menú físico con estética japonesa", "Diseñador", 700),
        ("Branding", "Carta específica de sakes y whiskies japoneses", "Diseñador", 400),
        ("Branding", "Packaging delivery con marca (cajas sushi premium)", "Diseñador", 500),
        ("Digital", "Web propia con menú, carta sakes, fotos y reservas", "Dev/Agencia", 2500),
        ("Digital", "Sistema de reservas online con depósito (omakase)", "Dev/Agencia", 800),
        ("Digital", "Google My Business: ficha completa y verificada", "Socio", 0),
        ("Digital", "Instagram: perfil profesional con 12+ publicaciones premium", "Community", 0),
        ("Digital", "TikTok: perfil activo con vídeos del itamae, ramen, robata", "Community", 0),
        ("Digital", "TheFork / TripAdvisor: perfiles completos con fotos", "Socio", 0),
        ("Digital", "Google Ads local configurado", "Agencia", 700),
        ("Digital", "Instagram/Facebook Ads configurados", "Agencia", 500),
        ("Contenido", "Sesión fotográfica profesional (sushi, ramen, robata, decoración)", "Fotógrafo", 600),
        ("Contenido", "Vídeo del itamae preparando nigiri (30-60 seg para TikTok/Reels)", "Videógrafo", 700),
        ("Contenido", "Vídeo ramen: montaje bowl con caldo y toppings", "Videógrafo", 500),
        ("Contenido", "Calendario de publicaciones primer mes", "Community", 0),
        ("Eventos JP", "Plan Hanami (marzo-abril): menú sakura + decoración cerezo", "Socio", 600),
        ("Eventos JP", "Plan Día del Sushi Internacional (18 junio): omakase degustación", "Socio", 500),
        ("Eventos JP", "Plan Año Nuevo japonés (Oshogatsu): menú osechi ryori", "Socio", 500),
        ("Eventos JP", "Plan Día cerveza japonesa (1 agosto): promos izakaya", "Socio", 300),
        ("Eventos JP", "Eventos cata de sake / whisky japonés", "Sommelier", 400),
        ("Pre-apertura", "Evento soft opening (prensa gastronómica, influencers, críticos)", "Socio", 800),
        ("Pre-apertura", "Invitación a prensa especializada (guías Michelin, Repsol)", "Socio", 300),
        ("Pre-apertura", "Flyers / cartelería barrio", "Diseñador", 300),
        ("Pre-apertura", "Colaboraciones con negocios vecinos premium", "Socio", 0),
        ("Delivery", "Alta en Glovo / Uber Eats / Just Eat", "Socio", 0),
        ("Delivery", "Fotos profesionales de sushi boxes para plataformas", "Fotógrafo", 0),
        ("Delivery", "Menú delivery optimizado (sushi boxes, bowls, katsu, gyoza)", "Jefe cocina", 0),
        ("Fidelización", "Programa fidelización (tarjeta, app o CRM — omakase VIP)", "Socio", 500),
        ("Fidelización", "Estrategia de reseñas (pedir valoraciones a clientes)", "Encargado", 0),
    ]
    checklist_ws(ws, "Checklist Marketing Pre-Apertura — Restaurante Japonés", items)
    path = os.path.join(OUTPUT_DIR, "checklist-marketing-preapertura.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 4. DOCUMENT MODELS (2)
# ═══════════════════════════════════════════════════════════

def gen_business_plan():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Business Plan"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Japonés — 60 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("[Ciudad, Mes Año]\nPlantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Resumen Ejecutivo", [
            "Descripción del proyecto: restaurante japonés auténtico de 60 plazas en [ciudad].",
            "Modelo: [sushi-ya / ramen-ya / izakaya / omakase / robatayaki / mixto / dark kitchen].",
            "Inversión total requerida: [250.000-500.000€].",
            "Facturación prevista año 1: [700.000-1.100.000€].",
            "Break-even estimado: mes [10-16].",
            "Equipo fundador: [nombres y experiencia].",
            "Diferencial: itamae cualificado + pescado sashimi-grade + barra de sake/whisky japonés 30-50 referencias.",
        ]),
        ("2. El Concepto", [
            "Tipo de cocina: sushi-ya / ramen-ya / izakaya / omakase / robata / mixto.",
            "Ticket medio objetivo: [35-55€] casual / [80-150€] omakase, sin bebidas.",
            "Propuesta de valor: itamae profesional, pescado diario sashimi-grade, sake premium.",
            "Público objetivo: 28-55 años, urbano, foodie, alto poder adquisitivo, viajero.",
            "Posicionamiento vs competencia: auténtico premium vs japonés genérico buffet.",
        ]),
        ("3. Análisis de Mercado", [
            "Restaurantes japoneses en España: +2.400, crecimiento 35% en 5 años.",
            "Menos del 30% tienen itamae cualificado — oportunidad premium clara.",
            "Tendencia sake premium y whisky japonés: crecimiento acelerado.",
            "Omakase de alta gama con listas de espera en Madrid y Barcelona.",
            "Competencia directa en zona: [listar 3-5 competidores con ticket medio].",
            "Gap de mercado identificado: [qué falta en la zona].",
        ]),
        ("4. Plan Operativo", [
            "Ubicación: [dirección, m², alquiler mensual].",
            "Distribución: 60 plazas interior + barra sushi 8-12 plazas + [X] plazas terraza.",
            "Horario: [ej. 13:00-16:00 y 20:00-00:00, lunes cerrado].",
            "Equipo cocina: [8-12] personas (incluye itamae, sushi chef, ramen, robata).",
            "Equipo sala: [5-8] personas (incluye sommelier sake).",
            "Proveedores clave: [pescado sashimi-grade, distribuidores japoneses, sake].",
        ]),
        ("5. Plan Financiero", [
            "Inversión inicial desglosada: ver plantilla Excel adjunta.",
            "P&L mensual con 3 escenarios: ver plantilla Excel adjunta.",
            "Cash flow 12 meses: ver plantilla Excel adjunta.",
            "Break-even: [mes estimado].",
            "ROI estimado: [X%] en año [2-3].",
        ]),
        ("6. Plan de Marketing", [
            "Estrategia pre-apertura: [acciones 3 meses antes de abrir].",
            "Canales digitales: Google My Business, Instagram, TikTok, TheFork.",
            "Eventos culturales: Hanami, Día del Sushi Internacional, Año Nuevo japonés.",
            "Presupuesto marketing mensual: [800-2.000€/mes].",
            "Estrategia de delivery: sushi boxes, bowls, katsu y gyoza.",
        ]),
        ("7. Equipo Fundador", [
            "[Nombre]: [experiencia en hostelería/cocina japonesa, rol en el proyecto].",
            "[Nombre]: [experiencia complementaria, rol en el proyecto].",
            "Itamae contratado: [experiencia, formación en Japón, años en el oficio].",
            "Asesores externos: [gestor, arquitecto, consultor gastronómico].",
        ]),
        ("8. Necesidades de Financiación", [
            "Inversión total: [X€].",
            "Fondos propios: [X€] ([X%]).",
            "Financiación solicitada: [X€] ([X%]).",
            "Destino fondos: obra [X€], equipamiento japonés [X€], barra sake [X€], maniobra [X€].",
            "Plan de amortización propuesto: [X años, cuota mensual X€].",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "business-plan-modelo.docx")
    doc.save(path); print(f"✓ {path}")

def gen_manual_operaciones():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual de Operaciones"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Japonés — 60 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Plantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Apertura del Restaurante", [
            "Hora de llegada del equipo: [3 horas antes del servicio].",
            "Itamae: recibir pescado fresco del día, verificar calidad sashimi-grade, revisar registros de congelación preventiva del pescado del día anterior para sashimi.",
            "Sushi chef: cocer arroz Koshihikari en suihanki, preparar sushi-zu (vinagre + azúcar + sal), aderezar shari, colocar en ohitsu y ventilar con uchiwa.",
            "Ramen cook: verificar caldos (tonkotsu con 12h de cocción continua), preparar toppings del día (chashu, ajitama, menma, negi), montar mise en place de tare.",
            "Robata cook: encender parrilla binchotan (40 min para alcanzar temperatura), preparar brochetas, mise en place de tare y sal.",
            "Tempura cook: filtrar y verificar aceite, preparar masa al momento según pedidos, mise en place de verduras y gambas.",
            "Cocineros: preparar guisos, salsas (ponzu, teriyaki, tare), mise en place hot kitchen.",
            "Verificar temperaturas de cámaras y REGISTRAR. Verificar congelador certificado -20°C para anisakis.",
            "Barra sake: verificar stock sakes y whiskies, enfriar sakes fríos, preparar tokkuri y ochoko, montar barra.",
            "Preparar sala: colocar mesas, noren, cartas, palillos, servilletas.",
            "Verificar limpieza de aseos.",
            "Encender TPV, música ambiente (jazz japonés a volumen bajo), iluminación tenue.",
            "Briefing de equipo: sashimis del día según pescado recibido, sake destacado, reservas omakase, alérgenos.",
        ]),
        ("2. Servicio de Mediodía", [
            "Horario: 13:00-16:00.",
            "Mediodía: público de business lunch, omakase corto, ramen para profesionales.",
            "Ofrecer menú del día japonés (si aplica) + carta completa + omakase almuerzo.",
            "Velocidad de servicio: sashimi en mesa en 8-10 min, ramen en 10-12 min, principales en 15 min.",
            "Itamae: ritmo constante en barra de sushi. Nigiri al momento.",
            "Ramen cook: montaje rápido de bowls, caldo siempre hirviendo.",
            "Gestión de turnos de mesa si hay espera (notificar tiempo estimado).",
            "Cobro y despedida: sugerir postre (mochi, dorayaki) y té verde final.",
        ]),
        ("3. Servicio de Cena", [
            "Horario: 20:00-00:00.",
            "Carta completa + omakase + carta destacada de sakes y whiskies japoneses.",
            "Iluminación más tenue, música ambient, volumen bajo.",
            "Sugerir omakase (si disponible), sashimi para compartir, vuelos de sake.",
            "Itamae: servicio de omakase en barra (si aplica). Interacción con cliente.",
            "Último pedido de cocina: [23:30].",
            "Barra robata activa durante toda la cena — plato emblema.",
        ]),
        ("4. Estación del Itamae / Barra de Sushi: Protocolo Diario", [
            "Verificar que TODO el pescado para sashimi fue previamente congelado -20°C/24h (registro OBLIGATORIO — anisakis).",
            "Descongelar pescado en cámara (0-4°C) la noche anterior. NUNCA a temperatura ambiente ni con agua.",
            "Cocer arroz Koshihikari en suihanki. Proporción aprox 1:1.1 agua.",
            "Preparar sushi-zu (vinagre arroz + azúcar + sal). Aderezar arroz caliente con movimientos suaves con shamoji.",
            "Transferir shari a ohitsu de cedro. Ventilar con uchiwa para enfriarlo a 20-25°C rápidamente.",
            "Montar vitrina de sashimi con pescado del día, decorado con shiso, daikon juliana, wasabi, gari.",
            "Cortar pescado al momento del pedido: dados/rectángulos precisos para sashimi, láminas finas para nigiri.",
            "Montar nigiri a mano: shari del tamaño exacto, wasabi en cantidad precisa, pescado encima con presión delicada.",
            "Limpieza continua de la estación: cuchillos yanagiba, tabla, superficies. Protocolo de higiene estricto por pescado crudo.",
            "Controlar temperatura del shari en ohitsu cada 2 horas (registro APPCC).",
            "Vida útil máxima del shari: 4 horas. Desechar sobrante y preparar nuevo lote si servicio largo.",
        ]),
        ("5. Estación de Ramen: Protocolo Diario", [
            "Verificar caldos (tonkotsu de 12-24h, shoyu, miso, shio) — niveles, temperatura >75°C servicio.",
            "Preparar toppings del día: chashu laminado, ajitama (huevo marinado en tare), menma, negi picado, nori cortado.",
            "Mise en place de tare (salsa base): shoyu tare, miso tare, shio tare en botes dedicados.",
            "Fideos: cocer al momento en ramen cooker (1-2 min según grosor).",
            "Montaje del bowl: tare al fondo → caldo caliente → fideos escurridos → toppings en disposición estética.",
            "Tiempo máximo montaje: 90 segundos (el ramen se debe servir a 75°C+).",
            "Limpieza de cestas y recambio frecuente del agua del ramen cooker.",
            "Enfriamiento rápido del sobrante de caldo al cierre (<4°C en <2h).",
        ]),
        ("6. Barra de Sake y Whisky: Protocolo", [
            "Verificar stock de sakes (30-50 refs) y whiskies japoneses (6-10 refs).",
            "Enfriar sakes fríos (Junmai Ginjo, Daiginjo, Nigori) a 8-12°C.",
            "Preparar sakes calientes (Junmai tradicional) con tokkuri warmer, temperatura 40-50°C (nunca hervir).",
            "Preparar hielo abundante (highballs consumen mucho).",
            "Verificar stock garnish: limón yuzu, shiso, cáscara de limón, cáscara de mandarina.",
            "Sugerir vuelos de sake (3 copas: junmai, ginjo, daiginjo) por 18-25€.",
            "Explicar grados de pulido de arroz (60% Ginjo, 50% Daiginjo) y diferencias.",
            "Servir con protocolo japonés: presentar botella al cliente, servir sin tocar copa con botella, ambas manos si cliente importante.",
            "Registrar consumo de botellas para control de stock (whisky japonés en alocación).",
            "Al cierre: limpiar barra, guardar sakes abiertos según su categoría, cerrar botellas de whisky.",
        ]),
        ("7. Cierre del Restaurante", [
            "Cocina: apagar ramen cooker, robata (dejar enfriar binchotan de forma segura), hornos, freidora, fogones.",
            "Itamae: DESECHAR todo el shari sobrante, limpiar vitrina de sashimi, guardar pescado no utilizado en cámara con trazabilidad.",
            "Verificar y registrar temperaturas de cámaras y congelador -20°C.",
            "Guardar salsas (ponzu, teriyaki, tare) en cámara con fecha y hora.",
            "Ramen: enfriar caldos rápidamente (<4°C en <2h), almacenar en cámara para día siguiente.",
            "Sala: limpiar mesas, barrer, recoger terraza, guardar noren.",
            "Barra: cerrar caja, cuadrar TPV, limpiar barra, guardar whisky bajo llave (alocación).",
            "Sacar basuras a contenedores (separado: aceite de tempura para recogida).",
            "Cerrar puertas, alarma, luces.",
        ]),
        ("8. Protocolo Anisakis (PCC Crítico)", [
            "La normativa RD 1420/2006 obliga a congelar preventivamente TODO pescado destinado a consumo crudo.",
            "Temperatura: -20°C durante mínimo 24 horas (o -35°C durante 15 horas).",
            "Equipamiento: congelador certificado que alcance y mantenga -20°C real (no un congelador doméstico).",
            "Registro OBLIGATORIO por lote: especie, proveedor, fecha entrada congelador, fecha salida, temperatura.",
            "Archivo de registros: mínimo 12 meses disponibles para inspección sanitaria.",
            "Excepción: pescado de acuicultura con certificado del proveedor garantizando ausencia de riesgo.",
            "Información al cliente: obligación de informar en carta (nota pie de página).",
            "NUNCA servir pescado para sashimi/sushi sin haber pasado por el protocolo — responsabilidad civil y penal.",
            "En caso de inspección: mostrar congelador, registros y etiquetas de lotes activos.",
            "Un caso de anisakiasis positivo puede cerrar tu restaurante y generar responsabilidad civil millonaria.",
        ]),
        ("9. Gestión de Delivery Japonés", [
            "Zona de packaging separada del pase de sala y de la barra de sushi.",
            "Sushi boxes: producto estrella. Caja negra mate con compartimentos, salsa soja en bote pequeño, wasabi, gari, palillos.",
            "Chirashi bowls: envases con tapa hermética, pescado previamente pesado.",
            "Poke bowls: base arroz + proteína + verduras + salsa separada en cápsula.",
            "Katsudon / Oyakodon: envases con tapa, arroz abajo, katsu arriba con salsa.",
            "Gyoza y karaage: envases con papel absorbente para mantener crujientes.",
            "Ramen delivery: SOLO si entrega < 15 min. Caldo separado del fideo (cliente ensambla).",
            "NO incluir en delivery: tempura (se ablanda), sashimi puro sin hielo (riesgo sanitario).",
            "Verificar cada pedido antes de entregar al rider.",
            "Incluir palillos, salsa soja extra, carta del restaurante para futuros pedidos.",
            "Tiempo máximo de preparación: 15 minutos.",
            "Responder a reseñas negativas en plataformas en <24h.",
        ]),
        ("10. Gestión de Alérgenos en Cocina Japonesa", [
            "14 alérgenos de declaración obligatoria (estándar UE).",
            "Alérgenos frecuentes en cocina japonesa: pescado (sashimi, sushi), crustáceos (tempura de gambas), moluscos (vieira, pulpo), soja (omnipresente en salsas, marinados, miso), sésamo (aceite, semillas, pastas), gluten (fideos udon, tempura, panko, salsa soja tradicional con trigo), huevo (tamago, tempura, ajitama).",
            "Carta con indicación de alérgenos por plato.",
            "La salsa soja tradicional contiene TRIGO: informar SIEMPRE a celíacos. Ofrecer tamari (soja sin trigo).",
            "El panko contiene gluten: alternativa sin gluten para katsu.",
            "Muchos platos japoneses son sin gluten si se sustituye shoyu por tamari: ventaja competitiva.",
            "Protocolo de cocina para platos sin alérgenos: utensilios separados, zona limpia.",
            "Ante duda: confirmar SIEMPRE con itamae o jefe de cocina.",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "manual-operaciones-japones.docx")
    doc.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 5. PDF placeholder
# ═══════════════════════════════════════════════════════════
def gen_guide_pdf():
    """Generate a minimal placeholder PDF using the DOCX content."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        path = os.path.join(OUTPUT_DIR, "guia-restaurante-japones.pdf")
        c = canvas.Canvas(path, pagesize=A4)
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(297, 600, "Como Montar un Restaurante Japones")
        c.setFont("Helvetica", 16)
        c.drawCentredString(297, 560, "60 Plazas de Aforo - Guia Espana 2026")
        c.setFont("Helvetica", 12)
        c.drawCentredString(297, 520, "Chef John Guerrero - AI Chef Pro - aichef.pro")
        c.drawCentredString(297, 480, "20 capitulos - 60+ paginas - 8 plantillas - 6 checklists")
        c.save()
        print(f"✓ {path}")
    except ImportError:
        import shutil
        src = os.path.join(OUTPUT_DIR, "guia-restaurante-japones.docx")
        dst = os.path.join(OUTPUT_DIR, "guia-restaurante-japones.pdf")
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f"✓ {dst} (placeholder from DOCX)")


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n🇯🇵 Generando archivos: Restaurante Japonés 60 Plazas\n")
    gen_guide_docx()
    gen_guide_pdf()
    gen_plan_financiero()
    gen_calculadora_capex()
    gen_pl_mensual()
    gen_cash_flow()
    gen_escandallo()
    gen_menu_engineering()
    gen_calculadora_ticket()
    gen_cronograma()
    gen_plantilla_turnos()
    gen_checklist_legal()
    gen_checklist_equipamiento()
    gen_checklist_appcc()
    gen_checklist_sala()
    gen_checklist_contratacion()
    gen_checklist_marketing()
    gen_business_plan()
    gen_manual_operaciones()
    print(f"\n✅ {len(os.listdir(OUTPUT_DIR))} archivos generados en {OUTPUT_DIR}\n")
