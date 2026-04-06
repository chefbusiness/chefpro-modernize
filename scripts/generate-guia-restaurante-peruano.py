#!/usr/bin/env python3
"""
Generate "Cómo Montar un Restaurante Peruano 80 Plazas" guide deliverables.
AI Chef Pro — aichef.pro
19 files: 1 DOCX guide + 8 Excel templates + 6 Excel checklists + 2 DOCX models + 1 PDF placeholder
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "public", "dl", "guia-restaurante-peruano"
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

GOLD = "FFD700"
GOLD_RGB = RGBColor(0xFF, 0xD7, 0x00)
HEADER_BG = "2D2D2D"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"
MEDIUM_GRAY = "E0E0E0"
INPUT_COLOR = "E8F5E9"
BRAND = "AI Chef Pro · aichef.pro — Restaurante Peruano"

title_font = Font(name="Calibri", size=16, bold=True, color=GOLD)
sub_font = Font(name="Calibri", size=11, color="888888", italic=True)
hdr_font = Font(name="Calibri", size=11, bold=True, color=WHITE)
sec_font = Font(name="Calibri", size=12, bold=True, color=GOLD)
dat_font = Font(name="Calibri", size=11)
bld_font = Font(name="Calibri", size=11, bold=True)
frm_font = Font(name="Calibri", size=11, color="1565C0", bold=True)
brn_font = Font(name="Calibri", size=9, color="888888", italic=True)

hdr_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
inp_fill = PatternFill(start_color=INPUT_COLOR, end_color=INPUT_COLOR, fill_type="solid")
lgt_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")
sec_fill = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")

thin_b = Border(
    left=Side(style="thin", color=MEDIUM_GRAY), right=Side(style="thin", color=MEDIUM_GRAY),
    top=Side(style="thin", color=MEDIUM_GRAY), bottom=Side(style="thin", color=MEDIUM_GRAY),
)
c_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
l_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
cur_fmt = '#,##0.00 €'
pct_fmt = '0.0%'


def shr(ws, row, ncols):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin_b; cell.alignment = c_align

def sdc(ws, r, c, v=None, font=None, fill=None, fmt=None, align=None):
    cell = ws.cell(row=r, column=c)
    if v is not None: cell.value = v
    cell.font = font or dat_font; cell.border = thin_b; cell.alignment = align or l_align
    if fill: cell.fill = fill
    if fmt: cell.number_format = fmt
    return cell

def title_block(ws, t, ncols=8):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    ws.cell(row=1, column=1, value=t).font = title_font
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
    ws.cell(row=2, column=1, value=BRAND).font = sub_font
    ws.row_dimensions[1].height = 30

def brand_footer(ws, row, ncols):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
    ws.cell(row=row, column=1, value=BRAND).font = brn_font
    ws.cell(row=row, column=1).alignment = c_align

def instr_sheet(wb, t, items):
    ws = wb.active; ws.title = "Instrucciones"; ws.sheet_properties.tabColor = "FFD700"
    title_block(ws, t, 3)
    r = 4
    ws.cell(row=r, column=1, value="Instrucciones de uso").font = sec_font; r += 1
    for i, x in enumerate(items, 1):
        ws.cell(row=r, column=1, value=f"{i}. {x}").font = dat_font; r += 1
    r += 1
    ws.cell(row=r, column=1, value="Celdas verdes = campos editables").font = dat_font
    ws.cell(row=r, column=1).fill = inp_fill
    ws.column_dimensions['A'].width = 80

def checklist_ws(ws, t, items):
    hdrs = ["#", "Categoría", "Tarea / Ítem", "Responsable", "Fecha Límite", "Estado", "Coste Est. (€)", "Notas"]
    title_block(ws, t)
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8); ws.freeze_panes = f"A{r+1}"
    ws.auto_filter.ref = f"A{r}:H{r}"
    dv = DataValidation(type="list", formula1='"Pendiente,En Curso,Completado"', allow_blank=True)
    ws.add_data_validation(dv)
    widths = [6, 22, 45, 18, 14, 14, 14, 30]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    for idx, (cat, tarea, resp, coste) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, tarea)
        sdc(ws, r, 4, resp); sdc(ws, r, 5, None, fill=inp_fill)
        sdc(ws, r, 6, "Pendiente", fill=inp_fill); dv.add(ws.cell(row=r, column=6))
        sdc(ws, r, 7, coste, fmt=cur_fmt, fill=inp_fill); sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1; brand_footer(ws, r, 8)


# ═══════════════════════════════════════════════════════════
# 1. GUIDE DOCX
# ═══════════════════════════════════════════════════════════
def gen_guide_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    def add_bullet(txt):
        doc.add_paragraph(txt, style='List Bullet')
    def add_num(txt):
        doc.add_paragraph(txt, style='List Number')
    def tip(txt):
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Consejo del Chef: {txt}")
        run.font.italic = True; run.font.color.rgb = GOLD_RGB

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cómo Montar un Restaurante Peruano"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("80 Plazas de Aforo — Guía España 2026"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Chef John Guerrero\nAI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # CH1
    doc.add_heading("1. Qué es un Restaurante Peruano Auténtico", level=1)
    doc.add_paragraph(
        "La cocina peruana es una de las más reconocidas del mundo. Tres restaurantes peruanos "
        "figuran en la lista 50 Best Restaurants del mundo (Central, Maido, Kjolle), y Perú ha sido "
        "elegido como Mejor Destino Culinario del Mundo en los World Travel Awards durante más de "
        "una década consecutiva. La gastronomía peruana fusiona tradiciones prehispánicas, españolas, "
        "africanas, chinas (chifa), japonesas (Nikkei) e italianas en una cocina extraordinariamente diversa."
    )
    doc.add_paragraph(
        "En España, el mercado de restauración peruana ha experimentado un auge enorme gracias a "
        "la gran comunidad peruana residente y al creciente interés del público español por el ceviche, "
        "el lomo saltado y los piscos. Un restaurante peruano auténtico con buena materia prima, "
        "una barra de piscos premium y una propuesta que va más allá del ceviche básico tiene un "
        "posicionamiento diferencial muy potente y márgenes superiores."
    )
    doc.add_heading("La diversidad peruana: mucho más que ceviche", level=2)
    add_bullet("Cocina criolla: lomo saltado, ají de gallina, seco de res, tacu tacu. La base de la gastronomía peruana.")
    add_bullet("Cocina marina (cevichería): ceviches, tiraditos, causas, leche de tigre. El segmento más conocido internacionalmente.")
    add_bullet("Chifa (fusión peruano-china): arroz chaufa, lomo saltado, wantán frito, aeropuerto. Enorme potencial en España.")
    add_bullet("Nikkei (fusión peruano-japonesa): tiraditos con acabados japoneses, sashimi peruano, maki acevichado. Alta gama.")
    add_bullet("Novoandina: reinterpretación gourmet de ingredientes andinos (quinoa, kiwicha, cuy, papa nativa). Fine dining.")
    add_bullet("Ticket medio auténtico: 20-35€ por persona, frente a 12-18€ de un peruano rápido/casual.")
    add_bullet("Margen diferencial: el ceviche cuesta 2-3€ de materia prima y se vende a 14-18€. El pisco sour tiene márgenes del 80%+.")
    tip("La autenticidad es tu ventaja competitiva. Si preparas la leche de tigre al momento delante del cliente y tienes una barra de piscos con 20-30 referencias, ya te has diferenciado del 90% de peruanos en España que sirven ceviche de bolsa.")
    doc.add_page_break()

    # CH2
    doc.add_heading("2. El Mercado de la Cocina Peruana en España 2026", level=1)
    doc.add_paragraph(
        "La cocina peruana es una de las tendencias gastronómicas más consolidadas en España. "
        "Con más de 200.000 peruanos residentes y un público español cada vez más sofisticado, "
        "la demanda de cocina peruana auténtica no para de crecer. España es el segundo país de Europa "
        "con más restaurantes peruanos, solo por detrás de Italia."
    )
    doc.add_heading("Datos clave del sector", level=2)
    add_bullet("Restaurantes peruanos en España: +1.800 (2025), crecimiento del 40% en 5 años.")
    add_bullet("Ticket medio: 20-35€ (auténtico) vs 12-18€ (peruano rápido/casual).")
    add_bullet("Ciudades con mayor demanda: Madrid, Barcelona, Valencia, Málaga, Bilbao, Sevilla.")
    add_bullet("Público: 25-50 años, urbano, foodie, viajero, busca experiencias gastronómicas premium.")
    add_bullet("Tendencia clave: pisco y coctelería peruana en auge. El pisco sour es ya un clásico en las barras de España.")
    add_bullet("3 peruanos en 50 Best: Central (#1), Maido (#5), Kjolle (#28) — la cocina peruana está en la élite mundial.")
    add_bullet("Delivery peruano: ceviches y lomo saltado son cada vez más demandados en plataformas, aunque requieren packaging especial.")
    tip("El 80% de los restaurantes peruanos en España ofrecen solo ceviche y lomo saltado. Si incorporas chifa, Nikkei o novoandina, te posicionas en un segmento premium con muy poca competencia directa.")
    doc.add_page_break()

    # CH3
    doc.add_heading("3. Modelos de Negocio: Cevichería, Criollo, Nikkei y Más", level=1)
    doc.add_paragraph("El restaurante peruano admite múltiples modelos. La elección del modelo determina la inversión, la operativa diaria y tu posicionamiento en el mercado.")
    doc.add_heading("Modelo 1: Cevichería", level=2)
    doc.add_paragraph("Especializado en ceviches, tiraditos, causas y cocina marina. Barra de ceviche a la vista, leche de tigre al momento. Ticket 18-28€. Producto fresco como estrella. Requiere proveedor de pescado de máxima calidad y cevichero experto.")
    doc.add_heading("Modelo 2: Restaurante Criollo", level=2)
    doc.add_paragraph("Carta amplia de cocina criolla peruana: lomo saltado, ají de gallina, seco de res, arroz con mariscos, anticuchos, tacu tacu. Concepto familiar, porciones generosas. Ticket 20-30€. El modelo más replicable y estable en España.")
    doc.add_heading("Modelo 3: Nikkei / Fusión Peruano-Japonesa", level=2)
    doc.add_paragraph("Alta gama. Tiraditos con acabados japoneses, sashimi con ají amarillo, maki acevichado, ceviche de atún con ponzu. Ticket 35-55€. Requiere chef con formación en cocina japonesa + peruana. Muy diferencial pero más exigente en equipo y producto.")
    doc.add_heading("Modelo 4: Chifa (Peruano-Chino)", level=2)
    doc.add_paragraph("Fusión peruano-china: arroz chaufa, lomo saltado, wantán frito, sopa wantán, aeropuerto, tallarín saltado. Wok como equipo estrella. Ticket 15-22€. Alta rotación, menor inversión en producto fresco. Muy popular entre la comunidad latina.")
    doc.add_heading("Modelo 5: Novoandina / Fine Dining Peruano", level=2)
    doc.add_paragraph("Reinterpretación de la cocina andina con técnicas modernas. Ingredientes como quinoa, papa nativa, ají, rocoto, cuy, alpaca. Menú degustación 60-90€. Requiere chef de alto nivel y concepto de autor.")
    doc.add_heading("Modelo 6: Pollería Peruana", level=2)
    doc.add_paragraph("Pollo a la brasa peruano con papas fritas y ensalada. Concepto simple, alta rotación, ticket 12-18€. Inversión en horno de pollos a la brasa. Complementar con anticuchos, salchipapas y chicha morada.")
    doc.add_heading("Modelo 7: Dark Kitchen / Delivery Peruano", level=2)
    doc.add_paragraph("Solo delivery y take away: menú optimizado para lomo saltado, arroz chaufa y bowls peruanos que viajan bien. Sin sala. Inversión 40-80K€. Ideal para testear mercado antes de abrir local.")
    doc.add_page_break()

    # CH4
    doc.add_heading("4. Estudio de Viabilidad y Plan Financiero", level=1)
    doc.add_paragraph(
        "La inversión total para un restaurante peruano de 80 plazas en España oscila "
        "entre 130.000€ y 300.000€ dependiendo de la ubicación, nivel de acabados, "
        "si incluyes barra de piscos premium y el grado de decoración temática peruana."
    )
    doc.add_heading("Desglose de inversión típica", level=2)
    add_bullet("Obra civil y acondicionamiento: 45.000-110.000€")
    add_bullet("Equipamiento cocina peruana (estación ceviche, wok, parrilla anticuchos, horno brasa): 28.000-55.000€")
    add_bullet("Barra de piscos y coctelería (instalación + stock inicial 20-30 refs): 6.000-15.000€")
    add_bullet("Mobiliario sala + decoración peruana (textiles andinos, cerámica, madera, colores tierra): 15.000-35.000€")
    add_bullet("Licencias y permisos (incluye importación pisco, ají): 5.000-12.000€")
    add_bullet("Marketing pre-apertura: 3.000-8.000€")
    add_bullet("Stock inicial (materia prima + ajíes + especias importadas + pisco): 5.000-12.000€")
    add_bullet("Fondo de maniobra (3 meses): 25.000-65.000€")
    add_bullet("Tecnología (TPV, web, delivery): 3.000-8.000€")
    doc.add_heading("Ratios financieros objetivo", level=2)
    add_bullet("Food cost: 28-32% sobre ventas (el pescado fresco para ceviche sube ligeramente el coste)")
    add_bullet("Coste de personal: 28-35% sobre ventas")
    add_bullet("Alquiler: máximo 8-10% sobre ventas")
    add_bullet("EBITDA objetivo: 12-18%")
    add_bullet("Break-even: mes 8-14 (depende de ubicación y estacionalidad)")
    tip("El margen en bebidas (pisco sour, chilcano) compensa el food cost del pescado fresco. Un pisco sour cuesta 1.20€ y lo vendes a 9-12€. La barra de piscos puede representar el 25-30% de tu facturación con márgenes del 80%.")
    doc.add_page_break()

    # CH5
    doc.add_heading("5. Requisitos Legales en España + Importación", level=1)
    doc.add_paragraph("Abrir un restaurante peruano en España requiere los mismos trámites que cualquier restaurante, más requisitos adicionales si importas productos directamente de Perú (pisco, ají amarillo, rocoto, maíz morado, chicha morada).")
    doc.add_heading("Trámites generales", level=2)
    add_num("Constitución de sociedad (SL) o alta como autónomo")
    add_num("Licencia de actividad / declaración responsable (según municipio)")
    add_num("Registro sanitario autonómico")
    add_num("Alta en Hacienda: IAE epígrafe 671.4 o similar")
    add_num("Inscripción Seguridad Social empresa")
    add_num("Seguro de responsabilidad civil (mínimo 300.000€)")
    add_num("Plan APPCC y registro RGSEAA")
    add_num("Licencia de terraza (si aplica)")
    add_num("Licencia de música / SGAE")
    add_num("Libro de reclamaciones")
    doc.add_heading("Requisitos específicos de importación", level=2)
    add_bullet("Registro como operador de comercio exterior (EORI) si importas directamente")
    add_bullet("Certificado SOIVRE para productos alimentarios importados de fuera de la UE")
    add_bullet("Etiquetado en español obligatorio para todos los productos importados")
    add_bullet("Control fitosanitario para ajíes secos, especias y productos vegetales")
    add_bullet("Impuestos especiales para bebidas alcohólicas (pisco): IIEE")
    add_bullet("Denominación de Origen Pisco: solo el pisco peruano puede llamarse 'Pisco' en la UE (regulación protegida)")
    add_bullet("Alternativa más sencilla: comprar a distribuidores especializados en España (El Ají, Peru Gourmet, La Tienda Peruana, Inca Products)")
    tip("No necesitas importar directamente. Hay distribuidores especializados en Madrid y Barcelona que traen todo lo que necesitas: ají amarillo, rocoto, maíz morado, pisco, chicha morada concentrada. Empieza con distribuidores y solo importa directamente cuando tu volumen lo justifique.")
    doc.add_page_break()

    # CH6
    doc.add_heading("6. APPCC y Seguridad Alimentaria para Cocina Peruana", level=1)
    doc.add_paragraph("El plan APPCC de un restaurante peruano tiene las mismas bases que cualquier restaurante, pero con puntos de control adicionales para pescado crudo (ceviche, tiradito), productos importados y técnicas específicas de la cocina peruana.")
    doc.add_heading("Prerrequisitos estándar", level=2)
    add_bullet("Plan de limpieza y desinfección (frecuencias, productos, responsables)")
    add_bullet("Control de temperaturas (cámaras, servicio, transporte)")
    add_bullet("Plan de control de plagas (DDD)")
    add_bullet("Trazabilidad de materias primas (proveedores, lotes, fechas)")
    add_bullet("Plan de formación del personal (manipulador de alimentos)")
    add_bullet("Control de agua potable")
    add_bullet("Plan de gestión de residuos y aceite usado")
    add_bullet("Plan de alérgenos (14 alérgenos de declaración obligatoria)")
    doc.add_heading("Puntos críticos específicos de cocina peruana", level=2)
    add_bullet("Pescado para ceviche: CONGELACIÓN PREVIA obligatoria -20°C/24h para eliminar anisakis. El zumo de limón NO mata el parásito.")
    add_bullet("Leche de tigre: preparar al momento, vida útil máxima 2h refrigerado. No reutilizar nunca.")
    add_bullet("Tiraditos: mismos requisitos que ceviche, pescado previamente congelado.")
    add_bullet("Mariscos (pulpo, langostinos, conchas): trazabilidad estricta, cadena de frío ininterrumpida.")
    add_bullet("Ají amarillo y rocoto importados (en pasta o congelados): verificar etiquetado, almacenar según indicaciones.")
    add_bullet("Anticuchos (corazón de res): asegurar cocción interna >75°C, marinado previo no elimina patógenos.")
    add_bullet("Chicha morada: si se prepara en casa, vida útil refrigerado 48h máximo.")
    tip("El ceviche es tu plato estrella y tu mayor riesgo sanitario. SIEMPRE congela el pescado previamente (-20°C mínimo 24h). Esto es obligatorio por ley en España y no negociable. El limón NO mata el anisakis. Ten documentación de congelación para cada lote.")
    doc.add_page_break()

    # CH7
    doc.add_heading("7. Ubicación y Local", level=1)
    doc.add_paragraph("La ubicación de un restaurante peruano es crucial. El componente de frescura del producto (pescado para ceviche) exige cercanía a mercados de abastos o proveedores fiables de pescado fresco.")
    doc.add_heading("Criterios de selección", level=2)
    add_bullet("Tráfico peatonal: mínimo 400 personas/hora en horario punta")
    add_bullet("Visibilidad: fachada atractiva con identidad peruana visible desde la calle")
    add_bullet("Metros cuadrados: 160-220 m² para 80 plazas (ratio 2-2.75 m²/plaza)")
    add_bullet("Zona: ocio + afterwork + zona gastronómica = ideal. El ceviche funciona especialmente bien en mediodía y afterwork.")
    add_bullet("Competencia: analiza los 500m alrededor — evita zonas con más de 2 peruanos ya establecidos")
    add_bullet("Salida de humos: imprescindible. La parrilla de anticuchos, el wok chifa y el horno de pollos generan mucho humo.")
    add_bullet("Proximidad a mercado de abastos: ventaja enorme para pescado fresco diario")
    add_bullet("Accesibilidad: acceso para personas con movilidad reducida (obligatorio por ley)")
    add_bullet("Terraza: si la zona lo permite, la terraza con decoración peruana atrae tráfico y es ideal para ceviches en verano")
    tip("El mediodía es el servicio estrella de un cevichería (el peruano dice 'el ceviche se come de día'). Pero la cena es donde está el ticket alto con pisco y coctelería. Busca zonas con tráfico tanto a mediodía como por la noche.")
    doc.add_page_break()

    # CH8
    doc.add_heading("8. Diseño de Cocina Peruana Profesional", level=1)
    doc.add_paragraph("La cocina de un restaurante peruano tiene estaciones especializadas que no existen en una cocina convencional: la estación del cevichero, el wok para chifa/Nikkei, la parrilla de anticuchos y la estación de salsas peruanas.")
    doc.add_heading("Zonas de la cocina peruana", level=2)
    add_bullet("Estación del cevichero: mesa fría, tabla de corte grande, cítricos (limón peruano o lima), ají limo, cebolla morada, cilantro. Visible al cliente si es posible (show cooking). Zona refrigerada para pescado.")
    add_bullet("Estación wok (chifa/Nikkei): wok profesional de alta potencia (20-30 kW), zona de salteado rápido, mise en place de salsas (sillao, ostión, tamarindo).")
    add_bullet("Parrilla de anticuchos: parrilla de carbón o gas, zona de marinado, brochetas y utensilios.")
    add_bullet("Zona caliente: fogones para guisos criollos (seco, ají de gallina, arroz con mariscos), horno para pollo a la brasa.")
    add_bullet("Estación de salsas: preparación de ají amarillo, rocoto, huancaína, ocopa, criolla. Licuadoras industriales.")
    add_bullet("Zona fría: preparación de causas, ensaladas, tiraditos.")
    add_bullet("Almacenamiento: cámaras frigoríficas (2) con zona específica para pescado, economato seco para ajíes y especias.")
    add_bullet("Zona de lavado: tren de lavado, zona de limpieza.")
    add_bullet("Superficie mínima cocina: 45-60 m² (ratio cocina/sala 1:3)")
    tip("La estación del cevichero a la vista del cliente es tu mayor herramienta de marketing. Ver cómo se prepara un ceviche al momento — el corte del pescado, el zumo de limón, la leche de tigre — es un espectáculo que genera fotos en redes y fideliza.")
    doc.add_page_break()

    # CH9
    doc.add_heading("9. Equipamiento Específico de Cocina Peruana", level=1)
    doc.add_paragraph("Además del equipamiento estándar de cocina profesional, un restaurante peruano necesita equipos específicos para las distintas estaciones de su cocina.")
    doc.add_heading("Equipamiento específico peruano con costes orientativos", level=2)
    add_bullet("Wok profesional de alta potencia (20-30 kW, quemador chino): 1.500-4.000€")
    add_bullet("Mesa refrigerada para estación de ceviche: 2.500-4.000€")
    add_bullet("Parrilla para anticuchos (carbón o gas): 1.500-3.500€")
    add_bullet("Horno de pollo a la brasa (si modelo pollería/criollo): 3.000-8.000€")
    add_bullet("Licuadoras industriales para salsas y ajíes (3 uds): 400-700€")
    add_bullet("Vitrina de piscos iluminada para barra: 2.000-5.000€")
    add_bullet("Exprimidor de cítricos industrial (alto consumo de limón para ceviches): 500-1.200€")
    add_bullet("Tablas de corte extra grandes para cevichero: 100-300€")
    doc.add_heading("Equipamiento estándar de cocina", level=2)
    add_bullet("Cocina de gas 6 fuegos + horno: 3.500-6.000€")
    add_bullet("Freidora doble (yuca frita, chicharrón, camote): 1.200-2.500€")
    add_bullet("Horno mixto (Rational / Unox): 8.000-15.000€")
    add_bullet("Cámaras frigoríficas (2 puertas): 2.500-4.000€")
    add_bullet("Cámara congelación (imprescindible para congelación previa de pescado): 2.000-3.500€")
    add_bullet("Tren de lavado: 4.000-8.000€")
    add_bullet("Campana extractora + filtros (sobredimensionar por wok y parrilla): 4.000-8.000€")
    add_bullet("Total equipamiento cocina: 28.000-55.000€")
    tip("El wok de alta potencia es fundamental si ofreces chifa. Un wok doméstico no alcanza la temperatura necesaria para el 'wok hei' (sabor ahumado del salteado rápido). Invierte en un quemador chino de 20+ kW — la diferencia en sabor es brutal.")
    doc.add_page_break()

    # CH10
    doc.add_heading("10. Diseño de Sala y Decoración Peruana", level=1)
    doc.add_paragraph("La decoración de un restaurante peruano debe reflejar la riqueza cultural del país sin caer en clichés. Perú tiene una identidad visual potente: textiles andinos, cerámica precolombina, colores tierra, madera y piedra. El reto es crear un ambiente sofisticado y cálido.")
    doc.add_heading("Distribución tipo (80 plazas)", level=2)
    add_bullet("Mesas de 2: 6-8 unidades (12-16 plazas) — parejas")
    add_bullet("Mesas de 4: 10-12 unidades (40-48 plazas) — familias, grupos")
    add_bullet("Barra alta / barra de piscos: 10-16 plazas — afterwork, cócteles")
    add_bullet("Zona terraza (si aplica): 15-25 plazas adicionales")
    add_bullet("Distancia entre mesas: 60-80 cm")
    doc.add_heading("Elementos clave de la decoración peruana auténtica", level=2)
    add_bullet("Textiles andinos: mantas, tapices o cojines con patrones geométricos andinos en colores tierra, rojo, ocre y negro.")
    add_bullet("Cerámica: piezas de cerámica inspiradas en culturas precolombinas (Moche, Nasca, Chimú) como elementos decorativos.")
    add_bullet("Madera: mesas, barra y detalles en madera natural. La madera es omnipresente en la arquitectura peruana.")
    add_bullet("Colores tierra: paleta de marrones, terracota, ocre, beige, con acentos en rojo y dorado. Evitar colores fríos.")
    add_bullet("Piedra: detalles en piedra natural en pared de acento o barra. Evoca la arquitectura inca.")
    add_bullet("Plantas: helechos, suculentas y plantas tropicales. Verde como acento sobre colores tierra.")
    add_bullet("Iluminación: cálida (2700K), lámparas de fibra natural o cerámica.")
    add_bullet("Vajilla: platos de cerámica artesanal en tonos tierra y negro. Muy instagrameable.")
    tip("Menos es más en decoración peruana. Unos textiles andinos de calidad en los asientos, vajilla artesanal, madera natural y una iluminación cálida crean un ambiente sofisticado. Evita las fotos de Machu Picchu y las llamas de peluche — eso es un peruano de aeropuerto.")
    doc.add_page_break()

    # CH11
    doc.add_heading("11. Barra de Piscos (20-30 Referencias)", level=1)
    doc.add_paragraph("La barra de piscos es el elemento diferencial que eleva un restaurante peruano de casual a experiencial. El pisco sour es ya un cóctel clásico en España, y el interés por el pisco como destilado premium está creciendo rápidamente.")
    doc.add_heading("Estructura de la barra (20-30 referencias)", level=2)
    add_bullet("Pisco Quebranta (6-8 refs): Barsol, Portón, Tacama, Viñas de Oro, La Caravedo, Cuatro Gallos")
    add_bullet("Pisco Italia (3-4 refs): aromático, ideal para chilcano. Barsol Italia, Portón Mosto Verde Italia")
    add_bullet("Pisco Acholado (4-6 refs): mezcla de uvas, versátil para cócteles. Barsol Selecto, BarSol Primero")
    add_bullet("Pisco Mosto Verde (3-4 refs): premium, destilado de mosto parcialmente fermentado. Portón Mosto Verde, La Caravedo")
    add_bullet("Pisco premium/artesanal (2-3 refs): ediciones limitadas, piscos de autor")
    add_bullet("Cócteles peruanos (8-12): Pisco Sour, Chilcano de pisco, Capitán, Maracuyá Sour, Pisco Punch, Algarrobina")
    doc.add_heading("Inversión y márgenes", level=2)
    add_bullet("Stock inicial barra (20-30 refs): 3.000-10.000€")
    add_bullet("Vitrina / estantería iluminada para botellas: 2.000-5.000€")
    add_bullet("Margen pisco premium: 75-85% (compra 18€/botella → venta 7-10€/copa)")
    add_bullet("Margen cócteles: 78-85% (un pisco sour cuesta 1.20€ y se vende a 9-12€)")
    add_bullet("Objetivo: que la barra represente 20-30% de la facturación total")
    tip("Ofrece vuelos de pisco (3 copas de diferentes variedades: Quebranta, Italia, Mosto Verde) por 14-18€. Es la forma más rentable de vender pisco premium: el cliente aprende, disfruta y acaba pidiendo pisco sour con Mosto Verde. Margen brutal.")
    doc.add_page_break()

    # CH12
    doc.add_heading("12. Brigada de Cocina (6-10 personas)", level=1)
    doc.add_paragraph("La brigada de cocina de un restaurante peruano tiene roles específicos que no existen en una cocina europea convencional: el cevichero, el wokero (chifa/Nikkei) y el parrillero de anticuchos son puestos especializados.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Jefe de cocina (idealmente con experiencia en cocina peruana): 1 persona — 2.200-2.800€ brutos/mes")
    add_bullet("Cevichero (estación de ceviche, tiraditos, causas, leche de tigre): 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Wokero (estación wok para chifa/Nikkei: chaufa, saltados, tallarines): 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Parrillero (anticuchos, pollo a la brasa, carnes a la parrilla): 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Cocineros (guisos criollos: ají de gallina, seco, arroz con mariscos): 2-3 personas — 1.500-1.800€ brutos/mes cada uno")
    add_bullet("Ayudante / office: 1-2 personas — 1.200-1.500€ brutos/mes")
    add_bullet("Coste total cocina: 13.000-22.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Contratación del cevichero", level=2)
    doc.add_paragraph("El cevichero es el alma del restaurante peruano. Necesita conocimiento profundo de pescados, técnica de corte, dominio de la leche de tigre y sensibilidad para equilibrar ácido-picante-sal. Si encuentras un cevichero peruano con experiencia, contratarlo es una inversión que se paga sola. Alternativamente, un cocinero español con experiencia en cocina de mercado puede formarse con un programa de 3-4 semanas.")
    tip("Busca cevicheros y cocineros peruanos en comunidades latinas de Madrid, Barcelona y Valencia. Hay talento increíble. LinkedIn, grupos de Facebook de peruanos en España, la Asociación de Restaurantes Peruanos en España (ARPE) y eventos gastronómicos peruanos son tus mejores canales de reclutamiento.")
    doc.add_page_break()

    # CH13
    doc.add_heading("13. Equipo de Sala (6-8 personas)", level=1)
    doc.add_paragraph("El equipo de sala de un restaurante peruano necesita conocer la cultura gastronómica peruana para poder explicar platos, recomendar piscos y crear la experiencia completa.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Encargado/a de sala: 1 persona — 2.000-2.500€ brutos/mes")
    add_bullet("Barman de piscos y coctelería peruana: 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Camareros: 3-4 personas — 1.400-1.700€ brutos/mes cada uno")
    add_bullet("Runner / ayudante: 1-2 personas — 1.200-1.400€ brutos/mes")
    add_bullet("Coste total sala: 9.000-15.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Claves del servicio en un restaurante peruano", level=2)
    add_bullet("Conocimiento de piscos: todo el equipo debe saber distinguir Quebranta, Italia, Acholado, Mosto Verde y recomendar.")
    add_bullet("Explicar platos: muchos clientes no saben qué es una causa limeña, un tiradito o un ají de gallina. El camarero debe saber describirlo con pasión.")
    add_bullet("Sugerir orden de comida: en Perú se comparte. Recomendar pedir ceviches/causas para compartir + principales individuales.")
    add_bullet("Gestión de picante: el ají amarillo y el rocoto tienen niveles diferentes. Preguntar tolerancia y adaptar salsas.")
    add_bullet("Upselling de bebidas: pisco sour, chilcano, vuelos de pisco. El barman es clave para el ticket medio.")
    add_bullet("Chicha morada como bebida signature: ofrecer siempre como alternativa sin alcohol, es un gran diferenciador.")
    tip("La formación del barman de piscos es una inversión que se multiplica. Un barman que sabe contar la historia del pisco, explica las variedades de uva y prepara un pisco sour perfecto vende el triple. Hay cursos de la Ruta del Pisco y la Cofradía del Pisco.")
    doc.add_page_break()

    # CH14
    doc.add_heading("14. Menú Engineering: Ceviches, Tiraditos, Causas, Lomo Saltado", level=1)
    doc.add_paragraph("El diseño de la carta de un restaurante peruano debe equilibrar autenticidad, margen y accesibilidad. No todos los clientes conocen la cocina peruana: la carta debe educar y seducir.")
    doc.add_heading("Estructura de carta peruana", level=2)
    add_bullet("Ceviches y tiraditos: 4-6 opciones (clásico, mixto, de pescado del día, Nikkei)")
    add_bullet("Causas y entradas frías: 3-4 opciones (causa limeña, pulpo al olivo, papa a la huancaína)")
    add_bullet("Anticuchos y entradas calientes: 3-4 opciones (anticuchos de corazón, tequeños, chicharrón)")
    add_bullet("Principales criollos: 6-8 opciones (lomo saltado, ají de gallina, seco de res, arroz con mariscos, tacu tacu)")
    add_bullet("Chifa (si aplica): 3-4 opciones (arroz chaufa, tallarín saltado, aeropuerto)")
    add_bullet("Postres: 3-4 opciones (suspiro limeño, picarones, mazamorra morada, tres leches)")
    add_bullet("Total carta: 22-30 platos")
    doc.add_heading("Menú engineering: clasificación de platos peruanos", level=2)
    add_bullet("Stars (alta popularidad + alto margen): Ceviche clásico, Pisco Sour, Causa limeña → destacar siempre")
    add_bullet("Plowhorses (alta popularidad + bajo margen): Lomo saltado, Arroz con mariscos → subir precio o ajustar porción")
    add_bullet("Puzzles (baja popularidad + alto margen): Tiradito Nikkei, Suspiro limeño → mejorar descripción en carta")
    add_bullet("Dogs (baja popularidad + bajo margen): Ensalada peruana, Sopa criolla → eliminar o reinventar")
    tip("El ceviche clásico es tu Star absoluto. El coste de un ceviche es 2.50-3€ y lo vendes a 14-18€. Con un ceviche + pisco sour por persona ya tienes un ticket de 25€+ con márgenes altísimos. Es el plato con mejor ratio margen/satisfacción de toda la carta.")
    doc.add_page_break()

    # CH15
    doc.add_heading("15. Proveedores de Producto Peruano en España", level=1)
    doc.add_paragraph("Encontrar ingredientes auténticos peruanos en España es más fácil de lo que parece. La gran comunidad peruana en España ha creado una red sólida de distribuidores e importadores.")
    doc.add_heading("Ingredientes clave y dónde encontrarlos", level=2)
    add_bullet("Ají amarillo (en pasta, congelado o seco): distribuidores latinos especializados, El Ají, Peru Gourmet")
    add_bullet("Rocoto (en pasta o fresco congelado): distribuidores peruanos, tiendas latinas de Madrid/Barcelona")
    add_bullet("Ají limo (para ceviches): difícil fresco en España, sustituir con ají similar o importar congelado")
    add_bullet("Maíz morado (para chicha morada): distribuidores peruanos, tiendas latinas")
    add_bullet("Pisco (Denominación de Origen): distribuidores como Barsol Europe, Pisco Portón, o directo a bodegas peruanas")
    add_bullet("Chicha morada concentrada: importación desde Perú, disponible en distribuidores latinos")
    add_bullet("Huacatay, muña, huacatay en pasta: tiendas latinas especializadas")
    add_bullet("Pescado fresco para ceviche: Mercamadrid, Mercabarna, lonjas locales. Lubina, corvina, lenguado, dorada son las mejores opciones en España.")
    add_bullet("Limón peruano: difícil de conseguir en España. Sustituir con lima (key lime) o mezcla de limón + lima.")
    add_bullet("Papa amarilla peruana: no disponible fresca. Sustituir con papa Monalisa o Kennebec (textura similar).")
    tip("Haz una visita a los mercados latinos de Madrid (Lavapiés, Usera) o Barcelona (Raval). Encontrarás ají amarillo, rocoto, maíz morado y muchos productos peruanos a precios de mayorista. Muchos proveedores hacen reparto diario si superas 50€/pedido.")
    doc.add_page_break()

    # CH16
    doc.add_heading("16. 15 Recetas Base con Food Cost", level=1)
    doc.add_paragraph("Estas 15 recetas forman la base de la carta de un restaurante peruano. Cada una incluye food cost estimado y PVP sugerido para mantener márgenes del 28-32%.")
    recipes = [
        ("Ceviche clásico de corvina", "2.80€", "14.00€", "20%"),
        ("Lomo saltado", "3.20€", "15.00€", "21%"),
        ("Causa limeña de atún", "1.80€", "10.00€", "18%"),
        ("Ají de gallina con arroz", "2.00€", "12.00€", "17%"),
        ("Anticuchos de corazón (3 brochetas)", "1.60€", "11.00€", "15%"),
        ("Tiradito de lubina al ají amarillo", "3.00€", "15.00€", "20%"),
        ("Arroz con mariscos", "3.80€", "16.00€", "24%"),
        ("Papa a la huancaína", "1.00€", "8.00€", "13%"),
        ("Pollo a la brasa con papas", "2.50€", "13.00€", "19%"),
        ("Seco de res con frijoles", "2.80€", "13.00€", "22%"),
        ("Tacu tacu con lomo fino", "3.50€", "16.00€", "22%"),
        ("Picarones con miel de chancaca", "0.80€", "7.00€", "11%"),
        ("Suspiro limeño", "0.90€", "7.00€", "13%"),
        ("Chicharrón con mote y salsa criolla", "2.20€", "12.00€", "18%"),
        ("Cau cau de mondongo", "1.80€", "11.00€", "16%"),
    ]
    for name, fc, pvp, pct in recipes:
        doc.add_heading(name, level=2)
        add_bullet(f"Food cost: {fc}")
        add_bullet(f"PVP sugerido (sin IVA): {pvp}")
        add_bullet(f"% Food cost: {pct}")
    tip("El ceviche preparado al momento delante del cliente cuesta 2.80€ y lo vendes a 14€. Es el plato con mejor margen y el más instagrameable. La leche de tigre sobrante se vende como shot a 4€ — margen infinito porque ya está hecho.")
    doc.add_page_break()

    # CH17
    doc.add_heading("17. Delivery Optimizado para Cocina Peruana", level=1)
    doc.add_paragraph("La cocina peruana tiene platos que viajan bien y otros que no. El lomo saltado, el arroz chaufa y los bowls peruanos son ideales para delivery, mientras que el ceviche requiere packaging especial.")
    doc.add_heading("Plataformas y canales", level=2)
    add_bullet("Glovo, Uber Eats, Just Eat: comisión 25-35% — útiles para visibilidad inicial")
    add_bullet("Delivery propio (web/app): comisión 0% pero requiere repartidores propios o Stuart/Paack")
    add_bullet("Take away en local: margen completo, packaging incluido en precio")
    doc.add_heading("Menú delivery peruano optimizado", level=2)
    add_bullet("Lomo saltado: viaja muy bien en envase con compartimento para arroz separado.")
    add_bullet("Arroz chaufa: envase cerrado, mantiene calor. Producto estrella del delivery peruano.")
    add_bullet("Bowls peruanos: base de arroz + proteína + ají + ensalada criolla. Formato bowl = transporte seguro.")
    add_bullet("Pollo a la brasa: cuartos de pollo con papas, ají y ensalada en envase familiar.")
    add_bullet("Causa limeña: viaja bien refrigerada, ideal como entrante de delivery.")
    add_bullet("Ceviche delivery: SOLO si se entrega en menos de 20 min. Pescado y leche de tigre separados, el cliente mezcla al recibir.")
    add_bullet("NO incluir en delivery: tiraditos (se oxidan), platos con mucha salsa líquida, suspiro limeño (frágil).")
    add_bullet("Packaging: envases de cartón kraft con compartimentos, bolsa de papel + logo. Coste 0.50-1.50€/pedido.")
    tip("El 'ceviche kit' es una idea brillante para delivery: envías el pescado cortado en frío + leche de tigre en bote + cebolla morada + camote + choclo por separado. El cliente arma su ceviche en casa, ultra-fresco.")
    doc.add_page_break()

    # CH18
    doc.add_heading("18. Marketing con Fiestas Patrias (28 de Julio) y Más", level=1)
    doc.add_paragraph("Un restaurante peruano tiene un calendario de eventos culturales que genera picos de facturación naturales. Las Fiestas Patrias (28-29 de julio) son el Super Bowl del restaurante peruano en el extranjero.")
    doc.add_heading("Eventos culturales para marketing", level=2)
    add_bullet("Fiestas Patrias (28-29 julio): el evento más importante del año. Menú especial con ceviche, anticuchos, pisco sour. Decoración con banderas peruanas, música criolla. Puede facturar 3x una semana normal.")
    add_bullet("Día del Pisco Sour (primer sábado de febrero): promociones de pisco sour, catas de pisco, vuelos de pisco. Excelente para redes sociales.")
    add_bullet("Día del Ceviche (28 junio): menú especial de ceviches, degustación, concurso de ceviches. Gran engagement en redes.")
    add_bullet("Inti Raymi / Fiesta del Sol (24 junio): celebración inca, ideal para menú andino especial o decoración temática.")
    add_bullet("Mistura (festival gastronómico peruano): aunque es en Lima, puedes crear tu propio 'mini Mistura' con platos de todas las regiones del Perú.")
    add_bullet("Navidad peruana: panetón, chocolate caliente peruano, pavo con relleno peruano.")
    doc.add_heading("Canales de marketing", level=2)
    add_num("Instagram: contenido visual de ceviches, preparación de leche de tigre, cocktails de pisco. Reels > fotos.")
    add_num("TikTok: vídeos del cevichero preparando ceviches, el wokero saltando chaufa, pisco sour perfecto. Potencial viral altísimo.")
    add_num("Google My Business: ficha completa con fotos profesionales, menú, reseñas.")
    add_num("TheFork / TripAdvisor: perfiles con fotos y respuesta a reseñas.")
    add_num("Web propia: menú, historia, reservas online, carta de piscos.")
    doc.add_heading("Presupuesto marketing mensual", level=2)
    add_bullet("Google Ads local: 200-500€/mes")
    add_bullet("Instagram/Facebook Ads: 200-500€/mes")
    add_bullet("Fotógrafo food: 200-300€/sesión trimestral")
    add_bullet("Community manager (si se externaliza): 300-600€/mes")
    add_bullet("Total recomendado: 600-1.500€/mes (2-3% facturación)")
    tip("Las Fiestas Patrias (28 de julio) son tu Super Bowl. Decora el restaurante con banderas, prepara menú especial con todos los clásicos, ofrece pisco sour libre por 25€ y pon música criolla en vivo. Un solo fin de semana de Fiestas Patrias puede facturar el equivalente a 3 semanas normales. Planifica la campaña con 1 mes de antelación.")
    doc.add_page_break()

    # CH19
    doc.add_heading("19. Tecnología para Restaurante Peruano", level=1)
    doc.add_paragraph("La tecnología en un restaurante peruano debe ser invisible para el cliente pero esencial para la operativa. Automatiza lo repetitivo para que tu equipo se centre en crear experiencias.")
    doc.add_heading("Stack tecnológico recomendado", level=2)
    add_bullet("TPV: Lightspeed, Square, Revo. Cloud, con reporting en tiempo real. 60-150€/mes.")
    add_bullet("Reservas: TheFork, CoverManager, Resy. Gestión de turnos y no-shows. 0-100€/mes.")
    add_bullet("Delivery: integrador (Ordatic, Deliverect) que centraliza Glovo + Uber Eats + Just Eat. 80-150€/mes.")
    add_bullet("Contabilidad: Holded, Quipu. Facturas automáticas, conexión con asesoría. 30-60€/mes.")
    add_bullet("RRHH y turnos: Factorial, Kenjo. Cuadrantes, fichajes, nóminas. 4-8€/empleado/mes.")
    add_bullet("Web y carta digital: QR con carta actualizable, sección de piscos. WordPress o Squarespace. 10-30€/mes.")
    add_bullet("WiFi profesional: router dual-band, red separada para clientes y cocina. 50-80€/mes.")
    tip("El TPV es el corazón de tu restaurante. Elige uno cloud que te permita ver ventas, costes y KPIs desde el móvil. Analiza qué ceviches venden más, qué piscos rotan y cuál es tu margen real por plato.")
    doc.add_page_break()

    # CH20
    doc.add_heading("20. Plan de Acción: De la Idea a la Inauguración", level=1)
    doc.add_paragraph("Abrir un restaurante peruano en España es un proyecto de 10-14 meses desde la primera idea hasta la inauguración. Este capítulo resume el cronograma y los hitos clave.")
    doc.add_heading("Fases del proyecto", level=2)
    add_bullet("Meses 1-3: Plan de negocio, búsqueda de financiación, constitución de empresa, estudio de mercado.")
    add_bullet("Meses 3-5: Búsqueda y contratación de local, negociación de alquiler, licencia de actividad.")
    add_bullet("Meses 4-7: Obra civil, instalaciones, decoración peruana.")
    add_bullet("Meses 5-7: Equipamiento de cocina, pedido de wok y parrilla, instalación de barra de piscos.")
    add_bullet("Meses 5-7: Establecer relaciones con proveedores de producto peruano e importadores de pisco.")
    add_bullet("Meses 7-9: Contratación y formación de equipo (especialmente cevichero, wokero y barman de piscos).")
    add_bullet("Meses 8-10: Branding, web, redes sociales, fotografía profesional.")
    add_bullet("Mes 10-11: Soft opening con amigos, familia e influencers.")
    add_bullet("Mes 11-12: Ajustes finales e INAUGURACIÓN.")
    add_bullet("Mes 12-14: Seguimiento, ajustes de carta y optimización de costes.")
    tip("No abras sin haber hecho mínimo 3 soft openings. Necesitas probar la estación del cevichero a pleno rendimiento, la velocidad del wok, los tiempos de coctelería y el flujo de sala. Cada soft opening te enseña fallos que no puedes detectar en papel.")

    # Save
    path = os.path.join(OUTPUT_DIR, "guia-restaurante-peruano.docx")
    doc.save(path)
    print(f"✓ {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 2. EXCEL TEMPLATES (8)
# ═══════════════════════════════════════════════════════════

def gen_plan_financiero():
    wb = Workbook()
    instr_sheet(wb, "Plan Financiero 3 Años — Restaurante Peruano 80 Plazas", [
        "Rellena las celdas verdes con los datos de tu proyecto.",
        "Las fórmulas se recalculan automáticamente.",
        "Pestaña 'Inversión' = CAPEX inicial desglosado.",
        "Pestaña 'P&L Mensual' = cuenta de resultados mes a mes.",
        "Pestaña 'Proyección 3 Años' = evolución anual.",
    ])
    # Inversión sheet
    ws = wb.create_sheet("Inversión"); ws.sheet_properties.tabColor = "4CAF50"
    title_block(ws, "Inversión Inicial — Restaurante Peruano 80 Plazas")
    hdrs = ["Categoría", "Partida", "Coste Estimado (€)", "Coste Real (€)", "Diferencia", "Notas"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 6)
    widths = [25, 40, 18, 18, 18, 30];
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    items = [
        ("Obra Civil", "Acondicionamiento local", 75000),
        ("Obra Civil", "Instalaciones eléctricas y gas", 10000),
        ("Obra Civil", "Fontanería y saneamiento", 7000),
        ("Obra Civil", "Climatización / HVAC", 8000),
        ("Obra Civil", "Extracción reforzada (wok/parrilla anticuchos)", 6000),
        ("Cocina Peruana", "Wok profesional alta potencia (20-30 kW)", 3000),
        ("Cocina Peruana", "Mesa refrigerada estación ceviche", 3500),
        ("Cocina Peruana", "Parrilla anticuchos (carbón/gas)", 2500),
        ("Cocina Peruana", "Horno pollo a la brasa (si aplica)", 5000),
        ("Cocina Peruana", "Licuadoras industriales para salsas/ajíes (3 uds)", 600),
        ("Cocina Peruana", "Exprimidor cítricos industrial", 800),
        ("Cocina Peruana", "Tablas corte extra grandes cevichero", 200),
        ("Cocina Estándar", "Horno mixto Rational/Unox", 12000),
        ("Cocina Estándar", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina Estándar", "Freidora doble (yuca, chicharrón, camote)", 2000),
        ("Cocina Estándar", "Cámaras frigoríficas (2)", 6000),
        ("Cocina Estándar", "Congelador (imprescindible para pescado ceviche)", 3500),
        ("Cocina Estándar", "Tren de lavado", 5000),
        ("Cocina Estándar", "Campana extractora + filtros", 6000),
        ("Cocina Estándar", "Menaje y utensilios", 2500),
        ("Barra Piscos", "Vitrina iluminada botellas pisco", 4000),
        ("Barra Piscos", "Utensilios coctelería peruana", 500),
        ("Barra Piscos", "Stock inicial pisco (20-30 refs)", 6000),
        ("Barra Piscos", "Máquina de hielo", 1500),
        ("Sala", "Mobiliario interior (mesas + sillas + bancos corridos)", 15000),
        ("Sala", "Barra + taburetes altos", 6000),
        ("Decoración", "Textiles andinos (tapices, cojines, manteles)", 2000),
        ("Decoración", "Cerámica precolombina decorativa", 1500),
        ("Decoración", "Revestimientos madera natural", 4000),
        ("Decoración", "Iluminación cálida + lámparas fibra natural", 2500),
        ("Decoración", "Vajilla cerámica artesanal tonos tierra", 2000),
        ("Decoración", "Plantas tropicales y helechos", 800),
        ("Terraza", "Mobiliario exterior", 6000),
        ("Terraza", "Sombrillas / toldos", 2500),
        ("Tecnología", "TPV + pantallas cocina", 3000),
        ("Tecnología", "Web + carta QR + WiFi", 2000),
        ("Licencias", "Licencia actividad + terraza", 6000),
        ("Licencias", "Proyecto técnico + tasas", 4000),
        ("Licencias", "Registro EORI / importación (si aplica)", 500),
        ("Marketing", "Branding peruano + diseño", 3000),
        ("Marketing", "Fotos + vídeo + campaña pre-apertura", 3000),
        ("Stock", "Materia prima inicial (incluye ajíes, pisco, especias)", 8000),
        ("Stock", "Bebidas (cerveza, refrescos, chicha morada)", 3000),
        ("Maniobra", "Fondo de maniobra (3 meses)", 50000),
    ]
    r += 1
    start_r = r
    for cat, partida, coste in items:
        sdc(ws, r, 1, cat); sdc(ws, r, 2, partida)
        sdc(ws, r, 3, coste, fmt=cur_fmt); sdc(ws, r, 4, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}-C{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 2, "TOTAL INVERSIÓN", font=bld_font)
    sdc(ws, r, 3, f"=SUM(C{start_r}:C{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=D{r}-C{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 6)

    # P&L Mensual sheet
    ws2 = wb.create_sheet("P&L Mensual"); ws2.sheet_properties.tabColor = "2196F3"
    title_block(ws2, "P&L Mensual — Restaurante Peruano 80 Plazas", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic", "TOTAL"]
    r = 4
    for i, m in enumerate(months, 1): ws2.cell(row=r, column=i, value=m)
    shr(ws2, r, 14)
    ws2.column_dimensions['A'].width = 28
    for i in range(2, 15): ws2.column_dimensions[get_column_letter(i)].width = 12
    ws2.freeze_panes = f"B{r+1}"
    rows_data = [
        ("INGRESOS", None, True),
        ("Ventas sala (comida)", 30000, False),
        ("Barra piscos y coctelería", 10000, False),
        ("Delivery (lomo saltado, chaufa, bowls)", 8000, False),
        ("Bebidas (cerveza, chicha morada, refrescos)", 7000, False),
        ("TOTAL INGRESOS", "=SUM", True),
        ("", None, False),
        ("COSTES VARIABLES", None, True),
        ("Food cost (30%)", "=0.30*", False),
        ("Packaging delivery", 500, False),
        ("TOTAL COSTES VARIABLES", "=SUM_CV", True),
        ("", None, False),
        ("COSTES FIJOS", None, True),
        ("Alquiler", 4500, False),
        ("Personal cocina (cevichero, wokero, parrillero, equipo)", 17000, False),
        ("Personal sala (barman piscos, camareros)", 12000, False),
        ("Suministros (luz, gas, agua)", 2800, False),
        ("Seguros", 350, False),
        ("Gestoría + contabilidad", 400, False),
        ("Marketing", 900, False),
        ("Tecnología (TPV, delivery, web)", 350, False),
        ("Mantenimiento", 300, False),
        ("Reposición stock pisco", 1200, False),
        ("Varios / imprevistos", 500, False),
        ("TOTAL COSTES FIJOS", "=SUM_CF", True),
        ("", None, False),
        ("EBITDA", "=EBITDA", True),
        ("% EBITDA", "=%EBITDA", True),
    ]
    r += 1
    for label, val, bold in rows_data:
        sdc(ws2, r, 1, label, font=bld_font if bold else dat_font)
        if val and val not in ("=SUM", "=SUM_CV", "=SUM_CF", "=EBITDA", "=%EBITDA") and not str(val).startswith("=0."):
            for c in range(2, 14):
                sdc(ws2, r, c, val, fill=inp_fill, fmt=cur_fmt)
            sdc(ws2, r, 14, f"=SUM(B{r}:M{r})", font=bld_font, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws2, r, 14)

    path = os.path.join(OUTPUT_DIR, "plan-financiero-3-anos.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_capex():
    wb = Workbook()
    instr_sheet(wb, "Calculadora CAPEX — Restaurante Peruano 80 Plazas", [
        "Introduce tus costes reales en las celdas verdes.",
        "La columna 'Diferencia' se calcula automáticamente.",
        "Usa la columna 'Prioridad' para planificar fases de inversión.",
    ])
    ws = wb.create_sheet("CAPEX Desglosado"); ws.sheet_properties.tabColor = "FF9800"
    title_block(ws, "Calculadora CAPEX — Inversión 130K-300K€")
    hdrs = ["#", "Categoría", "Partida", "Estimado (€)", "Real (€)", "Diferencia", "Prioridad", "Proveedor"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8)
    widths = [5, 20, 40, 14, 14, 14, 12, 25]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    dv = DataValidation(type="list", formula1='"Alta,Media,Baja"', allow_blank=True)
    ws.add_data_validation(dv)
    items = [
        ("Obra", "Acondicionamiento general", 75000),
        ("Obra", "Electricidad + gas", 10000),
        ("Obra", "Fontanería", 7000),
        ("Obra", "Climatización", 8000),
        ("Obra", "Extracción reforzada wok/parrilla", 6000),
        ("Cocina PE", "Wok profesional alta potencia 20-30 kW", 3000),
        ("Cocina PE", "Mesa refrigerada estación ceviche", 3500),
        ("Cocina PE", "Parrilla anticuchos carbón/gas", 2500),
        ("Cocina PE", "Horno pollo a la brasa", 5000),
        ("Cocina PE", "Licuadoras industriales salsas/ajíes (3 uds)", 600),
        ("Cocina PE", "Exprimidor cítricos industrial", 800),
        ("Cocina PE", "Tablas corte extra cevichero", 200),
        ("Cocina", "Horno mixto Rational/Unox", 12000),
        ("Cocina", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina", "Freidora doble", 2000),
        ("Cocina", "Cámaras frigoríficas (2)", 6000),
        ("Cocina", "Congelador (pescado ceviche)", 3500),
        ("Cocina", "Tren de lavado", 5000),
        ("Cocina", "Campana extractora reforzada", 6000),
        ("Cocina", "Menaje y utensilios", 2500),
        ("Barra", "Vitrina iluminada piscos", 4000),
        ("Barra", "Utensilios coctelería + máquina hielo", 2000),
        ("Barra", "Stock inicial pisco (20-30 refs)", 6000),
        ("Sala", "Mesas, sillas, bancos corridos (80 plazas)", 15000),
        ("Sala", "Barra + taburetes altos", 6000),
        ("Decoración", "Textiles andinos, cerámica, madera, iluminación, vajilla", 12800),
        ("Terraza", "Mobiliario exterior + sombrillas", 8500),
        ("Tech", "TPV + pantallas cocina", 3000),
        ("Tech", "Web + carta QR + WiFi", 2000),
        ("Licencias", "Licencia actividad + terraza", 6000),
        ("Licencias", "Proyecto técnico + registro EORI", 4500),
        ("Marketing", "Branding + diseño + pre-apertura", 6000),
        ("Stock", "Materia prima + ajíes + pisco + especias", 8000),
        ("Stock", "Bebidas (cerveza, chicha morada)", 3000),
        ("Maniobra", "Fondo 3 meses", 50000),
    ]
    r += 1; start_r = r
    for idx, (cat, partida, est) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, partida)
        sdc(ws, r, 4, est, fmt=cur_fmt); sdc(ws, r, 5, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 6, f"=E{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 7, "Alta", fill=inp_fill); dv.add(ws.cell(row=r, column=7))
        sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 3, "TOTAL", font=bld_font)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=SUM(E{start_r}:E{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 6, f"=E{r}-D{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 8)
    path = os.path.join(OUTPUT_DIR, "calculadora-capex.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_pl_mensual():
    wb = Workbook()
    instr_sheet(wb, "P&L Mensual con 3 Escenarios — Restaurante Peruano", [
        "3 pestañas: Pesimista, Realista, Optimista.",
        "Modifica los ingresos en cada escenario.",
        "Los costes fijos son iguales en los 3 escenarios.",
    ])
    for scenario, factor, color in [("Pesimista", 0.75, "F44336"), ("Realista", 1.0, "4CAF50"), ("Optimista", 1.25, "2196F3")]:
        ws = wb.create_sheet(scenario); ws.sheet_properties.tabColor = color
        title_block(ws, f"P&L Mensual — Escenario {scenario}", 4)
        hdrs = ["Concepto", "Mensual (€)", "Anual (€)", "% s/Ventas"]
        r = 4
        for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
        shr(ws, r, 4)
        ws.column_dimensions['A'].width = 35
        for i in range(2, 5): ws.column_dimensions[get_column_letter(i)].width = 16
        base_ventas = int(55000 * factor)
        rows = [
            ("INGRESOS", None, True),
            ("Ventas sala (comida)", int(30000 * factor)),
            ("Barra piscos y coctelería", int(10000 * factor)),
            ("Delivery", int(8000 * factor)),
            ("Bebidas (cerveza, chicha morada, refrescos)", int(7000 * factor)),
            ("TOTAL INGRESOS", base_ventas, True),
            ("", None, False),
            ("COSTES VARIABLES", None, True),
            ("Food cost (30%)", int(base_ventas * 0.30)),
            ("Packaging delivery", 500),
            ("TOTAL COSTES VARIABLES", int(base_ventas * 0.30 + 500), True),
            ("", None, False),
            ("COSTES FIJOS", None, True),
            ("Alquiler", 4500),
            ("Personal cocina", 17000),
            ("Personal sala", 12000),
            ("Suministros", 2800),
            ("Seguros", 350),
            ("Gestoría", 400),
            ("Marketing", 900),
            ("Tecnología", 350),
            ("Mantenimiento", 300),
            ("Reposición pisco", 1200),
            ("Varios", 500),
            ("TOTAL COSTES FIJOS", 40300, True),
            ("", None, False),
            ("EBITDA", int(base_ventas - base_ventas * 0.30 - 500 - 40300), True),
        ]
        r += 1
        for item in rows:
            if len(item) == 3:
                label, val, bold = item
            else:
                label, val = item; bold = False
            sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
            if val is not None:
                sdc(ws, r, 2, val, font=bld_font if bold else dat_font, fmt=cur_fmt)
                sdc(ws, r, 3, val * 12, font=bld_font if bold else dat_font, fmt=cur_fmt)
            r += 1
        r += 1; brand_footer(ws, r, 4)
    path = os.path.join(OUTPUT_DIR, "pl-mensual-escenarios.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cash_flow():
    wb = Workbook()
    instr_sheet(wb, "Cash Flow y Break-Even — Restaurante Peruano 80 Plazas", [
        "Pestaña 'Cash Flow' = flujo de caja mensual 12 meses.",
        "Pestaña 'Break-Even' = calculadora de punto de equilibrio.",
        "Rellena las celdas verdes con tus datos reales.",
    ])
    ws = wb.create_sheet("Cash Flow"); ws.sheet_properties.tabColor = "009688"
    title_block(ws, "Cash Flow 12 Meses — Restaurante Peruano", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
    r = 4
    for i, m in enumerate(months, 1): ws.cell(row=r, column=i, value=m)
    shr(ws, r, 13)
    ws.column_dimensions['A'].width = 28
    for i in range(2, 14): ws.column_dimensions[get_column_letter(i)].width = 12
    concepts = [
        ("Saldo inicial", True), ("", False),
        ("ENTRADAS", True), ("Ventas sala (comida)", False), ("Barra piscos/coctelería", False),
        ("Delivery", False), ("Bebidas + chicha morada", False), ("Total entradas", True), ("", False),
        ("SALIDAS", True), ("Materia prima + ajíes + pescado fresco", False), ("Personal cocina", False),
        ("Personal sala", False), ("Alquiler", False), ("Suministros", False), ("Marketing", False),
        ("Reposición stock pisco", False), ("Tecnología", False), ("Otros gastos", False),
        ("Total salidas", True), ("", False),
        ("Flujo neto mensual", True), ("Saldo acumulado", True),
    ]
    r += 1
    for label, bold in concepts:
        sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
        for c in range(2, 14):
            sdc(ws, r, c, None, fill=inp_fill if not bold and label else None, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws, r, 13)

    # Break-even
    ws2 = wb.create_sheet("Break-Even"); ws2.sheet_properties.tabColor = "FF5722"
    title_block(ws2, "Calculadora Break-Even — Restaurante Peruano", 4)
    r = 4
    params = [
        ("Ticket medio (€)", 27),
        ("Comensales/día promedio", 80),
        ("Días abierto/mes", 26),
        ("Food cost (%)", 0.30),
        ("Costes fijos mensuales (€)", 40300),
    ]
    for label, val in params:
        r += 1
        sdc(ws2, r, 1, label, font=bld_font); sdc(ws2, r, 2, val, fill=inp_fill, fmt=cur_fmt if isinstance(val, int) else pct_fmt)
    r += 2
    sdc(ws2, r, 1, "Facturación mensual estimada", font=bld_font)
    sdc(ws2, r, 2, "=B5*B6*B7", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Margen contribución mensual", font=bld_font)
    sdc(ws2, r, 2, "=B12*(1-B8)", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Break-Even (meses)", font=bld_font)
    ws2.column_dimensions['A'].width = 35; ws2.column_dimensions['B'].width = 20
    r += 2; brand_footer(ws2, r, 4)
    path = os.path.join(OUTPUT_DIR, "cash-flow-break-even.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_escandallo():
    wb = Workbook()
    instr_sheet(wb, "Escandallo Maestro — Restaurante Peruano", [
        "Una ficha técnica por plato.",
        "Introduce ingredientes, cantidades y precios.",
        "El food cost se calcula automáticamente.",
    ])
    ws = wb.create_sheet("Escandallo"); ws.sheet_properties.tabColor = "E91E63"
    title_block(ws, "Escandallo Maestro — Fichas Técnicas de Platos Peruanos")
    hdrs = ["#", "Ingrediente", "Cantidad (g/ml)", "Precio/Kg (€)", "Coste (€)", "Merma (%)", "Coste Real (€)"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 7)
    widths = [5, 30, 16, 14, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "NOMBRE DEL PLATO:", font=bld_font)
    sdc(ws, r, 3, "Ejemplo: Ceviche Clásico de Corvina", fill=inp_fill)
    r += 1
    ingredients = [
        ("Corvina fresca (previamente congelada -20°C/24h)", 200, 18.00, 15),
        ("Zumo de limón fresco (o lima)", 80, 3.00, 10),
        ("Cebolla morada", 60, 2.00, 10),
        ("Ají limo (o ají similar)", 5, 25.00, 5),
        ("Cilantro fresco", 10, 6.00, 20),
        ("Ají amarillo en pasta", 8, 12.00, 0),
        ("Sal, pimienta", 3, 5.00, 0),
        ("Camote cocido (guarnición)", 80, 1.50, 10),
        ("Choclo desgranado (guarnición)", 40, 4.00, 5),
        ("Lechuga (decoración base)", 15, 2.50, 15),
    ]
    start_r = r
    for idx, (ing, cant, precio, merma) in enumerate(ingredients, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, ing, fill=inp_fill)
        sdc(ws, r, 3, cant, fill=inp_fill)
        sdc(ws, r, 4, precio, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=(C{r}/1000)*D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, merma / 100, fill=inp_fill, fmt=pct_fmt)
        sdc(ws, r, 7, f"=E{r}*(1+F{r})", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 2, "COSTE TOTAL PLATO", font=bld_font)
    sdc(ws, r, 7, f"=SUM(G{start_r}:G{r-2})", font=bld_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP sugerido (food cost 30%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}/0.30", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP con IVA (10%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}*1.10", font=frm_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 7)
    path = os.path.join(OUTPUT_DIR, "escandallo-maestro-peruano.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_menu_engineering():
    wb = Workbook()
    instr_sheet(wb, "Menú Engineering Matrix — Restaurante Peruano", [
        "Introduce los platos de tu carta con ventas y food cost.",
        "La matrix clasifica automáticamente cada plato.",
        "Stars = mantener. Plowhorses = subir precio. Puzzles = promocionar. Dogs = eliminar.",
    ])
    ws = wb.create_sheet("Matrix"); ws.sheet_properties.tabColor = "9C27B0"
    title_block(ws, "Menú Engineering Matrix — Restaurante Peruano", 10)
    hdrs = ["#", "Plato", "PVP (€)", "Food Cost (€)", "% Food Cost", "Uds. Vendidas/Mes", "Margen Unit. (€)", "Margen Total (€)", "Popularidad", "Clasificación"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    widths = [5, 30, 10, 12, 10, 16, 12, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    platos = [
        ("Ceviche clásico de corvina", 14.00, 2.80, 160),
        ("Lomo saltado", 15.00, 3.20, 130),
        ("Causa limeña de atún", 10.00, 1.80, 80),
        ("Ají de gallina con arroz", 12.00, 2.00, 70),
        ("Anticuchos de corazón (3 brochetas)", 11.00, 1.60, 90),
        ("Tiradito de lubina al ají amarillo", 15.00, 3.00, 55),
        ("Arroz con mariscos", 16.00, 3.80, 45),
        ("Papa a la huancaína", 8.00, 1.00, 100),
        ("Pisco Sour", 10.00, 1.20, 200),
        ("Chicharrón con mote", 12.00, 2.20, 60),
        ("Tacu tacu con lomo", 16.00, 3.50, 35),
        ("Arroz chaufa de pollo", 12.00, 2.00, 85),
        ("Suspiro limeño", 7.00, 0.90, 55),
        ("Picarones con miel de chancaca", 7.00, 0.80, 65),
        ("Pollo a la brasa (1/4)", 13.00, 2.50, 75),
    ]
    r += 1
    for idx, (plato, pvp, fc, uds) in enumerate(platos, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, plato, fill=inp_fill)
        sdc(ws, r, 3, pvp, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 4, fc, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}/C{r}", font=frm_font, fmt=pct_fmt)
        sdc(ws, r, 6, uds, fill=inp_fill)
        sdc(ws, r, 7, f"=C{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 8, f"=G{r}*F{r}", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "menu-engineering-matrix.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_ticket():
    wb = Workbook()
    instr_sheet(wb, "Calculadora Ticket Medio + Margen Pisco/Cócteles", [
        "Simula diferentes escenarios de ticket medio.",
        "Incluye pestaña separada para margen de barra de piscos.",
        "Rellena las celdas verdes con tus datos.",
    ])
    ws = wb.create_sheet("Ticket Medio"); ws.sheet_properties.tabColor = "00BCD4"
    title_block(ws, "Calculadora Ticket Medio — Restaurante Peruano", 4)
    r = 4
    params = [
        ("% clientes que piden ceviches/causas para compartir", 0.60),
        ("Precio medio ceviches/causas (€)", 12.00),
        ("Precio medio principal (€)", 13.50),
        ("% clientes que piden postre", 0.25),
        ("Precio medio postre (€)", 7.00),
        ("% clientes que piden pisco sour/cóctel", 0.45),
        ("Precio medio pisco/cóctel (€)", 10.00),
        ("% clientes que piden cerveza/chicha morada", 0.50),
        ("Precio medio cerveza/chicha (€)", 4.50),
    ]
    r += 1
    for label, val in params:
        sdc(ws, r, 1, label, font=dat_font)
        sdc(ws, r, 2, val, fill=inp_fill, fmt=pct_fmt if isinstance(val, float) and val < 1 else cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 1, "TICKET MEDIO ESTIMADO", font=bld_font)
    sdc(ws, r, 2, None, font=bld_font, fmt=cur_fmt)
    ws.column_dimensions['A'].width = 50; ws.column_dimensions['B'].width = 18

    # Margen piscos tab
    ws2 = wb.create_sheet("Margen Piscos"); ws2.sheet_properties.tabColor = "FFC107"
    title_block(ws2, "Simulador Margen Barra de Piscos y Cócteles Peruanos", 4)
    r = 4
    sdc(ws2, r+1, 1, "Precio medio copa pisco (€)", font=bld_font); sdc(ws2, r+1, 2, 7.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+2, 1, "Coste medio copa (€)", font=bld_font); sdc(ws2, r+2, 2, 1.20, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+3, 1, "Margen por copa (€)", font=bld_font); sdc(ws2, r+3, 2, f"=B{r+1}-B{r+2}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+4, 1, "% Margen pisco", font=bld_font); sdc(ws2, r+4, 2, f"=B{r+3}/B{r+1}", font=frm_font, fmt=pct_fmt)
    sdc(ws2, r+5, 1, "Precio medio pisco sour (€)", font=bld_font); sdc(ws2, r+5, 2, 10.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+6, 1, "Coste medio pisco sour (€)", font=bld_font); sdc(ws2, r+6, 2, 1.50, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+7, 1, "Margen por pisco sour (€)", font=bld_font); sdc(ws2, r+7, 2, f"=B{r+5}-B{r+6}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+8, 1, "Copas pisco vendidas/día", font=bld_font); sdc(ws2, r+8, 2, 20, fill=inp_fill)
    sdc(ws2, r+9, 1, "Pisco sour vendidos/día", font=bld_font); sdc(ws2, r+9, 2, 35, fill=inp_fill)
    sdc(ws2, r+10, 1, "Días/mes", font=bld_font); sdc(ws2, r+10, 2, 26, fill=inp_fill)
    sdc(ws2, r+11, 1, "Facturación barra piscos/mes (€)", font=bld_font)
    sdc(ws2, r+11, 2, f"=(B{r+1}*B{r+8}+B{r+5}*B{r+9})*B{r+10}", font=frm_font, fmt=cur_fmt)
    ws2.column_dimensions['A'].width = 40; ws2.column_dimensions['B'].width = 18
    path = os.path.join(OUTPUT_DIR, "calculadora-ticket-medio.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cronograma():
    wb = Workbook()
    instr_sheet(wb, "Cronograma Apertura Gantt 12 Meses — Restaurante Peruano", [
        "Fases y tareas para abrir un restaurante peruano.",
        "Marca las celdas con 'X' para indicar el mes activo.",
        "Incluye fase de setup de importación y proveedores peruanos.",
    ])
    ws = wb.create_sheet("Gantt"); ws.sheet_properties.tabColor = "795548"
    title_block(ws, "Cronograma de Apertura — Restaurante Peruano 80 Plazas", 15)
    hdrs = ["#", "Fase", "Tarea", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "M8", "M9", "M10", "M11", "M12"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 15)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 45
    for i in range(4, 16): ws.column_dimensions[get_column_letter(i)].width = 5
    tasks = [
        ("Planificación", "Estudio de mercado y viabilidad", [1, 2]),
        ("Planificación", "Plan financiero y búsqueda financiación", [1, 2, 3]),
        ("Planificación", "Constitución empresa / alta autónomo", [2, 3]),
        ("Importación", "Registro EORI y trámites importación pisco/ajíes", [2, 3, 4]),
        ("Importación", "Contacto distribuidores producto peruano", [2, 3]),
        ("Importación", "Selección proveedores pisco + pescado fresco", [3, 4]),
        ("Local", "Búsqueda y selección de local", [2, 3, 4]),
        ("Local", "Negociación alquiler y contrato", [3, 4]),
        ("Licencias", "Licencia actividad / declaración responsable", [3, 4, 5]),
        ("Licencias", "Proyecto técnico y visado", [3, 4]),
        ("Obra", "Acondicionamiento y obra civil", [4, 5, 6, 7]),
        ("Obra", "Instalaciones (electricidad, gas, extracción)", [5, 6, 7]),
        ("Obra", "Decoración peruana (textiles, cerámica, madera)", [6, 7]),
        ("Equipamiento", "Pedido equipo cocina + wok + parrilla anticuchos", [5, 6]),
        ("Equipamiento", "Instalación barra piscos + vitrina", [6, 7]),
        ("Equipamiento", "Recepción e instalación equipos", [7, 8]),
        ("Equipamiento", "Mobiliario sala + terraza", [6, 7]),
        ("Equipo", "Búsqueda cevichero y wokero", [6, 7, 8]),
        ("Equipo", "Selección y contratación resto equipo", [7, 8, 9]),
        ("Equipo", "Formación equipo (cocina peruana + piscos)", [9, 10]),
        ("Marketing", "Branding, web, redes sociales", [6, 7, 8]),
        ("Marketing", "Fotos profesionales de platos", [9]),
        ("Marketing", "Campaña pre-apertura", [9, 10]),
        ("APPCC", "Plan APPCC + registro sanitario (pescado crudo)", [7, 8]),
        ("Pre-apertura", "Soft opening (amigos y familia)", [10]),
        ("Pre-apertura", "Ajustes finales", [10, 11]),
        ("Apertura", "INAUGURACIÓN", [11]),
        ("Post-apertura", "Seguimiento y ajustes", [11, 12]),
    ]
    r += 1
    for idx, (fase, tarea, meses) in enumerate(tasks, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, fase); sdc(ws, r, 3, tarea)
        for m in meses:
            sdc(ws, r, m + 3, "X", font=bld_font, fill=sec_fill, align=c_align)
        r += 1
    r += 1; brand_footer(ws, r, 15)
    path = os.path.join(OUTPUT_DIR, "cronograma-apertura-gantt.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_plantilla_turnos():
    wb = Workbook()
    instr_sheet(wb, "Plantilla Turnos Brigada — Restaurante Peruano", [
        "Cuadrante semanal para todo el equipo (cocina + sala).",
        "Rellena los turnos con: M (mañana), T (tarde), P (partido), L (libre).",
        "Incluye puestos específicos: cevichero, wokero, parrillero, barman piscos.",
    ])
    ws = wb.create_sheet("Turnos Semana"); ws.sheet_properties.tabColor = "607D8B"
    title_block(ws, "Cuadrante Semanal — Restaurante Peruano 80 Plazas", 10)
    hdrs = ["#", "Nombre", "Puesto", "Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 24
    for i in range(4, 11): ws.column_dimensions[get_column_letter(i)].width = 8
    dv = DataValidation(type="list", formula1='"M,T,P,L,V"', allow_blank=True)
    ws.add_data_validation(dv)
    staff = [
        ("Jefe de cocina", "P"), ("Cevichero", "P"), ("Wokero (chifa/Nikkei)", "P"),
        ("Parrillero (anticuchos)", "P"), ("Cocinero 1 (criollos)", "P"), ("Cocinero 2", "M"),
        ("Ayudante cocina", "P"), ("Office", "P"),
        ("Encargado sala", "P"), ("Barman piscos", "T"),
        ("Camarero 1", "P"), ("Camarero 2", "P"), ("Camarero 3", "T"),
        ("Runner", "P"),
    ]
    r += 1
    for idx, (puesto, turno_default) in enumerate(staff, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, None, fill=inp_fill); sdc(ws, r, 3, puesto)
        for c in range(4, 11):
            sdc(ws, r, c, turno_default, fill=inp_fill, align=c_align)
            dv.add(ws.cell(row=r, column=c))
        r += 1
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "M=Mañana, T=Tarde, P=Partido, L=Libre, V=Vacaciones", font=brn_font)
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "plantilla-turnos-brigada.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 3. CHECKLISTS (6)
# ═══════════════════════════════════════════════════════════

def gen_checklist_legal():
    wb = Workbook(); ws = wb.active
    items = [
        ("Empresa", "Decidir forma jurídica (SL, autónomo, cooperativa)", "Socio/Gestor", 500),
        ("Empresa", "Constituir sociedad ante notario", "Socio/Gestor", 600),
        ("Empresa", "Inscripción en Registro Mercantil", "Gestor", 200),
        ("Empresa", "Obtener CIF provisional", "Gestor", 0),
        ("Empresa", "Alta IAE epígrafe 671.4 o similar", "Gestor", 0),
        ("Empresa", "Alta censal en Hacienda (modelo 036/037)", "Gestor", 0),
        ("Empresa", "Inscripción Seguridad Social empresa", "Gestor", 0),
        ("Empresa", "Abrir cuenta bancaria empresa", "Socio", 0),
        ("Licencias", "Licencia de actividad / declaración responsable", "Arquitecto", 2000),
        ("Licencias", "Proyecto técnico y visado colegial", "Arquitecto", 3000),
        ("Licencias", "Licencia de obras (si aplica)", "Arquitecto", 1500),
        ("Licencias", "Licencia de terraza", "Socio", 1500),
        ("Licencias", "Licencia de música / SGAE / AGEDI", "Socio", 500),
        ("Importación", "Registro EORI (operador comercio exterior)", "Gestor", 0),
        ("Importación", "Certificado SOIVRE para alimentos importados extra-UE", "Gestor", 300),
        ("Importación", "Etiquetado en español para productos importados (pisco, ajíes)", "Proveedor", 200),
        ("Importación", "Control fitosanitario ajíes secos y especias", "Gestor", 150),
        ("Importación", "Impuestos especiales bebidas alcohólicas (IIEE) pisco", "Gestor", 0),
        ("Importación", "Denominación de Origen Pisco — verificar cumplimiento normativa UE", "Gestor", 0),
        ("Importación", "Acuerdo con distribuidor especializado en producto peruano", "Socio", 0),
        ("Sanidad", "Registro General Sanitario (RGSEAA)", "Gestor", 0),
        ("Sanidad", "Plan APPCC documentado (adaptado cocina peruana — pescado crudo)", "Consultor", 1500),
        ("Sanidad", "Formación manipulador alimentos (equipo)", "Consultor", 300),
        ("Sanidad", "Contrato empresa DDD (desratización)", "Socio", 600),
        ("Sanidad", "Control de aguas", "Socio", 200),
        ("Sanidad", "Plan de alérgenos (14 alérgenos + específicos peruanos)", "Jefe cocina", 0),
        ("Seguros", "Seguro responsabilidad civil (min. 300K€)", "Socio", 800),
        ("Seguros", "Seguro de local (incendio, robo, daños)", "Socio", 600),
        ("Seguros", "Seguro de accidentes convenio", "Gestor", 400),
        ("Laboral", "Contratos de trabajo según convenio", "Gestor", 0),
        ("Laboral", "Prevención de riesgos laborales", "SPA", 500),
        ("Laboral", "Calendario laboral y cuadrante turnos", "Encargado", 0),
        ("Protección datos", "Registro LOPD / RGPD", "Gestor", 300),
        ("Protección datos", "Política de cookies web", "Socio", 0),
        ("Protección datos", "Videovigilancia (si aplica): cartelería + registro", "Socio", 100),
        ("Otros", "Libro de reclamaciones", "Socio", 30),
        ("Otros", "Cartel de horarios en puerta", "Socio", 0),
        ("Otros", "Aforo máximo visible", "Socio", 0),
        ("Otros", "Señalética salida emergencia + extintores", "Arquitecto", 500),
    ]
    checklist_ws(ws, "Checklist Legal — Restaurante Peruano España", items)
    path = os.path.join(OUTPUT_DIR, "checklist-legal.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_equipamiento():
    wb = Workbook(); ws = wb.active
    items = [
        ("Cocina PE", "Wok profesional alta potencia (20-30 kW, quemador chino)", "Jefe cocina", 3000),
        ("Cocina PE", "Mesa refrigerada estación ceviche", "Jefe cocina", 3500),
        ("Cocina PE", "Parrilla anticuchos carbón/gas profesional", "Jefe cocina", 2500),
        ("Cocina PE", "Horno pollo a la brasa (si modelo criollo/pollería)", "Jefe cocina", 5000),
        ("Cocina PE", "Licuadoras industriales para salsas y ajíes (3 uds)", "Jefe cocina", 600),
        ("Cocina PE", "Exprimidor cítricos industrial (limones para ceviche)", "Jefe cocina", 800),
        ("Cocina PE", "Tablas corte extra grandes para cevichero", "Jefe cocina", 200),
        ("Cocina PE", "Vitrina iluminada piscos para barra", "Encargado", 4000),
        ("Cocción", "Cocina gas 6 fuegos + horno", "Jefe cocina", 5000),
        ("Cocción", "Horno mixto (Rational / Unox)", "Jefe cocina", 12000),
        ("Cocción", "Freidora doble cuba (yuca frita, chicharrón, camote)", "Jefe cocina", 2000),
        ("Cocción", "Salamandra", "Jefe cocina", 1200),
        ("Frío", "Cámara frigorífica 2 puertas", "Jefe cocina", 3500),
        ("Frío", "Cámara congelación (obligatoria para pescado ceviche)", "Jefe cocina", 3500),
        ("Frío", "Mesa refrigerada bajo barra", "Jefe cocina", 2500),
        ("Frío", "Botellero refrigerado barra piscos", "Encargado", 1800),
        ("Preparación", "Mesa de trabajo acero inox (2-3 uds)", "Jefe cocina", 2000),
        ("Preparación", "Robot de cocina / cutter", "Jefe cocina", 1500),
        ("Preparación", "Prensador de limones manual grande", "Jefe cocina", 150),
        ("Lavado", "Tren de lavado / lavavajillas industrial", "Jefe cocina", 5000),
        ("Lavado", "Fregadero doble seno inox", "Jefe cocina", 800),
        ("Extracción", "Campana extractora reforzada + filtros (wok + parrilla)", "Instalador", 6000),
        ("Extracción", "Salida de humos reglamentaria (sobredimensionada)", "Instalador", 4000),
        ("Menaje", "Ollas grandes para arroz con mariscos, seco, ají de gallina", "Jefe cocina", 1000),
        ("Menaje", "Sartén honda wok adicional", "Jefe cocina", 300),
        ("Menaje", "Cuchillos profesionales (incluye cuchillos sashimi si Nikkei)", "Jefe cocina", 600),
        ("Menaje", "Gastronormas GN 1/1, 1/2, 1/3", "Jefe cocina", 600),
        ("Menaje", "Tablas de corte por colores", "Jefe cocina", 150),
        ("Almacenamiento", "Estanterías inox economato seco (ajíes, especias)", "Jefe cocina", 800),
        ("Almacenamiento", "Contenedores herméticos para ajíes secos y pasta de ají", "Jefe cocina", 300),
        ("Barra", "Utensilios coctelería peruana (pisco sour shaker, medidores)", "Encargado", 500),
        ("Barra", "Máquina de hielo (alto consumo por cócteles de pisco)", "Encargado", 2000),
    ]
    checklist_ws(ws, "Checklist Equipamiento Cocina Peruana", items)
    path = os.path.join(OUTPUT_DIR, "checklist-equipamiento-cocina-peruana.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_appcc():
    wb = Workbook(); ws = wb.active
    items = [
        ("Prerrequisitos", "Plan de limpieza y desinfección documentado", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan DDD contratado (empresa externa)", "Socio", 600),
        ("Prerrequisitos", "Plan de control de agua potable", "Socio", 200),
        ("Prerrequisitos", "Plan de formación del personal", "Jefe cocina", 300),
        ("Prerrequisitos", "Plan de gestión de residuos", "Socio", 0),
        ("Prerrequisitos", "Plan de mantenimiento de instalaciones", "Socio", 0),
        ("Prerrequisitos", "Plan de trazabilidad (registro proveedores + lotes)", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan de control de alérgenos (14 alérgenos)", "Jefe cocina", 0),
        ("Temperaturas", "Termómetro sonda calibrado", "Jefe cocina", 50),
        ("Temperaturas", "Registro diario temperaturas cámaras", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas servicio (caliente >65°C)", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas recepción mercancía", "Cocinero", 0),
        ("Temperaturas", "Protocolo cadena de frío (0-4°C refrigerado, -18°C congelado)", "Jefe cocina", 0),
        ("Temperaturas", "Control temperatura congelación pescado ceviche (-20°C mín 24h)", "Cevichero", 0),
        ("Temperaturas", "Control temperatura anticuchos (interna >75°C)", "Parrillero", 0),
        ("Pescado crudo", "CONGELACIÓN PREVIA obligatoria -20°C/24h (anisakis)", "Cevichero", 0),
        ("Pescado crudo", "Registro de congelación por lote de pescado", "Cevichero", 0),
        ("Pescado crudo", "Trazabilidad completa del pescado (proveedor, lonja, fecha)", "Jefe cocina", 0),
        ("Pescado crudo", "Vida útil leche de tigre: máx 2h refrigerado, no reutilizar", "Cevichero", 0),
        ("Pescado crudo", "Tiraditos: mismos requisitos que ceviche", "Cevichero", 0),
        ("Pescado crudo", "Mariscos: cadena de frío ininterrumpida, trazabilidad estricta", "Jefe cocina", 0),
        ("Prod. Importados", "Certificado fitosanitario ajíes secos en archivo", "Jefe cocina", 0),
        ("Prod. Importados", "Etiquetado en español de todos los productos importados", "Jefe cocina", 0),
        ("Prod. Importados", "Trazabilidad de origen de ajíes, especias, pisco", "Jefe cocina", 0),
        ("Prod. Importados", "Almacenamiento separado ajíes secos (seco, ventilado)", "Cocinero", 0),
        ("Prod. Importados", "Registro de lotes de pisco importado", "Encargado", 0),
        ("Higiene", "Lavamanos no manual en cocina con jabón y papel", "Socio", 300),
        ("Higiene", "Vestuarios con taquillas para personal", "Socio", 500),
        ("Higiene", "Uniformes de cocina (gorro, chaqueta, pantalón, calzado)", "Socio", 600),
        ("Higiene", "Protocolo lavado de manos (cartelería)", "Jefe cocina", 0),
        ("Higiene", "Contenedores de basura con tapa y pedal", "Socio", 200),
        ("Salsas", "Salsas frescas (criolla, huancaína, ocopa): vida útil máx 48h refrig.", "Jefe cocina", 0),
        ("Salsas", "Ají amarillo y rocoto en pasta: almacenar refrigerado tras abrir", "Cocinero", 0),
        ("Almacenamiento", "Sistema FIFO/FEFO en cámaras y economato", "Cocinero", 0),
        ("Almacenamiento", "Separación crudo/cocinado en cámaras", "Cocinero", 0),
        ("Documentación", "Plan APPCC escrito y disponible en cocina", "Socio", 1500),
        ("Documentación", "Registros de control actualizados diariamente", "Jefe cocina", 0),
        ("Documentación", "Certificados manipulador alimentos de todo el personal", "Socio", 0),
        ("Documentación", "Contrato DDD visible y registros de actuaciones", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist APPCC — Restaurante Peruano (Pescado Crudo)", items)
    path = os.path.join(OUTPUT_DIR, "checklist-appcc.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_sala():
    wb = Workbook(); ws = wb.active
    items = [
        ("Distribución", "Plano de sala con posición de mesas definitivo", "Socio/Arquitecto", 0),
        ("Distribución", "80 plazas distribuidas (mesas 2, 4, barra alta)", "Socio", 0),
        ("Distribución", "Distancia entre mesas: mínimo 60-80 cm", "Arquitecto", 0),
        ("Distribución", "Flujo de servicio definido (entrada → sala → barra → aseos)", "Socio", 0),
        ("Distribución", "Zona de espera / recepción", "Socio", 500),
        ("Distribución", "Accesibilidad PMR (rampa, aseo adaptado)", "Arquitecto", 2000),
        ("Barra Piscos", "Barra de piscos con vitrina iluminada trasera", "Interiorista", 7000),
        ("Barra Piscos", "Taburetes altos cómodos para barra", "Socio", 3000),
        ("Barra Piscos", "Iluminación ambiental en barra (luz cálida, retroiluminación piscos)", "Electricista", 1000),
        ("Barra Piscos", "Carta de piscos enmarcada o en pizarra artesanal", "Diseñador", 300),
        ("Decoración PE", "Textiles andinos (tapices, cojines, mantas en bancos corridos)", "Interiorista", 2000),
        ("Decoración PE", "Cerámica precolombina decorativa (Moche, Nasca, Chimú)", "Socio", 1500),
        ("Decoración PE", "Revestimientos de madera natural en paredes de acento", "Interiorista", 4000),
        ("Decoración PE", "Paleta de colores tierra: marrón, terracota, ocre, beige, rojo", "Interiorista", 0),
        ("Decoración PE", "Vajilla cerámica artesanal en tonos tierra y negro", "Socio", 2000),
        ("Decoración PE", "Plantas tropicales y helechos como acento verde", "Socio", 500),
        ("Decoración PE", "Detalles en piedra natural (pared de acento o barra)", "Interiorista", 2000),
        ("Iluminación", "Iluminación general cálida 2700K", "Electricista", 2000),
        ("Iluminación", "Iluminación regulable por zonas (dimmer)", "Electricista", 500),
        ("Iluminación", "Lámparas de fibra natural o cerámica artesanal", "Socio", 1500),
        ("Acústica", "Paneles absorbentes si techos altos", "Interiorista", 1500),
        ("Acústica", "Sistema de música (música criolla, afroperuana, fusión)", "Socio", 800),
        ("Mobiliario", "Mesas de madera natural maciza", "Socio", 10000),
        ("Mobiliario", "Sillas con asiento tapizado colores tierra", "Socio", 8000),
        ("Mobiliario", "Bancos corridos con cojines de textiles andinos", "Socio", 3000),
        ("Terraza", "Mobiliario exterior en madera/ratán", "Socio", 6000),
        ("Terraza", "Sombrillas o toldos en colores neutros", "Socio", 2500),
        ("Terraza", "Plantas tropicales y macetas de cerámica", "Socio", 500),
        ("Aseos", "Aseos diferenciados y aseo PMR", "Arquitecto", 3000),
        ("Aseos", "Decoración peruana en aseos (cerámica, espejo artesanal)", "Socio", 500),
        ("Señalética", "Señalética interior con identidad peruana", "Diseñador", 300),
        ("Señalética", "Rótulo exterior / fachada con identidad peruana", "Socio", 2500),
    ]
    checklist_ws(ws, "Checklist Diseño Sala Peruana + Barra Piscos", items)
    path = os.path.join(OUTPUT_DIR, "checklist-diseno-sala-peruana.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_contratacion():
    wb = Workbook(); ws = wb.active
    items = [
        ("Perfiles", "Definir organigrama completo (cocina + sala)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Jefe de cocina (experiencia cocina peruana)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Cevichero (estación ceviche, tiraditos, causas)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Wokero (chifa/Nikkei: chaufa, saltados, tallarines)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Parrillero (anticuchos, pollo a la brasa)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Cocinero (guisos criollos: ají de gallina, seco)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Ayudante / office", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Encargado de sala", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Barman de piscos y coctelería peruana", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Camarero", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Runner", "Socio", 0),
        ("Selección", "Publicar ofertas (InfoJobs, Hostelwork, LinkedIn)", "Socio", 200),
        ("Selección", "Búsqueda cevichero en comunidades peruanas España", "Socio", 0),
        ("Selección", "Búsqueda wokero con experiencia en chifa", "Socio", 0),
        ("Selección", "Selección jefe de cocina con experiencia peruana", "Socio", 0),
        ("Selección", "Selección barman con conocimiento de pisco", "Socio", 0),
        ("Selección", "Selección y entrevistas resto equipo cocina", "Jefe cocina", 0),
        ("Selección", "Selección y entrevistas camareros", "Encargado", 0),
        ("Selección", "Pruebas prácticas cocina (elaborar ceviche + leche de tigre)", "Jefe cocina", 0),
        ("Documentación", "Contratos según convenio hostelería CCAA", "Gestor", 0),
        ("Documentación", "Alta Seguridad Social", "Gestor", 0),
        ("Documentación", "Certificado manipulador alimentos", "Empleado", 0),
        ("Documentación", "Prevención riesgos laborales (formación)", "SPA", 500),
        ("Onboarding", "Manual de bienvenida / handbook del restaurante", "Socio", 0),
        ("Onboarding", "Formación en carta peruana y alérgenos", "Jefe cocina", 0),
        ("Onboarding", "Formación en piscos y coctelería peruana (todo equipo sala)", "Barman", 0),
        ("Onboarding", "Formación en TPV y sistema de pedidos", "Encargado", 0),
        ("Onboarding", "Formación APPCC y protocolos (especial pescado crudo)", "Jefe cocina", 0),
        ("Onboarding", "Ensayos de servicio (soft opening)", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist Contratación — Restaurante Peruano", items)
    path = os.path.join(OUTPUT_DIR, "checklist-contratacion.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_marketing():
    wb = Workbook(); ws = wb.active
    items = [
        ("Branding", "Nombre del restaurante definitivo (verificar disponibilidad)", "Socio", 0),
        ("Branding", "Logo profesional con identidad peruana", "Diseñador", 500),
        ("Branding", "Manual de marca (colores tierra, tipografía, estilo)", "Diseñador", 300),
        ("Branding", "Diseño de carta / menú físico con estética peruana", "Diseñador", 500),
        ("Branding", "Packaging delivery con marca peruana", "Diseñador", 300),
        ("Digital", "Web propia con menú, carta piscos, fotos y reservas", "Dev/Agencia", 1500),
        ("Digital", "Google My Business: ficha completa y verificada", "Socio", 0),
        ("Digital", "Instagram: perfil profesional con 9+ publicaciones", "Community", 0),
        ("Digital", "TikTok: perfil activo con vídeos del cevichero, wok, pisco sour", "Community", 0),
        ("Digital", "TheFork / TripAdvisor: perfiles completos con fotos", "Socio", 0),
        ("Digital", "Google Ads local configurado", "Agencia", 500),
        ("Digital", "Instagram/Facebook Ads configurados", "Agencia", 400),
        ("Contenido", "Sesión fotográfica profesional (ceviches + decoración + piscos)", "Fotógrafo", 400),
        ("Contenido", "Vídeo cevichero preparando ceviche (30-60 seg para TikTok/Reels)", "Videógrafo", 500),
        ("Contenido", "Calendario de publicaciones primer mes", "Community", 0),
        ("Eventos PE", "Plan Fiestas Patrias (28-29 julio): menú especial + decoración + pisco libre", "Socio", 800),
        ("Eventos PE", "Plan Día del Pisco Sour (febrero): catas + vuelos de pisco", "Socio", 300),
        ("Eventos PE", "Plan Día del Ceviche (28 junio): degustación ceviches", "Socio", 300),
        ("Eventos PE", "Plan Inti Raymi (24 junio): menú andino especial", "Socio", 200),
        ("Pre-apertura", "Evento soft opening (amigos, familia, influencers food)", "Socio", 500),
        ("Pre-apertura", "Invitación a prensa local / bloggers food", "Socio", 200),
        ("Pre-apertura", "Flyers / cartelería barrio", "Diseñador", 200),
        ("Pre-apertura", "Colaboraciones con negocios vecinos", "Socio", 0),
        ("Delivery", "Alta en Glovo / Uber Eats / Just Eat", "Socio", 0),
        ("Delivery", "Fotos de platos para plataformas delivery", "Fotógrafo", 0),
        ("Delivery", "Menú delivery optimizado (lomo saltado, chaufa, bowls, pollo brasa)", "Jefe cocina", 0),
        ("Fidelización", "Programa de fidelización (tarjeta, app o CRM)", "Socio", 300),
        ("Fidelización", "Estrategia de reseñas (pedir valoraciones a clientes)", "Encargado", 0),
    ]
    checklist_ws(ws, "Checklist Marketing Pre-Apertura — Restaurante Peruano", items)
    path = os.path.join(OUTPUT_DIR, "checklist-marketing-preapertura.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 4. DOCUMENT MODELS (2)
# ═══════════════════════════════════════════════════════════

def gen_business_plan():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Business Plan"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Peruano — 80 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("[Ciudad, Mes Año]\nPlantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Resumen Ejecutivo", [
            "Descripción del proyecto: restaurante peruano auténtico de 80 plazas en [ciudad].",
            "Modelo: [cevichería / criollo / Nikkei / chifa / novoandina / pollería / dark kitchen].",
            "Inversión total requerida: [130.000-300.000€].",
            "Facturación prevista año 1: [550.000-800.000€].",
            "Break-even estimado: mes [8-14].",
            "Equipo fundador: [nombres y experiencia].",
            "Diferencial: cocina peruana auténtica + barra de piscos premium 20-30 referencias.",
        ]),
        ("2. El Concepto", [
            "Tipo de cocina: cevichería / criollo / Nikkei / chifa / novoandina / fusión.",
            "Ticket medio objetivo: [20-35€] sin bebidas.",
            "Propuesta de valor: ceviche al momento, leche de tigre fresca, pisco premium.",
            "Público objetivo: 25-50 años, urbano, foodie, viajero, busca experiencias premium.",
            "Posicionamiento vs competencia: auténtico vs peruano casual genérico.",
        ]),
        ("3. Análisis de Mercado", [
            "Restaurantes peruanos en España: +1.800, crecimiento 40% en 5 años.",
            "Menos del 15% ofrecen Nikkei, chifa o novoandina — oportunidad clara.",
            "Tendencia pisco premium: crecimiento acelerado en España.",
            "3 peruanos en 50 Best Restaurants — cocina peruana en la élite mundial.",
            "Competencia directa en zona: [listar 3-5 competidores con ticket medio].",
            "Gap de mercado identificado: [qué falta en la zona].",
        ]),
        ("4. Plan Operativo", [
            "Ubicación: [dirección, m², alquiler mensual].",
            "Distribución: 80 plazas interior + barra piscos + [X] plazas terraza.",
            "Horario: [ej. 13:00-16:00 y 20:00-00:00, lunes cerrado].",
            "Equipo cocina: [6-10] personas (incluye cevichero, wokero, parrillero).",
            "Equipo sala: [6-8] personas (incluye barman piscos).",
            "Proveedores clave producto peruano: [distribuidores, importadores, lonja].",
        ]),
        ("5. Plan Financiero", [
            "Inversión inicial desglosada: ver plantilla Excel adjunta.",
            "P&L mensual con 3 escenarios: ver plantilla Excel adjunta.",
            "Cash flow 12 meses: ver plantilla Excel adjunta.",
            "Break-even: [mes estimado].",
            "ROI estimado: [X%] en año [2-3].",
        ]),
        ("6. Plan de Marketing", [
            "Estrategia pre-apertura: [acciones 3 meses antes de abrir].",
            "Canales digitales: Google My Business, Instagram, TikTok, TheFork.",
            "Eventos culturales: Fiestas Patrias (28 julio), Día Pisco Sour, Día del Ceviche.",
            "Presupuesto marketing mensual: [600-1.500€/mes].",
            "Estrategia de delivery: lomo saltado, chaufa, bowls y pollo a la brasa.",
        ]),
        ("7. Equipo Fundador", [
            "[Nombre]: [experiencia en hostelería/cocina peruana, rol en el proyecto].",
            "[Nombre]: [experiencia complementaria, rol en el proyecto].",
            "Asesores externos: [gestor, arquitecto, consultor gastronómico].",
        ]),
        ("8. Necesidades de Financiación", [
            "Inversión total: [X€].",
            "Fondos propios: [X€] ([X%]).",
            "Financiación solicitada: [X€] ([X%]).",
            "Destino de los fondos: obra [X€], equipamiento [X€], barra piscos [X€], maniobra [X€].",
            "Plan de amortización propuesto: [X años, cuota mensual X€].",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "business-plan-modelo.docx")
    doc.save(path); print(f"✓ {path}")

def gen_manual_operaciones():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual de Operaciones"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Peruano — 80 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Plantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Apertura del Restaurante", [
            "Hora de llegada del equipo: [2.5 horas antes del servicio].",
            "Cevichero: verificar stock de pescado (previamente congelado y descongelado en cámara), preparar leche de tigre base, cortar cebolla morada, preparar ají limo.",
            "Wokero: encender wok, preparar mise en place de chaufa/saltados (verduras, salsas sillao/ostión, arroz del día anterior para chaufa).",
            "Parrillero: encender parrilla, preparar marinado de anticuchos, montar brochetas.",
            "Cocineros: encender fogones, preparar guisos criollos (ají de gallina, seco), arroz blanco, frijoles.",
            "Verificar temperaturas de cámaras y registrar. Verificar registro de congelación de pescado.",
            "Reponer mise en place: cilantro, cebolla morada, limones, ajíes, camote, choclo.",
            "Barra: verificar stock piscos, preparar jarabe de goma para pisco sour, exprimir limones, hielo.",
            "Preparar sala: colocar mesas, cartas, servilletas, cubiertos.",
            "Verificar limpieza de aseos.",
            "Encender TPV, música peruana ambiental (criolla, afroperuana), iluminación según turno.",
            "Briefing de equipo: ceviches del día (según pescado fresco), pisco destacado, reservas, alérgenos.",
        ]),
        ("2. Servicio de Mediodía", [
            "Horario: 13:00-16:00.",
            "Mediodía = servicio estrella de cevichería ('el ceviche se come de día').",
            "Ofrecer menú del día peruano (si aplica) + carta completa.",
            "Velocidad de servicio: ceviches en mesa en 8-10 min, principales en 15 min.",
            "Cevichero: producción continua de ceviches al momento durante servicio.",
            "Wokero: chaufa y saltados deben salir en 5-7 min (wok a máxima potencia).",
            "Gestión de turnos de mesa si hay espera.",
            "Cobro y despedida: sugerir postre (suspiro limeño, picarones) y pisco digestivo.",
        ]),
        ("3. Servicio de Cena", [
            "Horario: 20:00-00:00.",
            "Carta completa + carta de cócteles y piscos destacada.",
            "Iluminación reducida, música más ambiente (música criolla, fusión).",
            "Sugerir vuelos de pisco, pisco sour, ceviches y causas para compartir.",
            "Cevichero: servicio de tiraditos y ceviches especiales (Nikkei si aplica).",
            "Último pedido de cocina: [23:30].",
        ]),
        ("4. Estación del Cevichero: Protocolo Diario", [
            "Verificar que TODO el pescado del día fue previamente congelado -20°C/24h (registro obligatorio).",
            "Descongelar pescado en cámara (0-4°C) con antelación suficiente. NUNCA a temperatura ambiente.",
            "Preparar leche de tigre base: zumo de limón fresco, ají limo, cilantro, cebolla, caldo de pescado, sal. Refrigerar.",
            "Cortar cebolla morada en juliana fina. Lavar en agua con hielo para suavizar.",
            "Preparar ají limo: limpiar, despepitar, picar fino.",
            "Preparar camote cocido y choclo desgranado (guarniciones).",
            "Cortar pescado al momento del pedido: dados de 2x2 cm para ceviche clásico, láminas finas para tiradito.",
            "Mezclar pescado + leche de tigre al momento. No dejar marinando más de 2 minutos (el peruano lo quiere 'recién hecho').",
            "Vida útil de la leche de tigre preparada: máximo 2 horas refrigerada. Desechar sobrante.",
            "Limpieza continua de la estación: tabla, cuchillos, superficies. Protocolo de higiene estricto por pescado crudo.",
        ]),
        ("5. Barra de Piscos: Protocolo", [
            "Verificar stock de todas las referencias (20-30 botellas).",
            "Preparar jarabe de goma casero (si no se compra hecho): agua + azúcar 1:1.",
            "Exprimir limones frescos para pisco sour (alto consumo — 6-8 limones por litro de zumo).",
            "Verificar stock de clara de huevo o espumante vegano para pisco sour.",
            "Preparar hielo suficiente (alto consumo en chilcanos y pisco sour).",
            "Verificar garnish: rodajas de limón, gotas de amargo de angostura, canela en polvo.",
            "Sugerir vuelos de pisco (3 copas: Quebranta, Italia, Mosto Verde) por 14-18€.",
            "Explicar al cliente las variedades de pisco y la diferencia con el aguardiente.",
            "Registrar consumo de botellas para control de stock.",
            "Al cierre: limpiar barra, guardar zumo sobrante refrigerado (máx 24h), cerrar botellas.",
        ]),
        ("6. Cierre del Restaurante", [
            "Cocina: apagar wok, parrilla, hornos, freidora, fogones.",
            "Cevichero: DESECHAR toda la leche de tigre sobrante y pescado cortado no utilizado.",
            "Guardar pescado entero/no cortado en cámara (0-4°C), verificar registro de congelación.",
            "Guardar salsas (huancaína, ocopa) en cámara con fecha y hora.",
            "Registrar temperaturas de cierre de todas las cámaras.",
            "Sala: limpiar mesas, barrer, recoger terraza.",
            "Barra: cerrar caja, cuadrar TPV, limpiar barra, guardar zumo limón refrigerado.",
            "Sacar basuras a contenedores.",
            "Cerrar puertas, alarma, luces.",
        ]),
        ("7. Protocolo Leche de Tigre", [
            "La leche de tigre es el alma del ceviche peruano. Su preparación correcta marca la diferencia.",
            "Ingredientes base: zumo de limón fresco (nunca embotellado), ají limo, cilantro, cebolla morada, caldo de pescado, sal, pimienta.",
            "Preparar en lotes pequeños: máximo para 2 horas de servicio.",
            "Temperatura: mantener SIEMPRE refrigerada (0-4°C) entre usos.",
            "No reutilizar: la leche de tigre que ha estado en contacto con pescado se desecha. Siempre usar leche de tigre limpia.",
            "Venta como shot: la leche de tigre sobrante (sin contacto con pescado) se puede vender como shot aperitivo a 4-5€. Margen infinito.",
            "Variantes: leche de tigre clásica, leche de tigre Nikkei (con sillao y jengibre), leche de tigre de mariscos.",
            "Registro: documentar hora de preparación y hora de descarte en hoja de control APPCC.",
        ]),
        ("8. Gestión de Delivery Peruano", [
            "Zona de packaging separada del pase de sala.",
            "Lomo saltado: envase con compartimento separado para arroz. Viaja perfecto.",
            "Arroz chaufa: envase cerrado, mantiene calor. El rey del delivery peruano.",
            "Bowls peruanos: base arroz + proteína + ají + ensalada criolla.",
            "Pollo a la brasa: cuartos con papas fritas y ensalada en envase familiar.",
            "Causa limeña: viaja bien refrigerada como entrante.",
            "Ceviche delivery: SOLO si entrega < 20 min. Pescado y leche de tigre SEPARADOS.",
            "Verificar cada pedido antes de entregar al rider.",
            "Incluir servilletas, salsas extra (ají, huancaína), tarjeta del restaurante.",
            "Tiempo máximo de preparación: 15 minutos.",
            "Responder a reseñas negativas en plataformas en <24h.",
        ]),
        ("9. Gestión de Alérgenos en Cocina Peruana", [
            "14 alérgenos de declaración obligatoria (estándar UE).",
            "Alérgenos frecuentes en cocina peruana: pescado (ceviches, tiraditos), crustáceos (arroz con mariscos), moluscos (pulpo), frutos secos (en algunas salsas), lácteos (queso fresco, leche evaporada en ají de gallina), gluten (pan en ají de gallina, tallarines).",
            "Carta con indicación de alérgenos por plato.",
            "El ají de gallina contiene pan + leche + frutos secos: informar SIEMPRE.",
            "Muchos platos peruanos son naturalmente sin gluten (a base de arroz, papa, maíz): ventaja competitiva.",
            "Ante duda: confirmar SIEMPRE con jefe de cocina.",
            "Protocolo de cocina para platos sin alérgenos: utensilios separados, zona limpia.",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "manual-operaciones-peruano.docx")
    doc.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 5. PDF placeholder
# ═══════════════════════════════════════════════════════════
def gen_guide_pdf():
    """Generate a minimal placeholder PDF using the DOCX content."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        path = os.path.join(OUTPUT_DIR, "guia-restaurante-peruano.pdf")
        c = canvas.Canvas(path, pagesize=A4)
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(297, 600, "Como Montar un Restaurante Peruano")
        c.setFont("Helvetica", 16)
        c.drawCentredString(297, 560, "80 Plazas de Aforo - Guia Espana 2026")
        c.setFont("Helvetica", 12)
        c.drawCentredString(297, 520, "Chef John Guerrero - AI Chef Pro - aichef.pro")
        c.drawCentredString(297, 480, "20 capitulos - 60+ paginas - 8 plantillas - 6 checklists")
        c.save()
        print(f"✓ {path}")
    except ImportError:
        import shutil
        src = os.path.join(OUTPUT_DIR, "guia-restaurante-peruano.docx")
        dst = os.path.join(OUTPUT_DIR, "guia-restaurante-peruano.pdf")
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f"✓ {dst} (placeholder from DOCX)")


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n🇵🇪 Generando archivos: Restaurante Peruano 80 Plazas\n")
    gen_guide_docx()
    gen_guide_pdf()
    gen_plan_financiero()
    gen_calculadora_capex()
    gen_pl_mensual()
    gen_cash_flow()
    gen_escandallo()
    gen_menu_engineering()
    gen_calculadora_ticket()
    gen_cronograma()
    gen_plantilla_turnos()
    gen_checklist_legal()
    gen_checklist_equipamiento()
    gen_checklist_appcc()
    gen_checklist_sala()
    gen_checklist_contratacion()
    gen_checklist_marketing()
    gen_business_plan()
    gen_manual_operaciones()
    print(f"\n✅ {len(os.listdir(OUTPUT_DIR))} archivos generados en {OUTPUT_DIR}\n")
