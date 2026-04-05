#!/usr/bin/env python3
"""
Generate "Cómo Montar un Restaurante Gastronómico 65 Plazas" guide deliverables.
AI Chef Pro — aichef.pro
21 files: 1 DOCX guide + 10 Excel templates + 8 Excel checklists + 2 DOCX models
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "public", "dl", "guia-restaurante-gastronomico"
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

GOLD = "FFD700"
GOLD_RGB = RGBColor(0xFF, 0xD7, 0x00)
HEADER_BG = "2D2D2D"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"
MEDIUM_GRAY = "E0E0E0"
INPUT_COLOR = "E8F5E9"
BRAND = "AI Chef Pro · aichef.pro — Restaurante Gastronómico"

title_font = Font(name="Calibri", size=16, bold=True, color=GOLD)
sub_font = Font(name="Calibri", size=11, color="888888", italic=True)
hdr_font = Font(name="Calibri", size=11, bold=True, color=WHITE)
sec_font = Font(name="Calibri", size=12, bold=True, color=GOLD)
dat_font = Font(name="Calibri", size=11)
bld_font = Font(name="Calibri", size=11, bold=True)
frm_font = Font(name="Calibri", size=11, color="1565C0", bold=True)
brn_font = Font(name="Calibri", size=9, color="888888", italic=True)

hdr_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
inp_fill = PatternFill(start_color=INPUT_COLOR, end_color=INPUT_COLOR, fill_type="solid")
lgt_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")
sec_fill = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")

thin_b = Border(
    left=Side(style="thin", color=MEDIUM_GRAY), right=Side(style="thin", color=MEDIUM_GRAY),
    top=Side(style="thin", color=MEDIUM_GRAY), bottom=Side(style="thin", color=MEDIUM_GRAY),
)
c_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
l_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
cur_fmt = '#,##0.00 €'
pct_fmt = '0.0%'


def shr(ws, row, ncols):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin_b; cell.alignment = c_align

def sdc(ws, r, c, v=None, font=None, fill=None, fmt=None, align=None):
    cell = ws.cell(row=r, column=c)
    if v is not None: cell.value = v
    cell.font = font or dat_font; cell.border = thin_b; cell.alignment = align or l_align
    if fill: cell.fill = fill
    if fmt: cell.number_format = fmt
    return cell

def title_block(ws, t, ncols=8):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    ws.cell(row=1, column=1, value=t).font = title_font
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
    ws.cell(row=2, column=1, value=BRAND).font = sub_font
    ws.row_dimensions[1].height = 30

def brand_footer(ws, row, ncols):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
    ws.cell(row=row, column=1, value=BRAND).font = brn_font
    ws.cell(row=row, column=1).alignment = c_align

def instr_sheet(wb, t, items):
    ws = wb.active; ws.title = "Instrucciones"; ws.sheet_properties.tabColor = "FFD700"
    title_block(ws, t, 3)
    r = 4
    ws.cell(row=r, column=1, value="Instrucciones de uso").font = sec_font; r += 1
    for i, x in enumerate(items, 1):
        ws.cell(row=r, column=1, value=f"{i}. {x}").font = dat_font; r += 1
    r += 1
    ws.cell(row=r, column=1, value="Celdas verdes = campos editables").font = dat_font
    ws.cell(row=r, column=1).fill = inp_fill
    ws.column_dimensions['A'].width = 80

def checklist_ws(ws, t, items):
    """Create a checklist sheet with items: [(cat, tarea, resp, coste), ...]"""
    hdrs = ["#", "Categoría", "Tarea / Ítem", "Responsable", "Fecha Límite", "Estado", "Coste Est. (€)", "Notas"]
    title_block(ws, t)
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8); ws.freeze_panes = f"A{r+1}"
    ws.auto_filter.ref = f"A{r}:H{r}"
    dv = DataValidation(type="list", formula1='"Pendiente,En Curso,Completado"', allow_blank=True)
    ws.add_data_validation(dv)
    widths = [6, 22, 45, 18, 14, 14, 14, 30]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    for idx, (cat, tarea, resp, coste) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, tarea)
        sdc(ws, r, 4, resp); sdc(ws, r, 5, None, fill=inp_fill)
        sdc(ws, r, 6, "Pendiente", fill=inp_fill); dv.add(ws.cell(row=r, column=6))
        sdc(ws, r, 7, coste, fmt=cur_fmt, fill=inp_fill); sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1; brand_footer(ws, r, 8)


# ═══════════════════════════════════════════════════════════
# 1. GUIDE DOCX
# ═══════════════════════════════════════════════════════════
def gen_guide_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    def add_bullet(txt):
        p = doc.add_paragraph(txt, style='List Bullet')
    def add_num(txt):
        p = doc.add_paragraph(txt, style='List Number')
    def tip(txt):
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Consejo del Chef: {txt}")
        run.font.italic = True; run.font.color.rgb = GOLD_RGB

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cómo Montar un Restaurante Gastronómico"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("65 Plazas de Aforo — Guía España 2026\n(Michelin · Sol Repsol)"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Chef John Guerrero\nAI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # CH1
    doc.add_heading("1. Qué es un Restaurante Gastronómico", level=1)
    doc.add_paragraph(
        "Un restaurante gastronómico es un establecimiento de alta cocina donde la experiencia "
        "culinaria va más allá de la simple alimentación. Se distingue por la calidad excepcional "
        "de los ingredientes, el dominio técnico en las elaboraciones, la creatividad en la "
        "propuesta gastronómica y un servicio de sala impecable."
    )
    doc.add_paragraph(
        "A diferencia de un restaurante casual (ticket medio 15-25€) o fast casual (8-15€), "
        "el gastronómico opera con tickets de 80-120€ por comensal, menús degustación de "
        "90-180€ y una carta de vinos con 300-500 referencias. La inversión inicial oscila "
        "entre 500.000€ y 900.000€ para un local de 65 plazas."
    )
    doc.add_heading("Qué hace diferente al fine dining", level=2)
    add_bullet("Producto: materias primas de primera calidad, Km0, trazabilidad completa.")
    add_bullet("Técnica: dominio absoluto de cocciones, texturas, temperaturas y emplatado.")
    add_bullet("Servicio: ratio 1 camarero por cada 6-8 comensales (vs 1:15 en casual).")
    add_bullet("Experiencia: cada detalle cuenta — iluminación, acústica, vajilla, cristalería.")
    add_bullet("Consistencia: el mismo nivel de excelencia en cada servicio, cada día.")
    tip("El gastronómico no es un restaurante caro. Es un restaurante donde cada euro del ticket se justifica con producto, técnica y servicio. Si no puedes justificar el precio, no es gastronómico.")
    doc.add_page_break()

    # CH2
    doc.add_heading("2. El Mercado de la Alta Cocina en España 2026", level=1)
    doc.add_paragraph(
        "España es una potencia mundial en alta gastronomía. Con 306 restaurantes con Estrella "
        "Michelin (2026), es el tercer país del mundo tras Francia y Japón. La Guía Repsol "
        "reconoce a más de 600 restaurantes con Soles. El sector de la alta cocina española "
        "factura más de 3.000 millones de euros anuales."
    )
    doc.add_heading("Datos clave del sector", level=2)
    add_bullet("306 estrellas Michelin en España (2026): 12 restaurantes con 3 estrellas, 40 con 2 estrellas, 254 con 1 estrella.")
    add_bullet("Ticket medio: 80-120€ en 1 estrella, 120-200€ en 2 estrellas, 200-350€ en 3 estrellas.")
    add_bullet("Facturación anual típica: 1,5-3 millones EUR para un gastronómico de 65 plazas.")
    add_bullet("Ciudades con más oferta gastronómica: Madrid, Barcelona, San Sebastián, Valencia, Sevilla, Bilbao.")
    add_bullet("Crecimiento del turismo gastronómico: +12% anual. España es el segundo destino mundial de turismo culinario.")
    tip("El 85% de los restaurantes con Estrella Michelin en España están fuera de Madrid y Barcelona. Las ciudades medianas y zonas rurales ofrecen alquileres más bajos, producto más directo y menos competencia.")
    doc.add_page_break()

    # CH3
    doc.add_heading("3. Modelos de Negocio", level=1)
    doc.add_paragraph("No todos los restaurantes gastronómicos siguen el mismo modelo. La elección del modelo determina la inversión, la operativa y el posicionamiento.")
    doc.add_heading("Modelo 1: Menú degustación exclusivo", level=2)
    doc.add_paragraph("Sin carta. Solo menú degustación (8-12 pases). Máximo control del food cost (25-28%), menor desperdicio, experiencia inmersiva. Ejemplos: DiverXO, Mugaritz, Azurmendi.")
    doc.add_heading("Modelo 2: Carta + menú degustación", level=2)
    doc.add_paragraph("Carta reducida (4-6 entrantes, 4-6 principales, 3-4 postres) + menú degustación. Mayor flexibilidad, atrae público que no quiere menú cerrado. Food cost 28-32%.")
    doc.add_heading("Modelo 3: Barra / Counter + chef's table", level=2)
    doc.add_paragraph("Tendencia creciente. Cocina abierta con barra donde el chef sirve directamente. 8-12 plazas de barra + comedor. Experiencia inmersiva. Menor inversión en sala.")
    doc.add_heading("Modelo 4: Restaurante de producto", level=2)
    doc.add_paragraph("Centrado en un producto excepcional: pescado de lonja, carne madurada, arroz, verdura de huerta propia. Técnica al servicio del producto. Menor complejidad de carta.")
    doc.add_page_break()

    # CH4
    doc.add_heading("4. Estudio de Viabilidad y Plan Financiero", level=1)
    doc.add_paragraph(
        "La inversión total para un restaurante gastronómico de 65 plazas en España oscila "
        "entre 500.000€ y 900.000€. Es fundamental realizar un estudio de viabilidad riguroso "
        "antes de comprometer capital."
    )
    doc.add_heading("Desglose de inversión (CAPEX)", level=2)
    t = doc.add_table(rows=12, cols=4)
    t.style = 'Table Grid'
    hdrs = ["Concepto", "Rango Bajo", "Rango Medio", "Rango Alto"]
    for i, h in enumerate(hdrs): t.rows[0].cells[i].text = h
    data = [
        ("Obra civil y reforma", "150.000€", "250.000€", "450.000€"),
        ("Equipamiento cocina", "55.000€", "90.000€", "150.000€"),
        ("Mobiliario sala", "30.000€", "60.000€", "120.000€"),
        ("Vajilla/cristalería/cubertería", "15.000€", "35.000€", "65.000€"),
        ("Interiorismo y decoración", "20.000€", "50.000€", "100.000€"),
        ("Bodega inicial", "15.000€", "40.000€", "80.000€"),
        ("Proyecto técnico + licencias", "8.000€", "15.000€", "25.000€"),
        ("Tecnología", "5.000€", "10.000€", "20.000€"),
        ("Marketing lanzamiento", "5.000€", "15.000€", "30.000€"),
        ("Fondo de maniobra (3-6 meses)", "60.000€", "120.000€", "200.000€"),
        ("TOTAL", "363.000€", "685.000€", "1.240.000€"),
    ]
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data): t.rows[i+1].cells[j].text = val
    doc.add_paragraph()
    tip("El fondo de maniobra es la partida que más se subestima. Un gastronómico tarda 6-12 meses en alcanzar velocidad de crucero. Sin fondo, cierras antes de despegar.")
    doc.add_page_break()

    # CH5
    doc.add_heading("5. Requisitos Legales en España", level=1)
    doc.add_paragraph("Un restaurante gastronómico necesita licencia C3 (restaurante con cocina profesional completa separada físicamente del comedor).")
    doc.add_heading("Licencias y permisos obligatorios", level=2)
    add_bullet("Licencia de Actividad (C3): proyecto técnico + tasas municipales (1.500-3.000€). Plazo: 4-8 meses.")
    add_bullet("Proyecto Técnico: instalaciones eléctricas, ventilación, seguridad incendios, CTE (3.000-8.000€).")
    add_bullet("Licencia de Obras: si hay reforma (1.000-3.000€ tasas + coste obra).")
    add_bullet("Registro Sanitario: inscripción CCAA, requiere APPCC implementado. Sin coste directo.")
    add_bullet("Certificado Seguridad contra Incendios: extintores, señalización, detección humos.")
    add_bullet("Alta censal Hacienda (036/037): registro actividad, forma jurídica (SL recomendable).")
    add_bullet("Seguros obligatorios: responsabilidad civil (600-2.000€/año), convenio hostelería, local/contenido.")
    doc.add_heading("Diferencias por Comunidad Autónoma", level=2)
    add_bullet("Madrid: Ley LEPAR, dos vías (licencia completa o declaración responsable). Relativamente flexible.")
    add_bullet("Cataluña: más restrictiva. Moratoria en zonas saturadas. Estudio de impacto acústico obligatorio.")
    add_bullet("Valencia: normativa permisiva, tasas reducidas (~752€), procesos más rápidos.")
    add_bullet("Andalucía: plataformas telemáticas CIRCE simplifican trámites.")
    tip("El obstáculo técnico número 1 es la salida de humos. Muchos locales urbanos no tienen viabilidad de extracción hasta cubierta. Verifica ANTES de alquilar. La capacidad eléctrica también es crítica: una cocina profesional necesita 30-100 kW.")
    doc.add_page_break()

    # CH6
    doc.add_heading("6. APPCC y Seguridad Alimentaria", level=1)
    doc.add_paragraph("El sistema APPCC (Análisis de Peligros y Puntos Críticos de Control) es obligatorio para toda empresa alimentaria en España.")
    doc.add_heading("Los 7 principios APPCC", level=2)
    add_num("Identificar los peligros (biológicos, químicos, físicos) en cada proceso.")
    add_num("Determinar los Puntos Críticos de Control (PCC).")
    add_num("Establecer límites críticos para cada PCC.")
    add_num("Implementar sistema de vigilancia para cada PCC.")
    add_num("Establecer acciones correctivas cuando se superen los límites.")
    add_num("Verificar que el sistema funciona correctamente.")
    add_num("Documentar todos los procedimientos y registros.")
    doc.add_heading("Los 14 alérgenos de declaración obligatoria", level=2)
    doc.add_paragraph("Gluten, crustáceos, huevos, pescado, cacahuetes, soja, lácteos, frutos de cáscara, apio, mostaza, sésamo, sulfitos, altramuces, moluscos.")
    doc.add_heading("Control de temperaturas", level=2)
    add_bullet("Recepción: ≤4°C (refrigerados), ≤-18°C (congelados).")
    add_bullet("Almacenamiento: 0-4°C (cámaras), -18°C (congelación).")
    add_bullet("Cocción: ≥75°C en el centro del producto.")
    add_bullet("Enfriamiento rápido: de 60°C a 10°C en menos de 2 horas.")
    add_bullet("Servicio en caliente: ≥65°C. Servicio en frío: ≤8°C.")
    doc.add_page_break()

    # CH7
    doc.add_heading("7. Ubicación y Local", level=1)
    doc.add_paragraph("Un restaurante gastronómico de 65 plazas necesita un local de aproximadamente 250 m² (cocina + sala + servicios + almacén).")
    doc.add_heading("Ratio cocina/sala", level=2)
    doc.add_paragraph("En fine dining, la cocina debe ocupar entre el 25% y el 35% del espacio total. Para 65 plazas: sala 120-160 m² (1,8-2,5 m² por comensal) y cocina 60-80 m².")
    doc.add_heading("Verificaciones críticas antes de alquilar", level=2)
    add_bullet("Salida de humos: conducto independiente hasta cubierta. Sin esto, no hay licencia C3.")
    add_bullet("Capacidad eléctrica: mínimo 30 kW, idealmente 60-100 kW para cocina profesional.")
    add_bullet("Suministro de agua: caudal suficiente para lavavajillas industrial + cocina.")
    add_bullet("Accesibilidad: acceso sin barreras, baños adaptados (obligatorio por ley).")
    add_bullet("Carga en forjado: la cocina industrial requiere forjados que soporten equipos pesados.")
    tip("El alquiler en zona premium puede representar el 8-12% de tu facturación. En zona secundaria, 4-6%. La diferencia en un gastronómico es menor que en casual porque el cliente viene a ti, no te descubre al pasar.")
    doc.add_page_break()

    # CH8
    doc.add_heading("8. Diseño de Cocina Profesional", level=1)
    doc.add_paragraph("La cocina de un gastronómico se organiza en partidas (estaciones de trabajo especializadas), cada una con su jefe de partida.")
    doc.add_heading("Partidas esenciales", level=2)
    add_bullet("Cuarto frío: entrantes, ensaladas, tartares, ceviches. Mesa refrigerada, cortadora, abatidor.")
    add_bullet("Partida de pescados: limpieza, porcionado, cocción. Plancha, salamandra.")
    add_bullet("Partida de carnes: porcionado, cocción. Josper/parrilla, plancha.")
    add_bullet("Partida de calientes/salsas: fondos, salsas, guarniciones. Fogones, marmita, roner sous-vide.")
    add_bullet("Pastelería/Obrador: postres, petit fours, pan. Horno convección, amasadora, heladera.")
    add_bullet("Zona de emplatado (pase): montaje final. Mesa caliente, lámparas calor, campana.")
    add_bullet("Plonge: lavado vajilla. Túnel de lavado, fregaderos industriales.")
    doc.add_heading("El pase", level=2)
    doc.add_paragraph("El pase es el punto neurálgico. Mesa caliente con lámparas infrarrojo, espacio para que el jefe de cocina revise cada plato, comunicación directa con sala. Ancho mínimo 1,2m. Iluminación blanca neutra 4000K.")
    doc.add_page_break()

    # CH9
    doc.add_heading("9. Equipamiento de Cocina", level=1)
    doc.add_paragraph("Presupuesto de equipamiento cocina completa para 65 plazas: 90.000-150.000€.")
    t2 = doc.add_table(rows=11, cols=3)
    t2.style = 'Table Grid'
    for i, h in enumerate(["Equipo", "Modelo Referencia", "Precio"]): t2.rows[0].cells[i].text = h
    equip = [
        ("Horno racional combi 10 GN", "Rational iCombi Pro", "17.500€"),
        ("Horno brasa", "Josper HJX-25", "10.000-13.000€"),
        ("Abatidor 10 bandejas", "Gama profesional", "3.500-6.000€"),
        ("Cámaras frigoríficas (×2)", "Modulares 1,2×1,2m", "4.000-5.000€"),
        ("Cocina industrial 6 fuegos", "Gama alta + horno", "5.000-5.200€"),
        ("Plancha industrial", "4 quemadores", "1.500-2.500€"),
        ("Lavavajillas capota", "50×50 gama alta", "2.700-4.000€"),
        ("Roner / sous-vide", "Circulador profesional", "500-2.000€"),
        ("Campana extractora", "Con filtración", "1.200-3.000€"),
        ("Mesas trabajo inox (×6)", "2m, con balda", "2.400-3.600€"),
    ]
    for i, (a, b, c) in enumerate(equip):
        t2.rows[i+1].cells[0].text = a; t2.rows[i+1].cells[1].text = b; t2.rows[i+1].cells[2].text = c
    doc.add_page_break()

    # CH10-11
    doc.add_heading("10. Diseño de Sala para 65 Plazas", level=1)
    doc.add_paragraph("Superficie recomendada: 120-160 m² de comedor (1,8-2,5 m² por comensal en gastronómico).")
    doc.add_heading("Distribución típica", level=2)
    add_bullet("Comedor principal: 45-50 plazas.")
    add_bullet("Comedor privado / chef's table: 8-10 plazas.")
    add_bullet("Barra o counter: 6-8 plazas.")
    add_bullet("Zona de recepción/espera con aperitivo.")
    doc.add_heading("Iluminación", level=2)
    doc.add_paragraph("Iluminación suave y cálida (2700-3000K). Capas de luz: general + acento + tarea. Reguladores de intensidad para ajustar según servicio (comida vs cena). Luz directa sobre cada mesa.")
    doc.add_page_break()

    doc.add_heading("11. Vajilla, Cristalería y Cubertería Premium", level=1)
    doc.add_paragraph("Presupuesto vajilla/cristalería/cubertería para 65 plazas: 15.000-65.000€.")
    add_bullet("Vajilla: RAK, Villeroy & Boch (media-alta) o Bernardaud, Noritake (ultra-premium). 8.000-25.000€.")
    add_bullet("Cristalería: Riedel Restaurant (<25€/copa) o Zalto (~80€/copa). 3.000-15.000€.")
    add_bullet("Cubertería: WMF (media-alta) o Christofle (premium). 2.000-10.000€.")
    add_bullet("Mantelería: algodón 300 hilos o lino premium. 1.500-5.000€.")
    tip("Presupuesta un 20% adicional para reposición anual. El fine dining tiene alta tasa de rotura: copas Zalto de 80€ se rompen.")
    doc.add_page_break()

    # CH12
    doc.add_heading("12. Bodega y Servicio de Vinos", level=1)
    doc.add_paragraph("Un gastronómico necesita una carta de vinos con 300-500 referencias. Inversión bodega inicial: 15.000-80.000€.")
    add_bullet("Temperatura constante: 12-14°C (tintos), 8-10°C (blancos).")
    add_bullet("Humedad: 60-70%. Almacenamiento: 1.500-3.000 botellas.")
    add_bullet("Vitrina climatizada en sala: 5.000-25.000€.")
    add_bullet("Software gestión: Winerim (integración con TPV).")
    doc.add_heading("Margen de vinos", level=2)
    doc.add_paragraph("Multiplicador estándar: ×2,5 a ×3,5 sobre coste. Food cost vino: 28-40%. Los vinos por copa tienen mayor margen (×4-5). El maridaje de menú degustación es la venta cruzada más rentable.")
    doc.add_page_break()

    # CH13-14
    doc.add_heading("13. Brigada de Cocina (12-17 personas)", level=1)
    t3 = doc.add_table(rows=7, cols=4)
    t3.style = 'Table Grid'
    for i, h in enumerate(["Puesto", "Cantidad", "Salario Bruto Anual", "Función"]): t3.rows[0].cells[i].text = h
    brigade = [
        ("Chef Ejecutivo", "1", "35.000-80.000€", "Dirección gastronómica, menús, proveedores"),
        ("Sous Chef", "1", "28.000-40.000€", "Segundo al mando, supervisión diaria"),
        ("Jefe de Partida", "3-4", "22.000-30.000€/ud", "Responsable de sección"),
        ("Cocinero (Commis)", "4-6", "18.000-22.000€/ud", "Ejecución bajo supervisión"),
        ("Ayudante de cocina", "2-3", "16.000-19.000€/ud", "Mise en place, apoyo"),
        ("Plonge", "1-2", "15.000-17.000€/ud", "Lavado, limpieza"),
    ]
    for i, (a, b, c, d) in enumerate(brigade):
        t3.rows[i+1].cells[0].text = a; t3.rows[i+1].cells[1].text = b; t3.rows[i+1].cells[2].text = c; t3.rows[i+1].cells[3].text = d
    doc.add_paragraph("\nCoste total anual personal cocina: 250.000-450.000€ (incluyendo SS empresa ~30%).")
    doc.add_page_break()

    doc.add_heading("14. Equipo de Sala (9-12 personas)", level=1)
    t4 = doc.add_table(rows=6, cols=4)
    t4.style = 'Table Grid'
    for i, h in enumerate(["Puesto", "Cantidad", "Salario Bruto Anual", "Función"]): t4.rows[0].cells[i].text = h
    foh = [
        ("Maître / Jefe de Sala", "1", "28.000-42.000€", "Dirección sala, recepción, reservas"),
        ("Sommelier", "1", "30.000-45.000€", "Carta vinos, maridajes, servicio vino"),
        ("Camarero/a de rango", "4-6", "18.000-24.000€/ud", "Servicio mesas, atención cliente"),
        ("Runner / Commis sala", "2-3", "16.000-19.000€/ud", "Apoyo servicio, transporte platos"),
        ("Hostess", "1", "18.000-22.000€", "Reservas, recepción, coordinación"),
    ]
    for i, (a, b, c, d) in enumerate(foh):
        t4.rows[i+1].cells[0].text = a; t4.rows[i+1].cells[1].text = b; t4.rows[i+1].cells[2].text = c; t4.rows[i+1].cells[3].text = d
    doc.add_paragraph("\nCoste total anual personal sala: 200.000-350.000€. Plantilla total: 22-30 personas.")
    doc.add_page_break()

    # CH15-16
    doc.add_heading("15. Menú Engineering para Fine Dining", level=1)
    doc.add_paragraph("El menú engineering es la disciplina que optimiza la carta para maximizar rentabilidad y satisfacción del cliente.")
    doc.add_heading("Estructura de oferta", level=2)
    add_bullet("Menú degustación largo (8-12 pases): 90-180€. Food cost objetivo: 25-28%.")
    add_bullet("Menú degustación corto (5-7 pases): 55-90€. Servicio mediodía entre semana.")
    add_bullet("Carta: entrantes 18-28€, principales 32-48€, postres 14-22€.")
    add_bullet("Maridaje vinos: 45-90€ adicional. Maridaje sin alcohol: 25-40€.")
    doc.add_heading("Matrix de Kasavana & Smith", level=2)
    add_bullet("Stars: alta popularidad + alta rentabilidad → mantener y promocionar.")
    add_bullet("Plowhorses: alta popularidad + baja rentabilidad → reformular costes.")
    add_bullet("Puzzles: baja popularidad + alta rentabilidad → mejorar posición en carta.")
    add_bullet("Dogs: baja popularidad + baja rentabilidad → eliminar o reemplazar.")
    doc.add_page_break()

    doc.add_heading("16. Proveedores Km0 y Producto de Temporada", level=1)
    doc.add_paragraph("La relación directa con productores locales es un pilar del gastronómico moderno. Permite carta viva que cambia según temporada, trazabilidad completa y coherencia con los criterios de Guía Repsol.")
    add_bullet("Establecer relación directa con 10-15 productores locales (agricultores, pescadores, ganaderos).")
    add_bullet("Mínimo 4 cartas anuales, idealmente cambios mensuales en menú degustación.")
    add_bullet("Documentar origen de cada producto (requisito también para Soles Repsol).")
    add_bullet("Fermentaciones propias: miso, garum, kimchi, encurtidos — tendencia fuerte 2026.")
    doc.add_page_break()

    # CH17-18-19
    doc.add_heading("17. Cómo Aspirar a Estrella Michelin", level=1)
    doc.add_paragraph("La Guía Michelin envía inspectores anónimos que evalúan exclusivamente lo que hay en el plato.")
    doc.add_heading("Los 5 criterios universales Michelin", level=2)
    add_num("Calidad de los ingredientes: procedencia, frescura, trazabilidad.")
    add_num("Dominio de técnicas culinarias: precisión en cocciones, texturas, temperaturas.")
    add_num("Armonía de sabores: equilibrio, complejidad, limpieza de sabor.")
    add_num("Personalidad de la cocina: identidad propia, firma del chef, coherencia conceptual.")
    add_num("Consistencia: entre platos del menú y entre diferentes visitas.")
    doc.add_paragraph("\nLo que NO evalúan: servicio, vajilla, decoración, ambiente. Estos aspectos se documentan pero no influyen en la decisión.")
    doc.add_paragraph("Proceso: 12 inspectores en España. Visitas anónimas múltiples. Segundo inspector confirma. Decisiones en sesiones colegiadas.")
    tip("La consistencia es la clave. El restaurante debe rendir al mismo nivel TODOS los días. No solo cuando se espera visita. La reputación se construye 18-24 meses ANTES de aspirar a reconocimiento.")
    doc.add_page_break()

    doc.add_heading("18. Cómo Aspirar a Sol Repsol", level=1)
    doc.add_paragraph("La Guía Repsol evalúa la experiencia global del cliente, no solo el plato.")
    doc.add_heading("Criterios diferenciadores vs Michelin", level=2)
    add_bullet("Evalúan la experiencia completa: desde la reserva hasta la despedida.")
    add_bullet("Incorporan criterios de sostenibilidad, Km0, economía circular.")
    add_bullet("Valoran coherencia entre cocina, local, puesta en escena y bodega.")
    add_bullet("El servicio y ambiente SÍ cuentan (a diferencia de Michelin).")
    doc.add_heading("Niveles", level=2)
    add_bullet("1 Sol: restaurante recomendable. Ingredientes de calidad, cocina coherente.")
    add_bullet("2 Soles: cocina madura y ambiciosa. Servicio impecable, carta de vinos cuidada.")
    add_bullet("3 Soles: experiencia destino única. Cocina en continuo perfeccionamiento.")
    doc.add_paragraph("\n70 inspectores de perfiles diversos trabajan con el Basque Culinary Center.")
    doc.add_page_break()

    doc.add_heading("19. The World's 50 Best", level=1)
    doc.add_paragraph("No se puede solicitar candidatura. 1.120 expertos internacionales votan (50/50 género), divididos en 28 regiones. Deloitte audita el proceso.")
    doc.add_paragraph("Estrategia: invitar críticos internacionales, participar en eventos globales, presencia en redes internacionales. Un restaurante necesita votos de MÁS de un país/región para aparecer.")
    doc.add_page_break()

    # CH20-21-22
    doc.add_heading("20. Marketing, PR y Lanzamiento", level=1)
    doc.add_paragraph("El marketing de un gastronómico es fundamentalmente diferente al de un casual. Se basa en relaciones, prensa especializada y reputación.")
    add_bullet("Primeros 6 meses: prensa local y gastronómica nacional (El Comidista, Bon Viveur, Metrópoli).")
    add_bullet("Agencia de comunicación gastronómica: 1.000-3.000€/mes.")
    add_bullet("Participar en eventos: Madrid Fusión, San Sebastián Gastronomika, Salón de Gourmets.")
    add_bullet("Instagram: canal principal. Fotografía profesional, stories de cocina, reels de técnica.")
    add_bullet("Inversión marketing lanzamiento: 5.000-30.000€.")
    doc.add_page_break()

    doc.add_heading("21. Tecnología para Fine Dining", level=1)
    doc.add_heading("TPV / POS", level=2)
    add_bullet("Revo XEF: el más completo del mercado español. 80-200€/mes.")
    add_bullet("Last.app: moderno, buena integración. 60-150€/mes.")
    add_bullet("ICG FrontRest: robusto, implantado en hostelería tradicional. 100-250€/mes.")
    doc.add_heading("Reservas", level=2)
    add_bullet("TheFork: líder en España. Comisión 1,5-4€ por comensal.")
    add_bullet("Tock: experiencias y prepago. Ideal para menú degustación. Comisión 2-3%.")
    add_bullet("Resy: segmento premium. Propiedad de American Express. 300-500 USD/mes.")
    doc.add_page_break()

    doc.add_heading("22. Tendencias 2025-2026", level=1)
    add_bullet("Cocina al fuego y brasas: Josper, parrilla vasca, horno de leña como técnica central.")
    add_bullet("Fermentación: miso, garum, kimchi, encurtidos caseros. Sabores con mayor profundidad.")
    add_bullet("Km0 como práctica estructural: no consigna sino práctica diaria.")
    add_bullet("Zero waste / Cero desperdicio: aprovechamiento integral del producto.")
    add_bullet("Cocina abierta y counter dining: barra como experiencia premium.")
    add_bullet("Chef's table: experiencia inmersiva con menú exclusivo.")
    add_bullet("Vinos naturales y biodinámicos: presencia creciente en carta.")
    add_bullet("Maridaje no alcohólico: kombucha, jugos fermentados, infusiones frías.")
    add_bullet("Sostenibilidad certificada: energía renovable, huella de carbono documentada.")

    # Final
    doc.add_page_break()
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Guía elaborada por Chef John Guerrero"); r.font.size = Pt(14); r.font.bold = True
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AI Chef Pro · aichef.pro\n© 2026 Todos los derechos reservados"); r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    path = os.path.join(OUTPUT_DIR, "guia-restaurante-gastronomico.docx")
    doc.save(path)
    print(f"  ✅ Guía DOCX: {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 2-11. EXCEL TEMPLATES (10 files)
# ═══════════════════════════════════════════════════════════

def gen_plan_financiero():
    wb = Workbook()
    instr_sheet(wb, "Plan Financiero 3 Años — Restaurante Gastronómico", [
        "Rellena las celdas verdes con tus datos de inversión y costes previstos.",
        "Las fórmulas calcularán automáticamente P&L, margen y proyección a 3 años.",
        "La pestaña 'Inversión' desglosa el CAPEX por categoría.",
        "La pestaña 'P&L Mensual' calcula la cuenta de resultados mensual.",
        "La pestaña 'Proyección 3 Años' extrapola con crecimiento anual.",
    ])
    # Inversión tab
    ws = wb.create_sheet("Inversión")
    title_block(ws, "Inversión Inicial — CAPEX", 5)
    hdrs = ["#", "Concepto", "Presupuesto (€)", "Real (€)", "Desviación (%)"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 5)
    items = [
        "Obra civil y reforma integral", "Instalaciones (electricidad, fontanería, gas, ventilación)",
        "Equipamiento cocina caliente", "Equipamiento cocina fría", "Pastelería y obrador",
        "Zona de pase y expedición", "Plonge y lavado", "Almacenamiento y cámaras",
        "Mobiliario sala (mesas, sillas, sofás)", "Iluminación y decoración",
        "Vajilla, cristalería, cubertería", "Mantelería y textil",
        "Bodega inicial (vinos)", "Vitrina climatizada",
        "TPV, software, tablets", "Web, branding, diseño gráfico",
        "Marketing de lanzamiento", "Proyecto técnico y licencias",
        "Seguros (primer año)", "Stock inicial materias primas",
        "Uniformes equipo", "Fondo de maniobra (6 meses)",
    ]
    for i, item in enumerate(items, 1):
        r += 1
        sdc(ws, r, 1, i, align=c_align)
        sdc(ws, r, 2, item)
        sdc(ws, r, 3, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 4, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=IF(C{r}=0,0,(D{r}-C{r})/C{r})", font=frm_font, fmt=pct_fmt, align=c_align)
    r += 1
    sdc(ws, r, 2, "TOTAL", font=bld_font)
    sdc(ws, r, 3, f"=SUM(C5:C{r-1})", font=frm_font, fmt=cur_fmt)
    sdc(ws, r, 4, f"=SUM(D5:D{r-1})", font=frm_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=IF(C{r}=0,0,(D{r}-C{r})/C{r})", font=frm_font, fmt=pct_fmt, align=c_align)
    ws.column_dimensions['A'].width = 6; ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 18; ws.column_dimensions['D'].width = 18; ws.column_dimensions['E'].width = 16
    ws.freeze_panes = "A5"
    r += 2; brand_footer(ws, r, 5)

    # P&L tab
    ws2 = wb.create_sheet("P&L Mensual")
    title_block(ws2, "Cuenta de Resultados Mensual", 4)
    ws2.cell(row=4, column=1, value="Concepto").font = hdr_font; ws2.cell(row=4, column=1).fill = hdr_fill
    ws2.cell(row=4, column=2, value="Importe (€)").font = hdr_font; ws2.cell(row=4, column=2).fill = hdr_fill
    ws2.cell(row=4, column=3, value="% s/Ventas").font = hdr_font; ws2.cell(row=4, column=3).fill = hdr_fill
    ws2.cell(row=4, column=4, value="Notas").font = hdr_font; ws2.cell(row=4, column=4).fill = hdr_fill
    shr(ws2, 4, 4)
    r = 5
    lines = [
        ("INGRESOS", None, True),
        ("Comidas (cubiertos × ticket medio)", None, False),
        ("Cenas (cubiertos × ticket medio)", None, False),
        ("Vinos y bebidas", None, False),
        ("Eventos privados / chef's table", None, False),
        ("TOTAL INGRESOS", "=SUM(B6:B9)", True),
        ("", None, False),
        ("COSTES VARIABLES", None, True),
        ("Materia prima (food cost)", None, False),
        ("Bebidas (coste bodega)", None, False),
        ("TOTAL COSTES VARIABLES", "=SUM(B13:B14)", True),
        ("", None, False),
        ("MARGEN BRUTO", "=B10-B15", True),
        ("", None, False),
        ("COSTES FIJOS", None, True),
        ("Personal cocina", None, False),
        ("Personal sala", None, False),
        ("Alquiler", None, False),
        ("Suministros (luz, gas, agua)", None, False),
        ("Seguros", None, False),
        ("Marketing", None, False),
        ("Tecnología y software", None, False),
        ("Mantenimiento y reparaciones", None, False),
        ("Limpieza externa", None, False),
        ("Asesoría y gestoría", None, False),
        ("Amortización equipamiento", None, False),
        ("Otros gastos fijos", None, False),
        ("TOTAL COSTES FIJOS", "=SUM(B20:B31)", True),
        ("", None, False),
        ("EBITDA", "=B17-B32", True),
        ("Margen EBITDA", "=IF(B10=0,0,B34/B10)", True),
    ]
    for label, formula, is_bold in lines:
        font = bld_font if is_bold else dat_font
        sdc(ws2, r, 1, label, font=font)
        if formula:
            sdc(ws2, r, 2, formula, font=frm_font, fmt=cur_fmt)
            if "%" in label or "Margen" in label.lower():
                ws2.cell(row=r, column=2).number_format = pct_fmt
            sdc(ws2, r, 3, f"=IF($B$10=0,0,B{r}/$B$10)", font=frm_font, fmt=pct_fmt, align=c_align)
        elif label and not is_bold:
            sdc(ws2, r, 2, None, fill=inp_fill, fmt=cur_fmt)
            sdc(ws2, r, 3, f"=IF($B$10=0,0,B{r}/$B$10)", font=frm_font, fmt=pct_fmt, align=c_align)
        sdc(ws2, r, 4, None, fill=inp_fill)
        r += 1
    ws2.column_dimensions['A'].width = 40; ws2.column_dimensions['B'].width = 18
    ws2.column_dimensions['C'].width = 14; ws2.column_dimensions['D'].width = 30

    path = os.path.join(OUTPUT_DIR, "plan-financiero-3-anos.xlsx")
    wb.save(path); print(f"  ✅ Plan Financiero: {path}"); return path

def gen_capex():
    wb = Workbook()
    instr_sheet(wb, "Calculadora CAPEX — Restaurante Gastronómico", [
        "Cada categoría tiene rangos bajo/medio/alto para referencia.",
        "Rellena la columna 'Tu Presupuesto' con tu inversión prevista.",
        "El total se calcula automáticamente.",
    ])
    ws = wb.create_sheet("CAPEX")
    title_block(ws, "Calculadora CAPEX Restaurante Gastronómico 65 Plazas", 6)
    hdrs = ["#", "Categoría", "Rango Bajo (€)", "Rango Medio (€)", "Rango Alto (€)", "Tu Presupuesto (€)"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 6)
    items = [
        ("Obra civil y reforma integral", 150000, 250000, 450000),
        ("Equipamiento cocina profesional", 55000, 90000, 150000),
        ("Mobiliario sala (65 plazas)", 30000, 60000, 120000),
        ("Vajilla, cristalería, cubertería", 15000, 35000, 65000),
        ("Interiorismo y decoración", 20000, 50000, 100000),
        ("Bodega inicial (vinos)", 15000, 40000, 80000),
        ("Proyecto técnico + licencias", 8000, 15000, 25000),
        ("Tecnología (TPV, reservas, software)", 5000, 10000, 20000),
        ("Marketing lanzamiento", 5000, 15000, 30000),
        ("Stock inicial materias primas", 5000, 10000, 20000),
        ("Fondo de maniobra (6 meses)", 60000, 120000, 200000),
        ("Seguros y otros", 5000, 8000, 15000),
    ]
    for i, (cat, lo, mid, hi) in enumerate(items, 1):
        r += 1
        sdc(ws, r, 1, i, align=c_align); sdc(ws, r, 2, cat)
        sdc(ws, r, 3, lo, fmt=cur_fmt, align=c_align)
        sdc(ws, r, 4, mid, fmt=cur_fmt, align=c_align)
        sdc(ws, r, 5, hi, fmt=cur_fmt, align=c_align)
        sdc(ws, r, 6, None, fill=inp_fill, fmt=cur_fmt, align=c_align)
    r += 1
    sdc(ws, r, 2, "TOTAL", font=bld_font)
    sdc(ws, r, 3, f"=SUM(C5:C{r-1})", font=frm_font, fmt=cur_fmt, align=c_align)
    sdc(ws, r, 4, f"=SUM(D5:D{r-1})", font=frm_font, fmt=cur_fmt, align=c_align)
    sdc(ws, r, 5, f"=SUM(E5:E{r-1})", font=frm_font, fmt=cur_fmt, align=c_align)
    sdc(ws, r, 6, f"=SUM(F5:F{r-1})", font=frm_font, fmt=cur_fmt, align=c_align)
    for c, w in enumerate([6, 42, 18, 18, 18, 20], 1): ws.column_dimensions[get_column_letter(c)].width = w
    ws.freeze_panes = "A5"
    path = os.path.join(OUTPUT_DIR, "calculadora-capex.xlsx")
    wb.save(path); print(f"  ✅ Calculadora CAPEX: {path}"); return path

def gen_pl_escenarios():
    wb = Workbook()
    instr_sheet(wb, "P&L Mensual con 3 Escenarios", [
        "Tres escenarios: Pesimista, Realista, Optimista.",
        "Modifica los parámetros de entrada (celdas verdes).",
        "Los resultados se calculan automáticamente.",
    ])
    ws = wb.create_sheet("Escenarios")
    title_block(ws, "P&L Mensual — 3 Escenarios", 5)
    hdrs = ["Concepto", "Pesimista", "Realista", "Optimista", "Notas"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 5)
    params = [
        ("PARÁMETROS", True), ("Cubiertos/día comida", False), ("Cubiertos/día cena", False),
        ("Ticket medio comida (€)", False), ("Ticket medio cena (€)", False),
        ("Días abierto/mes", False), ("Food cost (%)", False), ("Coste personal mensual (€)", False),
        ("Alquiler mensual (€)", False), ("Otros costes fijos mensuales (€)", False),
        ("", False), ("RESULTADOS", True),
        ("Facturación mensual (€)", False), ("Coste materia prima (€)", False),
        ("Margen bruto (€)", False), ("Costes fijos totales (€)", False),
        ("EBITDA mensual (€)", False), ("Margen EBITDA (%)", False),
    ]
    for label, is_section in params:
        r += 1
        font = sec_font if is_section else dat_font
        sdc(ws, r, 1, label, font=font)
        if not is_section and label:
            for c in [2, 3, 4]:
                sdc(ws, r, c, None, fill=inp_fill if "RESULT" not in label else None, fmt=cur_fmt, align=c_align)
    for c, w in enumerate([40, 18, 18, 18, 25], 1): ws.column_dimensions[get_column_letter(c)].width = w
    path = os.path.join(OUTPUT_DIR, "pl-mensual-escenarios.xlsx")
    wb.save(path); print(f"  ✅ P&L Escenarios: {path}"); return path

def gen_cashflow():
    wb = Workbook()
    instr_sheet(wb, "Cash Flow y Break-Even", [
        "Rellena los ingresos y gastos mensuales previstos.",
        "El cash flow acumulado y el break-even se calculan automáticamente.",
    ])
    ws = wb.create_sheet("Cash Flow 12 Meses")
    title_block(ws, "Cash Flow Mensual — 12 Meses", 14)
    ws.cell(row=4, column=1, value="Concepto").font = hdr_font; ws.cell(row=4, column=1).fill = hdr_fill
    for m in range(1, 13):
        ws.cell(row=4, column=m+1, value=f"Mes {m}").font = hdr_font; ws.cell(row=4, column=m+1).fill = hdr_fill
    ws.cell(row=4, column=14, value="Total Anual").font = hdr_font; ws.cell(row=4, column=14).fill = hdr_fill
    shr(ws, 4, 14)
    rows_data = [
        ("INGRESOS", True), ("Facturación comidas", False), ("Facturación cenas", False),
        ("Bebidas y vinos", False), ("Otros ingresos", False), ("Total Ingresos", True),
        ("", False), ("GASTOS", True), ("Materia prima", False), ("Personal", False),
        ("Alquiler", False), ("Suministros", False), ("Marketing", False),
        ("Otros gastos", False), ("Total Gastos", True), ("", False),
        ("FLUJO DE CAJA NETO", True), ("FLUJO ACUMULADO", True),
    ]
    r = 5
    for label, is_bold in rows_data:
        font = bld_font if is_bold else dat_font
        sdc(ws, r, 1, label, font=font)
        if label and not is_bold:
            for c in range(2, 14): sdc(ws, r, c, None, fill=inp_fill, fmt=cur_fmt, align=c_align)
            sdc(ws, r, 14, f"=SUM(B{r}:M{r})", font=frm_font, fmt=cur_fmt, align=c_align)
        r += 1
    ws.column_dimensions['A'].width = 28
    for c in range(2, 15): ws.column_dimensions[get_column_letter(c)].width = 14
    ws.freeze_panes = "B5"
    path = os.path.join(OUTPUT_DIR, "cash-flow-break-even.xlsx")
    wb.save(path); print(f"  ✅ Cash Flow: {path}"); return path

def gen_escandallo():
    wb = Workbook()
    instr_sheet(wb, "Escandallo Maestro — Fichas Técnicas", [
        "Crea una ficha técnica por plato del menú.",
        "Introduce ingredientes, gramaje, precio y merma.",
        "El coste total y PVP sugerido se calculan automáticamente.",
    ])
    ws = wb.create_sheet("Escandallo")
    title_block(ws, "Ficha Técnica / Escandallo", 9)
    ws.cell(row=4, column=1, value="Nombre del plato:").font = bld_font
    ws.cell(row=4, column=2).fill = inp_fill
    ws.merge_cells("B4:D4")
    ws.cell(row=4, column=5, value="Raciones:").font = bld_font
    ws.cell(row=4, column=6).fill = inp_fill
    ws.cell(row=4, column=7, value="Food Cost Objetivo:").font = bld_font
    ws.cell(row=4, column=8, value="28%").fill = inp_fill
    hdrs = ["#", "Ingrediente", "Unidad", "Cantidad Bruta", "Precio/Ud (€)", "Merma (%)", "Cantidad Neta", "Coste (€)", "Notas"]
    r = 6
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 9)
    for i in range(1, 21):
        r += 1
        sdc(ws, r, 1, i, align=c_align)
        for c in range(2, 7): sdc(ws, r, c, None, fill=inp_fill)
        sdc(ws, r, 7, f"=D{r}*(1-F{r})", font=frm_font, fmt='#,##0.00')
        sdc(ws, r, 8, f"=D{r}*E{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 9, None, fill=inp_fill)
    r += 1
    sdc(ws, r, 7, "COSTE TOTAL:", font=bld_font)
    sdc(ws, r, 8, f"=SUM(H7:H{r-1})", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 7, "PVP Sugerido (28%):", font=bld_font)
    sdc(ws, r, 8, f"=H{r-1}/0.28", font=frm_font, fmt=cur_fmt)
    for c, w in enumerate([5, 28, 10, 14, 14, 10, 14, 14, 20], 1): ws.column_dimensions[get_column_letter(c)].width = w
    path = os.path.join(OUTPUT_DIR, "escandallo-maestro.xlsx")
    wb.save(path); print(f"  ✅ Escandallo: {path}"); return path

def gen_menu_eng():
    wb = Workbook()
    instr_sheet(wb, "Menu Engineering Matrix — Kasavana & Smith", [
        "Introduce los platos de tu carta con su popularidad y margen.",
        "La clasificación Stars/Plowhorses/Puzzles/Dogs se calcula automáticamente.",
    ])
    ws = wb.create_sheet("Menu Engineering")
    title_block(ws, "Menu Engineering Matrix", 8)
    hdrs = ["#", "Plato", "Categoría", "Uds Vendidas", "Coste (€)", "PVP (€)", "Margen (€)", "Clasificación"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8)
    for i in range(1, 26):
        r += 1
        sdc(ws, r, 1, i, align=c_align)
        for c in [2, 3]: sdc(ws, r, c, None, fill=inp_fill)
        for c in [4, 5, 6]: sdc(ws, r, c, None, fill=inp_fill, fmt=cur_fmt if c > 3 else '#,##0')
        sdc(ws, r, 7, f"=F{r}-E{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 8, None, font=frm_font, align=c_align)
    for c, w in enumerate([5, 30, 16, 14, 12, 12, 12, 16], 1): ws.column_dimensions[get_column_letter(c)].width = w
    ws.freeze_panes = "A5"
    path = os.path.join(OUTPUT_DIR, "menu-engineering-matrix.xlsx")
    wb.save(path); print(f"  ✅ Menu Engineering: {path}"); return path

def gen_bodega():
    wb = Workbook()
    instr_sheet(wb, "Budget de Bodega — Inventario de Vinos", [
        "Registra cada referencia de vino con coste, PVP y stock.",
        "El margen y la rotación se calculan automáticamente.",
    ])
    ws = wb.create_sheet("Bodega")
    title_block(ws, "Inventario y Budget de Bodega", 10)
    hdrs = ["#", "Referencia", "Tipo", "D.O. / Región", "Coste (€)", "PVP Carta (€)", "Margen (€)", "Margen (%)", "Stock (uds)", "Rotación/Mes"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    for i in range(1, 51):
        r += 1
        sdc(ws, r, 1, i, align=c_align)
        for c in [2, 3, 4]: sdc(ws, r, c, None, fill=inp_fill)
        sdc(ws, r, 5, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 6, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 7, f"=F{r}-E{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 8, f"=IF(E{r}=0,0,(F{r}-E{r})/E{r})", font=frm_font, fmt=pct_fmt, align=c_align)
        sdc(ws, r, 9, None, fill=inp_fill, fmt='#,##0', align=c_align)
        sdc(ws, r, 10, None, fill=inp_fill, fmt='#,##0', align=c_align)
    for c, w in enumerate([5, 30, 12, 18, 12, 14, 12, 12, 10, 12], 1): ws.column_dimensions[get_column_letter(c)].width = w
    ws.freeze_panes = "A5"
    path = os.path.join(OUTPUT_DIR, "budget-bodega.xlsx")
    wb.save(path); print(f"  ✅ Budget Bodega: {path}"); return path

def gen_ticket():
    wb = Workbook()
    instr_sheet(wb, "Calculadora Ticket Medio", [
        "Simula diferentes escenarios de mix de menús y vinos.",
        "Calcula el ticket medio ponderado y la facturación estimada.",
    ])
    ws = wb.create_sheet("Ticket Medio")
    title_block(ws, "Simulador Ticket Medio — Escenarios", 5)
    hdrs = ["Concepto", "Escenario 1", "Escenario 2", "Escenario 3", "Notas"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 5)
    rows = [
        "% comensales menú degustación largo", "Precio menú largo (€)",
        "% comensales menú degustación corto", "Precio menú corto (€)",
        "% comensales carta", "Ticket medio carta (€)",
        "% comensales con maridaje vinos", "Precio maridaje (€)",
        "% comensales con copa de vino", "Precio copa media (€)",
        "", "TICKET MEDIO PONDERADO (€)", "Cubiertos/día", "Facturación diaria (€)",
        "Días abierto/mes", "Facturación mensual (€)",
    ]
    for label in rows:
        r += 1
        font = bld_font if "TICKET" in label or "Facturación" in label else dat_font
        sdc(ws, r, 1, label, font=font)
        if label:
            for c in [2, 3, 4]: sdc(ws, r, c, None, fill=inp_fill, fmt=cur_fmt, align=c_align)
    for c, w in enumerate([42, 16, 16, 16, 20], 1): ws.column_dimensions[get_column_letter(c)].width = w
    path = os.path.join(OUTPUT_DIR, "calculadora-ticket-medio.xlsx")
    wb.save(path); print(f"  ✅ Calculadora Ticket: {path}"); return path

def gen_gantt():
    wb = Workbook()
    instr_sheet(wb, "Cronograma de Apertura — Gantt 18 Meses", [
        "Cada fase tiene una barra de duración estimada.",
        "Rellena las fechas de inicio y fin reales.",
        "Usa las columnas de estado para hacer seguimiento.",
    ])
    ws = wb.create_sheet("Gantt")
    title_block(ws, "Cronograma de Apertura — 18 Meses", 22)
    ws.cell(row=4, column=1, value="Fase / Tarea").font = hdr_font; ws.cell(row=4, column=1).fill = hdr_fill
    ws.cell(row=4, column=2, value="Responsable").font = hdr_font; ws.cell(row=4, column=2).fill = hdr_fill
    ws.cell(row=4, column=3, value="Estado").font = hdr_font; ws.cell(row=4, column=3).fill = hdr_fill
    ws.cell(row=4, column=4, value="Inicio").font = hdr_font; ws.cell(row=4, column=4).fill = hdr_fill
    ws.cell(row=4, column=5, value="Fin").font = hdr_font; ws.cell(row=4, column=5).fill = hdr_fill
    for m in range(1, 19):
        ws.cell(row=4, column=m+5, value=f"M{m}").font = hdr_font; ws.cell(row=4, column=m+5).fill = hdr_fill
    shr(ws, 4, 23)
    dv = DataValidation(type="list", formula1='"Pendiente,En Curso,Completado"', allow_blank=True)
    ws.add_data_validation(dv)
    phases = [
        ("FASE 1: PLANIFICACIÓN", True),
        ("Estudio de viabilidad financiera", False), ("Búsqueda y selección de local", False),
        ("Constitución de sociedad (SL)", False), ("Contratación arquitecto/interiorista", False),
        ("FASE 2: LICENCIAS Y PROYECTO", True),
        ("Proyecto técnico (arquitecto)", False), ("Solicitud licencia de obras", False),
        ("Solicitud licencia de actividad C3", False), ("Registro sanitario CCAA", False),
        ("FASE 3: OBRA Y EQUIPAMIENTO", True),
        ("Obra civil y reforma integral", False), ("Instalación cocina profesional", False),
        ("Mobiliario y decoración sala", False), ("Instalación tecnología (TPV, etc.)", False),
        ("FASE 4: EQUIPO Y BODEGA", True),
        ("Selección y contratación brigada cocina", False), ("Selección equipo de sala", False),
        ("Compra vajilla/cristalería/cubertería", False), ("Selección proveedores y bodega", False),
        ("FASE 5: PRE-APERTURA", True),
        ("Formación de equipo (2-4 semanas)", False), ("Pruebas de carta y menús", False),
        ("Fotografía profesional de platos", False), ("Marketing pre-apertura y prensa", False),
        ("Servicios de prueba (soft opening)", False),
        ("FASE 6: APERTURA", True),
        ("Apertura oficial", False), ("Primeros 30 días: monitorización", False),
        ("Ajustes de carta según feedback", False),
    ]
    r = 5
    for label, is_section in phases:
        font = sec_font if is_section else dat_font
        fill_row = sec_fill if is_section else None
        sdc(ws, r, 1, label, font=font, fill=fill_row)
        if not is_section:
            sdc(ws, r, 2, None, fill=inp_fill); sdc(ws, r, 3, "Pendiente", fill=inp_fill)
            dv.add(ws.cell(row=r, column=3))
            sdc(ws, r, 4, None, fill=inp_fill); sdc(ws, r, 5, None, fill=inp_fill)
        r += 1
    ws.column_dimensions['A'].width = 42; ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 14; ws.column_dimensions['D'].width = 12; ws.column_dimensions['E'].width = 12
    for c in range(6, 24): ws.column_dimensions[get_column_letter(c)].width = 5
    ws.freeze_panes = "F5"
    path = os.path.join(OUTPUT_DIR, "cronograma-apertura-gantt.xlsx")
    wb.save(path); print(f"  ✅ Cronograma Gantt: {path}"); return path

def gen_turnos():
    wb = Workbook()
    instr_sheet(wb, "Plantilla Turnos Brigada (25 personas)", [
        "Un cuadrante semanal para toda la brigada.",
        "Rellena los turnos (M=Mañana, T=Tarde, P=Partido, L=Libre).",
        "Las horas y costes se calculan automáticamente.",
    ])
    ws = wb.create_sheet("Turnos Semana")
    title_block(ws, "Cuadrante Semanal — Brigada Restaurante Gastronómico", 11)
    hdrs = ["#", "Nombre", "Puesto", "Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom", "Horas/Semana"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 11)
    dv = DataValidation(type="list", formula1='"M,T,P,L,V"', allow_blank=True)
    dv.error = "M=Mañana, T=Tarde, P=Partido, L=Libre, V=Vacaciones"
    ws.add_data_validation(dv)
    staff = [
        ("COCINA", True),
        ("Chef Ejecutivo", False), ("Sous Chef", False),
        ("Jefe Partida Carnes", False), ("Jefe Partida Pescados", False),
        ("Jefe Partida Fríos", False), ("Jefe Partida Pastelería", False),
        ("Commis 1", False), ("Commis 2", False), ("Commis 3", False),
        ("Commis 4", False), ("Commis 5", False),
        ("Ayudante 1", False), ("Ayudante 2", False),
        ("Plonge 1", False), ("Plonge 2", False),
        ("SALA", True),
        ("Maître", False), ("Sommelier", False),
        ("Camarero Rango 1", False), ("Camarero Rango 2", False),
        ("Camarero Rango 3", False), ("Camarero Rango 4", False),
        ("Runner 1", False), ("Runner 2", False),
        ("Hostess", False),
    ]
    for label, is_section in staff:
        r += 1
        if is_section:
            sdc(ws, r, 2, label, font=sec_font, fill=sec_fill)
            for c in range(1, 12): ws.cell(row=r, column=c).fill = sec_fill
        else:
            sdc(ws, r, 1, None, align=c_align)
            sdc(ws, r, 2, None, fill=inp_fill)
            sdc(ws, r, 3, label)
            for c in range(4, 11):
                sdc(ws, r, c, None, fill=inp_fill, align=c_align)
                dv.add(ws.cell(row=r, column=c))
            sdc(ws, r, 11, None, fill=inp_fill, fmt='#,##0', align=c_align)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 22
    for c in range(4, 11): ws.column_dimensions[get_column_letter(c)].width = 6
    ws.column_dimensions['K'].width = 14
    ws.freeze_panes = "D5"
    path = os.path.join(OUTPUT_DIR, "plantilla-turnos-brigada.xlsx")
    wb.save(path); print(f"  ✅ Plantilla Turnos: {path}"); return path


# ═══════════════════════════════════════════════════════════
# 12-19. CHECKLISTS (8 files)
# ═══════════════════════════════════════════════════════════

def gen_cl_legal():
    wb = Workbook(); ws = wb.active
    items = [
        ("Constitución", "Elegir forma jurídica (SL recomendada)", "Abogado/Notario", 600),
        ("Constitución", "Escritura pública ante notario", "Notario", 300),
        ("Constitución", "Inscripción en el Registro Mercantil", "Notario", 200),
        ("Constitución", "Obtener NIF provisional y definitivo", "Hacienda", 0),
        ("Constitución", "Alta censal en Hacienda (Modelo 036/037)", "Asesor fiscal", 0),
        ("Constitución", "Alta en el IAE (epígrafe hostelería)", "Asesor fiscal", 0),
        ("Constitución", "Certificación negativa de denominación social", "Registro Mercantil", 20),
        ("Licencias", "Informe de compatibilidad urbanística", "Arquitecto", 500),
        ("Licencias", "Proyecto técnico de actividad (arquitecto/ingeniero)", "Arquitecto", 5000),
        ("Licencias", "Licencia de obras (si hay reforma)", "Ayuntamiento", 2000),
        ("Licencias", "Licencia de actividad C3 (restaurante)", "Ayuntamiento", 2500),
        ("Licencias", "Certificado de seguridad contra incendios", "Instalador", 800),
        ("Licencias", "Certificado de instalación eléctrica (BT)", "Instalador", 400),
        ("Licencias", "Certificado de instalación de gas", "Instalador", 300),
        ("Licencias", "Proyecto de ventilación y extracción", "Ingeniero", 1500),
        ("Licencias", "Licencia de terraza (si aplica)", "Ayuntamiento", 1000),
        ("Licencias", "Licencia de música (SGAE/AGEDI)", "SGAE", 600),
        ("Sanidad", "Registro sanitario en la CCAA", "CCAA", 0),
        ("Sanidad", "Elaborar e implementar Plan APPCC", "Consultor APPCC", 1500),
        ("Sanidad", "Plan de prerrequisitos (limpieza, plagas, agua...)", "Consultor APPCC", 0),
        ("Sanidad", "Formación manipuladores alimentos (equipo)", "Centro formación", 300),
        ("Sanidad", "Contrato empresa control de plagas (DDD)", "Empresa DDD", 600),
        ("Sanidad", "Análisis microbiológico de agua", "Laboratorio", 150),
        ("Sanidad", "Carnet de manipulador de alimentos", "Online/Presencial", 50),
        ("Seguros", "Seguro de responsabilidad civil", "Aseguradora", 1200),
        ("Seguros", "Seguro de convenio colectivo hostelería", "Aseguradora", 500),
        ("Seguros", "Seguro de local y contenido", "Aseguradora", 1500),
        ("Seguros", "Seguro de interrupción de negocio", "Aseguradora", 800),
        ("Laboral", "Alta empresa en Seguridad Social", "Asesor laboral", 0),
        ("Laboral", "Contratos de trabajo (modelo hostelería)", "Asesor laboral", 200),
        ("Laboral", "Plan de prevención de riesgos laborales", "SPA", 500),
        ("Laboral", "Reconocimientos médicos del personal", "Mutua", 400),
        ("Laboral", "Libro de visitas de la Inspección de Trabajo", "Papelería", 20),
        ("Protección datos", "Registro de actividades de tratamiento (RGPD)", "Abogado", 300),
        ("Protección datos", "Política de privacidad y cookies (web)", "Abogado", 200),
        ("Protección datos", "Contrato de encargado de tratamiento (TPV, reservas)", "Abogado", 0),
        ("Fiscal", "Alta en el censo de empresarios (Modelo 036)", "Asesor", 0),
        ("Fiscal", "Comunicación del sistema de facturación electrónica", "Asesor", 0),
        ("Fiscal", "Hojas de reclamaciones (obligatorio en todas las CCAA)", "Papelería", 10),
        ("Fiscal", "Cartel de horario de apertura y aforo", "Imprenta", 30),
    ]
    checklist_ws(ws, "Checklist Legal — Restaurante Gastronómico", items)
    path = os.path.join(OUTPUT_DIR, "checklist-legal.xlsx")
    wb.save(path); print(f"  ✅ Checklist Legal: {path}"); return path

def gen_cl_equip():
    wb = Workbook(); ws = wb.active
    zones = {
        "Cuarto Frío": [
            ("Mesa refrigerada 2 puertas", "Jefe Partida", 3500),
            ("Cortadora de fiambres profesional", "Sous Chef", 1200),
            ("Abatidor de temperatura 10 GN", "Chef", 5000),
            ("Envasadora al vacío profesional", "Chef", 2500),
            ("Cutter/Robot de cocina industrial", "Jefe Partida", 3000),
            ("Balanza de precisión", "Jefe Partida", 200),
            ("Mandolina profesional", "Jefe Partida", 150),
            ("Set de cortadores y aros de emplatado", "Chef", 300),
        ],
        "Partida Carnes": [
            ("Horno Josper HJX-25 (brasa)", "Chef", 12000),
            ("Plancha industrial 4 quemadores", "Sous Chef", 2000),
            ("Salamandra profesional", "Sous Chef", 1500),
            ("Mesa de trabajo inox 2m", "Sous Chef", 500),
            ("Sierra de huesos", "Jefe Partida", 2000),
            ("Termómetro de sonda digital", "Jefe Partida", 80),
        ],
        "Partida Pescados": [
            ("Plancha lisa/ranurada combinada", "Sous Chef", 1800),
            ("Mesa de trabajo inox con balda", "Sous Chef", 500),
            ("Sartenes de hierro fundido (set)", "Jefe Partida", 300),
            ("Escamador eléctrico", "Jefe Partida", 250),
            ("Pinzas de espinas profesionales", "Jefe Partida", 40),
        ],
        "Cocina Caliente": [
            ("Cocina industrial 6 fuegos + horno", "Chef", 5200),
            ("Cocina industrial 4 fuegos auxiliar", "Chef", 3500),
            ("Roner / sous-vide profesional", "Sous Chef", 1500),
            ("Marmita basculante 50L", "Chef", 4000),
            ("Freidora doble cuba 10+10L", "Sous Chef", 1800),
            ("Baño maría 3 cubetas GN", "Sous Chef", 600),
            ("Sartenes y ollas profesionales (set)", "Chef", 2000),
        ],
        "Pastelería/Obrador": [
            ("Horno de convección 10 GN", "Chef Pastelero", 6000),
            ("Amasadora planetaria 20L", "Chef Pastelero", 2500),
            ("Heladera batch freezer", "Chef Pastelero", 5000),
            ("Temperadora de chocolate", "Chef Pastelero", 1500),
            ("Laminadora de masas", "Chef Pastelero", 3000),
            ("Abatidor de temperatura (compartido)", "Chef Pastelero", 0),
            ("Moldes, aros y utensilios pastelería", "Chef Pastelero", 1000),
        ],
        "Pase/Expedición": [
            ("Mesa caliente pass-through con lámparas IR", "Chef", 3000),
            ("Campana extractora industrial", "Instalador", 2500),
            ("Lámparas de calor infrarrojo (×4)", "Chef", 800),
            ("Timbre/campana de servicio", "Chef", 50),
            ("Iluminación de emplatado (4000K)", "Instalador", 500),
        ],
        "Almacenamiento": [
            ("Cámara frigorífica modular conservación", "Chef", 2500),
            ("Cámara frigorífica modular congelación", "Chef", 3000),
            ("Armario frigorífico 600L", "Sous Chef", 4000),
            ("Armario congelador 1400L", "Sous Chef", 5000),
            ("Estanterías inox almacén seco (×4)", "Sous Chef", 1500),
        ],
        "Plonge/Lavado": [
            ("Lavavajillas de capota 50×50", "Responsable", 3500),
            ("Fregadero industrial doble seno", "Responsable", 800),
            ("Estantería escurrido vajilla", "Responsable", 400),
            ("Contenedores reciclaje separado", "Responsable", 200),
        ],
        "Tecnología": [
            ("TPV principal (hardware + software)", "Gerente", 3000),
            ("TPV auxiliar sala", "Gerente", 1500),
            ("Impresora de comandas cocina (×2)", "Gerente", 600),
            ("Tablet para reservas/hostess", "Gerente", 500),
            ("Red WiFi profesional", "Instalador", 800),
            ("Sistema de sonido sala", "Interiorista", 2000),
            ("Cámaras de seguridad (CCTV)", "Instalador", 1500),
        ],
    }
    items = []
    for zone, equips in zones.items():
        for name, resp, cost in equips:
            items.append((zone, name, resp, cost))
    checklist_ws(ws, "Checklist Equipamiento Cocina", items)
    path = os.path.join(OUTPUT_DIR, "checklist-equipamiento-cocina.xlsx")
    wb.save(path); print(f"  ✅ Checklist Equipamiento: {path}"); return path

def gen_cl_vajilla():
    wb = Workbook(); ws = wb.active
    items = [
        ("Platos", "Plato llano principal (×100)", "Maître", 2000),
        ("Platos", "Plato hondo / sopero (×80)", "Maître", 1600),
        ("Platos", "Plato postre (×80)", "Maître", 1200),
        ("Platos", "Plato pan (×80)", "Maître", 800),
        ("Platos", "Platos presentación / show plate (×70)", "Chef", 2100),
        ("Platos", "Boles variados para entrantes (×60)", "Chef", 900),
        ("Platos", "Pizarras / bandejas presentación (×20)", "Chef", 600),
        ("Platos", "Vajilla petit fours / mignardises (×70)", "Chef Pastelero", 500),
        ("Cristalería", "Copa vino tinto Burgundy (×80)", "Sommelier", 1600),
        ("Cristalería", "Copa vino tinto Bordeaux (×80)", "Sommelier", 1600),
        ("Cristalería", "Copa vino blanco (×80)", "Sommelier", 1200),
        ("Cristalería", "Copa champán / espumoso (×60)", "Sommelier", 900),
        ("Cristalería", "Vaso de agua (×80)", "Maître", 400),
        ("Cristalería", "Copa de cóctel / aperitivo (×40)", "Sommelier", 600),
        ("Cristalería", "Copa de postre / licor (×40)", "Sommelier", 400),
        ("Cristalería", "Decantadores (×6)", "Sommelier", 600),
        ("Cristalería", "Jarra de agua cristal (×20)", "Maître", 400),
        ("Cubertería", "Cuchillo mesa principal (×80)", "Maître", 800),
        ("Cubertería", "Tenedor mesa principal (×80)", "Maître", 800),
        ("Cubertería", "Cuchara sopera (×80)", "Maître", 600),
        ("Cubertería", "Pala pescado (×80)", "Maître", 700),
        ("Cubertería", "Tenedor postre (×80)", "Maître", 500),
        ("Cubertería", "Cuchara postre (×80)", "Maître", 500),
        ("Cubertería", "Cuchillo mantequilla (×80)", "Maître", 400),
        ("Cubertería", "Cucharilla café / infusión (×80)", "Maître", 300),
        ("Cubertería", "Pinzas servicio (×10)", "Maître", 200),
        ("Mantelería", "Manteles individuales o de mesa (×40)", "Maître", 1200),
        ("Mantelería", "Servilletas tela (×120)", "Maître", 600),
        ("Mantelería", "Muletones protectores (×20)", "Maître", 400),
        ("Petit menage", "Saleros y pimenteros (×20 sets)", "Maître", 300),
        ("Petit menage", "Aceiteras / vinajeras (×10)", "Maître", 200),
        ("Petit menage", "Cestitas de pan (×20)", "Maître", 300),
        ("Petit menage", "Portavelas / centros mesa (×20)", "Maître", 400),
        ("Servicio vino", "Sacacorchos Pulltap / Laguiole (×6)", "Sommelier", 300),
        ("Servicio vino", "Enfriadores de mesa (×10)", "Sommelier", 500),
        ("Servicio vino", "Drop stops / antigoteo (×20)", "Sommelier", 40),
        ("Servicio vino", "Termómetro de vino digital", "Sommelier", 30),
        ("Cocina pase", "Pinzas emplatado profesionales (×10)", "Chef", 200),
        ("Cocina pase", "Biberones de salsa (×20)", "Chef", 60),
        ("Cocina pase", "Cloches / campanas (×10)", "Chef", 500),
        ("Reposición", "Stock seguridad platos (+20%)", "Maître", 1500),
        ("Reposición", "Stock seguridad copas (+25%)", "Sommelier", 1000),
        ("Reposición", "Stock seguridad cubertería (+15%)", "Maître", 500),
    ]
    checklist_ws(ws, "Checklist Vajilla, Cristalería y Cubertería", items)
    path = os.path.join(OUTPUT_DIR, "checklist-vajilla-cristaleria.xlsx")
    wb.save(path); print(f"  ✅ Checklist Vajilla: {path}"); return path

def gen_cl_appcc():
    wb = Workbook(); ws = wb.active
    items = [
        ("Plan L+D", "Programa de limpieza y desinfección por zonas", "Resp. APPCC", 0),
        ("Plan L+D", "Fichas técnicas de productos de limpieza", "Resp. APPCC", 0),
        ("Plan L+D", "Registro diario de limpieza (cada zona)", "Equipo", 0),
        ("Plan L+D", "Verificación analítica de superficies (trimestral)", "Laboratorio", 200),
        ("Plan L+D", "Protocolo de limpieza de cámaras frigoríficas", "Resp. APPCC", 0),
        ("Control plagas", "Contrato con empresa DDD autorizada", "Gerente", 600),
        ("Control plagas", "Plano con ubicación de dispositivos", "Empresa DDD", 0),
        ("Control plagas", "Registro de visitas e incidencias", "Empresa DDD", 0),
        ("Control plagas", "Mallas antimoscas en ventanas/puertas", "Instalador", 300),
        ("Control agua", "Análisis microbiológico anual del agua", "Laboratorio", 150),
        ("Control agua", "Plan de mantenimiento de grifos y filtros", "Mantenimiento", 0),
        ("Trazabilidad", "Sistema de registro de lotes de proveedores", "Resp. APPCC", 0),
        ("Trazabilidad", "Etiquetado de productos con fecha y lote", "Equipo cocina", 100),
        ("Trazabilidad", "Registro de entradas de mercancía", "Resp. compras", 0),
        ("Trazabilidad", "Sistema FIFO implementado en almacenes", "Sous Chef", 0),
        ("Temperaturas", "Termómetros calibrados en todas las cámaras", "Resp. APPCC", 200),
        ("Temperaturas", "Registro de temperaturas 2×/día (cámaras)", "Equipo", 0),
        ("Temperaturas", "Registro de temperaturas de cocción", "Jefes partida", 0),
        ("Temperaturas", "Registro de temperaturas de recepción", "Resp. compras", 0),
        ("Temperaturas", "Protocolo de enfriamiento rápido documentado", "Chef", 0),
        ("Temperaturas", "Termómetro de sonda en cada partida", "Chef", 200),
        ("Alérgenos", "Carta de alérgenos actualizada (14 obligatorios)", "Chef", 0),
        ("Alérgenos", "Ficha técnica de cada plato con alérgenos", "Chef", 0),
        ("Alérgenos", "Formación del personal en alérgenos", "Resp. APPCC", 200),
        ("Alérgenos", "Protocolo de gestión de alergia en servicio", "Maître", 0),
        ("Alérgenos", "Señalización visible en carta/menú", "Diseñador", 100),
        ("Formación", "Certificado manipulador alimentos (todo el equipo)", "RRHH", 500),
        ("Formación", "Formación inicial APPCC para nuevas incorporaciones", "Resp. APPCC", 0),
        ("Formación", "Reciclaje anual de formación en higiene", "Resp. APPCC", 200),
        ("Formación", "Registro de asistencia a formaciones", "RRHH", 0),
        ("Proveedores", "Homologación de proveedores (registro sanitario)", "Resp. compras", 0),
        ("Proveedores", "Fichas técnicas de materias primas", "Resp. compras", 0),
        ("Proveedores", "Evaluación anual de proveedores", "Chef", 0),
        ("Proveedores", "Condiciones de entrega y temperatura pactadas", "Resp. compras", 0),
        ("Residuos", "Contrato de gestión de residuos (aceite, orgánico)", "Gerente", 400),
        ("Residuos", "Contenedores separados (orgánico, vidrio, cartón, aceite)", "Resp. APPCC", 200),
        ("Residuos", "Registro de retirada de aceite usado", "Resp. APPCC", 0),
        ("PCC", "Identificación de todos los PCC del proceso", "Resp. APPCC", 0),
        ("PCC", "Límites críticos definidos para cada PCC", "Resp. APPCC", 0),
        ("PCC", "Acciones correctivas documentadas", "Resp. APPCC", 0),
        ("PCC", "Verificación mensual del sistema APPCC", "Resp. APPCC", 0),
        ("Documentación", "Manual APPCC completo y actualizado", "Resp. APPCC", 0),
        ("Documentación", "Cuaderno de registros diarios accesible", "Equipo", 50),
        ("Documentación", "Archivo de incidencias y acciones correctivas", "Resp. APPCC", 0),
        ("Documentación", "Plan de autocontrol presentado a la CCAA", "Resp. APPCC", 0),
    ]
    checklist_ws(ws, "Checklist APPCC — Prerrequisitos y Controles", items)
    path = os.path.join(OUTPUT_DIR, "checklist-appcc.xlsx")
    wb.save(path); print(f"  ✅ Checklist APPCC: {path}"); return path

def gen_cl_michelin():
    wb = Workbook(); ws = wb.active
    items = [
        ("Ingredientes", "Trazabilidad completa de cada producto", "Chef", 0),
        ("Ingredientes", "Relación directa con productores locales documentada", "Chef", 0),
        ("Ingredientes", "Producto de temporada (carta cambia con estaciones)", "Chef", 0),
        ("Ingredientes", "Calidad constante: mismo nivel de producto cada día", "Sous Chef", 0),
        ("Ingredientes", "Almacenamiento impecable (etiquetado, FIFO, temperaturas)", "Sous Chef", 0),
        ("Técnica", "Cocciones precisas y consistentes (termómetros calibrados)", "Jefes Partida", 0),
        ("Técnica", "Texturas variadas y bien ejecutadas en cada plato", "Chef", 0),
        ("Técnica", "Emplatados limpios, estéticos y coherentes", "Chef", 0),
        ("Técnica", "Técnicas modernas dominadas (sous-vide, fermentación, etc.)", "Chef", 0),
        ("Técnica", "Pan y petit fours de producción propia", "Chef Pastelero", 0),
        ("Armonía", "Equilibrio de sabores en cada plato", "Chef", 0),
        ("Armonía", "Complejidad sin confusión: cada elemento tiene razón de ser", "Chef", 0),
        ("Armonía", "Limpieza de sabor: sabores nítidos, no empastados", "Chef", 0),
        ("Armonía", "Coherencia entre los pases del menú degustación", "Chef", 0),
        ("Armonía", "Progresión lógica de intensidad en el menú", "Chef", 0),
        ("Personalidad", "Identidad culinaria clara y reconocible", "Chef", 0),
        ("Personalidad", "Firma del chef visible en cada plato", "Chef", 0),
        ("Personalidad", "Coherencia conceptual (historia, territorio, filosofía)", "Chef", 0),
        ("Personalidad", "Diferenciación clara respecto a competidores", "Chef", 0),
        ("Personalidad", "Evolución constante sin perder identidad", "Chef", 0),
        ("Consistencia", "Mismo nivel de calidad cada servicio, cada día", "Sous Chef", 0),
        ("Consistencia", "Fichas técnicas detalladas de cada plato", "Jefes Partida", 0),
        ("Consistencia", "Mise en place estandarizada y verificada", "Jefes Partida", 0),
        ("Consistencia", "Briefing pre-servicio diario obligatorio", "Chef", 0),
        ("Consistencia", "Formación continua del equipo", "Chef", 0),
        ("Repsol Extra", "Experiencia global: reserva → despedida impecable", "Maître", 0),
        ("Repsol Extra", "Sostenibilidad documentada (Km0, residuos, energía)", "Gerente", 0),
        ("Repsol Extra", "Coherencia cocina-local-servicio-bodega", "Chef + Maître", 0),
        ("Repsol Extra", "Servicio de sala excepcional y formado", "Maître", 0),
        ("Repsol Extra", "Carta de vinos coherente con la propuesta gastronómica", "Sommelier", 0),
        ("Reputación", "Participación en eventos gastronómicos (Madrid Fusión, etc.)", "Chef", 5000),
        ("Reputación", "Relación con prensa gastronómica nacional", "Marketing", 2000),
        ("Reputación", "Presencia en redes sociales con contenido profesional", "Marketing", 0),
        ("Reputación", "Invitaciones estratégicas a críticos y periodistas", "Chef", 3000),
        ("Reputación", "Mínimo 18-24 meses de operación consistente antes de aspirar", "Todo equipo", 0),
        ("Operativa", "Registros APPCC impecables (inspectores lo revisan)", "Resp. APPCC", 0),
        ("Operativa", "Limpieza visible en cocina y sala (incluido baños)", "Todo equipo", 0),
        ("Operativa", "Uniformes limpios y profesionales", "Todo equipo", 1000),
        ("Operativa", "Temperatura de servicio correcta en cada plato", "Jefes Partida", 0),
        ("Operativa", "Tiempos entre pases controlados y fluidos", "Chef + Maître", 0),
    ]
    checklist_ws(ws, "Checklist Preparación Inspección Michelin / Sol Repsol", items)
    path = os.path.join(OUTPUT_DIR, "checklist-inspeccion-michelin-repsol.xlsx")
    wb.save(path); print(f"  ✅ Checklist Michelin/Repsol: {path}"); return path

def gen_cl_sala():
    wb = Workbook(); ws = wb.active
    items = [
        ("Distribución", "Plano de sala con 65 plazas (1,8-2,5 m² por comensal)", "Interiorista", 3000),
        ("Distribución", "Zona comedor principal (45-50 plazas)", "Interiorista", 0),
        ("Distribución", "Comedor privado / chef's table (8-10 plazas)", "Interiorista", 5000),
        ("Distribución", "Barra / counter (6-8 plazas)", "Interiorista", 8000),
        ("Distribución", "Zona de recepción / espera con aperitivo", "Interiorista", 3000),
        ("Distribución", "Flujo de circulación camareros sin cruzar comensales", "Interiorista", 0),
        ("Distribución", "Acceso a baños sin cruzar cocina", "Arquitecto", 0),
        ("Distribución", "Espacio para carrito de gueridón / servicio mesa", "Maître", 2000),
        ("Iluminación", "Iluminación cálida 2700-3000K en comedor", "Interiorista", 5000),
        ("Iluminación", "Reguladores de intensidad (dimmer) por zonas", "Electricista", 1000),
        ("Iluminación", "Luz directa sobre cada mesa (pendientes/spots)", "Interiorista", 3000),
        ("Iluminación", "Iluminación específica para carta de vinos", "Interiorista", 500),
        ("Iluminación", "Sin reflejos ni deslumbramiento en ojos del comensal", "Interiorista", 0),
        ("Acústica", "Tratamiento acústico (paneles absorbentes)", "Interiorista", 3000),
        ("Acústica", "Nivel sonoro objetivo: 60-65 dB en servicio", "Interiorista", 0),
        ("Acústica", "Hilo musical: sistema de sonido discreto", "Instalador", 2000),
        ("Climatización", "Sistema de climatización silencioso y eficiente", "Instalador", 8000),
        ("Climatización", "Temperatura 21-23°C constante", "Instalador", 0),
        ("Climatización", "Extracción de olores de cocina sin afectar sala", "Instalador", 2000),
        ("Mobiliario", "Mesas: estables, tamaño adecuado, material premium", "Interiorista", 10000),
        ("Mobiliario", "Sillas: cómodas para servicio de 2h+, tapizadas", "Interiorista", 15000),
        ("Mobiliario", "Aparador/consola de servicio por zona", "Interiorista", 3000),
        ("Mobiliario", "Estación de sommelier (decantadores, copas, etc.)", "Sommelier", 2000),
        ("Mobiliario", "Vitrina de vinos climatizada (elemento visual)", "Proveedor", 15000),
        ("Baños", "Baños clientes: diseño premium, coherente con sala", "Interiorista", 5000),
        ("Baños", "Baño adaptado personas con movilidad reducida (obligatorio)", "Arquitecto", 2000),
        ("Baños", "Amenities: jabón, crema, ambientador de calidad", "Maître", 200),
        ("Accesibilidad", "Acceso sin barreras desde la calle (rampa o nivel)", "Arquitecto", 2000),
        ("Accesibilidad", "Señalización clara y visible", "Interiorista", 500),
        ("Terraza", "Mobiliario exterior coherente con interior (si aplica)", "Interiorista", 5000),
        ("Terraza", "Toldos/parasoles y calefacción exterior (si aplica)", "Instalador", 3000),
    ]
    checklist_ws(ws, "Checklist Diseño de Sala (FOH)", items)
    path = os.path.join(OUTPUT_DIR, "checklist-diseno-sala.xlsx")
    wb.save(path); print(f"  ✅ Checklist Sala: {path}"); return path

def gen_cl_contratacion():
    wb = Workbook(); ws = wb.active
    items = [
        ("Perfiles", "Definir organigrama completo (cocina + sala)", "Chef + Maître", 0),
        ("Perfiles", "Fichas de puesto: funciones, requisitos, salario", "RRHH", 0),
        ("Perfiles", "Chef Ejecutivo: experiencia mínima 5 años en fine dining", "Propietario", 0),
        ("Perfiles", "Sous Chef: experiencia mínima 3 años en gastronómico", "Chef", 0),
        ("Perfiles", "Sommelier: certificación WSET o equivalente", "Maître", 0),
        ("Perfiles", "Maître: experiencia mínima 3 años en sala premium", "Propietario", 0),
        ("Selección", "Publicar ofertas en portales especializados (Hostelería)", "RRHH", 200),
        ("Selección", "Contactar escuelas de hostelería (BCC, Hofmann, etc.)", "Chef", 0),
        ("Selección", "Pruebas prácticas para cocineros (cocinar un plato)", "Chef", 0),
        ("Selección", "Entrevistas de servicio para personal de sala", "Maître", 0),
        ("Selección", "Verificar referencias de empleos anteriores", "RRHH", 0),
        ("Selección", "Stage de prueba (1-3 días) antes de contrato definitivo", "Chef/Maître", 0),
        ("Documentación", "Contrato de trabajo (modalidad según puesto)", "Asesor laboral", 300),
        ("Documentación", "Alta en Seguridad Social", "Asesor laboral", 0),
        ("Documentación", "Certificado manipulador alimentos", "Empleado", 50),
        ("Documentación", "Cláusula de confidencialidad y no competencia (chef)", "Abogado", 200),
        ("Documentación", "Reconocimiento médico inicial", "Mutua", 60),
        ("Onboarding", "Manual de bienvenida y procedimientos", "Chef/Maître", 0),
        ("Onboarding", "Formación APPCC y alérgenos (primer día)", "Resp. APPCC", 0),
        ("Onboarding", "Tour completo de instalaciones", "Sous Chef", 0),
        ("Onboarding", "Asignación de uniforme y taquilla", "RRHH", 200),
        ("Onboarding", "Presentación al equipo completo", "Chef/Maître", 0),
        ("Formación", "Sesiones de entrenamiento pre-apertura (2-4 semanas)", "Chef", 0),
        ("Formación", "Formación en carta de vinos y maridajes (sala)", "Sommelier", 0),
        ("Formación", "Protocolo de servicio de sala (manual)", "Maître", 0),
        ("Formación", "Simulacro de servicio completo (soft opening)", "Todo equipo", 500),
        ("Legal", "Convenio colectivo de hostelería aplicable", "Asesor laboral", 0),
        ("Legal", "Cuadro horario visible en el centro de trabajo", "RRHH", 0),
        ("Legal", "Registro de jornada digital (obligatorio)", "RRHH", 100),
        ("Legal", "Plan de igualdad (si >50 empleados)", "RRHH", 0),
    ]
    checklist_ws(ws, "Checklist Contratación Equipo", items)
    path = os.path.join(OUTPUT_DIR, "checklist-contratacion.xlsx")
    wb.save(path); print(f"  ✅ Checklist Contratación: {path}"); return path

def gen_cl_marketing():
    wb = Workbook(); ws = wb.active
    items = [
        ("Identidad", "Nombre del restaurante y registro de marca", "Propietario", 300),
        ("Identidad", "Diseño de logo y branding profesional", "Diseñador", 3000),
        ("Identidad", "Línea gráfica: carta, tarjetas, uniformes, señalética", "Diseñador", 2000),
        ("Identidad", "Historia/narrativa del chef y del restaurante", "Chef + Marketing", 0),
        ("Digital", "Web propia optimizada SEO (restaurante + ciudad)", "Agencia web", 4000),
        ("Digital", "Perfil Google My Business completo y verificado", "Marketing", 0),
        ("Digital", "Instagram: perfil profesional con grid planificado", "Marketing", 0),
        ("Digital", "Sesión de fotos profesional de platos (15-25 platos)", "Fotógrafo", 500),
        ("Digital", "Vídeo del restaurante y cocina (60-90 segundos)", "Videógrafo", 1500),
        ("Digital", "Perfil en TheFork/ElTenedor optimizado", "Maître", 0),
        ("Digital", "Perfil en TripAdvisor reclamado", "Marketing", 0),
        ("PR/Prensa", "Nota de prensa profesional para medios", "Agencia PR", 500),
        ("PR/Prensa", "Envío a medios gastronómicos (El Comidista, Bon Viveur)", "Agencia PR", 0),
        ("PR/Prensa", "Envío a medios locales y regionales", "Agencia PR", 0),
        ("PR/Prensa", "Invitación a 5-10 periodistas/influencers gastronómicos", "Chef", 2000),
        ("PR/Prensa", "Contratar agencia de comunicación gastronómica", "Propietario", 3000),
        ("Eventos", "Evento de pre-apertura (prensa + amigos + industria)", "Marketing", 3000),
        ("Eventos", "Soft opening: 2 semanas de servicio de prueba", "Chef + Maître", 0),
        ("Eventos", "Planificar participación en Madrid Fusión / Gastronomika", "Chef", 2000),
        ("Eventos", "Colaboraciones con productores locales (eventos Km0)", "Chef", 500),
        ("Reservas", "Sistema de reservas configurado (Tock/Resy/TheFork)", "Maître", 500),
        ("Reservas", "Política de reservas, cancelación y no-show definida", "Maître", 0),
        ("Reservas", "Protocolo de gestión de lista de espera", "Hostess", 0),
        ("Guías", "Informar a la Guía Michelin de la apertura", "Chef", 0),
        ("Guías", "Informar a la Guía Repsol de la apertura", "Chef", 0),
        ("Guías", "Enviar información a guías locales y regionales", "Marketing", 0),
        ("Fidelización", "Base de datos de clientes (CRM básico)", "Marketing", 0),
        ("Fidelización", "Newsletter mensual / trimestral", "Marketing", 0),
        ("Fidelización", "Programa de recomendación boca a boca", "Maître", 0),
        ("Fidelización", "Seguimiento de reseñas online (responder todas)", "Marketing", 0),
    ]
    checklist_ws(ws, "Checklist Marketing Pre-Apertura", items)
    path = os.path.join(OUTPUT_DIR, "checklist-marketing-preapertura.xlsx")
    wb.save(path); print(f"  ✅ Checklist Marketing: {path}"); return path


# ═══════════════════════════════════════════════════════════
# 20-21. DOCX MODELS (2 files)
# ═══════════════════════════════════════════════════════════

def gen_business_plan():
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLAN DE NEGOCIO\nRestaurante Gastronómico"); r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("[Nombre del Restaurante]\n[Ciudad, Provincia]\n[Fecha]").font.size = Pt(14)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Plantilla modelo — AI Chef Pro · aichef.pro").font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Resumen Ejecutivo", [
            "Describe en 1-2 páginas el concepto del restaurante, la propuesta de valor, la inversión necesaria, la facturación prevista y el equipo fundador.",
            "[Tu resumen ejecutivo aquí — es lo primero que lee el banco/inversor]"
        ]),
        ("2. Concepto y Propuesta de Valor", [
            "Tipo de cocina y filosofía gastronómica.",
            "Público objetivo y posicionamiento.",
            "Diferenciación respecto a la competencia local.",
            "Modelo de servicio: menú degustación, carta, barra, chef's table.",
        ]),
        ("3. Análisis de Mercado", [
            "Análisis del mercado gastronómico local y nacional.",
            "Competencia directa e indirecta (listar 5-10 competidores).",
            "Perfil del cliente objetivo (edad, poder adquisitivo, hábitos).",
            "Tendencias del sector que apoyan la viabilidad.",
        ]),
        ("4. Ubicación y Local", [
            "Dirección o zona prevista.",
            "Superficie total y distribución (cocina/sala/almacén).",
            "Condiciones del alquiler o compra.",
            "Ventajas de la ubicación elegida.",
        ]),
        ("5. Plan de Operaciones", [
            "Estructura organizativa (organigrama).",
            "Brigada de cocina y equipo de sala (perfiles y salarios).",
            "Horarios de servicio y días de apertura.",
            "Proveedores clave y política de compras.",
        ]),
        ("6. Plan de Marketing", [
            "Estrategia de posicionamiento.",
            "Plan de lanzamiento (pre-apertura, soft opening, apertura).",
            "Canales: web, redes sociales, prensa, eventos.",
            "Presupuesto de marketing primer año.",
        ]),
        ("7. Plan Financiero", [
            "Inversión inicial total (CAPEX desglosado).",
            "Fuentes de financiación (fondos propios, préstamo, inversores).",
            "Cuenta de resultados previsional (3 años).",
            "Cash flow mensual primer año.",
            "Punto de equilibrio y análisis de sensibilidad.",
            "Indicadores: food cost objetivo, ratio personal, EBITDA target.",
        ]),
        ("8. Equipo Fundador", [
            "CV del chef y socios fundadores.",
            "Experiencia relevante en hostelería.",
            "Roles y responsabilidades de cada socio.",
        ]),
        ("9. Análisis de Riesgos", [
            "Principales riesgos identificados.",
            "Plan de contingencia para cada riesgo.",
            "Escenarios pesimista/realista/optimista.",
        ]),
        ("10. Calendario de Ejecución", [
            "Cronograma de apertura (12-18 meses).",
            "Hitos clave y fechas objetivo.",
        ]),
    ]
    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_paragraph()
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Plantilla elaborada por Chef John Guerrero\nAI Chef Pro · aichef.pro\n© 2026").font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    path = os.path.join(OUTPUT_DIR, "business-plan-modelo.docx")
    doc.save(path); print(f"  ✅ Business Plan: {path}"); return path

def gen_manual_sala():
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MANUAL DE SERVICIO DE SALA"); r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Restaurante Gastronómico\n[Nombre del Restaurante]").font.size = Pt(14)
    doc.add_page_break()

    doc.add_heading("1. Protocolo de Recepción", level=1)
    doc.add_paragraph("La primera impresión determina la experiencia. El comensal debe sentirse bienvenido desde que entra por la puerta.")
    doc.add_paragraph("1. Saludar al comensal por su nombre si tiene reserva.", style='List Number')
    doc.add_paragraph("2. Acompañar a la mesa, retirar la silla y ofrecer la carta.", style='List Number')
    doc.add_paragraph("3. Ofrecer aperitivo o agua inmediatamente.", style='List Number')
    doc.add_paragraph("4. Explicar brevemente el concepto y las opciones (carta, menú degustación).", style='List Number')

    doc.add_heading("2. Toma de Comanda", level=1)
    doc.add_paragraph("Registrar pedido completo incluyendo alergias, preferencias y ritmo deseado. Confirmar con el comensal antes de enviar a cocina.")

    doc.add_heading("3. Servicio de Vinos", level=1)
    doc.add_paragraph("El sommelier se acerca tras la toma de comanda para sugerir maridaje. Presenta la botella, descorcha en mesa, sirve pequeña muestra al anfitrión, sirve a los demás comenzando por las señoras.")

    doc.add_heading("4. Servicio de Platos", level=1)
    doc.add_paragraph("Servir por la izquierda, retirar por la derecha. Todos los comensales de una mesa reciben el plato simultáneamente. El maître o camarero de rango describe brevemente cada plato al servir.")

    doc.add_heading("5. Gestión de Quejas", level=1)
    doc.add_paragraph("Escuchar sin interrumpir. Disculparse sinceramente. Ofrecer solución inmediata (rehacer el plato, compensar). Informar al maître y al chef. Registrar la incidencia. Hacer seguimiento post-visita.")

    doc.add_heading("6. Protocolo VIP", level=1)
    doc.add_paragraph("Clientes VIP (repetidores, prensa, inspectores potenciales): revisar historial de visitas, preferencias, alergias. Pequeño detalle del chef (amuse-bouche extra). Mesa preferente. Despedida personalizada.")

    doc.add_heading("7. Despedida y Post-Servicio", level=1)
    doc.add_paragraph("Acompañar hasta la puerta. Agradecer la visita. Ofrecer tarjeta del restaurante. Registrar notas del servicio (preferencias, comentarios) en el CRM para la próxima visita.")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nPlantilla elaborada por Chef John Guerrero\nAI Chef Pro · aichef.pro\n© 2026").font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    path = os.path.join(OUTPUT_DIR, "manual-servicio-sala.docx")
    doc.save(path); print(f"  ✅ Manual Sala: {path}"); return path


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("🚀 Generando archivos — Guía Restaurante Gastronómico (21 archivos)\n")

    print("📖 Guía principal:")
    gen_guide_docx()

    print("\n📊 Plantillas Excel (10):")
    gen_plan_financiero()
    gen_capex()
    gen_pl_escenarios()
    gen_cashflow()
    gen_escandallo()
    gen_menu_eng()
    gen_bodega()
    gen_ticket()
    gen_gantt()
    gen_turnos()

    print("\n✅ Checklists Excel (8):")
    gen_cl_legal()
    gen_cl_equip()
    gen_cl_vajilla()
    gen_cl_appcc()
    gen_cl_michelin()
    gen_cl_sala()
    gen_cl_contratacion()
    gen_cl_marketing()

    print("\n📄 Documentos modelo (2):")
    gen_business_plan()
    gen_manual_sala()

    print("\n📦 Verificación:")
    total = 0
    for f in sorted(os.listdir(OUTPUT_DIR)):
        fp = os.path.join(OUTPUT_DIR, f)
        sz = os.path.getsize(fp)
        total += sz
        print(f"   {f}: {sz/1024:.1f} KB")
    print(f"\n   Total: {total/1024:.1f} KB ({len(os.listdir(OUTPUT_DIR))} archivos)")
    print("\n✅ ¡Todos los archivos generados correctamente!")
