#!/usr/bin/env python3
"""
motor.py — Motor COMÚN de la familia «Planes de Negocio» v2.0.

Implementa el **§1 entero** de `scripts/productos-digitales/planes-v2-SPEC.md`
sobre los **30 xlsx** de los 10 productos. NO toca ficheros: recibe un
`Workbook` ya cargado y lo modifica en memoria; quien copia, guarda y verifica
es `main.py`. Los docx son de `documentos.py` (§4); aquí sólo se ofrece un
lector de solo lectura para el gate de A4 (`censar_docx`).

REGLA DE ORO (§1.1 / §7-bis.9): **el motor mide antes de escribir y ABORTA si
no reconoce el molde.** La familia tiene CUATRO moldes de plan financiero
(A-α, A-β, B-γ, B-δ), CUATRO de checklist (C1, C2, C3, C4), DOS de calculadora
(β, γ) y DOS de plantilla de proveedores (P1, P2). Dar por hecho el molde del
representante rompe siete de los diez productos, que es el mismo error que
costó dinero con los «tres moldes de HTML» del blog.

Cobertura por sección de la SPEC:

  §1.1  `detectar()` — firma por nombres de hoja NORMALIZADOS (sin tildes, en
        minúsculas). La normalización no es cosmética: el propio motor
        renombra `'PyG 3 Anos'` → `'PyG 3 Años'` en §1.7, y la 2.ª pasada
        (idempotencia) tiene que seguir reconociendo el fichero. Si ninguna
        firma casa → `MoldeDesconocido` con el nombre del fichero.
  §1.2  `Parametros` — superficie única de parámetros: hoja `0. Supuestos`
        (línea A) o bloque `PARÁMETROS` en `Instrucciones` (B-γ) / `Resumen`
        (B-δ). El motor escribe ahí lo suyo (IVA, SS, SMI, IS) y expone
        `ref(clave)` para que los grupos cablen contra celda, nunca contra
        literal.
  §1.3  `cablear_sumas()` — las constantes tecleadas de las filas TOTAL pasan a
        `=SUM(...)`, **conservando el número de los inputs como ejemplo** y
        SÓLO cuando la aritmética ya cuadra (tolerancia 0,50 € / 0,5 %). Lo que
        no cuadra no se toca: se ANOTA como desajuste, que es justo el defecto
        que los grupos tienen que resolver con criterio.
  §1.4  `formatos_por_tipo()` — €/h entrecomillado, magnitudes guardadas como
        texto a número, `General` en las cifras estrella, decimales
        significativos y la regla de rótulos (recuento nunca lleva €; importe
        nunca lleva %).
  §1.5  `guardas()` + `validaciones()` + `dv_lista()` / `dv_numerica()` — toda
        división con `IFERROR(...,"")`, DV con `showErrorMessage=True`.
  §1.6  `semaforo_num()` / `semaforo_texto()` — formato condicional con la
        guarda `ISNUMBER`, porque las celdas nuevas pueden traer `""`.
  §1.7  `ortografia()` + `renombrar_hojas()` — tildes y eñes en textos y en
        nombres de hoja; el renombrado va DESPUÉS del cableado y reescribe las
        referencias entre hojas, con re-verificación (`gate_referencias`).
  §1.8  `altos_y_wrap()` — alto = ceil(len/ancho) × 15 pt en las combinadas con
        `wrapText`; las filas de datos con `wrap` pierden el alto fijo.
  §1.9  `asegurar_instrucciones()`, `cierre_instrucciones()`, `metadatos()`,
        `cross_sell_sin_precios()`.
  §1.10 `hipervinculos()` — índice → pestañas y URLs de texto plano a enlace,
        con los rótulos igualados al nombre exacto de la pestaña.

  Convenciones de familia: verdes `E8F5E9` desbloqueadas, calculadas sin
  relleno, «sin dato» = `""` (nunca `0`), protección de hoja **sin
  contraseña**, A4, bio anclada, «Versión 2.0 · agosto 2026 · …».

IDEMPOTENCIA: todo lo que se escribe es absoluto o va detrás de un centinela
estructural. DV y formato condicional se PURGAN antes de reescribirse (es donde
openpyxl acumula duplicados en la 2.ª pasada). El bloque de cierre de
`Instrucciones` se borra entero y se reescribe al final del texto vivo, así su
posición depende del contenido y no de dónde estaba antes.

pycel 1.0b30 (medido en este Mac): evalúa `SUM`, `SUMPRODUCT`, `SUMIF`,
`COUNTIF`, `IFERROR`, `IF`/`AND`, `TEXT`, `NPV`, `ROUND`, `MATCH`+`INDEX`.
**NO** implementa `IRR`, `PMT` ni `COUNTA`: `COUNTA(r)` → `COUNTIF(r,"<>")`,
la cuota va como anualidad algebraica `importe*i/(1-(1+i)^-n)` y la TIR se
cachea por Newton (`tir_newton`, portado de `kit-plan-financiero-v2_0`).

Los caracteres tipográficos de la familia se referencian SIEMPRE por escape
(`FINO = '\\u202f'`, `GUION = '\\u2011'`): escritos a pelo degeneran al pasar
por un heredoc del shell y ninguna sustitución encuentra su patrón.

API que consumen los grupos (`grupo_a.py`, `grupo_b.py`, `grupo_c.py`):

    det = motor.detectar(wb, fname)          # aborta si no reconoce el molde
    ws  = motor.hoja(wb, 'PyG 3 Anos')       # búsqueda insensible a tildes
    p   = motor.parametros(wb, det)          # superficie §1.2
    motor.f(ws, 'B10', '=SUM(B6:B9)', motor.FMT_EUR)
    motor.val(ws, 'B4', 55, motor.FMT_ENT, verde_=True)
    motor.iferror('B8/B5')                   # -> '=IFERROR(B8/B5,"")'
    p.ref('ss_empresa')                      # -> "'0. Supuestos'!$B$20"
"""
import copy
import datetime
import html
import math
import os
import re
import shutil
import unicodedata
import zipfile

from openpyxl.formatting.formatting import ConditionalFormattingList
from openpyxl.formatting.rule import FormulaRule, Rule
from openpyxl.cell.cell import MergedCell
from openpyxl.styles import Alignment, Font, PatternFill, Protection
from openpyxl.styles.differential import DifferentialStyle
from openpyxl.utils import column_index_from_string, get_column_letter
from openpyxl.worksheet.cell_range import CellRange
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.worksheet.page import PageMargins

# ==========================================================================
# Tipografía de la familia — SIEMPRE por escape (§ regla dura 5)
# ==========================================================================
FINO = ' '          # espacio fino antes de las unidades
GUION = '‑'         # guion no separable en los rangos

# ==========================================================================
# Paleta, formatos y textos fijos
# ==========================================================================
VERDE = 'E8F5E9'          # celda editable
CREMA = 'FFF8DC'          # subtotal
CAB = '2D2D2D'            # cabecera de tabla

CF_VERDE_BG, CF_VERDE_FG = 'C6EFCE', '006100'
CF_AMBAR_BG, CF_AMBAR_FG = 'FFEB9C', '9C6500'
CF_ROJO_BG, CF_ROJO_FG = 'FFC7CE', '9C0006'

FMT_EUR = '#,##0.00 €'
FMT_EUR0 = '#,##0 €'
FMT_EURH = '#,##0" €/h"'
FMT_PCT = '0.0%'
FMT_PCT0 = '0%'
FMT_ENT = '#,##0'
FMT_DEC = '#,##0.0'
FMT_DEC2 = '#,##0.00'

VERSION = '2.0'
VERSION_FMT = ('Versión ' + VERSION + ' · agosto 2026 · aichef.pro/{pid} · '
               'info@aichef.pro')
RX_VERSION = re.compile(r'^Versi[óo]n \d+\.\d+ · ')

#: Bio anclada de la familia (la Fase A la dejó en 2 de los 30 xlsx).
BIO_LINE = ('Diseñado por John Guerrero — chef y consultor gastronómico desde '
            '2010, en cocina desde los 17 años · johnguerrero.es')
RX_BIO = re.compile(r'John\s+Guerrero|En cocina desde los 17')

NOTA_DESPROTEGER = ('Para editar una celda que no esté en verde: Revisar → '
                    'Desproteger hoja (no tiene contraseña).')
RX_DESPROTEGER = re.compile(r'^Para editar una celda')

NOTA_VERDES = ('Las celdas VERDES son las editables: cambia esas y el resto '
               'del libro se recalcula solo.')
RX_VERDES = re.compile(r'^Las celdas VERDES')

#: §3.8 — la leyenda de la casilla que hoy falta en 12 xlsx.
LEYENDA_OK = ('Columna OK: elige ✓ (hecho), — (pendiente) o N/A (no aplica). '
              'Las N/A no cuentan en el total.')
RX_LEYENDA_OK = re.compile(r'^Columna OK:')

#: §1.9 — cross-sell SIN importes: nombre + una sola URL.
CROSS_SELL = ('Más plantillas y kits del catálogo en '
              'aichef.pro/productos-digitales')
RX_CROSS_SELL = re.compile(r'^Más plantillas y kits')

PIE = 'AI Chef Pro · aichef.pro · Página &P de &N'
CREATOR = 'AI Chef Pro'

# ==========================================================================
# Parámetros fiscales y laborales de la familia (§1.9 / §7-bis.2 y .5)
# ==========================================================================
# Ninguno se escribe dentro de una fórmula ni dentro de un rótulo: van a celda
# con su nota, para que el cliente pueda actualizarlos cuando cambien.
IVA_RESTAURACION = 0.10       # tipo reducido de restauración
IVA_GENERAL = 0.21
SS_EMPRESA = 0.33             # SS a cargo de la empresa (kit-gestion-personal)
SMI_ANUAL = 17094.0           # SMI 2026, RD 126/2026, 14 pagas
PAGAS = 14
IS_NUEVA_CREACION = 0.15      # art. 29.1 LIS, dos primeros ejercicios con base +
IS_GENERAL = 0.25

#: (clave, etiqueta, valor, formato, nota). El orden es FIJO: de él depende la
#: posición del bloque en línea B y, por tanto, la idempotencia.
PARAMS_MOTOR = (
    ('ss_empresa', 'Seguridad Social a cargo de la empresa', SS_EMPRESA,
     FMT_PCT, 'Contingencias comunes + AT/EP + desempleo + FOGASA + FP; '
     'ajusta a tu convenio y bonificaciones'),
    ('pagas', 'Número de pagas anuales', PAGAS, FMT_ENT,
     'Salario bruto anual = bruto mensual × pagas'),
    ('smi_anual', 'SMI anual de referencia (€)', SMI_ANUAL, FMT_EUR0,
     'SMI 2026 · RD 126/2026 · suelo legal por jornada completa'),
    ('is_nueva', 'Impuesto de Sociedades, nueva creación', IS_NUEVA_CREACION,
     FMT_PCT, 'Art. 29.1 LIS: los DOS primeros ejercicios con base imponible '
     'positiva'),
    ('is_general', 'Impuesto de Sociedades, tipo general', IS_GENERAL,
     FMT_PCT, 'A partir del tercer ejercicio con base positiva'),
    ('iva_reducido', 'IVA repercutido de restauración', IVA_RESTAURACION,
     FMT_PCT, 'Tipo reducido de hostelería'),
    ('iva_general', 'IVA repercutido/soportado general', IVA_GENERAL,
     FMT_PCT, 'Bebidas alcohólicas, suministros, equipamiento y servicios'),
)

HOJA_SUPUESTOS = '0. Supuestos'
SENTINELA_PARAMS = 'PARÁMETROS'

#: Celdas que el motor OCUPA en `0. Supuestos` (§2.1 las reserva para estos
#: parámetros). `grupo_a` construye los demás bloques y NO escribe aquí.
CELDAS_SUPUESTOS = {
    'ss_empresa': 'B20',
    'pagas': 'B21',
    'smi_anual': 'B22',
    'is_nueva': 'B37',
    'is_general': 'B38',
    'iva_reducido': 'B39',
    'iva_general': 'B40',
}
#: Rótulos de bloque de `0. Supuestos` (§2.1). El motor los deja escritos para
#: que `grupo_a` rellene debajo sin inventarse la rejilla.
BLOQUES_SUPUESTOS = (
    ('A3', 'ACTIVIDAD'),
    ('A10', 'MIX Y COSTE DE MERCANCÍA'),
    ('A19', 'PERSONAL'),
    ('A23', 'LOCAL Y FIJOS'),
    ('A29', 'FINANCIACIÓN'),
    ('A36', 'FISCAL'),
    ('A43', 'AMORTIZACIÓN'),
)

# ==========================================================================
# §1.1 — Moldes de la familia
# ==========================================================================
#: Hojas que crea el propio motor y que por tanto NO cuentan para la firma.
HOJAS_MOTOR = frozenset(('instrucciones', HOJA_SUPUESTOS.lower()))


class MoldeDesconocido(Exception):
    """El fichero no casa con ninguna firma: el motor NO adivina (§1.1)."""


def norm(texto):
    """Normaliza para comparar: sin tildes, minúsculas, espacios colapsados.

    Imprescindible porque el propio §1.7 renombra hojas (`'PyG 3 Anos'` →
    `'PyG 3 Años'`) y la 2.ª pasada tiene que reconocer el mismo molde.
    """
    if texto is None:
        return ''
    t = unicodedata.normalize('NFD', str(texto))
    t = ''.join(c for c in t if unicodedata.category(c) != 'Mn')
    return re.sub(r'\s+', ' ', t).strip().lower()


#: (molde, hojas obligatorias, hojas alternativas —al menos un grupo entero—).
FIRMAS_PF = (
    ('A-alfa',
     ('1. inversion inicial', '2. p&l 3 anos', '3. punto equilibrio',
      '4. escenarios', '5. personal'),
     ()),
    ('B-gamma',
     ('inversion inicial', 'pyg 3 anos', 'punto equilibrio'),
     (('personal freelance',), ('estacionalidad', 'pricing modelo'))),
    ('A-beta',
     ('inversion inicial', 'pyg 3 anos', 'punto equilibrio', 'escenarios',
      'personal'),
     ()),
    ('B-delta',
     ('resumen', 'p&l ano 1', 'proyeccion 3 anos', 'break-even',
      'mix b2c-b2b', 'kpis'),
     ()),
)

FIRMAS_CALC = (
    ('calc-beta', ('calculadora', 'ejemplos')),
    ('calc-gamma', ('calculadora pricing', '10 ejemplos validados')),
)

RX_PROV_MONO = re.compile(r'^\d+ proveedores$')
RX_FASE = re.compile(r'^f[1-6]\b')

#: Vocabulario de la fila de cabecera de cada molde de checklist (§1.1).
CAB_C1 = ('fase', 'tramite / tarea', 'ok')
CAB_C2 = ('ok', 'tramite / accion')
CAB_C3 = ('#', 'tarea', 'tiempo estimado')
CAB_C4 = ('fase', '#', 'hito')


def _titulos(wb, incluir_motor=False):
    fuera = []
    for ws in wb.worksheets:
        t = norm(ws.title)
        if not incluir_motor and t in HOJAS_MOTOR:
            continue
        fuera.append(t)
    return fuera


def _fila_cabecera(ws, limite=12):
    """Primera fila con >= 3 rótulos de texto: la cabecera de la tabla."""
    for r in range(1, min(limite, ws.max_row) + 1):
        textos = [norm(c.value) for c in ws[r]
                  if isinstance(c.value, str) and c.value.strip()]
        if len(textos) >= 3:
            return r, textos
    return None, []


def _molde_checklist(wb):
    titulos = _titulos(wb)
    fases = [t for t in titulos if RX_FASE.match(t)]
    hojas = [ws for ws in wb.worksheets if norm(ws.title) not in HOJAS_MOTOR]
    if len(fases) >= 6:
        # C2 y C3 comparten la geometría F1..F6 y se distinguen por la
        # CABECERA, no por «tiene fórmulas»: en cuanto el motor le añade el
        # contador, un C3 pasaría a parecer un C2 y la 2.ª pasada trataría el
        # fichero como otro molde (idempotencia rota y estructura pisada).
        ws = [w for w in hojas if RX_FASE.match(norm(w.title))][0]
        _, cab = _fila_cabecera(ws)
        cab = set(cab)
        if 'tiempo estimado' in cab or ('tarea' in cab and 'ok' not in cab):
            return 'C3'
        return 'C2'
    if len(hojas) == 1:
        t = norm(hojas[0].title)
        if t.startswith('checklist 6 fases'):
            return 'C4'
        if t.startswith('checklist'):
            return 'C1'
    return None


def detectar(wb, fname):
    """§1.1 — devuelve `{'tipo', 'molde', 'hojas'}` o ABORTA.

    `MoldeDesconocido` lleva el nombre del fichero y las hojas que tiene, que
    es lo que el orquestador necesita para decidir; el motor no adivina.
    """
    titulos = _titulos(wb)
    juego = set(titulos)

    for molde, obligatorias, alternativas in FIRMAS_PF:
        if not set(obligatorias) <= juego:
            continue
        if alternativas and not any(set(alt) <= juego for alt in alternativas):
            continue
        return {'tipo': 'plan_financiero', 'molde': molde,
                'hojas': list(titulos), 'fichero': fname}

    for molde, obligatorias in FIRMAS_CALC:
        if set(obligatorias) <= juego:
            return {'tipo': 'calculadora', 'molde': molde,
                    'hojas': list(titulos), 'fichero': fname}

    if 'indice' in juego and len(juego) >= 10:
        return {'tipo': 'proveedores', 'molde': 'P1', 'hojas': list(titulos),
                'fichero': fname}
    if len(juego) == 1 and RX_PROV_MONO.match(list(juego)[0]):
        return {'tipo': 'proveedores', 'molde': 'P2', 'hojas': list(titulos),
                'fichero': fname}

    chk = _molde_checklist(wb)
    if chk:
        return {'tipo': 'checklist', 'molde': chk, 'hojas': list(titulos),
                'fichero': fname}

    raise MoldeDesconocido(
        fname + ': ninguna firma de la familia casa con sus hojas '
        + repr([ws.title for ws in wb.worksheets])
        + '. El motor no adivina (§1.1): añade la firma a FIRMAS_* o revisa '
          'el fichero.')


def hoja(wb, nombre, obligatoria=False):
    """Busca una hoja por nombre insensible a tildes/mayúsculas.

    Los grupos la usan SIEMPRE en vez de `wb['PyG 3 Anos']`: después de §1.7 la
    hoja se llama `'PyG 3 Años'` y el acceso directo reventaría con `KeyError`
    en la 2.ª pasada.
    """
    objetivo = norm(nombre)
    for ws in wb.worksheets:
        if norm(ws.title) == objetivo:
            return ws
    if obligatoria:
        raise KeyError('no existe la hoja ' + repr(nombre) + ' en '
                       + repr(wb.sheetnames))
    return None


# ==========================================================================
# Registro de fórmulas — main.py verifica una a una que quedaron cacheadas
# ==========================================================================
REGISTRO = []
#: Renombrados de hoja de la pasada en curso (§1.7). `main.py` los aplica al
#: REGISTRO antes de verificar el caché con `data_only`.
RENOMBRES = {}


def reset(fname=None):
    """Estado por fichero. `main.py` lo llama antes de procesar cada uno."""
    del REGISTRO[:]
    RENOMBRES.clear()


def reg(ws, coord, formula):
    REGISTRO.append((ws.title, coord, formula))


def celda(ws, coord):
    """`ws[coord]` con la garantía de que se puede ESCRIBIR en ella.

    `MergedCell.value` es de sólo lectura, y openpyxl deja objetos
    `MergedCell` colgando cuando una combinada se deshace o cuando
    `delete_rows` desplaza el cuerpo sin tocar `merged_cells`: la celda ya no
    pertenece a ninguna combinada viva y aun así revienta al escribirla con
    `AttributeError: object attribute 'value' is read-only`. Se sustituye por
    una celda normal.
    """
    cel = ws[coord]
    if isinstance(cel, MergedCell):
        fila, col = cel.row, cel.column
        vivas = [CellRange(str(m)) for m in ws.merged_cells.ranges]
        dentro = any(cr.min_row <= fila <= cr.max_row
                     and cr.min_col <= col <= cr.max_col for cr in vivas)
        if dentro:
            for m in list(ws.merged_cells.ranges):
                cr = CellRange(str(m))
                if cr.min_row <= fila <= cr.max_row \
                        and cr.min_col <= col <= cr.max_col:
                    try:
                        ws.unmerge_cells(str(m))
                    except KeyError:
                        try:
                            ws.merged_cells.ranges.remove(m)
                        except Exception:                    # noqa: BLE001
                            pass
        ws._cells.pop((fila, col), None)
        cel = ws.cell(row=fila, column=col)
    return cel


def f(ws, coord, formula, fmt=None, align=None):
    """Escribe una FÓRMULA y la registra para la verificación `data_only`."""
    cel = celda(ws, coord)
    cel.value = formula
    if fmt:
        cel.number_format = fmt
    if align:
        cel.alignment = Alignment(horizontal=align)
    reg(ws, coord, formula)
    return cel


def val(ws, coord, valor, fmt=None, verde_=False, bold=None, align=None,
        wrap=None):
    """Escribe un VALOR constante. `verde_` lo marca como editable."""
    cel = celda(ws, coord)
    cel.value = valor
    if fmt:
        cel.number_format = fmt
    if verde_:
        cel.fill = PatternFill('solid', fgColor=VERDE)
        cel.protection = Protection(locked=False)
    if bold is not None:
        cel.font = Font(bold=bold, size=cel.font.size, color=cel.font.color)
    if align or wrap is not None:
        cel.alignment = Alignment(horizontal=align or 'general',
                                  vertical='top', wrap_text=bool(wrap))
    return cel


def iferror(expresion, alterna=''):
    """`=IFERROR(<expr>,"<alterna>")` — «sin dato» se escribe `""`, no `0`."""
    expr = expresion[1:] if expresion.startswith('=') else expresion
    return '=IFERROR(' + expr + ',"' + alterna + '")'


def es_verde(cel):
    relleno = cel.fill
    return (relleno is not None and relleno.fill_type == 'solid'
            and relleno.fgColor is not None
            and isinstance(relleno.fgColor.rgb, str)
            and relleno.fgColor.rgb.upper().endswith(VERDE))


def verde(ws, rango):
    for fila in ws[rango] if ':' in rango else [[ws[rango]]]:
        for cel in fila:
            cel.fill = PatternFill('solid', fgColor=VERDE)
            cel.protection = Protection(locked=False)


def marcar_editable(ws, rango):
    """Desbloquea por ROL, no por color: hay inputs con relleno propio."""
    marcadas = getattr(ws, '_pl_editables', None)
    if marcadas is None:
        marcadas = set()
        ws._pl_editables = marcadas
    for fila in ws[rango] if ':' in rango else [[ws[rango]]]:
        for cel in fila:
            marcadas.add(cel.coordinate)
            cel.protection = Protection(locked=False)


def permitir_negativo(ws, rango):
    """Excluye un rango de la DV «>= 0»: un resultado puede ser negativo."""
    libres = getattr(ws, '_pl_negativos', None)
    if libres is None:
        libres = set()
        ws._pl_negativos = libres
    for fila in ws[rango] if ':' in rango else [[ws[rango]]]:
        for cel in fila:
            libres.add(cel.coordinate)


def anchos(ws, mapa):
    for letra, ancho in mapa.items():
        actual = ws.column_dimensions[letra].width
        if actual is None or actual < ancho:
            ws.column_dimensions[letra].width = ancho


# ==========================================================================
# §1.6 — Formato condicional (SIEMPRE purgando antes de escribir)
# ==========================================================================
def _norm_rango(ref):
    try:
        return CellRange(ref).coord
    except Exception:                                        # noqa: BLE001
        return ref


def _limpiar_cf(ws, rango):
    objetivo = _norm_rango(rango)
    supervivientes = []
    for cf in ws.conditional_formatting:
        if _norm_rango(str(cf.sqref)) == objetivo:
            continue
        supervivientes.append((str(cf.sqref), list(cf.rules)))
    nueva = ConditionalFormattingList()
    for sqref, reglas in supervivientes:
        for r in reglas:
            nueva.add(sqref, r)
    ws.conditional_formatting = nueva


def _dxf(bg, fg):
    return DifferentialStyle(font=Font(color=fg, bold=True),
                             fill=PatternFill(start_color=bg, end_color=bg,
                                              fill_type='solid'))


def semaforo_num(ws, rango, verde_si=None, ambar_si=None, rojo_si=None):
    """Semáforo numérico con la guarda `ISNUMBER` (§1.6).

    Las expresiones se escriben en función del ANCLA del rango (p. ej.
    `'$B$39<=0.35'`). La guarda no es un adorno: estas celdas pueden traer
    `""` porque «sin dato» se escribe vacío, y sin `ISNUMBER` una comparación
    contra texto pinta de verde una celda sin dato.
    """
    _limpiar_cf(ws, rango)
    ancla = rango.split(':')[0]
    reglas = ((rojo_si, CF_ROJO_BG, CF_ROJO_FG),
              (ambar_si, CF_AMBAR_BG, CF_AMBAR_FG),
              (verde_si, CF_VERDE_BG, CF_VERDE_FG))
    puestas = 0
    for expr, bg, fg in reglas:
        if not expr:
            continue
        formula = 'AND(ISNUMBER(' + ancla + '),' + expr + ')'
        ws.conditional_formatting.add(
            rango, FormulaRule(formula=[formula], stopIfTrue=True,
                               font=Font(color=fg, bold=True),
                               fill=PatternFill(start_color=bg, end_color=bg,
                                                fill_type='solid')))
        puestas += 1
    return puestas


def semaforo_texto(ws, rango, vocabulario):
    """Semáforo por TEXTO contenido (checklists y veredictos)."""
    _limpiar_cf(ws, rango)
    ancla = rango.split(':')[0]
    for texto, bg, fg in vocabulario:
        regla = Rule(type='containsText', operator='containsText', text=texto,
                     dxf=_dxf(bg, fg), stopIfTrue=True)
        regla.formula = ['NOT(ISERROR(SEARCH("' + texto + '",' + ancla
                         + ')))']
        ws.conditional_formatting.add(rango, regla)
    return len(vocabulario)


# ==========================================================================
# §1.5 — Validación de datos
# ==========================================================================
def _purgar_dv(ws, sqref, tipo):
    ws.data_validations.dataValidation = [
        dv for dv in ws.data_validations.dataValidation
        if not (dv.type == tipo and str(dv.sqref) == sqref)]


def _purgar_dv_duplicadas(ws):
    vistas, fuera = set(), []
    for dv in ws.data_validations.dataValidation:
        clave = (dv.type, dv.formula1, dv.formula2, str(dv.sqref))
        if clave in vistas:
            continue
        vistas.add(clave)
        fuera.append(dv)
    ws.data_validations.dataValidation = fuera


def dv_lista(ws, rango, opciones, titulo='Valor no válido', mensaje=None,
             celdas=None):
    """Desplegable cerrado con `showErrorMessage=True` (§1.5).

    Motivo medido: `'Calculadora'!F16` cae en la rama premium ante cualquier
    valor distinto de 1 o 2 y `=IF(B9="sí",1.18,1)` cae en la rama barata si el
    usuario teclea «si» sin tilde. Ninguna de las dos avisa hoy.
    """
    formula = '"' + ','.join(opciones) + '"'
    _purgar_dv(ws, rango, 'list')
    dv = DataValidation(type='list', formula1=formula, allow_blank=True,
                        showErrorMessage=True, errorTitle=titulo,
                        error=mensaje or ('Elige un valor de la lista: '
                                          + ', '.join(opciones)))
    ws.add_data_validation(dv)
    # `celdas` permite colgar el desplegable SÓLO de las filas de ítem: si se
    # cuelga del rectángulo entero, las filas de encabezado de fase también
    # ofrecen la marca y el recuento del gate sale inflado (55 donde el
    # producto anuncia 51).
    for ref in (celdas or [rango]):
        dv.add(ref)
    _purgar_dv_duplicadas(ws)
    return dv


def dv_numerica(ws, coordenadas, minimo=0, maximo=None, entero=False,
                titulo=None, mensaje=None):
    if not coordenadas:
        return None
    coords = sorted(set(coordenadas),
                    key=lambda c: (column_index_from_string(
                        re.match(r'([A-Z]+)', c).group(1)),
                        int(re.search(r'(\d+)', c).group(1))))
    tipo = 'whole' if entero else 'decimal'
    if maximo is None:
        dv = DataValidation(type=tipo, operator='greaterThanOrEqual',
                            formula1=str(minimo), allow_blank=True,
                            showErrorMessage=True,
                            errorTitle=titulo or 'Importe no válido',
                            error=mensaje or ('Escribe un número mayor o igual '
                                              'que ' + str(minimo) + '.'))
    else:
        dv = DataValidation(type=tipo, operator='between',
                            formula1=str(minimo), formula2=str(maximo),
                            allow_blank=True, showErrorMessage=True,
                            errorTitle=titulo or 'Valor fuera de rango',
                            error=mensaje or ('Escribe un valor entre '
                                              + str(minimo) + ' y '
                                              + str(maximo) + '.'))
    ws.add_data_validation(dv)
    for c in coords:
        dv.add(c)
    _purgar_dv_duplicadas(ws)
    return dv


#: ⚠️ RT-09 / RC-17 — «Cubiertos/día» CONTIENE «día», así que el patrón de
#: los DÍAS DE APERTURA se tragaba el driver principal del modelo: quedaba
#: topado en 365 y, al rechazar el valor, enseñaba «Los días van de 1 a 365».
#: La regla se parte en dos: `RX_DIAS` sólo casa cuando el rótulo habla de un
#: NÚMERO de días (los que abres, los de cobro, los de pago), y `RX_POR_DIA`
#: desactiva la regla en cualquier rótulo que sea una magnitud POR día.
RX_DIAS = re.compile(r'^d[ií]as\b|d[ií]as\s+(de|al|por|h[aá]biles|naturales)'
                     r'|n[uú]mero de d[ií]as', re.I)
RX_POR_DIA = re.compile(r'[/ ]\s*d[ií]a\b|al d[ií]a\b|por d[ií]a\b', re.I)
#: RT-08 — recuentos y divisores: no admiten 0 ni decimales. Un 0 en «pagas»
#: pone el coste de personal a cero y TODOS los semáforos en CUMPLE; un 0 en
#: «vida útil» sube el resultado antes de impuestos, porque `IFERROR` convierte
#: la división en `""` y `SUM()` ignora el texto. La guarda no protegía:
#: escondía.
RX_RECUENTO_DV = re.compile(
    r'cubiertos|clientes|comensales|eventos|personas|unidades|pagas?\b|'
    r'n[uú]mero de|vida [uú]til|plazo|carencia|meses|a[ñn]os|aforo|plazas',
    re.I)
#: Techo por magnitud, para que el mensaje de error diga la verdad. El del
#: PLAZO es el número de filas del cuadro de amortización (RT-07): el cuadro
#: tiene tantas como `PLAZO_MAX`, y por encima quedaría capital sin amortizar
#: acusando al usuario de un error que es del generador.
PLAZO_MAX = 15
TECHOS_DV = (
    (re.compile(r'mes de la paga', re.I), 1, 12,
     'Mes del año, de 1 a 12.'),
    (re.compile(r'meses (hasta|de alquiler)', re.I), 0, 24,
     'Número entero de meses (0 a 24).'),
    (re.compile(r'vida [uú]til', re.I), 1, 50,
     'Vida útil en AÑOS ENTEROS (1 a 50). Los coeficientes máximos son los '
     'de la tabla del art. 12.1 LIS.'),
    (re.compile(r'plazo', re.I), 1, PLAZO_MAX,
     'Plazo en años enteros: el cuadro de amortización de este libro llega '
     'a ' + str(PLAZO_MAX) + ' años.'),
    (re.compile(r'carencia', re.I), 0, PLAZO_MAX,
     'Años enteros de carencia, de 0 a ' + str(PLAZO_MAX) + '. Si iguala o '
     'supera al plazo, la hoja la anula.'),
    (re.compile(r'pagas', re.I), 12, 16,
     'Número de pagas al año: 12, 14 o las que fije tu convenio (12 a 16).'),
)
#: Cabeceras de columna que NO son de dato: pintarles una DV numérica hace que
#: el comprador reciba «Escribe un número mayor o igual que 0» al escribir una
#: nota (RT-10 / RC-16: D22, D29 y D42 de la hoja de Inversión).
RX_COL_TEXTO = re.compile(r'^(notas?|comentarios?|observaciones|fuente|'
                          r'responsable|plazo|descripci)', re.I)


def validaciones(ws, informe=None):
    """§1.5 — DV sobre TODAS las celdas verdes, clasificadas por formato.

    `€`/`#,##0` → importe >= 0; `%` → 0-1; rótulo de «Días de apertura» →
    entero 1-365; recuentos y divisores → entero >= 1 (RT-08). Las verdes de
    texto, las de una columna de notas y las que ya llevan desplegable no se
    tocan.
    """
    con_lista = set()
    for dv in ws.data_validations.dataValidation:
        if dv.type == 'list':
            for r in dv.sqref.ranges:
                ref = str(r)
                for fila in ws[ref] if ':' in ref else [[ws[ref]]]:
                    for c in fila:
                        con_lista.add(c.coordinate)
    libres = set(getattr(ws, '_pl_negativos', set()))
    importes, porcentajes, dias, negativos, cuentas = [], [], [], [], []
    for row in ws.iter_rows():
        for c in row:
            if not es_verde(c) or c.coordinate in con_lista:
                continue
            if isinstance(c.value, str) and not c.value.startswith('='):
                continue
            if RX_COL_TEXTO.match(cabecera_de_columna(ws, c.column) or ''):
                continue
            fmt = c.number_format or ''
            rotulo = _rotulo_de_fila(ws, c.row) or ''
            es_dia = bool(RX_DIAS.search(rotulo)) \
                and not RX_POR_DIA.search(rotulo)
            if c.coordinate in libres:
                negativos.append(c.coordinate)
            elif '%' in fmt:
                porcentajes.append(c.coordinate)
            elif es_dia:
                dias.append(c.coordinate)
            elif RX_RECUENTO_DV.search(rotulo) and '€' not in fmt:
                for rx, mn, mx, msg in TECHOS_DV:
                    if rx.search(rotulo):
                        cuentas.append((c.coordinate, mn, mx, msg))
                        break
                else:
                    cuentas.append((c.coordinate, 1, 100000,
                                    'Escribe un número ENTERO de 1 en '
                                    'adelante: esta celda es un recuento y '
                                    'varias fórmulas dividen entre ella.'))
            elif '€' in fmt or fmt.startswith('#,##0') or fmt == '0':
                importes.append(c.coordinate)
    if importes:
        dv_numerica(ws, importes, minimo=0)
    if porcentajes:
        dv_numerica(ws, porcentajes, minimo=0, maximo=1,
                    titulo='Porcentaje no válido',
                    mensaje='Escribe un porcentaje entre 0 y 1 (0,35 = 35 %).')
    if dias:
        dv_numerica(ws, dias, minimo=1, maximo=365, entero=True,
                    titulo='Días fuera de rango',
                    mensaje='Los días van de 1 a 365.')
    grupos = {}
    for coord, mn, mx, msg in cuentas:
        grupos.setdefault((mn, mx, msg), []).append(coord)
    for (mn, mx, msg), coords in sorted(grupos.items(),
                                        key=lambda kv: kv[0][:2]):
        dv_numerica(ws, coords, minimo=mn, maximo=mx, entero=True,
                    titulo='Recuento no válido', mensaje=msg)
    if negativos:
        dv_numerica(ws, negativos, minimo=-1000000000000,
                    titulo='Importe no válido',
                    mensaje='Escribe un número (puede ser negativo).')
    if informe is not None and (importes or porcentajes or dias or cuentas):
        informe.append(ws.title + ': DV en ' + str(len(importes))
                       + ' importes, ' + str(len(porcentajes))
                       + ' porcentajes, ' + str(len(dias)) + ' días y '
                       + str(len(cuentas)) + ' recuentos')
    return (len(importes) + len(porcentajes) + len(dias) + len(negativos)
            + len(cuentas))


# ==========================================================================
# §1.2 — Superficie de parámetros
# ==========================================================================
class Parametros(object):
    """Hoja `0. Supuestos` (línea A) o bloque `PARÁMETROS` (línea B).

    Ningún literal sobrevive dentro de una fórmula: hoy `'P&L Año 1'!B9` obliga
    a editar 26 fórmulas para cambiar el food cost y `'Break-even'!B7` lo
    obliga otra vez en un tercer sitio con el mismo número escrito dos veces.
    """

    def __init__(self, wb, det):
        self.wb = wb
        self.det = det
        self.refs = {}
        self.valores = {}
        self.hoja = None
        self.fila0 = None
        self._construir()

    # -- construcción ----------------------------------------------------
    def _construir(self):
        molde = self.det['molde']
        if self.det['tipo'] != 'plan_financiero':
            # Sólo el plan financiero tiene tipos dentro de las fórmulas: una
            # checklist o una plantilla de proveedores no necesita bloque de
            # parámetros, y metérselo sería mobiliario que nadie usa.
            return
        if molde in ('A-alfa', 'A-beta'):
            self._supuestos()
        elif molde == 'B-delta':
            self.hoja = hoja(self.wb, 'Resumen')
            self._bloque()
        else:
            self.hoja = hoja(self.wb, 'Instrucciones') or hoja(self.wb,
                                                               'Resumen')
            self._bloque()

    def _supuestos(self):
        ws = hoja(self.wb, HOJA_SUPUESTOS)
        if ws is None:
            ws = self.wb.create_sheet(HOJA_SUPUESTOS, 0)
        else:
            self.wb.move_sheet(ws, offset=-self.wb.index(ws))
        self.hoja = ws
        val(ws, 'A1', 'SUPUESTOS — la ÚNICA hoja donde se teclean datos',
            bold=True)
        val(ws, 'A2', NOTA_VERDES)
        for coord, texto in BLOQUES_SUPUESTOS:
            val(ws, coord, texto, bold=True)
        for clave, etiqueta, valor, fmt, nota in PARAMS_MOTOR:
            coord = CELDAS_SUPUESTOS.get(clave)
            if not coord:
                continue
            fila = int(re.search(r'(\d+)', coord).group(1))
            val(ws, 'A' + str(fila), etiqueta)
            cel = ws[coord]
            if cel.value is None:
                val(ws, coord, valor, fmt, verde_=True)
            else:
                cel.number_format = fmt
                verde(ws, coord)
            val(ws, 'C' + str(fila), nota)
            self.refs[clave] = ("'" + ws.title + "'!$"
                                + re.match(r'([A-Z]+)', coord).group(1) + '$'
                                + str(fila))
            self.valores[clave] = ws[coord].value
        anchos(ws, {'A': 46, 'B': 16, 'C': 62})

    def _bloque(self):
        """Bloque `PARÁMETROS` anclado por centinela, no por posición."""
        ws = self.hoja
        if ws is None:
            return
        col = _col_texto(ws)
        letra = get_column_letter(col)
        letra_v = get_column_letter(col + 1)
        letra_n = get_column_letter(col + 2)
        fila0 = None
        for r in range(1, ws.max_row + 1):
            if norm(ws.cell(row=r, column=col).value) == norm(SENTINELA_PARAMS):
                fila0 = r
                break
        if fila0 is None:
            fila0 = _ultima_fila(ws, col) + 2
        val(ws, letra + str(fila0), SENTINELA_PARAMS, bold=True)
        val(ws, letra + str(fila0 + 1),
            'Cambia AQUÍ los tipos: las fórmulas del libro los leen de esta '
            'celda, no los llevan escritos dentro.')
        fila = fila0 + 2
        for clave, etiqueta, valor, fmt, nota in PARAMS_MOTOR:
            val(ws, letra + str(fila), etiqueta)
            coord = letra_v + str(fila)
            cel = ws[coord]
            if cel.value is None:
                val(ws, coord, valor, fmt, verde_=True)
            else:
                cel.number_format = fmt
                verde(ws, coord)
            val(ws, letra_n + str(fila), nota)
            self.refs[clave] = ("'" + ws.title + "'!$" + letra_v + '$'
                                + str(fila))
            self.valores[clave] = ws[coord].value
            fila += 1
        self.fila0 = fila0
        anchos(ws, {letra: 46, letra_v: 16, letra_n: 62})

    # -- API para los grupos --------------------------------------------
    def ref(self, clave):
        """Referencia ABSOLUTA con nombre de hoja: `'0. Supuestos'!$B$20`."""
        if clave not in self.refs:
            raise KeyError('parámetro desconocido: ' + repr(clave)
                           + ' (disponibles: ' + ', '.join(sorted(self.refs))
                           + ')')
        return self.refs[clave]

    def valor(self, clave):
        return self.valores.get(clave)

    def alta(self, clave, etiqueta, valor, fmt=None, nota=None, coord=None):
        """Registra un parámetro NUEVO (lo usan los grupos, §2.1 y §3.4).

        Con `coord` escribe en una celda concreta de la hoja de supuestos; sin
        él, lo cuelga al final del bloque `PARÁMETROS`.
        """
        ws = self.hoja
        if ws is None:
            raise RuntimeError('este fichero no tiene superficie de '
                               'parámetros: ' + self.det['fichero'])
        if coord is None:
            col = _col_texto(ws)
            letra_v = get_column_letter(col + 1)
            fila = _ultima_fila(ws, col + 1) + 1
            coord = letra_v + str(fila)
        fila = int(re.search(r'(\d+)', coord).group(1))
        letra = re.match(r'([A-Z]+)', coord).group(1)
        col_et = get_column_letter(max(1, column_index_from_string(letra) - 1))
        val(ws, col_et + str(fila), etiqueta)
        cel = ws[coord]
        if cel.value is None:
            val(ws, coord, valor, fmt, verde_=True)
        else:
            if fmt:
                cel.number_format = fmt
            verde(ws, coord)
        if nota:
            col_no = get_column_letter(column_index_from_string(letra) + 1)
            val(ws, col_no + str(fila), nota)
        self.refs[clave] = ("'" + ws.title + "'!$" + letra + '$' + str(fila))
        self.valores[clave] = ws[coord].value
        return self.refs[clave]


def parametros(wb, det):
    return Parametros(wb, det)


# ==========================================================================
# Utilidades de rejilla
# ==========================================================================
def _col_texto(ws):
    """Columna donde vive el texto de una hoja de instrucciones (A o B)."""
    for r in range(1, min(8, ws.max_row) + 1):
        if isinstance(ws.cell(row=r, column=2).value, str):
            if not isinstance(ws.cell(row=r, column=1).value, str):
                return 2
    return 1


def _ultima_fila(ws, col=None):
    ultima = 0
    for r in range(1, ws.max_row + 1):
        if col is None:
            fila = [c.value for c in ws[r]]
            if any(v is not None for v in fila):
                ultima = r
        elif ws.cell(row=r, column=col).value is not None:
            ultima = r
    return ultima


def cabecera_de_columna(ws, col):
    """Rótulo de la columna, leído en la fila de cabecera de la tabla.

    Hace falta para no confundir una columna de RATIOS con una fila de
    importes: `'PyG 3 Anos'!E` se titula «% s/VENTAS» y sus celdas están en la
    fila «TOTAL COSTES VARIABLES». Mirando sólo la fila, la regla del §1.4 le
    quitaba el formato de porcentaje a un porcentaje.
    """
    cache = getattr(ws, '_pl_cabeceras', None)
    if cache is None:
        cache = {}
        fila, _ = _fila_cabecera(ws)
        if fila:
            for c in ws[fila]:
                if isinstance(c.value, str):
                    cache[c.column] = c.value
        ws._pl_cabeceras = cache
    return cache.get(col, '')


def _rotulo_de_fila(ws, fila, max_col=3):
    """Rótulo de la fila: el primer texto de las columnas A..C."""
    for col in range(1, max_col + 1):
        v = ws.cell(row=fila, column=col).value
        if isinstance(v, str) and v.strip() and not v.startswith('='):
            return v.strip()
    return ''


def _filas_combinadas(ws):
    fuera = set()
    for m in ws.merged_cells.ranges:
        for r in range(m.min_row, m.max_row + 1):
            fuera.add(r)
    return fuera


def fila_libre(ws, desde, columnas):
    """Primera fila >= `desde` vacía en esas columnas y SIN combinar.

    Hace falta porque el pie de estos libros («ChefBusiness.co — Plan de
    Negocio: …») va COMBINADO a lo ancho de la tabla, y una `MergedCell` tiene
    el `value` de sólo lectura: escribir el contador encima reventaba con
    `AttributeError`. Se busca sitio en vez de romper la combinación, que es
    maquetación legítima del fichero.
    """
    combinadas = _filas_combinadas(ws)
    r = desde
    while r < desde + 200:
        if r not in combinadas and all(
                ws.cell(row=r, column=c).value is None for c in columnas):
            return r
        r += 1
    return ws.max_row + 2


def _es_numero(v):
    return isinstance(v, (int, float)) and not isinstance(v, bool)


def _es_formula(v):
    return isinstance(v, str) and v.startswith('=')


# ==========================================================================
# §1.3 — Constantes tecleadas → fórmulas (conservando el ejemplo)
# ==========================================================================
RX_TOTAL = re.compile(
    r'^(total|inversion total|inversi[oó]n total|subtotal|suma)\b', re.I)
RX_CAB_TABLA = re.compile(
    r'^(concepto|partida|variable|m[eé]trica|puesto|ratio|fase|item|'
    r'hito|proveedor)\b', re.I)
#: Rótulos que NO son una suma de lo de arriba aunque empiecen por «TOTAL»
#: (los resuelven los grupos con criterio, no una suma ciega).
RX_TOTAL_NO = re.compile(r'total\s+(ingresos?\s*[-−]|resultado)', re.I)


def _inicio_bloque(ws, cab, fila_total, excluir):
    """Primera fila de partidas del bloque al que pertenece esa fila TOTAL.

    Sube hasta el primer separador: otro total, un encabezado de bloque (texto
    sin cifras) o una fila en blanco. Es lo que impide que la suma de «TOTAL
    COSTES FIJOS» se coma el «MARGEN BRUTO» que hay más arriba.
    """
    for r in range(fila_total - 1, cab, -1):
        if r in excluir:
            return r + 1
        hay_numero = False
        for c in range(2, ws.max_column + 1):
            v = ws.cell(row=r, column=c).value
            if _es_numero(v) or _es_formula(v):
                hay_numero = True
                break
        if not hay_numero:
            return r + 1
    return cab + 1


def _segmentos_numericos(ws, col, r0, r1, excluir):
    """Tramos contiguos de celdas numéricas, saltando huecos y subtotales."""
    segmentos, ini = [], None
    for r in range(r0, r1 + 1):
        v = ws.cell(row=r, column=col).value
        bueno = _es_numero(v) and r not in excluir
        if bueno and ini is None:
            ini = r
        elif not bueno and ini is not None:
            segmentos.append((ini, r - 1))
            ini = None
    if ini is not None:
        segmentos.append((ini, r1))
    return segmentos


def cablear_sumas(ws, informe, desajustes):
    """Las filas TOTAL pasan de constante a `=SUM(...)` — §1.3.

    Cautela deliberada: **sólo se cabla cuando la aritmética YA cuadra**
    (tolerancia 0,50 € o 0,5 %). Un total que no cuadra con sus partidas es un
    defecto de contenido, no de formato: cablearlo cambiaría en silencio la
    cifra que el cliente ve y que la landing publica. Se anota como desajuste
    para que el grupo lo resuelva con criterio (§9 lo cruza).
    """
    cab, _ = _fila_cabecera(ws)
    if cab is None:
        return 0
    filas_total = []
    for r in range(cab + 1, ws.max_row + 1):
        rot = _rotulo_de_fila(ws, r)
        if rot and RX_TOTAL.match(rot) and not RX_TOTAL_NO.search(rot):
            filas_total.append(r)
    if not filas_total:
        return 0
    excluir = set(filas_total)
    for r in range(cab + 1, ws.max_row + 1):
        rot = _rotulo_de_fila(ws, r)
        if rot and RX_RESULTADO.match(rot):
            excluir.add(r)          # un margen o un EBITDA no es una partida
    puestas = 0
    for col in range(2, ws.max_column + 1):
        for r in filas_total:
            # El bloque empieza en su ENCABEZADO («COSTES FIJOS»), no en el
            # total anterior: arrancando en el total anterior, la suma de
            # `'2. P&L 3 Anos'!B30` se tragaba el MARGEN BRUTO de B17 y daba
            # 374.047 € contra los 198.108 € tecleados — un «descuadre» que
            # no existe y que habría mandado a los grupos a investigar humo.
            anterior = _inicio_bloque(ws, cab, r, excluir) - 1
            v = ws.cell(row=r, column=col).value
            if _es_formula(v) or not _es_numero(v):
                continue
            segmentos = _segmentos_numericos(ws, col, anterior + 1, r - 1,
                                             excluir)
            if not segmentos:
                continue
            suma = 0.0
            for a, b in segmentos:
                for rr in range(a, b + 1):
                    vv = ws.cell(row=rr, column=col).value
                    if _es_numero(vv):
                        suma += float(vv)
            # La tolerancia tiene que ser RELATIVA salvo en magnitudes
            # grandes: con 0,50 € fijos, una columna de porcentajes daba por
            # buena una «suma» de 0,52 contra un 0,58 tecleado y cableaba una
            # fórmula falsa en la fila de TOTAL COSTES VARIABLES. Medido en
            # `plan-financiero-cocteleria-eventos.xlsx:'PyG 3 Anos'!E18`.
            magnitud = abs(float(v))
            tolerancia = (max(0.5, magnitud * 0.005) if magnitud >= 100
                          else max(magnitud * 0.005, 1e-9))
            letra = get_column_letter(col)
            partes = ['SUM(' + letra + str(a) + ':' + letra + str(b) + ')'
                      for a, b in segmentos]
            formula = '=' + '+'.join(partes)
            if abs(suma - float(v)) <= tolerancia:
                fmt = ws.cell(row=r, column=col).number_format
                f(ws, letra + str(r), formula, fmt)
                puestas += 1
            else:
                desajustes.append({
                    'hoja': ws.title, 'celda': letra + str(r),
                    'rotulo': _rotulo_de_fila(ws, r),
                    'valor_tecleado': float(v), 'suma_de_partidas': round(suma, 2),
                    'diferencia': round(float(v) - suma, 2),
                    'formula_propuesta': formula})
    if puestas and informe is not None:
        informe.append(ws.title + ': ' + str(puestas)
                       + ' totales cableados con SUM (§1.3)')
    return puestas


# ==========================================================================
# Pseudo-fórmulas: notas que empiezan por «=» y Excel intenta calcular
# ==========================================================================
#: Medido en el representante A: `'3. Punto Equilibrio'!C14` guarda el TEXTO
#: `'= CF anuales / MC unitario'` y `C15`, `'= 17950 / 310 dias'`. openpyxl —y
#: Excel— los tratan como fórmula: el cliente ve `#¿NOMBRE?` en la columna de
#: notas de la hoja estrella. Y para el motor son trampas dobles: `guardas()`
#: las envolvía en `IFERROR(...)`, produciendo una fórmula inválida que pycel
#: no puede evaluar (`fallos_pycel=1`) y que se entrega sin caché.
FUNCIONES = frozenset((
    'SUM', 'SUMA', 'IF', 'SI', 'IFERROR', 'SUMIF', 'SUMIFS', 'COUNTIF',
    'COUNTIFS', 'COUNT', 'MAX', 'MIN', 'ROUND', 'ROUNDUP', 'ROUNDDOWN',
    'ABS', 'AVERAGE', 'PROMEDIO', 'NPV', 'VNA', 'INDEX', 'MATCH', 'TEXT',
    'TEXTO', 'AND', 'OR', 'NOT', 'ISNUMBER', 'ISERROR', 'SEARCH', 'LEN',
    'CONCATENATE', 'SUMPRODUCT', 'POWER', 'EXP', 'LN', 'YEAR', 'TODAY',
    'TRUE', 'FALSE'))
RX_TOKEN = re.compile(r'[A-Za-zÁÉÍÓÚÑáéíóúñ_][A-Za-z0-9ÁÉÍÓÚÑáéíóúñ_.]*')
RX_REF_SOLA = re.compile(r'^\$?[A-Z]{1,3}\$?\d+$')


def parece_formula_real(v):
    """¿Es una fórmula de Excel o una NOTA que empieza por «=»?

    Regla: fuera las cadenas entrecomilladas y las referencias con hoja; lo que
    quede en letras tiene que ser una función conocida o una referencia de
    celda. «CF anuales / MC unitario» falla en `anuales` y se declara nota.
    """
    if not isinstance(v, str) or not v.startswith('='):
        return False
    cuerpo = re.sub(r'"[^"]*"', '', v[1:])
    cuerpo = re.sub(r"'[^']*'!", '', cuerpo)
    cuerpo = re.sub(r'[A-Za-z_][A-Za-z0-9_.]*!', '', cuerpo)
    # El `$` de las referencias absolutas se quita ANTES de tokenizar: sin
    # esto, `=D6/$D$15` se troceaba en «D6» (referencia) y «D» (letra suelta),
    # la letra suelta no era ni función ni referencia y la fórmula se
    # declaraba «nota» — el motor le quitaba el `=` y borraba de un golpe las
    # 9 fórmulas de porcentaje de `'Resumen'!E6:E14`.
    cuerpo = cuerpo.replace('$', '')
    for token in RX_TOKEN.findall(cuerpo):
        if token.upper() in FUNCIONES:
            continue
        if RX_REF_SOLA.match(token.upper()):
            continue
        return False
    return True


def neutralizar_pseudoformulas(ws, informe, detalle):
    """La nota vuelve a ser texto (sin el «=» que Excel intenta calcular)."""
    n = 0
    for row in ws.iter_rows():
        for c in row:
            v = c.value
            if not isinstance(v, str) or not v.startswith('='):
                continue
            if parece_formula_real(v):
                continue
            nuevo = v[1:].strip()
            c.value = nuevo
            detalle.append((ws.title, c.coordinate, v[:70], nuevo[:70]))
            n += 1
    if n and informe is not None:
        informe.append(ws.title + ': ' + str(n) + ' pseudo-fórmulas '
                       'convertidas en texto (habrían dado #¿NOMBRE?)')
    return n


# ==========================================================================
# §1.5 — Guardas: IFERROR en toda división
# ==========================================================================
RX_YA_GUARDADA = re.compile(r'^=\s*IFERROR\s*\(', re.I)


def guardas(ws, informe):
    puestas = 0
    for row in ws.iter_rows():
        for c in row:
            v = c.value
            if not _es_formula(v) or '/' not in v:
                continue
            if RX_YA_GUARDADA.match(v) or not parece_formula_real(v):
                continue
            # una división dentro de un texto entrecomillado no es división
            cuerpo = re.sub(r'"[^"]*"', '', v)
            if '/' not in cuerpo:
                continue
            c.value = iferror(v)
            reg(ws, c.coordinate, c.value)
            puestas += 1
    if puestas and informe is not None:
        informe.append(ws.title + ': ' + str(puestas)
                       + ' divisiones con IFERROR (§1.5)')
    return puestas


# ==========================================================================
# §1.4 — Formatos por tipo de dato
# ==========================================================================
#: Rótulos que cuentan unidades: nunca llevan formato de euro.
RX_RECUENTO = re.compile(
    r'\b(eventos?|clientes?|cubiertos?|personas?|d[ií]as?|meses|mes\b|'
    r'unidades?|comensales?|pax|invitados?|plazas?|n[uú]mero de|pagas)\b', re.I)
#: Rótulos de importe: nunca llevan formato de porcentaje.
RX_IMPORTE = re.compile(
    r'\b(anticipo|precio|coste|costes|ingresos?|facturaci[oó]n|ticket|'
    r'importe|salario|cuota|inversi[oó]n|alquiler|resultado|ebitda|'
    r'presupuesto)\b', re.I)
#: …salvo que el rótulo declare que es un RATIO.
RX_RATIO = re.compile(r'\(\s*%\s*\)|%\s*$|/\s*(ingresos|ventas|facturaci)',
                      re.I)
#: Rótulos que son euros aunque contengan una palabra de recuento
#: («Ingresos por eventos», «Coste por invitado»).
RX_EURO_FUERTE = re.compile(
    r'\b(€|eur|ingresos?|coste|costes|precio|importe|facturaci[oó]n|ticket|'
    r'margen|salario|cuota)\b', re.I)

RX_TXT_PCT = re.compile(r'^\s*(\d{1,3}(?:[.,]\d+)?)\s*%\s*$')
RX_TXT_EUR = re.compile(
    r'^\s*(\d{1,3}(?:\.\d{3})+|\d+(?:[.,]\d{1,2})?)\s*(?:EUR|€)\s*$')
RX_RANGO = re.compile(r'\d\s*[-–—' + GUION + r']\s*\d')


def _a_numero_texto(v):
    """Magnitud escalar guardada como texto → (número, formato) o None.

    Los RANGOS («25-30 %», «4-6 EUR», «5-15K EUR») NO se convierten: son
    referencia, no dato. En los porcentajes el punto es decimal («66.1%» →
    0,661); en los euros con tres cifras detrás es separador de millares
    («1.200 €» → 1200).
    """
    if not isinstance(v, str) or RX_RANGO.search(v):
        return None
    m = RX_TXT_PCT.match(v)
    if m:
        return float(m.group(1).replace(',', '.')) / 100.0, FMT_PCT
    m = RX_TXT_EUR.match(v)
    if m:
        crudo = m.group(1)
        if re.match(r'^\d{1,3}(\.\d{3})+$', crudo):
            return float(crudo.replace('.', '')), FMT_EUR0
        return float(crudo.replace(',', '.')), FMT_EUR0
    return None


def formatos_por_tipo(ws, informe, cambios_detalle):
    """§1.4 — el formato lo decide el TIPO DE DATO, no el bloque."""
    n = 0
    for row in ws.iter_rows():
        for c in row:
            v = c.value
            fmt = c.number_format or ''
            rot = _rotulo_de_fila(ws, c.row)

            # (a) '#,##0 €/h' con la h sin entrecomillar: openpyxl lee
            #     datetime(1900,1,22) y el cliente ve una fecha.
            if '/h' in fmt and '"' not in fmt:
                c.number_format = FMT_EURH
                cambios_detalle.append((ws.title, c.coordinate, 'fmt €/h',
                                        fmt, FMT_EURH))
                n += 1
                fmt = FMT_EURH

            # (b) magnitudes guardadas como texto
            conv = _a_numero_texto(v)
            if conv is not None and c.column > 1:
                nuevo, nfmt = conv
                c.value = nuevo
                c.number_format = nfmt
                cambios_detalle.append((ws.title, c.coordinate, 'texto→número',
                                        v, nuevo))
                n += 1
                v, fmt = nuevo, nfmt

            if not (_es_numero(v) or _es_formula(v)) or c.column == 1:
                continue

            cab_col = cabecera_de_columna(ws, c.column)
            col_ratio = bool(RX_RATIO.search(cab_col)
                             or re.search(r'%|ratio', cab_col, re.I))
            # La CABECERA manda sobre el rótulo en el eje que ella define: la
            # columna «Premium (€)» es de euros aunque la fila se llame
            # «Vajilla eventos…» y la palabra «eventos» huela a recuento.
            col_euro = bool(re.search(r'€|\beur\b|importe|coste|precio',
                                      cab_col, re.I))
            # …y la cabecera también DEFINE el recuento: en `'5. Personal'` el
            # rótulo de fila es «Gerente / Propietario» y quien dice que la
            # columna B son personas es su cabecera, «Personas». Sin esto, la
            # plantilla imprime «7 €» donde hay 7 trabajadores (TEC-16).
            col_recuento = bool(RX_RECUENTO.search(cab_col)) and not col_euro
            recuento = ((bool(RX_RECUENTO.search(rot))
                         and not RX_EURO_FUERTE.search(rot)) or col_recuento) \
                and not col_ratio and not col_euro
            importe = bool(RX_IMPORTE.search(rot)) and not \
                RX_RATIO.search(rot) and not col_ratio

            # (c) recuento con formato de euro → «45 €» donde pone clientes/día
            if recuento and '€' in fmt:
                nuevo = FMT_DEC if (_es_numero(v)
                                    and float(v) != int(float(v))) else FMT_ENT
                c.number_format = nuevo
                cambios_detalle.append((ws.title, c.coordinate,
                                        'recuento sin €', fmt, nuevo))
                n += 1
                continue

            # (d) importe con formato de porcentaje → «48348,0 %» en el
            #     anticipo que el cliente copia al contrato. Con constante,
            #     el discriminante es el valor: un ratio va en 0-2.
            if importe and '%' in fmt:
                if _es_formula(v) or abs(float(v)) > 5:
                    c.number_format = FMT_EUR
                    cambios_detalle.append((ws.title, c.coordinate,
                                            'importe sin %', fmt, FMT_EUR))
                    n += 1
                    continue

            # (e) decimales significativos ocultos: 18,5 impreso «19 €»
            if _es_numero(v) and fmt == FMT_EUR0 and float(v) != int(float(v)):
                c.number_format = FMT_EUR
                cambios_detalle.append((ws.title, c.coordinate,
                                        'decimales visibles', fmt, FMT_EUR))
                n += 1
                continue

            # (f) `General` en cifras que son el resultado estrella
            if fmt == 'General' and _es_numero(v):
                if recuento:
                    nuevo = FMT_DEC if float(v) != int(float(v)) else FMT_ENT
                elif importe or RX_EURO_FUERTE.search(rot):
                    nuevo = FMT_EUR if float(v) != int(float(v)) else FMT_EUR0
                else:
                    nuevo = FMT_DEC2 if float(v) != int(float(v)) else FMT_ENT
                c.number_format = nuevo
                cambios_detalle.append((ws.title, c.coordinate,
                                        'General→formato', fmt, nuevo))
                n += 1
    n += _armonizar_euros_por_fila(ws, cambios_detalle)
    if n and informe is not None:
        informe.append(ws.title + ': ' + str(n) + ' formatos corregidos (§1.4)')
    return n


def _armonizar_euros_por_fila(ws, cambios_detalle):
    """Si un euro de la fila muestra decimales, TODOS los de esa fila.

    Sin esto la fila «Resultado antes impuestos» de `4. Escenarios` imprime
    «198.000 €», «−50.658,75 €» y «42.600 €»: la misma magnitud con dos
    precisiones distintas en tres celdas contiguas.
    """
    n = 0
    for row in ws.iter_rows():
        euros = [c for c in row
                 if '€' in (c.number_format or '')
                 and (_es_numero(c.value) or _es_formula(c.value))]
        if len(euros) < 2:
            continue
        if not any((c.number_format or '').startswith(FMT_EUR) for c in euros):
            continue
        for c in euros:
            if (c.number_format or '') != FMT_EUR:
                cambios_detalle.append((ws.title, c.coordinate,
                                        'euros armonizados',
                                        c.number_format, FMT_EUR))
                c.number_format = FMT_EUR
                n += 1
    return n


# ==========================================================================
# §1.7 — Ortografía: acentos, eñes y erratas
# ==========================================================================
#: Sustituciones con frontera de palabra. La clave va SIN tilde y el valor CON
#: ella; el reemplazo respeta mayúscula inicial y MAYÚSCULAS completas.
TILDES = {
    'ano': 'año', 'anos': 'años', 'anual': None, 'espana': 'España',
    'espanol': 'español', 'espanola': 'española', 'espanoles': 'españoles',
    'analisis': 'análisis', 'constitucion': 'constitución',
    'tramite': 'trámite', 'tramites': 'trámites', 'diseno': 'diseño',
    'disenos': 'diseños',
    # CRIT-05 — el diccionario tenía el sustantivo («diseño») pero no el
    # OFICIO ni el VERBO, y el corpus los usa: «Disenador» en 8 celdas de 6
    # ficheros, «Disenar identidad visual» en 2 y «Disena carta tapas» en 1.
    # El gate de ortografía salía en 0 sobre un fichero que el cliente imprime
    # con la errata delante.
    'disenador': 'diseñador', 'disenadora': 'diseñadora',
    'disenadores': 'diseñadores', 'disenadoras': 'diseñadoras',
    'disenar': 'diseñar', 'disena': 'diseña', 'disenan': 'diseñan',
    'disenado': 'diseñado', 'disenada': 'diseñada',
    'disenados': 'diseñados', 'disenadas': 'diseñadas',
    'nomina': 'nómina', 'nominas': 'nóminas',
    'amortizacion': 'amortización', 'comision': 'comisión',
    'comisiones': 'comisiones', 'gestoria': 'gestoría',
    'bolleria': 'bollería', 'cumpleanos': 'cumpleaños', 'danos': 'daños',
    'cataluna': 'Cataluña', 'resenas': 'reseñas', 'resena': 'reseña',
    'desempeno': 'desempeño', 'alergenos': 'alérgenos',
    'inversion': 'inversión', 'proyeccion': 'proyección',
    'facturacion': 'facturación', 'situacion': 'situación',
    'previsional': None, 'estimacion': 'estimación',
    'estimaciones': 'estimaciones', 'informacion': 'información',
    'formacion': 'formación', 'operacion': 'operación',
    'operaciones': 'operaciones', 'ocupacion': 'ocupación',
    'facturas': None, 'senaletica': 'señalética', 'ninos': 'niños',
    'canas': 'cañas', 'manana': 'mañana', 'pequeno': 'pequeño',
    'pequena': 'pequeña', 'ensenar': 'enseñar', 'companero': 'compañero',
    'campana': None,        # «campana extractora» es legítima: caso aparte
    'extraccion': 'extracción', 'refrigeracion': 'refrigeración',
    'climatizacion': 'climatización', 'iluminacion': 'iluminación',
    'decoracion': 'decoración', 'rotulacion': 'rotulación',
    'instalacion': 'instalación', 'instalaciones': 'instalaciones',
    'certificacion': 'certificación', 'inscripcion': 'inscripción',
    'denominacion': 'denominación', 'eleccion': 'elección',
    'cotizacion': 'cotización', 'adecuacion': 'adecuación',
    'evacuacion': 'evacuación', 'fontaneria': 'fontanería',
    'panaderia': 'panadería', 'pasteleria': 'pastelería',
    'cafeteria': 'cafetería', 'heladeria': 'heladería',
    'cocteleria': 'coctelería', 'cristaleria': 'cristalería',
    'cuchilleria': 'cuchillería', 'vajilla': None, 'cuberteria': 'cubertería',
    'carniceria': 'carnicería', 'hosteleria': 'hostelería',
    'telefonia': 'telefonía', 'categoria': 'categoría',
    'categorias': 'categorías', 'indice': 'índice', 'numero': 'número',
    'minimo': 'mínimo', 'minima': 'mínima', 'minimas': 'mínimas',
    'maximo': 'máximo', 'maxima': 'máxima', 'basico': 'básico',
    'basica': 'básica', 'tecnico': 'técnico', 'tecnica': 'técnica',
    'practico': 'práctico', 'automatico': 'automático',
    'organico': 'orgánico', 'economico': 'económico',
    'juridica': 'jurídica', 'publica': 'pública', 'publico': 'público',
    'electrica': 'eléctrica', 'acustica': 'acústica',
    'energetica': 'energética', 'logistica': 'logística',
    'estrategia': None, 'garantia': 'garantía', 'garantias': 'garantías',
    'dia': 'día', 'dias': 'días', 'mas': None, 'segun': 'según',
    'ademas': 'además', 'tambien': 'también', 'aqui': 'aquí',
    'asi': None, 'esta': None, 'este': None,
}
#: Excepción legítima: la campana de extracción convive en el mismo libro con
#: «Campaña lanzamiento RRSS», que sí es errata. Se resuelve por CONTEXTO.
EXCEPCIONES_CAMPANA = ('campana extractora', 'campana de extraccion',
                       'campana extraccion', 'campana industrial')

#: Erratas puntuales medidas (TEC-26, TEC-23).
ERRATAS = (
    ('Priorizarcexperiencia hosteleria', 'Priorizar experiencia hostelería'),
    ('Priorizarcexperiencia', 'Priorizar experiencia'),
    ('6. Frutas, herbas y garnish', '6. Frutas y garnish'),
    ('horno de convención', 'horno de convección'),
    ('mudança', 'mudanza'),
    # RC-20 — abreviatura sin punto ni tilde en la nota del capital social
    ('Capital social min 1', 'Capital social mín. 1'),
    ('capital social min 1', 'capital social mín. 1'),
)

#: RC-31 — dos convenciones tipográficas conviviendo en el mismo libro.
#: `EUR` detrás de una cifra pasa a `€` (el resto del libro ya usa `€`) y el
#: signo menos matemático U+2212 —que NO está en WinAnsi y se pierde al
#: exportar a PDF— pasa a guion normal. Los dos caracteres se referencian por
#: ESCAPE, como exige la nota de U+202F/U+2011 de la familia.
MENOS_MAT = u'\u2212'
RX_EUR_TEXTO = re.compile(r'(?<=[\d\s])EUR\b')

RX_PALABRA = re.compile(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+")
#: Lo que la ortografía NO toca: URLs, correos y **slugs** (`plan-negocio-
#: cocteleria-eventos`). El slug importa: la línea de versión y la de
#: instrucciones llevan el pid dentro, y «acentuarlo» produce un id que no
#: existe — además de romper la idempotencia, porque la 2.ª pasada corrige lo
#: que la 1.ª acababa de escribir.
RX_URL_MAIL = re.compile(r'(https?://\S+|www\.\S+|\S+@\S+|\S+\.(?:com|es|pro|'
                         r'co|net|org)\b|[a-z0-9]+(?:-[a-z0-9]+){2,})', re.I)


#: Regla generativa: en español NO existe palabra llana acabada en «-cion» o
#: «-sion» sin tilde (constitución, captación, retención, previsión). Cubre las
#: decenas de casos que el diccionario explícito no lista sin inventarse nada:
#: sólo actúa si la palabra no lleva ya una tilde.
RX_CION = re.compile(r'(?i)^([a-zñ]+)(ci|si)on$')
RX_YA_ACENTUADA = re.compile(r'[áéíóúÁÉÍÓÚ]')
#: Palabras sueltas que la regla generativa no alcanza.
TILDES_EXTRA = {
    'autonomo': 'autónomo', 'autonomos': 'autónomos',
    'regimen': 'régimen', 'regimenes': 'regímenes',
    'credito': 'crédito', 'debito': 'débito', 'margenes': 'márgenes',
    'articulo': 'artículo', 'periodo': 'periodo', 'ultimo': 'último',
    'ultima': 'última', 'proximo': 'próximo', 'proxima': 'próxima',
    'sabado': 'sábado', 'miercoles': 'miércoles', 'telefono': 'teléfono',
    'kilometro': 'kilómetro', 'kilometros': 'kilómetros',
    'metodo': 'método', 'analisis': 'análisis', 'organico': 'orgánico',
    'consultoria': 'consultoría', 'metrica': 'métrica',
    'metricas': 'métricas', 'facturacion': 'facturación',
    'valoracion': 'valoración', 'logistico': 'logístico',
    'itinerante': 'itinerante', 'tematico': 'temático',
    'tematica': 'temática', 'nomina': 'nómina',
    # -ico/-ica: esdrújulas frecuentes en este corpus. NO se generaliza por
    # sufijo (rompería «chica», «rica», «publica» del verbo publicar): se
    # listan una a una las que aparecen medidas en los 30 ficheros.
    'gastronomico': 'gastronómico', 'gastronomica': 'gastronómica',
    'gastronomicos': 'gastronómicos', 'gastronomicas': 'gastronómicas',
    'economico': 'económico', 'economica': 'económica',
    'estrategico': 'estratégico', 'estrategica': 'estratégica',
    'especifico': 'específico', 'especifica': 'específica',
    'electronico': 'electrónico', 'electronica': 'electrónica',
    'informatico': 'informático', 'informatica': 'informática',
    'logistica': 'logística', 'domestico': 'doméstico',
    'domestica': 'doméstica', 'higienico': 'higiénico',
    'higienica': 'higiénica',
    # RD-32 / RC-20 — nueve palabras que el gate daba por buenas y que el
    # cliente ve en la hoja que imprime para el banco. Cinco de ellas
    # convivían con su forma acentuada en la MISMA celda («Pagina básica»,
    # «Constitución SL (notaria)»): por eso `gate_ortografia` incorpora
    # además la heurística de convivencia (`_convive_acentuada`).
    'camara': 'cámara', 'camaras': 'cámaras',
    'frigorifico': 'frigorífico', 'frigorifica': 'frigorífica',
    'frigorificos': 'frigoríficos', 'frigorificas': 'frigoríficas',
    'estanteria': 'estantería', 'estanterias': 'estanterías',
    'almacen': 'almacén',
    'cafe': 'café', 'cafes': 'cafés',
    'pagina': 'página', 'paginas': 'páginas',
    'notaria': 'notaría',
    'busqueda': 'búsqueda', 'busquedas': 'búsquedas',
    'consejeria': 'consejería', 'consejerias': 'consejerías',
    'codigo': 'código', 'codigos': 'códigos',
    'clausula': 'cláusula', 'clausulas': 'cláusulas',
    # CRIT-05 — «crítico/crítica/críticos/críticas» llevan tilde SIEMPRE como
    # adjetivo y como sustantivo («la crítica»). Estaban clasificados como
    # homógrafos y el gate los perdonaba: cuatro erratas vivas en el checklist
    # («Critico para SEO local», «Primeras 4 semanas criticas») y 14 más en
    # los hermanos. Lo único sin tilde es el VERBO criticar («se critica el
    # servicio»), que no aparece ni una vez en los 30 xlsx (censo del
    # 2026-08-29) y que `RX_CRITICA_VERBO` protege igualmente.
    'critico': 'crítico', 'critica': 'crítica',
    'criticos': 'críticos', 'criticas': 'críticas',
}
#: Excepción documentada de `critica`/`criticas`: formas del verbo «criticar».
#: Se reconocen por el clítico o el relativo que las precede («se critica»,
#: «que critica», «lo critica», «la criticas»). Sin este guardián, acentuar
#: siempre sería correcto en todo el corpus medido, pero convertiría un futuro
#: «no se critica al proveedor» en «no se crítica al proveedor».
RX_CRITICA_VERBO = re.compile(
    r'(?i)\b(?:se|que|no|lo|la|le|les|me|te|nos|os)\s+criticas?\b')


def _reponer_caso(original, corregido):
    if original.isupper():
        return corregido.upper()
    if original[:1].isupper():
        return corregido[:1].upper() + corregido[1:]
    return corregido


def corregir_texto(texto):
    """Devuelve el texto con tildes y eñes puestas, respetando URLs y correos."""
    if not isinstance(texto, str) or not texto.strip():
        return texto
    for malo, bueno in ERRATAS:
        if malo in texto:
            texto = texto.replace(malo, bueno)
    # RC-31 — homogeneización tipográfica (ver MENOS_MAT / RX_EUR_TEXTO)
    if MENOS_MAT in texto:
        texto = texto.replace(MENOS_MAT, '-')
    if 'EUR' in texto:
        texto = RX_EUR_TEXTO.sub(u'\u20ac', texto)
    trozos = RX_URL_MAIL.split(texto)
    fuera = []
    for i, trozo in enumerate(trozos):
        if i % 2 == 1 or not trozo:
            fuera.append(trozo)
            continue
        bajo = norm(trozo)
        # CRIT-05 — tramos donde «critica/criticas» es el VERBO criticar y por
        # tanto NO lleva tilde. Se calculan sobre el trozo original porque
        # `_sub` recibe posiciones de ese mismo trozo.
        verbal = [(m.start(), m.end())
                  for m in RX_CRITICA_VERBO.finditer(trozo)]

        def _sub(m):
            palabra = m.group(0)
            clave = norm(palabra)
            if clave.startswith('campana'):
                return palabra
            if clave in ('critica', 'criticas') \
                    and any(a <= m.start() < b for a, b in verbal):
                return palabra
            correcto = TILDES.get(clave) or TILDES_EXTRA.get(clave)
            if not correcto and not RX_YA_ACENTUADA.search(palabra):
                mm = RX_CION.match(clave)
                if mm:
                    # `capta` + `ci` + `ón` = «captación». Ojo: quedarse con
                    # `mm.group(2)[0]` daba «captacón» — probado antes de
                    # soltarlo sobre los 30 ficheros.
                    correcto = mm.group(1) + mm.group(2) + 'ón'
            if not correcto:
                return palabra
            return _reponer_caso(palabra, correcto)

        nuevo = RX_PALABRA.sub(_sub, trozo)
        # «campana» sólo se corrige a «campaña» cuando NO es la extractora
        if 'campana' in bajo:
            def _campana(m):
                ini = max(0, m.start() - 0)
                cola = norm(nuevo[m.end():m.end() + 20])
                contexto = norm(m.group(0)) + ' ' + cola
                for exc in EXCEPCIONES_CAMPANA:
                    if contexto.startswith(norm(exc)):
                        return m.group(0)
                return _reponer_caso(m.group(0), 'campaña')
            nuevo = re.sub(r'(?i)\bcampanas?\b', _campana, nuevo)
        fuera.append(nuevo)
    return ''.join(fuera)


def ortografia(ws, informe, detalle):
    n = 0
    for row in ws.iter_rows():
        for c in row:
            v = c.value
            if not isinstance(v, str) or v.startswith('='):
                continue
            nuevo = corregir_texto(v)
            if nuevo != v:
                c.value = nuevo
                detalle.append((ws.title, c.coordinate, v[:60], nuevo[:60]))
                n += 1
    if n and informe is not None:
        informe.append(ws.title + ': ' + str(n) + ' textos con tildes (§1.7)')
    return n


# -- renombrado de hojas con reescritura de referencias --------------------
RX_REF_HOJA = re.compile(r"(?:'([^']+)'|([A-Za-z_][A-Za-z0-9_.]*))!")


def _reescribir_referencias(wb, viejo, nuevo):
    """Reescribe `Hoja!A1` y `'Hoja'!A1` en TODA fórmula del libro."""
    tocadas = 0
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                v = c.value
                if not _es_formula(v):
                    continue

                def _sub(m):
                    nombre = m.group(1) or m.group(2)
                    if nombre != viejo:
                        return m.group(0)
                    return "'" + nuevo + "'!"

                nuevo_v = RX_REF_HOJA.sub(_sub, v)
                if nuevo_v != v:
                    c.value = nuevo_v
                    tocadas += 1
    return tocadas


def _reescribir_enlaces_internos(wb, viejo, nuevo):
    """Los hipervínculos internos y los RÓTULOS que citan el nombre de la hoja.

    Renombrar sin esto deja el índice de la plantilla de proveedores
    apuntando a `#'3. Cristaleria'!A1`, una pestaña que ya no existe: el enlace
    muere en silencio (Excel no avisa) y encima la 2.ª pasada lo arregla, con
    lo que la idempotencia salta y delata que la 1.ª entregó mal.
    """
    tocados = 0
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if c.hyperlink is not None:
                    destino = getattr(c.hyperlink, 'location', None) \
                        or c.hyperlink.target or ''
                    if viejo in str(destino):
                        c.hyperlink = "#'" + nuevo + "'!A1"
                        tocados += 1
                # Comparación por AMBOS extremos recortados: hay pestañas con
                # espacio final («4. Cuchilleria, tablas, menaje ») y con
                # `c.value.strip() == viejo` el rótulo del índice nunca casaba
                # → la 1.ª pasada lo dejaba sin tilde y la 2.ª lo corregía.
                if isinstance(c.value, str) \
                        and c.value.strip() == viejo.strip():
                    c.value = nuevo
                    tocados += 1
    return tocados


def renombrar_hojas(wb, informe):
    """§1.7 — los nombres de hoja también llevan tilde.

    Va DESPUÉS del cableado: renombrar obliga a reescribir las referencias
    entre hojas que el motor y los grupos acaban de crear. `gate_referencias`
    re-verifica que no queda ninguna colgando.
    """
    hechos = {}
    for ws in list(wb.worksheets):
        viejo = ws.title
        nuevo = corregir_texto(viejo)
        if nuevo == viejo:
            continue
        if len(nuevo) > 31:
            continue
        if hoja(wb, nuevo) is not None and norm(nuevo) != norm(viejo):
            continue
        ws.title = nuevo
        _reescribir_referencias(wb, viejo, nuevo)
        _reescribir_enlaces_internos(wb, viejo, nuevo)
        hechos[viejo] = nuevo
        RENOMBRES[viejo] = nuevo
        if informe is not None:
            informe.append('hoja renombrada: ' + viejo + ' → ' + nuevo)
    # El REGISTRO guarda el nombre que la hoja tenía al escribir la fórmula.
    for i, (h, coord, formula) in enumerate(REGISTRO):
        if h in hechos:
            REGISTRO[i] = (hechos[h], coord, formula)
    return hechos


# ==========================================================================
# §1.8 — Anchos, altos y celdas combinadas con wrapText
# ==========================================================================
ANCHO_POR_DEFECTO = 8.43
ALTO_LINEA = 15.0


#: RC-15 / RT-11 — 86 celdas de TEXTO con formato de euro o de porcentaje: la
#: columna «Notas» de la inversión, la de comentarios del P&L y la de los
#: escenarios. `gate_formatos` no las veía porque saltaba las celdas de texto
#: (`motor.py`, «if not (_es_numero(v) or _es_formula(v))»), así que el run
#: informaba «formatos: 0» y era un verde falso.
RX_FORMULA_TEXTO = re.compile(r'^=\s*(IFERROR\s*\(\s*)?"'      # ="…"
                              r'|&\s*TEXT\s*\(', re.I)


def _devuelve_texto(v):
    """¿La celda imprime TEXTO? Vale para el literal y para la fórmula."""
    if isinstance(v, str) and not v.startswith('='):
        return True
    if isinstance(v, str) and RX_FORMULA_TEXTO.search(v):
        return True
    return False


def formatos_texto(ws, informe=None):
    """Pone `General` en toda celda cuyo contenido se imprime como texto."""
    n = 0
    for row in ws.iter_rows():
        for c in row:
            fmt = c.number_format or ''
            if fmt in ('General', '@') or not _devuelve_texto(c.value):
                continue
            if '\u20ac' not in fmt and '%' not in fmt:
                continue
            c.number_format = 'General'
            n += 1
    if n and informe is not None:
        informe.append(ws.title + ': ' + str(n) + ' celdas de texto con '
                       'formato numérico → General (§1.4)')
    return n


def _ancho_de(ws, col):
    d = ws.column_dimensions.get(get_column_letter(col))
    if d is not None and d.width:
        return d.width
    return ANCHO_POR_DEFECTO


def altos_y_wrap(ws, informe):
    """Alto = ceil(len / ancho) × 15 pt en toda combinada con `wrapText`.

    Medido: la `CONCLUSION` de `'Punto Equilibrio'!A18:C18` tiene 313
    caracteres y se ve 1 línea de 4; la nota de autoría de `Indice!B4:D7`, 606
    caracteres en 4 filas que necesitan 9.
    """
    n = 0
    for rango in [str(m) for m in ws.merged_cells.ranges]:
        cr = CellRange(rango)
        cel = ws.cell(row=cr.min_row, column=cr.min_col)
        texto = cel.value
        if not isinstance(texto, str) or len(texto) < 40:
            continue
        cel.alignment = Alignment(horizontal=cel.alignment.horizontal or 'left',
                                  vertical='top', wrap_text=True)
        ancho = sum(_ancho_de(ws, c) for c in range(cr.min_col, cr.max_col + 1))
        lineas = max(1, int(math.ceil(len(texto) / max(8.0, ancho * 1.05))))
        filas = cr.max_row - cr.min_row + 1
        alto = max(ALTO_LINEA, ALTO_LINEA * lineas / float(filas))
        for r in range(cr.min_row, cr.max_row + 1):
            actual = ws.row_dimensions[r].height
            if actual is None or actual < alto - 0.01:
                ws.row_dimensions[r].height = round(alto, 1)
                n += 1
    # RT-19 / RC-04 / RC-23 — §1.8 se aplicaba SÓLO a las combinadas y a las
    # filas de datos les ponía un tope de 34 pt. Las seis notas legales que
    # justifican la v2.0 tienen 123-221 caracteres en una columna de 28: con
    # 34 pt se leen DOS líneas de las ocho que necesitan, y el cliente tiene
    # que ensanchar la fila a mano en un libro protegido. Ahora la norma
    # (ceil(len/ancho) × 15 pt) vale para toda celda con `wrapText`, esté
    # combinada o no, sin tope.
    for r in range(1, ws.max_row + 1):
        alto = 0.0
        for c in ws[r]:
            if not isinstance(c.value, str) or not c.alignment.wrap_text:
                continue
            if len(c.value) < 40:
                continue
            ancho = _ancho_de(ws, c.column)
            lineas = max(1, int(math.ceil(len(c.value)
                                          / max(8.0, ancho * 1.05))))
            alto = max(alto, ALTO_LINEA * lineas)
        if not alto:
            continue
        actual = ws.row_dimensions[r].height
        if actual is None or actual < alto - 0.01:
            ws.row_dimensions[r].height = round(alto, 1)
            n += 1
    if n and informe is not None:
        informe.append(ws.title + ': ' + str(n) + ' altos ajustados (§1.8)')
    return n


# ==========================================================================
# §1.9 — Instrucciones, versión, bio, metadata y cross-sell
# ==========================================================================
def linea_texto(ws, texto, rx=None, col=None):
    """Sustituye la línea que case con `rx` o la añade al final. No duplica."""
    col = col or _col_texto(ws)
    for r in range(1, ws.max_row + 1):
        v = ws.cell(row=r, column=col).value
        if isinstance(v, str):
            if v == texto:
                return r
            if rx and rx.match(v):
                ws.cell(row=r, column=col).value = texto
                return r
    destino = _ultima_fila(ws, col) + 2
    origen = None
    for r in range(ws.max_row, 0, -1):
        if isinstance(ws.cell(row=r, column=col).value, str):
            origen = r
            break
    cel = ws.cell(row=destino, column=col, value=texto)
    if origen:
        cel._style = copy.copy(ws.cell(row=origen, column=col)._style)
    return destino


TITULOS_INSTRUCCIONES = 'INSTRUCCIONES DE USO'


#: RC-21 — el texto del contador estaba escrito para el molde C2 (seis
#: pestañas F1..F6 con un contador por hoja). En C1 y C4, que son
#: MONOLÍTICOS, prometía «el contador de cada bloque», que no existe: el
#: cliente que buscase el recuento de la FASE 3 no lo encontraba.
LINEA_CONTADOR = {
    'mono': 'El contador del pie cuenta sólo las ✓ de toda la lista; las N/A '
            'salen del total.',
    'bloques': 'El contador de cada pestaña cuenta sólo las ✓; las N/A salen '
               'del total.',
}
RX_LINEA_CONTADOR = re.compile(r'^El contador d')


def _linea_contador(det):
    return LINEA_CONTADOR['mono' if det.get('molde') in ('C1', 'C4')
                          else 'bloques']


def asegurar_instrucciones(wb, det, pid, informe):
    """Crea la hoja `Instrucciones` en los 12 xlsx que no la tienen (§1.9)."""
    ws = hoja(wb, 'Instrucciones')
    if ws is not None:
        if det['tipo'] == 'checklist':
            linea_texto(ws, _linea_contador(det), RX_LINEA_CONTADOR)
        return ws, False
    ws = wb.create_sheet('Instrucciones')
    val(ws, 'A1', TITULOS_INSTRUCCIONES, bold=True)
    fila = 3
    lineas = [
        'Este fichero forma parte del producto «' + pid + '» de AI Chef Pro.',
        NOTA_VERDES,
    ]
    if det['tipo'] == 'checklist':
        lineas.append(LEYENDA_OK)
        lineas.append(_linea_contador(det))
    if det['tipo'] == 'calculadora':
        lineas.append('Los campos con desplegable sólo admiten los valores de '
                      'la lista: una fórmula los compara por igualdad y un '
                      'valor tecleado a mano cae en la rama equivocada sin '
                      'avisar.')
    if det['tipo'] == 'proveedores':
        lineas.append('El índice enlaza con cada pestaña; las webs de la '
                      'columna correspondiente son enlaces directos.')
    if det['tipo'] == 'plan_financiero':
        lineas.append('Todas las cifras van SIN IVA salvo donde se diga lo '
                      'contrario.')
    for texto in lineas:
        val(ws, 'A' + str(fila), texto, wrap=True)
        fila += 1
    anchos(ws, {'A': 110})
    if informe is not None:
        informe.append('Instrucciones: hoja creada (§1.9)')
    return ws, True


def cierre_instrucciones(ws, pid):
    """Bloque de cierre: desproteger + cross-sell + BIO + VERSIÓN.

    Se BORRA entero y se reescribe al final del texto vivo: así su posición
    depende del contenido y no de dónde quedó en la pasada anterior, aunque un
    grupo haya añadido líneas por el medio.
    """
    col = _col_texto(ws)
    patrones = (RX_DESPROTEGER, RX_BIO, RX_VERSION, RX_CROSS_SELL)
    estilo = None
    for r in range(1, ws.max_row + 1):
        cel = ws.cell(row=r, column=col)
        v = cel.value
        if isinstance(v, str) and any(p.search(v) for p in patrones):
            if estilo is None:
                estilo = copy.copy(cel._style)
            cel.value = None
    ultima = _ultima_fila(ws, col)
    lineas = [NOTA_DESPROTEGER, CROSS_SELL, BIO_LINE,
              VERSION_FMT.format(pid=pid)]
    fila = ultima + 2
    for texto in lineas:
        cel = ws.cell(row=fila, column=col, value=texto)
        if estilo is not None:
            cel._style = copy.copy(estilo)
        cel.alignment = Alignment(horizontal='left', vertical='top')
        fila += 1
    return fila - 1


#: §1.9 — el cierre del docx del representante lista cuatro productos con
#: importe escrito a mano dentro de un fichero ya descargado. En xlsx el patrón
#: es el mismo: nombre de producto + precio. Se quita el importe, no el nombre
#: (la marca ChefBusiness se mantiene: es del grupo).
RX_PRECIO = re.compile(r'\s*[—–-]?\s*(?:por\s+)?\d{1,3}(?:[.,]\d{1,2})?\s*€'
                       r'(?:\s*\(.*?\))?')
RX_LINEA_CATALOGO = re.compile(
    r'(aichef\.pro/productos|kit\s|plan\s|gu[ií]a\s|pack\s|ebook)', re.I)


def cross_sell_sin_precios(ws, informe, detalle):
    n = 0
    for row in ws.iter_rows():
        for c in row:
            v = c.value
            if not isinstance(v, str) or '€' not in v:
                continue
            if not RX_LINEA_CATALOGO.search(v):
                continue
            if len(v) < 25:
                continue
            nuevo = RX_PRECIO.sub('', v).strip()
            if nuevo != v:
                c.value = nuevo
                detalle.append((ws.title, c.coordinate, v[:70], nuevo[:70]))
                n += 1
    if n and informe is not None:
        informe.append(ws.title + ': ' + str(n)
                       + ' precios fuera del cross-sell (§1.9)')
    return n


def metadatos(wb, pid, titulo=None):
    p = wb.properties
    p.creator = CREATOR
    p.lastModifiedBy = CREATOR
    nombre = titulo or (p.title if isinstance(p.title, str) and p.title
                        else pid)
    nombre = re.sub(r'\s*·\s*v\d+\.\d+\s*$', '', nombre).strip()
    p.title = nombre + ' · v' + VERSION
    p.subject = re.sub(r'\s*·\s*v\d+\.\d+\s*$', '',
                       p.subject or nombre).strip() + ' · v' + VERSION
    p.modified = datetime.datetime(2026, 8, 29)
    return p.title


# ==========================================================================
# §1.10 — Hipervínculos y coherencia de nombres
# ==========================================================================
RX_URL = re.compile(r'^\s*(https?://[^\s]+|www\.[^\s]+)\s*$', re.I)
RX_NUM_HOJA = re.compile(r'^\s*(\d{1,2})\s*[.)-]')


def hipervinculos(wb, det, informe):
    """§1.10 — índice → pestañas, URLs de texto plano → enlace, rótulos = tabs.

    `ws._hyperlinks` está vacío en las 11 hojas de la plantilla del
    representante B: el índice no enlaza a nada y las URLs de `10. Hosply.pro`
    son texto.
    """
    n = 0
    indice = hoja(wb, 'Indice') or hoja(wb, 'Índice')
    if indice is not None:
        titulos = [ws.title for ws in wb.worksheets if ws is not indice]
        por_numero = {}
        for t in titulos:
            m = RX_NUM_HOJA.match(t)
            if m:
                por_numero[int(m.group(1))] = t
        for row in indice.iter_rows():
            for c in row:
                v = c.value
                if not isinstance(v, str):
                    continue
                m = RX_NUM_HOJA.match(v)
                if not m:
                    continue
                destino = por_numero.get(int(m.group(1)))
                if not destino:
                    continue
                if v.strip() != destino:
                    c.value = destino          # rótulo = nombre exacto del tab
                c.hyperlink = "#'" + destino + "'!A1"
                c.font = Font(color='0563C1', underline='single',
                              size=c.font.size, bold=c.font.bold)
                n += 1
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                v = c.value
                if not isinstance(v, str) or not RX_URL.match(v):
                    continue
                url = v.strip()
                if url.lower().startswith('www.'):
                    url = 'https://' + url
                if c.hyperlink is None:
                    c.hyperlink = url
                    c.font = Font(color='0563C1', underline='single',
                                  size=c.font.size, bold=c.font.bold)
                    n += 1
    if n and informe is not None:
        informe.append('hipervínculos: ' + str(n) + ' enlaces (§1.10)')
    return n


# ==========================================================================
# Checklists — columna OK, desplegable y contador (§1.5 + §7-bis.21)
# ==========================================================================
MARCA_OK, MARCA_NO, MARCA_NA = '✓', '—', 'N/A'
VOCAB_MARCA = (MARCA_OK, MARCA_NO, MARCA_NA)
SEM_CHECK = ((MARCA_OK, CF_VERDE_BG, CF_VERDE_FG),)


def _dv_marca_existente(ws, purgar=False):
    """Vocabulario del desplegable de marca que ya trae la hoja.

    Con `purgar=True` lo ELIMINA además: el original cuelga de una lista de 51
    celdas sueltas y el nuevo, de un rango; como los `sqref` no coinciden, sin
    esto el fichero acaba con DOS desplegables sobre las mismas celdas y el
    gate de ítems cuenta el doble.
    """
    encontrado = None
    quedan = []
    for dv in ws.data_validations.dataValidation:
        if dv.type == 'list' and isinstance(dv.formula1, str) \
                and MARCA_OK in dv.formula1:
            if encontrado is None:
                encontrado = [o for o in dv.formula1.strip('"').split(',')
                              if o]
            if purgar:
                continue
        quedan.append(dv)
    if purgar:
        ws.data_validations.dataValidation = quedan
    return encontrado


def _col_ok(ws, cab_fila, cab_textos):
    for c in ws[cab_fila]:
        if norm(c.value) == 'ok':
            return c.column
    return None


def checklist_ok_y_contador(wb, det, informe):
    """La columna OK con desplegable y contador, en los CUATRO moldes.

    En C4 la columna existe pero está vacía y sin contador; en C3 la
    instrucción dice «Marca [X] cuando completes el item» y nombra una columna
    que no existe. Es lo que el propio fichero ya promete (§7-bis.21).
    """
    if det['tipo'] != 'checklist':
        return 0
    total_items = 0
    for ws in wb.worksheets:
        if norm(ws.title) in HOJAS_MOTOR:
            continue
        cab, textos = _fila_cabecera(ws)
        if cab is None:
            continue
        col_ok = _col_ok(ws, cab, textos)
        col_tarea = None
        for c in ws[cab]:
            if norm(c.value) in ('tarea', 'tramite / accion', 'hito',
                                 'tramite / tarea', 'accion', 'concepto'):
                col_tarea = c.column
                break
        if col_tarea is None:
            for c in ws[cab]:
                if isinstance(c.value, str) and norm(c.value) not in ('ok',):
                    col_tarea = c.column
                    break
        if col_tarea is None:
            continue
        if col_ok is None:
            col_ok = ws.max_column + 1
            val(ws, get_column_letter(col_ok) + str(cab), 'OK', bold=True)
            anchos(ws, {get_column_letter(col_ok): 8})
        letra_ok = get_column_letter(col_ok)
        letra_t = get_column_letter(col_tarea)
        # filas de ítem = las que tienen texto en la columna de tarea
        filas = [r for r in range(cab + 1, ws.max_row + 1)
                 if isinstance(ws.cell(row=r, column=col_tarea).value, str)
                 and ws.cell(row=r, column=col_tarea).value.strip()
                 and not RX_TOTAL.match(
                     str(ws.cell(row=r, column=col_tarea).value))]
        # se descartan los pies («Tareas completadas:»)
        filas = [r for r in filas
                 if 'completad' not in norm(
                     ws.cell(row=r, column=col_tarea).value)]
        if not filas:
            continue
        r0, r1 = min(filas), max(filas)
        rango = letra_ok + str(r0) + ':' + letra_ok + str(r1)
        opciones = _dv_marca_existente(ws, purgar=True) or list(VOCAB_MARCA)
        celdas_item = [letra_ok + str(r) for r in filas]
        dv_lista(ws, rango, opciones, titulo='Marca no válida',
                 mensaje='Usa el desplegable: ' + ', '.join(opciones)
                 + '. N/A = no aplica, sale del total.',
                 celdas=celdas_item)
        for coord in celdas_item:
            verde(ws, coord)
        semaforo_texto(ws, rango, SEM_CHECK)
        total_items += len(filas)
        # Contador con COUNTIF (pycel NO implementa COUNTA). Si la hoja YA
        # trae uno —los cinco checklists del molde C2 lo llevan en su pie—, se
        # REUTILIZA esa fila: crear otro dejaría dos contadores y, peor, la
        # 2.ª pasada añadiría uno más en cada ejecución.
        formula_ok = '=COUNTIF(' + rango + ',"' + MARCA_OK + '")'
        formula_total = ('=COUNTIF(' + letra_t + str(r0) + ':' + letra_t
                         + str(r1) + ',"?*")-COUNTIF(' + rango + ',"'
                         + MARCA_NA + '")')
        fila_c, celdas_countif = None, []
        for r in range(cab + 1, ws.max_row + 1):
            fila_txt = ' '.join(norm(c.value) for c in ws[r]
                                if isinstance(c.value, str))
            if 'completad' in fila_txt:
                fila_c = r
                celdas_countif = [c.coordinate for c in ws[r]
                                  if isinstance(c.value, str)
                                  and 'COUNTIF' in c.value.upper()]
                break
        if fila_c is None:
            fila_c = fila_libre(ws, r1 + 2, (col_tarea, col_ok, col_ok + 1,
                                             col_ok + 2))
            val(ws, letra_t + str(fila_c), 'Ítems completados:', bold=True)
            celdas_countif = []
        if len(celdas_countif) >= 2:
            f(ws, celdas_countif[0], formula_ok, FMT_ENT)
            f(ws, celdas_countif[1], formula_total, FMT_ENT)
        else:
            f(ws, letra_ok + str(fila_c), formula_ok, FMT_ENT)
            val(ws, get_column_letter(col_ok + 1) + str(fila_c), 'de')
            f(ws, get_column_letter(col_ok + 2) + str(fila_c), formula_total,
              FMT_ENT)
        # la instrucción que nombra una columna inexistente (molde C3)
        for r in range(1, cab):
            v = ws.cell(row=r, column=1).value
            if isinstance(v, str) and 'marca [x]' in norm(v):
                ws.cell(row=r, column=1).value = (
                    'Marca la columna OK con el desplegable (✓ hecho, — '
                    'pendiente, N/A no aplica). El contador del pie suma sólo '
                    'las ✓.')
    if informe is not None and total_items:
        informe.append('checklist ' + det['molde'] + ': ' + str(total_items)
                       + ' ítems con desplegable y contador')
    return total_items


# ==========================================================================
# Semáforos genéricos (§1.6) — toda fila de resultado que pueda ser negativa
# ==========================================================================
RX_RESULTADO = re.compile(
    r'^(resultado|ebitda|beneficio|margen de seguridad|cash ?flow|'
    r'flujo de caja|tesorer[ií]a)', re.I)


def semaforos_resultado(ws, informe):
    n = 0
    cab, _ = _fila_cabecera(ws)
    if cab is None:
        return 0
    for r in range(cab + 1, ws.max_row + 1):
        rot = _rotulo_de_fila(ws, r)
        if not rot or not RX_RESULTADO.match(rot):
            continue
        cols = [c.column for c in ws[r]
                if (_es_numero(c.value) or _es_formula(c.value))
                and c.column > 1]
        if not cols:
            continue
        letra0 = get_column_letter(min(cols))
        letra1 = get_column_letter(max(cols))
        rango = letra0 + str(r) + ':' + letra1 + str(r)
        ancla = letra0 + str(r)
        n += semaforo_num(ws, rango, verde_si=ancla + '>0',
                          rojo_si=ancla + '<0')
    if n and informe is not None:
        informe.append(ws.title + ': ' + str(n)
                       + ' reglas de semáforo con ISNUMBER (§1.6)')
    return n


# ==========================================================================
# Impresión A4 y protección
# ==========================================================================
def print_setup(ws, header_row=None, landscape=None, congelar=None):
    ws.page_setup.paperSize = 9                      # A4
    if landscape is not None:
        ws.page_setup.orientation = 'landscape' if landscape else 'portrait'
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    if ws.sheet_properties.pageSetUpPr is None:
        from openpyxl.worksheet.properties import PageSetupProperties
        ws.sheet_properties.pageSetUpPr = PageSetupProperties()
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins = PageMargins(left=0.59, right=0.59, top=0.59, bottom=0.59,
                                  header=0.3, footer=0.3)
    ws.oddFooter.center.text = PIE
    ws.oddFooter.center.size = 8
    if header_row:
        ws.print_title_rows = str(header_row) + ':' + str(header_row)
    if congelar:
        ws.freeze_panes = congelar


def _hay_numeros(ws):
    for row in ws.iter_rows():
        for c in row:
            if _es_numero(c.value):
                return True
    return False


def protegible(ws):
    """¿Se puede proteger esta hoja SIN dejarla inservible?

    La protección es del §1, pero marcar los inputs en verde es de §2/§3. Si
    se protege una hoja de datos a la que ningún grupo le ha declarado todavía
    una sola celda editable, el cliente abre la calculadora y **no puede
    teclear nada**: peor que no proteger. Se protege cuando hay inputs, o
    cuando la hoja es de texto puro (Instrucciones, Índice), y se anota el
    resto como pendiente para que el gate no cante verde sobre un fichero a
    medio construir.
    """
    extras = getattr(ws, '_pl_editables', set())
    if extras:
        return True
    for row in ws.iter_rows():
        for c in row:
            if es_verde(c):
                return True
    return not _hay_numeros(ws)


def proteger(ws, informe=None):
    """Protección SIN contraseña: se desbloquean sólo las verdes (y las que un
    grupo haya declarado editables por rol).

    Ojo con `password`: `= None` revienta openpyxl y `= ''` escribe el hash de
    la cadena vacía → Excel pediría contraseña justo donde las instrucciones
    dicen que no hay ninguna. Se deja sin asignar.
    """
    extras = getattr(ws, '_pl_editables', set())
    verdes = 0
    for row in ws.iter_rows():
        for c in row:
            if es_verde(c) or c.coordinate in extras:
                c.protection = Protection(locked=False)
                verdes += 1
            else:
                c.protection = Protection(locked=True)
    ws.protection.sheet = True
    ws.protection.selectLockedCells = False
    ws.protection.selectUnlockedCells = False
    ws.protection.formatColumns = False
    ws.protection.formatRows = False
    ws.protection.sort = False
    if informe is not None:
        informe.append(ws.title + ': protegida sin contraseña (' + str(verdes)
                       + ' celdas editables)')
    return verdes


# ==========================================================================
# Gates de §9 (los mide el motor; el veredicto lo da main.py)
# ==========================================================================
RX_NO_LATINO = re.compile(
    '[一-鿿぀-ヿ가-힯Ѐ-ӿ֐-׿'
    '؀-ۿ฀-๿]')
RX_SIN_TILDE = re.compile(
    r'\b(Ano|anos|Anos|Espana|Diseno|resenas|desempeno|alergenos|danos|'
    r'Cataluna|Cumpleanos|Analisis|Constitucion|Tramite|Nominas|'
    r'Amortizacion|Comision|Gestoria|Bolleria)\b')
#: RD-32 / RC-20 — la lista cerrada del gate se quedaba corta: nueve palabras
#: sin tilde seguían vivas en el fichero que el cliente imprime para el banco
#: («Camara frigorifica», «Estanterias inox almacen», «Molinillo cafe»,
#: «Pagina básica», «notaria»). Van en su propio patrón, sin distinguir
#: mayúsculas, porque aparecen tanto en rótulo como en nota.
#: CRIT-05 añade `critic[oa]s?` y la familia de «diseñ-»: las dos salían en
#: verde con erratas vivas en el fichero que el cliente imprime. La lista de
#: «diseñ-» se enumera (no `disen\w*`) para no inventar palabras.
RX_SIN_TILDE_2 = re.compile(
    r'(?i)\b(camaras?|frigorific[oa]s?|estanterias?|almacen|cafes?|paginas?|'
    r'notaria|busquedas?|consejerias?|codigos?|clausulas?|critic[oa]s?|'
    r'disen(?:o|os|ador|adora|adores|adoras|ar|a|an|ado|ada|ados|adas))\b')
#: RC-20 — heurística de CONVIVENCIA: la palabra sin tilde que comparte forma
#: normalizada con otra ACENTUADA del mismo libro. Es la que caza los casos
#: que ninguna lista cerrada prevé («Pagina básica» tiene la tilde dos
#: palabras más allá). Se excluyen los homógrafos legítimos del español, que
#: son pares reales y no erratas.
#: ⚠️ CRIT-05 — `critica`/`criticas` SALIERON de esta lista. No son homógrafos
#: útiles aquí: como sustantivo y como adjetivo llevan tilde («la crítica»,
#: «semanas críticas»), y lo único sin tilde es el verbo criticar, que
#: `RX_CRITICA_VERBO` reconoce por su clítico. Mientras estuvieron aquí, el
#: gate daba VERDE sobre «Primeras 4 semanas criticas» y «Critico para SEO
#: local» en el fichero que el cliente imprime.
HOMOGRAFAS = frozenset((
    'publica', 'publico', 'publicas', 'publicos', 'practica', 'practicas',
    'termino', 'terminos', 'continuo', 'continua',
    'medico', 'valido', 'calculo', 'calculos', 'trabajo', 'numero',
    'deposito', 'depositos', 'limite', 'limites', 'titulo', 'titulos',
    'transito', 'estimulo', 'domicilio', 'ejercito', 'liquido', 'liquidos',
    'capitulo', 'capitulos', 'articulo', 'articulos', 'intimo', 'animo',
))


def gate_no_latinos(wb, fname):
    fuera = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if isinstance(c.value, str) and RX_NO_LATINO.search(c.value):
                    m = RX_NO_LATINO.search(c.value)
                    ini = max(0, m.start() - 30)
                    fuera.append({'fichero': fname, 'hoja': ws.title,
                                  'celda': c.coordinate,
                                  'fragmento': c.value[ini:m.end() + 30]})
    return fuera


def _acentuadas_del_libro(wb):
    """Formas ACENTUADAS presentes en el libro, indexadas por forma sin tilde.

    Sirve a la heurística de convivencia de `gate_ortografia` (RC-20): si el
    libro escribe «Página» en una celda y «Pagina» en otra, la segunda es una
    errata aunque no esté en ninguna lista cerrada.
    """
    mapa = {}
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                v = c.value
                if not isinstance(v, str):
                    continue
                for palabra in RX_PALABRA.findall(v):
                    if RX_YA_ACENTUADA.search(palabra) or 'ñ' in palabra.lower():
                        mapa.setdefault(norm(palabra), palabra)
    return mapa


def gate_ortografia(wb, fname):
    fuera = []
    acentuadas = _acentuadas_del_libro(wb)
    for ws in wb.worksheets:
        if RX_SIN_TILDE.search(ws.title or '') \
                or RX_SIN_TILDE_2.search(ws.title or ''):
            fuera.append({'fichero': fname, 'hoja': ws.title, 'celda': '(hoja)',
                          'texto': ws.title})
        for row in ws.iter_rows():
            for c in row:
                v = c.value
                if not isinstance(v, str) or v.startswith('='):
                    continue
                limpio = RX_URL_MAIL.sub(' ', v)
                # CRIT-05 — tramos donde «critica/criticas» es el verbo
                # criticar y NO lleva tilde. Se calculan una vez por celda y
                # los usan las dos vías del gate (lista cerrada y convivencia).
                verbal = [(mm.start(), mm.end())
                          for mm in RX_CRITICA_VERBO.finditer(limpio)]

                def _excusada(clave, ini, _verbal=verbal):
                    # «campana extractora» convive a propósito con «Campaña
                    # lanzamiento RRSS»: es la excepción que documenta §1.7
                    if clave.startswith('campana'):
                        return True
                    return (clave in ('critica', 'criticas')
                            and any(a <= ini < b for a, b in _verbal))

                # ⚠️ antes se miraba SÓLO la primera coincidencia y, si era la
                # campana extractora, se saltaba la celda entera: una errata
                # posterior en el mismo texto se perdía sin avisar.
                m = None
                for rx in (RX_SIN_TILDE, RX_SIN_TILDE_2):
                    for mm in rx.finditer(limpio):
                        if _excusada(norm(mm.group(0)), mm.start()):
                            continue
                        m = mm
                        break
                    if m is not None:
                        break
                if m is not None:
                    fuera.append({'fichero': fname, 'hoja': ws.title,
                                  'celda': c.coordinate,
                                  'palabra': m.group(0), 'texto': v[:80]})
                    continue
                # heurística de convivencia (RC-20)
                for mp in RX_PALABRA.finditer(limpio):
                    palabra = mp.group(0)
                    clave = norm(palabra)
                    if len(clave) < 5 or clave in HOMOGRAFAS:
                        continue
                    if _excusada(clave, mp.start()):
                        continue
                    # «campana extractora» convive a propósito con «Campaña
                    # lanzamiento RRSS» en el mismo libro: es la excepción
                    # que documenta §1.7 y que un barrido genérico rompe.
                    if clave.startswith('campana'):
                        continue
                    if RX_YA_ACENTUADA.search(palabra) or 'ñ' in palabra.lower():
                        continue
                    gemela = acentuadas.get(clave)
                    if gemela and norm(gemela) == clave:
                        fuera.append({'fichero': fname, 'hoja': ws.title,
                                      'celda': c.coordinate,
                                      'palabra': palabra,
                                      'convive_con': gemela,
                                      'texto': v[:80]})
                        break
    return fuera


def gate_formatos(wb, fname):
    """0 celdas de euros con formato de porcentaje y 0 recuentos con euro."""
    fuera = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                v, fmt = c.value, (c.number_format or '')
                # RT-11 — el `continue` para las celdas de TEXTO dejaba fuera
                # del gate las 51 notas con formato de euro que el propio
                # motor escribía. Ahora se miran ANTES de saltarlas.
                if _devuelve_texto(v) and ('\u20ac' in fmt or '%' in fmt):
                    fuera.append({'fichero': fname, 'hoja': ws.title,
                                  'celda': c.coordinate,
                                  'tipo': 'texto con formato numérico',
                                  'rotulo': (_rotulo_de_fila(ws, c.row)
                                             or '')[:50],
                                  'formato': fmt})
                    continue
                if not (_es_numero(v) or _es_formula(v)) or c.column == 1:
                    continue
                rot = _rotulo_de_fila(ws, c.row)
                if '/h' in fmt and '"' not in fmt:
                    fuera.append({'fichero': fname, 'hoja': ws.title,
                                  'celda': c.coordinate,
                                  'tipo': '€/h sin entrecomillar',
                                  'rotulo': rot[:50], 'formato': fmt})
                if not rot:
                    continue
                cab_col = cabecera_de_columna(ws, c.column)
                col_ratio = bool(re.search(r'%|ratio', cab_col, re.I))
                col_euro = bool(re.search(r'€|\beur\b|importe|coste|precio',
                                          cab_col, re.I))
                if not col_ratio and RX_IMPORTE.search(rot) \
                        and not RX_RATIO.search(rot) and '%' in fmt:
                    if _es_formula(v) or abs(float(v)) > 5:
                        fuera.append({'fichero': fname, 'hoja': ws.title,
                                      'celda': c.coordinate, 'tipo':
                                      'importe con %', 'rotulo': rot[:50],
                                      'formato': fmt})
                # un rótulo que declara euros («Coste anual … (€)») manda
                # sobre la cabecera de su columna: si no, toda celda de
                # importe que caiga bajo una columna llamada «Personas» se
                # marcaba como recuento con € (RD-10, bloque nuevo de
                # crecimiento de plantilla)
                if not col_euro and not col_ratio and '€' in fmt \
                        and not RX_EURO_FUERTE.search(rot) \
                        and (RX_RECUENTO.search(rot)
                             or RX_RECUENTO.search(cab_col)):
                    fuera.append({'fichero': fname, 'hoja': ws.title,
                                  'celda': c.coordinate,
                                  'tipo': 'recuento con €',
                                  'rotulo': rot[:50], 'formato': fmt})
    return fuera


#: RC-29 — Excel corta los nombres de pestaña en 31 caracteres y openpyxl
#: sólo avisa por `UserWarning`, que se pierde entre la salida del run. Dos
#: hojas del checklist de parrillero se pasan (37 y 40 caracteres): al
#: reescribir el fichero, Excel puede pedir reparación.
LIMITE_TITULO = 31


def gate_nombres_hoja(wb, fname):
    fuera = []
    for ws in wb.worksheets:
        if len(ws.title or '') > LIMITE_TITULO:
            fuera.append({'fichero': fname, 'hoja': ws.title,
                          'celda': '(hoja)', 'longitud': len(ws.title),
                          'limite': LIMITE_TITULO})
    return fuera


#: RT-01 — un formato condicional de tipo `expression` cuya fórmula NO empieza
#: en la primera fila del `sqref` pinta la tabla DESPLAZADA. Es la regresión
#: que el motor introdujo al reanclar el resaltado del checklist: el fichero
#: de producción tenía `sqref=A4:F59` con `$E4`, y el generado `sqref=A3:F73`
#: con `$E4`, así que la cabecera se pintaba según un ítem.
RX_FILA_FORMULA = re.compile(r'\$?[A-Z]{1,3}\$?(\d+)')


def gate_cf_anclado(wb, fname):
    fuera = []
    for ws in wb.worksheets:
        for cf in ws.conditional_formatting:
            filas = [CellRange(str(r)).min_row for r in cf.sqref.ranges]
            if not filas:
                continue
            primera = min(filas)
            for regla in cf.rules:
                if regla.type != 'expression' or not regla.formula:
                    continue
                m = RX_FILA_FORMULA.search(str(regla.formula[0]))
                if not m:
                    continue
                if int(m.group(1)) != primera:
                    fuera.append({'fichero': fname, 'hoja': ws.title,
                                  'celda': str(cf.sqref),
                                  'formula': str(regla.formula[0])[:60],
                                  'fila_sqref': primera,
                                  'fila_formula': int(m.group(1))})
    return fuera


def gate_referencias(wb, fname):
    """Re-verificación del §1.7: ninguna fórmula apunta a una hoja que no está."""
    presentes = set(norm(ws.title) for ws in wb.worksheets)
    fuera = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if not _es_formula(c.value):
                    continue
                for m in RX_REF_HOJA.finditer(c.value):
                    nombre = m.group(1) or m.group(2)
                    if not nombre or nombre.upper() in ('IF', 'SUM', 'TRUE',
                                                        'FALSE'):
                        continue
                    if norm(nombre) not in presentes:
                        fuera.append({'fichero': fname, 'hoja': ws.title,
                                      'celda': c.coordinate,
                                      'hoja_referida': nombre,
                                      'formula': c.value[:80]})
    return fuera


def contadores(wb, fname, det):
    r = {'fichero': fname, 'tipo': det['tipo'], 'molde': det['molde'],
         'hojas': len(wb.worksheets), 'formulas': 0, 'verdes': 0, 'cf': 0,
         'dv': 0, 'protegidas': 0, 'divisiones_sin_iferror': 0,
         'ceros_constantes': 0, 'vacias': 0, 'a4': 0, 'items_checklist': 0}
    for ws in wb.worksheets:
        r['cf'] += sum(len(cf.rules) for cf in ws.conditional_formatting)
        r['dv'] += len(ws.data_validations.dataValidation)
        if ws.protection.sheet:
            r['protegidas'] += 1
        if ws.page_setup.paperSize == 9:
            r['a4'] += 1
        for row in ws.iter_rows():
            for c in row:
                if es_verde(c):
                    r['verdes'] += 1
                v = c.value
                if _es_formula(v):
                    r['formulas'] += 1
                    cuerpo = re.sub(r'"[^"]*"', '', v)
                    if '/' in cuerpo and not RX_YA_GUARDADA.match(v):
                        r['divisiones_sin_iferror'] += 1
                elif v == 0 and not es_verde(c):
                    r['ceros_constantes'] += 1
                elif v == '':
                    r['vacias'] += 1
    return r


def censar_docx(path):
    """Lectura de SOLO LECTURA para el gate de §9 (A4 + author).

    Los docx los produce `documentos.py` (§4); aquí sólo se miden, para que el
    informe diga qué queda pendiente y no se cante verde antes de tiempo.
    """
    try:
        from docx import Document
    except ImportError:
        return {'fichero': os.path.basename(path), 'error': 'sin python-docx'}
    doc = Document(path)
    sec = doc.sections[0]
    ancho = sec.page_width.inches if sec.page_width else None
    alto = sec.page_height.inches if sec.page_height else None
    a4 = (ancho is not None and abs(ancho - 8.27) < 0.05
          and alto is not None and abs(alto - 11.69) < 0.05)
    return {'fichero': os.path.basename(path), 'a4': a4,
            'ancho_in': round(ancho, 2) if ancho else None,
            'alto_in': round(alto, 2) if alto else None,
            'author': doc.core_properties.author,
            'title': doc.core_properties.title,
            'headings': sum(1 for p in doc.paragraphs
                            if (p.style.name or '').startswith('Heading')),
            'tablas': len(doc.tables),
            'palabras': sum(len(p.text.split()) for p in doc.paragraphs)}


# ==========================================================================
# Aritmética financiera (pycel NO implementa IRR ni PMT)
# ==========================================================================
def van(tasa, flujos):
    total = 0.0
    for i, fl in enumerate(flujos):
        total += float(fl) / ((1.0 + tasa) ** i)
    return total


def _dvan(tasa, flujos):
    total = 0.0
    for i, fl in enumerate(flujos):
        if i == 0:
            continue
        total += -i * float(fl) / ((1.0 + tasa) ** (i + 1))
    return total


def _biseccion(flujos, lo=-0.99, hi=10.0, iteraciones=200, tol=1e-10):
    f_lo, f_hi = van(lo, flujos), van(hi, flujos)
    if f_lo * f_hi > 0:
        return None
    for _ in range(iteraciones):
        medio = (lo + hi) / 2.0
        f_m = van(medio, flujos)
        if abs(f_m) < tol:
            return medio
        if f_lo * f_m < 0:
            hi, f_hi = medio, f_m
        else:
            lo, f_lo = medio, f_m
    return (lo + hi) / 2.0


def tir_newton(flujos, semilla=0.1, iteraciones=100, tol=1e-10):
    """TIR por Newton-Raphson con caída a bisección. `None` si no existe."""
    if not flujos or all(float(x) >= 0 for x in flujos) \
            or all(float(x) <= 0 for x in flujos):
        return None
    tasa = semilla
    for _ in range(iteraciones):
        v = van(tasa, flujos)
        if abs(v) < tol:
            return tasa
        d = _dvan(tasa, flujos)
        if abs(d) < 1e-12:
            break
        nueva = tasa - v / d
        if nueva <= -0.999999:
            break
        tasa = nueva
    return _biseccion(flujos)


def cuota_anualidad(importe, tipo_anual, anios, periodos_ano=12):
    """Cuota francesa SIN `PMT` (pycel no lo implementa)."""
    i = float(tipo_anual) / periodos_ano
    n = int(anios) * periodos_ano
    if i == 0:
        return float(importe) / n
    return float(importe) * i / (1.0 - (1.0 + i) ** (-n))


def payback(flujos):
    acumulado = float(flujos[0])
    for i in range(1, len(flujos)):
        previo = acumulado
        acumulado += float(flujos[i])
        if acumulado >= 0 and float(flujos[i]) != 0:
            return round(i - 1 + abs(previo) / float(flujos[i]), 2)
    return None


# ==========================================================================
# Inyección del valor cacheado de una fórmula (para TIR y compañía)
# ==========================================================================
def _mapa_hojas(z):
    wbxml = z.read('xl/workbook.xml').decode('utf-8')
    rels = z.read('xl/_rels/workbook.xml.rels').decode('utf-8')
    relmap = {}
    for rel in re.findall(r'<Relationship[^>]*/>', rels):
        rid = re.search(r'Id="([^"]+)"', rel)
        tgt = re.search(r'Target="([^"]+)"', rel)
        if rid and tgt:
            relmap[rid.group(1)] = tgt.group(1)
    fuera = {}
    for m in re.finditer(r'<sheet[^>]*name="([^"]+)"[^>]*r:id="([^"]+)"',
                         wbxml):
        tgt = relmap.get(m.group(2), '')
        if tgt:
            sfile = tgt.lstrip('/') if tgt.startswith('/') else 'xl/' + tgt
            fuera[html.unescape(m.group(1))] = sfile
    return fuera


def inyectar_valor(path, nombre_hoja, coord, valor):
    """Escribe el `<v>` de una celda de fórmula reescribiendo el zip.

    Va SIEMPRE al final del pipeline: cualquier `wb.save()` posterior borraría
    el caché.
    """
    z = zipfile.ZipFile(path)
    hojas = _mapa_hojas(z)
    sfile = hojas.get(nombre_hoja)
    parts = dict((n, z.read(n)) for n in z.namelist())
    z.close()
    if sfile is None or sfile not in parts:
        return False
    xml = parts[sfile].decode('utf-8')
    if isinstance(valor, (int, float)) and not isinstance(valor, bool):
        pat = re.compile(r'(<c r="' + coord + r'"[^>]*>)(<f[^>]*>[^<]*</f>)'
                         r'(?:<v>[^<]*</v>|<v\s*/>)?(</c>)')
        rep = r'\1\2<v>' + repr(float(valor)) + r'</v>\3'
        xml, n = pat.subn(rep, xml, count=1)
    else:
        esc = html.escape(str(valor))
        pat = re.compile(r'<c r="' + coord + r'"([^>]*)>(<f[^>]*>[^<]*</f>)'
                         r'(?:<v>[^<]*</v>|<v\s*/>)?</c>')

        def _rep(m):
            attrs = re.sub(r'\s+t="[^"]*"', '', m.group(1))
            return ('<c r="' + coord + '"' + attrs + ' t="str">' + m.group(2)
                    + '<v>' + esc + '</v></c>')

        xml, n = pat.subn(_rep, xml, count=1)
    if not n:
        return False
    parts[sfile] = xml.encode('utf-8')
    tmp = path + '.tmp'
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for nombre, data in parts.items():
            zout.writestr(nombre, data)
    shutil.move(tmp, path)
    return True


# ==========================================================================
# API principal
# ==========================================================================
def aplicar(wb, fname, det, pid, informe, detalle=None):
    """§1 transversal ANTES del trabajo de los grupos.

    Orden deliberado: ortografía y formatos primero (los grupos buscan
    etiquetas y escriben fórmulas sobre celdas ya saneadas), cableado de
    totales después, y la superficie de parámetros al final para que exista
    cuando el grupo la pida. El renombrado de hojas NO va aquí: §1.7 lo manda
    después del cableado, y está en `cerrar()`.
    """
    detalle = detalle if detalle is not None else {}
    detalle.setdefault('formatos', [])
    detalle.setdefault('ortografia', [])
    detalle.setdefault('desajustes_total', [])
    detalle.setdefault('cross_sell', [])
    detalle.setdefault('pseudo_formulas', [])
    metadatos(wb, pid)
    for ws in wb.worksheets:
        neutralizar_pseudoformulas(ws, informe, detalle['pseudo_formulas'])
        ortografia(ws, informe, detalle['ortografia'])
        formatos_por_tipo(ws, informe, detalle['formatos'])
    if det['tipo'] in ('plan_financiero', 'calculadora'):
        for ws in wb.worksheets:
            cablear_sumas(ws, informe, detalle['desajustes_total'])
    for ws in wb.worksheets:
        guardas(ws, informe)
    # La hoja de Instrucciones se crea AQUÍ, antes que la superficie de
    # parámetros: en B-γ el bloque `PARÁMETROS` vive dentro de ella, y crearla
    # después dejaba la 1.ª pasada sin bloque y la 2.ª con él (idempotencia
    # rota por orden de construcción, no por contenido).
    asegurar_instrucciones(wb, det, pid, informe)
    params = parametros(wb, det)
    return params, detalle


def cerrar(wb, fname, det, pid, informe, detalle=None, proteger_hojas=True):
    """§1 transversal DESPUÉS de los grupos.

    Si se protegiera antes, cada celda que un grupo creara después nacería
    bloqueada aunque fuese verde.
    """
    detalle = detalle if detalle is not None else {}
    detalle.setdefault('cross_sell', [])
    checklist_ok_y_contador(wb, det, informe)
    hipervinculos(wb, det, informe)
    for ws in wb.worksheets:
        semaforos_resultado(ws, informe)
        cross_sell_sin_precios(ws, informe, detalle['cross_sell'])
        for row in ws.iter_rows():
            for c in row:
                if c.value == '':
                    # el censo cuenta `''` como defecto `empty_str`
                    c.value = None
    ws_ins, _ = asegurar_instrucciones(wb, det, pid, informe)
    for ws in wb.worksheets:
        formatos_texto(ws, informe)
        validaciones(ws, informe)
        altos_y_wrap(ws, informe)
        print_setup(ws)
    cierre_instrucciones(ws_ins, pid)
    sin_inputs = []
    if proteger_hojas:
        for ws in wb.worksheets:
            if protegible(ws):
                proteger(ws, informe)
            else:
                sin_inputs.append(ws.title)
    detalle['hojas_sin_inputs_sin_proteger'] = sin_inputs
    if sin_inputs and informe is not None:
        informe.append('sin proteger (ningún input declarado todavía; lo hará '
                       'el grupo de §2/§3): ' + ', '.join(sin_inputs))
    # §1.7: el renombrado va el ÚLTIMO, con reescritura de referencias, y
    # `gate_referencias` (en main.py) lo re-verifica.
    renombrar_hojas(wb, informe)
    metadatos(wb, pid)
    return informe
