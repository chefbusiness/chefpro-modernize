#!/usr/bin/env python3
"""
documentos.py — Pipeline de DOCUMENTOS de la familia «Guías Cómo Montar» v2.0.

Implementa el **§5** de `guias-v2-SPEC.md` (§5.1-§5.6) y las decisiones
§7-bis.6 (la promesa de páginas SE CUMPLE), §7-bis.7 (una sola fuente de
cifras: los xlsx del propio producto), §7-bis.15 (el PDF y el DOCX salen del
MISMO Markdown), §7-bis.18 (la cifra de páginas se MIDE) y §7-bis.21 (las
cifras del sector sólo entran con fuente).

    python3 documentos.py --producto <pid> [--salida <dir>] [--json informe.json]
                          [--solo-capitulos 1,2,3] [--solo-bonus business-plan-modelo]
                          [--saltar-generacion] [--min-palabras-cap 900]

Es el ÚNICO bloque de la familia que **no post-procesa: produce**. Hoy no hay
nada que parchear —2.157 palabras para 22 capítulos en el representante, 255
palabras de «manual de servicio de sala», y en 6 de las 8 guías el PDF es una
portada de una página—, así que el texto se escribe de nuevo.

FLUJO (§5.3)
  guion_<pid>.py  →  por capítulo, N bloques de epígrafes
      → prompt a bridge.py **con las cifras y las tablas del xlsx ya
        RESUELTAS dentro del prompt** (el modelo no lee ficheros: se le dan
        los números medidos, y se le prohíbe inventar otros)
      → un .txt por bloque  →  .md por capítulo
      → barrido de no latinos + erratas (fechas caducas, «60 % cierra»)
      → ensamblado (portada, índice, PageBreak por capítulo, tablas reales)
      → DOCX (python-docx) **y** PDF (reportlab) desde el MISMO Markdown
      → gates §5.6 medidos con PyMuPDF / python-docx.

REGLAS DURAS QUE ESTE FICHERO RESPETA
  1. **PROHIBIDO escribir en `astro-site/public/dl/**`.** Los xlsx de ese
     directorio se abren SIEMPRE en modo lectura (`load_workbook(data_only=True)`)
     y son la única fuente de cifras del producto (§7-bis.7). La salida va
     donde diga `--salida`, y si esa ruta cae dentro de `public/dl/` el script
     ABORTA con exit 2.
  2. **Texto largo SIEMPRE con `bridge.py`.** Este módulo redacta el guion, las
     instrucciones al modelo, las tablas de cifras y el ensamblado; **ni una
     frase del cuerpo**. En este Mac el routing por defecto está desactualizado
     (enruta `content` a `deepseek-v4-pro` y trae `--max-tokens 4096`), así que
     `--model ~deepseek/deepseek-v4-flash-latest` y `--max-tokens 8192` NO son
     opcionales. Medido el 2026-08-29: una petición con `--max-tokens 8192`
     puede agotar el presupuesto razonando y devolver un fichero VACÍO; por eso
     `bridge()` reintenta con presupuestos menores y aborta si nunca hay texto.
  3. **Cifras del sector sólo desde `auditorias/guias-v2-research-sector.json`**,
     con su fuente citada en el texto. Un `id` sin `cifra` o con
     `fiabilidad: baja` es un hueco DELIBERADO: el prompt lo dice y el capítulo
     se reformula sin número.
  4. **Barrido de no latinos sobre cada .md antes de ensamblar** (§5.3.3), más
     `gate-no-latinos.py --only <salida>` al final, fuera de este script.

IDEMPOTENCIA / CACHÉ: cada bloque se cachea en `<salida>/txt/<pid>_capNN_bM.txt`
y no se vuelve a pedir si ya existe y cumple el mínimo de palabras. Para forzar
la regeneración de un capítulo corto basta con borrar sus .txt (o pasar
`--regenerar 9,13,14`), que es exactamente lo que pide §5.6.1: **se amplían los
capítulos más delgados; no se toca la cifra de la landing**.
"""
import argparse
import importlib.util
import json
import os
import re
import collections
import datetime
import html
import unicodedata
import textwrap
import subprocess
import sys
import time
import unicodedata

# Caracteres tipográficos de la familia, SIEMPRE por escape (CLAUDE.md): al
# pasar por un heredoc del shell degeneran en espacio y guion normales.
NARROW = ' '          # espacio fino antes de las unidades
NOBRK = '‑'           # guion no separable en los rangos

AQUI = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.normpath(os.path.join(AQUI, '..', '..', '..'))
DL = os.path.join(REPO, 'astro-site', 'public', 'dl')
RESEARCH = os.path.join(AQUI, '..', 'auditorias', 'guias-v2-research-sector.json')
BRIDGE = '/Users/johnguerrero/chefbusiness-ai/bridge.py'
MODELO = '~deepseek/deepseek-v4-flash-latest'
# 2026-09-04: regla de la memoria («si bridge devuelve vacío dos veces, no dupliques
# el presupuesto: cambia de motor»). A partir del 3.er intento el bloque se pide a
# un modelo que NO razona; `--modelo` permite arrancar directamente con él.
MODELO_FALLBACK = 'anthropic/claude-sonnet-4.6'

# --------------------------------------------------------------------------
# 0. Guardas de seguridad
# --------------------------------------------------------------------------


def guarda_salida(salida):
    """REGLA DURA 1: nunca se escribe dentro de astro-site/public/dl/."""
    real = os.path.realpath(salida)
    if real.startswith(os.path.realpath(DL)):
        print(f'ABORTADO: {real} cae dentro de public/dl/ — los entregables de '
              'producción no se tocan desde aquí (SPEC §5, regla dura 1).',
              file=sys.stderr)
        raise SystemExit(2)
    os.makedirs(salida, exist_ok=True)
    return salida


def istats():
    """REGLA TÉRMICA: el Mac se apaga a 65 °C. Devuelve la temperatura o None."""
    try:
        out = subprocess.run(['istats', 'cpu', 'temp'], capture_output=True,
                             text=True, timeout=20).stdout
        m = re.search(r'([\d.]+)\s*°?C', out)
        return float(m.group(1)) if m else None
    except Exception:
        return None


def respirar(umbral=60.0, espera=90):
    t = istats()
    if t is not None and t >= umbral:
        print(f'  [térmica] {t} °C ≥ {umbral} — durmiendo {espera} s',
              flush=True)
        time.sleep(espera)
    return t


# --------------------------------------------------------------------------
# 1. Idioma y erratas — gates propios sobre el texto crudo (§5.3.3, §5.6.5/7)
# --------------------------------------------------------------------------
RX_NO_LATINOS = re.compile(
    r'[　-〿぀-ヿ㐀-䶿一-鿿'
    r'가-힯Ѐ-ӿ؀-ۿ֐-׿฀-๿]')

# Ventana alrededor del año, como el valida() de fase8c-libreria-assemble.py:
# mirar el texto entero tumbaría un «1982» legítimo de un bloque de historia.
RX_PRECIO = re.compile(
    r'(precio|coste|cuesta|€|EUR|euros|tarifa|presupuesto|inversión|factura|'
    r'tendencia|tendencias)', re.I)
# Citas legales: «RD 1021/2022», «Reglamento (CE) 853/2004», «art. 34.9».
# (Ojo: «RD 1420/2006» ya NO es cita legal aceptada — quedó derogado el
#  22-dic-2022 por el RD 1021/2022; sólo vale dentro de «que derogó el
#  RD 1420/2006». Gate: `motor.PROHIBIDAS`.)
RX_LEGAL_ANTES = re.compile(
    r'(RD|R\.D\.|Real\s+Decreto|RDL|Reglamento|Directiva|Ley|LIS|BOE|Estatuto|'
    r'art\.|artículo|CTE|UNE|ISO|edición|desde|Reg\.)[^.]{0,60}$', re.I)


def guard_no_latinos(texto, etiqueta=''):
    """Aborta con el fragmento de contexto (§5.3.3). Es el mismo defecto que
    puso una Amanita phalloides dentro de un prompt de garum en producción."""
    fallos = []
    for m in RX_NO_LATINOS.finditer(texto):
        ini = max(0, m.start() - 40)
        fallos.append({'char': m.group(0), 'codepoint': f'U+{ord(m.group(0)):04X}',
                       'contexto': texto[ini:m.end() + 40].replace('\n', ' ')})
    if fallos:
        print(f'ABORTADO [{etiqueta}]: {len(fallos)} caracteres no latinos.',
              file=sys.stderr)
        for f in fallos[:10]:
            print(f'   {f["codepoint"]} → …{f["contexto"]}…', file=sys.stderr)
    return fallos


RX_CITA = re.compile(
    r'\(\s*[^)]{8,}\)'                 # «(Restauracion News, 2026-01-07)»
    r'|\(\s*(?:19|20)\d{2}\s*\)'       # «(2025)» — el ano de la fuente
    r'|\b\d{4}-\d{2}\b'                # fecha ISO de publicacion
    r'|seg\u00fan\s+(?:el\s+|la\s+|los\s+|datos\s+de\s+)?'
    r'[A-Z\u00c1\u00c9\u00cd\u00d3\u00da]')  # «segun Profesional Horeca»
# Solo anios REALMENTE pasados (<= 2025): «previsiones para 2026» es correcto.
RX_TENDENCIA_CADUCA = re.compile(
    r'(tendencias?|previsiones?|novedades)[^.\n]{0,25}'
    r'\b(?:19\d{2}|20[01]\d|202[0-5])\b'
    r'|\b(?:19\d{2}|20[01]\d|202[0-5])\s*[-/]\s*20[0-3]\d\b', re.I)


def erratas_fechas(texto):
    """Ningun anio anterior a 2026 a menos de 90 caracteres de lenguaje de
    precios (SPEC 5.6.7), sin tumbar las citas legales NI las fuentes fechadas.

    Un anio pasado ACOMPANADO DE SU FUENTE no es una fecha caduca: es un dato
    fechado, que es justamente lo que exige la decision 7-bis.21 («cada cifra
    con su fuente y su fecha de corte»). La fecha caduca del incidente de las
    librerias de prompts —«precios HORECA de mayo de 2025»— no llevaba
    ninguna. Lo que NO se salva por llevar cita es un ROTULO con anio pasado
    («Tendencias 2025-2026» en un producto de agosto de 2026), que es el
    defecto COM-29: por eso va aparte y sin exencion.
    """
    fallos = []
    for m in re.finditer(r'\b(19\d{2}|20[0-2]\d)\b', texto):
        anio = int(m.group(1))
        if anio >= 2026:
            continue
        antes = texto[max(0, m.start() - 60):m.start()]
        if RX_LEGAL_ANTES.search(antes) or re.search(r'\d+/$', antes):
            continue
        ventana = texto[max(0, m.start() - 90):m.end() + 90]
        if not RX_PRECIO.search(ventana):
            continue
        if RX_CITA.search(texto[max(0, m.start() - 250):m.end() + 250]):
            continue
        fallos.append({'anio': anio, 'contexto': ventana.replace('\n', ' ')})
    for m in RX_TENDENCIA_CADUCA.finditer(texto):
        frag = m.group(0)
        # un titulo de fuente («Perspectivas 2026, informe de 2025») no es un
        # rotulo caduco del libro: si lleva su cita al lado, se deja pasar
        if RX_CITA.search(texto[max(0, m.start() - 120):m.end() + 120]):
            continue
        anios = [int(x) for x in re.findall(r'\b(19\d{2}|20\d{2})\b', frag)]
        if anios and min(anios) < 2026:
            fallos.append({'anio': min(anios), 'contexto': frag,
                           'tipo': 'rotulo_con_anio_pasado'})
    return fallos


# «el 60 % de los restaurantes cierra el primer año» y familia: cifra de
# mortalidad sin fuente. El research trae supervivencia del INE (SECT-09) y
# cierres (SECT-08); cualquier otra formulación se caza aquí (§5.4).
RX_MORTALIDAD = re.compile(
    r'(\d{1,3})\s*%[^.]{0,80}?(cierran?|fracasan?|no\s+sobreviv|desaparec|'
    r'quiebran?)', re.I)
RX_MORTALIDAD2 = re.compile(
    r'(cierran?|fracasan?|no\s+sobreviv|mueren)[^.]{0,60}?(\d{1,3})\s*%', re.I)


# El capítulo 5 se publicó con «Debo asegurarme de no mencionar el libro de
# visitas. No mencionar "fracaso" de restaurantes...» DENTRO del texto: el
# modelo volcó su razonamiento en el cuerpo. Y cuatro bloques escribieron «la
# tabla de abajo, que el maquetador insertará»: vocabulario del taller, no del
# libro. Las dos cosas se venden tal cual si nadie las mide.
RX_META = re.compile(
    r'(debo asegurar|me piden|se me pide|el maquetador|el prompt|'
    r'las instrucciones|como modelo|no puedo escribir|voy a redactar|'
    r'el guion dice|el usuario quiere|epígrafes que me)', re.I)


def erratas_meta(texto):
    fallos = []
    for m in RX_META.finditer(texto):
        ini = max(0, m.start() - 70)
        fallos.append(texto[ini:m.end() + 90].replace('\n', ' '))
    return fallos


def erratas_mortalidad(texto, permitidos):
    fallos = []
    for rx in (RX_MORTALIDAD, RX_MORTALIDAD2):
        for m in rx.finditer(texto):
            frag = m.group(0)
            if any(p in frag for p in permitidos):
                continue
            fallos.append(frag.replace('\n', ' '))
    return fallos


# --------------------------------------------------------------------------
# 1-bis. Calidad del TEXTO (gates nuevos, 2026-08-29).
# El gate anterior contaba páginas, palabras, tablas y cifras y NO miraba ni
# una sola vez si el texto se podía leer: dio «verde: true» con tres páginas de
# «(2) (2) (2)», cuatro párrafos cortados a media palabra y 292 erratas dentro
# (RT-06). Estos cinco detectores son los que faltaban.
# --------------------------------------------------------------------------
LEXICO = os.path.join(AQUI, '..', 'lexico-es-corpus.txt')
_LEX = None


def lexico():
    """Léxico español del propio grupo: tokens de >=4 letras que aparecen 2+
    veces en los 325 posts ES del blog, los datos de producto y los guiones.
    En este Mac no hay hunspell ni aspell: este fichero ES el diccionario del
    gate, y es reproducible (se regenera del repo). Cada línea es
    «forma sin tildes <TAB> forma con tildes más frecuente», porque el
    reparador tiene que devolver «degustación» y no «degustacion»."""
    global _LEX
    if _LEX is None:
        _LEX = {}
        if os.path.exists(LEXICO):
            with open(LEXICO, encoding='utf-8') as f:
                for ln in f:
                    if ln.startswith('#') or not ln.strip():
                        continue
                    partes = ln.rstrip('\n').split('\t')
                    _LEX[partes[0]] = (partes[1] if len(partes) > 1 else partes[0],
                                       int(partes[2]) if len(partes) > 2 else 1)
        else:
            print(f'  [aviso] no existe {LEXICO}: el gate de erratas queda ciego',
                  file=sys.stderr)
    return _LEX


def _norm_tok(w):
    return ''.join(c for c in unicodedata.normalize('NFD', w.lower())
                   if unicodedata.category(c) != 'Mn')


RX_TOKEN = re.compile(r'[a-zA-ZáéíóúüñÁÉÍÓÚÜÑ]+')
RX_ENTIDAD_HTML = re.compile(r'&(?:[a-zA-Z]{2,10}|#\d{2,5});')
RX_EURO_DOBLE = re.compile('€[\\s  ]*[\\d.,]+[\\s  ]*€')
RX_EURO_PREFIJO = re.compile('€[\\s  ]*\\d')
# Correcciones aplicadas al salir del modelo (limpiar_bloque):
# 2026-09-04: la clase empezaba por [\\d.] y «€.» (euro + punto final de frase) casaba como
# «€» + cifra «.», y la corrección lo convertía en «. €»: cuatro «237,35 . €» en la guía
# de food cost. La cifra tiene que EMPEZAR por dígito.
RX_EURO_DOBLE_FIX = re.compile('€[\\s\u00a0\u202f]*(\\d[\\d.,]*[\\s\u00a0\u202f]*€)')
RX_EURO_PRE_FIX = re.compile('€[\\s\u00a0\u202f]*(\\d[\\d.,]*)')
RX_EURO_UNIDAD = re.compile('([\\d.,]+)\u202f€([\\s\u00a0\u202f]*(?:años|año|meses|mes|'
                            'días|día|semanas|semana|personas|persona|plazas|plaza|'
                            'cubiertos|puntos|%))')
# «(P&L Mensual!B13)»: sintaxis de prompt que se coló porque el gate viejo
# sólo buscaba «.xlsx!» y aquí la cita va sin nombre de fichero (RT-11).
RX_CITA_HOJA = re.compile(
    r'[A-Za-zÁÉÍÓÚÑáéíóúñ&][\w áéíóúñÁÉÍÓÚÑ&]{0,28}!\s*\$?[A-Z]{1,2}\$?\d{1,4}\b')
RX_FUGA_MODELO = re.compile(
    r'\b(ChatGPT|GPT-?[0-9]?|Claude|DeepSeek|Gemini|OpenAI|Anthropic|'
    r'como modelo de lenguaje|como IA\b)\b', re.I)
FIN_DE_FRASE = '.!?:»"”’)…'


def erratas_degeneracion(texto, max_rep=12, max_linea=1500):
    """Bucle del modelo: n-grama repetido SEGUIDO o línea kilométrica sin
    puntos. Es lo que imprimió tres páginas de «(2) (2) (2)» en el PDF de 85 €
    (RT-01).

    Se mide la racha CONSECUTIVA, no la frecuencia. La primera versión contaba
    el token más frecuente de la línea y rechazaba párrafos correctos porque en
    un párrafo de 300 palabras «de» sale más de veinte veces: un contador de
    frecuencia no distingue el español del bucle. Lo que no ocurre nunca en
    prosa es la misma secuencia repetida doce veces SEGUIDAS.
    """
    fuera = []
    for i, ln in enumerate(texto.split('\n'), 1):
        s = ln.strip()
        if len(s) > max_linea and s.count('.') < len(s) / 400.0:
            fuera.append({'motivo': 'linea_larga_sin_puntos', 'linea': i,
                          'chars': len(s), 'muestra': s[:110]})
            continue
        toks = s.split()
        if len(toks) < max_rep * 2:
            continue
        for n in (1, 2, 3, 4):
            gram = [' '.join(toks[k:k + n]) for k in range(0, len(toks) - n + 1, n)]
            racha = mejor = 1
            cual = ''
            for a, b in zip(gram, gram[1:]):
                racha = racha + 1 if a == b else 1
                if racha > mejor:
                    mejor, cual = racha, a
            if mejor > max_rep:
                fuera.append({'motivo': f'{n}-grama repetido {mejor} veces seguidas',
                              'linea': i, 'muestra': cual[:80]})
                break
    return fuera


def erratas_truncamiento(texto, min_chars=40):
    """Todo párrafo de prosa termina en signo de cierre. Cuatro bloques se
    publicaron cortados a media palabra («La fre», «el comensal no») porque
    bridge() aceptaba una salida corta como válida (RT-05/RD-02)."""
    fuera = []
    for i, ln in enumerate(texto.split('\n'), 1):
        s = ln.strip()
        if len(s) < min_chars or s.startswith(('#', '|', '-', '*', '>', '`')):
            continue
        if s.endswith('|') or re.match(r'^\d+\.\s', s):
            continue
        cola = s.rstrip('*_ ')
        if cola and cola[-1] not in FIN_DE_FRASE:
            fuera.append({'linea': i, 'cola': s[-70:]})
    return fuera


RX_ARRANQUE_OK = re.compile(r'^[A-ZÁÉÍÓÚÑÜ0-9«"\u201c¿¡—\-\*#|>`\[(]')


def erratas_arranque(texto, min_chars=40):
    """Párrafo cortado por el PRINCIPIO. Complementa a erratas_truncamiento(),
    que sólo mira el final.

    Cazado el 2026-08-31: junto a las tres páginas de «(2)» del cap. 15 había
    una línea suelta —«, no, el problema del cliente es que es una. (4) No hay
    una.»— que pasaba LOS CUATRO guards: no es larga, no repite n-gramas,
    termina en punto y no menciona ningún motor. Si el bloque sólo hubiera
    degenerado por el arranque, se habría publicado.

    En prosa española un párrafo empieza por mayúscula, cifra, comilla o signo
    de apertura; nunca por coma o minúscula. Validado sobre los 29 ficheros del
    corpus de esta guía (3 documentos + 26 bloques en caché): 2 hallazgos, que
    son exactamente las 2 líneas corruptas. Cero falsos positivos.
    """
    fuera = []
    for i, ln in enumerate(texto.split('\n'), 1):
        s = ln.strip()
        if len(s) < min_chars:
            continue
        if s.startswith(('#', '|', '-', '*', '>', '`', '!', '[')):
            continue
        if not RX_ARRANQUE_OK.match(s):
            fuera.append({'linea': i, 'arranque': s[:70]})
    return fuera


def _parentesis_desequilibrados(s):
    """Balance de paréntesis con tolerancia a enumeradores legales: un «)»
    sin «(» pendiente y precedido por una letra o un número de 1-2 caracteres
    («letras a) a g)», «apartado 2)») es notación, no una errata."""
    prof = 0
    for m, ch in enumerate(s):
        if ch == '(':
            prof += 1
        elif ch == ')':
            if prof > 0:
                prof -= 1
            else:
                ant = s[max(0, m - 3):m]
                if not re.search(r'(?:^|[^\w(])[a-zA-Z0-9]{1,2}$', ant):
                    return True          # cierre sin apertura y sin enumerador
    return prof != 0


def erratas_parentesis(texto):
    """Paréntesis sin pareja por línea. 2026-09-04: los enumeradores legales
    «letras a) a g)» o «apartado 2)» son notación española correcta y no un
    paréntesis abierto (cap. 4 de la guía de food cost abortó seis veces por
    «letra g)»); ver _parentesis_desequilibrados."""
    fuera = []
    for i, ln in enumerate(texto.split('\n'), 1):
        s = ln.strip()
        if s.startswith('|'):
            continue
        if _parentesis_desequilibrados(s):
            fuera.append({'linea': i, 'abre': s.count('('),
                          'cierra': s.count(')'), 'muestra': s[:130]})
    return fuera


def erratas_ortograficas(texto, permitidas=(), min_freq_base=6, minlen=5):
    """Errata por letra caída o transpuesta: token que NO está en el léxico y
    que, reponiéndole una letra en posición interior o intercambiando dos
    adyacentes, da una palabra del léxico. Las 292 de la v2.0 («degusación»,
    «sofware», «decreo legislaivo») salen todas aquí (RT-03).

    Cada hallazgo lleva `veces` (cuántas veces aparece la errata) y `apoyo`
    (cuántas veces aparece en ESTE MISMO texto la palabra correcta). El apoyo
    es lo que separa una errata de una palabra rara legítima: «renabilidad»
    convive con 30 «rentabilidad», mientras que «capota» no tiene detrás
    ninguna «capta».
    """
    lex = lexico()
    perm = {_norm_tok(w) for w in permitidas}
    toks = [_norm_tok(w) for w in RX_TOKEN.findall(texto)]
    freq = collections.Counter(toks)
    conocidas = {w for w, n in freq.items() if n >= min_freq_base} | set(lex)
    fuera = {}
    for w, n in freq.items():
        if len(w) < minlen or w in conocidas or w in perm:
            continue
        cand = _reconstruir(w, conocidas)
        if cand:
            fuera[w] = {'errata': w, 'probable': cand, 'veces': n,
                        'apoyo': freq.get(cand, 0),
                        'freq_lexico': lex.get(cand, ('', 0))[1]}
    return sorted(fuera.values(), key=lambda d: (-d['veces'], d['errata']))


_ABC = 'abcdefghijklmnopqrstuvwxyzñ'
MIN_FREQ_LEXICO = 5      # la palabra correcta tiene que ser corriente
MIN_APOYO_DOC = 3        # y tiene que estar YA escrita bien en este documento


def _reconstruir(w, conocidas):
    """La palabra del léxico a UNA pulsación de distancia, por dos caminos y
    sólo dos: letra caída en posición INTERIOR y dos letras intercambiadas.

    Los otros dos caminos evidentes se probaron y se descartaron con datos
    (2026-08-29, sobre los 64 bloques de la versión anterior): quitar una letra
    y añadirla al final no reparan, DESTROZAN —«tasado» → «asado» 32 veces,
    «nominal» → «nómina», «provincias» → «provincia»—, porque en español
    quitar una letra o cambiar la última suele dar otra palabra válida. Y la
    sustitución de una letra por otra multiplica los falsos positivos sin
    cubrir el defecto medido, que es la letra que se cae.
    """
    for k in range(1, len(w)):
        for ch in _ABC:
            c = w[:k] + ch + w[k:]
            if c in conocidas:
                return c
    for k in range(len(w) - 1):
        c = w[:k] + w[k + 1] + w[k] + w[k + 2:]
        if c in conocidas:
            return c
    return None


def _con_mayusculas(original, nueva):
    if original.isupper():
        return nueva.upper()
    if original[:1].isupper():
        return nueva[:1].upper() + nueva[1:]
    return nueva


def reparar_erratas(texto, permitidas=()):
    """Repone la letra caída, y SÓLO cuando no hay duda. MEDIDO el 2026-08-29:
    este modelo deja caer una letra cada 1.000-1.600 palabras (la «t» y la «c»,
    sobre todo), así que rechazar el bloque entero por una errata quemaba los
    seis reintentos y abortaba el capítulo; y reparar a la ligera es peor que
    la errata, porque cambia el significado sin que se note.

    Por eso la corrección exige las DOS pruebas a la vez: la palabra correcta
    es corriente en el léxico (>= MIN_FREQ_LEXICO) y ya está escrita bien en
    este mismo documento (>= MIN_APOYO_DOC). Lo que no pasa las dos pruebas no
    se toca: se queda marcado y el gate `sin_erratas` lo pone en rojo para que
    lo mire una persona. Devuelve (texto, [cambios]) y los cambios van al
    informe: nada se corrige en silencio.
    """
    fallos = erratas_ortograficas(texto, permitidas)
    # ERRATAS_FORZADAS: correcciones REVISADAS UNA A UNA por una persona, para
    # las erratas que el automatismo detecta pero no se atreve a tocar porque
    # la palabra buena no aparece bien escrita en el propio documento
    # («panalla» → «pantalla» cuando «pantalla» sólo sale una vez). Es una
    # lista corta y explícita, no una heurística.
    seguros = [d for d in fallos
               if (d['errata'] in ERRATAS_FORZADAS
                   or (d['freq_lexico'] >= MIN_FREQ_LEXICO
                       and d['apoyo'] >= MIN_APOYO_DOC))]
    if not seguros:
        return texto, []
    lex = lexico()
    mapa = {}
    for d in seguros:
        if d['errata'] in ERRATAS_FORZADAS:
            mapa[d['errata']] = ERRATAS_FORZADAS[d['errata']]
        else:
            mapa[d['errata']] = lex.get(d['probable'], (d['probable'], 0))[0]
    cambios = []

    def sub(m):
        w = m.group(0)
        n = _norm_tok(w)
        if n not in mapa:
            return w
        nueva = _con_mayusculas(w, mapa[n])
        cambios.append({'de': w, 'a': nueva})
        return nueva

    return RX_TOKEN.sub(sub, texto), cambios


def erratas_entidades(texto):
    """«acci&oacute;n» impreso tal cual en el PDF (RD-04)."""
    return sorted(set(RX_ENTIDAD_HTML.findall(texto)))


def erratas_euro(texto):
    """«  €1.189.944,24 €» y «€10 años» en la petición al banco (RD-08/RT-09).
    El símbolo va SIEMPRE detrás de la cifra y nunca delante."""
    fuera = []
    for rx, tipo in ((RX_EURO_DOBLE, 'doble'), (RX_EURO_PREFIJO, 'prefijo')):
        for m in rx.finditer(texto):
            fuera.append({'tipo': tipo,
                          'muestra': texto[max(0, m.start() - 45):m.end() + 25]
                          .replace('\n', ' ')})
    return fuera


def erratas_citas_hoja(texto):
    return sorted(set(m.group(0) for m in RX_CITA_HOJA.finditer(texto)))


def erratas_fuga_modelo(texto):
    return sorted(set(m.group(0) for m in RX_FUGA_MODELO.finditer(texto)))


def epigrafes_ausentes(md_text, capitulos):
    """Cada epígrafe del guion tiene que estar como ### en SU capítulo. Sin
    esto, el cap. 15 se publicó sin los dos epígrafes de la matriz de Kasavana
    y Smith —que es lo que la landing promete— y el gate no dijo nada
    (RT-02/RT-04)."""
    fuera = []
    por_cap = {}
    for tr in re.split(r'^## ', md_text, flags=re.M)[1:]:
        m = re.match(r'(\d+)\. ', tr)
        if m:
            por_cap[int(m.group(1))] = tr
    for cap in capitulos:
        cuerpo = por_cap.get(cap['n'], '')
        vistos = [_norm_tok(x.strip()) for x in
                  re.findall(r'^###\s+(.+)$', cuerpo, re.M)]
        for e in cap.get('epigrafes', []):
            ne = _norm_tok(e.strip())
            if not any(ne == v or ne in v or v in ne for v in vistos):
                fuera.append({'cap': cap['n'], 'epigrafe': e})
    return fuera


# --------------------------------------------------------------------------
# 2. Lectura de los xlsx — la única fuente de cifras del producto (§7-bis.7)
# --------------------------------------------------------------------------
_WB = {}


def wb(xlsx_dir, fichero):
    """Workbook cacheado, SIEMPRE en data_only y SIEMPRE de sólo lectura."""
    if fichero not in _WB:
        from openpyxl import load_workbook
        ruta = os.path.join(xlsx_dir, fichero)
        if not os.path.exists(ruta):
            raise SystemExit(f'ABORTADO: falta {ruta}')
        _WB[fichero] = load_workbook(ruta, data_only=True)
    return _WB[fichero]


def celda(xlsx_dir, ref):
    """ref = 'fichero.xlsx!Hoja!C27' → valor. Cita fichero:hoja:celda."""
    fichero, hoja, coord = ref.split('!')
    ws = wb(xlsx_dir, fichero)[hoja]
    return ws[coord].value


def eur(v, dec=0):
    if v is None:
        return ''
    try:
        v = float(v)
    except (TypeError, ValueError):
        return str(v)
    s = f'{v:,.{dec}f}'.replace(',', '\x00').replace('.', ',').replace('\x00', '.')
    return f'{s}{NARROW}€'


def pct(v, dec=1):
    if v is None:
        return ''
    try:
        v = float(v)
    except (TypeError, ValueError):
        return str(v)
    s = f'{v * 100:,.{dec}f}'.replace(',', '\x00').replace('.', ',').replace('\x00', '.')
    return f'{s}{NARROW}%'


def num(v, dec=0):
    if v is None:
        return ''
    try:
        v = float(v)
    except (TypeError, ValueError):
        return str(v)
    return f'{v:,.{dec}f}'.replace(',', '\x00').replace('.', ',').replace('\x00', '.')


FORMATOS = {'eur': eur, 'eur2': lambda v: eur(v, 2), 'pct': pct, 'pct1': pct,
            'pct2': lambda v: pct(v, 2),
            'pct0': lambda v: pct(v, 0), 'num': num, 'num1': lambda v: num(v, 1),
            'num2': lambda v: num(v, 2), 'txt': lambda v: '' if v is None else str(v)}


def formatear(v, fmt):
    if v is None:
        return ''
    if fmt in FORMATOS:
        return FORMATOS[fmt](v)
    return str(v)


def resolver_cifras(xlsx_dir, cifras):
    """[(etiqueta, ref, fmt)] → [(etiqueta, valor_formateado, ref)]."""
    fuera = []
    for etiqueta, ref, fmt in cifras:
        v = celda(xlsx_dir, ref)
        fuera.append((etiqueta, formatear(v, fmt), ref, v))
    return fuera


RX_ETIQUETA_PCT = re.compile(r'\(\s*%\s*\)|\ben\s*%|\(porcentaje\)|%\s*$')


def es_fila_porcentual(celdas):
    """La etiqueta de la fila (o el título de la columna) declara porcentaje."""
    return any(isinstance(c, str) and RX_ETIQUETA_PCT.search(c) for c in celdas)


def construir_tabla(xlsx_dir, t):
    """Devuelve (markdown, n_filas). Dos formas:
       - literal:  {'cabecera': [...], 'filas': [[...], ...]}
       - del xlsx: {'src': ('fichero.xlsx','Hoja'), 'cols': [(titulo, col, fmt)],
                    'filas': (fila_ini, fila_fin), 'saltar_vacias': True}
    """
    if 'cabecera' in t:
        cab, filas = t['cabecera'], [[str(c) for c in f] for f in t['filas']]
    else:
        fichero, hoja = t['src']
        ws = wb(xlsx_dir, fichero)[hoja]
        cab = [c[0] for c in t['cols']]
        filas = []
        ini, fin = t['filas']
        for r in range(ini, fin + 1):
            fila, crudos = [], []
            for _tit, col, fmt in t['cols']:
                if col.startswith('='):          # literal por columna
                    fila.append(col[1:])
                    crudos.append(None)
                    continue
                v = ws[f'{col}{r}'].value
                crudos.append(v)
                fila.append(formatear(v, fmt))
            # 2026-08-29 (RD-07/RD-13): una fila cuya ETIQUETA dice «(%)» se
            # imprimía con el formato de su columna, así que el cuadro de mando
            # que se le enseña al banco publicaba «Margen EBITDA (%) | 0 €» y
            # el tipo de interés como «0,06» bajo un encabezado de porcentaje.
            # Un solo bug envenenaba cuatro tablas: la etiqueta manda sobre la
            # columna.
            # La etiqueta no siempre es la primera celda: en la tabla de CAPEX
            # la columna A es el «#» y el concepto va en la B.
            if es_fila_porcentual(fila[:2]):
                for k in range(1, len(fila)):
                    v = crudos[k]
                    if isinstance(v, (int, float)) and not isinstance(v, bool):
                        fila[k] = pct(v, 1) if abs(v) <= 1.5 else num(v, 1) + NARROW + '%'
            if t.get('saltar_vacias', True) and not any(x for x in fila[1:]):
                continue
            filas.append(fila)
        for extra in t.get('extra_filas', []):
            filas.append([str(x) for x in extra])
    md = ['| ' + ' | '.join(cab) + ' |',
          '|' + '|'.join(['---'] * len(cab)) + '|']
    for f in filas:
        f = list(f) + [''] * (len(cab) - len(f))
        md.append('| ' + ' | '.join(x.replace('|', '/') for x in f[:len(cab)]) + ' |')
    return '\n'.join(md), len(filas)


# --------------------------------------------------------------------------
# 3. Research del sector (§7-bis.21) — sin fuente, no entra
# --------------------------------------------------------------------------


def cargar_research(path=RESEARCH, esperar_min=20):
    """Espera con sleep(60) hasta `esperar_min` minutos si aún no existe."""
    esperado = 0
    while not os.path.exists(path) and esperado < esperar_min * 60:
        print(f'  [research] aún no existe {path}; esperando 60 s '
              f'({esperado // 60}/{esperar_min} min)', flush=True)
        time.sleep(60)
        esperado += 60
    if not os.path.exists(path):
        print('  [research] NO llegó: los capítulos se escriben sin cifra de '
              'sector (reformuladas). Anotado en el informe.', flush=True)
        return {'_meta': {'ausente': True}, 'datos': []}
    with open(path, encoding='utf-8') as f:
        return json.load(f)


def indexar_research(res):
    return {d['id']: d for d in res.get('datos', [])}


def bloque_research(idx, ids, breve=False):
    """Texto para el prompt: cada cifra con su fuente. Los ids sin cifra o de
    fiabilidad baja se declaran como HUECO y el modelo debe reformular."""
    lineas, usados, huecos = [], [], []
    for i in ids:
        d = idx.get(i)
        if d is None:
            huecos.append(i)
            continue
        if d.get('cifra') is None or d.get('fiabilidad') == 'baja' or not d.get('url'):
            huecos.append(i)
            lineas.append(f'- [{i}] HUECO SIN FUENTE — «{d["dato"]}»: NO escribas '
                          'ninguna cifra sobre esto; formúlalo en cualitativo.')
            continue
        cifra = d['cifra']
        unidad = d.get('unidad') or ''
        # 2026-08-29: cortar por CARACTER metia en el libro «(sobre el
        # Anuario 2025 de Hostel, 2026-01-07)» y «(sobre el es, 2025-12)» —el
        # modelo copia el recorte literal—, y de paso dejaba paréntesis sin
        # cerrar (RT-07). Se corta por límite de palabra y con sufijo explícito.
        titulo = textwrap.shorten(d['fuente_titulo'], width=110, placeholder='…')
        nota = textwrap.shorten(d.get('nota') or '-', width=240, placeholder='…')
        lineas.append(
            f'- [{i}] {d["dato"]}: **{cifra} {unidad}**. Fuente obligatoria a citar '
            f'en el texto: «{titulo}» '
            f'({d.get("fecha_publicacion") or d.get("anio_del_dato")}). '
            + ('' if breve else f'Nota: {nota}'))
        usados.append(i)
    return '\n'.join(lineas), usados, huecos


# --------------------------------------------------------------------------
# 4. bridge.py — el motor de redacción (regla capital)
# --------------------------------------------------------------------------
SYSTEM = (
    'Eres John Guerrero, chef desde los 17 años y consultor gastronómico desde '
    '2010, con más de 200 aperturas asesoradas en España. Escribes en español '
    'de España, en primera persona del plural o impersonal, con tono humano y '
    'profesional, para un lector que va a jugarse su dinero. '
    'REGLAS ABSOLUTAS: (1) NO inventes NINGUNA cifra: usa exclusivamente las '
    'que te doy en el prompt, con el mismo formato; si necesitas un número que '
    'no está en la lista, escribe la frase en cualitativo. (2) Escribe SOLO en '
    'alfabeto latino: ni un carácter chino, japonés, coreano, cirílico, árabe, '
    'hebreo ni tailandés, ni siquiera dentro de un paréntesis o de una cita. '
    '(3) Nada de introducciones meta ("en este capítulo veremos"), nada de '
    'conclusiones que resuman lo dicho, nada de adjetivos de relleno. Prosa '
    'densa: cada párrafo tiene que aportar un dato, un procedimiento o una '
    'decisión. (4) No escribas tablas Markdown ni encabezados de nivel 1 o 2: '
    'las tablas y los títulos los pone el maquetador. (5) No cites años '
    'anteriores a 2026 junto a precios ni a tendencias. (6) Nunca digas que un '
    'porcentaje de restaurantes "cierra" o "fracasa" salvo que la cifra te la '
    'haya dado yo con su fuente. (7) TERMINA SIEMPRE la última frase: nada de '
    'párrafos que acaben a media palabra o sin punto. (8) El símbolo del euro '
    'va DETRÁS de la cifra y separado («1.250,00 €»), nunca delante y nunca '
    'dos veces; y no pongas «€» junto a algo que no sea dinero (años, meses, '
    'personas). (9) No escribas entidades HTML («&oacute;», «&nbsp;»): escribe '
    'el carácter. (10) No repitas nunca una palabra, un número ni un símbolo '
    'en bucle: si te encuentras repitiendo, cierra la frase y sigue. (11) '
    'Nombra los ficheros y las hojas del pack en lenguaje de libro («la hoja '
    '«Inversión» de plan-financiero-3-anos.xlsx»), NUNCA con la sintaxis de '
    'celda («P&L Mensual!B13»). (12) Cuida la ortografía carácter a carácter: '
    'la edición anterior se publicó con «degusación», «sofware» y «decreo '
    'legislaivo», y eso en un producto de pago es inaceptable.')


def reparar_cola(texto):
    """Recorta la frase incompleta del final de un párrafo. Un truncamiento de
    cola SE PUEDE reparar (se pierde media frase, no el capítulo); una
    degeneración o una entidad HTML, no. Se usa sólo cuando bridge ha agotado
    los reintentos y el ÚNICO defecto que queda es el corte."""
    fuera = []
    for ln in texto.split('\n'):
        s = ln.rstrip()
        cuerpo = s.strip()
        if (len(cuerpo) < 40 or cuerpo.startswith(('#', '|', '-', '*', '>', '`'))
                or cuerpo.endswith('|')):
            fuera.append(ln)
            continue
        if cuerpo.rstrip('*_ ')[-1:] in FIN_DE_FRASE:
            fuera.append(ln)
            continue
        corte = max(s.rfind('. '), s.rfind('! '), s.rfind('? '),
                    s.rfind('.\u00bb'), s.rfind(': '))
        if corte > 60:
            fuera.append(s[:corte + 1].rstrip())
    return '\n'.join(fuera)


# Palabras que el detector ortográfico marca y NO son erratas (nombres
# propios, términos del oficio, extranjerismos). Lo fija main() desde el guion.
ERRATAS_PERMITIDAS = ()
# Cada letra repuesta queda registrada aquí y sale en el informe: una
# corrección automática que no se puede auditar es peor que la errata.
ERRATAS_REPARADAS = []
# {errata_normalizada: forma correcta} revisadas a mano; lo fija main() desde
# el guion (clave `erratas_forzadas` de gates).
ERRATAS_FORZADAS = {}


def defectos_de_bloque(texto):
    """Lo que descalifica una salida de bridge ANTES de guardarla: bucle del
    modelo, párrafo cortado a media frase, entidad HTML cruda, mención al
    motor, paréntesis sin cerrar o erratas por letra caída. Se comprueba en
    cada intento, no al final: un bloque roto que llega al .md ya no lo ve
    nadie hasta el PDF impreso."""
    if not texto or not texto.strip():
        return [{'tipo': 'vacio', 'muestra': ''}]
    malos = []
    for d in erratas_degeneracion(texto):
        malos.append({'tipo': 'degeneracion', 'muestra': d['muestra']})
    for d in erratas_truncamiento(texto):
        malos.append({'tipo': 'truncamiento', 'muestra': d['cola']})
    for d in erratas_arranque(texto):
        malos.append({'tipo': 'arranque_cortado', 'muestra': d['arranque']})
    for e in erratas_entidades(texto):
        malos.append({'tipo': 'entidad_html', 'muestra': e})
    for f in erratas_fuga_modelo(texto):
        malos.append({'tipo': 'fuga_de_modelo', 'muestra': f})
    for d in erratas_parentesis(texto):
        malos.append({'tipo': 'parentesis', 'muestra': d['muestra']})
    return malos


_MT_QUE_FUNCIONA = []


def bridge(prompt, salida_txt, palabras_min, max_tokens=12000, intentos=6,
           temperatura=0.45, verbose=True, prompt_corto=None):
    """Llama a bridge.py. Medido el 2026-08-29: con --max-tokens 8192 el modelo
    puede agotar el presupuesto razonando y devolver un fichero VACÍO, y la API
    devuelve contenido vacío con rc=0 de forma intermitente (tres veces
    seguidas en el capítulo 1, y el MISMO prompt funcionó a la cuarta). Por eso
    se reintenta con presupuestos distintos, con espera entre intentos y
    variando la temperatura: un vacío no es un fallo del prompt."""
    # MEDIDO el 2026-08-29 en este Mac: con `--max-tokens 8192` este modelo
    # devuelve contenido VACÍO de forma sistemática (0 palabras en TODOS los
    # primeros intentos de los 4 primeros capítulos) y con 6000 responde. Se
    # intenta primero el presupuesto que ya funcionó en esta ejecución para no
    # quemar una llamada por bloque; si aún no hay ninguno, se prueba 8192
    # (que es lo que manda la regla) y se cae a 6000.
    # MEDIDO OTRA VEZ el 2026-08-29, y esta vez con el mismo prompt Y el mismo
    # --system que usa el pipeline: `deepseek-v4-flash` es un modelo de
    # razonamiento y con --max-tokens 6000 gasta los 6000 razonando y devuelve
    # el fichero VACÍO (2.556 in / 6.000 out / 0 palabras). El MISMO prompt sin
    # --system sí responde, y con --max-tokens 12000 responde en 60 s con 1.146
    # palabras gastando 3.887. O sea: el techo no sobraba, FALTABA. Esa es
    # también la causa de fondo de los cuatro párrafos cortados a media palabra
    # y de las tres páginas de «(2)»: presupuesto agotado a mitad de la salida.
    presupuestos = [max_tokens, 14000, 10000, 16000, 11000, 8192][:intentos]
    if _MT_QUE_FUNCIONA:
        pref = _MT_QUE_FUNCIONA[-1]
        presupuestos = [pref] + [x for x in presupuestos if x != pref]
    texto, mejor = '', ''
    for k, mt in enumerate(presupuestos):
        if k:
            time.sleep(15)
        respirar()
        modelo = MODELO if k < 2 else MODELO_FALLBACK
        cmd = [sys.executable, BRIDGE, '--task', 'content', '--domain', 'aichef',
               '--lang', 'es', '--model', modelo, '--max-tokens', str(mt),
               '--temperature', str(round(temperatura + 0.05 * (k % 3), 2)),
               '--system', SYSTEM,
               # a partir del 3.er intento se manda la versión CORTA del guion:
               # los prompts que más fallan son los que llevan 18 fichas de
               # research con sus notas, y acortar el prompt los desatasca.
               '--prompt', (prompt_corto or prompt) if k >= 2 else prompt,
               '--output', salida_txt]
        r = subprocess.run(cmd, capture_output=True, text=True)
        texto = ''
        if os.path.exists(salida_txt):
            with open(salida_txt, encoding='utf-8') as f:
                texto = f.read().strip()
        n = len(texto.split())
        if verbose:
            print(f'    bridge intento {k + 1}/{len(presupuestos)} '
                  f'(max_tokens={mt}, {modelo}) → {n} palabras', flush=True)
        # 2026-08-29: una salida LARGA puede seguir siendo basura. Cuatro
        # bloques se publicaron cortados a media palabra («La fre») y uno con
        # tres páginas de «(2) (2) (2)» porque aquí sólo se contaban palabras.
        malos = defectos_de_bloque(texto)
        if not malos and n > len(mejor.split()):
            mejor = texto              # la mejor salida LIMPIA vista hasta ahora
        if n >= palabras_min and not malos:
            if mt not in _MT_QUE_FUNCIONA:
                _MT_QUE_FUNCIONA.append(mt)
            return texto
        if malos and verbose:
            print(f'    bridge intento {k + 1}: RECHAZADO por '
                  f'{", ".join(sorted(set(m["tipo"] for m in malos)))} '
                  f'→ {malos[0]["muestra"][:90]!r}', flush=True)
        if r.returncode != 0 and verbose:
            print(f'    bridge rc={r.returncode} :: {r.stderr.strip()[:300]}',
                  flush=True)
    # Una salida LIMPIA pero algo corta vale más que otra llamada: el bloque se
    # queda con la mejor de las que no tenían ningún defecto, y quien decide si
    # el capítulo es demasiado corto es el gate de palabras por capítulo, que
    # mide el capítulo entero y no el tramo. Sin esto, un bloque que se queda a
    # cincuenta palabras del objetivo quema seis llamadas y aborta el build.
    if mejor:
        print(f'    bridge: se acepta la mejor salida limpia '
              f'({len(mejor.split())} palabras, objetivo {palabras_min})', flush=True)
        with open(salida_txt, 'w', encoding='utf-8') as f:
            f.write(mejor)
        return mejor
    # Último recurso: si lo único que queda es una frase cortada al final, se
    # recorta la frase; el resto del bloque es bueno. Cualquier otro defecto
    # (bucle, entidad HTML, fuga del modelo, paréntesis) aborta: ese texto no
    # se puede vender.
    malos = defectos_de_bloque(texto)
    if texto and malos and all(m['tipo'] == 'truncamiento' for m in malos):
        reparado = reparar_cola(texto)
        if reparado.strip() and not defectos_de_bloque(reparado):
            print('    bridge: cola truncada RECORTADA '
                  f'({len(texto.split())} → {len(reparado.split())} palabras)',
                  flush=True)
            with open(salida_txt, 'w', encoding='utf-8') as f:
                f.write(reparado)
            return reparado
    raise SystemExit(
        f'ABORTADO: bridge.py no devolvió texto LIMPIO para {salida_txt} '
        f'({len(texto.split())} palabras, defectos: '
        f'{[m["tipo"] for m in malos][:5]})')


# --------------------------------------------------------------------------
# 5. Prompts de capítulo (§5.2) — el guion cerrado, no un título
# --------------------------------------------------------------------------


def prompt_bloque(cap, bloque_epigrafes, palabras, ctx_cifras, ctx_sector,
                  guia, es_ultimo):
    partes = []
    partes.append(
        f'Escribe un tramo del capítulo {cap["n"]} — «{cap["titulo"]}» de la '
        f'{guia.get("tipo_doc", "guía")} profesional «{guia["titulo"]}» ({guia["subtitulo"]}).')
    partes.append(f'OBJETIVO DEL CAPÍTULO: {cap["objetivo"]}')
    partes.append(
        'ESCRIBE EXACTAMENTE ESTOS EPÍGRAFES, cada uno como encabezado '
        'Markdown de nivel 3 (### ) con ese título literal, en este orden:\n'
        + '\n'.join(f'  ### {e}' for e in bloque_epigrafes))
    partes.append(f'EXTENSIÓN: unas {palabras} palabras en total para este tramo, '
                  'repartidas entre esos epígrafes. Prosa + listas con viñetas '
                  'cuando aporten; nunca tablas.')
    if cap.get('puntos'):
        partes.append('PUNTOS OBLIGATORIOS (tienen que aparecer, con su detalle '
                      'operativo):\n' + '\n'.join(f'  - {p}' for p in cap['puntos']))
    if ctx_cifras:
        partes.append(
            'CIFRAS DEL PROPIO PRODUCTO — son las ÚNICAS cifras de negocio que '
            'puedes escribir, y van tal cual (mismo separador de miles y misma '
            'coma decimal). Salen de las plantillas Excel que el lector tiene '
            'en el mismo pack, así que el texto y las hojas de cálculo tienen '
            'que decir lo mismo:\n' + ctx_cifras)
    if ctx_sector:
        partes.append(
            'DATOS DEL SECTOR — cada uno se cita CON SU FUENTE dentro del texto '
            '(entre paréntesis, con el nombre de la fuente y el año). Los '
            'marcados como HUECO SIN FUENTE no se escriben con número:\n'
            + ctx_sector)
    if cap.get('tablas_anunciadas'):
        partes.append(
            'Justo DESPUÉS de tu texto, en el documento, aparecerán estas '
            'tablas ya montadas: ' + '; '.join(cap['tablas_anunciadas']) +
            '. NO las escribas tú. Puedes remitir a ellas con naturalidad («la '
            'tabla de abajo», «el cuadro siguiente»), pero NUNCA menciones el '
            'proceso de edición ni palabras como «maquetador», «prompt», '
            '«instrucciones» o «guion»: el lector compra un libro, no ve el '
            'taller. Tampoco escribas tu propio razonamiento.')
    if cap.get('prohibido'):
        partes.append('LO QUE NO DEBES DECIR (son errores reales de la edición '
                      'anterior de esta guía y no se pueden repetir):\n'
                      + '\n'.join(f'  - {p}' for p in cap['prohibido']))
    partes.append(
        'FORMATO: empieza directamente por el primer «### ». No pongas título '
        'de capítulo, ni resumen inicial, ni cierre.'
        + ('' if es_ultimo else ' No cierres el capítulo: sigue otro tramo.'))
    return '\n\n'.join(partes)


def trocear(epigrafes, n_bloques):
    n_bloques = max(1, min(n_bloques, len(epigrafes)))
    tam = (len(epigrafes) + n_bloques - 1) // n_bloques
    return [epigrafes[i:i + tam] for i in range(0, len(epigrafes), tam)]


def _hojas_conocidas(xlsx_dir):
    """Los nombres de hoja REALES del producto, de mas largo a mas corto: hay
    hojas con espacios («Turnos Semana», «Cash Flow 12 Meses») y adivinar donde
    acaba el nombre con un regex generico se comia la palabra siguiente."""
    import glob
    from openpyxl import load_workbook
    nombres = set()
    for ruta in glob.glob(os.path.join(xlsx_dir, '*.xlsx')):
        try:
            nombres.update(load_workbook(ruta, read_only=True).sheetnames)
        except Exception:
            pass
    return sorted(nombres, key=len, reverse=True)


def normalizar_citas(texto, hojas=()):
    """El prompt le da al modelo cada cifra con su celda —«[fuente:
    plan-financiero-3-anos.xlsx!Inversion!C46]»— y el modelo la copia tal cual
    al cuerpo del libro. La trazabilidad SE QUEDA (decision 7-bis.7: el texto
    cita la hoja de la que sale el numero), pero escrita como se lee un libro y
    no como se escribe un prompt: sin la exclamacion, sin la celda y sin el
    corchete."""
    if hojas:
        alt = '|'.join(re.escape(h) for h in hojas)
        rx = re.compile(
            r'\[\s*fuente:\s*([\w-]+\.xlsx)!(' + alt + r')(?:![A-Z]+\d+)?\s*\]'
            r'|\(\s*([\w-]+\.xlsx)!(' + alt + r')(?:![A-Z]+\d+)?\s*\)'
            r'|([\w-]+\.xlsx)!(' + alt + r')(?:![A-Z]+\d+)?')

        def sub(m):
            g = m.groups()
            fich, hoja = (g[0], g[1]) if g[0] else \
                ((g[2], g[3]) if g[2] else (g[4], g[5]))
            return f'({fich}, hoja \u00ab{hoja}\u00bb)'
        texto = rx.sub(sub, texto)
    # red de seguridad: cualquier «fichero.xlsx!Hoja!Celda» que quede
    texto = re.sub(r'\[\s*fuente:\s*([\w-]+\.xlsx)![^!\]]+![A-Z]+\d+\s*\]',
                   r'(\1)', texto)
    texto = re.sub(r'([\w-]+\.xlsx)![^!\s,;.)\]]+![A-Z]+\d+', r'\1', texto)
    # 2026-08-29 (RT-11): el modelo también escribe la cita SIN el fichero
    # —«(P&L Mensual!B13)»—, y el gate viejo, que sólo buscaba «.xlsx!», la
    # daba por limpia. Se deja el nombre de la hoja, que es lo que el lector
    # puede abrir, y se quita la celda.
    if hojas:
        alt2 = '|'.join(re.escape(h) for h in hojas)
        texto = re.sub(r'\(\s*(' + alt2 + r')!\s*\$?[A-Z]{1,2}\$?\d{1,4}\s*\)',
                       lambda m: f'(hoja \u00ab{m.group(1)}\u00bb)', texto)
        texto = re.sub(r'\b(' + alt2 + r')!\s*\$?[A-Z]{1,2}\$?\d{1,4}\b',
                       lambda m: f'la hoja \u00ab{m.group(1)}\u00bb', texto)
    return texto


def limpiar_bloque(texto, hojas=()):
    """El modelo a veces devuelve H1/H2, una tabla o un preámbulo meta."""
    lineas = []
    for ln in texto.split('\n'):
        s = ln.strip()
        if s.startswith('# ') or s.startswith('## '):
            s = '### ' + s.lstrip('#').strip()
            ln = s
        if s.startswith('```'):
            continue
        lineas.append(ln)
    txt = '\n'.join(lineas)
    txt = re.sub(r'\n{3,}', '\n\n', txt).strip()
    # 2026-08-29 (RD-04): «acci&oacute;n especifica» salió impreso en el PDF.
    # El modelo devuelve entidades HTML de vez en cuando y nadie las miraba.
    if RX_ENTIDAD_HTML.search(txt):
        txt = html.unescape(txt)
    # 2026-08-29 (RD-08/RT-09): «  €1.189.944,24 €» y «€10 años» en la petición
    # al banco. El símbolo va SIEMPRE detrás; el prefijo sobra siempre.
    txt = re.sub(RX_EURO_DOBLE_FIX, lambda m: m.group(1), txt)
    txt = re.sub(RX_EURO_PRE_FIX, lambda m: m.group(1) + NARROW + '€', txt)
    txt = re.sub(RX_EURO_UNIDAD, lambda m: m.group(1) + m.group(2), txt)
    txt = re.sub('[ \u00a0]{2,}', ' ', txt)
    txt = normalizar_citas(txt, hojas)
    # quita un preámbulo antes del primer ###
    i = txt.find('### ')
    if i > 0:
        txt = txt[i:]
    return txt


# --------------------------------------------------------------------------
# 6. Generación y ensamblado
# --------------------------------------------------------------------------


def generar_capitulo(cap, guia, xlsx_dir, idx_research, dir_txt, forzar=False,
                     hojas=()):
    cifras = resolver_cifras(xlsx_dir, cap.get('cifras', []))
    ctx_cifras = '\n'.join(f'  - {e}: {v}   [fuente: {r}]' for e, v, r, _ in cifras)
    ctx_sector, usados, huecos = bloque_research(idx_research, cap.get('sector', []))
    ctx_sector_breve, _, _ = bloque_research(idx_research, cap.get('sector', []),
                                             breve=True)

    tablas = []
    for t in cap.get('tablas', []):
        md, nf = construir_tabla(xlsx_dir, t)
        tablas.append((t.get('titulo', ''), md, nf, t.get('nota')))
    cap['tablas_anunciadas'] = [t[0] for t in tablas]

    bloques = trocear(cap['epigrafes'], cap.get('bloques', 2))
    por_bloque = max(500, int(cap['palabras'] / len(bloques)))
    piezas = []
    for i, epis in enumerate(bloques):
        ruta = os.path.join(dir_txt, f'cap{cap["n"]:02d}_b{i + 1}.txt')
        if os.path.exists(ruta) and not forzar:
            with open(ruta, encoding='utf-8') as f:
                t = f.read().strip()
            malos = defectos_de_bloque(t)
            if len(t.split()) >= por_bloque * 0.6 and not malos:
                print(f'  cap {cap["n"]:02d} bloque {i + 1}: caché '
                      f'({len(t.split())} palabras)', flush=True)
                piezas.append(limpiar_bloque(t, hojas))
                continue
            if malos:
                # 2026-08-29: la caché era intocable y por eso las tres páginas
                # de «(2)» sobrevivieron a todas las reconstrucciones. Un .txt
                # con defectos se REGENERA aunque nadie haya pedido --regenerar.
                print(f'  cap {cap["n"]:02d} bloque {i + 1}: caché DESCARTADA '
                      f'({", ".join(sorted(set(m["tipo"] for m in malos)))})',
                      flush=True)
        p = prompt_bloque(cap, epis, por_bloque, ctx_cifras, ctx_sector, guia,
                          es_ultimo=(i == len(bloques) - 1))
        print(f'  cap {cap["n"]:02d} bloque {i + 1}/{len(bloques)} '
              f'({por_bloque} palabras objetivo)', flush=True)
        p_corto = prompt_bloque(cap, epis, por_bloque, ctx_cifras,
                                ctx_sector_breve, guia,
                                es_ultimo=(i == len(bloques) - 1))
        t = bridge(p, ruta, palabras_min=int(por_bloque * 0.72),
                   prompt_corto=p_corto)
        piezas.append(limpiar_bloque(t, hojas))

    cuerpo = '\n\n'.join(piezas)
    md = [f'## {cap["n"]}. {cap["titulo"]}', '', cuerpo]
    for titulo, tabla_md, _nf, nota in tablas:
        md += ['', f'**{titulo}**', '', tabla_md]
        if nota:
            md += ['', f'*{nota}*']
    return '\n'.join(md) + '\n', {'n': cap['n'], 'titulo': cap['titulo'],
                                  'tablas': len(tablas), 'sector_usado': usados,
                                  'sector_huecos': huecos,
                                  'cifras': [(e, v, r) for e, v, r, _ in cifras]}


def portada_e_indice(guia, capitulos):
    hoy = guia.get('fecha', 'agosto de 2026')
    p = [f'# {guia["titulo"]}', '',
         f'**{guia["subtitulo"]}**', '',
         f'*{guia["autor_linea"]}*', '',
         guia['portada_texto'], '',
         f'**Versión {guia.get("version", "2.0")} · {hoy} · aichef.pro/{guia["pid"]}**', '',
         '---', '', '## Índice', '']
    for c in capitulos:
        p.append(f'{c["n"]}. **{c["titulo"]}** — {c["resumen_indice"]}')
    p += ['', '---', '']
    return '\n'.join(p)


def cierre(guia):
    return '\n'.join([
        '', '---', '', '## Sobre el autor y condiciones de uso', '',
        guia['bio'], '',
        f'**Versión {guia.get("version", "2.0")} · {guia.get("fecha", "agosto de 2026")} · aichef.pro/{guia["pid"]} · info@aichef.pro**', '',
        guia['legal'], ''])


# --------------------------------------------------------------------------
# 7. Maquetado — DOCX + PDF desde el MISMO Markdown (§7-bis.15)
#    Patrón de kit-escandallos-v2_0/bono_guia.py, extendido con PageBreak,
#    portada, índice y metadatos.
# --------------------------------------------------------------------------
GOLD = '#B8860B'
DARK = '#1A1A1A'

# RT-18 (2026-08-29) — DEJADO POR ESCRITO, que es lo que faltaba: las fuentes
# del PDF son las Type-1 base de reportlab (Helvetica, Helvetica-Bold,
# Helvetica-Oblique), que sólo cubren WinAnsi/cp1252. Este mapa degrada a
# propósito el ESPACIO FINO (U+202F) y el GUION NO SEPARABLE (U+2011) que sí
# llevan los .md a un espacio y un guion normales. Consecuencia práctica: un
# gate o un parche que busque U+202F o U+2011 DENTRO DEL PDF no los encontrará
# jamás —igual que ya pasó con los .md de los kits—, y hay que medirlos sobre
# el Markdown. Si algún día se quiere conservar el espacio fino en el PDF, hay
# que embeber una TrueType en vez de usar las Type-1 base; no basta con quitar
# el mapa, porque entonces salen cuadrados .notdef.
MAPA_WINANSI = {
    '☐': '[  ]', '☑': '[X]', '☒': '[X]', '✅': 'OK',
    '✔': 'OK', '❌': 'X', '✖': 'X', '⭐': '*', '★': '*',
    '→': '->', '←': '<-', '─': '-', '━': '-',
    '−': '-', ' ': ' ', ' ': ' ',
    NARROW: ' ', NOBRK: '-',
    '–': '-', '—': '-', '…': '...',
    '\U0001F534': '', '\U0001F7E2': '', '\U0001F7E1': '',
}


def sanear(texto):
    if not isinstance(texto, str):
        return texto
    fuera = []
    for ch in texto:
        if ch in MAPA_WINANSI:
            fuera.append(MAPA_WINANSI[ch])
            continue
        try:
            ch.encode('cp1252')
        except UnicodeEncodeError:
            # último recurso: descomponer (ǎ → a) antes de descartar
            d = unicodedata.normalize('NFKD', ch).encode('ascii', 'ignore').decode()
            fuera.append(d)
            continue
        fuera.append(ch)
    return ''.join(fuera)


def sanear_bloques(bloques):
    fuera = []
    for tipo, contenido in bloques:
        if isinstance(contenido, str):
            contenido = sanear(contenido)
        elif tipo == 'table':
            contenido = [[re.sub(r'\s{2,}', ' ', sanear(c)).strip() for c in fila]
                         for fila in contenido]
        elif isinstance(contenido, list):
            contenido = [sanear(x) for x in contenido]
        fuera.append((tipo, contenido))
    return fuera


def restos_no_winansi(bloques):
    malos = {}
    for _, contenido in bloques:
        if isinstance(contenido, str):
            trozos = [contenido]
        elif isinstance(contenido, list) and contenido and isinstance(contenido[0], list):
            trozos = [c for fila in contenido for c in fila]
        elif isinstance(contenido, list):
            trozos = [x for x in contenido if isinstance(x, str)]
        else:
            continue
        for t in trozos:
            for ch in t:
                try:
                    ch.encode('cp1252')
                except UnicodeEncodeError:
                    malos[ch] = malos.get(ch, 0) + 1
    return malos


def parsear(md_text):
    lineas = md_text.split('\n')
    bloques = []
    i, n = 0, len(lineas)
    while i < n:
        linea = lineas[i].rstrip()
        if not linea.strip():
            i += 1
            continue
        if linea.strip() == '---':
            bloques.append(('hr', None))
            i += 1
            continue
        if linea.startswith('#### '):
            bloques.append(('h4', linea[5:].strip())); i += 1; continue
        if linea.startswith('### '):
            bloques.append(('h3', linea[4:].strip())); i += 1; continue
        if linea.startswith('## '):
            bloques.append(('h2', linea[3:].strip())); i += 1; continue
        if linea.startswith('# '):
            bloques.append(('h1', linea[2:].strip())); i += 1; continue
        if linea.strip().startswith('```'):
            codigo = []
            i += 1
            while i < n and not lineas[i].strip().startswith('```'):
                codigo.append(lineas[i]); i += 1
            i += 1
            bloques.append(('code', '\n'.join(codigo)))
            continue
        if linea.strip().startswith('|') and i + 1 < n and \
                re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lineas[i + 1]) and '-' in lineas[i + 1]:
            filas = [linea]
            i += 2
            while i < n and lineas[i].strip().startswith('|'):
                filas.append(lineas[i]); i += 1
            tabla = [[c.strip() for c in f.strip().strip('|').split('|')] for f in filas]
            bloques.append(('table', tabla))
            continue
        if re.match(r'^\s*[-*]\s+', linea):
            items = []
            while i < n and re.match(r'^\s*[-*]\s+', lineas[i]):
                items.append(re.sub(r'^\s*[-*]\s+', '', lineas[i].rstrip())); i += 1
            bloques.append(('ul', items))
            continue
        if re.match(r'^\s*\d+\.\s+', linea):
            items = []
            while i < n and re.match(r'^\s*\d+\.\s+', lineas[i]):
                items.append(re.sub(r'^\s*\d+\.\s+', '', lineas[i].rstrip())); i += 1
            bloques.append(('ol', items))
            continue
        parrafo = [linea]
        i += 1
        while i < n and lineas[i].strip() and not lineas[i].startswith('#') and \
                lineas[i].strip() != '---' and not lineas[i].strip().startswith('```') and \
                not lineas[i].strip().startswith('|') and \
                not re.match(r'^\s*[-*]\s+', lineas[i]) and \
                not re.match(r'^\s*\d+\.\s+', lineas[i]):
            parrafo.append(lineas[i].rstrip()); i += 1
        bloques.append(('p', ' '.join(parrafo)))
    return bloques


RX_CURSIVA_PARRAFO = re.compile(r'^\*(?!\*)(.+?)(?<!\*)\*$', re.S)


def cursiva_entera(texto):
    if not isinstance(texto, str):
        return texto, False
    m = RX_CURSIVA_PARRAFO.match(texto.strip())
    if m and '*' not in m.group(1):
        return m.group(1).strip(), True
    return texto, False


RX_INLINE = re.compile(r'(\*\*[^*]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*))')


def negrita_runs(texto):
    """(texto, negrita, cursiva) por tramo. 2026-08-29 (RT-08): la versión
    anterior sólo troceaba «**negrita**», así que el DOCX imprimía los
    asteriscos de la cursiva inline —«*take away*», «*Clostridium botulinum*»—
    mientras el PDF sí la convertía en <i>. El mismo Markdown tiene que dar el
    mismo documento en los dos formatos (§7-bis.15)."""
    runs = []
    for p in RX_INLINE.split(texto):
        if not p:
            continue
        if p.startswith('**') and p.endswith('**') and len(p) > 4:
            runs.append((p[2:-2], True, False))
        elif p.startswith('*') and p.endswith('*') and len(p) > 2:
            runs.append((p[1:-1], False, True))
        else:
            runs.append((p, False, False))
    return runs


def construir_docx(bloques, salida_docx, meta):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Pt, Cm, RGBColor

    doc = Document()
    est = doc.styles['Normal']
    est.font.name = 'Calibri'
    est.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.2)

    def runs(p, texto, negrita_base=False):
        for t, b, i in negrita_runs(texto):
            r = p.add_run(t)
            r.bold = b or negrita_base
            r.italic = i

    primero_h1 = True
    for tipo, contenido in bloques:
        if tipo == 'h1':
            if primero_h1:
                h = doc.add_heading(contenido, level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                primero_h1 = False
            else:
                doc.add_heading(contenido, level=1)
        elif tipo == 'h2':
            doc.add_heading(contenido, level=1)
        elif tipo == 'h3':
            doc.add_heading(contenido, level=2)
        elif tipo == 'h4':
            doc.add_heading(contenido, level=3)
        elif tipo == 'p':
            limpio, es_cursiva = cursiva_entera(contenido)
            p = doc.add_paragraph()
            if es_cursiva:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(limpio)
                r.italic = True
                r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            else:
                runs(p, contenido)
        elif tipo == 'ul':
            for it in contenido:
                runs(doc.add_paragraph(style='List Bullet'), it)
        elif tipo == 'ol':
            for it in contenido:
                runs(doc.add_paragraph(style='List Number'), it)
        elif tipo == 'code':
            p = doc.add_paragraph()
            r = p.add_run(contenido)
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
        elif tipo == 'table':
            filas = contenido
            ncols = len(filas[0])
            tabla = doc.add_table(rows=0, cols=ncols)
            tabla.style = 'Light Grid Accent 1'
            for fi, fila in enumerate(filas):
                row = tabla.add_row()
                for ci in range(ncols):
                    valor = fila[ci] if ci < len(fila) else ''
                    cell = row.cells[ci]
                    cell.text = ''
                    runs(cell.paragraphs[0], valor, negrita_base=(fi == 0))
            doc.add_paragraph()
        elif tipo == 'hr':
            # §7-bis.15: el salto de página es el MISMO en los dos formatos
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    cp = doc.core_properties
    cp.author = meta['author']
    cp.last_modified_by = meta['author']
    cp.title = meta['title']
    cp.subject = meta['subject']
    cp.category = meta.get('category', 'Guía profesional')
    cp.comments = meta.get('comments', '')
    # 2026-08-29 (RT-10): los tres DOCX salían con la fecha por defecto de
    # python-docx, 2013-12-23 23:15:00, y el gate de metadata sólo miraba
    # author y title, así que no lo veía nadie.
    ahora = datetime.datetime.now().replace(microsecond=0)
    cp.created = ahora
    cp.modified = ahora
    doc.save(salida_docx)
    return salida_docx


def _md_a_html(texto):
    texto = texto.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    texto = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', texto)
    texto = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<i>\1</i>', texto)
    return texto


def construir_pdf(bloques, salida_pdf, meta):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, ListFlowable, ListItem,
                                    Preformatted, PageBreak, KeepTogether)

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('H1c', parent=styles['Title'], fontName='Helvetica-Bold',
                        fontSize=24, leading=29, textColor=colors.HexColor(DARK),
                        spaceAfter=6, alignment=1)
    h2 = ParagraphStyle('H2c', parent=styles['Heading2'], fontName='Helvetica-Bold',
                        fontSize=16, leading=20, textColor=colors.HexColor(GOLD),
                        spaceBefore=12, spaceAfter=9)
    h3 = ParagraphStyle('H3c', parent=styles['Heading3'], fontName='Helvetica-Bold',
                        fontSize=12.5, leading=16, textColor=colors.HexColor(DARK),
                        spaceBefore=10, spaceAfter=6)
    h4 = ParagraphStyle('H4c', parent=h3, fontSize=11, leading=14,
                        textColor=colors.HexColor('#444444'))
    pst = ParagraphStyle('Pc', parent=styles['Normal'], fontName='Helvetica',
                         fontSize=10.2, leading=14.5, spaceAfter=8, alignment=4)
    li = ParagraphStyle('Lic', parent=pst, spaceAfter=3)
    firma = ParagraphStyle('Firma', parent=pst, fontName='Helvetica-Oblique',
                           alignment=1, textColor=colors.HexColor('#444444'),
                           spaceBefore=10)
    code = ParagraphStyle('Codec', parent=styles['Code'], fontName='Courier',
                          fontSize=9, leading=12,
                          backColor=colors.HexColor('#F5F5F5'),
                          borderPadding=6, spaceAfter=10)

    flujo = []
    primero_h1 = True
    for tipo, contenido in bloques:
        if tipo == 'h1':
            flujo.append(Paragraph(_md_a_html(contenido), h1 if primero_h1 else h2))
            primero_h1 = False
        elif tipo == 'h2':
            flujo.append(Paragraph(_md_a_html(contenido), h2))
        elif tipo == 'h3':
            flujo.append(Paragraph(_md_a_html(contenido), h3))
        elif tipo == 'h4':
            flujo.append(Paragraph(_md_a_html(contenido), h4))
        elif tipo == 'p':
            limpio, es_cursiva = cursiva_entera(contenido)
            flujo.append(Paragraph(_md_a_html(limpio), firma if es_cursiva else pst))
        elif tipo == 'ul':
            flujo.append(ListFlowable(
                [ListItem(Paragraph(_md_a_html(x), li), leftIndent=6) for x in contenido],
                bulletType='bullet', start='•', leftIndent=16, spaceAfter=8))
        elif tipo == 'ol':
            flujo.append(ListFlowable(
                [ListItem(Paragraph(_md_a_html(x), li), leftIndent=6) for x in contenido],
                bulletType='1', leftIndent=16, spaceAfter=8))
        elif tipo == 'code':
            flujo.append(Preformatted(contenido, code))
        elif tipo == 'table':
            filas = contenido
            ncols = max(len(f) for f in filas)
            fs = 8.6 if ncols <= 5 else (7.6 if ncols <= 7 else 6.8)
            data = [[Paragraph(_md_a_html(f[c] if c < len(f) else ''),
                               ParagraphStyle('Cellc', parent=pst, fontSize=fs,
                                              leading=fs + 2.4, alignment=0,
                                              spaceAfter=0,
                                              fontName='Helvetica-Bold' if fi == 0
                                              else 'Helvetica'))
                     for c in range(ncols)] for fi, f in enumerate(filas)]
            ancho = A4[0] - 4 * cm
            t = Table(data, colWidths=[ancho / ncols] * ncols, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(GOLD)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#DDDDDD')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                 [colors.white, colors.HexColor('#FAFAFA')]),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))
            flujo.append(t)
            flujo.append(Spacer(1, 10))
        elif tipo == 'hr':
            flujo.append(PageBreak())

    def cabecera_pie(canv, doc_):
        canv.saveState()
        ancho, alto = A4
        canv.setFont('Helvetica', 8)
        canv.setFillColor(colors.HexColor('#888888'))
        canv.drawString(2 * cm, alto - 1.3 * cm, meta['cabecera'])
        canv.drawRightString(ancho - 2 * cm, 1.2 * cm, f'Página {doc_.page}')
        canv.drawString(2 * cm, 1.2 * cm, meta['pie'])
        canv.setStrokeColor(colors.HexColor('#DDDDDD'))
        canv.line(2 * cm, alto - 1.5 * cm, ancho - 2 * cm, alto - 1.5 * cm)
        canv.restoreState()

    doc = SimpleDocTemplate(salida_pdf, pagesize=A4, topMargin=2.2 * cm,
                            bottomMargin=2 * cm, leftMargin=2 * cm,
                            rightMargin=2 * cm, title=meta['title'],
                            author=meta['author'], subject=meta['subject'],
                            creator=meta['author'])
    doc.build(flujo, onFirstPage=cabecera_pie, onLaterPages=cabecera_pie)
    return salida_pdf


def maquetar(md_text, base_sin_ext, meta):
    """Markdown → (docx, pdf). Los DOS salen del MISMO Markdown (§7-bis.15)."""
    bloques = parsear(md_text)
    antes = restos_no_winansi(bloques)
    bloques = sanear_bloques(bloques)
    quedan = restos_no_winansi(bloques)
    if quedan:
        detalle = ', '.join(f'{c!r} (U+{ord(c):04X}) x{n}' for c, n in sorted(quedan.items()))
        print('ABORTADO: caracteres que Helvetica pintaría como cuadrado negro '
              f'→ {detalle}. Añádelos a MAPA_WINANSI.', file=sys.stderr)
        raise SystemExit(3)
    docx = construir_docx(bloques, base_sin_ext + '.docx', meta)
    pdf = construir_pdf(bloques, base_sin_ext + '.pdf', meta)
    return docx, pdf, sum(antes.values())


# --------------------------------------------------------------------------
# 8. Gates §5.6 — medidos, no estimados
# --------------------------------------------------------------------------


def texto_pdf(path):
    import fitz
    doc = fitz.open(path)
    return doc, '\n'.join(p.get_text() for p in doc)


def texto_docx(path):
    from docx import Document
    d = Document(path)
    partes = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            partes += [c.text for c in row.cells]
    return d, '\n'.join(partes)


def palabras(t):
    return len([w for w in re.split(r'\s+', t) if w.strip()])


def palabras_reales(t):
    """Recuento robusto a la degeneración: un token sin ninguna letra no es
    una palabra. Con el recuento antiguo, tres páginas de «(2)» aprobaban el
    gate de extensión (RT-12)."""
    return len([w for w in re.split(r'\s+', t)
                if w.strip() and re.search(r'[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]', w)])


def valores_admitidos(xlsx_dir, idx_research, extra=()):
    """Toda cifra del texto tiene que existir en una celda de los xlsx del
    producto o en el research (§5.6.8). Se generan las variantes de formato
    con las que un redactor la escribiría."""
    vals = set()

    def mete(v):
        if v is None or isinstance(v, bool):
            return
        try:
            f = float(v)
        except (TypeError, ValueError):
            return
        if abs(f) < 1e-9:
            return
        for dec in (0, 1, 2):
            vals.add(num(f, dec))
            vals.add(num(round(f), 0))
            # el texto escribe «-19.116» como «19.116» detras de un signo: el
            # buscador de cifras no captura el menos
            vals.add(num(abs(f), dec))
            vals.add(num(round(abs(f)), 0))
        if 0 < abs(f) <= 1:
            for dec in (0, 1, 2):
                vals.add(num(f * 100, dec))
        if abs(f) >= 1000:
            vals.add(num(f / 1000, 0)); vals.add(num(f / 1000, 1))
            vals.add(num(f / 1000, 2))
        if abs(f) >= 1_000_000:
            vals.add(num(f / 1_000_000, 0)); vals.add(num(f / 1_000_000, 1))
            vals.add(num(f / 1_000_000, 2))

    import glob
    from openpyxl import load_workbook
    for ruta in sorted(glob.glob(os.path.join(xlsx_dir, '*.xlsx'))):
        w = load_workbook(ruta, data_only=True)
        for ws in w.worksheets:
            for row in ws.iter_rows():
                for c in row:
                    mete(c.value)
                    if isinstance(c.value, str):
                        for m in re.finditer(r'\d[\d.,]*', c.value):
                            vals.add(m.group(0))
    # Del research entra la cifra Y los numeros de su texto (cita literal,
    # nota, titulo de la fuente): son datos CON FUENTE, y el redactor los cita
    # («el dato cerrado de 2024 fue de 29.800 millones» sale de la nota de
    # SECT-03, no de la cabeza del modelo).
    for d in idx_research.values():
        mete(d.get('cifra'))
        for campo in ('cifra', 'cita_literal', 'nota', 'fuente_titulo', 'dato',
                      'unidad'):
            v = d.get(campo)
            if isinstance(v, str):
                for m in re.finditer(r'\d[\d.,]*', v):
                    vals.add(m.group(0))
                    vals.add(m.group(0).rstrip('.,'))
    for e in extra:
        vals.add(str(e))
        mete(e)
    return vals


RX_CIFRA_GRANDE = re.compile(r'\b\d{1,3}(?:\.\d{3})+(?:,\d+)?\b')


def coherencia_cifras(md_text, admitidos, ignorar=()):
    """Cada cifra con separador de miles del texto tiene que existir en una
    celda de los xlsx del producto o en el research (SPEC 5.6.8). Devuelve DOS
    listas: las de DINERO (con moneda pegada) son BLOQUEANTES —una cifra de
    inversion inventada es el defecto que esta v2.0 viene a corregir— y las
    demas quedan como aviso, porque una temperatura de color de 2.700 K o una
    superficie no son cifras de negocio."""
    euro, otras, vistas = [], [], set()
    for m in RX_CIFRA_GRANDE.finditer(md_text):
        t = m.group(0)
        if t in admitidos or t in ignorar or t in vistas:
            continue
        vistas.add(t)
        ini = max(0, m.start() - 80)
        ctx = md_text[ini:m.end() + 80].replace('\n', ' ')
        cerca = md_text[m.end():m.end() + 32].lower()
        es_dinero = ('\u20ac' in cerca or 'eur' in cerca or 'euro' in cerca
                     or 'millones' in cerca)
        (euro if es_dinero else otras).append({'cifra': t, 'contexto': ctx})
    return euro, otras


def gates(md_text, docx_path, pdf_path, cfg, xlsx_dir, idx_research,
          capitulos=()):
    doc_pdf, t_pdf = texto_pdf(pdf_path)
    d_docx, t_docx = texto_docx(docx_path)
    n_pdf, n_docx = palabras(t_pdf), palabras(t_docx)
    caps = re.findall(r'^## (\d+)\. (.+)$', md_text, re.M)
    # Las tablas se cuentan con el MISMO parser que maqueta (no con un regex
    # propio): un regex distinto contaba 33 donde el maquetador emitia 32 y la
    # paridad fallaba por una tabla que no existia en ningun formato.
    tablas_md = sum(1 for t, _ in parsear(md_text) if t == 'table')

    # palabras por capítulo
    trozos = re.split(r'^## ', md_text, flags=re.M)
    por_cap = {}
    for tr in trozos[1:]:
        m = re.match(r'(\d+)\. (.+)', tr.split('\n')[0])
        if m:
            # 2026-08-29 (RT-12): el cap. 15 declaraba 4.155 palabras de las
            # que ~2.700 eran «(2)». Se cuentan sólo los tokens con letras.
            por_cap[int(m.group(1))] = palabras_reales(tr)
    cortos = {k: v for k, v in por_cap.items() if v < cfg.get('min_palabras_cap', 900)}

    # tablas ancladas: ninguna tabla después del último ##
    ult = md_text.rfind('\n## ')
    cola = md_text[ult:] if ult > 0 else ''
    tabla_en_cola = bool(re.search(r'^\|[^\n]*\|\s*\n\|[\s:|-]+\|', cola, re.M))
    ult_tabla = md_text.rfind('\n| ')
    tabla_tras_ultimo_cap = ult_tabla > ult and ult > 0 and \
        'Sobre el autor' not in md_text[ult:ult + 80]

    # títulos de capítulo presentes en el PDF
    norm = lambda s: re.sub(r'\s+', ' ', s).strip().lower()
    t_pdf_n = norm(t_pdf)
    faltan_en_pdf = [t for _, t in caps if norm(t)[:38] not in t_pdf_n]

    admitidos = valores_admitidos(xlsx_dir, idx_research, cfg.get('cifras_extra', ()))
    incoherentes, avisos_cifras = coherencia_cifras(
        md_text, admitidos, cfg.get('cifras_ignorar', ()))

    # El PDF lleva cabecera y pie EN CADA PAGINA; el DOCX no. Comparar los dos
    # en bruto acusaba un 3,3 % de diferencia que es exactamente el mobiliario
    # de pagina, no contenido que falte. La paridad se mide sobre el CUERPO.
    palabras_marco = palabras(cfg.get('cabecera', '')) + \
        palabras(cfg.get('pie', '')) + 2          # «Pagina N»
    n_pdf_cuerpo = n_pdf - doc_pdf.page_count * palabras_marco

    # ---------------------------------------------------------------- CALIDAD
    # Los cinco detectores que faltaban (RT-06). Se miden sobre el .md Y sobre
    # el texto extraído del PDF: lo que se vende es el PDF, no el Markdown.
    perm = cfg.get('erratas_permitidas', ())
    deg = erratas_degeneracion(md_text) + erratas_degeneracion(t_pdf)
    trunc = erratas_truncamiento(md_text)
    parent = erratas_parentesis(md_text)
    orto = erratas_ortograficas(md_text, perm)
    entid = erratas_entidades(md_text) + erratas_entidades(t_pdf) + \
        erratas_entidades(t_docx)
    # El euro se mide SOLO sobre el Markdown: al extraer el texto de un PDF las
    # celdas de una tabla se pegan con un espacio («1.250,00 € 3.400,00 €») y
    # eso dispara el patrón de «€ delante de cifra» sin que haya defecto. En el
    # .md las celdas van separadas por «|» y no hay confusión posible.
    euro = erratas_euro(md_text)
    citas_hoja = erratas_citas_hoja(md_text)
    fuga_modelo = erratas_fuga_modelo(md_text) + erratas_fuga_modelo(t_pdf)
    epi = epigrafes_ausentes(md_text, capitulos)
    # asteriscos de Markdown que se hayan colado impresos (RT-08)
    asteriscos = t_docx.count('*') + t_pdf.count('*')

    r = {
        'paginas_pdf': doc_pdf.page_count,
        'paginas_prometidas': cfg['paginas_prometidas'],
        'palabras_md': palabras(md_text),
        'palabras_pdf': n_pdf,
        'palabras_docx': n_docx,
        'palabras_objetivo': cfg['palabras_objetivo'],
        'capitulos': len(caps),
        'tablas_md': tablas_md,
        'tablas_docx': len(d_docx.tables),
        'capitulos_cortos': cortos,
        'no_latinos_md': len(guard_no_latinos(md_text, 'md')),
        'no_latinos_pdf': len(RX_NO_LATINOS.findall(t_pdf)),
        'no_latinos_docx': len(RX_NO_LATINOS.findall(t_docx)),
        'fechas_caducas': erratas_fechas(md_text),
        'mortalidad_sin_fuente': erratas_mortalidad(
            md_text, cfg.get('mortalidad_permitida', [])),
        'fugas_de_taller': erratas_meta(md_text),
        'citas_en_sintaxis_de_prompt': len(re.findall(r'\.xlsx!|\[fuente:', md_text))
        + len(citas_hoja),
        'citas_hoja_celda': citas_hoja,
        'degeneracion': deg,
        'truncamientos': trunc,
        'parentesis_desequilibrados': parent,
        'erratas_ortograficas': orto,
        'erratas_ocurrencias': sum(d['veces'] for d in orto),
        'entidades_html': entid,
        'euro_mal_escrito': euro,
        'fugas_de_modelo': fuga_modelo,
        'epigrafes_ausentes': epi,
        'asteriscos_impresos': asteriscos,
        'palabras_reales_pdf': palabras_reales(t_pdf),
        'tabla_tras_ultimo_capitulo': bool(tabla_tras_ultimo_cap and tabla_en_cola),
        'titulos_ausentes_en_pdf': faltan_en_pdf,
        'cifras_no_encontradas': incoherentes,
        'cifras_aviso_no_dinero': avisos_cifras,
        'palabras_pdf_cuerpo': n_pdf_cuerpo,
        'meta_pdf': {'author': doc_pdf.metadata.get('author'),
                     'title': doc_pdf.metadata.get('title')},
        # A4 = 595 x 842 pt. Se mide PAGINA A PAGINA: una sola pagina en otro
        # tamano rompe la impresion del cliente y no se ve en el visor.
        'paginas_no_a4': [i + 1 for i, pg in enumerate(doc_pdf)
                          if not (abs(pg.rect.width - 595.276) < 2
                                  and abs(pg.rect.height - 841.89) < 2)],
    }
    d = __import__('docx').Document(docx_path).core_properties
    r['meta_docx'] = {'author': d.author, 'title': d.title,
                      'created': str(d.created)}
    dif = abs(n_pdf_cuerpo - n_docx) / max(1, max(n_pdf_cuerpo, n_docx))
    r['paridad_palabras_pct'] = round(dif * 100, 2)

    r['ok'] = {
        'paginas': doc_pdf.page_count >= cfg['paginas_prometidas'],
        'palabras_pdf': n_pdf >= cfg['palabras_objetivo'] * 0.95,
        'palabras_docx': n_docx >= cfg['palabras_objetivo'] * 0.95,
        'sin_capitulos_cortos': not cortos,
        'paridad': dif < 0.02 and len(d_docx.tables) == tablas_md,
        'tablas_ancladas': not r['tabla_tras_ultimo_capitulo'] and not faltan_en_pdf,
        'no_latinos': (r['no_latinos_md'] + r['no_latinos_pdf'] + r['no_latinos_docx']) == 0,
        'fechas': not r['fechas_caducas'],
        'mortalidad': not r['mortalidad_sin_fuente'],
        'sin_fugas_de_taller': not r['fugas_de_taller'],
        'citas_legibles': r['citas_en_sintaxis_de_prompt'] == 0,
        'a4': not r['paginas_no_a4'],
        'metadata': (r['meta_pdf']['author'] == 'AI Chef Pro'
                     and bool(r['meta_pdf']['title'])
                     and r['meta_docx']['author'] == 'AI Chef Pro'
                     and bool(r['meta_docx']['title'])
                     # RT-10: la fecha por defecto de python-docx no vale
                     and not r['meta_docx']['created'].startswith('2013-12-23')),
        'coherencia_cifras': not incoherentes,
        # --- calidad del texto (nuevos, 2026-08-29) -----------------------
        'sin_degeneracion': not deg,
        'sin_truncamientos': not trunc,
        'sin_erratas': not orto,
        'parentesis_equilibrados': not parent,
        'epigrafes_completos': not epi,
        'sin_entidades_html': not entid,
        'euro_bien_escrito': not euro,
        'sin_fugas_de_modelo': not fuga_modelo,
        'sin_asteriscos_impresos': asteriscos == 0,
    }
    r['verde'] = all(r['ok'].values())
    return r


# --------------------------------------------------------------------------
# 9. Orquestación por producto
# --------------------------------------------------------------------------


def cargar_guion(pid):
    ruta = os.path.join(AQUI, f'guion_{pid.replace("-", "_")}.py')
    if not os.path.exists(ruta):
        raise SystemExit(f'ABORTADO: no existe {ruta}')
    spec = importlib.util.spec_from_file_location(f'guion_{pid}', ruta)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def construir_documento(nombre, guia, capitulos, xlsx_dir, idx_research,
                        salida, dir_txt, forzar, cfg_extra=None):
    """Genera <nombre>.md/.docx/.pdf y devuelve (gates, detalle)."""
    detalle = []
    hojas = _hojas_conocidas(xlsx_dir)
    partes = [portada_e_indice(guia, capitulos)]
    for cap in capitulos:
        md, det = generar_capitulo(cap, guia, xlsx_dir, idx_research, dir_txt,
                                   forzar=(cap['n'] in forzar), hojas=hojas)
        f = guard_no_latinos(md, f'cap {cap["n"]}')
        if f:
            raise SystemExit(4)
        partes.append(md)
        partes.append('\n---\n')
        detalle.append(det)
    partes.append(cierre(guia))
    md_text = '\n'.join(partes)

    # 2026-08-29 (RT-03): las erratas por letra caída se corrigen sobre el
    # DOCUMENTO ENTERO, no bloque a bloque. La prueba que autoriza el cambio es
    # que la palabra correcta ya esté escrita bien en el mismo documento, y en
    # un bloque de 800 palabras esa prueba casi nunca se cumple: «rentabilidad»
    # sale 30 veces en el libro y una sola vez en el bloque que la escribió mal.
    md_text, cambios = reparar_erratas(md_text, ERRATAS_PERMITIDAS)
    if cambios:
        ERRATAS_REPARADAS.extend([dict(c, doc=nombre) for c in cambios])
        print(f'  letras repuestas en {nombre}: {len(cambios)}', flush=True)

    ruta_md = os.path.join(salida, nombre + '.md')
    with open(ruta_md, 'w', encoding='utf-8') as f:
        f.write(md_text)

    meta = {'author': 'AI Chef Pro',
            # 2026-09-04: la línea de MANUALES reutiliza el pipeline; «Guía profesional»
            # deja de ir a fuego (el guion puede fijar categoria_doc y tipo_doc).
            'category': guia.get('categoria_doc', 'Guía profesional'),
            'title': guia['titulo'],
            'subject': guia['subtitulo'] + f' · Versión {guia.get("version", "2.0")} · {guia.get("fecha", "agosto de 2026").replace(" de ", " ")}',
            'cabecera': guia['cabecera'],
            'pie': f'Versión {guia.get("version", "2.0")} · aichef.pro/{guia["pid"]}',
            'comments': 'AI Chef Pro · aichef.pro'}
    if cfg_extra and cfg_extra.get('meta'):
        meta.update(cfg_extra['meta'])
    docx_path, pdf_path, saneados = maquetar(md_text,
                                             os.path.join(salida, nombre), meta)
    cfg = dict(guia['gates'])
    cfg['cabecera'] = meta['cabecera']
    cfg['pie'] = meta['pie']
    if cfg_extra:
        cfg.update({k: v for k, v in cfg_extra.items() if k != 'meta'})
    g = gates(md_text, docx_path, pdf_path, cfg, xlsx_dir, idx_research,
              capitulos=capitulos)
    g['saneados_winansi'] = saneados
    g['ficheros'] = {'md': ruta_md, 'docx': docx_path, 'pdf': pdf_path}
    return g, detalle


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--producto', required=True)
    ap.add_argument('--salida', default=None)
    ap.add_argument('--json', default=None)
    ap.add_argument('--solo-capitulos', default=None,
                    help='1,2,3 — genera sólo esos capítulos (para depurar)')
    ap.add_argument('--regenerar', default='',
                    help='ns de capítulo cuyo caché se ignora (capítulos cortos)')
    ap.add_argument('--sin-bonus', action='store_true')
    ap.add_argument('--solo-bonus', default=None)
    ap.add_argument('--min-palabras-cap', type=int, default=900)
    ap.add_argument('--modelo', default=None,
                    help='slug de OpenRouter para TODOS los intentos (p. ej. anthropic/claude-sonnet-4.6)')
    args = ap.parse_args()
    if args.modelo:
        global MODELO, MODELO_FALLBACK
        MODELO = MODELO_FALLBACK = args.modelo

    pid = args.producto
    guion = cargar_guion(pid)
    guia = guion.GUIA
    guia['gates']['min_palabras_cap'] = args.min_palabras_cap
    xlsx_dir = os.path.join(DL, pid)
    scratch = os.environ.get(
        'GUIAS_SCRATCH',
        '/private/tmp/claude-501/-Users-johnguerrero-chefpro-modernize/'
        '9a6ebdb7-5d45-4ab1-9e93-10b86cb95c42/scratchpad/guias/docs/build')
    salida = guarda_salida(args.salida or os.path.join(scratch, pid))
    dir_txt = os.path.join(salida, 'txt')
    os.makedirs(dir_txt, exist_ok=True)

    global ERRATAS_PERMITIDAS, ERRATAS_FORZADAS
    ERRATAS_PERMITIDAS = tuple(guia['gates'].get('erratas_permitidas', ()))
    ERRATAS_FORZADAS = dict(guia['gates'].get('erratas_forzadas', {}))

    res = cargar_research()
    idx = indexar_research(res)
    forzar = {int(x) for x in args.regenerar.split(',') if x.strip()}

    capitulos = guion.CAPITULOS
    if args.solo_capitulos:
        pedidos = {int(x) for x in args.solo_capitulos.split(',')}
        capitulos = [c for c in capitulos if c['n'] in pedidos]

    informe = {'_meta': {'pid': pid, 'generado': time.strftime('%Y-%m-%d %H:%M'),
                         'salida': salida, 'xlsx_dir': xlsx_dir,
                         'modelo': MODELO,
                         'research': res.get('_meta', {}).get('generado'),
                         'research_ausente': res.get('_meta', {}).get('ausente', False)},
               'documentos': {}}

    print(f'== GUÍA {pid} — {len(capitulos)} capítulos', flush=True)
    g, det = construir_documento(pid, guia, capitulos, xlsx_dir, idx, salida,
                                 dir_txt, forzar)
    informe['documentos'][pid] = {'gates': g, 'capitulos': det}
    print(json.dumps(g['ok'], ensure_ascii=False, indent=2), flush=True)

    if not args.sin_bonus:
        for b in guion.BONUS:
            if args.solo_bonus and b['nombre'] != args.solo_bonus:
                continue
            print(f'== BONUS {b["nombre"]}', flush=True)
            gb = dict(guia)
            gb.update(b['guia'])
            dtxt = os.path.join(dir_txt, b['nombre'])
            os.makedirs(dtxt, exist_ok=True)
            gg, dd = construir_documento(b['nombre'], gb, b['capitulos'], xlsx_dir,
                                         idx, salida, dtxt, forzar,
                                         cfg_extra=b['gates'])
            informe['documentos'][b['nombre']] = {'gates': gg, 'capitulos': dd}
            print(json.dumps(gg['ok'], ensure_ascii=False, indent=2), flush=True)

    informe['_meta']['erratas_reparadas'] = ERRATAS_REPARADAS
    informe['_meta']['erratas_reparadas_n'] = len(ERRATAS_REPARADAS)
    print(f'letras repuestas: {len(ERRATAS_REPARADAS)}', flush=True)

    ruta_json = args.json or os.path.join(salida, f'{pid}-documentos.json')
    with open(ruta_json, 'w', encoding='utf-8') as f:
        json.dump(informe, f, ensure_ascii=False, indent=2)
    print(f'informe: {ruta_json}', flush=True)
    verde = all(d['gates']['verde'] for d in informe['documentos'].values())
    raise SystemExit(0 if verde else 1)


if __name__ == '__main__':
    main()
