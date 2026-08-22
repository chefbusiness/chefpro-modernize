#!/usr/bin/env python3
"""
bono_guia.py — Maqueta el bono «Guía: Controla tu Food Cost en 30 Días»
(Kit de Escandallos Pro v2.0) desde Markdown a DOCX y PDF.

    python3 bono_guia.py <entrada.md> <salida_pdf_sin_extension>
                         [<salida_docx_sin_extension>]

Genera <salida_pdf>.pdf (reportlab, A4) y el .docx en la segunda ruta si se
pasa (COM-B18: el .docx NO es un entregable —no tiene clave en PRODUCT_FILES
ni tarjeta en el dashboard—, así que no puede quedarse en la carpeta de
entrega ensuciando el censo). Sin segundo argumento, ambos van al mismo sitio.
Soporta: título (H1), H2/H3, párrafos con **negrita**, listas (- / 1.),
tablas Markdown → tablas reales, bloques ``` (código/fórmula) → monoespaciado.

No toca `astro-site/public/dl/kit-escandallos/`: sólo escribe donde se le diga.
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, ListFlowable, ListItem,
                                 Preformatted)
from reportlab.pdfgen import canvas as pdfcanvas

GOLD = '#B8860B'
DARK = '#1A1A1A'

# ==========================================================================
# COM-A01 — Saneado de caracteres (reportlab + Helvetica = WinAnsi/cp1252)
# ==========================================================================
# El bono salía con TODOS los caracteres fuera de WinAnsi convertidos en un
# cuadrado negro ■. No era cosmético: en la tabla de desviación ✅ y ❌
# colapsaban en el MISMO glifo (la columna que distingue el plato dentro de
# objetivo del que pierde dinero dejaba de distinguir nada) y las 29 casillas
# ☐ de los checklists semanales se veían como cuadrados YA MARCADOS.
# Tildes, ñ, €, «» y – SÍ están en WinAnsi y se dejan tal cual.
MAPA_WINANSI = {
    '\u2610': '[  ]',      # ☐ casilla vacía
    '\u2611': '[X]',       # ☑
    '\u2612': '[X]',       # ☒
    '\u2705': 'OK',        # ✅
    '\u2714': 'OK',        # ✔
    '\u274c': 'X',         # ❌
    '\u2716': 'X',         # ✖
    '\u2b50': '*',         # ⭐
    '\u2605': '*',         # ★
    # Los tres van seguidos de su propia palabra en la tabla de ingeniería de
    # menú («🐴 Caballo»): traducirlos la duplicaría («Caballo Caballo»), así
    # que se quitan y la palabra se queda sola.
    '\U0001F434': '',         # 🐴 Caballo
    '\U0001F9E9': '',         # 🧩 Puzzle
    '\U0001F436': '',         # 🐶 Perro
    '\u2192': '->',        # →
    '\u2190': '<-',        # ←
    '\u2500': '-',         # ─ (línea de cuadro)
    '\u2501': '-',
    '\u2212': '-',         # − (menos matemático)
    '\u00a0': ' ',         # espacio duro
    '\u202f': ' ',         # espacio fino
    '\u2011': '-',         # guion no separable
}


def sanear(texto):
    """Sustituye lo que Helvetica no sabe pintar. Lo que quede fuera de cp1252
    y no esté en el mapa se elimina (mejor un hueco que un ■)."""
    if not isinstance(texto, str):
        return texto
    fuera = []
    for ch in texto:
        if ch in MAPA_WINANSI:
            fuera.append(MAPA_WINANSI[ch])
            continue
        try:
            ch.encode('cp1252')
        except UnicodeEncodeError:
            continue                      # se descarta; el gate lo contaría
        fuera.append(ch)
    return ''.join(fuera)


def sanear_bloques(bloques):
    fuera = []
    for tipo, contenido in bloques:
        if isinstance(contenido, str):
            contenido = sanear(contenido)
        elif tipo == 'table':
            # en una celda nunca hace falta la sangría, y quitar un emoji deja
            # un espacio suelto delante
            contenido = [[re.sub(r'\s{2,}', ' ', sanear(c)).strip()
                          for c in fila] for fila in contenido]
        elif isinstance(contenido, list):
            contenido = [sanear(x) for x in contenido]
        fuera.append((tipo, contenido))
    return fuera


def restos_no_winansi(bloques):
    """Gate: qué caracteres seguirían saliendo como ■ si se maquetase esto."""
    malos = {}
    for _, contenido in bloques:
        if isinstance(contenido, str):
            trozos = [contenido]
        elif isinstance(contenido, list) and contenido and \
                isinstance(contenido[0], list):
            trozos = [c for fila in contenido for c in fila]
        elif isinstance(contenido, list):
            trozos = [x for x in contenido if isinstance(x, str)]
        else:
            continue
        for t in trozos:
            for ch in t:
                try:
                    ch.encode('cp1252')
                except UnicodeEncodeError:
                    malos[ch] = malos.get(ch, 0) + 1
    return malos


# ==========================================================================
# Parser Markdown → bloques
# ==========================================================================
def parsear(md_text):
    """Devuelve una lista de bloques tipados: h1, h2, h3, p, ul, ol, table,
    code, hr."""
    lineas = md_text.split('\n')
    bloques = []
    i, n = 0, len(lineas)
    while i < n:
        linea = lineas[i].rstrip()

        if not linea.strip():
            i += 1
            continue

        if linea.strip() == '---':
            bloques.append(('hr', None))
            i += 1
            continue

        if linea.startswith('# '):
            bloques.append(('h1', linea[2:].strip()))
            i += 1
            continue
        if linea.startswith('## '):
            bloques.append(('h2', linea[3:].strip()))
            i += 1
            continue
        if linea.startswith('### '):
            bloques.append(('h3', linea[4:].strip()))
            i += 1
            continue

        if linea.strip().startswith('```'):
            codigo = []
            i += 1
            while i < n and not lineas[i].strip().startswith('```'):
                codigo.append(lineas[i])
                i += 1
            i += 1  # salta la cerradura
            bloques.append(('code', '\n'.join(codigo)))
            continue

        # tabla: cabecera | separador --- | filas
        if linea.strip().startswith('|') and i + 1 < n and \
                re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lineas[i + 1]) and \
                '-' in lineas[i + 1]:
            filas = [linea]
            i += 2  # se salta la fila separadora
            while i < n and lineas[i].strip().startswith('|'):
                filas.append(lineas[i])
                i += 1
            tabla = []
            for f in filas:
                celdas = [c.strip() for c in f.strip().strip('|').split('|')]
                tabla.append(celdas)
            bloques.append(('table', tabla))
            continue

        if re.match(r'^\s*[-*]\s+', linea):
            items = []
            while i < n and re.match(r'^\s*[-*]\s+', lineas[i]):
                items.append(re.sub(r'^\s*[-*]\s+', '', lineas[i].rstrip()))
                i += 1
            bloques.append(('ul', items))
            continue

        if re.match(r'^\s*\d+\.\s+', linea):
            items = []
            while i < n and re.match(r'^\s*\d+\.\s+', lineas[i]):
                items.append(re.sub(r'^\s*\d+\.\s+', '', lineas[i].rstrip()))
                i += 1
            bloques.append(('ol', items))
            continue

        # párrafo: acumula líneas hasta línea en blanco / nuevo bloque
        parrafo = [linea]
        i += 1
        while i < n and lineas[i].strip() and \
                not lineas[i].startswith('#') and \
                lineas[i].strip() != '---' and \
                not lineas[i].strip().startswith('```') and \
                not lineas[i].strip().startswith('|') and \
                not re.match(r'^\s*[-*]\s+', lineas[i]) and \
                not re.match(r'^\s*\d+\.\s+', lineas[i]):
            parrafo.append(lineas[i].rstrip())
            i += 1
        bloques.append(('p', ' '.join(parrafo)))

    return bloques


RX_CURSIVA_PARRAFO = re.compile(r'^\*(?!\*)(.+?)(?<!\*)\*$', re.S)


def cursiva_entera(texto):
    """COM-M11: la firma del autor —la ÚLTIMA línea que lee el cliente y la que
    lleva la marca personal de John— salía con los asteriscos de Markdown
    literales, porque el conversor de énfasis sólo cubría `**negrita**`.
    Devuelve (texto_sin_asteriscos, True) si el párrafo entero va en cursiva."""
    if not isinstance(texto, str):
        return texto, False
    m = RX_CURSIVA_PARRAFO.match(texto.strip())
    if m and '*' not in m.group(1):
        return m.group(1).strip(), True
    return texto, False


def negrita_runs(texto):
    """Trocea 'texto **negrita** normal' en pares (texto, es_negrita)."""
    partes = re.split(r'(\*\*[^*]+\*\*)', texto)
    runs = []
    for p in partes:
        if not p:
            continue
        if p.startswith('**') and p.endswith('**'):
            runs.append((p[2:-2], True))
        else:
            runs.append((p, False))
    return runs


# ==========================================================================
# DOCX
# ==========================================================================
def construir_docx(bloques, salida_docx, titulo_meta):
    doc = Document()

    estilo_normal = doc.styles['Normal']
    estilo_normal.font.name = 'Calibri'
    estilo_normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    def escribir_runs(p, texto, negrita_base=False):
        for texto_run, es_negrita in negrita_runs(texto):
            r = p.add_run(texto_run)
            r.bold = es_negrita or negrita_base

    primero_h1 = True
    for tipo, contenido in bloques:
        if tipo == 'h1':
            if primero_h1:
                h = doc.add_heading(contenido, level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub = doc.add_paragraph(titulo_meta)
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub.runs[0].italic = True
                sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                doc.add_paragraph()
                primero_h1 = False
            else:
                doc.add_heading(contenido, level=1)
        elif tipo == 'h2':
            doc.add_heading(contenido, level=1)
        elif tipo == 'h3':
            doc.add_heading(contenido, level=2)
        elif tipo == 'p':
            limpio, es_cursiva = cursiva_entera(contenido)
            if es_cursiva:                                    # COM-M11
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(limpio)
                r.italic = True
                r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            elif contenido.strip().startswith('**') and \
                    contenido.strip().endswith('**') and \
                    contenido.count('**') == 2:
                # línea de metadatos en negrita (p.ej. "Bono del Kit...")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                escribir_runs(p, contenido)
            else:
                p = doc.add_paragraph()
                escribir_runs(p, contenido)
        elif tipo == 'ul':
            for it in contenido:
                p = doc.add_paragraph(style='List Bullet')
                escribir_runs(p, it)
        elif tipo == 'ol':
            for it in contenido:
                p = doc.add_paragraph(style='List Number')
                escribir_runs(p, it)
        elif tipo == 'code':
            p = doc.add_paragraph()
            r = p.add_run(contenido)
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        elif tipo == 'table':
            filas = contenido
            n_cols = len(filas[0])
            tabla = doc.add_table(rows=0, cols=n_cols)
            tabla.style = 'Light Grid Accent 1'
            for idx_fila, fila in enumerate(filas):
                row = tabla.add_row()
                for idx_col in range(n_cols):
                    valor = fila[idx_col] if idx_col < len(fila) else ''
                    cell = row.cells[idx_col]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    escribir_runs(p, valor, negrita_base=(idx_fila == 0))
        elif tipo == 'hr':
            doc.add_paragraph()

    doc.save(salida_docx)
    return salida_docx


# ==========================================================================
# PDF (reportlab)
# ==========================================================================
def _cabecera_pie(canv, doc_):
    canv.saveState()
    ancho, alto = A4
    canv.setFont('Helvetica', 8)
    canv.setFillColor(colors.HexColor('#888888'))
    canv.drawString(2 * cm, alto - 1.3 * cm, 'AI Chef Pro · Kit de Escandallos Pro')
    canv.drawRightString(ancho - 2 * cm, 1.2 * cm, f'Página {doc_.page}')
    canv.setStrokeColor(colors.HexColor('#DDDDDD'))
    canv.line(2 * cm, alto - 1.5 * cm, ancho - 2 * cm, alto - 1.5 * cm)
    canv.restoreState()


def _md_a_html(texto):
    """**negrita** → <b>negrita</b>, escapando & < > antes de reinsertar tags."""
    texto = texto.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    texto = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', texto)
    return texto


def construir_pdf(bloques, salida_pdf, titulo_meta):
    styles = getSampleStyleSheet()
    estilo_h1 = ParagraphStyle('H1c', parent=styles['Title'], fontName='Helvetica-Bold',
                               fontSize=22, leading=26, textColor=colors.HexColor(DARK),
                               spaceAfter=4, alignment=1)
    estilo_meta = ParagraphStyle('Meta', parent=styles['Normal'], fontName='Helvetica-Oblique',
                                 fontSize=11, textColor=colors.HexColor('#666666'),
                                 alignment=1, spaceAfter=16)
    estilo_h2 = ParagraphStyle('H2c', parent=styles['Heading2'], fontName='Helvetica-Bold',
                               fontSize=15, leading=19, textColor=colors.HexColor(GOLD),
                               spaceBefore=14, spaceAfter=8)
    estilo_h3 = ParagraphStyle('H3c', parent=styles['Heading3'], fontName='Helvetica-Bold',
                               fontSize=12.5, leading=16, textColor=colors.HexColor(DARK),
                               spaceBefore=10, spaceAfter=6)
    estilo_p = ParagraphStyle('Pc', parent=styles['Normal'], fontName='Helvetica',
                              fontSize=10.2, leading=14.5, spaceAfter=8,
                              alignment=4)  # justify
    estilo_li = ParagraphStyle('Lic', parent=estilo_p, spaceAfter=3)
    estilo_firma = ParagraphStyle('Firma', parent=estilo_p,
                                  fontName='Helvetica-Oblique', alignment=1,
                                  textColor=colors.HexColor('#444444'),
                                  spaceBefore=10)
    estilo_code = ParagraphStyle('Codec', parent=styles['Code'], fontName='Courier',
                                 fontSize=9, leading=12, backColor=colors.HexColor('#F5F5F5'),
                                 borderPadding=6, spaceAfter=10)

    flujo = []
    primero_h1 = True
    for tipo, contenido in bloques:
        if tipo == 'h1':
            if primero_h1:
                flujo.append(Paragraph(_md_a_html(contenido), estilo_h1))
                flujo.append(Paragraph(_md_a_html(titulo_meta), estilo_meta))
                primero_h1 = False
            else:
                flujo.append(Paragraph(_md_a_html(contenido), estilo_h2))
        elif tipo == 'h2':
            flujo.append(Paragraph(_md_a_html(contenido), estilo_h2))
        elif tipo == 'h3':
            flujo.append(Paragraph(_md_a_html(contenido), estilo_h3))
        elif tipo == 'p':
            limpio, es_cursiva = cursiva_entera(contenido)    # COM-M11
            flujo.append(Paragraph(_md_a_html(limpio),
                                   estilo_firma if es_cursiva else estilo_p))
        elif tipo == 'ul':
            items = [ListItem(Paragraph(_md_a_html(it), estilo_li), leftIndent=6)
                    for it in contenido]
            flujo.append(ListFlowable(items, bulletType='bullet', start='•',
                                      leftIndent=16, spaceAfter=8))
        elif tipo == 'ol':
            items = [ListItem(Paragraph(_md_a_html(it), estilo_li), leftIndent=6)
                    for it in contenido]
            flujo.append(ListFlowable(items, bulletType='1', leftIndent=16,
                                      spaceAfter=8))
        elif tipo == 'code':
            flujo.append(Preformatted(contenido, estilo_code))
        elif tipo == 'table':
            filas = contenido
            data = [[Paragraph(_md_a_html(c), ParagraphStyle(
                'Cellc', parent=estilo_p, fontSize=8.6, leading=11,
                alignment=0, spaceAfter=0,
                fontName='Helvetica-Bold' if fi == 0 else 'Helvetica'))
                    for c in fila] for fi, fila in enumerate(filas)]
            n_cols = len(filas[0])
            ancho_disp = A4[0] - 4 * cm
            col_w = ancho_disp / n_cols
            t = Table(data, colWidths=[col_w] * n_cols, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(GOLD)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#DDDDDD')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1),
                 [colors.white, colors.HexColor('#FAFAFA')]),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            flujo.append(t)
            flujo.append(Spacer(1, 10))
        elif tipo == 'hr':
            flujo.append(Spacer(1, 6))

    doc = SimpleDocTemplate(salida_pdf, pagesize=A4,
                            topMargin=2.2 * cm, bottomMargin=2 * cm,
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            title='Guía: Controla tu Food Cost en 30 Días')
    doc.build(flujo, onFirstPage=_cabecera_pie, onLaterPages=_cabecera_pie)
    return salida_pdf


def main():
    if len(sys.argv) not in (3, 4):
        print('uso: bono_guia.py <entrada.md> <salida_pdf_sin_extension> '
              '[<salida_docx_sin_extension>]')
        raise SystemExit(2)
    entrada, salida_base = sys.argv[1], sys.argv[2]
    # COM-B18: el .docx se escribe FUERA de la carpeta de entrega salvo que se
    # pida lo contrario.
    docx_base = sys.argv[3] if len(sys.argv) == 4 else salida_base
    with open(entrada, encoding='utf-8') as f:
        md_text = f.read()
    bloques = parsear(md_text)
    titulo_meta = 'Bono del Kit de Escandallos Pro de AI Chef Pro'
    for tipo, contenido in bloques:
        if tipo == 'p' and 'Bono del Kit' in contenido:
            titulo_meta = re.sub(r'\*\*', '', contenido)
            break

    # el bloque de metadatos "**Bono del Kit...**" ya se pinta bajo el H1;
    # si se deja también como párrafo normal sale duplicado.
    bloques_cuerpo = [b for b in bloques
                      if not (b[0] == 'p' and re.sub(r'\*\*', '', b[1]) == titulo_meta)]

    # COM-A01: se sanea ANTES de maquetar y se comprueba que no queda nada que
    # Helvetica pintaría como ■. El gate ABORTA: un cuadrado negro en la
    # columna que distingue «dentro de objetivo» de «pierde dinero» no es un
    # detalle tipográfico.
    antes = restos_no_winansi(bloques_cuerpo)
    bloques_cuerpo = sanear_bloques(bloques_cuerpo)
    titulo_meta = sanear(titulo_meta)
    quedan = restos_no_winansi(bloques_cuerpo)
    if quedan:
        detalle = ', '.join(f'{c!r} (U+{ord(c):04X}) ×{n}'
                            for c, n in sorted(quedan.items()))
        print('ABORTADO: caracteres que Helvetica no puede pintar y que '
              f'saldrían como ■ → {detalle}. Añádelos a MAPA_WINANSI.',
              file=sys.stderr)
        raise SystemExit(3)

    docx_path = construir_docx(bloques_cuerpo, docx_base + '.docx', titulo_meta)
    pdf_path = construir_pdf(bloques_cuerpo, salida_base + '.pdf', titulo_meta)
    n_tablas = sum(1 for t, _ in bloques if t == 'table')
    n_h2 = sum(1 for t, _ in bloques if t == 'h2')
    print(f'OK: {docx_path}')
    print(f'OK: {pdf_path}')
    print(f'bloques={len(bloques)} h2={n_h2} tablas={n_tablas} '
          f'saneados={sum(antes.values())} restos=0')


if __name__ == '__main__':
    main()
