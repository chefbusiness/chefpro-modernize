#!/usr/bin/env python3
"""
Generate "Cómo Montar un Restaurante Nikkei 60 Plazas" guide deliverables.
AI Chef Pro — aichef.pro
19 files: 1 DOCX guide + 8 Excel templates + 6 Excel checklists + 2 DOCX models + 1 PDF placeholder
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "public", "dl", "guia-restaurante-nikkei"
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

GOLD = "FFD700"
GOLD_RGB = RGBColor(0xFF, 0xD7, 0x00)
HEADER_BG = "2D2D2D"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"
MEDIUM_GRAY = "E0E0E0"
INPUT_COLOR = "E8F5E9"
BRAND = "AI Chef Pro · aichef.pro — Restaurante Nikkei"

title_font = Font(name="Calibri", size=16, bold=True, color=GOLD)
sub_font = Font(name="Calibri", size=11, color="888888", italic=True)
hdr_font = Font(name="Calibri", size=11, bold=True, color=WHITE)
sec_font = Font(name="Calibri", size=12, bold=True, color=GOLD)
dat_font = Font(name="Calibri", size=11)
bld_font = Font(name="Calibri", size=11, bold=True)
frm_font = Font(name="Calibri", size=11, color="1565C0", bold=True)
brn_font = Font(name="Calibri", size=9, color="888888", italic=True)

hdr_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
inp_fill = PatternFill(start_color=INPUT_COLOR, end_color=INPUT_COLOR, fill_type="solid")
lgt_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")
sec_fill = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")

thin_b = Border(
    left=Side(style="thin", color=MEDIUM_GRAY), right=Side(style="thin", color=MEDIUM_GRAY),
    top=Side(style="thin", color=MEDIUM_GRAY), bottom=Side(style="thin", color=MEDIUM_GRAY),
)
c_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
l_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
cur_fmt = '#,##0.00 €'
pct_fmt = '0.0%'


def shr(ws, row, ncols):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin_b; cell.alignment = c_align

def sdc(ws, r, c, v=None, font=None, fill=None, fmt=None, align=None):
    cell = ws.cell(row=r, column=c)
    if v is not None: cell.value = v
    cell.font = font or dat_font; cell.border = thin_b; cell.alignment = align or l_align
    if fill: cell.fill = fill
    if fmt: cell.number_format = fmt
    return cell

def title_block(ws, t, ncols=8):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    ws.cell(row=1, column=1, value=t).font = title_font
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
    ws.cell(row=2, column=1, value=BRAND).font = sub_font
    ws.row_dimensions[1].height = 30

def brand_footer(ws, row, ncols):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
    ws.cell(row=row, column=1, value=BRAND).font = brn_font
    ws.cell(row=row, column=1).alignment = c_align

def instr_sheet(wb, t, items):
    ws = wb.active; ws.title = "Instrucciones"; ws.sheet_properties.tabColor = "FFD700"
    title_block(ws, t, 3)
    r = 4
    ws.cell(row=r, column=1, value="Instrucciones de uso").font = sec_font; r += 1
    for i, x in enumerate(items, 1):
        ws.cell(row=r, column=1, value=f"{i}. {x}").font = dat_font; r += 1
    r += 1
    ws.cell(row=r, column=1, value="Celdas verdes = campos editables").font = dat_font
    ws.cell(row=r, column=1).fill = inp_fill
    ws.column_dimensions['A'].width = 80

def checklist_ws(ws, t, items):
    hdrs = ["#", "Categoría", "Tarea / Ítem", "Responsable", "Fecha Límite", "Estado", "Coste Est. (€)", "Notas"]
    title_block(ws, t)
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8); ws.freeze_panes = f"A{r+1}"
    ws.auto_filter.ref = f"A{r}:H{r}"
    dv = DataValidation(type="list", formula1='"Pendiente,En Curso,Completado"', allow_blank=True)
    ws.add_data_validation(dv)
    widths = [6, 22, 45, 18, 14, 14, 14, 30]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    for idx, (cat, tarea, resp, coste) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, tarea)
        sdc(ws, r, 4, resp); sdc(ws, r, 5, None, fill=inp_fill)
        sdc(ws, r, 6, "Pendiente", fill=inp_fill); dv.add(ws.cell(row=r, column=6))
        sdc(ws, r, 7, coste, fmt=cur_fmt, fill=inp_fill); sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1; brand_footer(ws, r, 8)


# ═══════════════════════════════════════════════════════════
# 1. GUIDE DOCX
# ═══════════════════════════════════════════════════════════
def gen_guide_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    def add_bullet(txt):
        doc.add_paragraph(txt, style='List Bullet')
    def add_num(txt):
        doc.add_paragraph(txt, style='List Number')
    def tip(txt):
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Consejo del Chef: {txt}")
        run.font.italic = True; run.font.color.rgb = GOLD_RGB

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cómo Montar un Restaurante Nikkei"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("60 Plazas de Aforo — Guía España 2026"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Chef John Guerrero\nAI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # CH1
    doc.add_heading("1. Qué es un Restaurante Nikkei Auténtico", level=1)
    doc.add_paragraph(
        "La cocina nikkei es la fusión viva entre la tradición culinaria peruana y la técnica japonesa, "
        "nacida en el Perú de finales del siglo XIX con la inmigración japonesa y consolidada hoy como "
        "una de las cocinas de fusión más respetadas del mundo. Referentes globales como Nobu Matsuhisa, "
        "Mitsuharu 'Micha' Tsumura (Maido, Lima), Albert Adrià (Pakta, Barcelona), Kikko (Madrid) u "
        "Osaka (Madrid/Barcelona) han llevado el nikkei al altar de la alta gastronomía internacional. "
        "Lejos del 'japonés buffet' genérico, un restaurante nikkei auténtico combina el corte samurái "
        "del itamae, los ajíes peruanos (amarillo, limo, rocoto), la leche de tigre como preparación "
        "madre y un respeto absoluto por el producto fresco."
    )
    doc.add_paragraph(
        "En España, el mercado nikkei vive un momento explosivo. Propuestas como Kikko, Osaka, Pakta, "
        "Totora o Maizal han educado al público urbano en tiraditos, ceviches nikkei, makis acevichados "
        "y anticuchos con miso. Un restaurante nikkei auténtico con cocinero con formación peruano-japonesa, "
        "pescado sashimi-grade diario, ajíes importados de Perú y una barra de pisco, sake y coctelería "
        "nikkei tiene un posicionamiento diferencial potentísimo y márgenes de bebida del 75-85%."
    )
    doc.add_heading("La diversidad nikkei: mucho más que sushi peruano", level=2)
    add_bullet("Tiraditos y ceviches nikkei: sashimi peruano bañado en leche de tigre con ají amarillo, ají limo o ponzu nikkei. Eje identitario del concepto.")
    add_bullet("Makis acevichados: uramaki bañado en salsa acevichada, crocante de camote, ají amarillo y cilantro. Plato bandera nikkei.")
    add_bullet("Causa nikkei: causa limeña (puré de papa amarilla + ají amarillo + limón) rellena de atún, salmón o pulpo estilo japonés.")
    add_bullet("Anticuchos nikkei: brochetas de corazón, pulpo o chita marinados con miso, soja y ají panca. Parrilla robata/josper.")
    add_bullet("Arroz chaufa y yakimeshi nikkei: wok station con mariscos, ají amarillo, sillao, jengibre y huevo tortilla.")
    add_bullet("Omakase nikkei (alta gama): menú degustación de 10-16 pases del cocinero. Ticket 90-160€. Referencia: Kikko, Osaka Madrid.")
    add_bullet("Robata/Josper: chita a la parrilla, anticuchos, pulpo al olivo nikkei. Parrilla vista, producto premium.")
    add_bullet("Ticket medio auténtico: 45-75€ (casual premium) o 90-160€ (omakase nikkei).")
    tip("La autenticidad nikkei no está en copiar Tokyo ni en copiar Lima — está en dominar la leche de tigre con ajíes peruanos Y la técnica del corte japonés. Un cocinero con formación en Maido, Pakta, Kikko u Osaka, ají amarillo real importado de Perú y pescado sashimi-grade diario de Mercamadrid/Mercabarna te sitúa por encima del 80% de los 'nikkei' genéricos en España.")
    doc.add_page_break()

    # CH2
    doc.add_heading("2. El Mercado de la Cocina Nikkei en España 2026", level=1)
    doc.add_paragraph(
        "La cocina nikkei es una de las categorías gastronómicas con mayor crecimiento en España. "
        "Desde la apertura de Pakta (Albert Adrià, Barcelona) en 2013 y la llegada del grupo Osaka "
        "(Madrid, Barcelona, Marbella) y Kikko (Madrid), el público español ha descubierto un universo "
        "de tiraditos, ceviches nikkei y makis acevichados muy distinto del sushi convencional. La "
        "demanda crece a doble dígito anual y el ticket medio supera al japonés tradicional."
    )
    doc.add_heading("Datos clave del sector", level=2)
    add_bullet("Restaurantes nikkei auténticos en España: ~120 (2025), con explosión en los últimos 3 años.")
    add_bullet("Ticket medio: 45-70€ (nikkei casual) vs 90-160€ (omakase nikkei / alta gama).")
    add_bullet("Ciudades con mayor demanda: Madrid, Barcelona, Valencia, Málaga, Marbella, Palma, San Sebastián, Bilbao.")
    add_bullet("Público: 30-55 años, urbano, foodie, alto poder adquisitivo, viajero latinoamericano/asiático.")
    add_bullet("Tendencia clave: Pisco Sour, Chilcano y coctelería nikkei en auge paralelo al mezcal y al sake.")
    add_bullet("Referentes que mueven el mercado: Maido Lima (mejor restaurante de Latinoamérica), Pakta Barcelona, Osaka, Kikko, Totora.")
    add_bullet("Delivery nikkei: makis acevichados, tiraditos y causas bowl viajan muy bien. Arroz chaufa también.")
    tip("El 80% de los 'nikkei' en España son en realidad japoneses o peruanos que añaden un par de platos fusión a la carta. Un restaurante nikkei 100% auténtico — con cocinero formado en Perú o en referentes nikkei, ajíes reales importados, leche de tigre madre y parrilla robata — sigue siendo un nicho casi vacío con demanda enorme.")
    doc.add_page_break()

    # CH3
    doc.add_heading("3. Modelos de Negocio Nikkei", level=1)
    doc.add_paragraph("El restaurante nikkei admite varios modelos bien diferenciados. La elección determina la inversión, la operativa, el perfil del equipo y tu posicionamiento.")
    doc.add_heading("Modelo 1: Cevichería Nikkei / Barra de Tiraditos", level=2)
    doc.add_paragraph("Especializado en tiraditos, ceviches nikkei, causas y sashimis peruanizados. Barra de tiraditos vista donde el cocinero trabaja el pescado con leches de tigre al momento. Ticket 45-70€. Requiere cocinero con dominio de corte japonés y recetas peruanas, pescado sashimi-grade diario y batería de ajíes frescos.")
    doc.add_heading("Modelo 2: Nikkei Casual Premium (Carta Amplia)", level=2)
    doc.add_paragraph("Carta amplia 35-55 platos: tiraditos, ceviches, makis acevichados, causas, anticuchos, chaufa, chita a la parrilla. Ambiente contemporáneo. Ticket 50-75€. Modelo más replicable y estable en España, referencia Osaka o Kikko.")
    doc.add_heading("Modelo 3: Omakase Nikkei / Alta Gama", level=2)
    doc.add_paragraph("Barra de 10-16 asientos, menú degustación de 10-16 pases del cocinero. Ticket 90-160€. Requiere cocinero de máximo nivel (ideal con paso por Maido, Pakta, Osaka, Kikko), pescado premium y ajíes peruanos de calidad excepcional. Referencia: Pakta (Barcelona).")
    doc.add_heading("Modelo 4: Nikkei con Parrilla Robata / Josper", level=2)
    doc.add_paragraph("Parrilla robata o Josper a la vista con anticuchos, chita entera, pulpo al olivo, chuletón de wagyu con chimichurri nikkei. Ticket 60-90€. Inversión alta en extracción y parrilla, pero muy experiencial.")
    doc.add_heading("Modelo 5: Nikkei + Pisco Bar", level=2)
    doc.add_paragraph("Concepto híbrido cocina + coctelería. Barra de coctelería nikkei con Pisco Sour, Chilcano, Maracuyá Sour, Yuzu Sour como protagonista, maridada con tiraditos, causas y makis. Ticket 45-65€ con 35% del ticket en bebida. Modelo muy rentable.")
    doc.add_heading("Modelo 6: Dark Kitchen Nikkei", level=2)
    doc.add_paragraph("Solo delivery: makis acevichados, tiradito bowls, causa bowls, chaufa de mariscos, anticuchos con arroz. Sin sala. Inversión 50-90K€. Ideal para testear mercado antes de abrir local físico.")
    doc.add_page_break()

    # CH4
    doc.add_heading("4. Estudio de Viabilidad y Plan Financiero", level=1)
    doc.add_paragraph(
        "La inversión total para un restaurante nikkei de 60 plazas en España oscila "
        "entre 280.000€ y 520.000€ dependiendo de la ubicación, nivel de acabados, "
        "si incluyes barra de sushi tradicional y el grado de equipamiento especializado "
        "nikkei (barra de tiraditos, licuadoras industriales, parrilla robata/josper, wok station, abatidor)."
    )
    doc.add_heading("Desglose de inversión típica", level=2)
    add_bullet("Obra civil y acondicionamiento: 70.000-150.000€")
    add_bullet("Equipamiento cocina nikkei (barra tiraditos, licuadoras, robata/josper, wok station, abatidor, cuchillos japoneses): 55.000-110.000€")
    add_bullet("Barra de tiraditos artesanal + vitrina refrigerada sashimi: 15.000-35.000€")
    add_bullet("Barra de pisco, sake y coctelería nikkei (instalación + stock inicial 40-60 refs): 10.000-20.000€")
    add_bullet("Mobiliario sala + decoración nikkei contemporánea (madera, ajíes secos, cerámica peruano-japonesa): 25.000-55.000€")
    add_bullet("Licencias y permisos (incluye extracción reforzada robata, registro sanitario pescado crudo): 6.000-14.000€")
    add_bullet("Marketing pre-apertura: 5.000-12.000€")
    add_bullet("Stock inicial (pescado sashimi-grade + ajíes peruanos + arroz Koshihikari + pisco + sake): 12.000-25.000€")
    add_bullet("Fondo de maniobra (3 meses): 45.000-95.000€")
    add_bullet("Tecnología (TPV, web, delivery, pantallas cocina): 4.000-10.000€")
    doc.add_heading("Ratios financieros objetivo", level=2)
    add_bullet("Food cost: 30-35% sobre ventas (pescado premium + importación de ajíes peruanos sube ligeramente el coste)")
    add_bullet("Coste de personal: 30-38% sobre ventas")
    add_bullet("Alquiler: máximo 8-10% sobre ventas")
    add_bullet("EBITDA objetivo: 10-18%")
    add_bullet("Break-even: mes 10-16 (depende de ubicación, ticket y tipo de modelo)")
    tip("El margen en pisco, coctelería nikkei y sake compensa el food cost del pescado fresco. Un Pisco Sour cuesta 1.80€ y lo vendes a 11-14€. Una copa de sake premium cuesta 4€ y se vende a 13-18€. La barra nikkei (pisco + sake + coctelería) puede representar el 25-35% de tu facturación con márgenes del 75-85%.")
    doc.add_page_break()

    # CH5
    doc.add_heading("5. Requisitos Legales en España + Normativa Anisakis", level=1)
    doc.add_paragraph("Abrir un restaurante nikkei en España requiere los mismos trámites que cualquier restaurante, MÁS una atención especial a la normativa de congelación preventiva para pescado consumido crudo (anisakis), trazabilidad de pescado y, si importas productos japoneses directamente, trámites de comercio exterior.")
    doc.add_heading("Trámites generales", level=2)
    add_num("Constitución de sociedad (SL) o alta como autónomo")
    add_num("Licencia de actividad / declaración responsable (según municipio)")
    add_num("Registro sanitario autonómico")
    add_num("Alta en Hacienda: IAE epígrafe 671.4 o similar")
    add_num("Inscripción Seguridad Social empresa")
    add_num("Seguro de responsabilidad civil (mínimo 600.000€ por riesgo sanitario del pescado crudo)")
    add_num("Plan APPCC con PCC específico para pescado crudo (sashimi/sushi)")
    add_num("Registro RGSEAA")
    add_num("Licencia de terraza (si aplica)")
    add_num("Licencia de música / SGAE")
    add_num("Libro de reclamaciones")
    doc.add_heading("Normativa ANISAKIS (OBLIGATORIA — crítica para sushi y sashimi)", level=2)
    add_bullet("El Real Decreto 1420/2006 OBLIGA a congelar preventivamente todo pescado destinado a ser consumido crudo, marinado o poco hecho a -20°C durante al menos 24 horas (o -35°C durante 15 h).")
    add_bullet("Aplica a: sashimi, sushi, nigiri, tataki, marinados en cítricos o vinagre.")
    add_bullet("El congelador debe ser específico y alcanzar la temperatura certificada. Conservar registros de congelación por lote.")
    add_bullet("Excepciones: pescado de acuicultura alimentado en condiciones controladas SIN riesgo de anisakis puede venderse fresco (requiere certificado del proveedor).")
    add_bullet("Información al cliente: obligación de informar en carta que el pescado ha sido congelado preventivamente según normativa.")
    add_bullet("Sanciones: multas de 5.000-60.000€ por incumplimiento + cierre cautelar en caso de caso positivo de anisakiasis.")
    doc.add_heading("Requisitos específicos de importación de Perú y Japón", level=2)
    add_bullet("Registro como operador de comercio exterior (EORI) si importas directamente")
    add_bullet("Certificado SOIVRE para productos alimentarios importados de fuera de la UE (ajíes peruanos secos, pasta de ají amarillo, quinoa)")
    add_bullet("Etiquetado en español obligatorio para productos importados (sake, salsa soja, pasta de ají amarillo, ají panca, pisco)")
    add_bullet("Impuestos especiales para bebidas alcohólicas (pisco, sake, shochu): IIEE")
    add_bullet("Alternativa más sencilla: comprar a distribuidores en España (Inkawasi, Latin Products, Perú Market para ajíes; Japan Sushi Express, Comercial Minamoto para producto japonés; Viñas de Oro/Pisco 1615 para pisco)")
    tip("La normativa anisakis NO es negociable. Compra un congelador certificado que alcance -20°C real, documenta cada lote, e incluye en el plan APPCC un PCC (Punto de Control Crítico) específico para el pescado crudo. Guarda registros al menos 12 meses. Un caso positivo de anisakiasis puede cerrar tu restaurante y generar responsabilidad civil millonaria.")
    doc.add_page_break()

    # CH6
    doc.add_heading("6. APPCC y Seguridad Alimentaria para Cocina Nikkei", level=1)
    doc.add_paragraph("El plan APPCC de un restaurante nikkei tiene las mismas bases que cualquier restaurante, pero con puntos de control críticos (PCC) adicionales para sashimi/sushi (pescado crudo), arroz sushi (temperatura y acidificación), y productos importados.")
    doc.add_heading("Prerrequisitos estándar", level=2)
    add_bullet("Plan de limpieza y desinfección (frecuencias, productos, responsables)")
    add_bullet("Control de temperaturas (cámaras, servicio, transporte)")
    add_bullet("Plan de control de plagas (DDD)")
    add_bullet("Trazabilidad de materias primas (proveedores, lotes, fechas, documentación de congelación del pescado)")
    add_bullet("Plan de formación del personal (manipulador de alimentos, específico pescado crudo)")
    add_bullet("Control de agua potable")
    add_bullet("Plan de gestión de residuos y aceite usado (tempura)")
    add_bullet("Plan de alérgenos (14 alérgenos + especial atención a soja, sésamo, crustáceos, gluten)")
    doc.add_heading("Puntos críticos específicos de cocina nikkei", level=2)
    add_bullet("Pescado para tiradito/sashimi/ceviche nikkei: CONGELACIÓN PREVIA obligatoria -20°C/24h para eliminar anisakis. NO es opcional. Documentación de cada lote.")
    add_bullet("Leche de tigre: aunque el pH bajo por el limón actúa como barrera parcial, el pescado base debe haber pasado por el protocolo de congelación. Conservación máxima 24h a 0-4°C.")
    add_bullet("Arroz de sushi para makis acevichados: tras cocer y aderezar con sushi-zu, mantener en ohitsu o recipiente cerrado entre 20-25°C máximo 4 horas. Nunca refrigerar.")
    add_bullet("Descongelación de pescado: SIEMPRE en cámara (0-4°C) durante la noche. Nunca a temperatura ambiente ni con agua.")
    add_bullet("Separación estricta de utensilios para pescado crudo: cuchillos yanagiba exclusivos, tablas dedicadas, superficies limpias.")
    add_bullet("Vitrinas refrigeradas de pescado: mantenimiento constante 0-4°C, rotación máxima de pescado expuesto cada 2-3 horas.")
    add_bullet("Ajíes frescos (amarillo, limo, rocoto): lavado riguroso, cadena de frío, separación de otros productos por riesgo de contaminación cruzada con manos.")
    add_bullet("Anticuchos de corazón: trazabilidad del corazón de ternera, cocción a temperatura superior a 75°C en núcleo.")
    add_bullet("Wok y parrilla robata: temperaturas altas. Control de aceite del wok (filtrado diario), limpieza diaria de parrilla.")
    tip("La leche de tigre es tu preparación madre nikkei. Prepárala al momento en la mañana, refrigera en recipientes estériles y descarta sobrantes a 24h máximo. Registra temperatura cada 2 horas. La combinación de pescado crudo + ají fresco + manipulación directa es un triple riesgo que exige disciplina APPCC extrema.")
    doc.add_page_break()

    # CH7
    doc.add_heading("7. Ubicación y Local", level=1)
    doc.add_paragraph("La ubicación de un restaurante nikkei es crucial. El alto ticket medio, la dependencia de pescado fresco diario y la necesidad de un público sofisticado acotan las zonas viables a barrios urbanos premium y zonas gastronómicas consolidadas.")
    doc.add_heading("Criterios de selección", level=2)
    add_bullet("Tráfico peatonal: mínimo 300 personas/hora en horario punta, especialmente público 28-55 años")
    add_bullet("Visibilidad: fachada contemporánea con identidad nikkei (madera natural, ajíes secos decorativos, tipografía peruano-japonesa)")
    add_bullet("Metros cuadrados: 140-200 m² para 60 plazas (ratio 2.3-3.3 m²/plaza)")
    add_bullet("Zona: gastronómica premium + afterwork + público foodie. Muy mal en zonas populares de bajo ticket.")
    add_bullet("Competencia: analiza los 500m alrededor — ideal si hay propuestas latinoamericanas premium o japonesas cerca")
    add_bullet("Salida de humos: imprescindible y reforzada. La parrilla robata/josper y el wok station generan mucho humo.")
    add_bullet("Proximidad a mercados de abastos: ventaja enorme para pescado sashimi-grade diario (Mercamadrid, Mercabarna, lonjas)")
    add_bullet("Accesibilidad: acceso para personas con movilidad reducida (obligatorio por ley)")
    add_bullet("Barrios donde funciona muy bien: Chamberí/Chamartín/Salamanca (Madrid), Eixample/Sarrià (Barcelona), Ruzafa (Valencia), Indautxu (Bilbao), Ensanche (San Sebastián), Casco Antiguo Marbella")
    tip("El nikkei premium depende menos del tráfico peatonal y más de la reputación y las reservas. Un omakase nikkei en una calle secundaria con 14 plazas y lista de espera factura más por m² que un genérico en zona comercial. Prioriza zonas donde tu público objetivo (foodies 30-55 años con alto poder adquisitivo) vive y come.")
    doc.add_page_break()

    # CH8
    doc.add_heading("8. Diseño de Cocina Nikkei Profesional", level=1)
    doc.add_paragraph("La cocina de un restaurante nikkei tiene estaciones muy especializadas que no existen en una cocina convencional europea: la barra de tiraditos del cocinero, la estación de leches de tigre con licuadoras industriales, la parrilla robata/josper para anticuchos, la wok station para arroz chaufa y la zona de arroz sushi para makis acevichados.")
    doc.add_heading("Zonas de la cocina nikkei", level=2)
    add_bullet("Barra de tiraditos (sushi/tiradito bar): zona de trabajo del cocinero nikkei, visible al cliente. Mesa refrigerada para pescado, tabla de corte grande, cuchillos japoneses yanagiba/deba. Batería de ajíes frescos al alcance.")
    add_bullet("Estación de leches de tigre: licuadoras industriales (2-3 uds de alta potencia), batidoras de inmersión, recipientes estériles, báscula de precisión, ajíes en fresco. Leche de tigre como preparación madre.")
    add_bullet("Estación de arroz sushi (shari) para makis acevichados: arrocera industrial profesional (5-10L), recipiente de cedro (ohitsu), abanico. Aderezo con sushi-zu.")
    add_bullet("Parrilla robata / josper: parrilla de carbón binchotan o Josper para anticuchos, chita entera, pulpo al olivo. Requiere extracción extrema.")
    add_bullet("Wok station: 2-3 woks profesionales alta potencia para arroz chaufa, yakimeshi nikkei, saltados de mariscos con ají amarillo.")
    add_bullet("Plancha teppanyaki (opcional): para tatakis y sellados finales.")
    add_bullet("Abatidor de temperatura: CRÍTICO para enfriamiento rápido de pescado y cumplimiento del protocolo anisakis (-20°C/24h).")
    add_bullet("Zona fría / preparación: corte de verduras, choclo, camote, yuca, preparación de salsas (acevichada, ponzu nikkei, nikkei dressing, ají amarillo), mise en place.")
    add_bullet("Ahumador pequeño (opcional): para tatakis de atún y sabores con toque ahumado.")
    add_bullet("Almacenamiento: cámaras frigoríficas (2) con zona específica para pescado sashimi-grade, cámara de congelación CERTIFICADA -20°C, economato seco para arroz Koshihikari, ajíes secos, algas, salsas.")
    add_bullet("Zona de lavado: tren de lavado, zona de limpieza.")
    add_bullet("Superficie mínima cocina: 55-70 m² (ratio cocina/sala 1:2.5 — la cocina nikkei necesita más espacio)")
    tip("La barra de tiraditos a la vista del cliente es tu mayor herramienta de marketing. Ver al cocinero laminando pescado al momento, bañándolo con leche de tigre frente al cliente, decorándolo con ají limo y choclo crujiente — es teatro en vivo y muy fotogénico. Invierte en una barra profesional con vitrina refrigerada delantera y deja que tu cocinero sea el protagonista.")
    doc.add_page_break()

    # CH9
    doc.add_heading("9. Equipamiento Específico de Cocina Nikkei", level=1)
    doc.add_paragraph("Además del equipamiento estándar de cocina profesional, un restaurante nikkei necesita equipos específicos: licuadoras industriales para leches de tigre, parrilla robata o Josper, woks de alta potencia, abatidor certificado y cuchillería japonesa.")
    doc.add_heading("Equipamiento específico nikkei con costes orientativos", level=2)
    add_bullet("Arrocera industrial profesional (5-10 litros, ideal japonesa): 1.800-4.500€")
    add_bullet("Vitrina refrigerada de pescado/tiraditos (barra): 4.500-9.000€")
    add_bullet("Licuadoras industriales alta potencia para leches de tigre (2-3 uds): 1.800-3.500€")
    add_bullet("Batidora de inmersión profesional: 300-700€")
    add_bullet("Parrilla robata de carbón binchotan o Josper: 4.000-12.000€")
    add_bullet("Wok station con 2 woks alta potencia (chaufa, yakimeshi): 2.500-5.500€")
    add_bullet("Plancha teppanyaki profesional (opcional): 3.000-7.000€")
    add_bullet("Abatidor de temperatura certificado (crítico para anisakis): 3.500-7.000€")
    add_bullet("Cuchillos japoneses profesionales (yanagiba, deba, usuba): 1.500-4.500€ (set completo)")
    add_bullet("Tabla de corte madera / sintética para sushi: 300-800€")
    add_bullet("Ahumador pequeño (opcional para tatakis): 800-1.800€")
    add_bullet("Vitrina refrigerada de pisco, sake y vinos: 3.000-6.000€")
    doc.add_heading("Equipamiento estándar de cocina", level=2)
    add_bullet("Cocina de gas 6 fuegos + horno: 3.500-6.000€")
    add_bullet("Freidora doble (crocante de camote, tempura ligera, chicharrones): 1.500-3.000€")
    add_bullet("Horno mixto (Rational / Unox): 8.000-15.000€")
    add_bullet("Cámaras frigoríficas (2 puertas): 2.500-4.000€")
    add_bullet("Cámara congelación CERTIFICADA -20°C (imprescindible para anisakis): 3.000-5.500€")
    add_bullet("Tren de lavado: 4.000-8.000€")
    add_bullet("Campana extractora REFORZADA + filtros (robata/wok generan humo extremo): 6.000-12.000€")
    add_bullet("Total equipamiento cocina: 55.000-110.000€")
    tip("La parrilla robata con carbón binchotan (o un Josper) es el elemento más espectacular de un restaurante nikkei premium — pero también el más exigente en extracción. Alcanzan 900-1.000°C y generan humo intenso. Invierte en extracción sobredimensionada (mínimo 12.000 m³/h) y filtros de grado profesional. Una instalación deficiente te cerrará el restaurante por quejas vecinales.")
    doc.add_page_break()

    # CH10
    doc.add_heading("10. Diseño de Sala y Decoración Nikkei Contemporánea", level=1)
    doc.add_paragraph("La decoración de un restaurante nikkei bien resuelto combina el minimalismo japonés (madera, líneas limpias, luz tenue) con los acentos cálidos del Perú contemporáneo (texturas de textil andino, tonos tierra, cerámica artesanal). Referencias de estilo: Pakta (Barcelona), Osaka, Kikko, Maido (Lima).")
    doc.add_heading("Distribución tipo (60 plazas)", level=2)
    add_bullet("Barra de tiraditos/sushi: 10-14 plazas frente al cocinero — lo más prestigioso, plato a plato en vivo")
    add_bullet("Mesas de 2: 5-7 unidades (10-14 plazas) — parejas, cenas íntimas")
    add_bullet("Mesas de 4: 6-8 unidades (24-32 plazas) — grupos, familias")
    add_bullet("Zona privada (si aplica): 4-8 plazas — experiencia premium / omakase")
    add_bullet("Distancia entre mesas: 70-90 cm — el público nikkei premium valora el confort y la conversación")
    doc.add_heading("Elementos clave de la decoración nikkei contemporánea", level=2)
    add_bullet("Madera natural: mesas, barra y detalles en madera de roble, haya o algarrobo. Base estética nikkei.")
    add_bullet("Texturas peruanas sutiles: textiles andinos en acentos muy puntuales, cerámica artesanal peruana y japonesa mezclada.")
    add_bullet("Elementos peruanos visuales: cuelgas de ajíes secos (amarillo, panca, rocoto) como elemento decorativo auténtico.")
    add_bullet("Colores: paleta neutra de madera natural, blanco roto, negro mate, tierra cálida. Acentos en cobre, terracota o índigo.")
    add_bullet("Iluminación: tenue, cálida (2400-2700K), lámparas de papel o fibras naturales, puntos de luz dirigida a la barra.")
    add_bullet("Vajilla: cerámica artesanal de estilo mixto — platos de pizarra para tiraditos, cuencos de barro para causas, platos hondos para ceviches.")
    add_bullet("Arte de pared: fotografías de pescadores peruanos, mercados de Lima, cocineros en faena. Evitar clichés de banderas.")
    add_bullet("Música: bossa nova, cumbia psicodélica peruana (Los Mirlos), jazz latino, ambient. Volumen bajo durante el servicio.")
    tip("El público español confunde fácilmente nikkei con 'japonés' o con 'peruano'. Tu decoración debe transmitir inmediatamente la fusión — una pared con cuelga de ajíes peruanos secos, mesas de madera natural de corte japonés, cerámica mixta y una barra de tiraditos vista dejan claro el concepto sin palabras. Evita cualquier cliché (kimonos en la pared, banderas peruanas de plástico, farolillos).")
    doc.add_page_break()

    # CH11
    doc.add_heading("11. Barra de Pisco, Sake y Coctelería Nikkei (40-60 Referencias)", level=1)
    doc.add_paragraph("La barra de pisco, sake y coctelería nikkei es el elemento diferencial que eleva un restaurante nikkei de casual a experiencial. El pisco peruano (puro, acholado, aromático), el sake premium y la coctelería con Pisco Sour, Chilcano y versiones nikkei de sours con yuzu son productos de muy alto margen.")
    doc.add_heading("Estructura de la barra (40-60 referencias)", level=2)
    add_bullet("Pisco peruano puro (6-10 refs): Quebranta, Italia, Torontel, Moscatel, Negra Criolla. Referencias clave: Viñas de Oro, Pisco 1615, Barsol, Waqar.")
    add_bullet("Pisco acholado (3-5 refs): mezclas aromáticas de varias uvas. Muy versátiles en coctelería.")
    add_bullet("Sake Junmai (4-6 refs): sake puro de arroz. Maridaje con tiraditos.")
    add_bullet("Sake Junmai Ginjo/Daiginjo (4-6 refs): sakes premium para omakase nikkei.")
    add_bullet("Sake Nigori (2 refs): sake sin filtrar, lechoso, dulce. Marida con postres.")
    add_bullet("Cervezas peruanas y japonesas (4-5 refs): Cusqueña, Pilsen Callao, Asahi, Sapporo, Kirin.")
    add_bullet("Cócteles nikkei (10-14): Pisco Sour, Maracuyá Sour, Chilcano (pisco + ginger ale + limón), Yuzu Sour, Aji Limo Sour, Chicha Sour, Pisco Tonic con sakura, Coca Sour, Pisco Punch, Sake Mojito, Matcha Sour.")
    add_bullet("Refrescos peruanos (4-5 refs): Inca Kola, chicha morada casera, emoliente, limonada de hierba luisa.")
    add_bullet("Vinos (8-12 refs): selección corta de blancos minerales (albariño, godello, riesling) y tintos ligeros para maridar pescado crudo.")
    doc.add_heading("Inversión y márgenes", level=2)
    add_bullet("Stock inicial barra (40-60 refs): 8.000-18.000€")
    add_bullet("Vitrina refrigerada / estantería iluminada: 3.000-6.000€")
    add_bullet("Utensilios de coctelería (boston shaker, jigger, colador, pinzas, dispenser clara de huevo): 600-1.200€")
    add_bullet("Margen Pisco Sour: 85% (coste 1.80€ → venta 11-14€)")
    add_bullet("Margen sake premium: 70-80% (compra 25€/botella → venta 13-18€/copa de 120ml)")
    add_bullet("Margen cócteles nikkei: 80-85%")
    add_bullet("Objetivo: que la barra represente 25-35% de la facturación total")
    tip("Ofrece vuelos de Pisco Sour (3 versiones: clásico, maracuyá, ají limo) por 18-22€. Es la forma más rentable de vender el producto estrella nikkei: el cliente aprende, disfruta y entiende el alma peruana del concepto. Lo mismo con sake: vuelo junmai + ginjo + daiginjo por 22-28€. La combinación pisco + sake en la misma carta es el corazón líquido del nikkei.")
    doc.add_page_break()

    # CH12
    doc.add_heading("12. Brigada de Cocina (14-18 personas)", level=1)
    doc.add_paragraph("La brigada de cocina de un restaurante nikkei tiene roles específicos que no existen en una cocina europea convencional: el jefe de cocina nikkei con dominio de corte japonés y recetas peruanas, el encargado de la barra de tiraditos, el parrillero de robata para anticuchos y el wok cook para chaufa son puestos muy especializados.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Jefe de cocina nikkei (ideal con paso por Maido, Pakta, Osaka, Kikko o con formación peruano-japonesa): 1 persona — 3.000-4.500€ brutos/mes")
    add_bullet("Cocinero barra de tiraditos / segundo de cocina (corte japonés + leches de tigre): 1 persona — 2.200-2.800€ brutos/mes")
    add_bullet("Parrillero robata/josper (anticuchos, chita, pulpo al olivo): 1 persona — 1.900-2.400€ brutos/mes")
    add_bullet("Wok cook / hot kitchen (chaufa, yakimeshi, saltados, chicharrones): 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Cocinero causas, postres y guarniciones: 1 persona — 1.700-2.100€ brutos/mes")
    add_bullet("Cocineros de partida (mise en place, preparación): 1-2 personas — 1.600-2.000€ brutos/mes cada uno")
    add_bullet("Preparación pescado / escamado / limpieza: 1 persona — 1.500-1.800€ brutos/mes")
    add_bullet("Ayudante / office: 1-2 personas — 1.300-1.600€ brutos/mes")
    add_bullet("Coste total cocina: 17.000-28.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Contratación del jefe de cocina nikkei", level=2)
    doc.add_paragraph("El jefe de cocina es el alma del restaurante nikkei y la decisión más importante que tomarás. Necesita experiencia real en cocina nikkei (idealmente con paso por referentes como Maido en Lima, Pakta en Barcelona, Osaka o Kikko en Madrid), dominio del corte japonés (yanagiba, deba), conocimiento profundo de las leches de tigre, las causas, los ajíes y las parrillas. Coste en España: 3.000-4.500€/mes. Alternativa: contratar un sushi chef con pasión por el Perú y mandarlo 1-2 meses a formarse a Lima — inversión formación 3.000-5.000€ pero altísimo retorno.")
    tip("Busca cocineros nikkei a través de redes peruanas en España (Asociación Peruana, comunidades peruanas en Madrid y Barcelona), LinkedIn, Hostelwork, y contactando directamente con cocineros que hayan pasado por Pakta, Osaka, Kikko o Maizal. Una opción muy eficaz: contratar un cocinero peruano con base clásica y hacerlo viajar a Tokyo 2 semanas para aprender el corte japonés.")
    doc.add_page_break()

    # CH13
    doc.add_heading("13. Equipo de Sala (5-8 personas)", level=1)
    doc.add_paragraph("El equipo de sala de un restaurante nikkei necesita conocer la historia de la fusión peruano-japonesa, los ingredientes peruanos clave (ajíes, leche de tigre, causa, anticucho), los tipos de pisco y la carta de sakes. Debe saber contar por qué el tiradito no es un sashimi y por qué el ceviche nikkei no es un ceviche peruano tradicional.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Encargado/a de sala (conocimientos de pisco, sake y cocina nikkei): 1 persona — 2.200-2.800€ brutos/mes")
    add_bullet("Bartender / coctelero nikkei especialista en pisco y sake: 1 persona — 2.200-2.800€ brutos/mes")
    add_bullet("Camareros: 3-4 personas — 1.500-1.800€ brutos/mes cada uno")
    add_bullet("Runner / ayudante: 1-2 personas — 1.300-1.500€ brutos/mes")
    add_bullet("Coste total sala: 9.000-15.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Claves del servicio en un restaurante nikkei", level=2)
    add_bullet("Conocimiento de pisco: todo el equipo debe saber distinguir Quebranta, Italia, Torontel, acholado, y explicar Pisco Sour, Chilcano.")
    add_bullet("Conocimiento de sake: distinguir Junmai, Ginjo, Daiginjo, Nigori y recomendar maridajes con tiraditos y makis acevichados.")
    add_bullet("Explicar platos: muchos clientes no saben qué es un tiradito, una causa, un anticucho o un maki acevichado. El camarero debe saber describirlo con pasión y precisión.")
    add_bullet("Contar la historia nikkei: inmigración japonesa al Perú en el siglo XIX, nacimiento de la fusión, referentes (Nobu, Micha).")
    add_bullet("Orden sugerido: tiradito/ceviche nikkei primero (sabor limpio), luego causas y anticuchos, makis y chaufa al final.")
    add_bullet("Upselling de bebidas: Pisco Sour de bienvenida, vuelo de piscos, sake maridado con omakase. El bartender es clave para el ticket medio.")
    add_bullet("Ajíes: explicar al cliente el nivel de picante (ají amarillo suave, ají limo medio, rocoto fuerte) y adaptar.")
    tip("La formación del bartender nikkei es una inversión que se multiplica. Un bartender que sabe contar la historia del pisco, diferencia uvas pisqueras y prepara un Pisco Sour perfecto (con clara de huevo emulsionada correctamente) vende el triple en bebida. Considera enviar al bartender a un curso intensivo en Lima o en la Escuela Nacional de Bartenders del Perú.")
    doc.add_page_break()

    # CH14
    doc.add_heading("14. Menú Engineering Nikkei: Tiraditos, Causas, Anticuchos, Chaufa", level=1)
    doc.add_paragraph("El diseño de la carta de un restaurante nikkei debe equilibrar autenticidad peruano-japonesa, margen y accesibilidad. La carta debe educar al cliente no experto sin aburrir al experto, y ofrecer platos de todos los rangos de precio.")
    doc.add_heading("Estructura de carta nikkei", level=2)
    add_bullet("Tiraditos: 5-8 opciones (atún en leche de tigre de ají amarillo, corvina al ají limo, salmón al yuzu-ponzu, pulpo al olivo nikkei)")
    add_bullet("Ceviches nikkei: 3-5 opciones (ceviche clásico con toque nikkei, ceviche de pescado + langostinos, ceviche al rocoto)")
    add_bullet("Causas nikkei: 4-6 opciones (causa de atún, causa de salmón ahumado, causa de pulpo, causa de langostinos con salsa acevichada)")
    add_bullet("Makis acevichados y uramakis nikkei: 6-10 opciones (maki acevichado, crunch maki, maki de langostinos al ají amarillo, dragon roll nikkei)")
    add_bullet("Nigiris nikkei: 6-10 opciones (nigiri de atún al ají limo, nigiri de salmón al rocoto, nigiri de vieira al miso-yuzu)")
    add_bullet("Anticuchos: 4-6 opciones (corazón de ternera con miso, pulpo, chita, wagyu con chimichurri nikkei, brochetas de langostinos)")
    add_bullet("Calientes: 4-6 opciones (chita a la parrilla robata, arroz chaufa de mariscos, yakimeshi de langostinos, tataki de atún)")
    add_bullet("Postres: 3-4 opciones (suspiro a la limeña con yuzu, mochi de lúcuma, helado de chocolate con maracuyá, tarta de chirimoya)")
    add_bullet("Total carta: 40-55 platos")
    doc.add_heading("Menú engineering: clasificación de platos nikkei", level=2)
    add_bullet("Stars (alta popularidad + alto margen): Tiradito de atún, Maki acevichado, Pisco Sour, Causa nikkei → destacar siempre")
    add_bullet("Plowhorses (alta popularidad + bajo margen): Ceviche clásico, Uramaki California → subir precio o ajustar porción")
    add_bullet("Puzzles (baja popularidad + alto margen): Anticucho de wagyu, Chita entera robata, Omakase nikkei → mejorar descripción y fotografía")
    add_bullet("Dogs (baja popularidad + bajo margen): Ensaladas genéricas, sopas mal presentadas → eliminar o reinventar")
    tip("El tiradito de atún en leche de tigre de ají amarillo es tu Star absoluto. El coste es 4-5€ y lo vendes a 18-22€. Con un tiradito + maki acevichado + Pisco Sour por persona ya tienes un ticket de 45€+ con márgenes excelentes. Es el plato con mejor ratio margen/satisfacción/foto de Instagram de toda la carta nikkei.")
    doc.add_page_break()

    # CH15
    doc.add_heading("15. Proveedores Nikkei en España (Perú + Japón)", level=1)
    doc.add_paragraph("Encontrar ingredientes nikkei auténticos en España requiere combinar dos redes distintas: los importadores de producto peruano (para ajíes, pastas, quinoa, ingredientes clave) y los distribuidores de producto japonés (para arroz Koshihikari, salsas, algas, sake).")
    doc.add_heading("Ingredientes clave y dónde encontrarlos", level=2)
    add_bullet("Pescado sashimi-grade: Mercamadrid, Mercabarna, lonjas de San Sebastián y Galicia. Proveedores especializados (Krustagroup, Pescadería Universal, Albert Rovira, Pescados Oceanic).")
    add_bullet("Ajíes peruanos (amarillo, limo, rocoto, panca): importadores españoles de producto peruano — Inkawasi, Latin Products, Perú Market, Casa del Perú. Ajíes frescos, en pasta o deshidratados.")
    add_bullet("Mariscos (langostinos, vieiras, pulpo): Mercamadrid, Mercabarna, proveedores especializados.")
    add_bullet("Arroz Koshihikari (o equivalente para sushi): distribuidores japoneses en España (Japan Sushi Express, Comercial Minamoto). Origen España, Italia o Japón.")
    add_bullet("Alga nori premium (yaki nori): distribuidores japoneses. Grado A (oro) es el mejor.")
    add_bullet("Salsa de soja (shoyu): Kikkoman estándar. Premium: Yamasa, Yamaroku, Kishibori. Distribuidores japoneses.")
    add_bullet("Mirin, vinagre de arroz, sake de cocina: Kikkoman, Mitsukan. Distribuidores japoneses o tiendas asiáticas.")
    add_bullet("Limón peruano / limón sutil: difícil fresco; sustituir con mezcla de lima fresca + zumo de limón peruano congelado (importación peruana).")
    add_bullet("Quinoa, camote, yuca, choclo: importadores latinoamericanos en España, mercados centrales.")
    add_bullet("Carbón binchotan (para robata): importación desde Japón o alternativas de Laos/Vietnam. Distribuidores especializados.")
    add_bullet("Pisco peruano: distribuidores en España — Viñas de Oro, Pisco 1615, Barsol, Waqar. Algunos importan directamente desde Ica.")
    add_bullet("Sake y vinos: distribuidores Japan Sake (Madrid), Sake Distribución España. Sakes premium Junmai Daiginjo de referentes como Dassai, Hakkaisan.")
    add_bullet("Yuzu, shiso: difíciles frescos; mayoritariamente en pasta o congelado vía distribuidores japoneses.")
    add_bullet("Huacatay, hierba luisa, muña: importadores peruanos, congelados o en pasta.")
    tip("Establece relación con un proveedor de pescado sashimi-grade diario Y con un importador peruano de confianza ANTES de abrir. Sin leche de tigre con ají amarillo fresco y sin pescado de máxima calidad todos los días no hay nikkei viable. Los mejores proveedores de ají amarillo pasta en España: Inkawasi y Latin Products. Firma acuerdos de suministro y haz pedidos semanales programados.")
    doc.add_page_break()

    # CH16
    doc.add_heading("16. 15 Recetas Base con Food Cost", level=1)
    doc.add_paragraph("Estas 15 recetas forman la base de la carta de un restaurante nikkei. Cada una incluye food cost estimado y PVP sugerido para mantener márgenes del 30-35%.")
    recipes = [
        ("Tiradito de atún en leche de tigre de ají amarillo", "4.50€", "18.00€", "25%"),
        ("Tiradito de corvina al ají limo", "4.00€", "17.00€", "24%"),
        ("Ceviche nikkei de pescado + langostinos", "5.50€", "19.00€", "29%"),
        ("Causa nikkei de atún con salsa acevichada", "3.20€", "14.00€", "23%"),
        ("Causa nikkei de pulpo al olivo", "4.50€", "16.00€", "28%"),
        ("Maki acevichado (8 piezas)", "3.80€", "15.00€", "25%"),
        ("Uramaki crunch nikkei (8 piezas)", "3.50€", "14.00€", "25%"),
        ("Nigiri de atún al ají limo (2 uds)", "2.80€", "11.00€", "25%"),
        ("Anticucho de corazón con miso (3 brochetas)", "2.20€", "12.00€", "18%"),
        ("Anticucho de pulpo al olivo nikkei (2 uds)", "5.50€", "18.00€", "31%"),
        ("Chita a la parrilla robata (ración)", "9.00€", "28.00€", "32%"),
        ("Arroz chaufa de mariscos", "4.50€", "16.00€", "28%"),
        ("Yakimeshi nikkei de langostinos", "4.20€", "15.00€", "28%"),
        ("Tataki de atún con ají amarillo", "6.00€", "20.00€", "30%"),
        ("Suspiro a la limeña con yuzu", "1.80€", "8.50€", "21%"),
    ]
    for name, fc, pvp, pct in recipes:
        doc.add_heading(name, level=2)
        add_bullet(f"Food cost: {fc}")
        add_bullet(f"PVP sugerido (sin IVA): {pvp}")
        add_bullet(f"% Food cost: {pct}")
    tip("El tiradito de atún con leche de tigre de ají amarillo es tu plato bandera nikkei y el más instagrameable. Food cost 4.50€, PVP 18€, food cost 25%. Un cliente que pide tiradito + maki acevichado + Pisco Sour por persona ya tiene un ticket de 42€ con márgenes excelentes. Es el plato que convierte clientes nuevos en embajadores.")
    doc.add_page_break()

    # CH17
    doc.add_heading("17. Delivery Optimizado para Cocina Nikkei", level=1)
    doc.add_paragraph("La cocina nikkei tiene platos que viajan muy bien (makis acevichados, causas, arroz chaufa, bowls) y otros que no (tiraditos puros, ceviches con leche de tigre abundante). Adaptar el menú de delivery es clave para mantener la reputación de calidad.")
    doc.add_heading("Plataformas y canales", level=2)
    add_bullet("Glovo, Uber Eats, Just Eat: comisión 25-35% — útiles para visibilidad inicial")
    add_bullet("Delivery propio (web/app): comisión 0% pero requiere repartidores propios o Stuart/Paack")
    add_bullet("Take away en local: margen completo, packaging incluido en precio")
    doc.add_heading("Menú delivery nikkei optimizado", level=2)
    add_bullet("Maki boxes (12-24 piezas mixtas): producto estrella del delivery nikkei. Makis acevichados, crunch maki, uramaki de langostinos. Caja negra con compartimentos, salsa acevichada aparte.")
    add_bullet("Causa bowls: base de causa (puré de papa amarilla + ají amarillo) con topping de atún o salmón, ají limo y palta. Formato bowl cerrado, viaja perfecto.")
    add_bullet("Tiradito box (entrega <15 min): pescado laminado separado de la leche de tigre en cápsula aparte. Cliente la vierte al abrir. Instrucciones en el envase.")
    add_bullet("Arroz chaufa de mariscos: viaja excelente. Envase con tapa hermética, mantiene textura y temperatura.")
    add_bullet("Yakimeshi de langostinos: similar al chaufa, muy adaptable al delivery.")
    add_bullet("Anticuchos con arroz: brochetas + arroz blanco + chimichurri nikkei. Envase con tapa, mantiene calor 20-25 min.")
    add_bullet("Causas y ensaladas: viajan perfecto, frescor mantenido.")
    add_bullet("NO incluir en delivery: tiradito con leche de tigre ya bañada (el pescado se cocina en exceso), ceviches con muchas horas de viaje, pescados enteros de robata.")
    add_bullet("Packaging: cajas negras con logotipo nikkei, mini-botes de salsa acevichada y ají limo, palillos. Coste 1.20-2.80€/pedido.")
    tip("El 'maki box acevichado' premium es la mina de oro del delivery nikkei. Incluye 18-24 piezas variadas (acevichado, crunch, langostinos), más edamame y anticucho de pollo, y te va a 38-48€. Margen 65% con packaging premium. Plantea también 'nikkei para compartir' (48-60 piezas + causas + anticuchos) para cenas en casa. Ticket 65-95€ con food cost del 27%.")
    doc.add_page_break()

    # CH18
    doc.add_heading("18. Marketing y Calendario Cultural Nikkei (Perú + Japón)", level=1)
    doc.add_paragraph("Un restaurante nikkei tiene un calendario cultural con eventos peruanos y japoneses que generan picos de facturación. El Día del Pisco Sour (primer sábado de febrero), las Fiestas Patrias del Perú (28 julio), el Día de la Gastronomía Peruana y el Día Nacional del Ceviche son oportunidades excelentes para storytelling y reservas.")
    doc.add_heading("Eventos culturales para marketing", level=2)
    add_bullet("Día del Pisco Sour (primer sábado de febrero): evento bandera del nikkei. Menú especial maridado con vuelo de Pisco Sour, masterclass de coctelería, historia del pisco.")
    add_bullet("Día Nacional del Pisco (cuarto domingo de julio): celebración del destilado peruano. Cata de piscos Quebranta, Italia, acholados.")
    add_bullet("Fiestas Patrias del Perú (28-29 julio): menú especial de Fiestas Patrias con causa, anticuchos, chaufa y Pisco Sour. Máxima ocupación asegurada.")
    add_bullet("Día Nacional del Ceviche (28 junio): evento de ceviches nikkei, cata de leches de tigre, explicación de ajíes peruanos.")
    add_bullet("Día de la Gastronomía Peruana (septiembre): participación en eventos de la comunidad peruana en España, colaboraciones con Mistura España.")
    add_bullet("Año Nuevo japonés (Oshogatsu, 1 enero): menú fusión de soba de fin de año y plato nikkei de año nuevo con pescado premium.")
    add_bullet("Día internacional del sushi (18 junio): omakase nikkei, evento con el cocinero explicando maki acevichado y nigiris al ají limo.")
    doc.add_heading("Canales de marketing", level=2)
    add_num("Instagram: contenido visual de tiraditos, makis acevichados, Pisco Sour. Reels del cocinero bañando el tiradito con leche de tigre son virales.")
    add_num("TikTok: vídeos del cocinero preparando maki acevichado, licuando leche de tigre, flameando anticuchos en robata.")
    add_num("Google My Business: ficha completa con fotos profesionales, menú, reseñas.")
    add_num("TheFork / TripAdvisor: perfiles con fotos y respuesta a reseñas.")
    add_num("Web propia: menú, historia nikkei, reservas online, carta de piscos y sakes.")
    add_num("Comunidad peruana online: asociaciones peruanas en España, grupos de Facebook, embajada del Perú.")
    doc.add_heading("Presupuesto marketing mensual", level=2)
    add_bullet("Google Ads local: 300-700€/mes")
    add_bullet("Instagram/Facebook Ads: 300-700€/mes")
    add_bullet("Fotógrafo food: 300-500€/sesión trimestral")
    add_bullet("Community manager (si se externaliza): 400-800€/mes")
    add_bullet("Total recomendado: 800-2.000€/mes (2-3% facturación)")
    tip("El contenido visual del nikkei es explosivo: el momento exacto en que el cocinero baña el tiradito con leche de tigre de ají amarillo, el flambeado de un anticucho, la emulsión de un Pisco Sour con clara de huevo. Publica Reels diarios. El Día del Pisco Sour (primer sábado de febrero) y Fiestas Patrias (28 julio) son tus dos días de máxima facturación — reserva las mesas con 2 meses de antelación y comunica agresivamente.")
    doc.add_page_break()

    # CH19
    doc.add_heading("19. Tecnología para Restaurante Nikkei", level=1)
    doc.add_paragraph("La tecnología en un restaurante nikkei debe ser invisible pero fiable. En un omakase nikkei con reservas de 3 semanas, no puedes permitirte un error de gestión. En una cevichería nikkei con cola diaria de mediodía, necesitas gestión de espera digital.")
    doc.add_heading("Stack tecnológico recomendado", level=2)
    add_bullet("TPV: Lightspeed, Square, Revo. Cloud, con reporting en tiempo real. 60-150€/mes.")
    add_bullet("Reservas: TheFork, CoverManager, Resy. Gestión de turnos, depósitos para omakase, política de no-shows. 0-150€/mes.")
    add_bullet("Delivery: integrador (Ordatic, Deliverect) que centraliza Glovo + Uber Eats + Just Eat. 80-150€/mes.")
    add_bullet("Contabilidad: Holded, Quipu. Facturas automáticas, conexión con asesoría. 30-60€/mes.")
    add_bullet("RRHH y turnos: Factorial, Kenjo. Cuadrantes, fichajes, nóminas. 4-8€/empleado/mes.")
    add_bullet("Gestión de cola (si hay espera en horas punta): Waitwhile, Qminder. Clientes dan su nombre y reciben notificación. 30-80€/mes.")
    add_bullet("Web y carta digital: QR con carta actualizable, sección de piscos y sakes. WordPress o Squarespace. 10-30€/mes.")
    add_bullet("WiFi profesional: router dual-band, red separada para clientes y cocina. 50-80€/mes.")
    tip("En un omakase, la gestión de reservas con depósito es fundamental. Cobra 30-50€ por persona como depósito de reserva (reembolsable contra la cena) para evitar no-shows. Los no-shows en un omakase de 14 plazas pueden destruir tu facturación del día. CoverManager y Resy permiten gestionar esto automáticamente.")
    doc.add_page_break()

    # CH20
    doc.add_heading("20. Plan de Acción: De la Idea a la Inauguración", level=1)
    doc.add_paragraph("Abrir un restaurante nikkei en España es un proyecto de 12-16 meses desde la primera idea hasta la inauguración. Este capítulo resume el cronograma y los hitos clave.")
    doc.add_heading("Fases del proyecto", level=2)
    add_bullet("Meses 1-3: Plan de negocio, búsqueda de financiación, constitución de empresa, estudio de mercado.")
    add_bullet("Meses 3-5: Búsqueda y contratación de local, negociación de alquiler, licencia de actividad.")
    add_bullet("Meses 4-7: Obra civil, instalaciones (especialmente extracción reforzada para robata), decoración nikkei contemporánea.")
    add_bullet("Meses 5-7: Equipamiento de cocina nikkei (arrocera industrial, vitrina tiraditos, licuadoras industriales, parrilla robata/josper, wok station, abatidor, cuchillos japoneses).")
    add_bullet("Meses 5-8: Búsqueda y reclutamiento del jefe de cocina nikkei (la decisión más crítica, puede tardar meses).")
    add_bullet("Meses 5-8: Establecer relaciones con proveedores de pescado sashimi-grade e importadores peruanos (ajíes, pisco).")
    add_bullet("Meses 8-10: Contratación y formación del resto del equipo (cocinero tiraditos, parrillero robata, wok cook, bartender nikkei).")
    add_bullet("Meses 9-11: Branding, web, redes sociales, fotografía profesional.")
    add_bullet("Mes 11-12: Soft opening con amigos, familia, influencers, prensa gastronómica y comunidad peruana en España.")
    add_bullet("Mes 12-13: Ajustes finales e INAUGURACIÓN.")
    add_bullet("Mes 13-16: Seguimiento, ajustes de carta y optimización de costes.")
    tip("No abras sin haber hecho mínimo 4 soft openings. Necesitas probar la barra de tiraditos con el cocinero a pleno rendimiento, las leches de tigre emulsionadas en servicio real, el flujo de la parrilla robata con anticuchos simultáneos, la coctelería con Pisco Sour al momento y la operativa de delivery. Cada soft opening te enseña fallos que no puedes detectar en papel. Con un nikkei premium no puedes permitirte un mal servicio el primer día — el boca a boca en la comunidad peruana y foodie es brutal.")

    # Save
    path = os.path.join(OUTPUT_DIR, "guia-restaurante-nikkei.docx")
    doc.save(path)
    print(f"✓ {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 2. EXCEL TEMPLATES (8)
# ═══════════════════════════════════════════════════════════

def gen_plan_financiero():
    wb = Workbook()
    instr_sheet(wb, "Plan Financiero 3 Años — Restaurante Nikkei 60 Plazas", [
        "Rellena las celdas verdes con los datos de tu proyecto.",
        "Las fórmulas se recalculan automáticamente.",
        "Pestaña 'Inversión' = CAPEX inicial desglosado.",
        "Pestaña 'P&L Mensual' = cuenta de resultados mes a mes.",
        "Pestaña 'Proyección 3 Años' = evolución anual.",
    ])
    # Inversión sheet
    ws = wb.create_sheet("Inversión"); ws.sheet_properties.tabColor = "4CAF50"
    title_block(ws, "Inversión Inicial — Restaurante Nikkei 60 Plazas")
    hdrs = ["Categoría", "Partida", "Coste Estimado (€)", "Coste Real (€)", "Diferencia", "Notas"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 6)
    widths = [25, 40, 18, 18, 18, 30]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    items = [
        ("Obra Civil", "Acondicionamiento local", 110000),
        ("Obra Civil", "Instalaciones eléctricas y gas", 14000),
        ("Obra Civil", "Fontanería y saneamiento", 9000),
        ("Obra Civil", "Climatización / HVAC", 10000),
        ("Obra Civil", "Extracción reforzada (robata/josper, wok station)", 12000),
        ("Cocina NK", "Arrocera industrial profesional (5-10L)", 3200),
        ("Cocina NK", "Vitrina refrigerada de pescado para barra tiraditos", 6500),
        ("Cocina NK", "Licuadoras industriales leches de tigre (2-3 uds)", 2800),
        ("Cocina NK", "Batidora inmersión profesional + utensilios", 600),
        ("Cocina NK", "Parrilla robata carbón binchotan / Josper", 8500),
        ("Cocina NK", "Wok station 2 woks alta potencia (chaufa)", 4500),
        ("Cocina NK", "Cuchillos japoneses set (yanagiba, deba, usuba)", 2800),
        ("Cocina NK", "Abatidor de temperatura certificado", 5500),
        ("Cocina NK", "Ahumador pequeño (opcional para tatakis)", 1200),
        ("Cocina Estándar", "Horno mixto Rational/Unox", 12000),
        ("Cocina Estándar", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina Estándar", "Plancha teppanyaki (opcional)", 2500),
        ("Cocina Estándar", "Freidora doble (crocante camote, chicharrones)", 2500),
        ("Cocina Estándar", "Cámaras frigoríficas (2)", 6000),
        ("Cocina Estándar", "Congelador CERTIFICADO -20°C (anisakis)", 4500),
        ("Cocina Estándar", "Tren de lavado", 5500),
        ("Cocina Estándar", "Campana extractora REFORZADA (robata + wok)", 10000),
        ("Cocina Estándar", "Menaje y utensilios mixtos (peruano + japonés)", 3500),
        ("Barra Tiraditos", "Barra profesional madera/vitrina", 9000),
        ("Barra Pisco-Sake", "Vitrina refrigerada pisco, sake y coctelería", 4500),
        ("Barra Pisco-Sake", "Utensilios coctelería nikkei (shaker, jigger, dispenser clara)", 900),
        ("Barra Pisco-Sake", "Stock inicial pisco + sake + vinos (40-60 refs)", 12000),
        ("Barra Sake", "Máquina de hielo", 1800),
        ("Sala", "Mobiliario interior mesas madera + sillas (60 plazas)", 22000),
        ("Sala", "Barra sushi taburetes altos", 4000),
        ("Decoración", "Madera natural en paredes de acento", 8000),
        ("Decoración", "Cuelgas ajíes peruanos secos + detalles textiles andinos", 800),
        ("Decoración", "Lámparas de papel/fibras naturales e iluminación cálida", 3500),
        ("Decoración", "Elementos decorativos nikkei contemporáneos", 2500),
        ("Decoración", "Vajilla cerámica artesanal (mixta Perú/Japón)", 4500),
        ("Decoración", "Detalles decorativos en piedra y cobre", 1200),
        ("Terraza", "Mobiliario exterior (si aplica)", 5000),
        ("Tecnología", "TPV + pantallas cocina", 3500),
        ("Tecnología", "Web + carta QR + WiFi + sistema reservas", 3500),
        ("Licencias", "Licencia actividad + terraza + proyecto técnico", 12000),
        ("Licencias", "Registro EORI / importación (si aplica)", 500),
        ("Marketing", "Branding nikkei + diseño + logotipo", 5000),
        ("Marketing", "Fotos + vídeo + campaña pre-apertura", 5000),
        ("Stock", "Materia prima inicial (pescado, arroz, ajíes, salsas)", 15000),
        ("Stock", "Bebidas (cerveza Pilsen/Asahi, Inca Kola, refrescos)", 5000),
        ("Maniobra", "Fondo de maniobra (3 meses)", 70000),
    ]
    r += 1
    start_r = r
    for cat, partida, coste in items:
        sdc(ws, r, 1, cat); sdc(ws, r, 2, partida)
        sdc(ws, r, 3, coste, fmt=cur_fmt); sdc(ws, r, 4, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}-C{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 2, "TOTAL INVERSIÓN", font=bld_font)
    sdc(ws, r, 3, f"=SUM(C{start_r}:C{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=D{r}-C{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 6)

    # P&L Mensual sheet
    ws2 = wb.create_sheet("P&L Mensual"); ws2.sheet_properties.tabColor = "2196F3"
    title_block(ws2, "P&L Mensual — Restaurante Nikkei 60 Plazas", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic", "TOTAL"]
    r = 4
    for i, m in enumerate(months, 1): ws2.cell(row=r, column=i, value=m)
    shr(ws2, r, 14)
    ws2.column_dimensions['A'].width = 28
    for i in range(2, 15): ws2.column_dimensions[get_column_letter(i)].width = 12
    ws2.freeze_panes = f"B{r+1}"
    rows_data = [
        ("INGRESOS", None, True),
        ("Ventas sala (tiraditos + causas + anticuchos + chaufa)", 45000, False),
        ("Barra pisco + sake + coctelería nikkei", 15000, False),
        ("Delivery (maki boxes, bowls, anticuchos)", 10000, False),
        ("Bebidas (cerveza, refrescos, Inca Kola)", 8000, False),
        ("TOTAL INGRESOS", "=SUM", True),
        ("", None, False),
        ("COSTES VARIABLES", None, True),
        ("Food cost (33%)", "=0.33*", False),
        ("Packaging delivery", 700, False),
        ("TOTAL COSTES VARIABLES", "=SUM_CV", True),
        ("", None, False),
        ("COSTES FIJOS", None, True),
        ("Alquiler", 6500, False),
        ("Personal cocina (jefe nikkei, tiraditos, robata, wok)", 22000, False),
        ("Personal sala (bartender nikkei, camareros)", 13000, False),
        ("Suministros (luz, gas, agua)", 3500, False),
        ("Seguros", 500, False),
        ("Gestoría + contabilidad", 500, False),
        ("Marketing", 1200, False),
        ("Tecnología (TPV, delivery, reservas)", 450, False),
        ("Mantenimiento", 400, False),
        ("Reposición stock pisco/sake", 1800, False),
        ("Varios / imprevistos", 700, False),
        ("TOTAL COSTES FIJOS", "=SUM_CF", True),
        ("", None, False),
        ("EBITDA", "=EBITDA", True),
        ("% EBITDA", "=%EBITDA", True),
    ]
    r += 1
    for label, val, bold in rows_data:
        sdc(ws2, r, 1, label, font=bld_font if bold else dat_font)
        if val and val not in ("=SUM", "=SUM_CV", "=SUM_CF", "=EBITDA", "=%EBITDA") and not str(val).startswith("=0."):
            for c in range(2, 14):
                sdc(ws2, r, c, val, fill=inp_fill, fmt=cur_fmt)
            sdc(ws2, r, 14, f"=SUM(B{r}:M{r})", font=bld_font, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws2, r, 14)

    path = os.path.join(OUTPUT_DIR, "plan-financiero-3-anos.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_capex():
    wb = Workbook()
    instr_sheet(wb, "Calculadora CAPEX — Restaurante Nikkei 60 Plazas", [
        "Introduce tus costes reales en las celdas verdes.",
        "La columna 'Diferencia' se calcula automáticamente.",
        "Usa la columna 'Prioridad' para planificar fases de inversión.",
    ])
    ws = wb.create_sheet("CAPEX Desglosado"); ws.sheet_properties.tabColor = "FF9800"
    title_block(ws, "Calculadora CAPEX — Inversión 280K-520K€")
    hdrs = ["#", "Categoría", "Partida", "Estimado (€)", "Real (€)", "Diferencia", "Prioridad", "Proveedor"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8)
    widths = [5, 20, 40, 14, 14, 14, 12, 25]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    dv = DataValidation(type="list", formula1='"Alta,Media,Baja"', allow_blank=True)
    ws.add_data_validation(dv)
    items = [
        ("Obra", "Acondicionamiento general", 110000),
        ("Obra", "Electricidad + gas", 14000),
        ("Obra", "Fontanería", 9000),
        ("Obra", "Climatización", 10000),
        ("Obra", "Extracción reforzada robata/wok station", 12000),
        ("Cocina NK", "Arrocera industrial profesional", 3200),
        ("Cocina NK", "Vitrina refrigerada pescado (barra tiraditos)", 6500),
        ("Cocina NK", "Licuadoras industriales (leches de tigre)", 2800),
        ("Cocina NK", "Abatidor temperatura certificado", 5500),
        ("Cocina NK", "Parrilla robata binchotan / Josper", 8500),
        ("Cocina NK", "Wok station 2 woks alta potencia", 4500),
        ("Cocina NK", "Cuchillos japoneses set (yanagiba, deba, usuba)", 2800),
        ("Cocina NK", "Ahumador pequeño (opcional)", 1200),
        ("Cocina", "Horno mixto Rational/Unox", 12000),
        ("Cocina", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina", "Plancha teppanyaki (opcional)", 2500),
        ("Cocina", "Freidora doble (crocantes, chicharrones)", 2500),
        ("Cocina", "Cámaras frigoríficas (2)", 6000),
        ("Cocina", "Congelador CERTIFICADO -20°C anisakis", 4500),
        ("Cocina", "Tren de lavado", 5500),
        ("Cocina", "Campana extractora REFORZADA robata+wok", 10000),
        ("Cocina", "Menaje y utensilios mixtos", 3500),
        ("Barra Tiraditos", "Barra profesional madera", 9000),
        ("Barra Pisco-Sake", "Vitrina refrigerada pisco/sake/vinos", 4500),
        ("Barra Pisco-Sake", "Utensilios coctelería + máquina hielo", 2500),
        ("Barra Pisco-Sake", "Stock inicial pisco+sake+vinos (40-60 refs)", 12000),
        ("Sala", "Mesas madera, sillas, barra (60 plazas)", 26000),
        ("Decoración", "Madera, noren, washi, jardín zen, cerámica", 20000),
        ("Terraza", "Mobiliario exterior", 5000),
        ("Tech", "TPV + pantallas cocina + reservas", 7000),
        ("Licencias", "Licencia actividad + terraza + proyecto", 12500),
        ("Marketing", "Branding + fotos + vídeo + pre-apertura", 10000),
        ("Stock", "Materia prima + pescado + arroz + salsas", 15000),
        ("Stock", "Bebidas (cerveza, té, refrescos)", 5000),
        ("Maniobra", "Fondo 3 meses", 70000),
    ]
    r += 1; start_r = r
    for idx, (cat, partida, est) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, partida)
        sdc(ws, r, 4, est, fmt=cur_fmt); sdc(ws, r, 5, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 6, f"=E{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 7, "Alta", fill=inp_fill); dv.add(ws.cell(row=r, column=7))
        sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 3, "TOTAL", font=bld_font)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=SUM(E{start_r}:E{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 6, f"=E{r}-D{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 8)
    path = os.path.join(OUTPUT_DIR, "calculadora-capex.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_pl_mensual():
    wb = Workbook()
    instr_sheet(wb, "P&L Mensual con 3 Escenarios — Restaurante Nikkei", [
        "3 pestañas: Pesimista, Realista, Optimista.",
        "Modifica los ingresos en cada escenario.",
        "Los costes fijos son iguales en los 3 escenarios.",
    ])
    for scenario, factor, color in [("Pesimista", 0.75, "F44336"), ("Realista", 1.0, "4CAF50"), ("Optimista", 1.25, "2196F3")]:
        ws = wb.create_sheet(scenario); ws.sheet_properties.tabColor = color
        title_block(ws, f"P&L Mensual — Escenario {scenario}", 4)
        hdrs = ["Concepto", "Mensual (€)", "Anual (€)", "% s/Ventas"]
        r = 4
        for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
        shr(ws, r, 4)
        ws.column_dimensions['A'].width = 35
        for i in range(2, 5): ws.column_dimensions[get_column_letter(i)].width = 16
        base_ventas = int(78000 * factor)
        rows = [
            ("INGRESOS", None, True),
            ("Ventas sala (tiraditos + causas + anticuchos + chaufa)", int(45000 * factor)),
            ("Barra pisco + sake + coctelería nikkei", int(15000 * factor)),
            ("Delivery (maki boxes, bowls, anticuchos)", int(10000 * factor)),
            ("Bebidas (cerveza, Inca Kola, refrescos)", int(8000 * factor)),
            ("TOTAL INGRESOS", base_ventas, True),
            ("", None, False),
            ("COSTES VARIABLES", None, True),
            ("Food cost (33%)", int(base_ventas * 0.33)),
            ("Packaging delivery", 700),
            ("TOTAL COSTES VARIABLES", int(base_ventas * 0.33 + 700), True),
            ("", None, False),
            ("COSTES FIJOS", None, True),
            ("Alquiler", 6500),
            ("Personal cocina (jefe nikkei + equipo)", 22000),
            ("Personal sala (bartender + camareros)", 13000),
            ("Suministros", 3500),
            ("Seguros", 500),
            ("Gestoría", 500),
            ("Marketing", 1200),
            ("Tecnología", 450),
            ("Mantenimiento", 400),
            ("Reposición pisco/sake", 1800),
            ("Varios", 700),
            ("TOTAL COSTES FIJOS", 50550, True),
            ("", None, False),
            ("EBITDA", int(base_ventas - base_ventas * 0.33 - 700 - 50550), True),
        ]
        r += 1
        for item in rows:
            if len(item) == 3:
                label, val, bold = item
            else:
                label, val = item; bold = False
            sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
            if val is not None:
                sdc(ws, r, 2, val, font=bld_font if bold else dat_font, fmt=cur_fmt)
                sdc(ws, r, 3, val * 12, font=bld_font if bold else dat_font, fmt=cur_fmt)
            r += 1
        r += 1; brand_footer(ws, r, 4)
    path = os.path.join(OUTPUT_DIR, "pl-mensual-escenarios.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cash_flow():
    wb = Workbook()
    instr_sheet(wb, "Cash Flow y Break-Even — Restaurante Nikkei 60 Plazas", [
        "Pestaña 'Cash Flow' = flujo de caja mensual 12 meses.",
        "Pestaña 'Break-Even' = calculadora de punto de equilibrio.",
        "Rellena las celdas verdes con tus datos reales.",
    ])
    ws = wb.create_sheet("Cash Flow"); ws.sheet_properties.tabColor = "009688"
    title_block(ws, "Cash Flow 12 Meses — Restaurante Nikkei", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
    r = 4
    for i, m in enumerate(months, 1): ws.cell(row=r, column=i, value=m)
    shr(ws, r, 13)
    ws.column_dimensions['A'].width = 28
    for i in range(2, 14): ws.column_dimensions[get_column_letter(i)].width = 12
    concepts = [
        ("Saldo inicial", True), ("", False),
        ("ENTRADAS", True), ("Ventas sala (tiraditos + causas + anticuchos + chaufa)", False), ("Barra pisco/sake/coctelería", False),
        ("Delivery (maki boxes, bowls)", False), ("Bebidas + refrescos", False), ("Total entradas", True), ("", False),
        ("SALIDAS", True), ("Pescado sashimi-grade + ajíes peruanos + arroz + importados", False), ("Personal cocina (jefe nikkei)", False),
        ("Personal sala", False), ("Alquiler", False), ("Suministros", False), ("Marketing", False),
        ("Reposición stock pisco/sake", False), ("Tecnología", False), ("Otros gastos", False),
        ("Total salidas", True), ("", False),
        ("Flujo neto mensual", True), ("Saldo acumulado", True),
    ]
    r += 1
    for label, bold in concepts:
        sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
        for c in range(2, 14):
            sdc(ws, r, c, None, fill=inp_fill if not bold and label else None, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws, r, 13)

    # Break-even
    ws2 = wb.create_sheet("Break-Even"); ws2.sheet_properties.tabColor = "FF5722"
    title_block(ws2, "Calculadora Break-Even — Restaurante Nikkei", 4)
    r = 4
    params = [
        ("Ticket medio (€)", 45),
        ("Comensales/día promedio", 65),
        ("Días abierto/mes", 26),
        ("Food cost (%)", 0.33),
        ("Costes fijos mensuales (€)", 50550),
    ]
    for label, val in params:
        r += 1
        sdc(ws2, r, 1, label, font=bld_font); sdc(ws2, r, 2, val, fill=inp_fill, fmt=cur_fmt if isinstance(val, int) else pct_fmt)
    r += 2
    sdc(ws2, r, 1, "Facturación mensual estimada", font=bld_font)
    sdc(ws2, r, 2, "=B5*B6*B7", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Margen contribución mensual", font=bld_font)
    sdc(ws2, r, 2, "=B12*(1-B8)", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Break-Even (meses)", font=bld_font)
    ws2.column_dimensions['A'].width = 35; ws2.column_dimensions['B'].width = 20
    r += 2; brand_footer(ws2, r, 4)
    path = os.path.join(OUTPUT_DIR, "cash-flow-break-even.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_escandallo():
    wb = Workbook()
    instr_sheet(wb, "Escandallo Maestro — Restaurante Nikkei", [
        "Una ficha técnica por plato.",
        "Introduce ingredientes, cantidades y precios.",
        "El food cost se calcula automáticamente.",
    ])
    ws = wb.create_sheet("Escandallo"); ws.sheet_properties.tabColor = "E91E63"
    title_block(ws, "Escandallo Maestro — Fichas Técnicas de Platos Nikkei")
    hdrs = ["#", "Ingrediente", "Cantidad (g/ml)", "Precio/Kg (€)", "Coste (€)", "Merma (%)", "Coste Real (€)"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 7)
    widths = [5, 30, 16, 14, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "NOMBRE DEL PLATO:", font=bld_font)
    sdc(ws, r, 3, "Ejemplo: Tiradito de Atún en Leche de Tigre de Ají Amarillo", fill=inp_fill)
    r += 1
    ingredients = [
        ("Atún rojo sashimi-grade (previamente congelado -20°C/24h)", 120, 55.00, 18),
        ("Limón peruano / mezcla lima-limón", 40, 6.00, 10),
        ("Ají amarillo en pasta (importación Perú)", 20, 28.00, 0),
        ("Ajo fresco", 5, 7.00, 10),
        ("Jengibre (kion)", 4, 9.00, 15),
        ("Cilantro fresco", 3, 18.00, 10),
        ("Fumet de pescado concentrado", 25, 8.00, 0),
        ("Ají limo fresco (decoración)", 3, 32.00, 0),
        ("Choclo peruano cocido", 25, 12.00, 5),
        ("Camote dulce cocido", 30, 3.50, 10),
        ("Cebolla roja juliana", 15, 2.50, 12),
        ("Aceite de sésamo (toque japonés)", 3, 28.00, 0),
        ("Salsa soja premium", 5, 18.00, 0),
    ]
    start_r = r
    for idx, (ing, cant, precio, merma) in enumerate(ingredients, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, ing, fill=inp_fill)
        sdc(ws, r, 3, cant, fill=inp_fill)
        sdc(ws, r, 4, precio, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=(C{r}/1000)*D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, merma / 100, fill=inp_fill, fmt=pct_fmt)
        sdc(ws, r, 7, f"=E{r}*(1+F{r})", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 2, "COSTE TOTAL PLATO", font=bld_font)
    sdc(ws, r, 7, f"=SUM(G{start_r}:G{r-2})", font=bld_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP sugerido (food cost 33%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}/0.33", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP con IVA (10%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}*1.10", font=frm_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 7)
    path = os.path.join(OUTPUT_DIR, "escandallo-maestro-nikkei.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_menu_engineering():
    wb = Workbook()
    instr_sheet(wb, "Menú Engineering Matrix — Restaurante Nikkei", [
        "Introduce los platos de tu carta con ventas y food cost.",
        "La matrix clasifica automáticamente cada plato.",
        "Stars = mantener. Plowhorses = subir precio. Puzzles = promocionar. Dogs = eliminar.",
    ])
    ws = wb.create_sheet("Matrix"); ws.sheet_properties.tabColor = "9C27B0"
    title_block(ws, "Menú Engineering Matrix — Restaurante Nikkei", 10)
    hdrs = ["#", "Plato", "PVP (€)", "Food Cost (€)", "% Food Cost", "Uds. Vendidas/Mes", "Margen Unit. (€)", "Margen Total (€)", "Popularidad", "Clasificación"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    widths = [5, 30, 10, 12, 10, 16, 12, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    platos = [
        ("Tiradito atún ají amarillo", 18.00, 4.50, 200),
        ("Tiradito corvina ají limo", 17.00, 4.00, 160),
        ("Ceviche nikkei pescado+langostinos", 19.00, 5.50, 140),
        ("Causa nikkei de atún", 14.00, 3.20, 170),
        ("Causa nikkei de pulpo al olivo", 16.00, 4.50, 90),
        ("Maki acevichado (8 piezas)", 15.00, 3.80, 210),
        ("Uramaki crunch nikkei (8)", 14.00, 3.50, 180),
        ("Nigiri atún al ají limo (2)", 11.00, 2.80, 120),
        ("Anticucho corazón con miso (3)", 12.00, 2.20, 150),
        ("Anticucho pulpo al olivo nikkei", 18.00, 5.50, 70),
        ("Chita a la parrilla robata", 28.00, 9.00, 60),
        ("Arroz chaufa de mariscos", 16.00, 4.50, 140),
        ("Tataki atún con ají amarillo", 20.00, 6.00, 80),
        ("Pisco Sour clásico", 12.00, 1.80, 320),
        ("Chilcano de pisco", 11.00, 1.60, 180),
    ]
    r += 1
    for idx, (plato, pvp, fc, uds) in enumerate(platos, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, plato, fill=inp_fill)
        sdc(ws, r, 3, pvp, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 4, fc, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}/C{r}", font=frm_font, fmt=pct_fmt)
        sdc(ws, r, 6, uds, fill=inp_fill)
        sdc(ws, r, 7, f"=C{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 8, f"=G{r}*F{r}", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "menu-engineering-matrix.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_ticket():
    wb = Workbook()
    instr_sheet(wb, "Calculadora Ticket Medio + Margen Sake/Whisky", [
        "Simula diferentes escenarios de ticket medio.",
        "Incluye pestaña separada para margen de barra de pisco, sake y coctelería nikkei.",
        "Rellena las celdas verdes con tus datos.",
    ])
    ws = wb.create_sheet("Ticket Medio"); ws.sheet_properties.tabColor = "00BCD4"
    title_block(ws, "Calculadora Ticket Medio — Restaurante Nikkei", 4)
    r = 4
    params = [
        ("% clientes que piden tiradito/ceviche para compartir", 0.75),
        ("Precio medio tiradito/ceviche (€)", 18.00),
        ("Precio medio causa o maki acevichado (€)", 15.00),
        ("Precio medio principal caliente (chaufa, anticucho, chita) (€)", 17.00),
        ("% clientes que piden postre", 0.25),
        ("Precio medio postre (€)", 8.00),
        ("% clientes que piden Pisco Sour / cóctel nikkei", 0.65),
        ("Precio medio Pisco Sour / cóctel (€)", 11.00),
        ("% clientes que piden cerveza/refresco", 0.50),
        ("Precio medio cerveza/refresco (€)", 5.00),
    ]
    r += 1
    for label, val in params:
        sdc(ws, r, 1, label, font=dat_font)
        sdc(ws, r, 2, val, fill=inp_fill, fmt=pct_fmt if isinstance(val, float) and val < 1 else cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 1, "TICKET MEDIO ESTIMADO", font=bld_font)
    sdc(ws, r, 2, None, font=bld_font, fmt=cur_fmt)
    ws.column_dimensions['A'].width = 50; ws.column_dimensions['B'].width = 18

    # Margen pisco/sake tab
    ws2 = wb.create_sheet("Margen Pisco-Sake"); ws2.sheet_properties.tabColor = "FFC107"
    title_block(ws2, "Simulador Margen Barra de Pisco, Sake y Coctelería Nikkei", 4)
    r = 4
    sdc(ws2, r+1, 1, "Precio medio Pisco Sour / cóctel (€)", font=bld_font); sdc(ws2, r+1, 2, 11.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+2, 1, "Coste medio Pisco Sour (€)", font=bld_font); sdc(ws2, r+2, 2, 1.80, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+3, 1, "Margen por Pisco Sour (€)", font=bld_font); sdc(ws2, r+3, 2, f"=B{r+1}-B{r+2}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+4, 1, "% Margen Pisco Sour", font=bld_font); sdc(ws2, r+4, 2, f"=B{r+3}/B{r+1}", font=frm_font, fmt=pct_fmt)
    sdc(ws2, r+5, 1, "Precio medio copa sake premium (€)", font=bld_font); sdc(ws2, r+5, 2, 13.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+6, 1, "Coste medio copa sake (€)", font=bld_font); sdc(ws2, r+6, 2, 3.50, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+7, 1, "Margen por copa sake (€)", font=bld_font); sdc(ws2, r+7, 2, f"=B{r+5}-B{r+6}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+8, 1, "Pisco Sours vendidos/día", font=bld_font); sdc(ws2, r+8, 2, 55, fill=inp_fill)
    sdc(ws2, r+9, 1, "Copas sake vendidas/día", font=bld_font); sdc(ws2, r+9, 2, 20, fill=inp_fill)
    sdc(ws2, r+10, 1, "Días/mes", font=bld_font); sdc(ws2, r+10, 2, 26, fill=inp_fill)
    sdc(ws2, r+11, 1, "Facturación barra pisco/sake/mes (€)", font=bld_font)
    sdc(ws2, r+11, 2, f"=(B{r+1}*B{r+8}+B{r+5}*B{r+9})*B{r+10}", font=frm_font, fmt=cur_fmt)
    ws2.column_dimensions['A'].width = 42; ws2.column_dimensions['B'].width = 18
    path = os.path.join(OUTPUT_DIR, "calculadora-ticket-medio.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cronograma():
    wb = Workbook()
    instr_sheet(wb, "Cronograma Apertura Gantt 12 Meses — Restaurante Nikkei", [
        "Fases y tareas para abrir un restaurante nikkei.",
        "Marca las celdas con 'X' para indicar el mes activo.",
        "Incluye fase de búsqueda de jefe de cocina nikkei y proveedores (pescado sashimi-grade + ajíes peruanos + pisco).",
    ])
    ws = wb.create_sheet("Gantt"); ws.sheet_properties.tabColor = "795548"
    title_block(ws, "Cronograma de Apertura — Restaurante Nikkei 60 Plazas", 15)
    hdrs = ["#", "Fase", "Tarea", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "M8", "M9", "M10", "M11", "M12"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 15)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 45
    for i in range(4, 16): ws.column_dimensions[get_column_letter(i)].width = 5
    tasks = [
        ("Planificación", "Estudio de mercado y viabilidad", [1, 2]),
        ("Planificación", "Plan financiero y búsqueda financiación", [1, 2, 3]),
        ("Planificación", "Constitución empresa / alta autónomo", [2, 3]),
        ("Jefe cocina", "Búsqueda y reclutamiento jefe de cocina nikkei (crítico)", [2, 3, 4, 5]),
        ("Importación", "Contacto importadores ajíes peruanos + distribuidores japoneses", [2, 3]),
        ("Proveedores", "Selección proveedor pescado sashimi-grade diario + pisco", [3, 4]),
        ("Local", "Búsqueda y selección de local", [2, 3, 4]),
        ("Local", "Negociación alquiler y contrato", [3, 4]),
        ("Licencias", "Licencia actividad / declaración responsable", [3, 4, 5]),
        ("Licencias", "Proyecto técnico y visado", [3, 4]),
        ("Obra", "Acondicionamiento y obra civil", [4, 5, 6, 7]),
        ("Obra", "Instalaciones (electricidad, gas, extracción reforzada robata/wok)", [5, 6, 7]),
        ("Obra", "Decoración nikkei contemporánea (madera, cuelgas ajíes, cerámica)", [6, 7]),
        ("Equipamiento", "Pedido equipo cocina nikkei (arrocera, licuadoras, robata, wok, abatidor)", [5, 6]),
        ("Equipamiento", "Instalación barra tiraditos + vitrina pescado", [6, 7]),
        ("Equipamiento", "Recepción e instalación equipos", [7, 8]),
        ("Equipamiento", "Mobiliario sala + barra pisco-sake", [6, 7]),
        ("Equipo", "Selección cocinero tiraditos + parrillero robata + bartender nikkei", [7, 8, 9]),
        ("Equipo", "Selección y contratación resto equipo", [8, 9]),
        ("Equipo", "Formación equipo (técnicas nikkei + pisco + sake + servicio)", [9, 10]),
        ("Marketing", "Branding, web, redes sociales", [6, 7, 8]),
        ("Marketing", "Fotos profesionales de platos (tiraditos, makis, anticuchos)", [9]),
        ("Marketing", "Campaña pre-apertura", [9, 10]),
        ("APPCC", "Plan APPCC + registro sanitario (pescado crudo + anisakis)", [7, 8]),
        ("Pre-apertura", "Soft opening (amigos, prensa, influencers)", [10, 11]),
        ("Pre-apertura", "Ajustes finales", [11]),
        ("Apertura", "INAUGURACIÓN", [12]),
        ("Post-apertura", "Seguimiento y ajustes", [12]),
    ]
    r += 1
    for idx, (fase, tarea, meses) in enumerate(tasks, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, fase); sdc(ws, r, 3, tarea)
        for m in meses:
            sdc(ws, r, m + 3, "X", font=bld_font, fill=sec_fill, align=c_align)
        r += 1
    r += 1; brand_footer(ws, r, 15)
    path = os.path.join(OUTPUT_DIR, "cronograma-apertura-gantt.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_plantilla_turnos():
    wb = Workbook()
    instr_sheet(wb, "Plantilla Turnos Brigada — Restaurante Nikkei", [
        "Cuadrante semanal para todo el equipo (cocina + sala).",
        "Rellena los turnos con: M (mañana), T (tarde), P (partido), L (libre).",
        "Incluye puestos específicos: jefe cocina nikkei, cocinero tiraditos, parrillero robata, wok cook, bartender nikkei.",
    ])
    ws = wb.create_sheet("Turnos Semana"); ws.sheet_properties.tabColor = "607D8B"
    title_block(ws, "Cuadrante Semanal — Restaurante Nikkei 60 Plazas", 10)
    hdrs = ["#", "Nombre", "Puesto", "Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 28
    for i in range(4, 11): ws.column_dimensions[get_column_letter(i)].width = 8
    dv = DataValidation(type="list", formula1='"M,T,P,L,V"', allow_blank=True)
    ws.add_data_validation(dv)
    staff = [
        ("Jefe de cocina nikkei", "P"), ("Cocinero barra tiraditos", "P"),
        ("Parrillero robata/josper", "T"), ("Wok cook (chaufa, yakimeshi)", "P"),
        ("Cocinero causas / guarniciones", "P"), ("Cocinero de partida", "P"),
        ("Preparación pescado", "M"), ("Ayudante cocina / office", "P"),
        ("Office", "P"),
        ("Encargado sala", "P"), ("Bartender nikkei (pisco/sake)", "P"),
        ("Camarero 1", "P"), ("Camarero 2", "P"), ("Camarero 3", "T"),
        ("Runner", "P"),
    ]
    r += 1
    for idx, (puesto, turno_default) in enumerate(staff, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, None, fill=inp_fill); sdc(ws, r, 3, puesto)
        for c in range(4, 11):
            sdc(ws, r, c, turno_default, fill=inp_fill, align=c_align)
            dv.add(ws.cell(row=r, column=c))
        r += 1
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "M=Mañana, T=Tarde, P=Partido, L=Libre, V=Vacaciones", font=brn_font)
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "plantilla-turnos-brigada.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 3. CHECKLISTS (6)
# ═══════════════════════════════════════════════════════════

def gen_checklist_legal():
    wb = Workbook(); ws = wb.active
    items = [
        ("Empresa", "Decidir forma jurídica (SL, autónomo, cooperativa)", "Socio/Gestor", 500),
        ("Empresa", "Constituir sociedad ante notario", "Socio/Gestor", 600),
        ("Empresa", "Inscripción en Registro Mercantil", "Gestor", 200),
        ("Empresa", "Obtener CIF provisional", "Gestor", 0),
        ("Empresa", "Alta IAE epígrafe 671.4 o similar", "Gestor", 0),
        ("Empresa", "Alta censal en Hacienda (modelo 036/037)", "Gestor", 0),
        ("Empresa", "Inscripción Seguridad Social empresa", "Gestor", 0),
        ("Empresa", "Abrir cuenta bancaria empresa", "Socio", 0),
        ("Licencias", "Licencia de actividad / declaración responsable", "Arquitecto", 2500),
        ("Licencias", "Proyecto técnico y visado colegial", "Arquitecto", 3500),
        ("Licencias", "Licencia de obras (si aplica)", "Arquitecto", 2000),
        ("Licencias", "Licencia de terraza", "Socio", 1500),
        ("Licencias", "Licencia de música / SGAE / AGEDI", "Socio", 500),
        ("Anisakis", "Compra congelador certificado -20°C (obligatorio para sushi/sashimi)", "Socio", 4500),
        ("Anisakis", "Documento protocolo congelación preventiva pescado crudo", "Jefe cocina", 0),
        ("Anisakis", "Registros de congelación por lote de pescado (archivo 12 meses)", "Itamae", 0),
        ("Anisakis", "Información al cliente en carta sobre congelación preventiva", "Socio", 0),
        ("Anisakis", "Certificados de proveedores pescado de acuicultura (excepción)", "Jefe cocina", 0),
        ("Importación", "Registro EORI (operador comercio exterior) si aplica", "Gestor", 0),
        ("Importación", "Certificado SOIVRE para alimentos importados extra-UE", "Gestor", 300),
        ("Importación", "Etiquetado en español de productos peruanos y japoneses importados", "Proveedor", 200),
        ("Importación", "Impuestos especiales bebidas alcohólicas (IIEE) pisco/sake", "Gestor", 0),
        ("Sanidad", "Registro General Sanitario (RGSEAA)", "Gestor", 0),
        ("Sanidad", "Plan APPCC documentado (adaptado pescado crudo + arroz sushi)", "Consultor", 2000),
        ("Sanidad", "Formación manipulador alimentos (equipo — especial pescado crudo)", "Consultor", 500),
        ("Sanidad", "Contrato empresa DDD (desratización)", "Socio", 600),
        ("Sanidad", "Control de aguas", "Socio", 200),
        ("Sanidad", "Plan de alérgenos (14 alérgenos + soja, sésamo)", "Jefe cocina", 0),
        ("Seguros", "Seguro responsabilidad civil MÍNIMO 600K€ (riesgo sanitario pescado crudo)", "Socio", 1200),
        ("Seguros", "Seguro de local (incendio, robo, daños)", "Socio", 700),
        ("Seguros", "Seguro de accidentes convenio", "Gestor", 400),
        ("Laboral", "Contratos de trabajo según convenio", "Gestor", 0),
        ("Laboral", "Contrato específico del jefe de cocina nikkei (cláusulas de propiedad intelectual recetas)", "Gestor", 0),
        ("Laboral", "Prevención de riesgos laborales", "SPA", 500),
        ("Laboral", "Calendario laboral y cuadrante turnos", "Encargado", 0),
        ("Protección datos", "Registro LOPD / RGPD", "Gestor", 300),
        ("Protección datos", "Política de cookies web", "Socio", 0),
        ("Protección datos", "Videovigilancia (si aplica): cartelería + registro", "Socio", 100),
        ("Otros", "Libro de reclamaciones", "Socio", 30),
        ("Otros", "Cartel de horarios en puerta", "Socio", 0),
        ("Otros", "Aforo máximo visible", "Socio", 0),
        ("Otros", "Señalética salida emergencia + extintores", "Arquitecto", 600),
    ]
    checklist_ws(ws, "Checklist Legal — Restaurante Nikkei España", items)
    path = os.path.join(OUTPUT_DIR, "checklist-legal.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_equipamiento():
    wb = Workbook(); ws = wb.active
    items = [
        ("Cocina NK", "Arrocera industrial profesional 5-10L (sushi + chaufa)", "Jefe cocina", 3200),
        ("Cocina NK", "Vitrina refrigerada de pescado (barra tiraditos)", "Jefe cocina", 6500),
        ("Cocina NK", "Licuadoras industriales alta potencia (leches de tigre, 2-3 uds)", "Jefe cocina", 2800),
        ("Cocina NK", "Batidora de inmersión profesional", "Jefe cocina", 600),
        ("Cocina NK", "Parrilla robata de carbón binchotan / Josper", "Parrillero", 8500),
        ("Cocina NK", "Wok station 2 woks alta potencia (chaufa, yakimeshi)", "Wok cook", 4500),
        ("Cocina NK", "Plancha teppanyaki (opcional, tatakis)", "Jefe cocina", 2500),
        ("Cocina NK", "Set cuchillos japoneses (yanagiba, deba, usuba, santoku)", "Jefe cocina", 2800),
        ("Cocina NK", "Tabla corte madera Hiba / sintética (2 uds)", "Jefe cocina", 500),
        ("Cocina NK", "Abatidor de temperatura certificado (crítico anisakis)", "Jefe cocina", 5500),
        ("Cocina NK", "Ahumador pequeño (opcional, tatakis ahumados)", "Jefe cocina", 1200),
        ("Cocina NK", "Vitrina refrigerada pisco/sake/vinos barra", "Bartender", 4500),
        ("Cocción", "Cocina gas 6 fuegos + horno", "Jefe cocina", 5000),
        ("Cocción", "Horno mixto (Rational / Unox)", "Jefe cocina", 12000),
        ("Cocción", "Freidora doble cuba (crocante camote, chicharrones)", "Jefe cocina", 2500),
        ("Cocción", "Salamandra", "Jefe cocina", 1200),
        ("Frío", "Cámara frigorífica 2 puertas (pescado fresco)", "Jefe cocina", 4000),
        ("Frío", "Cámara congelación CERTIFICADA -20°C (anisakis obligatorio)", "Jefe cocina", 4500),
        ("Frío", "Mesa refrigerada bajo barra tiraditos", "Jefe cocina", 3000),
        ("Frío", "Botellero refrigerado barra pisco-sake", "Bartender", 2000),
        ("Preparación", "Mesa de trabajo acero inox (2-3 uds)", "Jefe cocina", 2000),
        ("Preparación", "Robot de cocina / cutter", "Jefe cocina", 1500),
        ("Preparación", "Escamadora eléctrica de pescado", "Jefe cocina", 400),
        ("Preparación", "Báscula de precisión (leches de tigre)", "Jefe cocina", 200),
        ("Lavado", "Tren de lavado / lavavajillas industrial", "Jefe cocina", 5500),
        ("Lavado", "Fregadero doble seno inox", "Jefe cocina", 800),
        ("Extracción", "Campana extractora REFORZADA + filtros (robata + wok station)", "Instalador", 10000),
        ("Extracción", "Salida de humos reglamentaria (sobredimensionada)", "Instalador", 6000),
        ("Menaje", "Ollas grandes para fondos de pescado, arroces, guisos", "Jefe cocina", 1200),
        ("Menaje", "Wok adicional de reserva", "Jefe cocina", 300),
        ("Menaje", "Cuchillos adicionales y afiladores", "Jefe cocina", 800),
        ("Menaje", "Gastronormas GN 1/1, 1/2, 1/3", "Jefe cocina", 600),
        ("Menaje", "Tablas de corte por colores (separación crudo/cocinado)", "Jefe cocina", 150),
        ("Almacenamiento", "Estanterías inox economato seco (arroz, ajíes, salsas)", "Jefe cocina", 800),
        ("Almacenamiento", "Contenedores herméticos para arroz Koshihikari y ajíes", "Jefe cocina", 300),
        ("Barra", "Utensilios coctelería nikkei (boston shaker, jiggers, dispenser clara de huevo)", "Bartender", 900),
        ("Barra", "Máquina de hielo (alto consumo por coctelería)", "Bartender", 2200),
    ]
    checklist_ws(ws, "Checklist Equipamiento Cocina Nikkei", items)
    path = os.path.join(OUTPUT_DIR, "checklist-equipamiento-cocina-nikkei.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_appcc():
    wb = Workbook(); ws = wb.active
    items = [
        ("Prerrequisitos", "Plan de limpieza y desinfección documentado", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan DDD contratado (empresa externa)", "Socio", 600),
        ("Prerrequisitos", "Plan de control de agua potable", "Socio", 200),
        ("Prerrequisitos", "Plan de formación del personal (especial pescado crudo)", "Jefe cocina", 500),
        ("Prerrequisitos", "Plan de gestión de residuos (incluye aceite de fritura)", "Socio", 0),
        ("Prerrequisitos", "Plan de mantenimiento de instalaciones", "Socio", 0),
        ("Prerrequisitos", "Plan de trazabilidad (registro proveedores + lotes)", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan de control de alérgenos (14 + soja, sésamo)", "Jefe cocina", 0),
        ("Anisakis (PCC)", "CONGELACIÓN PREVIA obligatoria -20°C/24h todo pescado crudo (tiradito/ceviche/sashimi)", "Jefe cocina", 0),
        ("Anisakis (PCC)", "Congelador/abatidor certificado con registro de temperatura automático", "Jefe cocina", 5500),
        ("Anisakis (PCC)", "Registro de congelación por lote (hora entrada/salida/temperatura)", "Jefe cocina", 0),
        ("Anisakis (PCC)", "Archivo de registros mínimo 12 meses", "Jefe cocina", 0),
        ("Anisakis (PCC)", "Certificado proveedor para pescado de acuicultura (excepción)", "Jefe cocina", 0),
        ("Anisakis (PCC)", "Información al cliente en carta: 'pescado congelado preventivamente'", "Socio", 0),
        ("Leche de tigre (PCC)", "Preparación diaria en recipiente estéril, descarte a 24h", "Cocinero tiraditos", 0),
        ("Leche de tigre (PCC)", "Conservación 0-4°C con registro temperatura cada 2h", "Cocinero tiraditos", 0),
        ("Arroz sushi (PCC)", "Temperatura shari entre 20-25°C (no refrigerar, no calentar)", "Jefe cocina", 0),
        ("Arroz sushi (PCC)", "Máxima permanencia 4 horas en recipiente cerrado", "Jefe cocina", 0),
        ("Arroz sushi (PCC)", "Registro de temperatura cada 2 horas", "Cocinero", 0),
        ("Arroz sushi (PCC)", "Acidificación con sushi-zu (vinagre + azúcar + sal) verificada", "Jefe cocina", 0),
        ("Temperaturas", "Termómetro sonda calibrado", "Jefe cocina", 50),
        ("Temperaturas", "Registro diario temperaturas cámaras", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas servicio (caliente >65°C)", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas recepción mercancía pescado", "Jefe cocina", 0),
        ("Temperaturas", "Control cadena de frío (0-4°C refrigerado, -20°C congelado)", "Jefe cocina", 0),
        ("Temperaturas", "Control temperatura fondos y guisos (enfriamiento rápido <4°C)", "Jefe cocina", 0),
        ("Pescado crudo", "Descongelación SIEMPRE en cámara 0-4°C (nunca a temp. ambiente)", "Jefe cocina", 0),
        ("Pescado crudo", "Trazabilidad completa del pescado (proveedor, lonja, fecha)", "Jefe cocina", 0),
        ("Pescado crudo", "Vitrina de pescado a 0-4°C con control continuo", "Jefe cocina", 0),
        ("Pescado crudo", "Rotación pescado expuesto en vitrina cada 2-3 horas", "Cocinero tiraditos", 0),
        ("Pescado crudo", "Cuchillos y tablas exclusivos para pescado crudo", "Jefe cocina", 0),
        ("Prod. Importados", "Etiquetado en español (ajíes peruanos, sake, soja, mirin, panko)", "Jefe cocina", 0),
        ("Prod. Importados", "Trazabilidad de origen productos peruanos y japoneses", "Jefe cocina", 0),
        ("Prod. Importados", "Almacenamiento separado productos secos (arroz, ajíes)", "Cocinero", 0),
        ("Prod. Importados", "Registro de lotes de pisco/sake importado", "Bartender", 0),
        ("Aceite fritura", "Filtrado aceite diario, registro de cambio por polaridad", "Cocinero", 0),
        ("Aceite fritura", "Temperatura aceite controlada 170-180°C (crocantes, chicharrones)", "Cocinero", 0),
        ("Higiene", "Lavamanos no manual en cocina con jabón y papel", "Socio", 300),
        ("Higiene", "Vestuarios con taquillas para personal", "Socio", 500),
        ("Higiene", "Uniformes cocina (jefe cocina nikkei + equipo)", "Socio", 800),
        ("Higiene", "Protocolo lavado de manos (cartelería)", "Jefe cocina", 0),
        ("Higiene", "Contenedores de basura con tapa y pedal", "Socio", 200),
        ("Almacenamiento", "Sistema FIFO/FEFO en cámaras y economato", "Cocinero", 0),
        ("Almacenamiento", "Separación crudo/cocinado en cámaras (especialmente pescado)", "Cocinero", 0),
        ("Documentación", "Plan APPCC escrito y disponible en cocina", "Socio", 2000),
        ("Documentación", "Registros de control actualizados diariamente", "Jefe cocina", 0),
        ("Documentación", "Certificados manipulador alimentos de todo el personal", "Socio", 0),
        ("Documentación", "Contrato DDD visible y registros de actuaciones", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist APPCC — Restaurante Nikkei (Tiradito + Ceviche + Anisakis)", items)
    path = os.path.join(OUTPUT_DIR, "checklist-appcc.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_sala():
    wb = Workbook(); ws = wb.active
    items = [
        ("Distribución", "Plano de sala con posición de mesas y barra sushi definitivo", "Socio/Arquitecto", 0),
        ("Distribución", "60 plazas distribuidas (barra tiraditos 10-14 + mesas)", "Socio", 0),
        ("Distribución", "Distancia entre mesas: mínimo 70-90 cm (confort premium)", "Arquitecto", 0),
        ("Distribución", "Flujo de servicio definido (entrada → sala → barra → aseos)", "Socio", 0),
        ("Distribución", "Zona de espera / recepción acogedora", "Socio", 800),
        ("Distribución", "Accesibilidad PMR (rampa, aseo adaptado)", "Arquitecto", 2000),
        ("Barra Tiraditos", "Barra profesional madera + vitrina refrigerada de pescado delantera", "Interiorista", 10000),
        ("Barra Tiraditos", "Taburetes altos cómodos (10-14 plazas)", "Socio", 3500),
        ("Barra Tiraditos", "Iluminación dirigida al cocinero (teatro en vivo)", "Electricista", 1200),
        ("Barra Tiraditos", "Carta de tiraditos / omakase nikkei enmarcada", "Diseñador", 400),
        ("Barra Pisco-Sake", "Vitrina iluminada trasera para pisco, sake y vinos", "Interiorista", 5000),
        ("Barra Pisco-Sake", "Iluminación ambiental barra (retroiluminación botellas)", "Electricista", 1000),
        ("Decoración NK", "Revestimientos de madera natural en paredes de acento", "Interiorista", 8000),
        ("Decoración NK", "Cuelgas de ajíes peruanos secos (amarillo, panca, rocoto)", "Socio", 400),
        ("Decoración NK", "Textiles andinos sutiles como acento", "Interiorista", 1500),
        ("Decoración NK", "Paleta: madera, blanco roto, negro mate, tierra, cobre", "Interiorista", 0),
        ("Decoración NK", "Vajilla cerámica artesanal mixta (Perú + Japón)", "Socio", 4500),
        ("Decoración NK", "Fotografías artísticas (mercados Lima, pescadores, cocineros)", "Socio", 1200),
        ("Decoración NK", "Detalles piedra natural y cobre (pared acento, barra)", "Interiorista", 2500),
        ("Iluminación", "Iluminación general cálida 2400-2700K (tenue)", "Electricista", 2500),
        ("Iluminación", "Lámparas de papel/fibras naturales sobre mesas", "Socio", 2500),
        ("Iluminación", "Iluminación regulable por zonas (dimmer)", "Electricista", 600),
        ("Acústica", "Paneles absorbentes para ambiente confortable", "Interiorista", 2000),
        ("Acústica", "Sistema de música (bossa, jazz latino, cumbia peruana, ambient, volumen bajo)", "Socio", 800),
        ("Mobiliario", "Mesas de madera natural maciza (roble/haya/algarrobo)", "Socio", 14000),
        ("Mobiliario", "Sillas tapizadas en tonos neutros", "Socio", 8000),
        ("Mobiliario", "Zona privada omakase nikkei (si aplica)", "Interiorista", 3500),
        ("Terraza", "Mobiliario exterior en madera natural", "Socio", 5000),
        ("Aseos", "Aseos diferenciados y aseo PMR", "Arquitecto", 3500),
        ("Aseos", "Decoración nikkei en aseos (madera, detalles cálidos)", "Socio", 600),
        ("Señalética", "Señalética interior con identidad nikkei contemporánea", "Diseñador", 400),
        ("Señalética", "Rótulo exterior / fachada con identidad nikkei sutil", "Socio", 3000),
    ]
    checklist_ws(ws, "Checklist Diseño Sala Nikkei + Barra Tiraditos + Barra Pisco-Sake", items)
    path = os.path.join(OUTPUT_DIR, "checklist-diseno-sala-nikkei.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_contratacion():
    wb = Workbook(); ws = wb.active
    items = [
        ("Perfiles", "Definir organigrama completo (cocina + sala)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Jefe de cocina nikkei (paso por Maido/Pakta/Osaka/Kikko ideal)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Cocinero barra tiraditos (corte + leches de tigre)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Parrillero robata/josper (anticuchos, chita)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Wok cook (chaufa, yakimeshi, saltados)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Cocinero causas / guarniciones / postres", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Cocinero de partida", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Preparación pescado / escamado", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Ayudante / office", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Encargado sala", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Bartender nikkei (pisco + sake + coctelería)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Camarero (con conocimiento pisco/sake)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Runner", "Socio", 0),
        ("Selección", "Publicar ofertas (InfoJobs, Hostelwork, LinkedIn)", "Socio", 300),
        ("Selección", "Búsqueda en comunidad peruana España (Madrid, Barcelona)", "Socio", 0),
        ("Selección", "Búsqueda jefe cocina internacional (visado si procede)", "Socio", 2000),
        ("Selección", "Contacto con ex-empleados Pakta, Osaka, Kikko, Maizal, Totora", "Socio", 0),
        ("Selección", "Búsqueda bartender certificado en coctelería peruana (Escuela Nacional de Bartenders Perú)", "Socio", 0),
        ("Selección", "Selección resto equipo cocina y sala", "Jefe cocina", 0),
        ("Selección", "Pruebas prácticas cocina (elaborar tiradito, leche de tigre, anticucho, maki acevichado)", "Jefe cocina", 0),
        ("Documentación", "Contratos según convenio hostelería CCAA", "Gestor", 0),
        ("Documentación", "Contrato específico jefe cocina nikkei (cláusulas premium)", "Gestor", 200),
        ("Documentación", "Alta Seguridad Social", "Gestor", 0),
        ("Documentación", "Certificado manipulador alimentos (especial pescado crudo)", "Empleado", 0),
        ("Documentación", "Prevención riesgos laborales (formación)", "SPA", 500),
        ("Onboarding", "Manual de bienvenida / handbook del restaurante", "Socio", 0),
        ("Onboarding", "Formación en carta nikkei y alérgenos (soja, sésamo, gluten, pescado)", "Jefe cocina", 0),
        ("Onboarding", "Formación en pisco, sake y coctelería nikkei (todo equipo sala)", "Bartender", 0),
        ("Onboarding", "Formación en TPV y sistema de pedidos", "Encargado", 0),
        ("Onboarding", "Formación APPCC (especial pescado crudo + anisakis + leche de tigre)", "Jefe cocina", 0),
        ("Onboarding", "Formación servicio nikkei (historia de la fusión, relato al cliente)", "Encargado", 0),
        ("Onboarding", "Ensayos de servicio (soft opening)", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist Contratación — Restaurante Nikkei", items)
    path = os.path.join(OUTPUT_DIR, "checklist-contratacion.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_marketing():
    wb = Workbook(); ws = wb.active
    items = [
        ("Branding", "Nombre del restaurante definitivo (verificar disponibilidad)", "Socio", 0),
        ("Branding", "Logo profesional con identidad nikkei contemporánea", "Diseñador", 700),
        ("Branding", "Manual de marca (paleta neutra + tierras cálidas, tipografía, estilo)", "Diseñador", 400),
        ("Branding", "Diseño de carta / menú físico con estética nikkei", "Diseñador", 700),
        ("Branding", "Carta específica de piscos, sakes y coctelería nikkei", "Diseñador", 400),
        ("Branding", "Packaging delivery con marca (cajas maki premium + botes ajíes)", "Diseñador", 500),
        ("Digital", "Web propia con menú, carta piscos/sakes, fotos y reservas", "Dev/Agencia", 2500),
        ("Digital", "Sistema de reservas online con depósito (omakase nikkei)", "Dev/Agencia", 800),
        ("Digital", "Google My Business: ficha completa y verificada", "Socio", 0),
        ("Digital", "Instagram: perfil profesional con 12+ publicaciones premium", "Community", 0),
        ("Digital", "TikTok: perfil activo con vídeos del cocinero, tiraditos, robata, Pisco Sour", "Community", 0),
        ("Digital", "TheFork / TripAdvisor: perfiles completos con fotos", "Socio", 0),
        ("Digital", "Google Ads local configurado", "Agencia", 700),
        ("Digital", "Instagram/Facebook Ads configurados", "Agencia", 500),
        ("Contenido", "Sesión fotográfica profesional (tiraditos, makis, anticuchos, decoración)", "Fotógrafo", 600),
        ("Contenido", "Vídeo del cocinero bañando tiradito con leche de tigre (Reels/TikTok)", "Videógrafo", 700),
        ("Contenido", "Vídeo preparación de Pisco Sour con emulsión de clara", "Videógrafo", 500),
        ("Contenido", "Calendario de publicaciones primer mes", "Community", 0),
        ("Eventos NK", "Plan Día del Pisco Sour (1er sábado febrero): vuelos + masterclass", "Socio", 700),
        ("Eventos NK", "Plan Fiestas Patrias Perú (28 julio): menú especial + Pisco Sour", "Socio", 800),
        ("Eventos NK", "Plan Día Nacional del Ceviche (28 junio): cata leches de tigre", "Socio", 500),
        ("Eventos NK", "Plan Día Nacional del Pisco (4º domingo julio): cata de piscos", "Socio", 500),
        ("Eventos NK", "Plan Día Gastronomía Peruana / Mistura España (septiembre)", "Socio", 400),
        ("Eventos NK", "Eventos cata pisco + sake maridados con omakase nikkei", "Bartender", 400),
        ("Pre-apertura", "Evento soft opening (prensa gastronómica, influencers, críticos)", "Socio", 800),
        ("Pre-apertura", "Invitación a prensa especializada (guías Michelin, Repsol, medios peruanos)", "Socio", 300),
        ("Pre-apertura", "Colaboraciones con embajada del Perú y comunidad peruana en España", "Socio", 0),
        ("Pre-apertura", "Flyers / cartelería barrio", "Diseñador", 300),
        ("Pre-apertura", "Colaboraciones con negocios vecinos premium", "Socio", 0),
        ("Delivery", "Alta en Glovo / Uber Eats / Just Eat", "Socio", 0),
        ("Delivery", "Fotos profesionales de maki boxes y causas bowls para plataformas", "Fotógrafo", 0),
        ("Delivery", "Menú delivery optimizado (maki boxes, causas, chaufa, anticuchos)", "Jefe cocina", 0),
        ("Fidelización", "Programa fidelización (tarjeta, app o CRM — omakase VIP)", "Socio", 500),
        ("Fidelización", "Estrategia de reseñas (pedir valoraciones a clientes)", "Encargado", 0),
    ]
    checklist_ws(ws, "Checklist Marketing Pre-Apertura — Restaurante Nikkei", items)
    path = os.path.join(OUTPUT_DIR, "checklist-marketing-preapertura.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 4. DOCUMENT MODELS (2)
# ═══════════════════════════════════════════════════════════

def gen_business_plan():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Business Plan"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Nikkei — 60 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("[Ciudad, Mes Año]\nPlantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Resumen Ejecutivo", [
            "Descripción del proyecto: restaurante nikkei auténtico de 60 plazas en [ciudad].",
            "Modelo: [cevichería nikkei / casual premium / omakase / robata / pisco bar / dark kitchen].",
            "Inversión total requerida: [280.000-520.000€].",
            "Facturación prevista año 1: [750.000-1.200.000€].",
            "Break-even estimado: mes [10-16].",
            "Equipo fundador: [nombres y experiencia].",
            "Diferencial: jefe de cocina nikkei cualificado + pescado sashimi-grade + ajíes peruanos reales + barra de pisco/sake/coctelería nikkei 40-60 referencias.",
        ]),
        ("2. El Concepto", [
            "Tipo de cocina: nikkei auténtico fusión peruano-japonesa (tiraditos, causas, anticuchos, chaufa, robata).",
            "Ticket medio objetivo: [45-70€] casual / [90-160€] omakase nikkei, sin bebidas.",
            "Propuesta de valor: jefe de cocina con paso por referentes nikkei (Maido, Pakta, Osaka, Kikko), pescado diario sashimi-grade, ajíes peruanos frescos, coctelería Pisco Sour.",
            "Público objetivo: 30-55 años, urbano, foodie, alto poder adquisitivo, comunidad peruana/latinoamericana en España.",
            "Posicionamiento vs competencia: auténtico nikkei peruano-japonés vs japonés genérico o peruano tradicional.",
        ]),
        ("3. Análisis de Mercado", [
            "Restaurantes nikkei auténticos en España: ~120 (2025), nicho en fuerte crecimiento.",
            "Referentes españoles: Pakta (Barcelona), Osaka (Madrid/Barcelona/Marbella), Kikko (Madrid), Maizal, Totora.",
            "Tendencia pisco y coctelería nikkei: crecimiento acelerado.",
            "Omakase nikkei de alta gama con listas de espera en Madrid y Barcelona.",
            "Competencia directa en zona: [listar 3-5 competidores con ticket medio].",
            "Gap de mercado identificado: [qué falta en la zona].",
        ]),
        ("4. Plan Operativo", [
            "Ubicación: [dirección, m², alquiler mensual].",
            "Distribución: 60 plazas interior + barra tiraditos 10-14 plazas + [X] plazas terraza.",
            "Horario: [ej. 13:00-16:00 y 20:00-00:00, lunes cerrado].",
            "Equipo cocina: [8-12] personas (incluye jefe nikkei, cocinero tiraditos, parrillero robata, wok cook).",
            "Equipo sala: [5-8] personas (incluye bartender nikkei especialista en pisco y sake).",
            "Proveedores clave: [pescado sashimi-grade, importadores ajíes peruanos (Inkawasi, Latin Products), distribuidores japoneses, pisco (Viñas de Oro, Pisco 1615)].",
        ]),
        ("5. Plan Financiero", [
            "Inversión inicial desglosada: ver plantilla Excel adjunta.",
            "P&L mensual con 3 escenarios: ver plantilla Excel adjunta.",
            "Cash flow 12 meses: ver plantilla Excel adjunta.",
            "Break-even: [mes estimado].",
            "ROI estimado: [X%] en año [2-3].",
        ]),
        ("6. Plan de Marketing", [
            "Estrategia pre-apertura: [acciones 3 meses antes de abrir].",
            "Canales digitales: Google My Business, Instagram, TikTok, TheFork.",
            "Eventos culturales: Día del Pisco Sour, Fiestas Patrias Perú, Día del Ceviche, Día de la Gastronomía Peruana.",
            "Presupuesto marketing mensual: [800-2.000€/mes].",
            "Estrategia de delivery: maki boxes acevichados, causas bowls, chaufa y anticuchos.",
        ]),
        ("7. Equipo Fundador", [
            "[Nombre]: [experiencia en hostelería/cocina nikkei, rol en el proyecto].",
            "[Nombre]: [experiencia complementaria, rol en el proyecto].",
            "Jefe de cocina nikkei contratado: [experiencia, paso por referentes (Maido, Pakta, Osaka, Kikko), años en el oficio].",
            "Asesores externos: [gestor, arquitecto, consultor gastronómico].",
        ]),
        ("8. Necesidades de Financiación", [
            "Inversión total: [X€].",
            "Fondos propios: [X€] ([X%]).",
            "Financiación solicitada: [X€] ([X%]).",
            "Destino fondos: obra [X€], equipamiento nikkei [X€], barra pisco-sake [X€], maniobra [X€].",
            "Plan de amortización propuesto: [X años, cuota mensual X€].",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "business-plan-modelo.docx")
    doc.save(path); print(f"✓ {path}")

def gen_manual_operaciones():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual de Operaciones"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Nikkei — 60 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Plantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Apertura del Restaurante", [
            "Hora de llegada del equipo: [3 horas antes del servicio].",
            "Jefe de cocina nikkei: recibir pescado fresco del día, verificar calidad sashimi-grade, revisar registros de congelación preventiva (anisakis) del pescado del día anterior.",
            "Cocinero barra tiraditos: preparar leches de tigre del día (ají amarillo, ají limo, clásica) en licuadora industrial con todos los ingredientes frescos; cocer arroz Koshihikari en arrocera industrial y aderezar con sushi-zu para los makis acevichados.",
            "Parrillero robata: encender parrilla binchotan o Josper (40 min para alcanzar temperatura), preparar brochetas de anticucho (corazón marinado con miso y ají panca), pulpo, chita.",
            "Wok cook: preparar mise en place de arroz chaufa (arroz cocido, langostinos, vieiras, ají amarillo, sillao, huevo, cebolla china), verificar estado de los woks.",
            "Cocinero de causas y guarniciones: preparar puré de papa amarilla con ají amarillo para causas, cocinar choclo, camote, yuca; preparar salsa acevichada, salsa nikkei dressing y ponzu nikkei.",
            "Verificar temperaturas de cámaras y REGISTRAR. Verificar congelador/abatidor certificado -20°C para anisakis.",
            "Barra pisco-sake: verificar stock de piscos (Quebranta, Italia, acholados), sakes y vinos; preparar emulsión base de Pisco Sour (jarabe goma + limón); enfriar sakes.",
            "Preparar sala: colocar mesas, cartas, cubiertos, servilletas.",
            "Verificar limpieza de aseos.",
            "Encender TPV, música ambiente (bossa, cumbia peruana, jazz latino a volumen bajo), iluminación tenue.",
            "Briefing de equipo: tiraditos del día según pescado recibido, pisco o sake destacado, reservas omakase, alérgenos.",
        ]),
        ("2. Servicio de Mediodía", [
            "Horario: 13:00-16:00.",
            "Mediodía: público de business lunch, menú nikkei corto, tiraditos para profesionales.",
            "Ofrecer menú del día nikkei (si aplica) + carta completa + omakase almuerzo.",
            "Velocidad de servicio: tiradito en mesa en 8-10 min, causa en 10 min, principales calientes en 15 min.",
            "Cocinero barra tiraditos: ritmo constante, baña el tiradito al momento frente al cliente en la barra.",
            "Parrillero robata: anticuchos al momento, nunca pre-cocinar.",
            "Gestión de turnos de mesa si hay espera (notificar tiempo estimado).",
            "Cobro y despedida: sugerir postre (suspiro limeño con yuzu, mochi de lúcuma) y café pasado peruano.",
        ]),
        ("3. Servicio de Cena", [
            "Horario: 20:00-00:00.",
            "Carta completa + omakase nikkei + carta destacada de piscos y sakes.",
            "Iluminación más tenue, música ambient/latina, volumen bajo.",
            "Sugerir omakase (si disponible), tiraditos para compartir, vuelos de Pisco Sour.",
            "Cocinero barra tiraditos: servicio omakase en barra (si aplica). Interacción con cliente.",
            "Último pedido de cocina: [23:30].",
            "Barra robata activa durante toda la cena — plato emblema (chita, anticuchos, pulpo).",
        ]),
        ("4. Barra de Tiraditos: Protocolo Diario", [
            "Verificar que TODO el pescado para tiradito/ceviche/sashimi fue previamente congelado -20°C/24h (registro OBLIGATORIO — anisakis).",
            "Descongelar pescado en cámara (0-4°C) la noche anterior. NUNCA a temperatura ambiente ni con agua.",
            "Preparar las leches de tigre del día en licuadora industrial: leche de tigre clásica, leche de tigre al ají amarillo, leche de tigre al ají limo. Base: fumet de pescado + limón + ajo + jengibre + cilantro + ají + sal.",
            "Conservar leches de tigre a 0-4°C en recipientes estériles; registro cada 2 horas; descarte a las 24h.",
            "Montar vitrina de pescado del día, decorado con ajíes frescos, cilantro, choclo, camote.",
            "Cortar pescado al momento del pedido: láminas finas de tiradito, dados precisos para ceviche, láminas para nigiris nikkei.",
            "Bañar el tiradito con la leche de tigre SOLO al momento de servir (el ácido cocina el pescado en minutos).",
            "Montar causas: base de puré de papa amarilla con ají amarillo + topping de pescado/mariscos + decoración con palta, ají limo, salsa acevichada.",
            "Limpieza continua de la estación: cuchillos yanagiba, tabla, superficies. Protocolo de higiene estricto por pescado crudo y ajíes frescos.",
            "Controlar temperatura del arroz sushi en recipiente cerrado cada 2 horas (registro APPCC).",
            "Vida útil máxima del arroz sushi: 4 horas. Desechar sobrante y preparar nuevo lote si servicio largo.",
        ]),
        ("5. Estación Robata y Wok Station: Protocolo Diario", [
            "Robata: encender binchotan o Josper 40 min antes del servicio. Temperatura de trabajo 700-900°C.",
            "Preparar anticuchos: corazón de ternera limpio y cortado, marinado con miso, soja, ají panca, ajo, vinagre.",
            "Preparar chita/corvina entera: escamada, limpia, marcada al grill con toque de soja, yuzu y shichimi.",
            "Pulpo al olivo nikkei: pulpo pre-cocido, sellado en robata, salsa olivo con mayonesa nikkei.",
            "Wok station: woks siempre a máxima potencia antes del servicio. Arroz chaufa preparado al momento (nunca pre-hecho).",
            "Tiempo máximo montaje chaufa: 4-5 minutos. Yakimeshi nikkei similar.",
            "Limpieza del wok tras cada servicio intensivo (sin jabón, solo calor y trapo).",
            "Enfriamiento rápido al cierre de cualquier preparación que sobre (<4°C en <2h).",
        ]),
        ("6. Barra de Pisco y Sake: Protocolo", [
            "Verificar stock de piscos (12-16 refs), sakes (10-14 refs) y vinos (8-12 refs).",
            "Preparar base de Pisco Sour del día: jarabe de goma + limón fresco exprimido (nunca embotellado). Emulsionar clara de huevo correctamente.",
            "Enfriar sakes fríos (Junmai Ginjo, Daiginjo, Nigori) a 8-12°C.",
            "Preparar hielo abundante (coctelería nikkei consume mucho).",
            "Verificar stock garnish: limón peruano, yuzu, shiso, hierba luisa, cáscara de lima.",
            "Sugerir vuelos de Pisco Sour (3 versiones: clásico, maracuyá, ají limo) por 18-22€.",
            "Sugerir vuelos de sake (junmai + ginjo + daiginjo) por 22-28€.",
            "Explicar tipos de uva pisquera (Quebranta neutral, Italia floral, Torontel aromático).",
            "Servir con protocolo: presentar botella de pisco premium al cliente, servir al momento la coctelería.",
            "Registrar consumo de botellas para control de stock (piscos premium y sakes en control).",
            "Al cierre: limpiar barra, guardar sakes abiertos según su categoría, cerrar botellas.",
        ]),
        ("7. Cierre del Restaurante", [
            "Cocina: apagar wok station, robata (dejar enfriar binchotan de forma segura), hornos, freidora, fogones.",
            "Barra tiraditos: DESECHAR todo el arroz sushi sobrante y la leche de tigre sobrante de 24h, limpiar vitrina de pescado, guardar pescado no utilizado en cámara con trazabilidad.",
            "Verificar y registrar temperaturas de cámaras y congelador/abatidor -20°C.",
            "Guardar salsas (salsa acevichada, ponzu nikkei, nikkei dressing, pasta de ají amarillo) en cámara con fecha y hora.",
            "Wok station: limpiar woks (solo calor y trapo), guardar.",
            "Sala: limpiar mesas, barrer, recoger terraza.",
            "Barra: cerrar caja, cuadrar TPV, limpiar barra, guardar piscos premium bajo llave.",
            "Sacar basuras a contenedores (separado: aceite de fritura para recogida).",
            "Cerrar puertas, alarma, luces.",
        ]),
        ("8. Protocolo Anisakis (PCC Crítico)", [
            "La normativa RD 1420/2006 obliga a congelar preventivamente TODO pescado destinado a consumo crudo.",
            "Temperatura: -20°C durante mínimo 24 horas (o -35°C durante 15 horas).",
            "Equipamiento: congelador certificado que alcance y mantenga -20°C real (no un congelador doméstico).",
            "Registro OBLIGATORIO por lote: especie, proveedor, fecha entrada congelador, fecha salida, temperatura.",
            "Archivo de registros: mínimo 12 meses disponibles para inspección sanitaria.",
            "Excepción: pescado de acuicultura con certificado del proveedor garantizando ausencia de riesgo.",
            "Información al cliente: obligación de informar en carta (nota pie de página).",
            "NUNCA servir pescado para sashimi/sushi sin haber pasado por el protocolo — responsabilidad civil y penal.",
            "En caso de inspección: mostrar congelador, registros y etiquetas de lotes activos.",
            "Un caso de anisakiasis positivo puede cerrar tu restaurante y generar responsabilidad civil millonaria.",
        ]),
        ("9. Gestión de Delivery Nikkei", [
            "Zona de packaging separada del pase de sala y de la barra de tiraditos.",
            "Maki boxes acevichados: producto estrella. Caja negra mate con compartimentos, salsa acevichada aparte, ají limo, palillos.",
            "Causa bowls: envases con tapa hermética, causa + topping + decoración.",
            "Tiradito box (<15 min entrega): pescado laminado en envase principal + leche de tigre en cápsula separada. Cliente vierte al abrir. Instrucciones en tapa.",
            "Arroz chaufa de mariscos: envases con tapa, calor mantenido 20-25 min.",
            "Anticuchos con arroz: envases con tapa, brocheta + arroz + salsa aparte.",
            "Crocantes y chicharrones: envases con papel absorbente para mantener crujientes.",
            "NO incluir en delivery: tiradito con leche de tigre ya bañada (se cocina en exceso), pescados enteros de robata, sashimi puro sin hielo.",
            "Verificar cada pedido antes de entregar al rider.",
            "Incluir salsa acevichada extra, palillos, carta del restaurante para futuros pedidos.",
            "Tiempo máximo de preparación: 15 minutos.",
            "Responder a reseñas negativas en plataformas en <24h.",
        ]),
        ("10. Gestión de Alérgenos en Cocina Nikkei", [
            "14 alérgenos de declaración obligatoria (estándar UE).",
            "Alérgenos frecuentes en cocina nikkei: pescado (tiradito, ceviche, sashimi, makis), crustáceos (langostinos en chaufa y makis), moluscos (vieira, pulpo), soja (omnipresente en salsas, marinados, salsa acevichada, nikkei dressing), sésamo (aceite, semillas), gluten (salsa soja tradicional con trigo, panko), huevo (Pisco Sour con clara, mayonesa nikkei, yakimeshi), sulfitos (vinos, pisco en algunos).",
            "Carta con indicación de alérgenos por plato.",
            "La salsa soja tradicional contiene TRIGO: informar SIEMPRE a celíacos. Ofrecer tamari (soja sin trigo).",
            "El panko contiene gluten: alternativa sin gluten para crocantes.",
            "Pisco Sour contiene clara de huevo CRUDA: informar al cliente sensible al huevo o embarazadas.",
            "Muchos platos nikkei son sin gluten si se sustituye shoyu por tamari: ventaja competitiva.",
            "Protocolo de cocina para platos sin alérgenos: utensilios separados, zona limpia.",
            "Ante duda: confirmar SIEMPRE con el jefe de cocina nikkei.",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "manual-operaciones-nikkei.docx")
    doc.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 5. PDF placeholder
# ═══════════════════════════════════════════════════════════
def gen_guide_pdf():
    """Generate a minimal placeholder PDF using the DOCX content."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        path = os.path.join(OUTPUT_DIR, "guia-restaurante-nikkei.pdf")
        c = canvas.Canvas(path, pagesize=A4)
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(297, 600, "Como Montar un Restaurante Nikkei")
        c.setFont("Helvetica", 16)
        c.drawCentredString(297, 560, "60 Plazas de Aforo - Guia Espana 2026")
        c.setFont("Helvetica", 12)
        c.drawCentredString(297, 520, "Chef John Guerrero - AI Chef Pro - aichef.pro")
        c.drawCentredString(297, 480, "20 capitulos - 60+ paginas - 8 plantillas - 6 checklists")
        c.save()
        print(f"✓ {path}")
    except ImportError:
        import shutil
        src = os.path.join(OUTPUT_DIR, "guia-restaurante-nikkei.docx")
        dst = os.path.join(OUTPUT_DIR, "guia-restaurante-nikkei.pdf")
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f"✓ {dst} (placeholder from DOCX)")


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n🇯🇵 Generando archivos: Restaurante Nikkei 60 Plazas\n")
    gen_guide_docx()
    gen_guide_pdf()
    gen_plan_financiero()
    gen_calculadora_capex()
    gen_pl_mensual()
    gen_cash_flow()
    gen_escandallo()
    gen_menu_engineering()
    gen_calculadora_ticket()
    gen_cronograma()
    gen_plantilla_turnos()
    gen_checklist_legal()
    gen_checklist_equipamiento()
    gen_checklist_appcc()
    gen_checklist_sala()
    gen_checklist_contratacion()
    gen_checklist_marketing()
    gen_business_plan()
    gen_manual_operaciones()
    print(f"\n✅ {len(os.listdir(OUTPUT_DIR))} archivos generados en {OUTPUT_DIR}\n")
