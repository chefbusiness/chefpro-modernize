#!/usr/bin/env python3
"""
Generate "Cómo Montar un Restaurante Mexicano 80 Plazas" guide deliverables.
AI Chef Pro — aichef.pro
19 files: 1 DOCX guide + 8 Excel templates + 6 Excel checklists + 2 DOCX models + 1 PDF placeholder
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "public", "dl", "guia-restaurante-mexicano"
)
os.makedirs(OUTPUT_DIR, exist_ok=True)

GOLD = "FFD700"
GOLD_RGB = RGBColor(0xFF, 0xD7, 0x00)
HEADER_BG = "2D2D2D"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F5F5"
MEDIUM_GRAY = "E0E0E0"
INPUT_COLOR = "E8F5E9"
BRAND = "AI Chef Pro · aichef.pro — Restaurante Mexicano"

title_font = Font(name="Calibri", size=16, bold=True, color=GOLD)
sub_font = Font(name="Calibri", size=11, color="888888", italic=True)
hdr_font = Font(name="Calibri", size=11, bold=True, color=WHITE)
sec_font = Font(name="Calibri", size=12, bold=True, color=GOLD)
dat_font = Font(name="Calibri", size=11)
bld_font = Font(name="Calibri", size=11, bold=True)
frm_font = Font(name="Calibri", size=11, color="1565C0", bold=True)
brn_font = Font(name="Calibri", size=9, color="888888", italic=True)

hdr_fill = PatternFill(start_color=HEADER_BG, end_color=HEADER_BG, fill_type="solid")
inp_fill = PatternFill(start_color=INPUT_COLOR, end_color=INPUT_COLOR, fill_type="solid")
lgt_fill = PatternFill(start_color=LIGHT_GRAY, end_color=LIGHT_GRAY, fill_type="solid")
sec_fill = PatternFill(start_color="FFF8E1", end_color="FFF8E1", fill_type="solid")

thin_b = Border(
    left=Side(style="thin", color=MEDIUM_GRAY), right=Side(style="thin", color=MEDIUM_GRAY),
    top=Side(style="thin", color=MEDIUM_GRAY), bottom=Side(style="thin", color=MEDIUM_GRAY),
)
c_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
l_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
cur_fmt = '#,##0.00 €'
pct_fmt = '0.0%'


def shr(ws, row, ncols):
    for c in range(1, ncols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = hdr_font; cell.fill = hdr_fill; cell.border = thin_b; cell.alignment = c_align

def sdc(ws, r, c, v=None, font=None, fill=None, fmt=None, align=None):
    cell = ws.cell(row=r, column=c)
    if v is not None: cell.value = v
    cell.font = font or dat_font; cell.border = thin_b; cell.alignment = align or l_align
    if fill: cell.fill = fill
    if fmt: cell.number_format = fmt
    return cell

def title_block(ws, t, ncols=8):
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    ws.cell(row=1, column=1, value=t).font = title_font
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols)
    ws.cell(row=2, column=1, value=BRAND).font = sub_font
    ws.row_dimensions[1].height = 30

def brand_footer(ws, row, ncols):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
    ws.cell(row=row, column=1, value=BRAND).font = brn_font
    ws.cell(row=row, column=1).alignment = c_align

def instr_sheet(wb, t, items):
    ws = wb.active; ws.title = "Instrucciones"; ws.sheet_properties.tabColor = "FFD700"
    title_block(ws, t, 3)
    r = 4
    ws.cell(row=r, column=1, value="Instrucciones de uso").font = sec_font; r += 1
    for i, x in enumerate(items, 1):
        ws.cell(row=r, column=1, value=f"{i}. {x}").font = dat_font; r += 1
    r += 1
    ws.cell(row=r, column=1, value="Celdas verdes = campos editables").font = dat_font
    ws.cell(row=r, column=1).fill = inp_fill
    ws.column_dimensions['A'].width = 80

def checklist_ws(ws, t, items):
    hdrs = ["#", "Categoría", "Tarea / Ítem", "Responsable", "Fecha Límite", "Estado", "Coste Est. (€)", "Notas"]
    title_block(ws, t)
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8); ws.freeze_panes = f"A{r+1}"
    ws.auto_filter.ref = f"A{r}:H{r}"
    dv = DataValidation(type="list", formula1='"Pendiente,En Curso,Completado"', allow_blank=True)
    ws.add_data_validation(dv)
    widths = [6, 22, 45, 18, 14, 14, 14, 30]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    for idx, (cat, tarea, resp, coste) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, tarea)
        sdc(ws, r, 4, resp); sdc(ws, r, 5, None, fill=inp_fill)
        sdc(ws, r, 6, "Pendiente", fill=inp_fill); dv.add(ws.cell(row=r, column=6))
        sdc(ws, r, 7, coste, fmt=cur_fmt, fill=inp_fill); sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1; brand_footer(ws, r, 8)


# ═══════════════════════════════════════════════════════════
# 1. GUIDE DOCX
# ═══════════════════════════════════════════════════════════
def gen_guide_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    def add_bullet(txt):
        doc.add_paragraph(txt, style='List Bullet')
    def add_num(txt):
        doc.add_paragraph(txt, style='List Number')
    def tip(txt):
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Consejo del Chef: {txt}")
        run.font.italic = True; run.font.color.rgb = GOLD_RGB

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cómo Montar un Restaurante Mexicano"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("80 Plazas de Aforo — Guía España 2026"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Chef John Guerrero\nAI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # CH1
    doc.add_heading("1. Qué es un Restaurante Mexicano Auténtico", level=1)
    doc.add_paragraph(
        "Un restaurante mexicano auténtico va mucho más allá de los nachos y los burritos tex-mex. "
        "La cocina mexicana fue declarada Patrimonio Inmaterial de la Humanidad por la UNESCO en 2010, "
        "y su riqueza incluye técnicas prehispánicas como el nixtamal, el uso del comal, la diversidad "
        "de chiles (más de 60 variedades) y elaboraciones complejas como el mole poblano."
    )
    doc.add_paragraph(
        "En España, el mercado de restauración mexicana ha crecido enormemente, pero la mayoría de "
        "establecimientos ofrecen una versión tex-mex simplificada. Un restaurante mexicano auténtico "
        "con producto de calidad, tortillas hechas a mano y una barra de tequilas y mezcales premium "
        "tiene un posicionamiento diferencial enorme y márgenes superiores."
    )
    doc.add_heading("Autenticidad vs Tex-Mex: posicionamiento clave", level=2)
    add_bullet("Auténtico: tortillas de maíz nixtamalizado hechas en el momento, salsas frescas en molcajete, chiles importados.")
    add_bullet("Tex-Mex: nachos con queso industrial, burritos XXL, fajitas con pimiento. Más sencillo pero menos diferencial.")
    add_bullet("Fusión moderna: tacos gourmet con ingredientes locales españoles, ceviche con producto gallego, mole con chocolate de Astorga.")
    add_bullet("Ticket medio auténtico: 18-30€ por persona, frente a 10-15€ del tex-mex rápido.")
    add_bullet("Margen diferencial: las tortillas artesanas cuestan céntimos y se venden a precio premium. El tequila premium tiene márgenes del 80%+.")
    tip("La autenticidad es tu ventaja competitiva. Si haces tortillas a mano delante del cliente con una tortillera/prensa, ya has ganado. Eso no lo puede copiar un tex-mex de franquicia.")
    doc.add_page_break()

    # CH2
    doc.add_heading("2. El Mercado de la Cocina Mexicana en España 2026", level=1)
    doc.add_paragraph(
        "La cocina mexicana es una de las tendencias gastronómicas más fuertes en España. "
        "El número de restaurantes mexicanos en España ha crecido un 35% en los últimos 5 años, "
        "superando los 2.500 establecimientos. Sin embargo, menos del 10% ofrecen cocina "
        "auténticamente mexicana — el resto es tex-mex o cocina fusión."
    )
    doc.add_heading("Datos clave del sector", level=2)
    add_bullet("Restaurantes mexicanos en España: +2.500 (2025), crecimiento del 35% en 5 años.")
    add_bullet("Ticket medio: 18-30€ (auténtico) vs 10-15€ (tex-mex rápido).")
    add_bullet("Ciudades con mayor demanda: Madrid, Barcelona, Valencia, Málaga, Sevilla, Bilbao.")
    add_bullet("Público: 25-45 años, urbano, viajero, foodie, busca experiencias gastronómicas diferentes.")
    add_bullet("Tendencia clave: tequila y mezcal premium como categoría de bebida en auge (crecimiento 25% anual en España).")
    add_bullet("Delivery mexicano: tacos y burritos son los platos #1 en pedidos de comida étnica en plataformas.")
    tip("El 90% de españoles que han viajado a México dicen que echan de menos la comida. Ese es tu mercado: personas que ya conocen la diferencia entre auténtico y tex-mex.")
    doc.add_page_break()

    # CH3
    doc.add_heading("3. Modelos de Negocio: Taquería Gourmet, Cantina y Más", level=1)
    doc.add_paragraph("El restaurante mexicano admite múltiples modelos. La elección del modelo determina la inversión, la operativa diaria y tu posicionamiento en el mercado.")
    doc.add_heading("Modelo 1: Taquería Gourmet", level=2)
    doc.add_paragraph("Tacos de alta calidad con tortilla artesana, proteínas premium (cochinita pibil, carnitas, suadero, al pastor) y salsas frescas. Formato barra + mesas, taquero a la vista. Ticket 15-22€. Inversión contenida, alta rotación.")
    doc.add_heading("Modelo 2: Cantina Mexicana", level=2)
    doc.add_paragraph("Restaurante con ambiente de cantina: platos para compartir, margaritas, tequila y mezcal, música en vivo los fines de semana. Carta amplia con antojitos, ceviches, enchiladas, moles. Ticket 22-30€. Fuerte componente de barra y coctelería.")
    doc.add_heading("Modelo 3: Restaurante Mexicano Casual", level=2)
    doc.add_paragraph("Concepto familiar con carta variada: entrantes para compartir (guacamole, nachos, quesadillas), principales (enchiladas, burritos, fajitas) y postres (churros, flan). Ticket 18-25€. El modelo más replicable y estable.")
    doc.add_heading("Modelo 4: Dark Kitchen / Delivery Mexicano", level=2)
    doc.add_paragraph("Solo delivery y take away: menú optimizado para tacos, burritos y bowls que viajan bien. Sin sala. Inversión 40-80K€. Ideal como segundo concepto o para testear mercado antes de abrir local.")
    doc.add_page_break()

    # CH4
    doc.add_heading("4. Estudio de Viabilidad y Plan Financiero", level=1)
    doc.add_paragraph(
        "La inversión total para un restaurante mexicano de 80 plazas en España oscila "
        "entre 120.000€ y 280.000€ dependiendo de la ubicación, nivel de acabados, "
        "si incluyes barra de tequilas premium y el grado de decoración temática."
    )
    doc.add_heading("Desglose de inversión típica", level=2)
    add_bullet("Obra civil y acondicionamiento: 40.000-100.000€")
    add_bullet("Equipamiento cocina mexicana (comal, tortillera, ahumador, parrilla): 25.000-50.000€")
    add_bullet("Barra de tequilas y mezcales (instalación + stock inicial 30-50 refs): 8.000-20.000€")
    add_bullet("Mobiliario sala + decoración mexicana (Talavera, murales, artesanía): 15.000-35.000€")
    add_bullet("Licencias y permisos (incluye importación): 5.000-12.000€")
    add_bullet("Marketing pre-apertura: 3.000-8.000€")
    add_bullet("Stock inicial (materia prima + chiles + especias importadas): 5.000-10.000€")
    add_bullet("Fondo de maniobra (3 meses): 25.000-60.000€")
    add_bullet("Tecnología (TPV, web, delivery): 3.000-8.000€")
    doc.add_heading("Ratios financieros objetivo", level=2)
    add_bullet("Food cost: 28-33% sobre ventas (los chiles importados suben ligeramente el coste)")
    add_bullet("Coste de personal: 28-35% sobre ventas")
    add_bullet("Alquiler: máximo 8-10% sobre ventas")
    add_bullet("EBITDA objetivo: 12-18%")
    add_bullet("Break-even: mes 8-14 (depende de ubicación y estacionalidad)")
    tip("El margen en bebidas compensa el food cost ligeramente superior. Un tequila premium que compras a 25€/botella lo vendes a 8-12€ el chupito. La barra de tequilas puede ser el 30% de tu facturación con márgenes del 80%.")
    doc.add_page_break()

    # CH5
    doc.add_heading("5. Requisitos Legales en España + Importación", level=1)
    doc.add_paragraph("Abrir un restaurante mexicano en España requiere los mismos trámites que cualquier restaurante, más requisitos adicionales si importas productos directamente de México (chiles secos, salsas, tequila, mezcal).")
    doc.add_heading("Trámites generales", level=2)
    add_num("Constitución de sociedad (SL) o alta como autónomo")
    add_num("Licencia de actividad / declaración responsable (según municipio)")
    add_num("Registro sanitario autonómico")
    add_num("Alta en Hacienda: IAE epígrafe 671.4 o similar")
    add_num("Inscripción Seguridad Social empresa")
    add_num("Seguro de responsabilidad civil (mínimo 300.000€)")
    add_num("Plan APPCC y registro RGSEAA")
    add_num("Licencia de terraza (si aplica)")
    add_num("Licencia de música / SGAE")
    add_num("Libro de reclamaciones")
    doc.add_heading("Requisitos específicos de importación", level=2)
    add_bullet("Registro como operador de comercio exterior (EORI) si importas directamente")
    add_bullet("Certificado SOIVRE para productos alimentarios importados de fuera de la UE")
    add_bullet("Etiquetado en español obligatorio para todos los productos importados")
    add_bullet("Control fitosanitario para chiles secos, especias y productos vegetales")
    add_bullet("Impuestos especiales para bebidas alcohólicas (tequila, mezcal): IIEE")
    add_bullet("Alternativa más sencilla: comprar a distribuidores especializados en España (Mexgrocer, La Tienda Mexicana, Importadora del Sur)")
    tip("No necesitas importar directamente. Hay distribuidores especializados en Madrid y Barcelona que traen todo lo que necesitas: chiles ancho, guajillo, chipotle, masa de maíz nixtamalizado, tortillas, tequila 100% agave. Empieza con distribuidores y solo importa directamente cuando tu volumen lo justifique.")
    doc.add_page_break()

    # CH6
    doc.add_heading("6. APPCC y Seguridad Alimentaria para Cocina Mexicana", level=1)
    doc.add_paragraph("El plan APPCC de un restaurante mexicano tiene las mismas bases que cualquier restaurante, pero con puntos de control adicionales para productos importados y técnicas específicas (nixtamalización, fermentación, ahumado).")
    doc.add_heading("Prerrequisitos estándar", level=2)
    add_bullet("Plan de limpieza y desinfección (frecuencias, productos, responsables)")
    add_bullet("Control de temperaturas (cámaras, servicio, transporte)")
    add_bullet("Plan de control de plagas (DDD)")
    add_bullet("Trazabilidad de materias primas (proveedores, lotes, fechas)")
    add_bullet("Plan de formación del personal (manipulador de alimentos)")
    add_bullet("Control de agua potable")
    add_bullet("Plan de gestión de residuos y aceite usado")
    add_bullet("Plan de alérgenos (14 alérgenos de declaración obligatoria)")
    doc.add_heading("Puntos críticos específicos de cocina mexicana", level=2)
    add_bullet("Chiles secos importados: verificar certificado fitosanitario, almacenar en seco, controlar insectos")
    add_bullet("Masa de maíz nixtamalizado: conservación refrigerada máximo 48h, controlar pH")
    add_bullet("Tortillas frescas: producción y consumo en el día, temperatura de comal >200°C")
    add_bullet("Salsas frescas (pico de gallo, guacamole): vida útil máxima 24h refrigerado")
    add_bullet("Carne para tacos al pastor (trompo): rotación del trompo, temperatura interna >75°C")
    add_bullet("Aguacate: control de maduración, prevención de oxidación, trazabilidad de origen")
    add_bullet("Ceviche mexicano: pescado ultra-fresco, marinado ácido no elimina anisakis — congelar previamente")
    tip("El aguacate es tu ingrediente estrella y tu mayor riesgo de desperdicio. Compra en 3 estados de madurez (verde, madurando, listo) para tener siempre stock. Un aguacate que se pasa son 2€ a la basura.")
    doc.add_page_break()

    # CH7
    doc.add_heading("7. Ubicación y Local", level=1)
    doc.add_paragraph("La ubicación de un restaurante mexicano es crucial. A diferencia de un casual genérico, el componente temático y experiencial atrae a clientes de más lejos, pero necesitas visibilidad para captar tráfico espontáneo.")
    doc.add_heading("Criterios de selección", level=2)
    add_bullet("Tráfico peatonal: mínimo 400 personas/hora en horario punta")
    add_bullet("Visibilidad: fachada llamativa con decoración mexicana visible desde la calle")
    add_bullet("Metros cuadrados: 160-220 m² para 80 plazas (ratio 2-2.75 m²/plaza)")
    add_bullet("Zona: ocio nocturno + afterwork = ideal (el concepto mexicano funciona especialmente bien en cenas)")
    add_bullet("Competencia: analiza los 500m alrededor — evita zonas con más de 2 mexicanos ya establecidos")
    add_bullet("Salida de humos: imprescindible. El comal, la parrilla y el trompo generan mucho humo.")
    add_bullet("Accesibilidad: acceso para personas con movilidad reducida (obligatorio por ley)")
    add_bullet("Terraza: si la zona lo permite, la terraza mexicana con colores vivos atrae mucho tráfico")
    tip("La cena es el servicio estrella de un mexicano. Busca zonas con vida nocturna y afterwork. Los viernes y sábados noche pueden representar el 40% de tu facturación semanal.")
    doc.add_page_break()

    # CH8
    doc.add_heading("8. Diseño de Cocina Mexicana Profesional", level=1)
    doc.add_paragraph("La cocina de un restaurante mexicano tiene estaciones especializadas que no existen en una cocina convencional: la estación del taquero (comal + tortillera), la parrilla/trompo, y la estación de salsas.")
    doc.add_heading("Zonas de la cocina mexicana", level=2)
    add_bullet("Estación del taquero: comal grande (80-100 cm), prensa de tortillas o tortillera eléctrica, mise en place de proteínas y toppings. Visible al cliente si es posible (show cooking).")
    add_bullet("Estación de parrilla/trompo: parrilla de carbón o gas, trompo de pastor vertical, zona de corte de carne.")
    add_bullet("Estación de salsas: molcajetes, licuadoras industriales, mesa fría para ingredientes frescos (cilantro, cebolla, limón, chiles).")
    add_bullet("Zona caliente: fogones para guisos (mole, pozole, frijoles), plancha, freidora (para churros, totopos).")
    add_bullet("Zona fría: preparación de ceviches, guacamole, ensaladas.")
    add_bullet("Almacenamiento: cámaras frigoríficas (2), economato seco especial para chiles y especias.")
    add_bullet("Zona de lavado: tren de lavado, zona de limpieza de molcajetes y comales.")
    add_bullet("Superficie mínima cocina: 45-60 m² (ratio cocina/sala 1:3)")
    tip("La estación del taquero a la vista del cliente es tu mayor herramienta de marketing. Ver cómo se hacen las tortillas a mano es un espectáculo que genera fotos en Instagram y fideliza. Invierte en un buen diseño de show cooking.")
    doc.add_page_break()

    # CH9
    doc.add_heading("9. Equipamiento Específico de Cocina Mexicana", level=1)
    doc.add_paragraph("Además del equipamiento estándar de cocina profesional, un restaurante mexicano necesita equipos específicos que no encontrarás en proveedores convencionales.")
    doc.add_heading("Equipamiento específico mexicano con costes orientativos", level=2)
    add_bullet("Comal profesional grande (80-100 cm, plancha de hierro): 800-2.000€")
    add_bullet("Prensa de tortillas manual profesional: 150-400€")
    add_bullet("Tortillera eléctrica semiautomática (si volumen alto): 2.000-5.000€")
    add_bullet("Trompo de pastor vertical (gas): 500-1.500€")
    add_bullet("Molcajetes grandes (piedra volcánica, para preparación y servicio): 50-150€/ud × 6")
    add_bullet("Ahumador / smoker (para chipotle, carnes ahumadas): 1.500-4.000€")
    add_bullet("Parrilla de carbón/leña profesional: 2.000-5.000€")
    add_bullet("Máquina de margaritas / granizados (para margaritas frozen): 1.000-3.000€")
    doc.add_heading("Equipamiento estándar de cocina", level=2)
    add_bullet("Cocina de gas 6 fuegos + horno: 3.500-6.000€")
    add_bullet("Freidora doble (churros, totopos, empanadas): 1.200-2.500€")
    add_bullet("Horno mixto (Rational / Unox): 8.000-15.000€")
    add_bullet("Cámaras frigoríficas (2 puertas): 2.500-4.000€")
    add_bullet("Cámara congelación: 2.000-3.500€")
    add_bullet("Tren de lavado: 4.000-8.000€")
    add_bullet("Campana extractora + filtros (sobredimensionar por humo de comal y parrilla): 4.000-8.000€")
    add_bullet("Licuadoras industriales (para salsas, 2-3 uds): 300-600€")
    add_bullet("Total equipamiento cocina: 25.000-50.000€")
    tip("El comal y la tortillera son inversiones pequeñas con impacto enorme. Un comal profesional de hierro cuesta 1.000€ y dura 20 años. Las tortillas hechas al momento cuestan 0.05€ cada una y las vendes como experiencia premium.")
    doc.add_page_break()

    # CH10
    doc.add_heading("10. Diseño de Sala y Decoración Mexicana", level=1)
    doc.add_paragraph("La decoración de un restaurante mexicano es fundamental para la experiencia. México tiene una riqueza visual extraordinaria: colores vivos, artesanía, Talavera, murales, textiles. El reto es crear un ambiente auténtico sin caer en el cliché de sombrero y cactus.")
    doc.add_heading("Distribución tipo (80 plazas)", level=2)
    add_bullet("Mesas de 2: 6-8 unidades (12-16 plazas) — parejas")
    add_bullet("Mesas de 4: 10-12 unidades (40-48 plazas) — familias, grupos")
    add_bullet("Barra alta / barra de tequilas: 10-16 plazas — afterwork, cócteles")
    add_bullet("Zona terraza (si aplica): 15-25 plazas adicionales con colores vivos")
    add_bullet("Distancia entre mesas: 60-80 cm")
    doc.add_heading("Elementos clave de la decoración mexicana auténtica", level=2)
    add_bullet("Azulejos de Talavera: en barra, paredes o detalles. Pueden ser mexicanos importados o réplicas españolas.")
    add_bullet("Colores: paleta de rosas mexicanos, azul cobalto, naranja, amarillo. Evitar el exceso: 2-3 colores principales.")
    add_bullet("Murales: un buen mural de artista (calaveras de Día de Muertos, paisaje mexicano, frida-inspired) transforma el espacio. Coste: 1.500-5.000€.")
    add_bullet("Vajilla de barro: platos de barro negro de Oaxaca o barro rojo para servicio. Muy instagrameable.")
    add_bullet("Textiles: manteles individuales de tela mexicana, cojines en bancos corridos.")
    add_bullet("Plantas: cactus y suculentas reales en la entrada y barra.")
    add_bullet("Iluminación: cálida (2700K), lámparas de latón o papel picado como elementos decorativos.")
    tip("Menos es más en decoración mexicana. Un mural espectacular + vajilla de barro + Talavera en la barra es suficiente. Si pones sombreros en la pared, calaveras de plástico y cactus de mentira, pareces un tex-mex de centro comercial.")
    doc.add_page_break()

    # CH11
    doc.add_heading("11. Barra de Tequilas y Mezcales (30-50 Referencias)", level=1)
    doc.add_paragraph("La barra de tequilas y mezcales es el elemento diferencial que eleva un restaurante mexicano de casual a experiencial. El mercado de tequila y mezcal premium en España está en pleno auge, con crecimientos del 25% anual.")
    doc.add_heading("Estructura de la barra (30-50 referencias)", level=2)
    add_bullet("Tequila blanco (6-8 refs): Patrón, Don Julio, Herradura, Fortaleza, Ocho, Casa Dragones")
    add_bullet("Tequila reposado (5-7 refs): Don Julio Reposado, Clase Azul, Herradura Reposado")
    add_bullet("Tequila añejo (4-6 refs): Don Julio 1942, Clase Azul Añejo, Patrón Añejo")
    add_bullet("Mezcal (6-10 refs): Del Maguey, Montelobos, Ilegal, Alipús, Los Danzantes")
    add_bullet("Tequilas premium/ultra (2-3 refs): Clase Azul Ultra, Don Julio Real, Lobos 1707")
    add_bullet("Cócteles mexicanos (8-12): Margarita clásica, Paloma, Michelada, Mezcal Negroni, Oaxaca Old Fashioned")
    doc.add_heading("Inversión y márgenes", level=2)
    add_bullet("Stock inicial barra (30-50 refs): 5.000-15.000€")
    add_bullet("Expositor / estantería iluminada para botellas: 2.000-5.000€")
    add_bullet("Margen tequila premium: 75-85% (compra 25€/botella → venta 8-12€/chupito)")
    add_bullet("Margen cócteles: 78-85% (una margarita cuesta 1.50-2€ y se vende a 9-12€)")
    add_bullet("Objetivo: que la barra represente 25-35% de la facturación total")
    tip("Ofrece vuelos de tequila (3 chupitos de diferentes categorías por 15-20€). Es la forma más rentable de vender tequila premium: el cliente aprende, disfruta y acaba pidiendo una botella para la mesa. Margen brutal.")
    doc.add_page_break()

    # CH12
    doc.add_heading("12. Brigada de Cocina (6-10 personas)", level=1)
    doc.add_paragraph("La brigada de cocina de un restaurante mexicano tiene roles específicos que no existen en una cocina europea convencional: el taquero y el parrillero/salsero son puestos especializados.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Jefe de cocina (idealmente con experiencia en cocina mexicana): 1 persona — 2.200-2.800€ brutos/mes")
    add_bullet("Taquero (estación de comal, tortillas, tacos): 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Parrillero (parrilla, trompo, carnes): 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Salsero / cocinero frío (salsas, ceviches, guacamole): 1 persona — 1.500-1.800€ brutos/mes")
    add_bullet("Cocineros (guisos, moles, arroces, frijoles): 2-3 personas — 1.500-1.800€ brutos/mes cada uno")
    add_bullet("Ayudante / office: 1-2 personas — 1.200-1.500€ brutos/mes")
    add_bullet("Coste total cocina: 12.000-20.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Contratación del taquero", level=2)
    doc.add_paragraph("El taquero es el alma del restaurante mexicano. Si encuentras un taquero mexicano con experiencia, contratarlo es una inversión que se paga sola. Alternativamente, un cocinero español puede formarse en la estación del taquero con un buen programa de entrenamiento de 2-4 semanas.")
    tip("Busca taqueros y cocineros mexicanos en comunidades latinas de Madrid, Barcelona y Valencia. Hay talento increíble trabajando en otros sectores que sueña con volver a la cocina. LinkedIn, grupos de Facebook de mexicanos en España y asociaciones culturales son tus mejores canales de reclutamiento.")
    doc.add_page_break()

    # CH13
    doc.add_heading("13. Equipo de Sala (6-8 personas)", level=1)
    doc.add_paragraph("El equipo de sala de un restaurante mexicano necesita conocer la cultura gastronómica mexicana para poder explicar platos, recomendar tequilas y crear la experiencia completa.")
    doc.add_heading("Organigrama tipo", level=2)
    add_bullet("Encargado/a de sala: 1 persona — 2.000-2.500€ brutos/mes")
    add_bullet("Barman de tequilas y coctelería: 1 persona — 1.800-2.200€ brutos/mes")
    add_bullet("Camareros: 3-4 personas — 1.400-1.700€ brutos/mes cada uno")
    add_bullet("Runner / ayudante: 1-2 personas — 1.200-1.400€ brutos/mes")
    add_bullet("Coste total sala: 9.000-15.000€ brutos/mes (incluye SS empresa)")
    doc.add_heading("Claves del servicio en un restaurante mexicano", level=2)
    add_bullet("Conocimiento de tequilas y mezcales: todo el equipo debe saber distinguir blanco, reposado, añejo y recomendar.")
    add_bullet("Explicar platos: muchos clientes no saben qué es una enchilada suiza o un mole poblano. El camarero debe saber describirlo.")
    add_bullet("Sugerir orden de comida: en México se comparte. Recomendar pedir entrantes para compartir + principales individuales.")
    add_bullet("Gestión de picante: preguntar nivel de tolerancia al picante y adaptar salsas.")
    add_bullet("Upselling de bebidas: margaritas, vuelos de tequila, mezcal. El barman es clave para el ticket medio.")
    tip("La formación del barman de tequilas es una inversión que se multiplica. Un barman que sabe contar la historia de cada tequila vende el triple. Invierte en un curso de Tequila y Mezcal (hay certificaciones online de la CRT — Consejo Regulador del Tequila).")
    doc.add_page_break()

    # CH14
    doc.add_heading("14. Menú Engineering: Tacos, Burritos, Enchiladas y Más", level=1)
    doc.add_paragraph("El diseño de la carta de un restaurante mexicano debe equilibrar autenticidad, margen y accesibilidad. No todos los clientes conocen la cocina mexicana: la carta debe educar y seducir.")
    doc.add_heading("Estructura de carta mexicana", level=2)
    add_bullet("Antojitos / para compartir: 6-8 opciones (guacamole, nachos, quesadillas, elotes, totopos con salsas)")
    add_bullet("Tacos (3-5 variedades): al pastor, carnitas, cochinita, barbacoa, pescado")
    add_bullet("Principales: 6-8 opciones (enchiladas, burritos, fajitas, mole con pollo, chile relleno)")
    add_bullet("Ceviches mexicanos: 2-3 opciones")
    add_bullet("Postres: 3-4 opciones (churros con chocolate mexicano, flan de cajeta, pastel de tres leches)")
    add_bullet("Total carta: 20-28 platos")
    doc.add_heading("Menú engineering: clasificación de platos mexicanos", level=2)
    add_bullet("Stars (alta popularidad + alto margen): Tacos al pastor, Guacamole fresco, Margaritas → destacar siempre")
    add_bullet("Plowhorses (alta popularidad + bajo margen): Burritos XXL, Nachos con queso → subir precio o reducir porción")
    add_bullet("Puzzles (baja popularidad + alto margen): Mole poblano, Ceviche premium → mejorar descripción en carta")
    add_bullet("Dogs (baja popularidad + bajo margen): Ensalada mexicana, Sopa azteca → eliminar o reinventar")
    tip("Los tacos al pastor son tu Star absoluto. El coste de un taco es 0.80-1.20€ y lo vendes a 3.50-5€. Con 3 tacos por persona a 12-15€, es el plato con mejor ratio margen/satisfacción de toda la carta.")
    doc.add_page_break()

    # CH15
    doc.add_heading("15. Proveedores de Producto Mexicano en España", level=1)
    doc.add_paragraph("Encontrar ingredientes auténticos mexicanos en España es más fácil de lo que parece. Hay distribuidores especializados y la comunidad mexicana ha creado una red de importación sólida.")
    doc.add_heading("Ingredientes clave y dónde encontrarlos", level=2)
    add_bullet("Chiles secos (ancho, guajillo, chipotle, pasilla, morita, de árbol): Mexgrocer, La Tienda Mexicana, Importadora del Sur")
    add_bullet("Masa de maíz nixtamalizado (harina Maseca o masa fresca): distribuidores latinos en Madrid/Barcelona")
    add_bullet("Tortillas de maíz (si no haces artesanas): La Tortillería (Madrid), Tacos de Canasta (Barcelona)")
    add_bullet("Tequila y mezcal: distribuidores como Zamora Company, Pernod Ricard, Moët Hennessy, o directo a destilerías")
    add_bullet("Aguacate Hass: cualquier proveedor de frutas/verduras (España importa 100.000 toneladas/año)")
    add_bullet("Frijoles negros, arroz: supermercados latinos o distribuidores étnicos")
    add_bullet("Cilantro fresco: cualquier proveedor de verduras (pedir suministro diario)")
    add_bullet("Limón mexicano (lima): sustituir por lima persa o lima key, disponible en mercados de abastos")
    add_bullet("Queso Oaxaca / panela: difícil de importar, pero se puede sustituir con mozzarella fresca para quesadillas")
    add_bullet("Chapulines, huitlacoche, nopal: productos gourmet importados, disponibles en tiendas especializadas")
    tip("Haz una visita a los mercados latinos de Madrid (Lavapiés) o Barcelona (Raval). Encontrarás chiles frescos, tomatillos, nopales y masa de maíz a precios de mayorista. Muchos proveedores hacen reparto diario si superas 50€/pedido.")
    doc.add_page_break()

    # CH16
    doc.add_heading("16. 15 Recetas Base con Food Cost", level=1)
    doc.add_paragraph("Estas 15 recetas forman la base de la carta de un restaurante mexicano. Cada una incluye food cost estimado y PVP sugerido para mantener márgenes del 28-33%.")
    recipes = [
        ("Tacos al pastor (3 uds)", "1.20€", "4.50€", "27%"),
        ("Carnitas (3 tacos)", "1.40€", "4.50€", "31%"),
        ("Guacamole fresco en molcajete", "1.80€", "8.00€", "23%"),
        ("Mole poblano con pollo y arroz", "2.50€", "13.00€", "19%"),
        ("Enchiladas suizas (3 uds)", "1.60€", "10.50€", "15%"),
        ("Quesadillas de queso y flor de calabaza", "1.00€", "7.00€", "14%"),
        ("Burritos de carne asada", "2.20€", "11.00€", "20%"),
        ("Nachos con queso, jalapeños y guacamole", "1.50€", "9.00€", "17%"),
        ("Ceviche mexicano de lubina", "3.50€", "13.50€", "26%"),
        ("Tamales de pollo en hoja de maíz (2 uds)", "1.30€", "8.50€", "15%"),
        ("Pozole rojo con cerdo", "2.00€", "11.00€", "18%"),
        ("Churros con chocolate mexicano (6 uds)", "0.80€", "6.50€", "12%"),
        ("Flan de cajeta", "0.90€", "6.00€", "15%"),
        ("Elotes con mayonesa, chile y limón", "0.60€", "5.50€", "11%"),
        ("Chile relleno de queso con salsa de tomate", "1.80€", "11.00€", "16%"),
    ]
    for name, fc, pvp, pct in recipes:
        doc.add_heading(name, level=2)
        add_bullet(f"Food cost: {fc}")
        add_bullet(f"PVP sugerido (sin IVA): {pvp}")
        add_bullet(f"% Food cost: {pct}")
    tip("El guacamole fresco preparado en molcajete delante del cliente cuesta 1.80€ y lo vendes a 8€. Es el plato con mejor margen y el más instagrameable. Hazlo SIEMPRE en mesa como show.")
    doc.add_page_break()

    # CH17
    doc.add_heading("17. Delivery Optimizado para Tacos y Burritos", level=1)
    doc.add_paragraph("Los tacos y burritos son los reyes del delivery de comida étnica en España. Pero hay que adaptar el formato para que el producto llegue bien al cliente.")
    doc.add_heading("Plataformas y canales", level=2)
    add_bullet("Glovo, Uber Eats, Just Eat: comisión 25-35% — útiles para visibilidad inicial")
    add_bullet("Delivery propio (web/app): comisión 0% pero requiere repartidores propios o Stuart/Paack")
    add_bullet("Take away en local: margen completo, packaging incluido en precio")
    doc.add_heading("Menú delivery mexicano optimizado", level=2)
    add_bullet("Burritos: viajan perfecto envueltos en aluminio + papel. El producto estrella del delivery mexicano.")
    add_bullet("Bowls mexicanos: base de arroz + proteína + frijoles + toppings. Formato bowl = transporte seguro.")
    add_bullet("Tacos kit: tortillas separadas de proteínas y salsas en recipientes individuales. El cliente monta en casa.")
    add_bullet("Nachos loaded: envase ancho, queso y toppings en compartimento separado para que no se empapen.")
    add_bullet("Quesadillas: cortadas en triángulos, envueltas en papel. Viajan bien.")
    add_bullet("NO incluir en delivery: guacamole (se oxida), ceviches (cadena de frío), platos con salsa líquida abundante.")
    add_bullet("Packaging: envases de cartón kraft con compartimentos, bolsa de papel + logo. Coste 0.50-1.50€/pedido.")
    tip("El 'taco kit' es una idea brillante para delivery: envías tortillas en una bolsa térmica, proteínas en un envase, salsas en botecitos y el cliente arma sus tacos en casa. Experiencia interactiva + producto fresco al llegar.")
    doc.add_page_break()

    # CH18
    doc.add_heading("18. Marketing con Día de Muertos, Cinco de Mayo y Más", level=1)
    doc.add_paragraph("Un restaurante mexicano tiene un calendario de eventos culturales único que genera picos de facturación naturales. Aprovéchalo.")
    doc.add_heading("Eventos culturales para marketing", level=2)
    add_bullet("Día de Muertos (1-2 noviembre): decoración de altar de muertos, menú especial, pan de muerto, calaveras de azúcar. El evento más potente del año.")
    add_bullet("Cinco de Mayo (5 mayo): margaritas especiales, promociones, fiesta mexicana. Aunque es una fiesta menor en México, en España funciona muy bien.")
    add_bullet("Día de la Independencia (16 septiembre): Grito de Independencia, pozole, chiles en nogada (solo en temporada sep-oct).")
    add_bullet("Día de la Candelaria (2 febrero): tamales y atole. Tradición de 'quien saca el muñequito de la rosca invita los tamales'.")
    add_bullet("Navidades mexicanas: ponche navideño, tamales, bacalao a la vizcaína.")
    doc.add_heading("Canales de marketing", level=2)
    add_num("Instagram: contenido visual de platos coloridos, preparación de tortillas, cocktails. Reels > fotos.")
    add_num("TikTok: vídeos de taquero preparando tacos, el trompo al pastor, guacamole en molcajete. Potencial viral altísimo.")
    add_num("Google My Business: ficha completa con fotos profesionales, menú, reseñas.")
    add_num("TheFork / TripAdvisor: perfiles con fotos y respuesta a reseñas.")
    add_num("Web propia: menú, historia, reservas online, carta de tequilas.")
    doc.add_heading("Presupuesto marketing mensual", level=2)
    add_bullet("Google Ads local: 200-500€/mes")
    add_bullet("Instagram/Facebook Ads: 200-500€/mes")
    add_bullet("Fotógrafo food: 200-300€/sesión trimestral")
    add_bullet("Community manager (si se externaliza): 300-600€/mes")
    add_bullet("Total recomendado: 600-1.500€/mes (2-3% facturación)")
    tip("El Día de Muertos es tu Super Bowl. Decora el restaurante con un altar espectacular 2 semanas antes, publica contenido diario en redes, ofrece menú especial y cena con mariachi o DJ. Un solo fin de semana de Día de Muertos puede facturar el equivalente a 2 semanas normales.")
    doc.add_page_break()

    # CH19
    doc.add_heading("19. Tecnología para Restaurante Mexicano", level=1)
    doc.add_paragraph("La tecnología en un restaurante mexicano debe ser invisible para el cliente pero esencial para la operativa. Automatiza lo repetitivo para que tu equipo se centre en crear experiencias.")
    doc.add_heading("Stack tecnológico recomendado", level=2)
    add_bullet("TPV: Lightspeed, Square, Revo. Cloud, con reporting en tiempo real. 60-150€/mes.")
    add_bullet("Reservas: TheFork, CoverManager, Resy. Gestión de turnos y no-shows. 0-100€/mes.")
    add_bullet("Delivery: integrador (Ordatic, Deliverect) que centraliza Glovo + Uber Eats + Just Eat. 80-150€/mes.")
    add_bullet("Contabilidad: Holded, Quipu. Facturas automáticas, conexión con asesoría. 30-60€/mes.")
    add_bullet("RRHH y turnos: Factorial, Kenjo. Cuadrantes, fichajes, nóminas. 4-8€/empleado/mes.")
    add_bullet("Web y carta digital: QR con carta actualizable, sección de tequilas. WordPress o Squarespace. 10-30€/mes.")
    add_bullet("WiFi profesional: router dual-band, red separada para clientes y cocina. 50-80€/mes.")
    tip("El TPV es el corazón de tu restaurante. Elige uno cloud que te permita ver ventas, costes y KPIs desde el móvil. Analiza qué tacos venden más, qué tequilas rotan y cuál es tu margen real por plato.")
    doc.add_page_break()

    # CH20
    doc.add_heading("20. Plan de Acción: De la Idea a la Inauguración", level=1)
    doc.add_paragraph("Abrir un restaurante mexicano en España es un proyecto de 10-14 meses desde la primera idea hasta la inauguración. Este capítulo resume el cronograma y los hitos clave.")
    doc.add_heading("Fases del proyecto", level=2)
    add_bullet("Meses 1-3: Plan de negocio, búsqueda de financiación, constitución de empresa, estudio de mercado.")
    add_bullet("Meses 3-5: Búsqueda y contratación de local, negociación de alquiler, licencia de actividad.")
    add_bullet("Meses 4-7: Obra civil, instalaciones, decoración mexicana.")
    add_bullet("Meses 5-7: Equipamiento de cocina, pedido de comal y tortillera, instalación de barra de tequilas.")
    add_bullet("Meses 6-8: Establecer relaciones con proveedores de producto mexicano, primeras importaciones.")
    add_bullet("Meses 7-9: Contratación y formación de equipo (especialmente taquero y barman de tequilas).")
    add_bullet("Meses 8-10: Branding, web, redes sociales, fotografía profesional.")
    add_bullet("Mes 10-11: Soft opening con amigos, familia e influencers.")
    add_bullet("Mes 11-12: Ajustes finales e INAUGURACIÓN.")
    add_bullet("Mes 12-14: Seguimiento, ajustes de carta y optimización de costes.")
    tip("No abras sin haber hecho mínimo 3 soft openings. Necesitas probar la estación del taquero a pleno rendimiento, la velocidad de salsas, los tiempos de coctelería y el flujo de sala. Cada soft opening te enseña fallos que no puedes detectar en papel.")

    # Save
    path = os.path.join(OUTPUT_DIR, "guia-restaurante-mexicano.docx")
    doc.save(path)
    print(f"✓ {path}")
    return path


# ═══════════════════════════════════════════════════════════
# 2. EXCEL TEMPLATES (8)
# ═══════════════════════════════════════════════════════════

def gen_plan_financiero():
    wb = Workbook()
    instr_sheet(wb, "Plan Financiero 3 Años — Restaurante Mexicano 80 Plazas", [
        "Rellena las celdas verdes con los datos de tu proyecto.",
        "Las fórmulas se recalculan automáticamente.",
        "Pestaña 'Inversión' = CAPEX inicial desglosado.",
        "Pestaña 'P&L Mensual' = cuenta de resultados mes a mes.",
        "Pestaña 'Proyección 3 Años' = evolución anual.",
    ])
    # Inversión sheet
    ws = wb.create_sheet("Inversión"); ws.sheet_properties.tabColor = "4CAF50"
    title_block(ws, "Inversión Inicial — Restaurante Mexicano 80 Plazas")
    hdrs = ["Categoría", "Partida", "Coste Estimado (€)", "Coste Real (€)", "Diferencia", "Notas"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 6)
    widths = [25, 40, 18, 18, 18, 30];
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    items = [
        ("Obra Civil", "Acondicionamiento local", 70000),
        ("Obra Civil", "Instalaciones eléctricas y gas", 10000),
        ("Obra Civil", "Fontanería y saneamiento", 7000),
        ("Obra Civil", "Climatización / HVAC", 8000),
        ("Obra Civil", "Extracción reforzada (humo comal/parrilla)", 6000),
        ("Cocina Mexicana", "Comal profesional grande (80-100 cm)", 1500),
        ("Cocina Mexicana", "Prensa de tortillas / tortillera eléctrica", 3000),
        ("Cocina Mexicana", "Trompo de pastor vertical", 1000),
        ("Cocina Mexicana", "Ahumador / smoker", 3000),
        ("Cocina Mexicana", "Parrilla carbón/gas profesional", 3500),
        ("Cocina Mexicana", "Molcajetes grandes (6 uds)", 600),
        ("Cocina Mexicana", "Licuadoras industriales para salsas (3 uds)", 500),
        ("Cocina Estándar", "Horno mixto Rational/Unox", 12000),
        ("Cocina Estándar", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina Estándar", "Freidora doble (churros, totopos)", 2000),
        ("Cocina Estándar", "Cámaras frigoríficas (2)", 6000),
        ("Cocina Estándar", "Congelador", 3000),
        ("Cocina Estándar", "Tren de lavado", 5000),
        ("Cocina Estándar", "Campana extractora + filtros", 6000),
        ("Cocina Estándar", "Menaje y utensilios", 2500),
        ("Barra Tequilas", "Expositor iluminado botellas tequila/mezcal", 4000),
        ("Barra Tequilas", "Máquina de margaritas / granizados", 2000),
        ("Barra Tequilas", "Stock inicial tequila y mezcal (30-50 refs)", 10000),
        ("Barra Tequilas", "Coctelera, medidores, utensilios barra", 500),
        ("Barra Tequilas", "Máquina de hielo", 1500),
        ("Sala", "Mobiliario interior (mesas + sillas + bancos corridos)", 15000),
        ("Sala", "Barra + taburetes altos", 6000),
        ("Decoración", "Azulejos Talavera (barra + paredes)", 3000),
        ("Decoración", "Mural artístico mexicano", 3000),
        ("Decoración", "Vajilla de barro mexicano", 2000),
        ("Decoración", "Iluminación decorativa + papel picado", 2500),
        ("Decoración", "Plantas, cactus, textiles mexicanos", 1000),
        ("Terraza", "Mobiliario exterior con colores mexicanos", 6000),
        ("Terraza", "Sombrillas / toldos", 2500),
        ("Tecnología", "TPV + pantallas cocina", 3000),
        ("Tecnología", "Web + carta QR + WiFi", 2000),
        ("Licencias", "Licencia actividad + terraza", 6000),
        ("Licencias", "Proyecto técnico + tasas", 4000),
        ("Licencias", "Registro EORI / importación (si aplica)", 500),
        ("Marketing", "Branding mexicano + diseño", 3000),
        ("Marketing", "Fotos + vídeo + campaña pre-apertura", 3000),
        ("Stock", "Materia prima inicial (incluye chiles, especias importadas)", 6000),
        ("Stock", "Bebidas (cerveza, refrescos, vinos)", 3000),
        ("Maniobra", "Fondo de maniobra (3 meses)", 45000),
    ]
    r += 1
    start_r = r
    for cat, partida, coste in items:
        sdc(ws, r, 1, cat); sdc(ws, r, 2, partida)
        sdc(ws, r, 3, coste, fmt=cur_fmt); sdc(ws, r, 4, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}-C{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 2, "TOTAL INVERSIÓN", font=bld_font)
    sdc(ws, r, 3, f"=SUM(C{start_r}:C{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=D{r}-C{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 6)

    # P&L Mensual sheet
    ws2 = wb.create_sheet("P&L Mensual"); ws2.sheet_properties.tabColor = "2196F3"
    title_block(ws2, "P&L Mensual — Restaurante Mexicano 80 Plazas", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic", "TOTAL"]
    r = 4
    for i, m in enumerate(months, 1): ws2.cell(row=r, column=i, value=m)
    shr(ws2, r, 14)
    ws2.column_dimensions['A'].width = 28
    for i in range(2, 15): ws2.column_dimensions[get_column_letter(i)].width = 12
    ws2.freeze_panes = f"B{r+1}"
    rows_data = [
        ("INGRESOS", None, True),
        ("Ventas sala (comida)", 28000, False),
        ("Barra tequilas y coctelería", 12000, False),
        ("Delivery (tacos, burritos, bowls)", 6000, False),
        ("Bebidas (cerveza, refrescos, vinos)", 6000, False),
        ("TOTAL INGRESOS", "=SUM", True),
        ("", None, False),
        ("COSTES VARIABLES", None, True),
        ("Food cost (31%)", "=0.31*", False),
        ("Packaging delivery", 400, False),
        ("TOTAL COSTES VARIABLES", "=SUM_CV", True),
        ("", None, False),
        ("COSTES FIJOS", None, True),
        ("Alquiler", 4000, False),
        ("Personal cocina (taquero, parrillero, equipo)", 16000, False),
        ("Personal sala (barman tequilas, camareros)", 12000, False),
        ("Suministros (luz, gas, agua)", 2800, False),
        ("Seguros", 350, False),
        ("Gestoría + contabilidad", 400, False),
        ("Marketing", 900, False),
        ("Tecnología (TPV, delivery, web)", 350, False),
        ("Mantenimiento", 300, False),
        ("Reposición stock tequila/mezcal", 1500, False),
        ("Varios / imprevistos", 500, False),
        ("TOTAL COSTES FIJOS", "=SUM_CF", True),
        ("", None, False),
        ("EBITDA", "=EBITDA", True),
        ("% EBITDA", "=%EBITDA", True),
    ]
    r += 1
    for label, val, bold in rows_data:
        sdc(ws2, r, 1, label, font=bld_font if bold else dat_font)
        if val and val not in ("=SUM", "=SUM_CV", "=SUM_CF", "=EBITDA", "=%EBITDA") and not str(val).startswith("=0."):
            for c in range(2, 14):
                sdc(ws2, r, c, val, fill=inp_fill, fmt=cur_fmt)
            sdc(ws2, r, 14, f"=SUM(B{r}:M{r})", font=bld_font, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws2, r, 14)

    path = os.path.join(OUTPUT_DIR, "plan-financiero-3-anos.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_capex():
    wb = Workbook()
    instr_sheet(wb, "Calculadora CAPEX — Restaurante Mexicano 80 Plazas", [
        "Introduce tus costes reales en las celdas verdes.",
        "La columna 'Diferencia' se calcula automáticamente.",
        "Usa la columna 'Prioridad' para planificar fases de inversión.",
    ])
    ws = wb.create_sheet("CAPEX Desglosado"); ws.sheet_properties.tabColor = "FF9800"
    title_block(ws, "Calculadora CAPEX — Inversión 120K-280K€")
    hdrs = ["#", "Categoría", "Partida", "Estimado (€)", "Real (€)", "Diferencia", "Prioridad", "Proveedor"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 8)
    widths = [5, 20, 40, 14, 14, 14, 12, 25]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    dv = DataValidation(type="list", formula1='"Alta,Media,Baja"', allow_blank=True)
    ws.add_data_validation(dv)
    items = [
        ("Obra", "Acondicionamiento general", 70000),
        ("Obra", "Electricidad + gas", 10000),
        ("Obra", "Fontanería", 7000),
        ("Obra", "Climatización", 8000),
        ("Obra", "Extracción reforzada comal/parrilla", 6000),
        ("Cocina MX", "Comal profesional grande 80-100 cm", 1500),
        ("Cocina MX", "Prensa tortillas / tortillera eléctrica", 3000),
        ("Cocina MX", "Trompo de pastor vertical", 1000),
        ("Cocina MX", "Ahumador / smoker", 3000),
        ("Cocina MX", "Parrilla carbón/gas profesional", 3500),
        ("Cocina MX", "Molcajetes grandes (6 uds)", 600),
        ("Cocina MX", "Licuadoras industriales salsas (3 uds)", 500),
        ("Cocina", "Horno mixto Rational/Unox", 12000),
        ("Cocina", "Cocina gas 6 fuegos + horno", 5000),
        ("Cocina", "Freidora doble", 2000),
        ("Cocina", "Cámaras frigoríficas (2)", 6000),
        ("Cocina", "Congelador", 3000),
        ("Cocina", "Tren de lavado", 5000),
        ("Cocina", "Campana extractora reforzada", 6000),
        ("Cocina", "Menaje y utensilios", 2500),
        ("Barra", "Expositor iluminado tequila/mezcal", 4000),
        ("Barra", "Máquina margaritas / granizados", 2000),
        ("Barra", "Stock inicial tequila + mezcal (30-50 refs)", 10000),
        ("Barra", "Utensilios coctelería + máquina hielo", 2000),
        ("Sala", "Mesas, sillas, bancos corridos (80 plazas)", 15000),
        ("Sala", "Barra + taburetes altos", 6000),
        ("Decoración", "Talavera, mural, vajilla barro, iluminación", 11500),
        ("Terraza", "Mobiliario exterior + sombrillas", 8500),
        ("Tech", "TPV + pantallas cocina", 3000),
        ("Tech", "Web + carta QR + WiFi", 2000),
        ("Licencias", "Licencia actividad + terraza", 6000),
        ("Licencias", "Proyecto técnico + registro EORI", 4500),
        ("Marketing", "Branding + diseño + pre-apertura", 6000),
        ("Stock", "Materia prima + chiles + especias", 6000),
        ("Stock", "Bebidas (cerveza, refrescos)", 3000),
        ("Maniobra", "Fondo 3 meses", 45000),
    ]
    r += 1; start_r = r
    for idx, (cat, partida, est) in enumerate(items, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, cat); sdc(ws, r, 3, partida)
        sdc(ws, r, 4, est, fmt=cur_fmt); sdc(ws, r, 5, None, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 6, f"=E{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 7, "Alta", fill=inp_fill); dv.add(ws.cell(row=r, column=7))
        sdc(ws, r, 8, None, fill=inp_fill)
        r += 1
    r += 1
    sdc(ws, r, 3, "TOTAL", font=bld_font)
    sdc(ws, r, 4, f"=SUM(D{start_r}:D{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 5, f"=SUM(E{start_r}:E{r-2})", font=bld_font, fmt=cur_fmt)
    sdc(ws, r, 6, f"=E{r}-D{r}", font=bld_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 8)
    path = os.path.join(OUTPUT_DIR, "calculadora-capex.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_pl_mensual():
    wb = Workbook()
    instr_sheet(wb, "P&L Mensual con 3 Escenarios — Restaurante Mexicano", [
        "3 pestañas: Pesimista, Realista, Optimista.",
        "Modifica los ingresos en cada escenario.",
        "Los costes fijos son iguales en los 3 escenarios.",
    ])
    for scenario, factor, color in [("Pesimista", 0.75, "F44336"), ("Realista", 1.0, "4CAF50"), ("Optimista", 1.25, "2196F3")]:
        ws = wb.create_sheet(scenario); ws.sheet_properties.tabColor = color
        title_block(ws, f"P&L Mensual — Escenario {scenario}", 4)
        hdrs = ["Concepto", "Mensual (€)", "Anual (€)", "% s/Ventas"]
        r = 4
        for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
        shr(ws, r, 4)
        ws.column_dimensions['A'].width = 35
        for i in range(2, 5): ws.column_dimensions[get_column_letter(i)].width = 16
        base_ventas = int(52000 * factor)
        rows = [
            ("INGRESOS", None, True),
            ("Ventas sala (comida)", int(28000 * factor)),
            ("Barra tequilas y coctelería", int(12000 * factor)),
            ("Delivery", int(6000 * factor)),
            ("Bebidas (cerveza, refrescos, vinos)", int(6000 * factor)),
            ("TOTAL INGRESOS", base_ventas, True),
            ("", None, False),
            ("COSTES VARIABLES", None, True),
            ("Food cost (31%)", int(base_ventas * 0.31)),
            ("Packaging delivery", 400),
            ("TOTAL COSTES VARIABLES", int(base_ventas * 0.31 + 400), True),
            ("", None, False),
            ("COSTES FIJOS", None, True),
            ("Alquiler", 4000),
            ("Personal cocina", 16000),
            ("Personal sala", 12000),
            ("Suministros", 2800),
            ("Seguros", 350),
            ("Gestoría", 400),
            ("Marketing", 900),
            ("Tecnología", 350),
            ("Mantenimiento", 300),
            ("Reposición tequila/mezcal", 1500),
            ("Varios", 500),
            ("TOTAL COSTES FIJOS", 39100, True),
            ("", None, False),
            ("EBITDA", int(base_ventas - base_ventas * 0.31 - 400 - 39100), True),
        ]
        r += 1
        for item in rows:
            if len(item) == 3:
                label, val, bold = item
            else:
                label, val = item; bold = False
            sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
            if val is not None:
                sdc(ws, r, 2, val, font=bld_font if bold else dat_font, fmt=cur_fmt)
                sdc(ws, r, 3, val * 12, font=bld_font if bold else dat_font, fmt=cur_fmt)
            r += 1
        r += 1; brand_footer(ws, r, 4)
    path = os.path.join(OUTPUT_DIR, "pl-mensual-escenarios.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cash_flow():
    wb = Workbook()
    instr_sheet(wb, "Cash Flow y Break-Even — Restaurante Mexicano 80 Plazas", [
        "Pestaña 'Cash Flow' = flujo de caja mensual 12 meses.",
        "Pestaña 'Break-Even' = calculadora de punto de equilibrio.",
        "Rellena las celdas verdes con tus datos reales.",
    ])
    ws = wb.create_sheet("Cash Flow"); ws.sheet_properties.tabColor = "009688"
    title_block(ws, "Cash Flow 12 Meses — Restaurante Mexicano", 14)
    months = ["Concepto", "Ene", "Feb", "Mar", "Abr", "May", "Jun", "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"]
    r = 4
    for i, m in enumerate(months, 1): ws.cell(row=r, column=i, value=m)
    shr(ws, r, 13)
    ws.column_dimensions['A'].width = 28
    for i in range(2, 14): ws.column_dimensions[get_column_letter(i)].width = 12
    concepts = [
        ("Saldo inicial", True), ("", False),
        ("ENTRADAS", True), ("Ventas sala (comida)", False), ("Barra tequilas/coctelería", False),
        ("Delivery", False), ("Bebidas", False), ("Total entradas", True), ("", False),
        ("SALIDAS", True), ("Materia prima + chiles importados", False), ("Personal cocina", False),
        ("Personal sala", False), ("Alquiler", False), ("Suministros", False), ("Marketing", False),
        ("Reposición tequila/mezcal", False), ("Tecnología", False), ("Otros gastos", False),
        ("Total salidas", True), ("", False),
        ("Flujo neto mensual", True), ("Saldo acumulado", True),
    ]
    r += 1
    for label, bold in concepts:
        sdc(ws, r, 1, label, font=bld_font if bold else dat_font)
        for c in range(2, 14):
            sdc(ws, r, c, None, fill=inp_fill if not bold and label else None, fmt=cur_fmt)
        r += 1
    r += 1; brand_footer(ws, r, 13)

    # Break-even
    ws2 = wb.create_sheet("Break-Even"); ws2.sheet_properties.tabColor = "FF5722"
    title_block(ws2, "Calculadora Break-Even — Restaurante Mexicano", 4)
    r = 4
    params = [
        ("Ticket medio (€)", 24),
        ("Comensales/día promedio", 85),
        ("Días abierto/mes", 26),
        ("Food cost (%)", 0.31),
        ("Costes fijos mensuales (€)", 39100),
    ]
    for label, val in params:
        r += 1
        sdc(ws2, r, 1, label, font=bld_font); sdc(ws2, r, 2, val, fill=inp_fill, fmt=cur_fmt if isinstance(val, int) else pct_fmt)
    r += 2
    sdc(ws2, r, 1, "Facturación mensual estimada", font=bld_font)
    sdc(ws2, r, 2, "=B5*B6*B7", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Margen contribución mensual", font=bld_font)
    sdc(ws2, r, 2, "=B12*(1-B8)", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws2, r, 1, "Break-Even (meses)", font=bld_font)
    ws2.column_dimensions['A'].width = 35; ws2.column_dimensions['B'].width = 20
    r += 2; brand_footer(ws2, r, 4)
    path = os.path.join(OUTPUT_DIR, "cash-flow-break-even.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_escandallo():
    wb = Workbook()
    instr_sheet(wb, "Escandallo Maestro — Restaurante Mexicano", [
        "Una ficha técnica por plato.",
        "Introduce ingredientes, cantidades y precios.",
        "El food cost se calcula automáticamente.",
    ])
    ws = wb.create_sheet("Escandallo"); ws.sheet_properties.tabColor = "E91E63"
    title_block(ws, "Escandallo Maestro — Fichas Técnicas de Platos Mexicanos")
    hdrs = ["#", "Ingrediente", "Cantidad (g/ml)", "Precio/Kg (€)", "Coste (€)", "Merma (%)", "Coste Real (€)"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 7)
    widths = [5, 30, 16, 14, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "NOMBRE DEL PLATO:", font=bld_font)
    sdc(ws, r, 3, "Ejemplo: Tacos al Pastor (3 uds)", fill=inp_fill)
    r += 1
    ingredients = [
        ("Tortillas de maíz nixtamalizado (3 uds)", 90, 3.00, 2),
        ("Carne de cerdo adobada al pastor", 180, 8.50, 8),
        ("Piña fresca (para trompo)", 40, 2.00, 15),
        ("Cebolla blanca picada", 20, 1.50, 10),
        ("Cilantro fresco", 10, 6.00, 20),
        ("Limón (zumo)", 15, 3.00, 30),
        ("Chile guajillo (adobo)", 5, 18.00, 0),
        ("Achiote (pasta)", 3, 15.00, 0),
        ("Salsa verde taquera", 30, 4.00, 0),
    ]
    start_r = r
    for idx, (ing, cant, precio, merma) in enumerate(ingredients, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, ing, fill=inp_fill)
        sdc(ws, r, 3, cant, fill=inp_fill)
        sdc(ws, r, 4, precio, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=(C{r}/1000)*D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 6, merma / 100, fill=inp_fill, fmt=pct_fmt)
        sdc(ws, r, 7, f"=E{r}*(1+F{r})", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 2, "COSTE TOTAL PLATO", font=bld_font)
    sdc(ws, r, 7, f"=SUM(G{start_r}:G{r-2})", font=bld_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP sugerido (food cost 30%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}/0.30", font=frm_font, fmt=cur_fmt)
    r += 1
    sdc(ws, r, 2, "PVP con IVA (10%)", font=bld_font)
    sdc(ws, r, 7, f"=G{r-1}*1.10", font=frm_font, fmt=cur_fmt)
    r += 2; brand_footer(ws, r, 7)
    path = os.path.join(OUTPUT_DIR, "escandallo-maestro-mexicano.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_menu_engineering():
    wb = Workbook()
    instr_sheet(wb, "Menú Engineering Matrix — Restaurante Mexicano", [
        "Introduce los platos de tu carta con ventas y food cost.",
        "La matrix clasifica automáticamente cada plato.",
        "Stars = mantener. Plowhorses = subir precio. Puzzles = promocionar. Dogs = eliminar.",
    ])
    ws = wb.create_sheet("Matrix"); ws.sheet_properties.tabColor = "9C27B0"
    title_block(ws, "Menú Engineering Matrix — Restaurante Mexicano", 10)
    hdrs = ["#", "Plato", "PVP (€)", "Food Cost (€)", "% Food Cost", "Uds. Vendidas/Mes", "Margen Unit. (€)", "Margen Total (€)", "Popularidad", "Clasificación"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    widths = [5, 30, 10, 12, 10, 16, 12, 14, 12, 14]
    for i, w in enumerate(widths, 1): ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = f"A{r+1}"
    platos = [
        ("Tacos al pastor (3 uds)", 4.50, 1.20, 150),
        ("Guacamole fresco en molcajete", 8.00, 1.80, 120),
        ("Enchiladas suizas (3 uds)", 10.50, 1.60, 70),
        ("Burrito de carne asada", 11.00, 2.20, 90),
        ("Nachos loaded con queso y jalapeños", 9.00, 1.50, 100),
        ("Quesadillas flor de calabaza", 7.00, 1.00, 55),
        ("Mole poblano con pollo", 13.00, 2.50, 40),
        ("Ceviche mexicano de lubina", 13.50, 3.50, 35),
        ("Margarita clásica", 9.00, 1.50, 200),
        ("Churros con chocolate (6 uds)", 6.50, 0.80, 85),
        ("Pozole rojo con cerdo", 11.00, 2.00, 30),
        ("Chile relleno de queso", 11.00, 1.80, 25),
        ("Elotes con mayo, chile y limón", 5.50, 0.60, 65),
        ("Carnitas (3 tacos)", 4.50, 1.40, 110),
        ("Flan de cajeta", 6.00, 0.90, 50),
    ]
    r += 1
    for idx, (plato, pvp, fc, uds) in enumerate(platos, 1):
        sdc(ws, r, 1, idx, align=c_align)
        sdc(ws, r, 2, plato, fill=inp_fill)
        sdc(ws, r, 3, pvp, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 4, fc, fill=inp_fill, fmt=cur_fmt)
        sdc(ws, r, 5, f"=D{r}/C{r}", font=frm_font, fmt=pct_fmt)
        sdc(ws, r, 6, uds, fill=inp_fill)
        sdc(ws, r, 7, f"=C{r}-D{r}", font=frm_font, fmt=cur_fmt)
        sdc(ws, r, 8, f"=G{r}*F{r}", font=frm_font, fmt=cur_fmt)
        r += 1
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "menu-engineering-matrix.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_calculadora_ticket():
    wb = Workbook()
    instr_sheet(wb, "Calculadora Ticket Medio + Margen Tequila/Cócteles", [
        "Simula diferentes escenarios de ticket medio.",
        "Incluye pestaña separada para margen de barra de tequilas.",
        "Rellena las celdas verdes con tus datos.",
    ])
    ws = wb.create_sheet("Ticket Medio"); ws.sheet_properties.tabColor = "00BCD4"
    title_block(ws, "Calculadora Ticket Medio — Restaurante Mexicano", 4)
    r = 4
    params = [
        ("% clientes que piden antojitos/compartir", 0.55),
        ("Precio medio antojitos (€)", 8.00),
        ("Precio medio principal (€)", 10.50),
        ("% clientes que piden postre", 0.30),
        ("Precio medio postre (€)", 6.00),
        ("% clientes que piden tequila/margarita", 0.45),
        ("Precio medio tequila/cóctel (€)", 9.00),
        ("% clientes que piden cerveza/refresco", 0.50),
        ("Precio medio cerveza/refresco (€)", 4.00),
    ]
    r += 1
    for label, val in params:
        sdc(ws, r, 1, label, font=dat_font)
        sdc(ws, r, 2, val, fill=inp_fill, fmt=pct_fmt if isinstance(val, float) and val < 1 else cur_fmt)
        r += 1
    r += 1
    sdc(ws, r, 1, "TICKET MEDIO ESTIMADO", font=bld_font)
    sdc(ws, r, 2, None, font=bld_font, fmt=cur_fmt)
    ws.column_dimensions['A'].width = 45; ws.column_dimensions['B'].width = 18

    # Margen tequilas tab
    ws2 = wb.create_sheet("Margen Tequilas"); ws2.sheet_properties.tabColor = "FFC107"
    title_block(ws2, "Simulador Margen Barra de Tequilas y Cócteles", 4)
    r = 4
    sdc(ws2, r+1, 1, "Precio medio chupito tequila (€)", font=bld_font); sdc(ws2, r+1, 2, 8.00, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+2, 1, "Coste medio chupito (€)", font=bld_font); sdc(ws2, r+2, 2, 1.50, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+3, 1, "Margen por chupito (€)", font=bld_font); sdc(ws2, r+3, 2, f"=B{r+1}-B{r+2}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+4, 1, "% Margen tequila", font=bld_font); sdc(ws2, r+4, 2, f"=B{r+3}/B{r+1}", font=frm_font, fmt=pct_fmt)
    sdc(ws2, r+5, 1, "Precio medio margarita (€)", font=bld_font); sdc(ws2, r+5, 2, 9.50, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+6, 1, "Coste medio margarita (€)", font=bld_font); sdc(ws2, r+6, 2, 1.80, fill=inp_fill, fmt=cur_fmt)
    sdc(ws2, r+7, 1, "Margen por margarita (€)", font=bld_font); sdc(ws2, r+7, 2, f"=B{r+5}-B{r+6}", font=frm_font, fmt=cur_fmt)
    sdc(ws2, r+8, 1, "Chupitos vendidos/día", font=bld_font); sdc(ws2, r+8, 2, 25, fill=inp_fill)
    sdc(ws2, r+9, 1, "Margaritas vendidas/día", font=bld_font); sdc(ws2, r+9, 2, 30, fill=inp_fill)
    sdc(ws2, r+10, 1, "Días/mes", font=bld_font); sdc(ws2, r+10, 2, 26, fill=inp_fill)
    sdc(ws2, r+11, 1, "Facturación barra tequilas/mes (€)", font=bld_font)
    sdc(ws2, r+11, 2, f"=(B{r+1}*B{r+8}+B{r+5}*B{r+9})*B{r+10}", font=frm_font, fmt=cur_fmt)
    ws2.column_dimensions['A'].width = 40; ws2.column_dimensions['B'].width = 18
    path = os.path.join(OUTPUT_DIR, "calculadora-ticket-medio.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_cronograma():
    wb = Workbook()
    instr_sheet(wb, "Cronograma Apertura Gantt 12 Meses — Restaurante Mexicano", [
        "Fases y tareas para abrir un restaurante mexicano.",
        "Marca las celdas con 'X' para indicar el mes activo.",
        "Incluye fase de setup de importación y proveedores mexicanos.",
    ])
    ws = wb.create_sheet("Gantt"); ws.sheet_properties.tabColor = "795548"
    title_block(ws, "Cronograma de Apertura — Restaurante Mexicano 80 Plazas", 15)
    hdrs = ["#", "Fase", "Tarea", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "M8", "M9", "M10", "M11", "M12"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 15)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 45
    for i in range(4, 16): ws.column_dimensions[get_column_letter(i)].width = 5
    tasks = [
        ("Planificación", "Estudio de mercado y viabilidad", [1, 2]),
        ("Planificación", "Plan financiero y búsqueda financiación", [1, 2, 3]),
        ("Planificación", "Constitución empresa / alta autónomo", [2, 3]),
        ("Importación", "Registro EORI y trámites importación", [2, 3, 4]),
        ("Importación", "Contacto distribuidores producto mexicano", [2, 3]),
        ("Importación", "Selección proveedores tequila/mezcal", [3, 4]),
        ("Local", "Búsqueda y selección de local", [2, 3, 4]),
        ("Local", "Negociación alquiler y contrato", [3, 4]),
        ("Licencias", "Licencia actividad / declaración responsable", [3, 4, 5]),
        ("Licencias", "Proyecto técnico y visado", [3, 4]),
        ("Obra", "Acondicionamiento y obra civil", [4, 5, 6, 7]),
        ("Obra", "Instalaciones (electricidad, gas, extracción)", [5, 6, 7]),
        ("Obra", "Decoración mexicana (Talavera, mural, etc.)", [6, 7]),
        ("Equipamiento", "Pedido equipo cocina + comal + tortillera", [5, 6]),
        ("Equipamiento", "Instalación barra tequilas + expositor", [6, 7]),
        ("Equipamiento", "Recepción e instalación equipos", [7, 8]),
        ("Equipamiento", "Mobiliario sala + terraza", [6, 7]),
        ("Equipo", "Búsqueda taquero y parrillero", [6, 7, 8]),
        ("Equipo", "Selección y contratación resto equipo", [7, 8, 9]),
        ("Equipo", "Formación equipo (cocina mexicana + tequilas)", [9, 10]),
        ("Marketing", "Branding, web, redes sociales", [6, 7, 8]),
        ("Marketing", "Fotos profesionales de platos", [9]),
        ("Marketing", "Campaña pre-apertura", [9, 10]),
        ("APPCC", "Plan APPCC + registro sanitario", [7, 8]),
        ("Pre-apertura", "Soft opening (amigos y familia)", [10]),
        ("Pre-apertura", "Ajustes finales", [10, 11]),
        ("Apertura", "INAUGURACIÓN", [11]),
        ("Post-apertura", "Seguimiento y ajustes", [11, 12]),
    ]
    r += 1
    for idx, (fase, tarea, meses) in enumerate(tasks, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, fase); sdc(ws, r, 3, tarea)
        for m in meses:
            sdc(ws, r, m + 3, "X", font=bld_font, fill=sec_fill, align=c_align)
        r += 1
    r += 1; brand_footer(ws, r, 15)
    path = os.path.join(OUTPUT_DIR, "cronograma-apertura-gantt.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_plantilla_turnos():
    wb = Workbook()
    instr_sheet(wb, "Plantilla Turnos Brigada — Restaurante Mexicano", [
        "Cuadrante semanal para todo el equipo (cocina + sala).",
        "Rellena los turnos con: M (mañana), T (tarde), P (partido), L (libre).",
        "Incluye puestos específicos: taquero, parrillero, barman tequilas.",
    ])
    ws = wb.create_sheet("Turnos Semana"); ws.sheet_properties.tabColor = "607D8B"
    title_block(ws, "Cuadrante Semanal — Restaurante Mexicano 80 Plazas", 10)
    hdrs = ["#", "Nombre", "Puesto", "Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]
    r = 4
    for i, h in enumerate(hdrs, 1): ws.cell(row=r, column=i, value=h)
    shr(ws, r, 10)
    ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 22; ws.column_dimensions['C'].width = 24
    for i in range(4, 11): ws.column_dimensions[get_column_letter(i)].width = 8
    dv = DataValidation(type="list", formula1='"M,T,P,L,V"', allow_blank=True)
    ws.add_data_validation(dv)
    staff = [
        ("Jefe de cocina", "P"), ("Taquero", "P"), ("Parrillero", "P"),
        ("Salsero / cocinero frío", "P"), ("Cocinero 1", "P"), ("Cocinero 2", "M"),
        ("Ayudante cocina", "P"), ("Office", "P"),
        ("Encargado sala", "P"), ("Barman tequilas", "T"),
        ("Camarero 1", "P"), ("Camarero 2", "P"), ("Camarero 3", "T"),
        ("Runner", "P"),
    ]
    r += 1
    for idx, (puesto, turno_default) in enumerate(staff, 1):
        sdc(ws, r, 1, idx, align=c_align); sdc(ws, r, 2, None, fill=inp_fill); sdc(ws, r, 3, puesto)
        for c in range(4, 11):
            sdc(ws, r, c, turno_default, fill=inp_fill, align=c_align)
            dv.add(ws.cell(row=r, column=c))
        r += 1
    r += 1
    sdc(ws, r, 1, None); sdc(ws, r, 2, "M=Mañana, T=Tarde, P=Partido, L=Libre, V=Vacaciones", font=brn_font)
    r += 2; brand_footer(ws, r, 10)
    path = os.path.join(OUTPUT_DIR, "plantilla-turnos-brigada.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 3. CHECKLISTS (6)
# ═══════════════════════════════════════════════════════════

def gen_checklist_legal():
    wb = Workbook(); ws = wb.active
    items = [
        ("Empresa", "Decidir forma jurídica (SL, autónomo, cooperativa)", "Socio/Gestor", 500),
        ("Empresa", "Constituir sociedad ante notario", "Socio/Gestor", 600),
        ("Empresa", "Inscripción en Registro Mercantil", "Gestor", 200),
        ("Empresa", "Obtener CIF provisional", "Gestor", 0),
        ("Empresa", "Alta IAE epígrafe 671.4 o similar", "Gestor", 0),
        ("Empresa", "Alta censal en Hacienda (modelo 036/037)", "Gestor", 0),
        ("Empresa", "Inscripción Seguridad Social empresa", "Gestor", 0),
        ("Empresa", "Abrir cuenta bancaria empresa", "Socio", 0),
        ("Licencias", "Licencia de actividad / declaración responsable", "Arquitecto", 2000),
        ("Licencias", "Proyecto técnico y visado colegial", "Arquitecto", 3000),
        ("Licencias", "Licencia de obras (si aplica)", "Arquitecto", 1500),
        ("Licencias", "Licencia de terraza", "Socio", 1500),
        ("Licencias", "Licencia de música / SGAE / AGEDI", "Socio", 500),
        ("Importación", "Registro EORI (operador comercio exterior)", "Gestor", 0),
        ("Importación", "Certificado SOIVRE para alimentos importados extra-UE", "Gestor", 300),
        ("Importación", "Etiquetado en español para productos importados", "Proveedor", 200),
        ("Importación", "Control fitosanitario chiles secos y especias", "Gestor", 150),
        ("Importación", "Impuestos especiales bebidas alcohólicas (IIEE) tequila/mezcal", "Gestor", 0),
        ("Importación", "Acuerdo con distribuidor especializado en producto mexicano", "Socio", 0),
        ("Sanidad", "Registro General Sanitario (RGSEAA)", "Gestor", 0),
        ("Sanidad", "Plan APPCC documentado (adaptado cocina mexicana)", "Consultor", 1500),
        ("Sanidad", "Formación manipulador alimentos (equipo)", "Consultor", 300),
        ("Sanidad", "Contrato empresa DDD (desratización)", "Socio", 600),
        ("Sanidad", "Control de aguas", "Socio", 200),
        ("Sanidad", "Plan de alérgenos (14 alérgenos + específicos mexicanos)", "Jefe cocina", 0),
        ("Seguros", "Seguro responsabilidad civil (min. 300K€)", "Socio", 800),
        ("Seguros", "Seguro de local (incendio, robo, daños)", "Socio", 600),
        ("Seguros", "Seguro de accidentes convenio", "Gestor", 400),
        ("Laboral", "Contratos de trabajo según convenio", "Gestor", 0),
        ("Laboral", "Prevención de riesgos laborales", "SPA", 500),
        ("Laboral", "Calendario laboral y cuadrante turnos", "Encargado", 0),
        ("Protección datos", "Registro LOPD / RGPD", "Gestor", 300),
        ("Protección datos", "Política de cookies web", "Socio", 0),
        ("Protección datos", "Videovigilancia (si aplica): cartelería + registro", "Socio", 100),
        ("Otros", "Libro de reclamaciones", "Socio", 30),
        ("Otros", "Cartel de horarios en puerta", "Socio", 0),
        ("Otros", "Aforo máximo visible", "Socio", 0),
        ("Otros", "Señalética salida emergencia + extintores", "Arquitecto", 500),
    ]
    checklist_ws(ws, "Checklist Legal — Restaurante Mexicano España", items)
    path = os.path.join(OUTPUT_DIR, "checklist-legal.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_equipamiento():
    wb = Workbook(); ws = wb.active
    items = [
        ("Cocina MX", "Comal profesional grande (80-100 cm, hierro fundido)", "Jefe cocina", 1500),
        ("Cocina MX", "Prensa de tortillas manual profesional", "Jefe cocina", 300),
        ("Cocina MX", "Tortillera eléctrica semiautomática (opcional alto volumen)", "Jefe cocina", 4000),
        ("Cocina MX", "Trompo de pastor vertical (gas)", "Jefe cocina", 1000),
        ("Cocina MX", "Molcajetes grandes piedra volcánica (6 uds)", "Jefe cocina", 600),
        ("Cocina MX", "Ahumador / smoker profesional", "Jefe cocina", 3000),
        ("Cocina MX", "Parrilla carbón/gas profesional", "Jefe cocina", 3500),
        ("Cocina MX", "Licuadoras industriales para salsas (3 uds)", "Jefe cocina", 500),
        ("Cocina MX", "Máquina de margaritas / granizados", "Encargado", 2000),
        ("Cocina MX", "Vaporera industrial para tamales", "Jefe cocina", 400),
        ("Cocción", "Cocina gas 6 fuegos + horno", "Jefe cocina", 5000),
        ("Cocción", "Horno mixto (Rational / Unox)", "Jefe cocina", 12000),
        ("Cocción", "Freidora doble cuba (churros, totopos, empanadas)", "Jefe cocina", 2000),
        ("Cocción", "Salamandra", "Jefe cocina", 1200),
        ("Frío", "Cámara frigorífica 2 puertas", "Jefe cocina", 3500),
        ("Frío", "Cámara congelación", "Jefe cocina", 3000),
        ("Frío", "Mesa refrigerada bajo barra", "Jefe cocina", 2500),
        ("Frío", "Botellero refrigerado barra tequilas", "Encargado", 1800),
        ("Preparación", "Mesa de trabajo acero inox (2-3 uds)", "Jefe cocina", 2000),
        ("Preparación", "Robot de cocina / cutter", "Jefe cocina", 1500),
        ("Preparación", "Prensador de limones industrial", "Jefe cocina", 200),
        ("Lavado", "Tren de lavado / lavavajillas industrial", "Jefe cocina", 5000),
        ("Lavado", "Fregadero doble seno inox", "Jefe cocina", 800),
        ("Extracción", "Campana extractora reforzada + filtros", "Instalador", 6000),
        ("Extracción", "Salida de humos reglamentaria (sobredimensionada)", "Instalador", 4000),
        ("Menaje", "Ollas grandes para pozole, mole, frijoles", "Jefe cocina", 1000),
        ("Menaje", "Sartenes de hierro para tortillas", "Jefe cocina", 300),
        ("Menaje", "Cuchillos profesionales", "Jefe cocina", 500),
        ("Menaje", "Gastronormas GN 1/1, 1/2, 1/3", "Jefe cocina", 600),
        ("Menaje", "Tablas de corte por colores", "Jefe cocina", 150),
        ("Almacenamiento", "Estanterías inox economato seco (chiles, especias)", "Jefe cocina", 800),
        ("Almacenamiento", "Contenedores herméticos para chiles secos", "Jefe cocina", 300),
        ("Barra", "Expositor iluminado botellas tequila/mezcal", "Encargado", 4000),
        ("Barra", "Coctelera, medidores, utensilios barra", "Encargado", 500),
        ("Barra", "Máquina de hielo (alto consumo por cócteles)", "Encargado", 2000),
    ]
    checklist_ws(ws, "Checklist Equipamiento Cocina Mexicana", items)
    path = os.path.join(OUTPUT_DIR, "checklist-equipamiento-cocina-mexicana.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_appcc():
    wb = Workbook(); ws = wb.active
    items = [
        ("Prerrequisitos", "Plan de limpieza y desinfección documentado", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan DDD contratado (empresa externa)", "Socio", 600),
        ("Prerrequisitos", "Plan de control de agua potable", "Socio", 200),
        ("Prerrequisitos", "Plan de formación del personal", "Jefe cocina", 300),
        ("Prerrequisitos", "Plan de gestión de residuos", "Socio", 0),
        ("Prerrequisitos", "Plan de mantenimiento de instalaciones", "Socio", 0),
        ("Prerrequisitos", "Plan de trazabilidad (registro proveedores + lotes)", "Jefe cocina", 0),
        ("Prerrequisitos", "Plan de control de alérgenos (14 alérgenos)", "Jefe cocina", 0),
        ("Temperaturas", "Termómetro sonda calibrado", "Jefe cocina", 50),
        ("Temperaturas", "Registro diario temperaturas cámaras", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas servicio (caliente >65°C)", "Cocinero", 0),
        ("Temperaturas", "Registro temperaturas recepción mercancía", "Cocinero", 0),
        ("Temperaturas", "Protocolo cadena de frío (0-4°C refrigerado, -18°C congelado)", "Jefe cocina", 0),
        ("Temperaturas", "Control temperatura trompo de pastor (interna >75°C)", "Taquero", 0),
        ("Temperaturas", "Control temperatura comal para tortillas (>200°C)", "Taquero", 0),
        ("Prod. Importados", "Certificado fitosanitario chiles secos en archivo", "Jefe cocina", 0),
        ("Prod. Importados", "Etiquetado en español de todos los productos importados", "Jefe cocina", 0),
        ("Prod. Importados", "Trazabilidad de origen de chiles, masa, especias", "Jefe cocina", 0),
        ("Prod. Importados", "Almacenamiento separado chiles secos (seco, ventilado)", "Cocinero", 0),
        ("Prod. Importados", "Control de insectos en productos secos importados", "Cocinero", 0),
        ("Prod. Importados", "Registro de lotes de tequila y mezcal importado", "Encargado", 0),
        ("Higiene", "Lavamanos no manual en cocina con jabón y papel", "Socio", 300),
        ("Higiene", "Vestuarios con taquillas para personal", "Socio", 500),
        ("Higiene", "Uniformes de cocina (gorro, chaqueta, pantalón, calzado)", "Socio", 600),
        ("Higiene", "Protocolo lavado de manos (cartelería)", "Jefe cocina", 0),
        ("Higiene", "Contenedores de basura con tapa y pedal", "Socio", 200),
        ("Salsas frescas", "Vida útil salsas frescas: máx 24h refrigerado", "Jefe cocina", 0),
        ("Salsas frescas", "Guacamole: producción máxima 4h antes del servicio", "Salsero", 0),
        ("Salsas frescas", "Pico de gallo: producción diaria, no reutilizar", "Salsero", 0),
        ("Pescado", "Congelación previa -20°C/24h para ceviche (anisakis)", "Jefe cocina", 0),
        ("Almacenamiento", "Sistema FIFO/FEFO en cámaras y economato", "Cocinero", 0),
        ("Almacenamiento", "Separación crudo/cocinado en cámaras", "Cocinero", 0),
        ("Almacenamiento", "Masa nixtamal refrigerada máx 48h", "Taquero", 0),
        ("Documentación", "Plan APPCC escrito y disponible en cocina", "Socio", 1500),
        ("Documentación", "Registros de control actualizados diariamente", "Jefe cocina", 0),
        ("Documentación", "Certificados manipulador alimentos de todo el personal", "Socio", 0),
        ("Documentación", "Contrato DDD visible y registros de actuaciones", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist APPCC — Restaurante Mexicano", items)
    path = os.path.join(OUTPUT_DIR, "checklist-appcc.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_sala():
    wb = Workbook(); ws = wb.active
    items = [
        ("Distribución", "Plano de sala con posición de mesas definitivo", "Socio/Arquitecto", 0),
        ("Distribución", "80 plazas distribuidas (mesas 2, 4, barra alta)", "Socio", 0),
        ("Distribución", "Distancia entre mesas: mínimo 60-80 cm", "Arquitecto", 0),
        ("Distribución", "Flujo de servicio definido (entrada → sala → barra → aseos)", "Socio", 0),
        ("Distribución", "Zona de espera / recepción", "Socio", 500),
        ("Distribución", "Accesibilidad PMR (rampa, aseo adaptado)", "Arquitecto", 2000),
        ("Barra Tequilas", "Barra de tequilas con expositor iluminado trasero", "Interiorista", 8000),
        ("Barra Tequilas", "Taburetes altos cómodos para barra", "Socio", 3000),
        ("Barra Tequilas", "Iluminación ambiental en barra (luz cálida, neón sutil)", "Electricista", 1000),
        ("Barra Tequilas", "Carta de tequilas y mezcales enmarcada o en pizarra", "Diseñador", 300),
        ("Decoración MX", "Azulejos de Talavera en barra y/o paredes destacadas", "Interiorista", 3000),
        ("Decoración MX", "Mural artístico mexicano (Día de Muertos, paisaje, Frida-inspired)", "Artista", 4000),
        ("Decoración MX", "Vajilla de barro mexicano (platos negros Oaxaca o barro rojo)", "Socio", 2000),
        ("Decoración MX", "Paleta de colores: rosa mexicano, azul cobalto, naranja", "Interiorista", 0),
        ("Decoración MX", "Textiles mexicanos (individuales, cojines bancos corridos)", "Socio", 800),
        ("Decoración MX", "Cactus y suculentas reales en entrada y barra", "Socio", 300),
        ("Decoración MX", "Papel picado decorativo (techo o pared)", "Socio", 200),
        ("Iluminación", "Iluminación general cálida 2700K", "Electricista", 2000),
        ("Iluminación", "Iluminación regulable por zonas (dimmer)", "Electricista", 500),
        ("Iluminación", "Lámparas de latón o vidrio soplado estilo mexicano", "Socio", 1500),
        ("Acústica", "Paneles absorbentes si techos altos", "Interiorista", 1500),
        ("Acústica", "Sistema de música (ranchera suave, cumbia, son jarocho)", "Socio", 800),
        ("Mobiliario", "Mesas sólidas (madera, hierro, o combinación)", "Socio", 10000),
        ("Mobiliario", "Sillas con asiento tapizado colores mexicanos", "Socio", 8000),
        ("Mobiliario", "Bancos corridos con cojines (si espacio lo permite)", "Socio", 3000),
        ("Terraza", "Mobiliario exterior en colores vivos", "Socio", 6000),
        ("Terraza", "Sombrillas o toldos con colores mexicanos", "Socio", 2500),
        ("Terraza", "Plantas / macetas de barro", "Socio", 500),
        ("Aseos", "Aseos diferenciados y aseo PMR", "Arquitecto", 3000),
        ("Aseos", "Decoración mexicana en aseos (Talavera, espejo artesanal)", "Socio", 500),
        ("Señalética", "Señalética interior temática mexicana", "Diseñador", 300),
        ("Señalética", "Rótulo exterior / fachada con identidad mexicana", "Socio", 2500),
    ]
    checklist_ws(ws, "Checklist Diseño Sala Mexicana + Barra Tequilas", items)
    path = os.path.join(OUTPUT_DIR, "checklist-diseno-sala-mexicana.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_contratacion():
    wb = Workbook(); ws = wb.active
    items = [
        ("Perfiles", "Definir organigrama completo (cocina + sala)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Jefe de cocina (experiencia cocina mexicana)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Taquero (estación comal + tortillas)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Parrillero (parrilla + trompo pastor)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Salsero / cocinero frío", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Cocinero (guisos, moles, arroces)", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Ayudante / office", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Encargado de sala", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Barman de tequilas y coctelería", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Camarero", "Socio", 0),
        ("Perfiles", "Ficha de puesto: Runner", "Socio", 0),
        ("Selección", "Publicar ofertas (InfoJobs, Hostelwork, LinkedIn)", "Socio", 200),
        ("Selección", "Búsqueda taquero en comunidades mexicanas España", "Socio", 0),
        ("Selección", "Selección jefe de cocina con experiencia mexicana", "Socio", 0),
        ("Selección", "Selección barman con conocimiento de tequila/mezcal", "Socio", 0),
        ("Selección", "Selección y entrevistas resto equipo cocina", "Jefe cocina", 0),
        ("Selección", "Selección y entrevistas camareros", "Encargado", 0),
        ("Selección", "Pruebas prácticas cocina (elaborar tacos + salsa)", "Jefe cocina", 0),
        ("Documentación", "Contratos según convenio hostelería CCAA", "Gestor", 0),
        ("Documentación", "Alta Seguridad Social", "Gestor", 0),
        ("Documentación", "Certificado manipulador alimentos", "Empleado", 0),
        ("Documentación", "Prevención riesgos laborales (formación)", "SPA", 500),
        ("Onboarding", "Manual de bienvenida / handbook del restaurante", "Socio", 0),
        ("Onboarding", "Formación en carta mexicana y alérgenos", "Jefe cocina", 0),
        ("Onboarding", "Formación en tequilas y mezcales (todo equipo sala)", "Barman", 0),
        ("Onboarding", "Formación en TPV y sistema de pedidos", "Encargado", 0),
        ("Onboarding", "Formación APPCC y protocolos (incluye producto importado)", "Jefe cocina", 0),
        ("Onboarding", "Ensayos de servicio (soft opening)", "Socio", 0),
    ]
    checklist_ws(ws, "Checklist Contratación — Restaurante Mexicano", items)
    path = os.path.join(OUTPUT_DIR, "checklist-contratacion.xlsx")
    wb.save(path); print(f"✓ {path}")

def gen_checklist_marketing():
    wb = Workbook(); ws = wb.active
    items = [
        ("Branding", "Nombre del restaurante definitivo (verificar disponibilidad)", "Socio", 0),
        ("Branding", "Logo profesional con identidad mexicana", "Diseñador", 500),
        ("Branding", "Manual de marca (colores mexicanos, tipografía, estilo)", "Diseñador", 300),
        ("Branding", "Diseño de carta / menú físico con ilustraciones mexicanas", "Diseñador", 500),
        ("Branding", "Packaging delivery con marca y colores mexicanos", "Diseñador", 300),
        ("Digital", "Web propia con menú, carta tequilas, fotos y reservas", "Dev/Agencia", 1500),
        ("Digital", "Google My Business: ficha completa y verificada", "Socio", 0),
        ("Digital", "Instagram: perfil profesional con 9+ publicaciones", "Community", 0),
        ("Digital", "TikTok: perfil activo con vídeos del taquero, salsas, cocktails", "Community", 0),
        ("Digital", "TheFork / TripAdvisor: perfiles completos con fotos", "Socio", 0),
        ("Digital", "Google Ads local configurado", "Agencia", 500),
        ("Digital", "Instagram/Facebook Ads configurados", "Agencia", 400),
        ("Contenido", "Sesión fotográfica profesional (platos + decoración + tequilas)", "Fotógrafo", 400),
        ("Contenido", "Vídeo taquero preparando tacos (30-60 seg para TikTok/Reels)", "Videógrafo", 500),
        ("Contenido", "Calendario de publicaciones primer mes", "Community", 0),
        ("Eventos MX", "Plan Día de Muertos (noviembre): decoración altar + menú especial", "Socio", 500),
        ("Eventos MX", "Plan Cinco de Mayo: margaritas especiales + promociones", "Socio", 300),
        ("Eventos MX", "Plan Día Independencia (16 sep): pozole + chiles en nogada", "Socio", 300),
        ("Pre-apertura", "Evento soft opening (amigos, familia, influencers food)", "Socio", 500),
        ("Pre-apertura", "Invitación a prensa local / bloggers food", "Socio", 200),
        ("Pre-apertura", "Flyers / cartelería barrio", "Diseñador", 200),
        ("Pre-apertura", "Colaboraciones con negocios vecinos", "Socio", 0),
        ("Delivery", "Alta en Glovo / Uber Eats / Just Eat", "Socio", 0),
        ("Delivery", "Fotos de platos para plataformas delivery", "Fotógrafo", 0),
        ("Delivery", "Menú delivery optimizado (burritos, bowls, taco kits)", "Jefe cocina", 0),
        ("Fidelización", "Programa de fidelización (tarjeta, app o CRM)", "Socio", 300),
        ("Fidelización", "Estrategia de reseñas (pedir valoraciones a clientes)", "Encargado", 0),
    ]
    checklist_ws(ws, "Checklist Marketing Pre-Apertura — Restaurante Mexicano", items)
    path = os.path.join(OUTPUT_DIR, "checklist-marketing-preapertura.xlsx")
    wb.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 4. DOCUMENT MODELS (2)
# ═══════════════════════════════════════════════════════════

def gen_business_plan():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Business Plan"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Mexicano — 80 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("[Ciudad, Mes Año]\nPlantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Resumen Ejecutivo", [
            "Descripción del proyecto: restaurante mexicano auténtico de 80 plazas en [ciudad].",
            "Modelo: [taquería gourmet / cantina mexicana / casual mexicano / dark kitchen].",
            "Inversión total requerida: [120.000-280.000€].",
            "Facturación prevista año 1: [500.000-750.000€].",
            "Break-even estimado: mes [8-14].",
            "Equipo fundador: [nombres y experiencia].",
            "Diferencial: cocina mexicana auténtica + barra de tequilas premium 30-50 referencias.",
        ]),
        ("2. El Concepto", [
            "Tipo de cocina: mexicana auténtica / fusión / cantina / taquería gourmet.",
            "Ticket medio objetivo: [18-30€] sin bebidas.",
            "Propuesta de valor: tortillas artesanas, salsas frescas, tequila premium.",
            "Público objetivo: 25-45 años, urbano, foodie, viajero, busca experiencias.",
            "Posicionamiento vs competencia: auténtico vs tex-mex de franquicia.",
        ]),
        ("3. Análisis de Mercado", [
            "Restaurantes mexicanos en España: +2.500, crecimiento 35% en 5 años.",
            "Menos del 10% ofrecen cocina auténtica mexicana — oportunidad clara.",
            "Tendencia tequila/mezcal premium: +25% anual en España.",
            "Competencia directa en zona: [listar 3-5 competidores con ticket medio].",
            "Gap de mercado identificado: [qué falta en la zona].",
        ]),
        ("4. Plan Operativo", [
            "Ubicación: [dirección, m², alquiler mensual].",
            "Distribución: 80 plazas interior + barra tequilas + [X] plazas terraza.",
            "Horario: [ej. 13:00-16:00 y 20:00-00:00, lunes cerrado].",
            "Equipo cocina: [6-10] personas (incluye taquero, parrillero, salsero).",
            "Equipo sala: [6-8] personas (incluye barman tequilas).",
            "Proveedores clave producto mexicano: [distribuidores, importadores].",
        ]),
        ("5. Plan Financiero", [
            "Inversión inicial desglosada: ver plantilla Excel adjunta.",
            "P&L mensual con 3 escenarios: ver plantilla Excel adjunta.",
            "Cash flow 12 meses: ver plantilla Excel adjunta.",
            "Break-even: [mes estimado].",
            "ROI estimado: [X%] en año [2-3].",
        ]),
        ("6. Plan de Marketing", [
            "Estrategia pre-apertura: [acciones 3 meses antes de abrir].",
            "Canales digitales: Google My Business, Instagram, TikTok, TheFork.",
            "Eventos culturales: Día de Muertos, Cinco de Mayo, Independencia México.",
            "Presupuesto marketing mensual: [600-1.500€/mes].",
            "Estrategia de delivery: burritos, bowls y taco kits para plataformas.",
        ]),
        ("7. Equipo Fundador", [
            "[Nombre]: [experiencia en hostelería/cocina mexicana, rol en el proyecto].",
            "[Nombre]: [experiencia complementaria, rol en el proyecto].",
            "Asesores externos: [gestor, arquitecto, consultor gastronómico].",
        ]),
        ("8. Necesidades de Financiación", [
            "Inversión total: [X€].",
            "Fondos propios: [X€] ([X%]).",
            "Financiación solicitada: [X€] ([X%]).",
            "Destino de los fondos: obra [X€], equipamiento [X€], barra tequilas [X€], maniobra [X€].",
            "Plan de amortización propuesto: [X años, cuota mensual X€].",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "business-plan-modelo.docx")
    doc.save(path); print(f"✓ {path}")

def gen_manual_operaciones():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)

    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual de Operaciones"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GOLD_RGB
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Restaurante Mexicano — 80 Plazas\n[Nombre del Restaurante]"); r2.font.size = Pt(16)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Plantilla AI Chef Pro · aichef.pro"); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    sections = [
        ("1. Apertura del Restaurante", [
            "Hora de llegada del equipo: [2.5 horas antes del servicio].",
            "Taquero: encender comal, preparar masa de tortillas, montar estación.",
            "Parrillero: encender parrilla/trompo, preparar adobos, cortar carne.",
            "Salsero: preparar salsas frescas del día (verde, roja, pico de gallo, guacamole).",
            "Encender equipos de cocina: hornos, freidora (totopos, churros).",
            "Verificar temperaturas de cámaras y registrar.",
            "Reponer mise en place: cilantro, cebolla, limones, chiles, toppings.",
            "Barra: verificar stock tequilas, preparar garnish para cócteles, hielo.",
            "Preparar sala: colocar mesas, cartas, servilletas, cubiertos.",
            "Verificar limpieza de aseos.",
            "Encender TPV, música mexicana ambiental, iluminación según turno.",
            "Briefing de equipo: platos del día, tequila destacado, reservas, alérgenos.",
        ]),
        ("2. Servicio de Mediodía", [
            "Horario: 13:00-16:00.",
            "Ofrecer menú del día mexicano (si aplica) + carta completa.",
            "Velocidad de servicio: antojitos en mesa en 8-10 min, principales en 15 min.",
            "Taquero: producción continua de tortillas durante servicio.",
            "Salsero: reponer salsas frescas según consumo.",
            "Gestión de turnos de mesa si hay espera.",
            "Cobro y despedida: sugerir postre (churros) y tequila digestivo.",
        ]),
        ("3. Servicio de Cena", [
            "Horario: 20:00-00:00.",
            "Carta completa + carta de cócteles y tequilas destacada.",
            "Iluminación reducida, música más ambiente (cumbia, son jarocho).",
            "Sugerir vuelos de tequila, margaritas, antojitos para compartir.",
            "Taquero: show cooking a la vista del cliente (viernes y sábado).",
            "Último pedido de cocina: [23:30].",
        ]),
        ("4. Estación del Taquero: Protocolo Diario", [
            "Preparar masa de maíz nixtamalizado (o recibir masa fresca del proveedor).",
            "Calentar comal a temperatura óptima (200-220°C).",
            "Preparar bolas de masa para tortillas (porciones de 30g).",
            "Prensar y cocer tortillas al momento durante servicio.",
            "Mantener mise en place de proteínas: pastor, carnitas, cochinita, barbacoa.",
            "Cortar carne del trompo al pastor al momento del pedido.",
            "Mantener toppings frescos: cebolla, cilantro, piña, limón.",
            "Limpiar comal entre servicios (raspar y aceitar).",
        ]),
        ("5. Barra de Tequilas: Protocolo", [
            "Verificar stock de todas las referencias (30-50 botellas).",
            "Preparar mise en place de cócteles: jarabe de agave, zumo limón fresco, sal, garnish.",
            "Llenar máquina de margaritas (si aplica).",
            "Preparar hielo suficiente para servicio (alto consumo en cócteles mexicanos).",
            "Sugerir vuelos de tequila (3 chupitos de categorías diferentes).",
            "Explicar diferencias entre blanco, reposado, añejo y mezcal al cliente.",
            "Registrar consumo de botellas para control de stock.",
            "Al cierre: limpiar barra, guardar garnish, cerrar botellas abiertas.",
        ]),
        ("6. Cierre del Restaurante", [
            "Cocina: apagar comal, parrilla, trompo, hornos, freidora.",
            "Guardar masa de tortillas sobrante en cámara (máx 48h).",
            "Desechar salsas frescas sobrantes (no reutilizar pico de gallo ni guacamole).",
            "Registrar temperaturas de cierre.",
            "Sala: limpiar mesas, barrer, recoger terraza.",
            "Barra: cerrar caja, cuadrar TPV, limpiar barra, guardar garnish.",
            "Sacar basuras a contenedores.",
            "Cerrar puertas, alarma, luces.",
        ]),
        ("7. Gestión de Delivery Mexicano", [
            "Zona de packaging separada del pase de sala.",
            "Burritos: envolver en aluminio + papel kraft.",
            "Taco kits: tortillas en bolsa térmica, proteínas y salsas en envases separados.",
            "Bowls: envase ancho con tapa hermética.",
            "Nachos: queso y toppings en compartimento separado.",
            "Verificar cada pedido antes de entregar al rider.",
            "Incluir servilletas, salsas extra, tarjeta del restaurante.",
            "Tiempo máximo de preparación: 15 minutos.",
            "Responder a reseñas negativas en plataformas en <24h.",
        ]),
        ("8. Protocolo de Reclamaciones", [
            "Escuchar al cliente sin interrumpir.",
            "Pedir disculpas y ofrecer solución inmediata.",
            "Si el plato está mal: retirar y ofrecer alternativa sin cargo.",
            "Si el picante es excesivo: ofrecer crema agria o agua de horchata para suavizar.",
            "Nunca discutir con el cliente delante de otros comensales.",
            "Registrar incidencia en libro interno para análisis.",
            "Si solicita hoja de reclamaciones: entregarla sin poner objeciones.",
        ]),
        ("9. Gestión de Alérgenos en Cocina Mexicana", [
            "14 alérgenos de declaración obligatoria (estándar UE).",
            "Alérgenos frecuentes en cocina mexicana: gluten (tortilla de trigo), lácteos (queso, crema), frutos de cáscara (mole con cacahuete/almendra), sésamo (en algunos panes), crustáceos (en ceviches).",
            "Carta con indicación de alérgenos por plato.",
            "El mole poblano contiene chocolate + frutos secos: informar SIEMPRE.",
            "Las tortillas de maíz son SIN GLUTEN (ventaja competitiva).",
            "Ante duda: confirmar SIEMPRE con jefe de cocina.",
            "Protocolo de cocina para platos sin alérgenos: utensilios separados, zona limpia.",
        ]),
    ]

    for title, bullets in sections:
        doc.add_heading(title, level=1)
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')
        doc.add_page_break()

    path = os.path.join(OUTPUT_DIR, "manual-operaciones-mexicano.docx")
    doc.save(path); print(f"✓ {path}")


# ═══════════════════════════════════════════════════════════
# 5. PDF placeholder
# ═══════════════════════════════════════════════════════════
def gen_guide_pdf():
    """Generate a minimal placeholder PDF using the DOCX content."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        path = os.path.join(OUTPUT_DIR, "guia-restaurante-mexicano.pdf")
        c = canvas.Canvas(path, pagesize=A4)
        c.setFont("Helvetica-Bold", 24)
        c.drawCentredString(297, 600, "Como Montar un Restaurante Mexicano")
        c.setFont("Helvetica", 16)
        c.drawCentredString(297, 560, "80 Plazas de Aforo - Guia Espana 2026")
        c.setFont("Helvetica", 12)
        c.drawCentredString(297, 520, "Chef John Guerrero - AI Chef Pro - aichef.pro")
        c.drawCentredString(297, 480, "20 capitulos - 60+ paginas - 8 plantillas - 6 checklists")
        c.save()
        print(f"✓ {path}")
    except ImportError:
        import shutil
        src = os.path.join(OUTPUT_DIR, "guia-restaurante-mexicano.docx")
        dst = os.path.join(OUTPUT_DIR, "guia-restaurante-mexicano.pdf")
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f"✓ {dst} (placeholder from DOCX)")


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n🌮 Generando archivos: Restaurante Mexicano 80 Plazas\n")
    gen_guide_docx()
    gen_guide_pdf()
    gen_plan_financiero()
    gen_calculadora_capex()
    gen_pl_mensual()
    gen_cash_flow()
    gen_escandallo()
    gen_menu_engineering()
    gen_calculadora_ticket()
    gen_cronograma()
    gen_plantilla_turnos()
    gen_checklist_legal()
    gen_checklist_equipamiento()
    gen_checklist_appcc()
    gen_checklist_sala()
    gen_checklist_contratacion()
    gen_checklist_marketing()
    gen_business_plan()
    gen_manual_operaciones()
    print(f"\n✅ {len(os.listdir(OUTPUT_DIR))} archivos generados en {OUTPUT_DIR}\n")
